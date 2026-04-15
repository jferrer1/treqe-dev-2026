#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script para actualizar el Plan de Negocio Treqe con las soluciones analizadas.
Modifica SOLO las áreas concernientes a las problemáticas y sistema de scoring.
Respetando el contenido y estilo existente.
"""

import os
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def actualizar_plan_con_soluciones():
    """Actualiza el documento con las soluciones analizadas"""
    
    # Rutas de archivos
    workspace_path = Path(r"C:\Users\Shadow\.openclaw\workspace")
    doc_path = workspace_path / "projects" / "Treqe" / "plan_negocio" / "Plan_Negocio_Treqe_100%_PROFESIONAL_FINAL.docx"
    nuevo_doc_path = workspace_path / "projects" / "Treqe" / "plan_negocio" / "Plan_Negocio_Treqe_CON_SOLUCIONES_ACTUALIZADO.docx"
    
    print(f"📄 Leyendo documento: {doc_path}")
    
    try:
        # Cargar documento
        doc = Document(str(doc_path))
        
        # Contar secciones
        print(f"📊 Documento cargado: {len(doc.paragraphs)} párrafos")
        
        # Buscar secciones relevantes para actualizar
        # 1. Buscar sección sobre "Sistema de comunicación" o "Chat"
        # 2. Buscar sección sobre "Garantías" o "Protección al usuario"
        # 3. Buscar sección sobre "Sistema de reputación" o "Scoring"
        # 4. Buscar sección sobre "Gestión de envíos" o "Logística"
        
        secciones_actualizadas = []
        
        # SOLUCIONES A INCORPORAR (basadas en el análisis)
        soluciones = {
            "problema_chat": {
                "titulo": "Sistema de Comunicación y Negociación",
                "contenido": """
                **SISTEMA DE OFERTAS ESTRUCTURADAS (SIN CHAT GRUPAL)**

                Tras un análisis exhaustivo de las necesidades de Treqe, se ha determinado que el chat grupal tradicional NO es la opción óptima para las "ruedas de intercambio". En su lugar, se implementará un sistema de ofertas/aceptaciones estructuradas que ofrece las siguientes ventajas:

                **1. Arquitectura del Sistema:**
                - Cada usuario publica OFRECE (artículo + valoración) / QUIERE (artículo deseado)
                - El algoritmo de Treqe encuentra coincidencias automáticamente
                - Propone rueda completa con compensaciones calculadas matemáticamente
                - Usuarios aceptan/rechazan en silencio (sin necesidad de chat)

                **2. Ventajas Competitivas:**
                ✅ **Elimina fricción social:** Usuarios tímidos o menos experimentados no se sienten presionados
                ✅ **Más rápido:** No hay esperas por respuestas en tiempo real
                ✅ **Escala mejor:** Permite miles de ruedas simultáneas sin moderación humana
                ✅ **Coherencia:** Alineado con el modelo automatizado de cálculo de compensaciones
                ✅ **Menos conflictos:** Elimina malentendidos y negociaciones tensas

                **3. Interfaz Propuesta:**
                ```
                [RUEDA #4832 - PROPUESTA DEL SISTEMA]
                Ana → Beatriz: Bicicleta (valoración: 450€)
                Carlos → Ana: Sofá (valoración: 600€)  
                Beatriz → Carlos: Portátil (valoración: 800€)

                COMPENSACIONES CALCULADAS:
                • Ana paga 150€ a Carlos
                • Carlos paga 200€ a Beatriz
                • Beatriz recibe 350€ neto

                [ ] ACEPTO participar en esta rueda
                [ ] RECHAZAR (sin necesidad de explicación)
                [ ] CONTRAOFERTA (ajustar valores de compensación)
                ```

                **4. Validación del Enfoque:**
                Este sistema ha sido validado como superior al chat grupal por:
                - **Coherencia con el modelo:** Si el sistema ya calcula compensaciones automáticamente, no necesita chat para negociación
                - **Escalabilidad:** Treqe proyecta 120,000 transacciones/año en año 3
                - **Experiencia consistente:** Todos los usuarios tienen la misma experiencia
                - **Menos problemas operativos:** Sin moderación, sin conflictos en chat, sin esperas
                """
            },
            
            "problema_desistimiento": {
                "titulo": "Sistema de Garantías y Desistimiento 24h",
                "contenido": """
                **SISTEMA DE 2 CAPAS: FONDO DE GARANTÍA + SISTEMA DE REPUTACIÓN**

                Para resolver el problema crítico del desistimiento en las primeras 24 horas, se ha diseñado un sistema robusto de 2 capas que protege la integridad de las ruedas de intercambio:

                **CAPA 1: FONDO DE GARANTÍA TREQE**
                - **Financiación:** 0.1% de cada transacción va al fondo de garantía
                - **Cobertura:** Si un usuario desiste en 24h, el fondo paga las compensaciones a los otros participantes
                - **Gestión del artículo:** El artículo devuelto va al "inventario Treqe" para futuras ruedas
                - **Proceso simplificado:** Usuario puede devolver sin necesidad de explicaciones complejas

                **CAPA 2: SISTEMA DE REPUTACIÓN MEJORADO**
                - **Incentivos:** Usuarios con alta reputación tienen prioridad en matching de ruedas
                - **Límites progresivos:** Usuarios que usan la garantía frecuentemente tienen límites temporales
                - **Beneficios escalonados:** Mayor reputación = menores comisiones + acceso preferente

                **IMPLEMENTACIÓN TÉCNICA:**
                ```python
                class SistemaGarantiaTreqe:
                    def __init__(self):
                        self.fondo = 0.001  # 0.1% de cada transacción
                        self.inventario_treqe = []
                    
                    def manejar_desistimiento(self, usuario_id, rueda_id):
                        # 1. Pagar compensaciones desde el fondo
                        self.pagar_desde_fondo(rueda_id)
                        
                        # 2. Mover artículo a inventario Treqe
                        self.mover_a_inventario(rueda_id, usuario_id)
                        
                        # 3. Ajustar reputación del usuario
                        self.ajustar_reputacion(usuario_id, -10)
                        
                        # 4. Buscar reemplazo automático
                        reemplazo = self.buscar_reemplazo(rueda_id)
                        if reemplazo:
                            self.reactivar_rueda(rueda_id, reemplazo)
                ```

                **VENTAJAS DEL SISTEMA:**
                ✅ **Protege la rueda:** No se rompe aunque un usuario desista
                ✅ **Confianza:** Todos los participantes están protegidos
                ✅ **Sostenible:** El fondo se autofinancia con las transacciones
                ✅ **Justo:** Sistema de reputación incentiva comportamiento responsable
                """
            },
            
            "problema_envios": {
                "titulo": "Sistema de Gestión de Envíos y Protecciones",
                "contenido": """
                **TRIPLE SISTEMA DE PROTECCIÓN PARA FALLOS DE ENVÍO**

                Para resolver el problema crítico de fallos en envíos (envíos no realizados, pérdidas, retrasos), se ha diseñado un sistema triple que ofrece múltiples niveles de protección:

                **SOLUCIÓN 1: PAGOS EN ESCROW (RECOMENDACIÓN PRINCIPAL)**
                - **Proceso:** Las compensaciones van a una cuenta de Treqe (escrow)
                - **Verificación:** Sistema verifica que TODOS los participantes han enviado
                - **Liberación:** Solo entonces se liberan las compensaciones
                - **Protección:** Si falla algún envío → reversión completa de la transacción

                **SOLUCIÓN 2: SEGURO DE ENVÍO OBLIGATORIO (OPCIÓN SIMPLE)**
                - **Costo:** Cada usuario paga pequeño seguro (ej: 1€ por transacción)
                - **Cobertura:** Si su envío falla, el seguro cubre las compensaciones
                - **Incentivo:** Usuario que falla pierde seguro + reputación
                - **Autofinanciación:** El fondo de seguros se autofinancia

                **SOLUCIÓN 3: VERIFICACIÓN POR PASOS (OPCIÓN EXTRA SEGURA)**
                ```
                Día 1: Usuario A envía → Usuario B confirma recepción ✓
                Día 2: Usuario B envía → Usuario C confirma recepción ✓  
                Día 3: Usuario C envía → Usuario A confirma recepción ✓
                Día 4: Sistema libera compensaciones a todos
                ```
                - **Seguridad máxima:** Si falla cualquier paso, todo se revierte
                - **Verificación humana:** Confirmación en cada etapa
                - **Para transacciones de alto valor:** Opcional para artículos >500€

                **SISTEMA HÍBRICO RECOMENDADO:**
                - **Transacciones <100€:** Solución 2 (seguro simple)
                - **Transacciones 100-500€:** Solución 1 (escrow)
                - **Transacciones >500€:** Solución 3 (verificación por pasos)
                - **Usuarios Elite:** Pueden elegir sistema preferido
                """
            },
            
            "sistema_scoring": {
                "titulo": "Sistema de Puntuación y Reputación Treqe",
                "contenido": """
                **ALGORITMO DE SCORING TREQE - MÉTRICA COMPUESTA**

                El sistema de puntuación/reputación es fundamental para incentivar comportamiento positivo y ofrecer beneficios escalonados:

                **FÓRMULA DE PUNTUACIÓN:**
                ```
                PUNTUACIÓN_TREQE = 
                  (Transacciones exitosas × 10) +
                  (Valor total intercambiado / 100) +
                  (Tiempo promedio de envío < 48h × 5) -
                  (Fallos de envío × 50) -
                  (Desistimientos × 30) -
                  (Reclamos recibidos × 20)
                ```

                **NIVELES Y BENEFICIOS:**

                **🌟 NOVATO (0-100 puntos)**
                - Solo puede recibir ofertas, no iniciar ruedas
                - Comisión estándar: 1.0%
                - Envío propio con riesgo completo
                - Límite: 100€ por transacción

                **🌟🌟 MIEMBRO (101-500 puntos)**
                - Puede iniciar ruedas hasta 200€
                - Comisión: 1.0%
                - Acceso a seguro básico de envío
                - Prioridad media en matching

                **🌟🌟🌟 CONFIABLE (501-1000 puntos)**
                - Puede iniciar ruedas hasta 500€
                - Comisión: 0.9%
                - Acceso a logística asociada con 15% descuento
                - Menor depósito en escrow (50% vs 100%)
                - Prioridad alta en matching

                **🌟🌟🌟🌟 ELITE (1001+ puntos)**
                - Sin límite de valor en ruedas
                - Comisión: 0.8%
                - Logística garantizada incluida en transacciones >300€
                - Prioridad máxima en matching
                - Representante de la comunidad Treqe
                - Acceso a beta testing de nuevas funcionalidades

                **SISTEMA DE INCENTIVOS:**
                - **Transparencia:** Puntuación visible en perfil
                - **Progresión clara:** Usuarios ven cómo mejorar su nivel
                - **Recompensas tangibles:** Ahorro real en comisiones
                - **Reconocimiento social:** Estatus visible en la comunidad

                **IMPACTO ESPERADO:**
                - **Retención:** Usuarios buscan mantener/m mejorar reputación
                - **Calidad:** Menos incidencias con usuarios de alta reputación
                - **Crecimiento orgánico:** Usuarios recomiendan para ganar puntos
                - **Sostenibilidad:** Menor necesidad de moderación humana
                """
            }
        }
        
        # Buscar donde insertar las soluciones
        # Estrategia: Buscar secciones existentes y actualizarlas, o crear nuevas secciones al final
        
        # 1. Buscar índice o tabla de contenidos para entender estructura
        estructura_documento = []
        for i, para in enumerate(doc.paragraphs):
            if para.text and (para.text.startswith("1.") or para.text.startswith("2.") or 
                            para.text.startswith("3.") or para.text.startswith("4.") or
                            para.text.startswith("5.") or para.text.startswith("6.") or
                            para.text.startswith("7.") or para.text.startswith("8.") or
                            para.text.startswith("9.") or para.text.startswith("10.")):
                estructura_documento.append((i, para.text))
                print(f"  Sección encontrada: {para.text}")
        
        print(f"\n🔍 Encontradas {len(estructura_documento)} secciones numeradas")
        
        # 2. Buscar palabras clave para actualizar secciones existentes
        palabras_clave = {
            "chat": ["chat", "comunicación", "negociación", "conversación"],
            "garantía": ["garantía", "devolución", "desistimiento", "reembolso"],
            "envío": ["envío", "logística", "transporte", "entrega"],
            "reputación": ["reputación", "puntuación", "scoring", "confianza"]
        }
        
        secciones_actualizables = []
        
        for i, para in enumerate(doc.paragraphs):
            texto = para.text.lower()
            for clave, palabras in palabras_clave.items():
                for palabra in palabras:
                    if palabra in texto and len(para.text) < 200:  # Probablemente título
                        secciones_actualizables.append((i, clave, para.text))
                        print(f"  Sección '{clave}' encontrada en párrafo {i}: {para.text[:50]}...")
                        break
        
        # 3. Estrategia de actualización
        # Si encontramos secciones existentes, las actualizamos
        # Si no, añadimos nuevas secciones al final antes de los anexos
        
        # Buscar "ANEXOS" o "Anexos" para insertar antes
        indice_anexos = -1
        for i, para in enumerate(doc.paragraphs):
            if "ANEXO" in para.text.upper() or "Anexo" in para.text:
                indice_anexos = i
                print(f"📌 Anexos encontrados en párrafo {i}")
                break
        
        # Si no encontramos anexos, añadimos al final
        if indice_anexos == -1:
            indice_anexos = len(doc.paragraphs)
        
        # 4. Añadir nuevas secciones con las soluciones
        print(f"\n📝 Añadiendo soluciones analizadas...")
        
        # Añadir título de sección nueva
        doc.paragraphs[indice_anexos].insert_paragraph_before("")
        doc.paragraphs[indice_anexos].insert_paragraph_before("ANÁLISIS DE PROBLEMÁTICAS CRÍTICAS Y SOLUCIONES IMPLEMENTADAS")
        titulo_seccion = doc.paragraphs[indice_anexos - 1]
        titulo_seccion.style = doc.styles['Heading 1']
        
        # Añadir introducción
        intro = doc.paragraphs[indice_anexos].insert_paragraph_before("""
        Esta sección documenta el análisis exhaustivo realizado sobre las problemáticas críticas identificadas 
        en el modelo de negocio Treqe y las soluciones diseñadas para cada una de ellas. 
        El análisis incluye evaluación de múltiples alternativas y selección de la solución óptima basada en 
        criterios de escalabilidad, seguridad, experiencia de usuario y viabilidad técnica.
        """)
        
        # Añadir cada solución
        for solucion_key, solucion_data in soluciones.items():
            # Título de la solución
            doc.paragraphs[indice_anexos].insert_paragraph_before("")
            doc.paragraphs[indice_anexos].insert_paragraph_before(solucion_data["titulo"])
            titulo_solucion = doc.paragraphs[indice_anexos - 1]
            titulo_solucion.style = doc.styles['Heading 2']
            
            # Contenido de la solución
            contenido = doc.paragraphs[indice_anexos].insert