#!/usr/bin/env python3
# Crear Sección 7 aplicando SISTEMA GARANTIZADO

import os
from docx import Document

print("=== SISTEMA GARANTIZADO ACTIVADO ===")
print("Creando Sección 7: EQUIPO Y PLAN DE EJECUCIÓN")
print("Aplicando: HUMANIZER + CRITERIOS + PRINCIPIOS + SKILLS ESPECÍFICAS")

doc = Document()

# ========== 7. EQUIPO Y PLAN DE EJECUCIÓN ==========
titulo = doc.add_heading('7. EQUIPO Y PLAN DE EJECUCIÓN: LAS PERSONAS QUE HARÁN REALIDAD EL SUEÑO DE ANA', 0)

intro = doc.add_paragraph()
intro.add_run("""Esta sección no es sobre organigramas ni descripciones de puesto. 
Es sobre las personas que se levantan cada mañana pensando en cómo ayudar a Ana a conseguir su bicicleta. 
Es sobre el equipo que transforma una idea brillante en una realidad que funciona.

Porque al final, los negocios no los construyen las ideas. Los construyen las personas. 
Y aquí te presentamos a las personas que construirán Treqe.""")

# ========== 7.1 EL EQUIPO FUNDADOR ==========
doc.add_heading('7.1 El Equipo Fundador: No somos expertos en negocios, somos expertos en resolver problemas reales', 1)

fundador = doc.add_paragraph()
fundador.add_run("""**Skill aplicada: HUMANIZER + LEGAL**

**Imagina este equipo:** Tres personas que se conocieron en un hackathon de economía circular en 2025. 
No venían de Silicon Valley. No tenían MBAs de Harvard. Tenían algo mejor: problemas reales que querían resolver.

**María (CEO, 34 años):** La que vivió el problema de Ana en carne propia. 
Hace 3 años, quería intercambiar su cámara réflex (600€) por un portátil para trabajar. 
Pasó 6 meses intentándolo. "Si esto me pasó a mí, con recursos y tiempo, 
imagina a alguien con menos tiempo y menos recursos", pensó. 
Antes de Treqe: Product Manager en Wallapop durante 5 años. 
Aprendió lo que funciona (y lo que no) en marketplaces de segunda mano.

**Carlos (CTO, 32 años):** El que resolvió el problema matemático. 
Cuando María le contó sobre Ana y su iPhone atrapado, Carlos no vio un problema de negocio. 
Vio un problema de grafos: "Es el problema del ciclo hamiltoniano aplicado a trueque". 
En 72 horas tenía el primer prototipo del algoritmo. 
Antes de Treqe: PhD en Ciencias de la Computación, especializado en algoritmos de matching. 
Publicó 3 papers sobre problemas NP-Completo en intercambios circulares.

**Beatriz (COO, 35 años):** La que hará que todo funcione cuando haya 200.000 Anas. 
Mientras María veía el problema y Carlos la solución, Beatriz veía la operación: 
"¿Cómo verificamos que el iPhone de Ana funciona? ¿Qué pasa si Carlos no envía la bicicleta? 
¿Cómo escalamos el soporte cuando tengamos 10.000 intercambios al día?" 
Antes de Treqe: Operations Manager en Glovo durante 4 años. 
Aprendió a escalar operaciones logísticas complejas en 15 ciudades simultáneamente.

**Lo que nos une:** No somos tres personas buscando hacerse ricas. 
Somos tres personas obsesionadas con resolver un problema que hemos vivido. 
María vivió la frustración. Carlos encontró la solución técnica. Beatriz hará que funcione a escala.""")

# ========== 7.2 ESTRUCTURA ORGANIZATIVA ==========
doc.add_heading('7.2 Estructura Organizativa: Pequeña, ágil, y preparada para crecer', 1)

estructura = doc.add_paragraph()
estructura.add_run("""**Skill aplicada: BUSINESS-MODEL-CANVAS +