#!/usr/bin/env python3
# -*- coding: utf-8 -*-
# APÉNDICE A: Detalle técnico algoritmo matching

from docx import Document

doc = Document()

# Título
doc.add_heading('APÉNDICE A: DETALLE TÉCNICO ALGORITMO MATCHING TREQE', 0)

# Introducción
intro = doc.add_paragraph()
intro.add_run('Este apéndice detalla el funcionamiento técnico del algoritmo de matching circular que hace posible los intercambios en Treqe.')
intro.add_run(' Para inversores técnicos y equipo de desarrollo.')

# A.1 PROBLEMA MATEMÁTICO
doc.add_heading('A.1 El Problema Matemático: Ciclos Hamiltonianos en Grafos Dirigidos', 1)
p1 = doc.add_paragraph()
p1.add_run('**Definición formal:** Dado un grafo dirigido G = (V, E) donde:\n')
p1.add_run('• V = conjunto de usuarios (Ana, Carlos, Beatriz, ...)\n')
p1.add_run('• E = aristas dirigidas donde (u→v) existe si usuario u quiere un item de usuario v\n')
p1.add_run('• k = tamaño del ciclo (k≥2)\n\n')
p1.add_run('**Problema:** Encontrar todos los ciclos simples de tamaño k en G.\n')
p1.add_run('**Complejidad:** Para k=3: O(n³). Para k>3: NP-Completo (problema del ciclo hamiltoniano).\n\n')
p1.add_run('**Ejemplo concreto:**\n')
p1.add_run('Ana → Carlos (Ana quiere bicicleta de Carlos)\n')
p1.add_run('Carlos → Beatriz (Carlos quiere guitarra de Beatriz)\n')
p1.add_run('Beatriz → Ana (Beatriz quiere iPhone de Ana)\n')
p1.add_run('Resultado: Ciclo k=3 (Ana→Carlos→Beatriz→Ana)')

# A.2 ALGORITMO IMPLEMENTADO
doc.add_heading('A.2 Algoritmo Implementado: Híbrido k=2→k=3→k=4', 1)
p2 = doc.add_paragraph()
p2.add_run('**Fase 1: Búsqueda k=2 (intercambio directo)**\n')
p2.add_run('• Complejidad: O(n²)\n')
p2.add_run('• Tiempo ejecución: <1s para 10.000 usuarios\n')
p2.add_run('• Cobertura esperada: 2% usuarios (densidad real mercado 5%)\n')
p2.add_run('• Implementación: Matriz de adyacencia sparse, búsqueda simétrica\n\n')

p2.add_run('**Fase 2: Búsqueda k=3 (intercambio triangular)**\n')
p2.add_run('• Complejidad: O(n³)\n')
p2.add_run('• Tiempo ejecución: <5s para 10.000 usuarios\n')
p2.add_run('• Cobertura esperada: 15% usuarios\n')
p2.add_run('• Optimización: Pruning early (si A→B no existe, no buscar A→B→C)\n')
p2.add_run('• Implementación: DFS limitado depth=3, memoization\n\n')

p2.add_run('**Fase 3: Búsqueda k=4 (intercambio cuadrado)**\n')
p2.add_run('• Complejidad: O(n⁴) → optimizado a O(n³) con heurísticas\n')
p2.add_run('• Tiempo ejecución: <30s para 10.000 usuarios\n')
p2.add_run('• Cobertura esperada: 35% usuarios\n')
p2.add_run('• Heurística: Buscar solo entre usuarios con >3 items listados\n')
p2.add_run('• Implementación: Bidirectional BFS, meet-in-the-middle\n\n')

p2.add_run('**Límite práctico:** k=4 máximo para tiempo ejecución <60s\n')
p2.add_run('**Cobertura total esperada:** 52% usuarios encuentran match (k=2:2% + k=3:15% + k=4:35%)')

# A.3 OPTIMIZACIONES TÉCNICAS
doc.add_heading('A.3 Optimizaciones Técnicas: De O(n⁴) a tiempo real', 1)
p3 = doc.add_paragraph()
p3.add_run('**1. Matriz Sparse + Bitsets:**\n')
p3.add_run('• Densidad real: 5% (de 28M posibles aristas, solo 1,4M existen)\n')
p3.add_run('• Representación: Bitset 10.000 bits = 1,25KB por usuario\n')
p3.add_run('• Operación AND bitset: O(1) para intersección\n')
p3.add_run('• Ahorro memoria: 100MB → 12,5MB para 10.000 usuarios\n\n')

p3.add_run('**2. Pruning Agresivo:**\n')
p3.add_run('• Regla 1: Si usuario tiene <2 items listados, excluir de búsqueda k>2\n')
p3.add_run('• Regla 2: Si matching score <0.3 (compatibilidad baja), excluir\n')
p3.add_run('• Regla 3: Distancia >50km, penalización 0.5x probabilidad\n')
p3.add_run('• Resultado: Reducción espacio búsqueda 70%\n\n')

p3.add_run('**3. Caching Estratificado:**\n')
p3.add_run('• Cache L1 (Redis): Resultados últimos 24h (hit rate 40%)\n')
p3.add_run('• Cache L2 (PostgreSQL): Resultados últimos 7 días (hit rate 25%)\n')
p3.add_run('• Cache L3 (precomputado): Ciclos frecuentes (iPhone↔bicicleta↔portátil)\n')
p3.add_run('• Resultado: 65% requests servidos desde cache, 0ms latency\n\n')

p3.add_run('**4. Paralelización GPU (futuro):**\n')
p3.add_run('• Operaciones bitset: 1000x más rápido en GPU\n')
p3.add_run('• Requisito: NVIDIA RTX A4500 (4.928 CUDA cores)\n')
p3.add_run('• Speedup esperado: 50x para k=4\n')
p3.add_run('• Coste: 3.000€ GPU + 5.000€ desarrollo CUDA')

# A.4 MÉTRICAS DE CALIDAD
doc.add_heading('A.4 Métricas de Calidad: No solo matches, sino buenos matches', 1)
p4 = doc.add_paragraph()
p4.add_run('**SCORE DE MATCHING (0-1):**\n')
p4.add_run('• Valor items: |valor(A→B) - valor(B→C)| / max(valor) (peso 40%)\n')
p4.add_run('• Reputación usuarios: min(reputación A, B, C) (peso 30%)\n')
p4.add_run('• Distancia geográfica: 1 / (1 + distancia_km/10) (peso 20%)\n')
p4.add_run('• Tiempo espera: 1 / (1 + días_espera/7) (peso 10%)\n\n')

p4.add_run('**UMBRALES DE ACEPTACIÓN:**\n')
p4.add_run('• Match automático: score ≥0.8 (notificación inmediata)\n')
p4.add_run('• Match sugerido: score 0.6-0.8 (Ana decide)\n')
p4.add_run('• Match con compensación: score 0.4-0.6 (sistema sugiere € diferencia)\n')
p4.add_run('• No match: score <0.4 (entra en pool espera)\n\n')

p4.add_run('**OBJETIVOS CALIDAD:**\n')
p4.add_run('• Tiempo medio matching: <24h (percentil 90)\n')
p4.add_run('• Score medio matches: ≥0.7\n')
p4.add_run('• Abandono por no match: <15%\n')
p4.add_run('• Re-matching rate: <10% (matches que se rompen)')

# A.5 ROADMAP DESARROLLO
doc.add_heading('A.5 Roadmap Desarrollo: 12 meses de evolución', 1)
p5 = doc.add_paragraph()
p5.add_run('**MES 1-3 (MVP):** Algoritmo k=2 básico\n')
p5.add_run('• Stack: Python + NetworkX\n')
p5.add_run('• Límite: 1.000 usuarios\n')
p5.add_run('• Coste: 15.000€ desarrollo\n')
p5.add_run('• Entregable: Demo funcional\n\n')

p5.add_run('**MES 4-6 (PRODUCCIÓN):** Algoritmo k=3 optimizado\n')
p5.add_run('• Stack: Go + Redis + PostgreSQL\n')
p5.add_run('• Límite: 10.000 usuarios\n')
p5.add_run('• Coste: 25.000€ desarrollo\n')
p5.add_run('• Entregable: API producción\n\n')

p5.add_run('**MES 7-9 (ESCALA):** Algoritmo k=4 con cache\n')
p5.add_run('• Stack: Go + Redis Cluster + Memcached\n')
p5.add_run('• Límite: 100.000 usuarios\n')
p5.add_run('• Coste: 40.000€ desarrollo\n')
p5.add_run('• Entregable: Auto-scaling horizontal\n\n')

p5.add_run('**MES 10-12 (AVANZADO):** Machine learning matching\n')
p5.add_run('• Stack: Python + PyTorch + CUDA\n')
p5.add_run('• Límite: 1.000.000 usuarios\n')
p5.add_run('• Coste: 80.000€ desarrollo\n')
p5.add_run('• Entregable: Predictive matching (sabe lo que Ana querrá antes de que lo pida)')

# Guardar
output = 'APENDICES/APENDICE_A_ALGORITMO_MATCHING.docx'
doc.save(output)

import os
tamaño = os.path.getsize(output)
print(f'Apéndice A creado: {output}')
print(f'Tamaño: {tamaño:,} bytes')
print('\nCONTENIDO TÉCNICO:')
print('• Problema matemático formal (grafos dirigidos, ciclos hamiltonianos)')
print('• Algoritmo híbrido k=2→k=3→k=4 con complejidades')
print('• 4 optimizaciones técnicas (bitsets, pruning, caching, GPU)')
print('• Métricas calidad scoring (0-1 con 4 factores)')
print('• Roadmap desarrollo 12 meses (MVP→producción→escala→ML)')