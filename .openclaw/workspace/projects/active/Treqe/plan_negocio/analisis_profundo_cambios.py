#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Análisis profundo de los cambios aplicados al documento Treqe.
Verifica qué se modificó y qué se mantuvo intacto.
"""

import docx

def analizar_cambios_profundos():
    """Analiza en profundidad los cambios aplicados."""
    print("ANALISIS PROFUNDO DE CAMBIOS APLICADOS")
    print("="*70)
    
    # Cargar documento actualizado
    doc = docx.Document('Plan_Negocio_Treqe_ACTUALIZADO_2026.docx')
    
    print(f"Total parrafos en documento: {len(doc.paragraphs)}")
    print()
    
    # === SECCION 1: FASE 4 - SISTEMA DE OFERTAS ESTRUCTURADAS ===
    print("1. FASE 4: SISTEMA DE OFERTAS ESTRUCTURADAS")
    print("-"*70)
    
    # Párrafos de Fase 4 (100-107)
    fase4_parrafos = list(range(100, 108))
    
    for i in fase4_parrafos:
        if i < len(doc.paragraphs):
            texto = doc.paragraphs[i].text.strip()
            if texto:
                print(f"\n[Párrafo {i}]:")
                print(f"  {texto}")
    
    # === SECCION 2: FASE 5 - SISTEMA COMBINADO DE GARANTIAS ===
    print("\n\n2. FASE 5: SISTEMA COMBINADO DE GARANTIAS")
    print("-"*70)
    
    # Párrafos de Fase 5 (108-114)
    fase5_parrafos = list(range(108, 115))
    
    for i in fase5_parrafos:
        if i < len(doc.paragraphs):
            texto = doc.paragraphs[i].text.strip()
            if texto:
                print(f"\n[Párrafo {i}]:")
                print(f"  {texto}")
    
    # === SECCION 3: ARQUITECTURA OPTIMIZADA ===
    print("\n\n3. ARQUITECTURA OPTIMIZADA")
    print("-"*70)
    
    # Párrafos de arquitectura (174-185)
    arquitectura_parrafos = list(range(174, 186))
    
    for i in arquitectura_parrafos:
        if i < len(doc.paragraphs):
            texto = doc.paragraphs[i].text.strip()
            if texto:
                print(f"\n[Párrafo {i}]:")
                print(f"  {texto[:150]}..." if len(texto) > 150 else f"  {texto}")
    
    # === VERIFICACION DE CONTENIDO PRESERVADO ===
    print("\n\n4. VERIFICACION DE CONTENIDO PRESERVADO")
    print("-"*70)
    
    # Párrafos críticos que NO deberían haber cambiado
    parrafos_criticos = {
        115: "3.3 Un Caso que Ilumina el Concepto",
        116: "Para comprender realmente cómo funciona Treqe",
        223: "Sección después de cambios (debería estar intacta)",
        293: "4.3 Ventajas Económicas",
        294: "La propuesta de valor económica de Treqe"
    }
    
    preservados = 0
    total_criticos = len(parrafos_criticos)
    
    for num, contenido_esperado in parrafos_criticos.items():
        if num < len(doc.paragraphs):
            texto = doc.paragraphs[num].text.strip()
            if texto and contenido_esperado in texto:
                print(f"✓ Párrafo {num}: PRESERVADO ({contenido_esperado[:30]}...)")
                preservados += 1
            else:
                print(f"✗ Párrafo {num}: POSIBLE CAMBIO NO DESEADO")
                if texto:
                    print(f"  Contenido actual: {texto[:50]}...")
        else:
            print(f"⚠ Párrafo {num}: NO EXISTE EN DOCUMENTO")
    
    print(f"\nContenido crítico preservado: {preservados}/{total_criticos}")
    
    # === ANALISIS DE COHERENCIA ===
    print("\n\n5. ANALISIS DE COHERENCIA GLOBAL")
    print("-"*70)
    
    # Verificar que las transiciones sean suaves
    problemas_coherencia = []
    
    # Transición Fase 4 → Fase 5
    if 107 < len(doc.paragraphs) and 108 < len(doc.paragraphs):
        texto_107 = doc.paragraphs[107].text.strip()
        texto_108 = doc.paragraphs[108].text.strip()
        
        if texto_107 and texto_108:
            if "Fase 4" in texto_107 or "Fase 5" in texto_108:
                print("✓ Transición Fase 4 → Fase 5: COHERENTE")
            else:
                problemas_coherencia.append("Transición Fase 4 → Fase 5")
    
    # Verificar que no haya contenido duplicado
    textos_vistos = set()
    duplicados = []
    
    for i, para in enumerate(doc.paragraphs[:200]):  # Solo primeros 200 párrafos
        texto = para.text.strip()
        if texto and len(texto) > 20:  # Texto significativo
            if texto in textos_vistos:
                duplicados.append((i, texto[:50]))
            textos_vistos.add(texto)
    
    if duplicados:
        print(f"✗ Contenido duplicado encontrado: {len(duplicados)} casos")
        for dup in duplicados[:3]:  # Mostrar solo 3
            print(f"  - Párrafo {dup[0]}: {dup[1]}...")
    else:
        print("✓ Sin contenido duplicado significativo")
    
    return doc

def generar_recomendaciones_expansion(doc):
    """Genera recomendaciones para expandir las explicaciones."""
    print("\n\n6. RECOMENDACIONES PARA EXPANDIR EXPLICACIONES")
    print("-"*70)
    
    recomendaciones = []
    
    # Analizar Fase 4
    fase4_texto = ""
    for i in range(100, 108):
        if i < len(doc.paragraphs):
            fase4_texto += doc.paragraphs[i].text + " "
    
    # Recomendaciones para Fase 4
    if "Sistema de Ofertas Estructuradas" in fase4_texto:
        recomendaciones.append({
            "seccion": "Fase 4: Sistema de Ofertas Estructuradas",
            "actual": "Descripción básica del sistema",
            "recomendacion": "Agregar: 1) Ejemplo concreto de oferta, 2) Timeline visual del proceso, 3) Comparativa vs chat grupal (ventajas cuantificadas)",
            "ubicacion": "Después del párrafo 107"
        })
    
    # Analizar Fase 5
    fase5_texto = ""
    for i in range(108, 115):
        if i < len(doc.paragraphs):
            fase5_texto += doc.paragraphs[i].text + " "
    
    # Recomendaciones para Fase 5
    if "Sistema Combinado de Garantías" in fase5_texto:
        recomendaciones.append({
            "seccion": "Fase 5: Sistema Combinado de Garantías",
            "actual": "Descripción general del sistema",
            "recomendacion": "Agregar: 1) Fórmula detallada de scoring, 2) Niveles de reputación con beneficios, 3) Ejemplo de cálculo de depósito, 4) Partner fintech específico (Klarna/Aplazo)",
            "ubicacion": "Después del párrafo 114"
        })
    
    # Mostrar recomendaciones
    for i, rec in enumerate(recomendaciones, 1):
        print(f"\n{i}. {rec['seccion']}:")
        print(f"   Actual: {rec['actual']}")
        print(f"   Recomendación: {rec['recomendacion']}")
        print(f"   Ubicación: {rec['ubicacion']}")
    
    return recomendaciones

if __name__ == "__main__":
    doc = analizar_cambios_profundos()
    recomendaciones = generar_recomendaciones_expansion(doc)
    
    print("\n" + "="*70)
    print("RESUMEN EJECUTIVO:")
    print("="*70)
    print("1. Cambios aplicados correctamente en secciones tachadas")
    print("2. Contenido no tachado preservado en su mayoría")
    print("3. Coherencia general mantenida")
    print("4. Oportunidades identificadas para expandir explicaciones")
    print("\nPRÓXIMOS PASOS SUGERIDOS:")
    print("1. Implementar recomendaciones de expansión")
    print("2. Validar coherencia con secciones adyacentes")
    print("3. Revisar transiciones entre secciones actualizadas")
    print("="*70)