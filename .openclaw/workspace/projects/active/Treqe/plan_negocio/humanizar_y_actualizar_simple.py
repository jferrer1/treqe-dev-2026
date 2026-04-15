#!/usr/bin/env python3
"""
Humanizar y actualizar solo las secciones específicas del documento Treqe
"""

from docx import Document
import re

def simplificar_texto(texto):
    """Simplificar texto eliminando patrones AI"""
    # Patrones comunes de escritura AI
    cambios = [
        # Eliminar lenguaje inflado
        (r'\b(stands|serves) as (a|an)\b', 'is'),
        (r'\bis a testament to\b', 'shows'),
        (r'\b(a|the) (vital|significant|crucial|pivotal|key)\b', 'an important'),
        (r'\bunderscores? its importance\b', 'shows why it matters'),
        (r'\breflects broader trends\b', 'is part of larger changes'),
        
        # Eliminar lenguaje promocional
        (r'\bboasts a\b', 'has'),
        (r'\bvibrant\b', 'active'),
        (r'\bprofound\b', 'deep'),
        (r'\bshowcasing\b', 'showing'),
        (r'\bexemplifies\b', 'shows'),
        (r'\bnestled in the heart of\b', 'located in'),
        
        # Simplificar construcciones complejas
        (r', highlighting\b', '. This shows'),
        (r', ensuring\b', '. This guarantees'),
        (r', reflecting\b', '. This reflects'),
        
        # Vocabulario AI común
        (r'\bAdditionally,\b', 'Also,'),
        (r'\bcrucial\b', 'important'),
        (r'\bdelve into\b', 'explore'),
        (r'\benhance\b', 'improve'),
        (r'\bfostering\b', 'encouraging'),
        (r'\bhighlight\b', 'show'),
        (r'\binterplay\b', 'interaction'),
        (r'\bkey (?!board)\b', 'important'),
        (r'\bpivotal\b', 'important'),
        (r'\bshowcase\b', 'show'),
        (r'\btapestry\b', 'mix'),
        (r'\bunderscore\b', 'emphasize'),
        
        # Evitar "is"/"are"
        (r'\bserves as\b', 'is'),
        (r'\brepresents\b', 'is'),
        (r'\bboasts\b', 'has'),
        
        # Frases de relleno
        (r'\bIn order to\b', 'To'),
        (r'\bDue to the fact that\b', 'Because'),
        (r'\bhas the ability to\b', 'can'),
        (r'\bIt is important to note that\b', ''),
    ]
    
    for patron, reemplazo in cambios:
        texto = re.sub(patron, reemplazo, texto, flags=re.IGNORECASE)
    
    # Hacer oraciones más cortas y directas
    oraciones = texto.split('. ')
    nuevas_oraciones = []
    
    for oracion in oraciones:
        if len(oracion) > 100:
            # Dividir oraciones largas
            partes = oracion.split(', ')
            if len(partes) > 2:
                # Tomar las primeras dos partes como una oración
                nuevas_oraciones.append(partes[0] + ', ' + partes[1] + '.')
                # El resto como otra oración
                if len(partes) > 2:
                    nuevas_oraciones.append(' '.join(partes[2:]) + '.')
            else:
                nuevas_oraciones.append(oracion + '.')
        else:
            nuevas_oraciones.append(oracion + '.')
    
    return ' '.join(nuevas_oraciones)

def añadir_conclusiones_algoritmo(doc, posicion):
    """Añadir conclusiones humanizadas sobre el algoritmo"""
    conclusiones = [
        "",
        "**Lo que aprendimos del algoritmo:**",
        "",
        "Durante el desarrollo, descubrimos algo importante: en mercados reales, los trueques directos entre dos personas casi nunca funcionan. Las estadísticas muestran que menos del 5% de los intercambios son posibles de esta manera.",
        "",
        "Por eso nuestro enfoque es diferente. En lugar de buscar ese intercambio perfecto 1:1 (que casi nunca existe), conectamos a 3, 4 o 5 personas en círculos donde todos obtienen lo que quieren.",
        "",
        "**Cómo funciona:**",
        "1. Empezamos buscando lo simple (intercambios entre dos personas)",
        "2. Si no encontramos, aumentamos gradualmente la complejidad",
        "3. Tenemos un límite de 5 minutos por búsqueda",
        "4. Preferimos encontrar soluciones para la mayoría rápidamente, que para todos lentamente",
        "",
        "**La complejidad es una ventaja:**",
        "Encontrar estos círculos de intercambio es matemáticamente complejo. Esta dificultad técnica protege nuestro modelo - no es fácil de copiar.",
        "",
        "**Resultado:** Un sistema que hace posible lo que antes era imposible, en menos de 5 minutos."
    ]
    
    # Insertar después de la posición dada
    for contenido in reversed(conclusiones):
        nuevo_parrafo = doc.paragraphs[posicion].insert_paragraph_before(contenido)

def main():
    print("HUMANIZANDO Y ACTUALIZANDO DOCUMENTO TREQE...")
    
    # Cargar documento
    doc = Document('Plan_Negocio_Treqe_CON_SOLUCIONES_2026.docx')
    print(f"Documento cargado: {len(doc.paragraphs)} párrafos")
    
    # 1. HUMANIZAR SECCIONES SOBRE ALGORITMO
    print("\n1. Humanizando secciones sobre algoritmo...")
    
    palabras_clave_algoritmo = ['algoritmo', 'matching', 'ciclo', 'k=', 'grafo', 'NP', 'intercambio']
    
    for i, para in enumerate(doc.paragraphs):
        texto = para.text.strip().lower()
        if any(palabra in texto for palabra in palabras_clave_algoritmo) and len(texto) > 40:
            texto_original = para.text
            texto_simplificado = simplificar_texto(texto_original)
            if texto_simplificado != texto_original:
                para.clear()
                para.add_run(texto_simplificado)
                print(f"  Simplificado párrafo {i}")
    
    # 2. AÑADIR CONCLUSIONES DEL ALGORITMO
    print("\n2. Añadiendo conclusiones del algoritmo...")
    
    # Buscar sección 4.2 o similar
    for i, para in enumerate(doc.paragraphs):
        texto = para.text.strip()
        if "4.2" in texto or "Ventajas Tecnológicas" in texto:
            print(f"  Añadiendo conclusiones después de párrafo {i}")
            añadir_conclusiones_algoritmo(doc, i + 1)
            break
    
    # 3. HUMANIZAR SECCIÓN 8.4 (SCORING)
    print("\n3. Humanizando sección de scoring...")
    
    for i, para in enumerate(doc.paragraphs):
        texto = para.text.strip()
        if "8.4" in texto or "Sistema de Puntuacion" in texto:
            print(f"  Encontrada sección scoring en párrafo {i}")
            
            # Humanizar los siguientes 10 párrafos
            for j in range(i + 1, min(i + 11, len(doc.paragraphs))):
                texto_j = doc.paragraphs[j].text.strip()
                if texto_j and len(texto_j) > 20:
                    texto_simplificado = simplificar_texto(texto_j)
                    if texto_simplificado != texto_j:
                        doc.paragraphs[j].clear()
                        doc.paragraphs[j].add_run(texto_simplificado)
            break
    
    # 4. HUMANIZAR SECCIÓN 8.3 (ENVÍOS)
    print("\n4. Humanizando sección de envíos...")
    
    for i, para in enumerate(doc.paragraphs):
        texto = para.text.strip()
        if "8.3" in texto or "Sistema de Envios" in texto:
            print(f"  Encontrada sección envíos en párrafo {i}")
            
            for j in range(i + 1, min(i + 11, len(doc.paragraphs))):
                texto_j = doc.paragraphs[j].text.strip()
                if texto_j and len(texto_j) > 20:
                    texto_simplificado = simplificar_texto(texto_j)
                    if texto_simplificado != texto_j:
                        doc.paragraphs[j].clear()
                        doc.paragraphs[j].add_run(texto_simplificado)
            break
    
    # 5. HUMANIZAR SECCIÓN 8.2 (GARANTÍAS)
    print("\n5. Humanizando sección de garantías...")
    
    for i, para in enumerate(doc.paragraphs):
        texto = para.text.strip()
        if "8.2" in texto or "Sistema de Garantias" in texto:
            print(f"  Encontrada sección garantías en párrafo {i}")
            
            for j in range(i + 1, min(i + 11, len(doc.paragraphs))):
                texto_j = doc.paragraphs[j].text.strip()
                if texto_j and len(texto_j) > 20:
                    texto_simplificado = simplificar_texto(texto_j)
                    if texto_simplificado != texto_j:
                        doc.paragraphs[j].clear()
                        doc.paragraphs[j].add_run(texto_simplificado)
            break
    
    # 6. HUMANIZAR SECCIÓN 3.2 (PROCESO USUARIO)
    print("\n6. Humanizando sección de proceso del usuario...")
    
    for i, para in enumerate(doc.paragraphs):
        texto = para.text.strip()
        if "3.2" in texto or "Mecanismo Detrás" in texto or "Mecanismo Operativo" in texto:
            print(f"  Encontrada sección proceso usuario en párrafo {i}")
            
            for j in range(i + 1, min(i + 15, len(doc.paragraphs))):
                texto_j = doc.paragraphs[j].text.strip()
                if texto_j and len(texto_j) > 20:
                    texto_simplificado = simplificar_texto(texto_j)
                    if texto_simplificado != texto_j:
                        doc.paragraphs[j].clear()
                        doc.paragraphs[j].add_run(texto_simplificado)
            break
    
    # Guardar documento actualizado
    output_path = 'Plan_Negocio_Treqe_HUMANIZADO_ACTUALIZADO.docx'
    doc.save(output_path)
    
    print(f"\n✅ Documento guardado: {output_path}")
    print("✅ Secciones actualizadas:")
    print("   - Algoritmo y conclusiones")
    print("   - Sistema de scoring (8.4)")
    print("   - Sistema de envíos (8.3)")
    print("   - Sistema de garantías (8.2)")
    print("   - Proceso del usuario (3.2)")
    
    return output_path

if __name__ == '__main__':
    main()