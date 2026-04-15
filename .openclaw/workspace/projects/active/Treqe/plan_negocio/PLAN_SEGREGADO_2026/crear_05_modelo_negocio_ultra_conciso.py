#!/usr/bin/env python3
# Script para crear la Sección 05 ULTRA CONCISA: MODELO_NEGOCIO.docx
# Versión que realmente sigue el estándar del documento de referencia (6-8KB)

import os
from docx import Document

def crear_seccion_05_ultra_conciso():
    """Crear la sección 05 ultra concisa"""
    
    doc = Document()
    
    # 5. MODELO DE NEGOCIO
    doc.add_heading('5. MODELO DE NEGOCIO', 1)
    
    p1 = doc.add_paragraph()
    p1.add_run('Treqe resuelve la paradoja de liquidez del mercado de segunda mano (8.200M€ España 2026) mediante ruedas de intercambio inteligentes y un modelo de ingresos multicapa.')
    
    # 5.1 Business Model Canvas
    doc.add_heading('5.1 Business Model Canvas', 2)
    
    p2 = doc.add_paragraph()
    p2.add_run('• Segmentos: 28M usuarios + comercios + empresas\n')
    p2.add_run('• Propuesta Valor: Transformar objetos no usados en deseados\n')
    p2.add_run('• Canales: App móvil (94%) + web + API\n')
    p2.add_run('• Relación: Reputación gamificada + soporte\n')
    p2.add_run('• Ingresos: 3% comisión + 9,99€/mes + SaaS\n')
    p2.add_run('• Recursos: Algoritmo patentado + equipo\n')
    p2.add_run('• Actividades: Matching + verificación\n')
    p2.add_run('• Socios: Stripe + AWS + legal\n')
    p2.add_run('• Costes: Desarrollo 40% + marketing 30%')
    
    # 5.2 Flujos de Ingresos
    doc.add_heading('5.2 Flujos de Ingresos', 2)
    
    p3 = doc.add_paragraph()
    p3.add_run('Ejemplo: Ana (iPhone 12, 400€) ↔ Carlos (bicicleta, 450€ + 50€). Treqe: 3% de 450€ = 13,50€.\n\n')
    p3.add_run('1. Transacciones: 3% comisión (70% año 1)\n')
    p3.add_run('2. Suscripción: 9,99€/mes, matching 30% más rápido\n')
    p3.add_run('3. Empresas: SaaS white-label (IKEA, Decathlon)\n')
    p3.add_run('4. Data: Insights mercado para retailers')
    
    # 5.3 Estructura Legal
    doc.add_heading('5.3 Estructura Legal', 2)
    
    p4 = doc.add_paragraph()
    p4.add_run('• Sociedad Limitada: 3.000€ capital, 25% impuestos\n')
    p4.add_run('• Propiedad Intelectual: Patente algoritmo + marca\n')
    p4.add_run('• Contratos: Ley 7/2022 economía colaborativa + GDPR\n')
    p4.add_run('• Fondo Garantía: 50.000€ + seguro 1M€')
    
    # 5.4 Proyecciones
    doc.add_heading('5.4 Proyecciones 2026-2029', 2)
    
    p5 = doc.add_paragraph()
    p5.add_run('Año 1: 10k usuarios, 750k€ volumen, 52,5k€ ingresos\n')
    p5.add_run('Año 2: 50k usuarios, 4,5M€ volumen, 382,5k€ ingresos\n')
    p5.add_run('Año 3: 200k usuarios, 18M€ volumen, 1,53M€ ingresos\n\n')
    p5.add_run('Punto equilibrio: Mes 18 (15k usuarios)\n')
    p5.add_run('ROI: 5,2x en 5 años (TIR 38%)\n')
    p5.add_run('Valoración Año 3: 12-15M€')
    
    # 5.5 Estrategia
    doc.add_heading('5.5 Estrategia Crecimiento', 2)
    
    p6 = doc.add_paragraph()
    p6.add_run('5 fases: Internal (3m) → Alpha (6m) → Beta (9m) → Early Access (12m) → Full Launch (18m)\n\n')
    p6.add_run('Métricas: CAC 15€, LTV 360€, LTV:CAC 24:1, Retención 65%')
    
    # 5.6 Tecnología
    doc.add_heading('5.6 Ventaja Tecnológica', 2)
    
    p7 = doc.add_paragraph()
    p7.add_run('Algoritmo matching circular patentable (k=4 máximo, 5min ejecución).\n')
    p7.add_run('Experiencia mobile-first (94% transacciones móvil).\n')
    p7.add_run('Sistema reputación blockchain-inspired.\n')
    p7.add_run('Interfaz "brutalista digital con toques orgánicos".')
    
    # 5.7 Conclusión
    doc.add_heading('5.7 Conclusión', 2)
    
    p8 = doc.add_paragraph()
    p8.add_run('Círculo virtuoso: Más usuarios → más matches → más valor → más comisiones → más mejora plataforma → más usuarios.\n\n')
    p8.add_run('1% mercado español = 280k usuarios = 8,4M€ ingresos anuales.\n')
    p8.add_run('No vendemos una app. Hacemos posible que lo que ya tienes se convierta en lo que realmente quieres.')
    
    # Guardar
    output_path = "05_MODELO_NEGOCIO_ULTRACONCISO.docx"
    doc.save(output_path)
    
    tamaño = os.path.getsize(output_path)
    
    print(f"SECCION 05 ULTRA CONCISA CREADA")
    print(f"TAMAÑO: {tamaño:,} bytes")
    
    # Reemplazar si es suficientemente pequeña
    if tamaño < 10000:  # Menos de 10KB
        if os.path.exists("05_MODELO_NEGOCIO.docx"):
            backup = "05_MODELO_NEGOCIO_ORIGINAL.docx"
            os.rename("05_MODELO_NEGOCIO.docx", backup)
            print(f"Backup creado: {backup}")
        
        os.rename(output_path, "05_MODELO_NEGOCIO.docx")
        print(f"✓ NUEVA VERSIÓN: 05_MODELO_NEGOCIO.docx ({tamaño:,} bytes)")
        return "05_MODELO_NEGOCIO.docx", tamaño
    else:
        print(f"✗ DEMASIADO GRANDE: {tamaño:,} bytes (manteniendo versión anterior)")
        return output_path, tamaño

if __name__ == "__main__":
    os.chdir(os.path.dirname(os.path.abspath(__file__)))
    archivo, tamaño = crear_seccion_05_ultra_conciso()