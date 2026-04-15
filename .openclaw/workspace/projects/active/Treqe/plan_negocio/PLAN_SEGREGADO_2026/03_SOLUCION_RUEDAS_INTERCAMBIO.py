from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from datetime import datetime

# Crear documento
doc = Document()

# Portada
title = doc.add_heading('DOCUMENTO 03: SOLUCIÓN', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('PLAN DE NEGOCIO TREQE - ESTRUCTURA SEGREGADA\n')
run.font.size = Pt(18)
run = sub.add_run('Sección 3 de 9: Ruedas de Intercambio Inteligente\n')
run.font.size = Pt(14)
run = sub.add_run('Generado: 27/02/2026 10:30\n')
run.font.size = Pt(11)
run = sub.add_run('Basado en documento original del 25 febrero 9:29\n')
run.font.size = Pt(11)
run = sub.add_run('Skills: algorithm-solver + frontend-design + humanizer\n')
run.font.size = Pt(11)
run = sub.add_run('Estado: PARA REVISIÓN\n')
run.font.size = Pt(11)

doc.add_page_break()

# Sección 3.1
doc.add_heading('3. LA SOLUCIÓN TREQE: RUEDAS DE INTERCAMBIO INTELIGENTE', 1)

doc.add_heading('3.1 Un Concepto que Supera Limitaciones Históricas', 2)
doc.add_paragraph('Si hay algo que la historia nos ha enseñado sobre el trueque es que, a pesar de su aparente simplicidad, tiene una limitación fundamental que ha impedido que escale como modelo comercial viable. Nos referimos a lo que los economistas llaman "el problema de la doble coincidencia de deseos".')

doc.add_paragraph('Durante siglos, esta limitación ha confinado el trueque a escalas reducidas, a intercambios ocasionales entre personas que casualmente querían exactamente lo que la otra ofrecía. Pero, ¿y si pudiéramos superar esta limitación? ¿Y si pudiéramos transformar lo que históricamente ha sido un proceso azaroso en uno sistemático y escalable?')

doc.add_paragraph('Es precisamente aquí donde Treqe introduce lo que consideramos una innovación conceptual significativa: las "ruedas de intercambio inteligente". La idea es tan elegante en su concepción como poderosa en su aplicación. En lugar de limitarnos a buscar coincidencias directas entre dos personas (una tarea estadísticamente improbable), permitimos que tres, cuatro o incluso cinco usuarios participen en cadenas circulares donde cada uno recibe exactamente lo que desea, aunque no directamente de la persona a la que entrega su artículo.')

doc.add_paragraph('Lo que hace especial a este sistema no es solo su capacidad para resolver un problema matemático, sino su comprensión profunda de la naturaleza humana. Reconoce que nuestras necesidades son complejas, que los objetos que poseemos tienen tanto valor sentimental como económico, y que la falta de liquidez no debería ser una barrera insalvable para satisfacer nuestras necesidades de renovación y cambio.')

# Sección 3.2 (con algorithm-solver y frontend-design)
doc.add_heading('3.2 El Mecanismo Operativo Paso a Paso', 2)
doc.add_paragraph('Para comprender cómo Treqe transforma una idea conceptual en una experiencia tangible, es necesario analizar el mecanismo operativo paso a paso. Cada fase ha sido diseñada no solo pensando en la eficiencia técnica, sino especialmente en la experiencia humana.')

doc.add_heading('Fase 1: El Arte de Declarar lo que Tenemos y lo que Deseamos (skill: frontend-design)', 3)
doc.add_paragraph('El proceso comienza con algo aparentemente simple pero profundamente significativo: invitar a los usuarios a reflexionar sobre lo que poseen y lo que realmente necesitan. No se trata de un inventario frío, sino de un ejercicio de autoconocimiento práctico.', style='List Bullet')
doc.add_paragraph('Cuando un usuario ingresa a Treqe, se le pide que complete dos listas:', style='List Bullet')
doc.add_paragraph('• "Esto tengo": Los artículos que actualmente posee pero que ya no se alinean con su vida actual', style='List Bullet')
doc.add_paragraph('• "Esto quiero": Los objetos que necesita para mejorar su espacio, su comodidad, su funcionalidad', style='List Bullet')
doc.add_paragraph('Para cada artículo, el sistema solicita información que va más allá de lo meramente descriptivo. Se piden fotografías desde múltiples ángulos, no por mero formalismo, sino porque entendemos que ver un objeto es comenzar a establecer una relación con él.', style='List Bullet')

doc.add_heading('Fase 2: Cuando las Matemáticas se Convierten en Conectores Humanos (skill: algorithm-solver)', 3)
doc.add_paragraph('Una vez que los usuarios han declarado sus preferencias, el sistema transforma esta información en una estructura matemática: un grafo dirigido. Puede sonar técnico, pero en esencia es bastante simple:', style='List Bullet')
doc.add_paragraph('• Cada usuario se convierte en un "nodo"', style='List Bullet')
doc.add_paragraph('• Las flechas (aristas) van desde lo que una persona tiene hacia lo que otra persona quiere', style='List Bullet')
doc.add_paragraph('Utilizando algoritmos de teoría de grafos (específicamente búsqueda en profundidad optimizada), el sistema busca ciclos cerrados de 3 a 5 nodos. Cada ciclo representa una posible cadena de intercambio donde todos obtienen lo que desean.', style='List Bullet')
doc.add_paragraph('Lo notable es la velocidad: cada búsqueda tiene un timeout de 500 milisegundos, garantizando que el sistema sea rápido y responsivo incluso a medida que crece.', style='List Bullet')

doc.add_heading('Fase 3: La Equidad como Principio de Diseño (skill: algorithm-solver)', 3)
doc.add_paragraph('En el mundo real, es raro que dos artículos tengan exactamente el mismo valor económico. Por eso, cuando Treqe identifica un ciclo viable, no se limita a proponer intercambios directos. En su lugar, utiliza programación lineal (con el algoritmo PuLP) para calcular las compensaciones monetarias óptimas.', style='List Bullet')
doc.add_paragraph('La optimización considera múltiples factores:', style='List Bullet')
doc.add_paragraph('• Minimizar el flujo monetario total (reducir al máximo las transferencias)', style='List Bullet')
doc.add_paragraph('• Distribuir las compensaciones de manera equitativa', style='List Bullet')
doc.add_paragraph('• Respetar las preferencias individuales (algunos usuarios prefieren pagar menos aunque reciban menos valor, otros prefieren lo contrario)', style='List Bullet')
doc.add_paragraph('• Garantizar que nadie se sienta en desventaja en la transacción', style='List Bullet')

doc.add_heading('Fase 4: La Negociación como Conversación, no como Confrontación (skill: frontend-design)', 3)
doc.add_paragraph('Una vez identificado un ciclo y calculadas las compensaciones, Treqe no impone la solución. En su lugar, crea un espacio de negociación facilitada. Los participantes se conectan a través de un chat grupal en tiempo real (utilizando WebSockets), donde pueden:', style='List Bullet')
doc.add_paragraph('• Ver la propuesta inicial del sistema', style='List Bullet')
doc.add_paragraph('• Hacer preguntas sobre los artículos involucrados', style='List Bullet')
doc.add_paragraph('• Solicitar fotografías adicionales o información complementaria', style='List Bullet')
doc.add_paragraph('• Proponer ajustes menores a los términos', style='List Bullet')
doc.add_paragraph('• Confirmar explícitamente su acuerdo', style='List Bullet')
doc.add_paragraph('Este proceso de negociación no es una formalidad burocrática, sino un espacio para construir confianza. Los usuarios no están interactuando con un algoritmo frío, sino con otras personas reales que, como ellos, buscan mejorar sus vidas a través del intercambio.', style='List Bullet')

doc.add_page_break()

# Sección 3.3 (con humanizer)
doc.add_heading('3.3 Un Caso que Ilumina el Concepto: Ana, Carlos y Beatriz (skill: humanizer)', 2)
doc.add_paragraph('Para comprender realmente cómo funciona Treqe, nada mejor que analizar un caso concreto. No se trata de un ejemplo teórico, sino de una situación que refleja la realidad de millones de personas.')

doc.add_heading('Los Protagonistas y sus Historias:', 3)

doc.add_heading('Ana, 32 años, arquitecta en Barcelona:', 4)
doc.add_paragraph('Ana compró su bicicleta de montaña Specialized Rockhopper durante sus años universitarios, cuando el ciclismo era más que un hobby: era una pasión. Ahora, con un trabajo exigente y menos tiempo libre, la bicicleta pasa la mayor parte del tiempo aparcada en el pequeño balcón de su apartamento.', style='List Bullet')
doc.add_paragraph('Lo que Ana realmente necesita es un sofá de diseño contemporáneo que se adapte mejor a su espacio y estilo actual.', style='List Bullet')
doc.add_paragraph('El problema no es de deseo, sino de recursos: aunque la bicicleta tiene un valor considerable (450€), Ana no tiene los 600€ que costaría el sofá que quiere.', style='List Bullet')

doc.add_heading('Carlos, 41 años, profesor universitario en Hospitalet de Llobregat:', 4)
doc.add_paragraph('Carlos invirtió en un sofá de diseño de calidad hace dos años, pero su vida ha cambiado. Ha comenzado un proyecto personal: un canal educativo en YouTube sobre historia del arte.', style='List Bullet')
doc.add_paragraph('Para producir contenido de calidad, necesita un ordenador portátil potente para edición de video (valorado en 800€).', style='List Bullet')
doc.add_paragraph('El sofá, aunque perfectamente funcional y valioso (600€), ya no es una prioridad. Como muchos proyectos personales, el suyo tiene un presupuesto limitado.', style='List Bullet')

doc.add_heading('Beatriz, 28 años, diseñadora gráfica freelance en Badalona:', 4)
doc.add_paragraph('Beatriz invirtió en un MacBook Pro M2 para su trabajo como diseñadora freelance. Recientemente, ha comenzado a trabajar de forma híbrida, yendo a una oficina compartida varios días a la semana.', style='List Bullet')
doc.add_paragraph('Necesita una bicicleta para sus desplazamientos urbanos (presupuesto: 450€), pero como freelance con ingresos variables, un desembolso de esa magnitud afectaría significativamente su liquidez mensual.', style='List Bullet')

doc.add_heading('El Problema Tradicional:', 3)
doc.add_paragraph('Por separado, cada uno enfrenta una situación frustrante:', style='List Bullet')
doc.add_paragraph('• Ana no puede conseguir su sofá porque Carlos no quiere su bicicleta', style='List Bullet')
doc.add_paragraph('• Carlos no puede conseguir su ordenador porque Beatriz no quiere su sofá', style='List Bullet')
doc.add_paragraph('• Beatriz no puede conseguir su bicicleta porque Ana no quiere su ordenador', style='List Bullet')
doc.add_paragraph('Es el clásico problema de la doble coincidencia de deseos: cada par podría intercambiar si sus deseos coincidieran, pero no lo hacen.', style='List Bullet')

doc.add_heading('La Solución Treqe:', 3)
doc.add_paragraph('El algoritmo identifica que, aunque ningún intercambio directo es posible, existe un ciclo cerrado perfecto entre los tres. Propone:', style='List Bullet')

doc.add_heading('Intercambios físicos:', 4)
doc.add_paragraph('1. Ana entrega su bicicleta a Beatriz', style='List Number')
doc.add_paragraph('2. Carlos entrega su sofá a Ana', style='List Number')
doc.add_paragraph('3. Beatriz entrega su ordenador a Carlos', style='List Number')

doc.add_heading('Compensaciones monetarias:', 4)
doc.add_paragraph('• Ana paga 150€ a Carlos (la diferencia entre el sofá que recibe y la bicicleta que entrega)', style='List Bullet')
doc.add_paragraph('• Carlos paga 200€ a Beatriz (la diferencia entre el ordenador que recibe y el sofá que entrega)', style='List Bullet')
doc.add_paragraph('• Beatriz recibe 350€ neto (150€ de Ana + 200€ de Carlos)', style='List Bullet')

doc.add_heading('Los Resultados Transformadores:', 3)
doc.add_heading('Ana:', 4)
doc.add_paragraph('Obtiene el sofá de diseño que necesita por solo 150€, en lugar de los 600€ que costaría nuevo. Su ahorro del 75% no es solo económico: es la posibilidad de tener un espacio que realmente refleja quién es ahora.', style='List Bullet')

doc.add_heading('Carlos:', 4)
doc.add_paragraph('Consigue el ordenador que necesita para su proyecto YouTube por 200€, en lugar de 800€. El ahorro del 75% representa meses de trabajo que puede dedicar a crear contenido en lugar de ahorrar para equipamiento.', style='List Bullet')

doc.add_heading('Beatriz:', 4)
doc.add_paragraph('Recibe la bicicleta que necesita para sus desplazamientos más 350€ de liquidez adicional. El valor total recibido (800€) le permite mejorar tanto su movilidad como su situación financiera.', style='List Bullet')

doc.add_paragraph('Lo que este caso nos enseña: Este ejemplo no es una anomalía estadística, sino una ilustración de cómo Treqe transforma problemas aparentemente insolubles en oportunidades de ganancia mutua. Lo notable no es solo el ahorro económico (aunque es significativo), sino cómo el sistema permite a cada persona avanzar en sus proyectos de vida sin las restricciones tradicionales de liquidez.')

doc.add_page_break()

# Sección 3.4
doc.add_heading('3.4 Innovaciones Diferenciales del Modelo Treqe', 2)
doc.add_paragraph('Treqe no es simplemente otra plataforma de trueque. Introduce varias innovaciones que lo diferencian fundamentalmente de cualquier solución existente:')

doc.add_heading('1. Escalabilidad algorítmica', 3)
doc.add_paragraph('• Mientras el trueque tradicional colapsa con más de 100 usuarios, nuestro sistema maneja eficientemente miles de usuarios simultáneos', style='List Bullet')
doc.add_paragraph('• El algoritmo está optimizado para encontrar soluciones en menos de 5 minutos incluso en mercados densos', style='List Bullet')
doc.add_paragraph('• La arquitectura serverless permite escalar automáticamente según la demanda', style='List Bullet')

doc.add_heading('2. Equidad económica automatizada', 3)
doc.add_paragraph('• No requiere que los usuarios negocien manualmente las diferencias de valor', style='List Bullet')
doc.add_paragraph('• El sistema calcula automáticamente las compensaciones óptimas', style='List Bullet')
doc.add_paragraph('• Garantiza que nadie salga perdiendo económicamente en la transacción', style='List Bullet')

doc.add_heading('3. Experiencia mobile-first completa (skill: frontend-design)', 3)
doc.add_paragraph('• Diseñado específicamente para móvil desde el primer día', style='List Bullet')
doc.add_paragraph('• Proceso completo (registro, matching, negociación, confirmación) optimizado para pantallas táctiles', style='List Bullet')
doc.add_paragraph('• PWA instalable que funciona como app nativa sin necesidad de tiendas de aplicaciones', style='List Bullet')

doc.add_heading('4. Sistema de confianza construido, no asumido', 3)
doc.add_paragraph('• Reputación granular basada en múltiples dimensiones (puntualidad, precisión descripciones, calidad embalaje)', style='List Bullet')
doc.add_paragraph('• Verificación por pasos que aumenta la confianza progresivamente', style='List Bullet')
doc.add_paragraph('• Mecanismos de disputa escalonados que evitan la necesidad de intervención legal inmediata', style='List Bullet')

doc.add_heading('