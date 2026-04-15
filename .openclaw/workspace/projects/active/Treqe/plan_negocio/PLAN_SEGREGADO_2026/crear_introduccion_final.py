#!/usr/bin/env python3
"""
Crear 01_INTRODUCCION.docx para Treqe
Basado EXACTAMENTE en Plan_Negocio_Treqe_100%_PROFESIONAL_FINAL.docx del 25 febrero
Con mejoras aplicando skills: humanizer + marketing-mode
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
import os
from datetime import datetime

def crear_introduccion():
    """Crear documento 01_INTRODUCCION.docx basado en documento original"""
    
    print("ðŸ“ Creando 01_INTRODUCCION.docx...")
    print("ðŸŽ¯ Referencia: Plan_Negocio_Treqe_100%_PROFESIONAL_FINAL.docx (25 febrero 9:29)")
    print("ðŸ”§ Skills aplicadas: humanizer (lenguaje natural), marketing-mode (datos persuasivos)")
    
    # Crear documento
    doc = Document()
    
    # ===== CONFIGURAR ESTILOS PROFESIONALES =====
    
    # Estilo Normal
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Calibri'
    style_normal.font.size = Pt(11)
    style_normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    style_normal.paragraph_format.space_after = Pt(6)
    
    # Estilo TÃ­tulo 1
    style_heading1 = doc.styles['Heading 1']
    style_heading1.font.name = 'Calibri'
    style_heading1.font.size = Pt(16)
    style_heading1.font.bold = True
    style_heading1.font.color.rgb = RGBColor(0, 0, 0)
    style_heading1.paragraph_format.space_before = Pt(18)
    style_heading1.paragraph_format.space_after = Pt(12)
    
    # Estilo TÃ­tulo 2
    style_heading2 = doc.styles['Heading 2']
    style_heading2.font.name = 'Calibri'
    style_heading2.font.size = Pt(14)
    style_heading2.font.bold = True
    style_heading2.font.color.rgb = RGBColor(0, 0, 0)
    style_heading2.paragraph_format.space_before = Pt(12)
    style_heading2.paragraph_format.space_after = Pt(6)
    
    # Estilo TÃ­tulo 3
    style_heading3 = doc.styles['Heading 3']
    style_heading3.font.name = 'Calibri'
    style_heading3.font.size = Pt(12)
    style_heading3.font.bold = True
    style_heading3.font.color.rgb = RGBColor(0, 0, 0)
    style_heading3.paragraph_format.space_before = Pt(8)
    style_heading3.paragraph_format.space_after = Pt(4)
    
    # ===== PORTADA DEL DOCUMENTO INDIVIDUAL =====
    
    # TÃ­tulo principal
    title = doc.add_heading('DOCUMENTO 01: INTRODUCCIÃ“N', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(28)
    title.runs[0].font.bold = True
    title.runs[0].font.name = 'Calibri'
    
    # SubtÃ­tulo
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run('PLAN DE NEGOCIO TREQE - ESTRUCTURA SEGREGADA\n').font.size = Pt(18)
    sub.add_run('SecciÃ³n 1 de 9: Contexto Actual del Mercado\n').font.size = Pt(14)
    
    # InformaciÃ³n
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run(f'Generado: {datetime.now().strftime("%d/%m/%Y %H:%M")}\n').font.size = Pt(11)
    info.add_run('Basado en: Plan_Negocio_Treqe_100%_PROFESIONAL_FINAL.docx (25 febrero 9:29)\n').font.size = Pt(11)
    info.add_run('Skills aplicadas: humanizer + marketing-mode\n').font.size = Pt(11)
    info.add_run('Estado: PARA REVISIÃ“N\n').font.size = Pt(11)
    
    doc.add_page_break()
    
    # ===== ÃNDICE DE ESTE DOCUMENTO =====
    doc.add_heading('ÃNDICE DE ESTE DOCUMENTO', 1)
    
    # Tabla Ã­ndice
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    # Encabezados
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'SECCIÃ“N'
    hdr_cells[1].text = 'DESCRIPCIÃ“N'
    hdr_cells[2].text = 'PÃGINA'
    
    # Contenido
    secciones = [
        ('1', 'INTRODUCCIÃ“N: EL CONTEXTO ACTUAL DEL MERCADO', '3'),
        ('1.1', 'La TransformaciÃ³n de un Sector Tradicional', '3'),
        ('1.2', 'Datos Cuantitativos Actualizados (2025-2026)', '4'),
        ('1.3', 'El Panorama Competitivo Actual', '6'),
        ('1.4', 'Tendencias Emergentes que Definen el Futuro', '8'),
        ('', 'Resumen Ejecutivo de esta SecciÃ³n', '10')
    ]
    
    for seccion in secciones:
        row_cells = table.add_row().cells
        row_cells[0].text = seccion[0]
        row_cells[1].text = seccion[1]
        row_cells[2].text = seccion[2]
    
    doc.add_page_break()
    
    # ===== SECCIÃ“N 1.1 (MEJORADA CON HUMANIZER) =====
    doc.add_heading('1. INTRODUCCIÃ“N: EL CONTEXTO ACTUAL DEL MERCADO DE SEGUNDA MANO EN ESPAÃ‘A', 1)
    
    doc.add_heading('1.1 La TransformaciÃ³n de un Sector Tradicional', 2)
    
    # Texto original mejorado con humanizer
    p = doc.add_paragraph()
    p.add_run('Si miramos atrÃ¡s una dÃ©cada, el mercado de segunda mano en EspaÃ±a era muy diferente. ').bold = False
    p.add_run('Lo que antes se asociaba principalmente con Ã©pocas de crisis econÃ³mica o con personas con presupuestos limitados, hoy se ha transformado en algo mucho mÃ¡s significativo: un componente fundamental de cÃ³mo consumimos en el siglo XXI.').bold = False
    
    p = doc.add_paragraph()
    p.add_run('Esta evoluciÃ³n no ha sido casual. ').bold = False
    p.add_run('Lo que empezÃ³ como una soluciÃ³n prÃ¡ctica ante restricciones econÃ³micas se ha convertido en un autÃ©ntico movimiento cultural. ').bold = True
    p.add_run('Hoy, comprar de segunda mano no es solo una opciÃ³n para ahorrar dinero; es una declaraciÃ³n de valores: sostenibilidad medioambiental, consumo consciente, y una relaciÃ³n mÃ¡s inteligente con las cosas que nos rodean.').bold = False
    
    p = doc.add_paragraph()
    p.add_run('La transformaciÃ³n responde a cambios profundos en cÃ³mo pensamos como sociedad. ').bold = False
    p.add_run('Hay una mayor conciencia sobre el impacto ambiental de nuestro consumo, y estamos reevaluando lo que realmente significa "valor" en un mundo donde tenemos abundancia material pero recursos limitados.').bold = False
    
    p = doc.add_paragraph()
    p.add_run('El mercado de segunda mano ha dejado de ser un refugio en tiempos difÃ­ciles para convertirse en una elecciÃ³n activa y valorada por millones de personas. ').bold = False
    p.add_run('Personas que buscan alternativas mÃ¡s inteligentes, mÃ¡s sostenibles y mÃ¡s satisfactorias que el simple "comprar y tirar" del modelo tradicional.').bold = False
    
    # ===== SECCIÃ“N 1.2 (MEJORADA CON MARKETING-MODE) =====
    doc.add_heading('1.2 Datos Cuantitativos Actualizados (2025-2026)', 2)
    
    p = doc.add_paragraph()
    p.add_run('Para entender la magnitud real de este cambio, necesitamos mirar los nÃºmeros mÃ¡s recientes. ').bold = False
    p.add_run('Las cifras que presentamos aquÃ­ vienen de mÃºltiples fuentes confiables: informes del sector, estudios de mercado independientes, y datos oficiales de las principales plataformas.').bold = False
    
    doc.add_heading('ðŸ“ˆ Volumen econÃ³mico del mercado (skill: marketing-mode):', 3)
    doc.add_paragraph('â€¢ EstimaciÃ³n para 2026: 8.200 millones de euros en transacciones', style='List Bullet')
    doc.add_paragraph('â€¢ Crecimiento desde 2020: +42% (Â¡casi la mitad mÃ¡s en solo 6 aÃ±os!)', style='List Bullet')
    doc.add_paragraph('â€¢ ProyecciÃ³n 2027: SuperarÃ¡ los 9.500 millones de euros', style='List Bullet')
    doc.add_paragraph('â€¢ PosiciÃ³n internacional: EspaÃ±a es el 4Âº mercado europeo, solo detrÃ¡s de Reino Unido, Alemania y Francia', style='List Bullet')
    
    doc.add_heading('ðŸ‘¥ PenetraciÃ³n en la poblaciÃ³n:', 3)
    doc.add_paragraph('â€¢ Usuarios activos: 28 millones de espaÃ±oles (47% de la poblaciÃ³n total)', style='List Bullet')
    doc.add_paragraph('â€¢ Frecuencia de uso: 62% consulta plataformas semanalmente', style='List Bullet')
    doc.add_paragraph('â€¢ Edad media: 34 aÃ±os (el grueso estÃ¡ entre 25 y 45 aÃ±os)', style='List Bullet')
    doc.add_paragraph('â€¢ DistribuciÃ³n geogrÃ¡fica: Mayor en ciudades (65%) que en zonas rurales (35%)', style='List Bullet')
    
    doc.add_heading('ðŸ’° Comportamiento de gasto:', 3)
    doc.add_paragraph('â€¢ Gasto medio anual por usuario: 1.850 euros', style='List Bullet')
    doc.add_paragraph('â€¢ Incremento vs 2016: +142% (en 2016 era 766 euros)', style='List Bullet')
    doc.add_paragraph('â€¢ Ticket medio por transacciÃ³n: 85-100 euros (depende categorÃ­a)', style='List Bullet')
    doc.add_paragraph('â€¢ Frecuencia: 3,4 transacciones por usuario al aÃ±o', style='List Bullet')
    
    doc.add_heading('ðŸ“± TecnologÃ­a y hÃ¡bitos (skill: marketing-mode):', 3)
    doc.add_paragraph('â€¢ Mobile-first: 94% de transacciones empiezan en el mÃ³vil', style='List Bullet')
    doc.add_paragraph('â€¢ Preferencia apps: 88% usa apps especÃ­ficas, no el navegador web', style='List Bullet')
    doc.add_paragraph('â€¢ Horarios pico: 20:00-23:00 horas, especialmente domingos', style='List Bullet')
    doc.add_paragraph('â€¢ Tiempo decisiÃ³n: 3,2 dÃ­as de media desde primer contacto hasta compra', style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run('Estos nÃºmeros no solo muestran un cambio cuantitativo, sino cualitativo. ').bold = False
    p.add_run('El usuario de segunda mano actual es mÃ¡s activo, mÃ¡s informado y mÃ¡s exigente. ').bold = True
    p.add_run('Ya no se conforma con encontrar algo barato; busca calidad, autenticidad, y una experiencia de compra que alinee con sus valores personales.').bold = False
    
    doc.add_page_break()
    
    # ===== SECCIÃ“N 1.3 (MEJORADA) =====
    doc.add_heading('1.3 El Panorama Competitivo Actual', 2)
    
    p = doc.add_paragraph()
    p.add_run('El mercado espaÃ±ol tiene una estructura competitiva bien definida, con actores claros que dominan diferentes segmentos:').bold = False
    
    doc.add_heading('ðŸ† Wallapop: El lÃ­der indiscutible', 3)
    doc.add_paragraph('â€¢ Usuarios: ~15 millones en EspaÃ±a', style='List Bullet')
    doc.add_paragraph('â€¢ Modelo: Amplitud de categorÃ­as + efecto de red masivo', style='List Bullet')
    doc.add_paragraph('â€¢ Comisiones: 5% + 0,90â‚¬ fijo por venta', style='List Bullet')
    doc.add_paragraph('â€¢ Fortaleza: Todos lo conocen, todos lo usan', style='List Bullet')
    
    doc.add_heading('ðŸ‘— Vinted: Especialista en moda', 3)
    doc.add_paragraph('â€¢ Usuarios activos: ~4,5 millones', style='List Bullet')
    doc.add_paragraph('â€¢ EspecializaciÃ³n: Moda de segunda mano', style='List Bullet')
    doc.add_paragraph('â€¢ Comunidad: Muy sÃ³lida, sistema de reputaciÃ³n sofisticado', style='List Bullet')
    doc.add_paragraph('â€¢ Comisiones: Pueden llegar al 8-9% total', style='List Bullet')
    
    doc.add_heading('ðŸ“± Facebook Marketplace: El gigante dormido', 3)
    doc.add_paragraph('â€¢ Potencial: ~20 millones de usuarios (integraciÃ³n Facebook)', style='List Bullet')
    doc.add_paragraph('â€¢ Modelo: Gratuito para particulares', style='List Bullet')
    doc.add_paragraph('â€¢ Debilidad: Experiencia bÃ¡sica, seguridad limitada', style='List Bullet')
    doc.add_paragraph('â€¢ Oportunidad: Masivo pero poco optimizado', style='List Bullet')
    
    doc.add_heading('ðŸ“° Milanuncios: El tradicional', 3)
    doc.add_paragraph('â€¢ Cuota: ~10% del mercado', style='List Bullet')
    doc.add_paragraph('â€¢ Usuarios: Mayor edad, menos digitalizados', style='List Bullet')
    doc.add_paragraph('â€¢ Ventaja: Reconocimiento de marca histÃ³rico', style='List Bullet')
    doc.add_paragraph('â€¢ Debilidad: Interfaz anticuada', style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run('El espacio que NINGUNA de estas plataformas cubre actualmente: ').bold = False
    p.add_run('el trueque estructurado y escalable. ').bold = True
    p.add_run('Todas se centran exclusivamente en compraventa con dinero, dejando completamente desatendida la demanda real de intercambios directos.').bold = False
    
    # ===== SECCIÃ“N 1.4 (MEJORADA) =====
    doc.add_heading('1.4 Tendencias Emergentes que Definen el Futuro', 2)
    
    p = doc.add_paragraph()
    p.add_run('Cinco tendencias clave estÃ¡n reconfigurando el sector y creando oportunidades Ãºnicas:').bold = False
    
    doc.add_heading('1. PremiumizaciÃ³n acelerada', 3)
    doc.add_paragraph('â€¢ Crecimiento: +125% interanual en artÃ­culos de segunda mano de alta calidad', style='List Bullet')
    doc.add_paragraph('â€¢ Oportunidad: Marcas de lujo y productos premium encuentran mercado Ã¡vido', style='List Bullet')
    doc.add_paragraph('â€¢ Insight: La segunda mano ya no es "lo barato", es "lo inteligente"', style='List Bullet')
    
    doc.add_heading('2. Sostenibilidad como driver principal', 3)
    doc.add_paragraph('â€¢ MotivaciÃ³n: 68% de usuarios menciona razones ecolÃ³gicas como factor clave', style='List Bullet')
    doc.add_paragraph('â€¢ Cambio: La economÃ­a circular pasa de concepto teÃ³rico a prÃ¡ctica cotidiana', style='List Bullet')
    doc.add_paragraph('â€¢ ImplicaciÃ³n: Valores personales guÃ­an decisiones de consumo', style='List Bullet')
    
    doc.add_heading('3. Importancia de comunidades locales', 3)
    doc.add_paragraph('â€¢ Realidad: Las transacciones mÃ¡s exitosas ocurren en radios reducidos (<50km)', style='List Bullet')
    doc.add_paragraph('â€¢ Insight: La confianza geogrÃ¡fica supera a la escala global pura', style='List Bullet')
    doc.add_paragraph('â€¢ Oportunidad: Mercados hiperlocales con alta densidad de confianza', style='List Bullet')
    
    doc.add_heading('4.
 RegulaciÃ³n emergente', 3)
    doc.add_paragraph('â€¢ Contexto: Nuevas normativas fiscales para ventas entre particulares (2025+)', style='List Bullet')
    doc.add_paragraph('â€¢ Tendencia: ProfesionalizaciÃ³n progresiva del sector', style='List Bullet')
    doc.add_paragraph('â€¢ ImplicaciÃ³n: Mayor formalidad, mayor confianza del usuario', style='List Bullet')
    
    doc.add_heading('5. Experiencia mobile-first absoluta', 3)
    doc.add_paragraph('â€¢ DemografÃ­a: 75% del volumen viene de millennials y generaciÃ³n Z', style='List Bullet')
    doc.add_paragraph('â€¢ Expectativa: Experiencia perfectamente optimizada para mÃ³vil', style='List Bullet')
    doc.add_paragraph('â€¢ Imperativo: No tener app mÃ³vil excelente = no existir', style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run('Estas tendencias crean un contexto especialmente favorable para Treqe. ').bold = False
    p.add_run('Nuestra plataforma se diseÃ±a especÃ­ficamente para atender estas demandas emergentes del consumidor actual, no las necesidades del consumidor de hace una dÃ©cada.').bold = True
    
    # ===== RESUMEN EJECUTIVO DE ESTA SECCIÃ“N =====
    doc.add_page_break()
    doc.add_heading('ðŸ“‹ RESUMEN EJECUTIVO DE ESTA SECCIÃ“N', 1)
    
    p = doc.add_paragraph()
    p.add_run('Esta secciÃ³n establece el contexto fundamental para entender por quÃ© Treqe tiene sentido AHORA:').bold = True
    
    doc.add_heading('ðŸŽ¯ PUNTOS CLAVE:', 2)
    
    doc.add_heading('1. El mercado ha cambiado radicalmente', 3)
    doc.add_paragraph('â€¢ De "opciÃ³n econÃ³mica en crisis" a "movimiento cultural con valores"', style='List Bullet')
    doc.add_paragraph('â€¢ Usuarios mÃ¡s exigentes: buscan calidad + autenticidad + valores', style='List Bullet')
    doc.add_paragraph('â€¢ Crecimiento sostenido: +42% desde 2020, 8.200Mâ‚¬ en 2026', style='List Bullet')
    
    doc.add_heading('2. La competencia deja un espacio vacante', 3)
    doc.add_paragraph('â€¢ Wallapop/Vinted: Compraventa monetaria (pierdes valor)', style='List Bullet')
    doc.add_paragraph('â€¢ Facebook Marketplace: Gratis pero inseguro', style='List Bullet')
    doc.add_paragraph('â€¢ Milanuncios: Tradicional y desactualizado', style='List Bullet')
    doc.add_paragraph('â€¢ Espacio vacante: Trueque estructurado y escalable', style='List Bullet')
    
    doc.add_heading('3. Las tendencias juegan a favor', 3)
    doc.add_paragraph('â€¢ PremiumizaciÃ³n: La segunda mano ya no es "lo barato"', style='List Bullet')
    doc.add_paragraph('â€¢ Sostenibilidad: Valores ecolÃ³gicos guÃ­an decisiones', style='List Bullet')
    doc.add_paragraph('â€¢ Mobile-first: Experiencia mÃ³vil perfecta es obligatoria', style='List Bullet')
    doc.add_paragraph('â€¢ Comunidades locales: Confianza geogrÃ¡fica > escala global', style='List Bullet')
    
    doc.add_heading('4. La oportunidad es cuantificable', 3)
    doc.add_paragraph('â€¢ Mercado total: 8.200Mâ‚¬ (2026)', style='List Bullet')
    doc.add_paragraph('â€¢ Segmento trueque potencial: 1.230Mâ‚¬ (15%)', style='List Bullet')
    doc.add_paragraph('â€¢ Usuarios insatisfechos: 17M (60% de usuarios activos)', style='List Bullet')
    doc.add_paragraph('â€¢ Valor medio disponible: 1.200â‚¬ por usuario', style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run('Esta introducciÃ³n establece el "por quÃ© ahora" de Treqe. ').bold = False
    p.add_run('No estamos resolviendo un problema pequeÃ±o o marginal; estamos abordando una oportunidad masiva en un mercado en transformaciÃ³n, con competidores que han dejado un espacio claro sin cubrir.').bold = True
    
    # ===== INFORMACIÃ“N DE VERSIÃ“N =====
    doc.add_page_break()
    info_final = doc.add_paragraph()
    info_final.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_final.add_run('--- FIN DEL DOCUMENTO 01 ---\n').font.size = Pt(14)
    info_final.add_run('\n')
    info_final.add_run('DOCUMENTO: 01_INTRODUCCION.docx\n').font.size = Pt(12)
    info_final.add_run('PARTE DE: Plan de Negocio Treqe - Estructura Segregada\n').font.size = Pt(12)
    info_final.add_run(f'GENERADO: {datetime.now().strftime("%d/%m/%Y %H:%M")}\n').font.size = Pt(10)
    info_final.add_run('BASADO EN: Plan_Negocio_Treqe_100%_PROFESIONAL_FINAL.docx (25 febrero 9:29)\n').font.size = Pt(10)
    info_final.add_run('SKILLS APLICADAS: humanizer (lenguaje natural), marketing-mode (datos persuasivos)\n').font.size = Pt(10)
    info_final.add_run('ESTADO: Para revisiÃ³n - VersiÃ³n 1.0\n').font.size = Pt(10)
    
    # Guardar documento
    output_path = os.path.join(os.path.dirname(__file__), '01_INTRODUCCION.docx')
    doc.save(output_path)
    
    return output_path

if __name__ == '__main__':
    try:
        print("ðŸš€ CREANDO 01_INTRODUCCION.DOCX")
        print("ðŸŽ¯ Referencia exacta: Plan_Negocio_Treqe_100%_PROFESIONAL_FINAL.docx (25 febrero 9:29)")
        print("ðŸ”§ Skills aplicadas: humanizer + marketing-mode")
        print("ðŸ“Š Objetivo: Mantener estructura original + mejorar lenguaje + hacer datos mÃ¡s persuasivos")
        
        output_file = crear_introduccion()
        file_size = os.path.getsize(output_file)
        
        print(f"\nâœ… Â¡01_INTRODUCCION.DOCX CREADO!")
        print(f"ðŸ“„ Archivo: {output_file}")
        print(f"ðŸ“ TamaÃ±o: {file_size:,} bytes")
        print(f"ðŸ“… Generado: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
        
        print("\nðŸŽ¯ CONTENIDO (4 secciones + resumen):")
        print("   1.1 La TransformaciÃ³n de un Sector Tradicional (humanizer aplicado)")
        print("   1.2 Datos Cuantitativos Actualizados 2025-2026 (marketing-mode aplicado)")
        print("   1.3 El Panorama Competitivo Actual (anÃ¡lisis competencia)")
        print("   1.4 Tendencias Emergentes que Definen el Futuro (5 tendencias clave)")
        print("   + Resumen ejecutivo de la secciÃ³n")
        
        print("\nðŸ”§ MEJORAS APLICADAS:")
        print("   â€¢ humanizer: Lenguaje mÃ¡s natural (\"Si miramos atrÃ¡s\" vs \"Si echamos la vista atrÃ¡s\")")
        print("   â€¢ humanizer: Frases mÃ¡s conversacionales pero profesionales")
        print("   â€¢ marketing-mode: Datos presentados de manera mÃ¡s persuasiva")
        print("   â€¢ marketing-mode: Ã‰nfasis en oportunidades (Â¡casi la mitad mÃ¡s en solo 6 aÃ±os!)")
        print("   â€¢ Estructura: Ãndice interno + resumen ejecutivo al final")
        print("   â€¢ Formato: Portada especÃ­fica para este documento individual")
        
        print("\nðŸ“‹ PRÃ“XIMO PASO:")
        print("   Revisar este documento y dar feedback antes de continuar con 02_PROBLEMA_PARADOJA_LIQUIDEZ.docx")
        
    except Exception as e:
        print(f"\nâŒ Error: {e}")
        import traceback
        traceback.print_exc()
