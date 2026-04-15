#!/usr/bin/env python3
# Script para analizar la calidad de la Sección 05 vs documento de referencia (sin Unicode)

import os
from docx import Document

def analizar_documento_referencia():
    """Analizar el documento de referencia del 25 de febrero"""
    print("=== ANALISIS DOCUMENTO REFERENCIA (25 febrero) ===")
    
    ref_path = "../Plan_Negocio_Treqe_100%_PROFESIONAL_FINAL.docx"
    if os.path.exists(ref_path):
        tamaño = os.path.getsize(ref_path)
        print(f"Tamaño: {tamaño:,} bytes")
        print(f"Fecha: 25 de febrero, 9:29 AM")
        
        # Intentar leer estructura básica
        try:
            doc_ref = Document(ref_path)
            print(f"\nEstructura aproximada del documento de referencia:")
            print(f"- 9 secciones principales")
            print(f"- ~55KB / 9 = ~6.1KB por seccion en promedio")
            print(f"- Contenido: Introduccion, Problema, Solucion, Ventaja Competitiva, Modelo Negocio, etc.")
        except:
            print("(No se pudo leer contenido interno del documento de referencia)")
    else:
        print(f"ERROR: No se encuentra el documento de referencia: {ref_path}")
    
    return tamaño if os.path.exists(ref_path) else 0

def analizar_seccion_05():
    """Analizar la Sección 05 actual"""
    print("\n=== ANALISIS SECCION 05 ACTUAL ===")
    
    sec_path = "05_MODELO_NEGOCIO.docx"
    if os.path.exists(sec_path):
        tamaño = os.path.getsize(sec_path)
        print(f"Tamaño: {tamaño:,} bytes")
        
        # Leer contenido
        doc = Document(sec_path)
        
        # Contar párrafos y subsecciones
        total_parrafos = len([p for p in doc.paragraphs if p.text.strip()])
        subsecciones = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if text and any(c.isdigit() for c in text[:3]) and '.' in text[:5]:
                subsecciones.append(text)
        
        print(f"Total parrafos con contenido: {total_parrafos}")
        print(f"Numero de subsecciones: {len(subsecciones)}")
        
        print("\nSubsecciones identificadas:")
        for i, sec in enumerate(subsecciones[:10]):  # Mostrar primeras 10
            print(f"  {i+1}. {sec[:80]}...")
        
        # Analizar calidad de contenido
        print("\n=== ANALISIS DE CALIDAD ===")
        
        # Verificar presencia de elementos clave
        contenido = " ".join([p.text for p in doc.paragraphs if p.text.strip()])
        
        elementos_clave = {
            "business model canvas": "business model canvas" in contenido.lower(),
            "sociedad limitada": "sociedad limitada" in contenido.lower(),
            "flujos de ingresos": "flujos de ingresos" in contenido.lower() or "monetizacion" in contenido.lower(),
            "proyecciones financieras": "proyecciones" in contenido.lower() or "financieras" in contenido.lower(),
            "ejemplos concretos": "ana" in contenido.lower() or "carlos" in contenido.lower() or "beatriz" in contenido.lower(),
            "datos cuantitativos": any(x in contenido for x in ["2025", "2026", "€", "%", "usuarios"]),
            "lenguaje humano": not any(robotic in contenido.lower() for robotic in [
                "se procedera a", "es necesario", "optimizacion", "sinergia", "paradigma"
            ])
        }
        
        print("Elementos clave encontrados:")
        for elemento, encontrado in elementos_clave.items():
            status = "OK" if encontrado else "FALTA"
            print(f"  [{status}] {elemento}")
        
        # Skills aplicadas
        print("\n=== SKILLS APLICADAS ===")
        skills = {
            "business-model-canvas": "business model canvas" in contenido.lower(),
            "legal": any(term in contenido.lower() for term in ["sociedad limitada", "patente", "contrato", "fiscal"]),
            "humanizer": elementos_clave["lenguaje humano"],
            "marketing-mode": any(term in contenido.lower() for term in ["estrategia", "marketing", "crecimiento"]),
            "frontend-design": any(term in contenido.lower() for term in ["app", "movil", "usuario", "experiencia"]),
            "algorithm-solver": any(term in contenido.lower() for term in ["algoritmo", "matching", "circular"]),
        }
        
        for skill, aplicada in skills.items():
            status = "OK" if aplicada else "FALTA"
            print(f"  [{status}] {skill}")
        
        return tamaño, len(subsecciones), total_parrafos, skills, elementos_clave
    else:
        print(f"ERROR: No se encuentra la seccion 05: {sec_path}")
        return 0, 0, 0, {}, {}

def recomendar_mejoras(tamaño_ref, tamaño_sec, subsecciones, skills, elementos_clave):
    """Recomendar mejoras basadas en el análisis"""
    print("\n=== RECOMENDACIONES DE MEJORA ===")
    
    # Comparar tamaño
    tamaño_ideal_por_seccion = tamaño_ref / 9  # 9 secciones en documento completo
    print(f"Tamaño ideal por seccion: {tamaño_ideal_por_seccion:,.0f} bytes")
    print(f"Tamaño actual seccion 05: {tamaño_sec:,.0f} bytes")
    
    if tamaño_sec < tamaño_ideal_por_seccion * 0.7:
        print("  [PROBLEMA] DEMASIADO CORTA: Anadir mas detalle y ejemplos")
    elif tamaño_sec > tamaño_ideal_por_seccion * 1.5:
        print("  [PROBLEMA] DEMASIADO LARGA: Considerar dividir en subsecciones mas manejables")
    else:
        print("  [OK] TAMAÑO ADECUADO: Bien balanceada")
    
    # Verificar skills
    skills_requeridas = ["business-model-canvas", "legal", "humanizer"]
    skills_faltantes = [skill for skill in skills_requeridas if not skills.get(skill, False)]
    
    if skills_faltantes:
        print(f"  [PROBLEMA] SKILLS FALTANTES: {', '.join(skills_faltantes)}")
    else:
        print("  [OK] TODAS LAS SKILLS REQUERIDAS APLICADAS")
    
    # Verificar elementos clave
    elementos_faltantes = [elem for elem, encontrado in elementos_clave.items() if not encontrado]
    if elementos_faltantes:
        print(f"  [PROBLEMA] ELEMENTOS FALTANTES: {', '.join(elementos_faltantes[:3])}")
    else:
        print("  [OK] TODOS LOS ELEMENTOS CLAVE PRESENTES")
    
    # Verificar subsecciones
    if subsecciones < 4:
        print(f"  [PROBLEMA] POCAS SUBSECCIONES: Solo {subsecciones}. Ideal: 5-7 subsecciones por seccion")
    else:
        print(f"  [OK] SUBSECCIONES ADECUADAS: {subsecciones} subsecciones")
    
    # Recomendaciones específicas
    print("\n=== ACCIONES RECOMENDADAS ===")
    
    if not elementos_clave.get("ejemplos concretos", False):
        print("1. ANADIR EJEMPLOS CONCRETOS:")
        print("   - Usar Ana, Carlos, Beatriz (como en documento referencia)")
        print("   - Ejemplo: 'Ana intercambia su iPhone 12 (400€) por bicicleta de Carlos (450€) + 50€ compensacion'")
        print("   - Mostrar calculo comision: 3% de 450€ = 13,50€ para Treqe")
    
    if not skills.get("marketing-mode", False):
        print("\n2. APLICAR SKILL MARKETING-MODE:")
        print("   - Anadir estrategia de crecimiento 5 fases")
        print("   - Metricas clave: CAC, LTV, retencion")
        print("   - Canales de adquisicion especificos")
    
    if not skills.get("algorithm-solver", False):
        print("\n3. APLICAR SKILL ALGORITHM-SOLVER:")
        print("   - Explicar complejidad NP-Completo del matching")
        print("   - k=4 maximo por restriccion 5 minutos")
        print("   - Ventaja tecnologica patentable")
    
    print("\n4. VERIFICAR FORMATO DOCUMENTO REFERENCIA:")
    print("   - Revisar estructura exacta del documento original")
    print("   - Mantener mismo nivel de detalle cuantitativo")
    print("   - Usar mismos ejemplos y casos de estudio")

if __name__ == "__main__":
    print("ANALISIS DE CALIDAD: SECCION 05 vs DOCUMENTO REFERENCIA")
    print("=" * 60)
    
    # Cambiar al directorio correcto
    os.chdir(os.path.dirname(os.path.abspath(__file__)))
    
    # Analizar
    tamaño_ref = analizar_documento_referencia()
    tamaño_sec, subsecciones, total_parrafos, skills, elementos_clave = analizar_seccion_05()
    
    # Recomendar mejoras
    if tamaño_ref > 0 and tamaño_sec > 0:
        recomendar_mejoras(tamaño_ref, tamaño_sec, subsecciones, skills, elementos_clave)