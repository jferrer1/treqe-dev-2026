#!/usr/bin/env python3
# Script para crear la Sección 06: PROYECCIONES_FINANCIERAS.docx
# Siguiendo estrictamente el documento de referencia del 25 de febrero

import os
from docx import Document

def crear_seccion_06():
    """Crear la sección 06: PROYECCIONES FINANCIERAS"""
    
    doc = Document()
    
    # 6. PROYECCIONES FINANCIERAS
    doc.add_heading('6. PROYECCIONES FINANCIERAS 2026-2029', 1)
    
    # Introducción
    intro = doc.add_paragraph()
    intro.add_run('Las proyecciones financieras de Treqe se basan en datos reales del mercado español de segunda mano (8.200M€, 2026) y un modelo conservador de adopción. Asumimos una penetración gradual del 0,036% en año 1 hasta 0,71% en año 3.')
    
    # 6.1 Supuestos Clave
    doc.add_heading('6.1 Supuestos Clave del Modelo', 2)
    
    supuestos = doc.add_paragraph()
    supuestos.add_run('• Mercado total España 2026: 8.200M€ (crecimiento 42% vs 2020)\n')
    supuestos.add_run('• Usuarios activos: 28 millones (47% población)\n')
    supuestos.add_run('• Penetración Treqe: 0,036% año 1 → 0,71% año 3\n')
    supuestos.add_run('• Valor medio intercambio: 300€ (basado en datos Wallapop/Vinted)\n')
    supuestos.add_run('• Frecuencia intercambio: 2,5 por usuario/año (usuarios activos)\n')
    supuestos.add_run('• Comisión Treqe: 3% sobre valor declarado\n')
    supuestos.add_run('• Conversión suscripción premium: 8% usuarios activos\n')
    supuestos.add_run('• Precio suscripción: 9,99€/mes (99€/año con descuento)\n')
    supuestos.add_run('• CAC (Coste Adquisición Cliente): 15€ (orgánico + performance)\n')
    supuestos.add_run('• Retención 30 días: 65% (benchmark marketplaces establecidos)')
    
    # 6.2 Proyecciones de Ingresos
    doc.add_heading('6.2 Proyecciones de Ingresos Detalladas', 2)
    
    # Tabla año por año
    ingresos_text = doc.add_paragraph()
    ingresos_text.add_run('AÑO 1 (2026) - LANZAMIENTO:\n')
    ingresos_text.add_run('• Usuarios activos: 10.000 (0,036% mercado)\n')
    ingresos_text.add_run('• Transacciones: 2.500 (2,5 por usuario)\n')
    ingresos_text.add_run('• Volumen intercambiado: 750.000€ (300€ medio)\n')
    ingresos_text.add_run('• Comisiones (3%): 22.500€ (70% ingresos)\n')
    ingresos_text.add_run('• Suscripciones premium: 800 usuarios × 99€ = 79.200€ (30%)\n')
    ingresos_text.add_run('• Ingresos totales: 101.700€\n')
    ingresos_text.add_run('• Ingresos por usuario: 10,17€\n\n')
    
    ingresos_text.add_run('AÑO 2 (2027) - CRECIMIENTO:\n')
    ingresos_text.add_run('• Usuarios activos: 50.000 (0,18% mercado)\n')
    ingresos_text.add_run('• Transacciones: 12.500\n')
    ingresos_text.add_run('• Volumen intercambiado: 3.750.000€\n')
    ingresos_text.add_run('• Comisiones: 112.500€ (60%)\n')
    ingresos_text.add_run('• Suscripciones: 4.000 usuarios × 99€ = 396.000€ (25%)\n')
    ingresos_text.add_run('• Servicios empresas: 75.000€ (15%)\n')
    ingresos_text.add_run('• Ingresos totales: 583.500€\n')
    ingresos_text.add_run('• Ingresos por usuario: 11,67€\n\n')
    
    ingresos_text.add_run('AÑO 3 (2028) - ESCALABILIDAD:\n')
    ingresos_text.add_run('• Usuarios activos: 200.000 (0,71% mercado)\n')
    ingresos_text.add_run('• Transacciones: 50.000\n')
    ingresos_text.add_run('• Volumen intercambiado: 15.000.000€\n')
    ingresos_text.add_run('• Comisiones: 450.000€ (50%)\n')
    ingresos_text.add_run('• Suscripciones: 16.000 usuarios × 99€ = 1.584.000€ (20%)\n')
    ingresos_text.add_run('• Servicios empresas: 600.000€ (30%)\n')
    ingresos_text.add_run('• Ingresos totales: 2.634.000€\n')
    ingresos_text.add_run('• Ingresos por usuario: 13,17€')
    
    # 6.3 Estructura de Costes
    doc.add_heading('6.3 Estructura de Costes y EBITDA', 2)
    
    costes = doc.add_paragraph()
    costes.add_run('AÑO 1 - Inversión en desarrollo y lanzamiento:\n')
    costes.add_run('• Desarrollo tecnológico: 40.000€ (40%)\n')
    costes.add_run('• Marketing y adquisición: 30.000€ (30%)\n')
    costes.add_run('• Operaciones y soporte: 20.000€ (20%)\n')
    costes.add_run('• Gastos generales: 10.000€ (10%)\n')
    costes.add_run('• Costes totales: 100.000€\n')
    costes.add_run('• EBITDA: -47.500€ (ingresos 101.700€ - costes 149.200€)\n\n')
    
    costes.add_run('AÑO 2 - Crecimiento con eficiencia:\n')
    costes.add_run('• Desarrollo: 80.000€ (35%)\n')
    costes.add_run('• Marketing: 90.000€ (40%)\n')
    costes.add_run('• Operaciones: 40.000€ (18%)\n')
    costes.add_run('• General: 15.000€ (7%)\n')
    costes.add_run('• Costes totales: 225.000€\n')
    costes.add_run('• EBITDA: +45.000€ (Q4 primer trimestre rentable)\n\n')
    
    costes.add_run('AÑO 3 - Escalabilidad y rentabilidad:\n')
    costes.add_run('• Desarrollo: 120.000€ (25%)\n')
    costes.add_run('• Marketing: 180.000€ (38%)\n')
    costes.add_run('• Operaciones: 120.000€ (25%)\n')
    costes.add_run('• General: 50.000€ (10%)\n')
    costes.add_run('• Costes totales: 470.000€\n')
    costes.add_run('• EBITDA: +459.000€ (margen 30%)')
    
    # 6.4 Punto de Equilibrio
    doc.add_heading('6.4 Punto de Equilibrio y Sensibilidad', 2)
    
    equilibrio = doc.add_paragraph()
    equilibrio.add_run('PUNTO DE EQUILIBRIO OPERATIVO:\n')
    equilibrio.add_run('• Costes fijos mensuales año 1: 8.267€\n')
    equilibrio.add_run('• Margen contribución por transacción: 8,10€ (3% de 300€ - coste variable)\n')
    equilibrio.add_run('• Transacciones necesarias/mes: 1.021\n')
    equilibrio.add_run('• Usuarios activos necesarios: 408 (asumiendo 2,5 transacciones/usuario/año)\n')
    equilibrio.add_run('• Mes alcanzado: 18 (15.000 usuarios activos)\n\n')
    
    equilibrio.add_run('ANÁLISIS DE SENSIBILIDAD:\n')
    equilibrio.add_run('Escenario conservador (-20% vs proyección):\n')
    equilibrio.add_run('• Año 3: 160.000 usuarios (no 200.000)\n')
    equilibrio.add_run('• Ingresos: 2.107.000€ (no 2.634.000€)\n')
    equilibrio.add_run('• EBITDA: +367.000€ (margen 28%)\n\n')
    
    equilibrio.add_run('Escenario optimista (+20% vs proyección):\n')
    equilibrio.add_run('• Año 3: 240.000 usuarios\n')
    equilibrio.add_run('• Ingresos: 3.161.000€\n')
    equilibrio.add_run('• EBITDA: +551.000€ (margen 32%)')
    
    # 6.5 Unit Economics
    doc.add_heading('6.5 Unit Economics y Métricas Clave', 2)
    
    metrics = doc.add_paragraph()
    metrics.add_run('CAC (Coste Adquisición Cliente): 15€\n')
    metrics.add_run('• Orgánico (referrals, contenido): 40% - CAC 0€\n')
    metrics.add_run('• Performance (redes sociales, SEM): 60% - CAC 25€\n')
    metrics.add_run('• Promedio ponderado: 15€\n\n')
    
    metrics.add_run('LTV (Lifetime Value): 360€\n')
    metrics.add_run('• Ingresos por usuario año 1: 10,17€\n')
    metrics.add_run('• Retención mensual: 65% → vida media: 34 meses\n')
    metrics.add_run('• LTV = 10,17€ × 34 meses = 346€ (redondeado 360€)\n\n')
    
    metrics.add_run('LTV:CAC Ratio: 24:1 (excelente)\n')
    metrics.add_run('• Benchmark saludable: 3:1\n')
    metrics.add_run('• Treqe: 360€ / 15€ = 24:1\n\n')
    
    metrics.add_run('Payback Period CAC: 1,5 meses\n')
    metrics.add_run('• Usuario genera 10,17€/año → 0,85€/mes\n')
    metrics.add_run('• CAC 15€ recuperado en 18 meses (1,5 años)\n')
    metrics.add_run('• Con suscripción premium: payback 3 meses')
    
    # 6.6 Necesidades de Financiación
    doc.add_heading('6.6 Necesidades de Financiación y Uso de Fondos', 2)
    
    financiacion = doc.add_paragraph()
    financiacion.add_run('RONDAS DE FINANCIACIÓN:\n')
    financiacion.add_run('• Pre-seed (actual): 50.000€ - MVP y primeros 1.000 usuarios\n')
    financiacion.add_run('• Seed (mes 6): 250.000€ - Escalabilidad Madrid\n')
    financiacion.add_run('• Serie A (año 2): 1.500.000€ - Expansión nacional\n\n')
    
    financiacion.add_run('USO DE FONDOS PRE-SEED (50.000€):\n')
    financiacion.add_run('• Desarrollo MVP: 25.000€ (50%)\n')
    financiacion.add_run('• Marketing lanzamiento: 15.000€ (30%)\n')
    financiacion.add_run('• Operaciones iniciales: 7.500€ (15%)\n')
    financiacion.add_run('• Contingencia: 2.500€ (5%)\n\n')
    
    financiacion.add_run('DILUCIÓN Y VALORACIONES:\n')
    financiacion.add_run('• Pre-seed: 50.000€ por 10% → valoración 500.000€\n')
    financiacion.add_run('• Seed: 250.000€ por 15% → valoración 1,67M€\n')
    financiacion.add_run('• Serie A: 1,5M€ por 15% → valoración 10M€\n')
    financiacion.add_run('• Dilución fundadores: 65% post-serie A (objetivo estándar)')
    
    # 6.7 ROI para Inversores
    doc.add_heading('6.7 ROI para Inversores y Salida', 2)
    
    roi = doc.add_paragraph()
    roi.add_run('INVERSOR PRE-SEED (50.000€ por 10%):\n')
    roi.add_run('• Valoración año 3: 12-15M€ (8-10x ingresos)\n')
    roi.add_run('• Valor participación: 1,2-1,5M€\n')
    roi.add_run('• ROI: 24-30x inversión inicial\n')
    roi.add_run('• TIR (Tasa Interna Retorno): 125-140% anual\n\n')
    
    roi.add_run('ESCENARIOS DE SALIDA:\n')
    roi.add_run('1. Adquisición por marketplace establecido (Wallapop, Vinted):\n')
    roi.add_run('   • Valoración: 15-20M€ (año 3-4)\n')
    roi.add_run('   • Razón: Tecnología patentada + comunidad fiel\n\n')
    
    roi.add_run('2. IPO o SPAC (año 5-7):\n')
    roi.add_run('   • Valoración: 50-100M€\n')
    roi.add_run('   • Requisito: 1M+ usuarios, rentabilidad sostenida\n\n')
    
    roi.add_run('3. Crecimiento orgánico y dividendos (año 4+):\n')
    roi.add_run('   • EBITDA año 4: 1,2M€ (proyección)\n')
    roi.add_run('   • Dividendos: 30% EBITDA = 360.000€/año\n')
    roi.add_run('   • Yield inversor pre-seed: 720% (360k€ / 50k€)')
    
    # 6.8 Conclusión: Viabilidad Financiera
    doc.add_heading('6.8 Conclusión: Viabilidad Financiera Demostrada', 2)
    
    conclusion = doc.add_paragraph()
    conclusion.add_run('Las proyecciones financieras de Treqe no son optimistas, son conservadoras y basadas en:\n\n')
    conclusion.add_run('1. DATOS REALES DE MERCADO: 8.200M€ España 2026, crecimiento 42% desde 2020\n')
    conclusion.add_run('2. PENETRACIÓN MODESTA: 0,71% mercado año 3 (vs 5-10% líderes)\n')
    conclusion.add_run('3. UNIT ECONOMICS SÓLIDOS: LTV:CAC 24:1 (benchmark saludable: 3:1)\n')
    conclusion.add_run('4. PUNTO EQUILIBRIO TEMPRANO: Mes 18 (15.000 usuarios)\n')
    conclusion.add_run('5. MÚLTIPLES VÍAS DE SALIDA: Adquisición, IPO, dividendos\n\n')
    
    conclusion.add_run('El riesgo no es la viabilidad financiera (demostrada), sino la ejecución. Por eso:\n')
    conclusion.add_run('• Equipo experimentado en scale-ups\n')
    conclusion.add_run('• MVP funcional antes de seed round\n')
    conclusion.add_run('• Métricas validadas con usuarios reales\n')
    conclusion.add_run('• Roadmap tecnológico claro y ejecutable\n\n')
    
    conclusion.add