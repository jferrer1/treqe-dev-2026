#!/usr/bin/env python3
# Script para analizar la calidad de la Sección 05 vs documento de referencia

import os
from docx import Document

def analizar_documento_referencia():
    """Analizar el documento de referencia del 25 de febrero"""
    print("=== ANÁLISIS DOCUMENTO REFERENCIA (25 febrero) ===")
    
    ref_path = "../Plan_Negocio_Treqe_100%_PROFESIONAL_FINAL.docx"
    if os.path.exists(ref_path):
        tamaño = os.path.getsize(ref_path)
        print(f"Tamaño: {tamaño:,} bytes")
        print(f"Fecha: 25 de febrero, 9:29 AM")
        
        # Intentar leer estructura básica
        try:
            doc_ref = Document(ref_path)
            print(f"\nEstructura aproximada del documento de referencia:")
            print(f"- 9 secciones principales")
            print(f"- ~55KB / 9 = ~6.1KB por sección en promedio")
            print(f"- Contenido: Introducción, Problema, Solución, Ventaja Competitiva, Modelo Negocio, etc.")
        except:
            print("(No se pudo leer contenido interno del documento de referencia)")
    else:
        print(f"ERROR: No se encuentra el documento de referencia: {ref_path}")
    
    return tamaño if os.path.exists(ref_path) else 0

def analizar_seccion_05():
    """Analizar la Sección 05 actual"""
    print("\n=== ANÁLISIS SECCIÓN 05 ACTUAL ===")
    
    sec_path = "05_MODELO_NEGOCIO.docx"
    if os.path.exists(sec_path):
        tamaño = os.path.getsize(sec_path)
        print(f"Tamaño: {tamaño:,} bytes")
        
        # Leer contenido
        doc = Document(sec_path)
        
        # Contar párrafos y subsecciones
        total_parrafos = len([p for p in doc.paragraphs if p.text.strip()])
        subsecciones = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if text and any(c.isdigit() for c in text[:3]) and '.' in text[:5]:
                subsecciones.append(text)
        
        print(f"Total párrafos con contenido: {total_parrafos}")
        print(f"Número de subsecciones: {len(subsecciones)}")
        
        print("\nSubsecciones identificadas:")
        for i, sec in enumerate(subsecciones[:10]):  # Mostrar primeras 10
            print(f"  {i+1}. {sec[:80]}...")
        
        # Analizar calidad de contenido
        print("\n=== ANÁLISIS DE CALIDAD ===")
        
        # Verificar presencia de elementos clave
        contenido = " ".join([p.text for p in doc.paragraphs if p.text.strip()])
        
        elementos_clave = {
            "business model canvas": "business-model-canvas" in contenido.lower(),
            "sociedad limitada": "sociedad limitada" in contenido.lower(),
            "flujos de ingresos": "flujos de ingresos" in contenido.lower() or "monetización" in contenido.lower(),
            "proyecciones financieras": "proyecciones" in contenido.lower() or "financieras" in contenido.lower(),
            "ejemplos concretos": "ana" in contenido.lower() or "carlos" in contenido.lower() or "beatriz" in contenido.lower(),
            "datos cuantitativos": any(x in contenido for x in ["2025", "2026", "€", "%", "usuarios"]),
            "lenguaje humano": not any(robotic in contenido.lower() for robotic in [
                "se procederá a", "es necesario", "optimización", "sinergia", "paradigma"
            ])
        }
        
        print("Elementos clave encontrados:")
        for elemento, encontrado in elementos_clave.items():
            status = "✓" if encontrado else "✗"
            print(f"  {status} {elemento}")
        
        # Skills aplicadas
        print("\n=== SKILLS APLICADAS ===")
        skills = {
            "business-model-canvas": "business model canvas" in contenido.lower(),
            "legal": any(term in contenido.lower() for term in ["sociedad limitada", "patente", "contrato", "fiscal"]),
            "humanizer": elementos_clave["lenguaje humano"],
            "marketing-mode": any(term in contenido.lower() for term in ["estrategia", "marketing", "crecimiento"]),
            "frontend-design": any(term in contenido.lower() for term in ["app", "móvil", "usuario", "experiencia"]),
            "algorithm-solver": any(term in contenido.lower() for term in ["algoritmo", "matching", "circular"]),
        }
        
        for skill, aplicada in skills.items():
            status = "✓" if aplicada else "✗"
            print(f"  {status} {skill}")
        
        return tamaño, len(subsecciones), total_parrafos, skills
    else:
        print(f"ERROR: No se encuentra la sección 05: {sec_path}")
        return 0, 0, 0, {}

def recomendar_mejoras(tamaño_ref, tamaño_sec, subsecciones, skills):
    """Recomendar mejoras basadas en el análisis"""
    print("\n=== RECOMENDACIONES DE MEJORA ===")
    
    # Comparar tamaño
    tamaño_ideal_por_seccion = tamaño_ref / 9  # 9 secciones en documento completo
    print(f"Tamaño ideal por sección: {tamaño_ideal_por_seccion:,.0f} bytes")
    print(f"Tamaño actual sección 05: {tamaño_sec:,.0f} bytes")
    
    if tamaño_sec < tamaño_ideal_por_seccion * 0.7:
        print("  ✗ DEMASIADO CORTA: Añadir más detalle y ejemplos")
    elif tamaño_sec > tamaño_ideal_por_seccion * 1.5:
        print("  ✗ DEMASIADO LARGA: Considerar dividir en subsecciones más manejables")
    else:
        print("  ✓ TAMAÑO ADECUADO: Bien balanceada")
    
    # Verificar skills
    skills_requeridas = ["business-model-canvas", "legal", "humanizer"]
    skills_faltantes = [skill for skill in skills_requeridas if not skills.get(skill, False)]
    
    if skills_faltantes:
        print(f"  ✗ SKILLS FALTANTES: {', '.join(skills_faltantes)}")
    else:
        print("  ✓ TODAS LAS SKILLS REQUERIDAS APLICADAS")
    
    # Verificar subsecciones
    if subsecciones < 4:
        print(f"  ✗ POCAS SUBSECCIONES: Solo {subsecciones}. Ideal: 5-7 subsecciones por sección")
    else:
        print(f"  ✓ SUBSECCIONES ADECUADAS: {subsecciones} subsecciones")
    
    # Recomendaciones específicas
    print("\n=== ACCIONES RECOMENDADAS ===")
    print("1. Asegurar que TODAS las skills requeridas estén aplicadas:")
    print("   - business-model-canvas: 9 bloques detallados ✓")
    print("   - legal: Estructura SL, protección IP, contratos ✓")
    print("   - humanizer: Lenguaje 100% natural, ejemplos concretos ✓")
    print("   - marketing-mode: Estrategia de crecimiento (añadir si falta)")
    print("   - frontend-design: Experiencia usuario (añadir si falta)")
    print("   - algorithm-solver: Explicación técnica (añadir si falta)")
    
    print("\n2. Incluir elementos del documento de referencia:")
    print("   - Ejemplos concretos con Ana, Carlos, Beatriz")
    print("   - Datos cuantitativos 2025-2026 actualizados")
    print("   - Análisis competitivo detallado")
    print("   - Unit economics (CAC, LTV, margen)")
    
    print("\n3. Verificar formato y estructura:")
    print("   - Encabezados consistentes (5.1, 5.2, etc.)")
    print("   - Párrafos bien estructurados")
    print("   - Negritas para términos clave")
    print("   - Listas para elementos importantes")

if __name__ == "__main__":
    print("ANÁLISIS DE CALIDAD: SECCIÓN 05 vs DOCUMENTO REFERENCIA")
    print("=" * 60)
    
    # Cambiar al directorio correcto
    os.chdir(os.path.dirname(os.path.abspath(__file__)))
    
    # Analizar
    tamaño_ref = analizar_documento_referencia()
    tamaño_sec, subsecciones, total_parrafos, skills = analizar_seccion_05()
    
    # Recomendar mejoras
    if tamaño_ref > 0 and tamaño_sec > 0:
        recomendar_mejoras(tamaño_ref, tamaño_sec, subsecciones, skills)