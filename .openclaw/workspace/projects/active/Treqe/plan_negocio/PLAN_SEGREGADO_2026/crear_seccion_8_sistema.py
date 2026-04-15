#!/usr/bin/env python3
# -*- coding: utf-8 -*-
# Crear Sección 8: RIESGOS Y MITIGACIÓN con sistema garantizado

from docx import Document

print("=== SISTEMA GARANTIZADO ACTIVADO ===")
print("Creando Sección 8: RIESGOS Y MITIGACIÓN")
print("Skills aplicadas: HUMANIZER + LEGAL + BUSINESS-MODEL-CANVAS + ALGORITHM-SOLVER")

doc = Document()

# ========== 8. RIESGOS Y MITIGACIÓN ==========
titulo = doc.add_heading('8. RIESGOS Y MITIGACIÓN: PREPARADOS PARA CUANDO LAS COSAS SALEN MAL (PORQUE ALGUNA VEZ SALDRÁN MAL)', 0)

intro = doc.add_paragraph()
intro.add_run('Esta sección aplica: HUMANIZER + LEGAL + BUSINESS-MODEL-CANVAS + ALGORITHM-SOLVER')
intro.add_run('\n\nImagina este escenario: Ana acaba de conseguir su bicicleta después de meses intentándolo. ')
intro.add_run('Está feliz. Y de repente, Carlos no envía la bicicleta. O peor: envía una bicicleta rota. ')
intro.add_run('O aún peor: el algoritmo no encuentra a nadie que quiera el iPhone de Ana.\n\n')
intro.add_run('Esta sección no es sobre tener miedo. Es sobre estar preparado. ')
intro.add_run('Porque los buenos negocios no son los que nunca tienen problemas. ')
intro.add_run('Son los que tienen planes para cuando los problemas lleguen.')

# ========== 8.1 RIESGOS TÉCNICOS ==========
doc.add_heading('8.1 Riesgos Técnicos: Cuando el algoritmo falla (y fallará)', 1)

p1 = doc.add_paragraph()
p1.add_run('**Skill aplicada: ALGORITHM-SOLVER**\n\n')
p1.add_run('**RIESGO 1: El algoritmo no encuentra matches (k=2 imposible, k=3 muy lento)**\n')
p1.add_run('• **Probabilidad año 1:** 35% (mercado pequeño, densidad baja)\n')
p1.add_run('• **Impacto:** Ana se frustra, abandona, cuenta mala experiencia\n')
p1.add_run('• **MITIGACIÓN:**\n')
p1.add_run('  - Algoritmo híbrido: k=2 primero, si no hay match → k=3 en 5 segundos\n')
p1.add_run('  - "Match sugerido": Si no hay ciclo perfecto, sugerir intercambio con compensación €\n')
p1.add_run('  - Pool de espera: Ana entra en lista, notificación cuando haya match\n')
p1.add_run('• **COSTE MITIGACIÓN:** 15.000€ desarrollo algoritmo avanzado\n\n')

p1.add_run('**RIESGO 2: Escalabilidad técnica (100.000 Anas colapsan servidores)**\n')
p1.add_run('• **Probabilidad año 2:** 25% (crecimiento rápido)\n')
p1.add_run('• **Impacto:** Tiempos respuesta >10s, abandonos 40%\n')
p1.add_run('• **MITIGACIÓN:**\n')
p1.add_run('  - Arquitectura microservicios desde día 1 (coste +20%)\n')
p1.add_run('  - Auto-scaling AWS: 50% más caro, pero escala automáticamente\n')
p1.add_run('  - Cache distribuido Redis: matches frecuentes precalculados\n')
p1.add_run('• **COSTE MITIGACIÓN:** 8.000€/mes infraestructura premium\n\n')

p1.add_run('**RIESGO 3: Seguridad (hackeo, datos Ana robados)**\n')
p1.add_run('• **Probabilidad:** 5% anual (standard tech)\n')
p1.add_run('• **Impacto:** Multa GDPR hasta 4% facturación, pérdida confianza\n')
p1.add_run('• **MITIGACIÓN:**\n')
p1.add_run('  - Pentesting trimestral: 12.000€/año\n')
p1.add_run('  - Bug bounty program: 50.000€ presupuesto recompensas\n')
p1.add_run('  - Encryption end-to-end: datos Ana nunca en texto plano\n')
p1.add_run('• **COSTE MITIGACIÓN:** 62.000€/año seguridad premium')

# ========== 8.2 RIESGOS OPERATIVOS ==========
doc.add_heading('8.2 Riesgos Operativos: Cuando Beatriz no puede con todo', 1)

p2 = doc.add_paragraph()
p2.add_run('**Skill aplicada: BUSINESS-MODEL-CANVAS**\n\n')
p2.add_run('**RIESGO 4: Fraudes y estafas (Carlos envía bicicleta rota)**\n')
p2.add_run('• **Probabilidad año 1:** 8% (mercado nuevo, controles básicos)\n')
p2.add_run('• **Impacto:** Ana pierde iPhone + confianza, demanda posible\n')
p2.add_run('• **MITIGACIÓN:**\n')
p2.add_run('  - Sistema triple verificación: foto + video + código serial\n')
p2.add_run('  - Fondo garantía: 100.000€ reservado para reembolsos\n')
p2.add_run('  - Seguro fraude: 15.000€/año, cubre hasta 50.000€ pérdidas\n')
p2.add_run('  - Reputación pública: Carlos con 50 transacciones = "Verificado Oro"\n')
p2.add_run('• **COSTE MITIGACIÓN:** 127.000€ año 1 (fondo + seguro + equipo)\n\n')

p2.add_run('**RIESGO 5: Logística falla (transporte pierde paquete)**\n')
p2.add_run('• **Probabilidad:** 2% por transacción (estadística Correos)\n')
p2.add_run('• **Impacto:** Ana y Carlos frustrados, reembolso 100%\n')
p2.add_run('• **MITIGACIÓN:**\n')
p2.add_run('  - Seguro envío obligatorio: 2,50€/transacción, cubre hasta 500€\n')
p2.add_run('  - Acuerdos con 3 transportistas: Correos, SEUR, DHL\n')
p2.add_run('  - Tracking tiempo real: Ana ve dónde está su bicicleta minuto a minuto\n')
p2.add_run('• **COSTE MITIGACIÓN:** 2,50€/transacción (absorbido en comisión 3%)\n\n')

p2.add_run('**RIESGO 6: Soporte colapsa (10.000 Anas preguntando a la vez)**\n')
p2.add_run('• **Probabilidad año 2:** 40% (crecimiento rápido)\n')
p2.add_run('• **Impacto:** Tiempo respuesta >48h, NPS cae a 20\n')
p2.add_run('• **MITIGACIÓN:**\n')
p2.add_run('  - Chatbot IA: responde 70% preguntas frecuentes automáticamente\n')
p2.add_run('  - Equipo escalable: 1 soporte/2.500 Anas (contratación progresiva)\n')
p2.add_run('  - Comunidad auto-gestionada: Anas expertas ayudan a nuevas Anas\n')
p2.add_run('• **COSTE MITIGACIÓN:** 45.000€/año equipo soporte + 20.000€ desarrollo chatbot')

# ========== 8.3 RIESGOS DE MERCADO ==========
doc.add_heading('8.3 Riesgos de Mercado: Cuando Ana prefiere seguir usando Wallapop', 1)

p3 = doc.add_paragraph()
p3.add_run('**Skill aplicada: HUMANIZER + BUSINESS-MODEL-CANVAS**\n\n')
p3.add_run('**RIESGO 7: Adopción lenta (Ana no entiende por qué cambiar)**\n')
p3.add_run('• **Probabilidad año 1:** 60% (nuevo concepto)\n')
p3.add_run('• **Impacto:** CAC sube a 25€, crecimiento 50% más lento\n')
p3.add_run('• **MITIGACIÓN:**\n')
p3.add_run('  - Programa early adopters: primeros 5.000 Anas, comisión 0% año 1\n')
p3.add_run('  - "Prueba con un amigo": Ana invita a Carlos, ambos bonus 15€\n')
p3.add_run('  - Contenido educativo: "Guía Ana: cómo intercambiar tu iPhone por lo que realmente quieres"\n')
p3.add_run('• **COSTE MITIGACIÓN:** 75.000€ marketing educativo + 75.000€ bonificaciones\n\n')

p3.add_run('**RIESGO 8: Competencia reacciona (Wallapop añade trueque)**\n')
p3.add_run('• **Probabilidad año 2:** 30%\n')
p2.add_run('• **Impacto:** Pérdida cuota mercado 20%, CAC sube 40%\n')
p2.add_run('• **MITIGACIÓN:**\n')
p2.add_run('  - Ventaja técnica: nuestro algoritmo 3 años desarrollo vs su feature añadida\n')
p2.add_run('  - Foco comunidad: Wallapop es marketplace, Treqe es comunidad de intercambio\n')
p2.add_run('  - Especialización: solo segunda mano calidad (no coches, no casas, no servicios)\n')
p2.add_run('• **COSTE MITIGACIÓN:** 50.000€ branding "especialistas en trueque inteligente"\n\n')

p3.add_run('**RIESGO 9: Cambio regulación (leyes trueque más estrictas)**\n')
p3.add_run('• **Probabilidad:** 15% anual (UE regulando plataformas)\n')
p3.add_run('• **Impacto:** Costes compliance +200.000€/año, restricciones operativas\n')
p3.add_run('• **MITIGACIÓN:**\n')
p3.add_run('  - Asesor legal desde día 1: 24.000€/año\n')
p3.add_run('  - Lobby pro-trueque: unirse a asociaciones economía circular\n')
p3.add_run('  - Diseño compliant by design: KYC básico desde inicio\n')
p3.add_run('• **COSTE MITIGACIÓN:** 74.000€/año (legal + compliance + asociaciones)')

# ========== 8.4 RIESGOS FINANCIEROS ==========
doc.add_heading('8.4 Riesgos Financieros: Cuando el dinero se acaba antes de tiempo', 1)

p4 = doc.add_paragraph()
p4.add_run('**Skill aplicada: BUSINESS-MODEL-CANVAS + LEGAL**\n\n')
p4.add_run('**RIESGO 10: Runway insuficiente (150.000€ duran 14 meses, no 18)**\n')
p4.add_run('• **Probabilidad:** 25% (sobreestimación ingresos, subestimación costes)\n')
p4.add_run('• **Impacto:** Cierre mes 14, equipo despedido, Ana sin plataforma\n')
p4.add_run('• **MITIGACIÓN:**\n')
p4.add_run('  - Buffer 25%: pedir 187.500€, no 150.000€\n')
p4.add_run('  - Hitos financieros claros: si CAC >20€ mes 6, pivot inmediato\n')
p4.add_run('  - Plan profitability acelerado: reducir costes 30% si crecimiento <70% esperado\n')
p4.add_run('• **COSTE MITIGACIÓN:** Pedir 37.500€ extra (dilución 2,5% adicional)\n\n')

p4.add_run('**RIESGO 11: Inversores no aparecen (seed round falla)**\n')
p4.add_run('• **Probabilidad:** 20% (mercado VC volátil)\n')
p4.add_run('• **Impacto:** Sin dinero mes 7, crecimiento se estanca\n')
p4.add_run('• **MITIGACIÓN:**\n')
p4.add_run('  - Plan bootstrap: si no hay seed, reducir equipo 50%, focus profitability\n')
p4.add_run('  - Revenue-based financing: préstamos contra ingresos futuros\n')
p4.add_run('  - Strategic angels: 5-10 inversores pequeños vs 1 VC grande\n')
p4.add_run('• **COSTE MITIGACIÓN:** Crecimiento 60% más lento, pero supervivencia asegurada\n\n')

p4.add_run('**RIESGO 12: Unit economics negativos (CAC > LTV)**\n')
p4.add_run('• **Probabilidad año 1:** 15% (mercado prueba error)\n')
p4.add_run('• **Impacto:** Cada Ana nueva nos cuesta dinero, negocio insostenible\n')
p4.add_run('• **MITIGACIÓN:**\n')
p4.add_run('  - Test CAC/LTV mes 1-3 con 1.000€, no 50.000€\n')
p4.add_run('  - Pivot rápido: si CAC 25€, cambiar a growth orgánico inmediato\n')
p4.add_run('  - Monetización temprana: suscripciones desde día 1, no esperar escala\n')
p4.add_run('• **COSTE MITIGACIÓN:** 5.000€ pruebas pequeñas antes de escalar')

# ========== 8.5 MATRIZ DE RIESGOS ==========
doc.add_heading('8.5 Matriz de Riesgos: Lo que realmente importa y lo que realmente haremos', 1)

p5 = doc.add_paragraph()
p5.add_run('**ALTA PROBABILIDAD + ALTO IMPACTO (GESTIÓN INMEDIATA):**\n')
p5.add_run('1. **Adopción lenta (60% prob, impacto alto)** - Mitigación: 150.000€ educativo\n')
p5.add_run('2. **Soporte colapsa (40% prob, impacto alto)** - Mitigación: 65.000€ IA + equipo\n')
p5.add_run('3. **Algoritmo no matches (35% prob, impacto medio)** - Mitigación: 15.000€ desarrollo\n\n')

p5.add_run('**BAJA PROBABILIDAD + ALTO IMPACTO (SEGUROS + FONDOS):**\n')
p5.add_run('1. **Fraudes (8% prob, impacto muy alto)** - Mitigación: 127.000€ fondo + seguro\n')
p5.add_run('2. **Hackeo (5% prob, impacto muy alto)** - Mitigación: 62.000€ seguridad\n')
p5.add_run('3. **Regulación (15% prob, impacto alto)** - Mitigación: 74.000€ legal\n\n')

p5.add_run('**ALTA PROBABILIDAD + BAJO IMPACTO (MONITOREO):**\n')
p5.add_run('1. **Logística falla (2% prob, impacto bajo)** - Mitigación: 2,50€/transacción\n')
p5.add_run('2. **Escalabilidad técnica (25% prob, impacto medio)** - Mitigación: 8.000€/mes infra\n\n')

p5.add_run('**PRINCIPIO CLAVE:** No gastamos 200.000€ en prevenir riesgos de 2% probabilidad. ')
p5.add_run('Gastamos 150.000€ en prevenir riesgos de 60% probabilidad. ')
p5.add_run('Y tenemos seguros para riesgos de 5% probabilidad pero impacto catastrófico.')

# ========== 8.6 PLAN DE CONTINUIDAD ==========
doc.add_heading('8.6 Plan de Continuidad: Si todo falla, esto es lo que haremos', 1)

p6 = doc.add_paragraph()
p6.add_run('**ESCENARIO CATASTRÓFICO 1: CAC 30€, LTV 100€ (unit economics rotos)**\n')
p6.add_run('• **Acción día 1:** Congelar marketing performance (ahorro 20.000€/mes)\n')
p6.add_run('• **Acción día 7:** Pivot a modelo B2B (empresas pagan por programa recompra)\n')
p6.add_run('• **Acción día 30:** Reducir equipo 70%, focus profitability con 5.000 usuarios\n')
p6.add_run('• **Resultado:** Empresa más pequeña (10 empleados, 500.000€ ingresos) pero rentable\n\n')

p6.add_run('**ESCENARIO CATASTRÓFICO 2: Hackeo masivo, datos 50.000 Anas robados**\n')
p6.add_run('• **Acción hora 1:** Notificar autoridades, activar seguro cibernético\n')
p6.add_run('• **Acción hora 6:** Comunicación transparente a todas las Anas\n')
p6.add_run('• **Acción día 7:** Auditoría seguridad externa,