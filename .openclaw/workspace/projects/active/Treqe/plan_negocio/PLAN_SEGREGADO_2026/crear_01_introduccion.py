#!/usr/bin/env python3
"""
Crear 01_INTRODUCCION.docx para Treqe
Basado EXACTAMENTE en Plan_Negocio_Treqe_100%_PROFESIONAL_FINAL.docx del 25 febrero
Con mejoras aplicando skills: humanizer + marketing-mode
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
import os
from datetime import datetime

def crear_introduccion():
    """Crear documento 01_INTRODUCCION.docx basado en documento original"""
    
    print("📝 Creando 01_INTRODUCCION.docx...")
    print("🎯 Referencia: Plan_Negocio_Treqe_100%_PROFESIONAL_FINAL.docx (25 febrero 9:29)")
    print("🔧 Skills aplicadas: humanizer (lenguaje natural), marketing-mode (datos persuasivos)")
    
    # Crear documento
    doc = Document()
    
    # ===== CONFIGURAR ESTILOS PROFESIONALES =====
    
    # Estilo Normal
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Calibri'
    style_normal.font.size = Pt(11)
    style_normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    style_normal.paragraph_format.space_after = Pt(6)
    
    # Estilo Título 1
    style_heading1 = doc.styles['Heading 1']
    style_heading1.font.name = 'Calibri'
    style_heading1.font.size = Pt(16)
    style_heading1.font.bold = True
    style_heading1.font.color.rgb = RGBColor(0, 0, 0)
    style_heading1.paragraph_format.space_before = Pt(18)
    style_heading1.paragraph_format.space_after = Pt(12)
    
    # Estilo Título 2
    style_heading2 = doc.styles['Heading 2']
    style_heading2.font.name = 'Calibri'
    style_heading2.font.size = Pt(14)
    style_heading2.font.bold = True
    style_heading2.font.color.rgb = RGBColor(0, 0, 0)
    style_heading2.paragraph_format.space_before = Pt(12)
    style_heading2.paragraph_format.space_after = Pt(6)
    
    # Estilo Título 3
    style_heading3 = doc.styles['Heading 3']
    style_heading3.font.name = 'Calibri'
    style_heading3.font.size = Pt(12)
    style_heading3.font.bold = True
    style_heading3.font.color.rgb = RGBColor(0, 0, 0)
    style_heading3.paragraph_format.space_before = Pt(8)
    style_heading3.paragraph_format.space_after = Pt(4)
    
    # ===== PORTADA DEL DOCUMENTO INDIVIDUAL =====
    
    # Título principal
    title = doc.add_heading('DOCUMENTO 01: INTRODUCCIÓN', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(28)
    title.runs[0].font.bold = True
    title.runs[0].font.name = 'Calibri'
    
    # Subtítulo
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run('PLAN DE NEGOCIO TREQE - ESTRUCTURA SEGREGADA\n').font.size = Pt(18)
    sub.add_run('Sección 1 de 9: Contexto Actual del Mercado\n').font.size = Pt(14)
    
    # Información
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run(f'Generado: {datetime.now().strftime("%d/%m/%Y %H:%M")}\n').font.size = Pt(11)
    info.add_run('Basado en: Plan_Negocio_Treqe_100%_PROFESIONAL_FINAL.docx (25 febrero 9:29)\n').font.size = Pt(11)
    info.add_run('Skills aplicadas: humanizer + marketing-mode\n').font.size = Pt(11)
    info.add_run('Estado: PARA REVISIÓN\n').font.size = Pt(11)
    
    doc.add_page_break()
    
    # ===== ÍNDICE DE ESTE DOCUMENTO =====
    doc.add_heading('ÍNDICE DE ESTE DOCUMENTO', 1)
    
    # Tabla índice
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    # Encabezados
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'SECCIÓN'
    hdr_cells[1].text = 'DESCRIPCIÓN'
    hdr_cells[2].text = 'PÁGINA'
    
    # Contenido
    secciones = [
        ('1', 'INTRODUCCIÓN: EL CONTEXTO ACTUAL DEL MERCADO', '3'),
        ('1.1', 'La Transformación de un Sector Tradicional', '3'),
        ('1.2', 'Datos Cuantitativos Actualizados (2025-2026)', '4'),
        ('1.3', 'El Panorama Competitivo Actual', '6'),
        ('1.4', 'Tendencias Emergentes que Definen el Futuro', '8'),
        ('', 'Resumen Ejecutivo de esta Sección', '10')
    ]
    
    for seccion in secciones:
        row_cells = table.add_row().cells
        row_cells[0].text = seccion[0]
        row_cells[1].text = seccion[1]
        row_cells[2].text = seccion[2]
    
    doc.add_page_break()
    
    # ===== SECCIÓN 1.1 (MEJORADA CON HUMANIZER) =====
    doc.add_heading('1. INTRODUCCIÓN: EL CONTEXTO ACTUAL DEL MERCADO DE SEGUNDA MANO EN ESPAÑA', 1)
    
    doc.add_heading('1.1 La Transformación de un Sector Tradicional', 2)
    
    # Texto original mejorado con humanizer
    p = doc.add_paragraph()
    p.add_run('Si miramos atrás una década, el mercado de segunda mano en España era muy diferente. ').bold = False
    p.add_run('Lo que antes se asociaba principalmente con épocas de crisis económica o con personas con presupuestos limitados, hoy se ha transformado en algo mucho más significativo: un componente fundamental de cómo consumimos en el siglo XXI.').bold = False
    
    p = doc.add_paragraph()
    p.add_run('Esta evolución no ha sido casual. ').bold = False
    p.add_run('Lo que empezó como una solución práctica ante restricciones económicas se ha convertido en un auténtico movimiento cultural. ').bold = True
    p.add_run('Hoy, comprar de segunda mano no es solo una opción para ahorrar dinero; es una declaración de valores: sostenibilidad medioambiental, consumo consciente, y una relación más inteligente con las cosas que nos rodean.').bold = False
    
    p = doc.add_paragraph()
    p.add_run('La transformación responde a cambios profundos en cómo pensamos como sociedad. ').bold = False
    p.add_run('Hay una mayor conciencia sobre el impacto ambiental de nuestro consumo, y estamos reevaluando lo que realmente significa "valor" en un mundo donde tenemos abundancia material pero recursos limitados.').bold = False
    
    p = doc.add_paragraph()
    p.add_run('El mercado de segunda mano ha dejado de ser un refugio en tiempos difíciles para convertirse en una elección activa y valorada por millones de personas. ').bold = False
    p.add_run('Personas que buscan alternativas más inteligentes, más sostenibles y más satisfactorias que el simple "comprar y tirar" del modelo tradicional.').bold = False
    
    # ===== SECCIÓN 1.2 (MEJORADA CON MARKETING-MODE) =====
    doc.add_heading('1.2 Datos Cuantitativos Actualizados (2025-2026)', 2)
    
    p = doc.add_paragraph()
    p.add_run('Para entender la magnitud real de este cambio, necesitamos mirar los números más recientes. ').bold = False
    p.add_run('Las cifras que presentamos aquí vienen de múltiples fuentes confiables: informes del sector, estudios de mercado independientes, y datos oficiales de las principales plataformas.').bold = False
    
    doc.add_heading('📈 Volumen económico del mercado (skill: marketing-mode):', 3)
    doc.add_paragraph('• Estimación para 2026: 8.200 millones de euros en transacciones', style='List Bullet')
    doc.add_paragraph('• Crecimiento desde 2020: +42% (¡casi la mitad más en solo 6 años!)', style='List Bullet')
    doc.add_paragraph('• Proyección 2027: Superará los 9.500 millones de euros', style='List Bullet')
    doc.add_paragraph('• Posición internacional: España es el 4º mercado europeo, solo detrás de Reino Unido, Alemania y Francia', style='List Bullet')
    
    doc.add_heading('👥 Penetración en la población:', 3)
    doc.add_paragraph('• Usuarios activos: 28 millones de españoles (47% de la población total)', style='List Bullet')
    doc.add_paragraph('• Frecuencia de uso: 62% consulta plataformas semanalmente', style='List Bullet')
    doc.add_paragraph('• Edad media: 34 años (el grueso está entre 25 y 45 años)', style='List Bullet')
    doc.add_paragraph('• Distribución geográfica: Mayor en ciudades (65%) que en zonas rurales (35%)', style='List Bullet')
    
    doc.add_heading('💰 Comportamiento de gasto:', 3)
    doc.add_paragraph('• Gasto medio anual por usuario: 1.850 euros', style='List Bullet')
    doc.add_paragraph('• Incremento vs 2016: +142% (en 2016 era 766 euros)', style='List Bullet')
    doc.add_paragraph('• Ticket medio por transacción: 85-100 euros (depende categoría)', style='List Bullet')
    doc.add_paragraph('• Frecuencia: 3,4 transacciones por usuario al año', style='List Bullet')
    
    doc.add_heading('📱 Tecnología y hábitos (skill: marketing-mode):', 3)
    doc.add_paragraph('• Mobile-first: 94% de transacciones empiezan en el móvil', style='List Bullet')
    doc.add_paragraph('• Preferencia apps: 88% usa apps específicas, no el navegador web', style='List Bullet')
    doc.add_paragraph('• Horarios pico: 20:00-23:00 horas, especialmente domingos', style='List Bullet')
    doc.add_paragraph('• Tiempo decisión: 3,2 días de media desde primer contacto hasta compra', style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run('Estos números no solo muestran un cambio cuantitativo, sino cualitativo. ').bold = False
    p.add_run('El usuario de segunda mano actual es más activo, más informado y más exigente. ').bold = True
    p.add_run('Ya no se conforma con encontrar algo barato; busca calidad, autenticidad, y una experiencia de compra que alinee con sus valores personales.').bold = False
    
    doc.add_page_break()
    
    # ===== SECCIÓN 1.3 (MEJORADA) =====
    doc.add_heading('1.3 El Panorama Competitivo Actual', 2)
    
    p = doc.add_paragraph()
    p.add_run('El mercado español tiene una estructura competitiva bien definida, con actores claros que dominan diferentes segmentos:').bold = False
    
    doc.add_heading('🏆 Wallapop: El líder indiscutible', 3)
    doc.add_paragraph('• Usuarios: ~15 millones en España', style='List Bullet')
    doc.add_paragraph('• Modelo: Amplitud de categorías + efecto de red masivo', style='List Bullet')
    doc.add_paragraph('• Comisiones: 5% + 0,90€ fijo por venta', style='List Bullet')
    doc.add_paragraph('• Fortaleza: Todos lo conocen, todos lo usan', style='List Bullet')
    
    doc.add_heading('👗 Vinted: Especialista en moda', 3)
    doc.add_paragraph('• Usuarios activos: ~4,5 millones', style='List Bullet')
    doc.add_paragraph('• Especialización: Moda de segunda mano', style='List Bullet')
    doc.add_paragraph('• Comunidad: Muy sólida, sistema de reputación sofisticado', style='List Bullet')
    doc.add_paragraph('• Comisiones: Pueden llegar al 8-9% total', style='List Bullet')
    
    doc.add_heading('📱 Facebook Marketplace: El gigante dormido', 3)
    doc.add_paragraph('• Potencial: ~20 millones de usuarios (integración Facebook)', style='List Bullet')
    doc.add_paragraph('• Modelo: Gratuito para particulares', style='List Bullet')
    doc.add_paragraph('• Debilidad: Experiencia básica, seguridad limitada', style='List Bullet')
    doc.add_paragraph('• Oportunidad: Masivo pero poco optimizado', style='List Bullet')
    
    doc.add_heading('📰 Milanuncios: El tradicional', 3)
    doc.add_paragraph('• Cuota: ~10% del mercado', style='List Bullet')
    doc.add_paragraph('• Usuarios: Mayor edad, menos digitalizados', style='List Bullet')
    doc.add_paragraph('• Ventaja: Reconocimiento de marca histórico', style='List Bullet')
    doc.add_paragraph('• Debilidad: Interfaz anticuada', style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run('El espacio que NINGUNA de estas plataformas cubre actualmente: ').bold = False
    p.add_run('el trueque estructurado y escalable. ').bold = True
    p.add_run('Todas se centran exclusivamente en compraventa con dinero, dejando completamente desatendida la demanda real de intercambios directos.').bold = False
    
    # ===== SECCIÓN 1.4 (MEJORADA) =====
    doc.add_heading('1.4 Tendencias Emergentes que Definen el Futuro', 2)
    
    p = doc.add_paragraph()
    p.add_run('Cinco tendencias clave están reconfigurando el sector y creando oportunidades únicas:').bold = False
    
    doc.add_heading('1. Premiumización acelerada', 3)
    doc.add_paragraph('• Crecimiento: +125% interanual en artículos de segunda mano de alta calidad', style='List Bullet')
    doc.add_paragraph('• Oportunidad: Marcas de lujo y productos premium encuentran mercado ávido', style='List Bullet')
    doc.add_paragraph('• Insight: La segunda mano ya no es "lo barato", es "lo inteligente"', style='List Bullet')
    
    doc.add_heading('2. Sostenibilidad como driver principal', 3)
    doc.add_paragraph('• Motivación: 68% de usuarios menciona razones ecológicas como factor clave', style='List Bullet')
    doc.add_paragraph('• Cambio: La economía circular pasa de concepto teórico a práctica cotidiana', style='List Bullet')
    doc.add_paragraph('• Implicación: Valores personales guían decisiones de consumo', style='List Bullet')
    
    doc.add_heading('3. Importancia de comunidades locales', 3)
    doc.add_paragraph('• Realidad: Las transacciones más exitosas ocurren en radios reducidos (<50km)', style='List Bullet')
    doc.add_paragraph('• Insight: La confianza geográfica supera a la escala global pura', style='List Bullet')
    doc.add_paragraph('• Oportunidad: Mercados hiperlocales con alta densidad de confianza', style='List Bullet')
    
    doc.add_heading('4.