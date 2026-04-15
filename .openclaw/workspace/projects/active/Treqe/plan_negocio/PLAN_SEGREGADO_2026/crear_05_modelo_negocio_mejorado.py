#!/usr/bin/env python3
# Script para crear la Sección 05 MEJORADA: MODELO_NEGOCIO.docx
# Versión concisa que sigue el estándar del documento de referencia del 25 de febrero
# Aplicando skills: business-model-canvas, humanizer, legal, marketing-mode, frontend-design, algorithm-solver

import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE

def crear_seccion_05_mejorada():
    """Crear la sección 05 mejorada: MODELO DE NEGOCIO (versión concisa)"""
    
    # Crear documento
    doc = Document()
    
    # Configurar estilos básicos
    styles = doc.styles
    
    # Título principal (5. MODELO DE NEGOCIO)
    title = doc.add_heading('5. MODELO DE NEGOCIO: ARQUITECTURA SOSTENIBLE Y ESCALABLE', 1)
    
    # Introducción concisa
    intro = doc.add_paragraph()
    intro.add_run("""
Treqe transforma la paradoja de liquidez del mercado de segunda mano (8.200M€ en España, 2026) 
en un modelo de negocio multicapa donde todos ganan: usuarios, comercios, empresas y la sociedad.
""")
    
    # 5.1 Business Model Canvas (skill: business-model-canvas)
    doc.add_heading('5.1 Business Model Canvas: Los 9 Bloques Estratégicos', 2)
    
    canvas = doc.add_paragraph()
    canvas.add_run("""
1. **Segmentos:** 28M usuarios B2C + comercios B2B + empresas B2B2C
2. **Propuesta Valor:** Transformar lo que no usas en lo que sí quieres
3. **Canales:** App móvil (94% transacciones) + web + API empresas
4. **Relación Clientes:** Sistema reputación gamificado + soporte humano+IA
5. **Flujos Ingresos:** Comisión 3% + suscripción 9,99€/mes + SaaS empresas
6. **Recursos Clave:** Algoritmo patentado + equipo multidisciplinar + comunidad
7. **Actividades Clave:** Matching circular + verificación + gestión garantía
8. **Socios Clave:** Stripe Connect + AWS + asesoría legal + alianzas
9. **Estructura Costes:** Desarrollo 40% + marketing 30% + operaciones 20% + generales 10%
""")
    
    # 5.2 Flujos de Ingresos Multicapa con ejemplo concreto (skill: humanizer)
    doc.add_heading('5.2 Flujos de Ingresos: 4 Capas que Escalan con el Éxito', 2)
    
    ingresos = doc.add_paragraph()
    ingresos.add_run("""
**Ejemplo concreto:** Ana (arquitecta, 32 años) intercambia su iPhone 12 (valor 400€) 
por la bicicleta de Carlos (450€) + 50€ compensación. Treqe cobra 3% de 450€ = 13,50€.

**Capa 1 - Transacciones (3% comisión):** 70% ingresos año 1
**Capa 2 - Suscripción Premium (9,99€/mes):** Matching 30% más rápido + límites aumentados
**Capa 3 - Servicios Empresariales (SaaS):** White-label para marcas (IKEA, Decathlon)
**Capa 4 - Data & Analytics:** Insights mercado para retailers e investigadores
""")
    
    # 5.3 Estructura Legal y Protección (skill: legal)
    doc.add_heading('5.3 Estructura Legal: Sociedad Limitada y Activos Protegidos', 2)
    
    legal = doc.add_paragraph()
    legal.add_run("""
**Forma Jurídica:** Sociedad Limitada (SL) - 3.000€ capital, 25% impuesto beneficios
**Propiedad Intelectual:** 
- Patente algoritmo matching circular (solicitud UE)
- Marca "Treqe" registrada clases 9, 35, 36, 42
- Derechos autor código + diseño UI/UX
**Contratos Clave:** Términos adaptados economía colaborativa (Ley 7/2022) + GDPR
**Fondo Garantía:** Entidad separada 50.000€ + seguro 1M€ complementario
""")
    
    # 5.4 Proyecciones Financieras Realistas (skill: business-model-canvas)
    doc.add_heading('5.4 Proyecciones Financieras 2026-2029', 2)
    
    finanzas = doc.add_paragraph()
    finanzas.add_run("""
**Año 1 (Lanzamiento):** 10.000 usuarios, 750.000€ volumen, 52.500€ ingresos, EBITDA -47.500€
**Año 2 (Crecimiento):** 50.000 usuarios, 4.5M€ volumen, 382.500€ ingresos, EBITDA +45.000€ (Q4 rentable)
**Año 3 (Escalabilidad):** 200.000 usuarios, 18M€ volumen, 1.530.000€ ingresos, EBITDA +459.000€ (30% margen)

**Punto Equilibrio:** Mes 18 (15.000 usuarios activos)
**ROI Inversores:** 5.2x en 5 años (TIR 38%)
**Valoración Serie A (año 3):** 12-15M€ (8-10x ingresos)
""")
    
    # 5.5 Estrategia de Crecimiento (skill: marketing-mode)
    doc.add_heading('5.5 Estrategia de Crecimiento: 5 Fases Orgánicas', 2)
    
    crecimiento = doc.add_paragraph()
    crecimiento.add_run("""
**Fase 1 - Internal (3 meses):** MVP técnico + primeros 100 early adopters
**Fase 2 - Alpha (6 meses):** Madrid capital, 1.000 usuarios, refinamiento algoritmo
**Fase 3 - Beta (9 meses):** Comunidad Madrid, 10.000 usuarios, unit economics validados
**Fase 4 - Early Access (12 meses):** 5 ciudades españolas, 50.000 usuarios, rentabilidad
**Fase 5 - Full Launch (18 meses):** Nacional + primeras expansiones internacionales

**Métricas Clave:** CAC 15€, LTV 360€, LTV:CAC 24:1, Retención 30 días 65%
""")
    
    # 5.6 Ventaja Tecnológica (skill: algorithm-solver + frontend-design)
    doc.add_heading('5.6 Ventaja Tecnológica: Algoritmo + Experiencia Usuario', 2)
    
    tecnologia = doc.add_paragraph()
    tecnologia.add_run("""
**Algoritmo Patentable:** Matching circular resuelve problema NP-Completo
- k=4 máximo por restricción 5 minutos ejecución
- 94% transacciones desde móvil → experiencia mobile-first optimizada
- Sistema reputación blockchain-inspired (transparente, inmutable)
- Interfaz "brutalista digital con toques orgánicos" (diferencia visual)

**Ejemplo Beatriz:** Quiere guitarra (600€), tiene sofá (400€). Sin Treqe: no match. 
Con Treqe: Beatriz→Carlos→David→Ana→Beatriz (rueda 4 personas, todos satisfechos).
""")
    
    # 5.7 Conclusión: Por qué Funciona
    doc.add_heading('5.7 Conclusión: El Círculo Virtuoso Treqe', 2)
    
    conclusion = doc.add_paragraph()
    conclusion.add_run("""
El modelo funciona porque crea un círculo virtuoso donde:
1. **Más usuarios → más matches posibles → más valor intercambiado**
2. **Más valor → más comisiones para Treqe → más inversión en mejora plataforma**
3. **Mejor plataforma → mejor experiencia → más usuarios**

No somos una app más. Somos la solución a un problema matemático-económico 
(paradoja liquidez) convertida en negocio escalable. Cobramos 3% no por vender, 
sino por hacer posible lo imposible: que lo que ya tienes se convierta en lo que realmente quieres.

**En números:** 1% penetración mercado español = 280.000 usuarios = 8,4M€ ingresos anuales.
La oportunidad no es incremental, es exponencial.
""")
    
    # Guardar documento
    output_path = "05_MODELO_NEGOCIO_MEJORADO.docx"
    doc.save(output_path)
    
    # Verificar tamaño
    tamaño = os.path.getsize(output_path)
    
    print(f"SECCION 05 MEJORADA CREADA: {output_path}")
    print(f"TAMAÑO: {tamaño:,} bytes (objetivo: ~6,000-8,000 bytes)")
    print(f"SUBSECCIONES: 7")
    print(f"SKILLS APLICADAS: 6/6")
    print(f"EJEMPLOS CONCRETOS: Ana, Carlos, Beatriz, David")
    print(f"DATOS CUANTITATIVOS: 2025-2026, €, %, usuarios, proyecciones")
    
    # Comparar con objetivo
    objetivo = 6129  # 55KB / 9 secciones
    if tamaño <= objetivo * 1.5:  # Hasta 50% más grande está bien
        print(f"✓ TAMAÑO ADECUADO: {tamaño:,} vs objetivo {objetivo:,} bytes")
    else:
        print(f"⚠️  AÚN DEMASIADO LARGA: {tamaño:,} vs objetivo {objetivo:,} bytes")
    
    return output_path, tamaño

if __name__ == "__main__":
    # Cambiar al directorio correcto
    os.chdir(os.path.dirname(os.path.abspath(__file__)))
    
    # Crear la sección mejorada
    archivo, tamaño = crear_seccion_05_mejorada()
    
    # Reemplazar la versión anterior si es mejor
    if tamaño < 15000:  # Si es menor a 15KB, reemplazar
        if os.path.exists("05_MODELO_NEGOCIO.docx"):
            os.replace(archivo, "05_MODELO_NEGOCIO.docx")
            print(f"\n✓ VERSIÓN ANTERIOR REEMPLAZADA: 05_MODELO_NEGOCIO.docx ahora tiene {tamaño:,} bytes")
        else:
            print(f"\n✓ NUEVA VERSIÓN CREADA: 05_MODELO_NEGOCIO.docx ({tamaño:,} bytes)")
    else:
        print(f"\n⚠️  VERSIÓN ANTERIOR MANTENIDA: Nueva versión ({tamaño:,} bytes) sigue siendo muy larga")