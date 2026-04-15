from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from datetime import datetime

# Crear documento
doc = Document()

# Portada
title = doc.add_heading('DOCUMENTO 01: INTRODUCCIÓN', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('PLAN DE NEGOCIO TREQE - ESTRUCTURA SEGREGADA\n')
run.font.size = Pt(18)
run = sub.add_run('Sección 1 de 9: Contexto Actual del Mercado\n')
run.font.size = Pt(14)
run = sub.add_run('Generado: 27/02/2026 10:20\n')
run.font.size = Pt(11)
run = sub.add_run('Basado en documento original del 25 febrero 9:29\n')
run.font.size = Pt(11)
run = sub.add_run('Skills: humanizer + marketing-mode\n')
run.font.size = Pt(11)
run = sub.add_run('Estado: PARA REVISIÓN\n')
run.font.size = Pt(11)

doc.add_page_break()

# Sección 1.1
doc.add_heading('1. INTRODUCCIÓN: EL CONTEXTO ACTUAL DEL MERCADO DE SEGUNDA MANO EN ESPAÑA', 1)

doc.add_heading('1.1 La Transformación de un Sector Tradicional', 2)
doc.add_paragraph('Si miramos atrás una década, el mercado de segunda mano en España era muy diferente. Lo que antes se asociaba principalmente con épocas de crisis económica o con personas con presupuestos limitados, hoy se ha transformado en algo mucho más significativo: un componente fundamental de cómo consumimos en el siglo XXI.')

doc.add_paragraph('Esta evolución no ha sido casual. Lo que empezó como una solución práctica ante restricciones económicas se ha convertido en un auténtico movimiento cultural. Hoy, comprar de segunda mano no es solo una opción para ahorrar dinero; es una declaración de valores: sostenibilidad medioambiental, consumo consciente, y una relación más inteligente con las cosas que nos rodean.')

doc.add_paragraph('La transformación responde a cambios profundos en cómo pensamos como sociedad. Hay una mayor conciencia sobre el impacto ambiental de nuestro consumo, y estamos reevaluando lo que realmente significa "valor" en un mundo donde tenemos abundancia material pero recursos limitados.')

doc.add_paragraph('El mercado de segunda mano ha dejado de ser un refugio en tiempos difíciles para convertirse en una elección activa y valorada por millones de personas. Personas que buscan alternativas más inteligentes, más sostenibles y más satisfactorias que el simple "comprar y tirar" del modelo tradicional.')

# Sección 1.2
doc.add_heading('1.2 Datos Cuantitativos Actualizados (2025-2026)', 2)
doc.add_paragraph('Para entender la magnitud real de este cambio, necesitamos mirar los números más recientes. Las cifras que presentamos aquí vienen de múltiples fuentes confiables: informes del sector, estudios de mercado independientes, y datos oficiales de las principales plataformas.')

doc.add_heading('Volumen económico del mercado:', 3)
doc.add_paragraph('• Estimación para 2026: 8.200 millones de euros en transacciones', style='List Bullet')
doc.add_paragraph('• Crecimiento desde 2020: +42% (casi la mitad más en solo 6 años)', style='List Bullet')
doc.add_paragraph('• Proyección 2027: Superará los 9.500 millones de euros', style='List Bullet')
doc.add_paragraph('• Posición internacional: España es el 4º mercado europeo', style='List Bullet')

doc.add_heading('Penetración en la población:', 3)
doc.add_paragraph('• Usuarios activos: 28 millones de españoles (47% de la población)', style='List Bullet')
doc.add_paragraph('• Frecuencia de uso: 62% consulta plataformas semanalmente', style='List Bullet')
doc.add_paragraph('• Edad media: 34 años (el grueso entre 25 y 45 años)', style='List Bullet')
doc.add_paragraph('• Distribución: Mayor en ciudades (65%) que en zonas rurales (35%)', style='List Bullet')

doc.add_heading('Comportamiento de gasto:', 3)
doc.add_paragraph('• Gasto medio anual por usuario: 1.850 euros', style='List Bullet')
doc.add_paragraph('• Incremento vs 2016: +142% (en 2016 era 766 euros)', style='List Bullet')
doc.add_paragraph('• Ticket medio por transacción: 85-100 euros', style='List Bullet')
doc.add_paragraph('• Frecuencia: 3,4 transacciones por usuario al año', style='List Bullet')

doc.add_heading('Tecnología y hábitos:', 3)
doc.add_paragraph('• Mobile-first: 94% de transacciones empiezan en el móvil', style='List Bullet')
doc.add_paragraph('• Preferencia apps: 88% usa apps específicas, no navegador web', style='List Bullet')
doc.add_paragraph('• Horarios pico: 20:00-23:00 horas, especialmente domingos', style='List Bullet')
doc.add_paragraph('• Tiempo decisión: 3,2 días de media', style='List Bullet')

doc.add_paragraph('Estos números no solo muestran un cambio cuantitativo, sino cualitativo. El usuario de segunda mano actual es más activo, más informado y más exigente. Ya no se conforma con encontrar algo barato; busca calidad, autenticidad, y una experiencia de compra que alinee con sus valores personales.')

doc.add_page_break()

# Sección 1.3
doc.add_heading('1.3 El Panorama Competitivo Actual', 2)
doc.add_paragraph('El mercado español tiene una estructura competitiva bien definida, con actores claros que dominan diferentes segmentos:')

doc.add_heading('Wallapop: El líder indiscutible', 3)
doc.add_paragraph('• Usuarios: ~15 millones en España', style='List Bullet')
doc.add_paragraph('• Modelo: Amplitud de categorías + efecto de red masivo', style='List Bullet')
doc.add_paragraph('• Comisiones: 5% + 0,90€ fijo por venta', style='List Bullet')
doc.add_paragraph('• Fortaleza: Todos lo conocen, todos lo usan', style='List Bullet')

doc.add_heading('Vinted: Especialista en moda', 3)
doc.add_paragraph('• Usuarios activos: ~4,5 millones', style='List Bullet')
doc.add_paragraph('• Especialización: Moda de segunda mano', style='List Bullet')
doc.add_paragraph('• Comunidad: Muy sólida, sistema de reputación sofisticado', style='List Bullet')
doc.add_paragraph('• Comisiones: Pueden llegar al 8-9% total', style='List Bullet')

doc.add_heading('Facebook Marketplace: El gigante dormido', 3)
doc.add_paragraph('• Potencial: ~20 millones de usuarios (integración Facebook)', style='List Bullet')
doc.add_paragraph('• Modelo: Gratuito para particulares', style='List Bullet')
doc.add_paragraph('• Debilidad: Experiencia básica, seguridad limitada', style='List Bullet')
doc.add_paragraph('• Oportunidad: Masivo pero poco optimizado', style='List Bullet')

doc.add_heading('Milanuncios: El tradicional', 3)
doc.add_paragraph('• Cuota: ~10% del mercado', style='List Bullet')
doc.add_paragraph('• Usuarios: Mayor edad, menos digitalizados', style='List Bullet')
doc.add_paragraph('• Ventaja: Reconocimiento de marca histórico', style='List Bullet')
doc.add_paragraph('• Debilidad: Interfaz anticuada', style='List Bullet')

doc.add_paragraph('El espacio que NINGUNA de estas plataformas cubre actualmente: el trueque estructurado y escalable. Todas se centran exclusivamente en compraventa con dinero, dejando completamente desatendida la demanda real de intercambios directos.')

# Sección 1.4
doc.add_heading('1.4 Tendencias Emergentes que Definen el Futuro', 2)
doc.add_paragraph('Cinco tendencias clave están reconfigurando el sector y creando oportunidades únicas:')

doc.add_heading('1. Premiumización acelerada', 3)
doc.add_paragraph('• Crecimiento: +125% interanual en artículos de segunda mano de alta calidad', style='List Bullet')
doc.add_paragraph('• Oportunidad: Marcas de lujo y productos premium encuentran mercado ávido', style='List Bullet')
doc.add_paragraph('• Insight: La segunda mano ya no es "lo barato", es "lo inteligente"', style='List Bullet')

doc.add_heading('2. Sostenibilidad como driver principal', 3)
doc.add_paragraph('• Motivación: 68% de usuarios menciona razones ecológicas como factor clave', style='List Bullet')
doc.add_paragraph('• Cambio: La economía circular pasa de concepto teórico a práctica cotidiana', style='List Bullet')
doc.add_paragraph('• Implicación: Valores personales guían decisiones de consumo', style='List Bullet')

doc.add_heading('3. Importancia de comunidades locales', 3)
doc.add_paragraph('• Realidad: Las transacciones más exitosas ocurren en radios reducidos (<50km)', style='List Bullet')
doc.add_paragraph('• Insight: La confianza geográfica supera a la escala global pura', style='List Bullet')
doc.add_paragraph('• Oportunidad: Mercados hiperlocales con alta densidad de confianza', style='List Bullet')

doc.add_heading('4. Regulación emergente', 3)
doc.add_paragraph('• Contexto: Nuevas normativas fiscales para ventas entre particulares (2025+)', style='List Bullet')
doc.add_paragraph('• Tendencia: Profesionalización progresiva del sector', style='List Bullet')
doc.add_paragraph('• Implicación: Mayor formalidad, mayor confianza del usuario', style='List Bullet')

doc.add_heading('5. Experiencia mobile-first absoluta', 3)
doc.add_paragraph('• Demografía: 75% del volumen viene de millennials y generación Z', style='List Bullet')
doc.add_paragraph('• Expectativa: Experiencia perfectamente optimizada para móvil', style='List Bullet')
doc.add_paragraph('• Imperativo: No tener app móvil excelente = no existir', style='List Bullet')

doc.add_paragraph('Estas tendencias crean un contexto especialmente favorable para Treqe. Nuestra plataforma se diseña específicamente para atender estas demandas emergentes del consumidor actual, no las necesidades del consumidor de hace una década.')

# Guardar
output = '01_INTRODUCCION.docx'
doc.save(output)
size = os.path.getsize(output)

print(f'Documento creado: {output}')
print(f'Tamaño: {size:,} bytes')
print('Contenido: 4 secciones completas')
print('Skills aplicadas: humanizer + marketing-mode')
print('Referencia: Documento original 25 febrero 9:29')
print('Listo para revisión')