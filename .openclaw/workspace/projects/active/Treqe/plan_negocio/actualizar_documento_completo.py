#!/usr/bin/env python3
"""
Actualizar el documento Treqe con las conclusiones del algoritmo y contenido humanizado
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re

def actualizar_documento():
    print("ACTUALIZANDO DOCUMENTO TREQE CON CONCLUSIONES DEL ALGORITMO...")
    
    try:
        # Cargar documento existente
        doc = Document('Plan_Negocio_Treqe_CON_SOLUCIONES_2026.docx')
        
        print(f"Documento cargado: {len(doc.paragraphs)} párrafos")
        
        # ============================================
        # 1. ACTUALIZAR SECCIÓN 3.2: PROCESO DEL USUARIO
        # ============================================
        print("\n1. Actualizando Sección 3.2: Proceso del usuario...")
        
        # Buscar la sección 3.2
        for i, para in enumerate(doc.paragraphs):
            texto = para.text.strip()
            if "3.2 El Mecanismo Detrás de la Magia" in texto or "3.2 El Mecanismo Operativo Paso a Paso" in texto:
                print(f"  Encontrada sección 3.2 en párrafo {i}")
                
                # Eliminar contenido existente hasta la siguiente sección
                j = i + 1
                while j < len(doc.paragraphs):
                    texto_j = doc.paragraphs[j].text.strip()
                    if texto_j.startswith("3.3") or texto_j.startswith("3.4") or texto_j.startswith("4."):
                        break
                    # Marcar para eliminación
                    doc.paragraphs[j].clear()
                    j += 1
                
                # Insertar nuevo contenido humanizado
                nuevo_contenido = [
                    "Para comprender cómo Treqe transforma una idea conceptual en una experiencia tangible, es necesario analizar el mecanismo operativo paso a paso. Cada fase ha sido diseñada no solo pensando en la eficiencia técnica, sino especialmente en la experiencia humana.",
                    "",
                    "**Para el usuario, Treqe funciona así:**",
                    "",
                    "**1. Cuenta la historia de lo que tienes**",
                    "No subes un 'artículo'. Cuentas la historia de ese objeto que ya no encaja en tu vida, pero que todavía tiene valor. Subes fotos que muestran por qué lo apreciabas. Le pones un precio justo, no porque sea una transacción fría, sino porque reconoces lo que vale.",
                    "",
                    "**2. Descubre lo que te emociona**",
                    "Navegas, miras, te dejas sorprender. No buscas solo lo que 'necesitas' - descubres lo que te gusta. Marcas ese jersey que tiene exactamente el color que buscabas, esa lámpara que iluminaría perfectamente tu rincón de lectura, esa cámara con la que siempre has querido aprender fotografía.",
                    "",
                    "**3. Vive tu vida mientras la magia ocurre**",
                    "Aquí está lo mejor: tú no esperas. El sistema trabaja en silencio, buscando combinaciones inteligentes cada 10 minutos. Mientras tú trabajas, cenas, duermes, Treqe busca caminos donde antes solo había paredes.",
                    "",
                    "**4. La notificación que cambia todo**",
                    "Un día, tu teléfono vibra: '¡Encontramos un intercambio para ti!'",
                    "No es un algoritmo frío el que te escribe. Es el resultado de conectar tus deseos con los de otras personas reales.",
                    "",
                    "**5. Conocerse y confiar**",
                    "El resto es humano: coordináis envíos, activáis las garantías, os enviáis mensajes. La parte difícil - encontrar a las personas adecuadas - ya está hecha.",
                    "",
                    "**Lo que hace especial este proceso:**",
                    "- **No esperas activamente:** El sistema trabaja en segundo plano",
                    "- **No negocias desde cero:** El algoritmo ya ha encontrado una solución viable",
                    "- **No arriesgas:** Las garantías protegen cada paso",
                    "- **No es transaccional:** Es una conversación entre personas",
                    "",
                    "**Aprendizaje clave del desarrollo:** En mercados reales, los trueques directos (k=2) funcionan menos del 5% del tiempo. Por eso Treqe no busca intercambios 1:1 - busca círculos de 3, 4 o 5 personas donde todos obtienen lo que quieren."
                ]
                
                # Insertar después del título
                for contenido in nuevo_contenido:
                    nuevo_parrafo = doc.paragraphs[i].insert_paragraph_before(contenido)
                    # Mantener estilo similar
                    if contenido.startswith("**") and contenido.endswith("**"):
                        for run in nuevo_parrafo.runs:
                            run.bold = True
        
        # ============================================
        # 2. ACTUALIZAR SECCIÓN 4.2: ALGORITMO
        # ============================================
        print("\n2. Actualizando Sección 4.2: Algoritmo y conclusiones...")
        
        # Buscar la sección 4.2
        for i, para in enumerate(doc.paragraphs):
            texto = para.text.strip()
            if "4.2 Ventajas Tecnológicas" in texto or "Nivel 1: El Algoritmo de Matching" in texto:
                print(f"  Encontrada sección de algoritmo en párrafo {i}")
                
                # Buscar dónde insertar las conclusiones del algoritmo
                # Buscar después de "Nivel 1: El Algoritmo de Matching"
                for j in range(i, min(i+20, len(doc.paragraphs))):
                    if "Nivel 2:" in doc.paragraphs[j].text or "4.3" in doc.paragraphs[j].text:
                        posicion_insercion = j
                        break
                else:
                    posicion_insercion = i + 10
                
                # Contenido sobre el algoritmo actualizado
                contenido_algoritmo = [
                    "",
                    "**Conclusiones del Desarrollo del Algoritmo:**",
                    "",
                    "Durante el desarrollo del algoritmo de matching de Treqe, aprendimos lecciones cruciales que han dado forma a nuestro enfoque:",
                    "",
                    "**1. k>2 no es el problema, es la solución**",
                    "En mercados reales, la densidad de compatibilidad para intercambios directos (k=2) es inferior al 5%. Esto significa que los trueques 1:1 casi nunca funcionan. Treqe resuelve esto permitiendo círculos de 3, 4 o 5 personas, aumentando exponencialmente las probabilidades de encontrar matches viables.",
                    "",
                    "**2. La estrategia ascendente: k=2 → k_max**",
                    "Nuestro algoritmo no comienza buscando círculos complejos. Empieza con k=2 (intercambios directos) y solo aumenta el tamaño del círculo si no encuentra matches. Este enfoque incremental garantiza que encontramos la solución más simple posible primero.",
                    "",
                    "**3. Límite de tiempo: 5 minutos por batch**",
                    "Hemos establecido un timeout de 5 minutos para cada ejecución del algoritmo. Preferimos encontrar matches para el 80% de los usuarios en 5 minutos, que para el 100% en 5 horas. La perfección es enemiga de lo posible.",
                    "",
                    "**4. Complejidad como ventaja estratégica**",
                    "El problema de encontrar ciclos de intercambio óptimos es matemáticamente complejo (NP-Completo para k>3). Esta complejidad no es un accidente - es una barrera de entrada deliberada que protege nuestro modelo de negocio.",
                    "",
                    "**5. Optimizaciones implementadas:**",
                    "- **Búsqueda heurística:** En lugar de explorar todas las combinaciones (factorial), usamos heurísticas inteligentes",
                    "- **Paralelización:** Ejecución en múltiples núcleos de CPU (8 threads en nuestro hardware)",
                    "- **Poda temprana:** Descartamos caminos inviables rápidamente",
                    "- **Caché de resultados:** Reutilizamos cálculos similares",
                    "",
                    "**6. Validación con datos realistas:**",
                    "Nuestras simulaciones iniciales eran demasiado optimistas (12-14% densidad). Con datos de mercado reales (5% densidad), ajustamos el algoritmo para funcionar en condiciones realistas, no ideales.",
                    "",
                    "**Resultado:** Un algoritmo que encuentra intercambios viables para la mayoría de usuarios en menos de 5 minutos, escalando hasta miles de usuarios simultáneos."
                ]
                
                # Insertar contenido
                for contenido in reversed(contenido_algoritmo):
                    nuevo_parrafo = doc.paragraphs[posicion_insercion].insert_paragraph_before(contenido)
        
        # ============================================
        # 3. ACTUALIZAR SECCIÓN 8.4: SCORING
        # ============================================
        print("\n3. Actualizando Sección 8.4: Sistema de Scoring...")
        
        # Buscar la sección 8.4
        for i, para in enumerate(doc.paragraphs):
            texto = para.text.strip()
            if "8.4 Sistema de Puntuacion" in texto or "ALGORITMO DE SCORING" in texto:
                print(f"  Encontrada sección 8.4 en párrafo {i}")
                
                # Eliminar contenido existente (tabla de scoring)
                # Buscar hasta la siguiente sección
                j = i + 1
                while j < len(doc.paragraphs):
                    texto_j = doc.paragraphs[j].text.strip()
                    if texto_j.startswith("8.3") or texto_j.startswith("8.2") or texto_j.startswith("9."):
                        break
                    doc.paragraphs[j].clear()
                    j += 1
                
                # Nuevo contenido humanizado
                contenido_scoring = [
                    "",
                    "**En Treqe, tu puntuación no es un número. Es la historia de cómo intercambias.**",
                    "",
                    "**Cómo crece tu reputación:**",
                    "- **Cada intercambio exitoso** suma a tu historia",
                    "- **El valor que mueves** muestra tu seriedad",
                    "- **La puntualidad en los envíos** demuestra tu respeto por los demás",
                    "- **La claridad en tus descripciones** construye confianza",
                    "",
                    "**Los niveles no son categorías. Son capítulos de tu historia en Treqe:**",
                    "",
                    "**📖 Capítulo 1: Empiezas**",
                    "(Eres nuevo, aprendes cómo funciona)",
                    "- Puedes recibir propuestas de intercambio",
                    "- Comisión del 1%",
                    "- Envías los paquetes por tu cuenta",
                    "",
                    "**🌱 Capítulo 2: Creces**",
                    "(Ya has completado algunos intercambios)",
                    "- Puedes iniciar intercambios de hasta 200€",
                    "- Acceso al seguro básico",
                    "- La comunidad empieza a reconocerte",
                    "",
                    "**🤝 Capítulo 3: Eres confiable**",
                    "(Tu historial habla por ti)",
                    "- Intercambios de hasta 500€",
                    "- Comisión reducida al 0.9%",
                    "- Descuentos en logística",
                    "- Menos depósito requerido",
                    "",
                    "**🏆 Capítulo 4: Eres parte esencial de la comunidad**",
                    "(Eres un referente en Treqe)",
                    "- Sin límites de valor",
                    "- Comisión del 0.8%",
                    "- Logística incluida",
                    "- Prioridad en búsquedas",
                    "- Posibilidad de ayudar a nuevos usuarios",
                    "",
                    "**La filosofía:** No castigamos los errores ocasionales. Reconocemos el aprendizaje. Valoramos la consistencia. Celebramos la contribución a la comunidad."
                ]
                
                # Insertar después del título
                for contenido in contenido_scoring:
                    nuevo_parrafo = doc.paragraphs[i].insert_paragraph_before(contenido)
        
        # ============================================
        # 4. ACTUALIZAR SECCIÓN 8.3: ENVÍOS
        # ============================================
        print("\n4. Actualizando Sección 8.3: Sistema de Envíos...")
        
        # Buscar la sección 8.3
        for i, para in enumerate(doc.paragraphs):
            texto = para.text.strip()
            if "8.3 Sistema de Envios" in texto or "PROBLEMA: Rueda acordada" in texto:
                print(f"  Encontrada sección 8.3 en párrafo {i}")
                
                # Eliminar contenido existente
                j = i + 1
                while j < len(doc.paragraphs):
                    texto_j = doc.paragraphs[j].text.strip()
                    if texto_j.startswith("8.2") or texto_j.startswith("8.1") or texto_j.startswith("9."):
                        break
                    doc.paragraphs[j].clear()
                    j += 1
                
                # Nuevo contenido humanizado
                contenido_envios = [
                    "",
                    "**Los envíos no son solo paquetes que viajan. Son promesas que se cumplen.**",
                    "",
                    "**Cómo funciona:**",
                    "1. **Acordáis cómo enviar** (cada uno elige lo que le conviene)",
                    "2. **Compartís los códigos de seguimiento** (todos pueden ver dónde está cada paquete)",
                    "3. **Confirmáis la recepción** (un simple '¡llegó!' activa el siguiente paso)",
                    "4. **El sistema libera los pagos** (el dinero cambia de manos solo cuando todos están satisfechos)",
                    "",
                    "**Las opciones:**",
                    "- **Envías tú mismo** (si prefieres control personal)",
                    "- **Usas nuestro servicio asociado** (con descuentos según tu reputación)",
                    "- **Recogida en punto designado** (para ahorrar y conocer a otros usuarios)",
                    "",
                    "**La triple protección:**",
                    "1. **Escrow:** El dinero está seguro hasta que todos cumplen",
                    "2. **Seguro:** Si algo sale mal en el envío, hay cobertura",
                    "3. **Verificación:** Cada paso se confirma antes de continuar",
                    "",
                    "**La magia está en lo simple:** No es un proceso burocrático. Es una conversación entre personas que se están ayudando mutuamente."
                ]
                
                # Insertar después del título
                for contenido in contenido_envios:
                    nuevo_parrafo = doc.paragraphs[i].insert_paragraph_before(contenido)
        
        # ============================================
        # 5. ACTUALIZAR SECCIÓN 8.2: GARANTÍAS
        # ============================================
        print("\n5. Actualizando Sección 8.2: Sistema de Garantías...")
        
        # Buscar la sección 8.2
        for i, para in enumerate(doc.paragraphs):
            texto = para.text.strip()
            if "8.2 Sistema de Garantias" in texto or "PROBLEMA: Usuario recibe" in texto:
                print(f"  Encontrada sección 8.2 en párrafo {i}")
                
                # Eliminar contenido existente
                j = i + 1
                while j < len(doc.paragraphs):
                    texto_j = doc.paragraphs[j].text.strip()
                    if texto_j.startswith("8.1") or texto_j.startswith("9."):
                        break
                    doc.paragraphs[j].clear()
                    j += 1
                
                # Nuevo contenido humanizado
                contenido_garantias = [
                    "",
                    "**En Treqe, la confianza no se da por sentado. Se construye, capa por capa.**",
                    "",
                    "**Capa 1: El dinero está seguro hasta que todos cumplen**",
                    "Usamos un