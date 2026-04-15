#!/usr/bin/env python3
"""
Script simple para crear documento Word del plan de negocio Treqe
"""

import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

# Crear documento
doc = Document()

# Configurar estilos
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Título principal
title = doc.add_heading('PLAN DE NEGOCIO TREQE', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.runs[0]
title_run.font.size = Pt(24)
title_run.font.bold = True

# Subtítulo
subtitle = doc.add_paragraph('Documento único definitivo - Todo en un solo documento')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle.runs[0]
subtitle_run.font.size = Pt(14)
subtitle_run.font.italic = True

doc.add_paragraph(f'Creado el {datetime.now().strftime("%d/%m/%Y")}')
doc.add_page_break()

# ===== PARTE 1: RESUMEN EJECUTIVO =====
doc.add_heading('📋 RESUMEN EJECUTIVO (2 minutos)', level=1)

p = doc.add_paragraph()
p.add_run('El problema que resolvemos: ').bold = True
p.add_run('72% de españoles tiene cosas guardadas que no usa. Intercambiar entre 2 personas tiene solo 5% de probabilidad.')

p = doc.add_paragraph()
p.add_run('Nuestra solución: ').bold = True
p.add_run('Treqe conecta a 3+ personas para intercambios circulares (A→B→C→A). Probabilidad sube al 20-35%.')

p = doc.add_paragraph()
p.add_run('Cómo ganamos dinero: ').bold = True
p.add_run('3% de comisión cuando el intercambio funciona. Solo cobramos cuando tú ganas.')

doc.add_heading('Los números clave', level=2)
doc.add_paragraph('• Inversión inicial: 58.000€', style='List Bullet')
doc.add_paragraph('• Año 1: 8.000 usuarios activos, 360.000€ ingresos, 84.000€ beneficio', style='List Bullet')
doc.add_paragraph('• Año 2: 25.000 usuarios activos, 1.728.000€ ingresos', style='List Bullet')
doc.add_paragraph('• Año 3: 60.000 usuarios activos, 5.040.000€ ingresos', style='List Bullet')

doc.add_heading('Próximos pasos inmediatos', level=2)
doc.add_paragraph('1. Constituir Sociedad Limitada (3.000€)', style='List Number')
doc.add_paragraph('2. Registrar marca "Treqe" España + UE (850€)', style='List Number')
doc.add_paragraph('3. Desarrollar algoritmo v0.1', style='List Number')
doc.add_paragraph('4. Primeros 50 usuarios (personas conocidas)', style='List Number')

doc.add_page_break()

# ===== PARTE 2: EL PROBLEMA REAL =====
doc.add_heading('🎯 PARTE 1: EL PROBLEMA REAL', level=1)

doc.add_heading('Caso real: Ana, Carlos y Beatriz', level=2)

p = doc.add_paragraph()
p.add_run('Situación inicial (todos frustrados):').bold = True

doc.add_paragraph('• Ana: Bicicleta Orbea (300€) → quiere sofá IKEA (300€)', style='List Bullet')
doc.add_paragraph('• Carlos: Sofá IKEA (300€) → quiere ordenador portátil (300€)', style='List Bullet')
doc.add_paragraph('• Beatriz: Ordenador portátil (300€) → quiere bicicleta (300€)', style='List Bullet')

p = doc.add_paragraph()
p.add_run('Con Treqe (72 horas después):').bold = True

doc.add_paragraph('• Día 1: Registro + algoritmo encuentra combinación', style='List Bullet')
doc.add_paragraph('• Día 2: Ana envía bici a Carlos', style='List Bullet')
doc.add_paragraph('• Día 3: Beatriz envía ordenador a Carlos', style='List Bullet')
doc.add_paragraph('• Día 4: Carlos envía sofá a Ana, Beatriz recibe bici', style='List Bullet')

p = doc.add_paragraph()
p.add_run('Resultado final:').bold = True

doc.add_paragraph('• Todos felices con lo que querían', style='List Bullet')
doc.add_paragraph('• Valor intercambiado: 900€', style='List Bullet')
doc.add_paragraph('• Comisión Treqe (3%): 27€ (9€ cada uno)', style='List Bullet')
doc.add_paragraph('• Tiempo invertido: 15 minutos cada uno', style='List Bullet')
doc.add_paragraph('• Tiempo ahorrado vs vender: 10 horas cada uno', style='List Bullet')

doc.add_heading('Datos del mercado español', level=2)
doc.add_paragraph('• Mercado segunda mano: 5.000 millones € anuales', style='List Bullet')
doc.add_paragraph('• Cosas guardadas en hogares: 15.000 millones €', style='List Bullet')
doc.add_paragraph('• Personas que preferirían intercambiar: 65% población', style='List Bullet')
doc.add_paragraph('• Intercambios que no ocurren por dificultad: 80%', style='List Bullet')

doc.add_heading('El problema matemático', level=2)
p = doc.add_paragraph()
p.add_run('Para intercambio directo (2 personas): ').bold = True
p.add_run('5% probabilidad')

p = doc.add_paragraph()
p.add_run('Con Treqe (3 personas): ').bold = True
p.add_run('20% probabilidad (4x más)')

p = doc.add_paragraph()
p.add_run('Con Treqe (4 personas): ').bold = True
p.add_run('35% probabilidad (7x más)')

p = doc.add_paragraph()
p.add_run('Cada persona adicional ').bold = True
p.add_run('multiplica las posibilidades, no las suma.')

doc.add_page_break()

# ===== PARTE 3: LA SOLUCIÓN TREQE =====
doc.add_heading('🚀 PARTE 2: LA SOLUCIÓN TREQE', level=1)

doc.add_heading('Así funciona (4 pasos simples)', level=2)

doc.add_paragraph('1. Cuentas tu historia', style='List Number')
doc.add_paragraph('   • Subes foto de lo que tienes', style='List Bullet')
doc.add_paragraph('   • Describes lo que te emocionaría tener', style='List Bullet')
doc.add_paragraph('   • Estimas el valor', style='List Bullet')

doc.add_paragraph('2. Descubres posibilidades', style='List Number')
doc.add_paragraph('   • Nuestro algoritmo busca combinaciones', style='List Bullet')
doc.add_paragraph('   • No solo 2 personas, sino 3, 4, 5...', style='List Bullet')
doc.add_paragraph('   • En 24 horas tienes opciones reales', style='List Bullet')

doc.add_paragraph('3. Vives tu vida', style='List Number')
doc.add_paragraph('   • Aceptas el intercambio que te convence', style='List Bullet')
doc.add_paragraph('   • Nosotros coordinamos el resto', style='List Bullet')
doc.add_paragraph('   • Tú sigues con tu vida normal', style='List Bullet')

doc.add_paragraph('4. La magia ocurre', style='List Number')
doc.add_paragraph('   • En 72 horas (promedio) recibes lo que querías', style='List Bullet')
doc.add_paragraph('   • Das lo que ya no usabas', style='List Bullet')
doc.add_paragraph('   • Conoces personas reales', style='List Bullet')

doc.add_heading('Comparación con competencia', level=2)

p = doc.add_paragraph()
p.add_run('Wallapop/Vinted (hoy):').bold = True
doc.add_paragraph('   "Vende tu bici por 210€ (30% menos), luego compra sofá por 300€"', style='List Bullet')
doc.add_paragraph('   • 2 transacciones separadas', style='List Bullet')
doc.add_paragraph('   • Regateos interminables', style='List Bullet')
doc.add_paragraph('   • 10+ horas perdidas', style='List Bullet')
doc.add_paragraph('   • Riesgo de estafa', style='List Bullet')
doc.add_paragraph('   • Coste real: 90€ + 10 horas', style='List Bullet')

p = doc.add_paragraph()
p.add_run('Treqe (nuestra propuesta):').bold = True
doc.add_paragraph('   "Tu bici ya es tu sofá. Solo necesitábamos encontrar a Carlos y Beatriz."', style='List Bullet')
doc.add_paragraph('   • 1 intercambio circular', style='List Bullet')
doc.add_paragraph('   • Sin regateos', style='List Bullet')
doc.add_paragraph('   • 15 minutos de tu tiempo', style='List Bullet')
doc.add_paragraph('   • Cero riesgo (triple protección)', style='List Bullet')
doc.add_paragraph('   • Coste real: 9€ + 15 minutos', style='List Bullet')

p = doc.add_paragraph()
p.add_run('Ahorro con Treqe: ').bold = True
p.add_run('81€ + 9,75 horas por intercambio')

doc.add_page_break()

# ===== PARTE 4: MODELO DE NEGOCIO =====
doc.add_heading('💰 PARTE 3: MODELO DE NEGOCIO', level=1)

doc.add_heading('Comisión 3% (transparente)', level=2)
p = doc.add_paragraph()
p.add_run('Te decimos en qué se va tu 3%:').bold = True

doc.add_paragraph('• 1% para el cerebro (algoritmo): 0,30€ por intercambio 100€', style='List Bullet')
doc.add_paragraph('• 1% para la confianza (seguridad): 0,30€ por intercambio 100€', style='List Bullet')
doc.add_paragraph('• 1% para el futuro (mejoras): 0,30€ por intercambio 100€', style='List Bullet')

doc.add_heading('Segmentos de clientes', level=2)
doc.add_paragraph('1. Millennials urbanos (25-35 años): 5 millones en España', style='List Number')
doc.add_paragraph('2. Familias con hijos (35-50 años): 3 millones de hogares', style='List Number')
doc.add_paragraph('3. Estudiantes universitarios (18-25 años): 1,5 millones', style='List Number')

doc.add_heading('Flujos de ingresos', level=2)
doc.add_paragraph('• Principal: Comisión 3% sobre intercambios', style='List Bullet')
doc.add_paragraph('• Secundario: Suscripción Premium 9,99€/mes', style='List Bullet')
doc.add_paragraph('• Terciario: Publicidad contextual relevante', style='List Bullet')
doc.add_paragraph('• Futuro: Data insights y estudios de mercado', style='List Bullet')

doc.add_page_break()

# ===== PARTE 5: PROYECCIONES FINANCIERAS =====
doc.add_heading('📈 PARTE 4: PROYECCIONES FINANCIERAS', level=1)

doc.add_heading('Inversión inicial: 58.000€', level=2)
doc.add_paragraph('• Tecnología (40%): 23.200€', style='List Bullet')
doc.add_paragraph('• Marketing (35%): 20.300€', style='List Bullet')
doc.add_paragraph('• Operaciones (25%): 14.500€', style='List Bullet')

doc.add_heading('Proyecciones año a año', level=2)

p = doc.add_paragraph()
p.add_run('Año 1 (2026) - La prueba:').bold = True
doc.add_paragraph('   • Usuarios totales: 50.000', style='List Bullet')
doc.add_paragraph('   • Usuarios activos: 8.000 (16%)', style='List Bullet')
doc.add_paragraph('   • Intercambios/mes: 10.000', style='List Bullet')
doc.add_paragraph('   • Valor medio: 100€', style='List Bullet')
doc.add_paragraph('   • Ingresos: 360.000€/año', style='List Bullet')
doc.add_paragraph('   • Beneficio: 84.000€/año', style='List Bullet')

p = doc.add_paragraph()
p.add_run('Año 2 (2027) - El crecimiento:').bold = True
doc.add_paragraph('   • Usuarios totales: 150.000', style='List Bullet')
doc.add_paragraph('   • Usuarios activos: 25.000 (17%)', style='List Bullet')
doc.add_paragraph('   • Intercambios/mes: 40.000', style='List Bullet')
doc.add_paragraph('   • Valor medio: 120€', style='List Bullet')
doc.add_paragraph('   • Ingresos: 1.728.000€/año', style='List Bullet')
doc.add_paragraph('   • Beneficio: 960.000€/año', style='List Bullet')

p = doc.add_paragraph()
p.add_run('Año 3 (2028) - La consolidación:').bold = True
doc.add_paragraph('   • Usuarios totales: 350.000', style='List Bullet')
doc.add_paragraph('   • Usuarios activos: 60.000 (17%)', style='List Bullet')
doc.add_paragraph('   • Intercambios/mes: 100.000', style='List Bullet')
doc.add_paragraph('   • Valor medio: 140€', style='List Bullet')
doc.add_paragraph('   • Ingresos: 5.040.000€/año', style='List Bullet')
doc.add_paragraph('   • Beneficio: 3.000.000€/año', style='List Bullet')

doc.add_heading('Unit Economics', level=2)
doc.add_paragraph('• CAC (Costo Adquisición Usuario): 15€ año 1 → 10€ año 3', style='List Bullet')
doc.add_paragraph('• LTV (Valor Vida Usuario): 360€ año 1 → 1.200€ año 3', style='List Bullet')
doc.add_paragraph('• Ratio LTV:CAC: 24:1 año 1 → 120:1 año 3', style='List Bullet')
doc.add_paragraph('• Payback Period: 15 días por usuario, 8,3 meses inversión total', style='List Bullet')

doc.add_page_break()

# ===== PARTE 6: CRONOGRAMA 12 MESES =====
doc.add_heading('📅 PARTE 5: CRONOGRAMA 12 MESES', level=1)

doc.add_heading('Mes 1-3: Los cimientos', level=2)
doc.add_paragraph('• Semana 1-2: Fundación legal (SL + marca)', style='List Bullet')
doc.add_paragraph('• Semana 3-4: Tecnología básica (dominio + landing)', style='List Bullet')
doc.add_paragraph('• Mes 2: Primeros usuarios reales (50 personas)', style='List Bullet')
doc.add_paragraph('• Mes 3: Automatización básica (algoritmo v1.0)', style='List Bullet')
doc.add_paragraph('• Objetivo mes 3: 100 usuarios, 20 intercambios', style='List Bullet')

doc.add_heading('Mes 4-6: Primera ciudad (Madrid)', level=2)
doc.add_paragraph('• Mes 4: Escala Madrid (1.000 usuarios)', style='List Bullet')
doc.add_paragraph('• Mes 5: Optimización Madrid (2.000 usuarios)', style='List Bullet')
doc.add_paragraph('• Mes 6: Consolidación Madrid (3.000 usuarios)', style='List Bullet')
doc.add_paragraph('• Presupuesto gastado: 25.000€ (de 58.000€)', style='List Bullet')

doc.add_heading('Mes 7-9: Expansión a 3 ciudades', level=2)
doc.add_paragraph('• Ciudades: Barcelona, Valencia, Sevilla', style='List Bullet')
doc.add_paragraph('• Mes 7: Barcelona (1.000 usuarios)', style='List Bullet')
doc.add_paragraph('• Mes 8: Valencia + Sevilla (1.000 cada una)', style='List Bullet')
doc.add_paragraph('• Mes 9: Integración multi-ciudad', style='List Bullet')
doc.add_paragraph('• Objetivo: 5.000 usuarios total', style='List Bul