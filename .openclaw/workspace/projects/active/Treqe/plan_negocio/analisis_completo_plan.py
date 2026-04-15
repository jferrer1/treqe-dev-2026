#!/usr/bin/env python3
"""
Análisis completo del plan de negocio Treqe
"""

from docx import Document
import re

def analizar_documento_completo():
    print("ANÁLISIS COMPLETO DEL PLAN DE NEGOCIO TREQE")
    print("="*60)
    
    # Cargar documento
    try:
        doc = Document('Plan_Negocio_Treqe_DEFINITIVO_FINAL.docx')
        print(f"Documento cargado: {len(doc.paragraphs)} párrafos")
    except:
        print("Error cargando documento definitivo")
        return
    
    # Análisis de estructura
    print("\n1. ESTRUCTURA DEL DOCUMENTO:")
    print("-"*40)
    
    secciones_principales = []
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text and len(text) < 100:
            # Buscar secciones numeradas
            if re.match(r'^\d+\.\s+[A-Z]', text):
                secciones_principales.append((i, text))
    
    print(f"Secciones principales encontradas: {len(secciones_principales)}")
    for idx, (num, texto) in enumerate(secciones_principales[:15]):
        print(f"  {idx+1}. {texto[:60]}...")
    
    # Análisis de contenido por sección
    print("\n2. ANÁLISIS DE CONTENIDO POR SECCIÓN:")
    print("-"*40)
    
    secciones_contenido = {
        'introduccion': 0,
        'problema': 0,
        'solucion': 0,
        'ventaja_competitiva': 0,
        'modelo_negocio': 0,
        'financiero': 0,
        'equipo': 0,
        'riesgos': 0,
        'conclusiones': 0,
        'marketing': 0,
        'legal': 0,
        'anexos': 0
    }
    
    keywords = {
        'introduccion': ['introducción', 'contexto', 'mercado', 'tendencias'],
        'problema': ['problema', 'limitación', 'dificultad', 'desafío'],
        'solucion': ['solución', 'treqe', 'algoritmo', 'matching', 'rueda'],
        'ventaja_competitiva': ['ventaja', 'competitivo', 'diferenciación', 'barrera'],
        'modelo_negocio': ['modelo', 'ingresos', 'precio', 'comisión', 'tarifa'],
        'financiero': ['financiero', 'proyección', 'inversión', 'presupuesto', 'roi'],
        'equipo': ['equipo', 'fundador', 'ejecución', 'plan'],
        'riesgos': ['riesgo', 'mitigación', 'contingencia', 'problema'],
        'conclusiones': ['conclusión', 'resumen', 'recomendación', 'próximo'],
        'marketing': ['marketing', 'promoción', 'adquisición', 'cac', 'ltv'],
        'legal': ['legal', 'jurídico', 'regulatorio', 'contrato', 'propiedad'],
        'anexos': ['anexo', 'apéndice', 'wireframe', 'estudio', 'investigación']
    }
    
    for para in doc.paragraphs:
        text = para.text.lower()
        for seccion, palabras in keywords.items():
            for palabra in palabras:
                if palabra in text:
                    secciones_contenido[seccion] += 1
                    break
    
    print("Contenido por sección (párrafos con keywords):")
    for seccion, count in secciones_contenido.items():
        print(f"  {seccion.replace('_', ' ').title():20} : {count:3} párrafos")
    
    # Análisis de calidad
    print("\n3. ANÁLISIS DE CALIDAD:")
    print("-"*40)
    
    # Buscar elementos clave
    elementos_clave = {
        'algoritmo_detallado': any('algoritmo' in p.text.lower() and 'k=' in p.text for p in doc.paragraphs),
        'presupuesto_detallado': any('presupuesto' in p.text.lower() and '€' in p.text for p in doc.paragraphs),
        'metricas_kpi': any('kpi' in p.text.lower() or 'métrica' in p.text.lower() for p in doc.paragraphs),
        'plan_marketing': any('marketing' in p.text.lower() and ('5-phase' in p.text.lower() or 'fase' in p.text.lower()) for p in doc.paragraphs),
        'analisis_riesgos': any('riesgo' in p.text.lower() and 'matriz' in p.text.lower() for p in doc.paragraphs),
        'proyecciones_financieras': any('proyección' in p.text.lower() and ('2026' in p.text or '2027' in p.text) for p in doc.paragraphs),
        'equipo_detallado': any('equipo' in p.text.lower() and ('experiencia' in p.text.lower() or 'habilidad' in p.text.lower()) for p in doc.paragraphs),
        'aspectos_legales': any('legal' in p.text.lower() and ('propiedad' in p.text.lower() or 'intelectual' in p.text.lower()) for p in doc.paragraphs)
    }
    
    print("Elementos clave presentes:")
    for elemento, presente in elementos_clave.items():
        estado = "✅" if presente else "❌"
        print(f"  {estado} {elemento.replace('_', ' ').title()}")
    
    # Análisis de fortalezas y debilidades
    print("\n4. FORTALEZAS Y DEBILIDADES:")
    print("-"*40)
    
    fortalezas = []
    debilidades = []
    
    # Fortalezas basadas en análisis
    if elementos_clave['algoritmo_detallado']:
        fortalezas.append("Algoritmo detallado y bien explicado")
    if elementos_clave['plan_marketing']:
        fortalezas.append("Plan de marketing completo (5 fases, tácticas específicas)")
    if elementos_clave['presupuesto_detallado']:
        fortalezas.append("Presupuesto detallado con ROI proyectado")
    if elementos_clave['analisis_riesgos']:
        fortalezas.append("Análisis de riesgos estructurado con matriz")
    
    # Debilidades basadas en análisis
    if secciones_contenido['legal'] < 5:
        debilidades.append("Sección legal poco desarrollada")
    if secciones_contenido['anexos'] < 3:
        debilidades.append("Faltan anexos (wireframes, estudios de mercado)")
    if not elementos_clave['aspectos_legales']:
        debilidades.append("Aspectos legales (propiedad intelectual) no detallados")
    
    print("FORTALEZAS:")
    for f in fortalezas:
        print(f"  ✅ {f}")
    
    print("\nDEBILIDADES:")
    for d in debilidades:
        print(f"  ❌ {d}")
    
    # Recomendaciones
    print("\n5. RECOMENDACIONES:")
    print("-"*40)
    
    recomendaciones = []
    
    if 'Sección legal poco desarrollada' in debilidades:
        recomendaciones.append("Aplicar skill 'legal' para fortalecer aspectos jurídicos")
    
    if 'Faltan anexos (wireframes, estudios de mercado)' in debilidades:
        recomendaciones.append("Aplicar skill 'frontend-design' para crear wireframes")
        recomendaciones.append("Añadir estudios de mercado detallados en anexos")
    
    if 'Aspectos legales (propiedad intelectual) no detallados' in debilidades:
        recomendaciones.append("Detallar estrategia de propiedad intelectual y patentes")
    
    # Verificar modelo de negocio
    if secciones_contenido['modelo_negocio'] < 10:
        recomendaciones.append("Aplicar skill 'business-model-canvas' para refinar modelo")
    
    for i, rec in enumerate(recomendaciones, 1):
        print(f"  {i}. {rec}")
    
    # Resumen ejecutivo
    print("\n6. RESUMEN EJECUTIVO:")
    print("-"*40)
    
    calidad_general = (sum(1 for p in elementos_clave.values() if p) / len(elementos_clave)) * 100
    completitud = (sum(secciones_contenido.values()) / len(doc.paragraphs)) * 100
    
    print(f"Calidad general: {calidad_general:.1f}%")
    print(f"Completitud: {completitud:.1f}%")
    print(f"Fortalezas identificadas: {len(fortalezas)}")
    print(f"Debilidades identificadas: {len(debilidades)}")
    
    if calidad_general > 70 and len(debilidades) < 3:
        print("\n🎯 VEREDICTO: PLAN DE NEGOCIO SÓLIDO, LISTO PARA MEJORAS INCREMENTALES")
    elif calidad_general > 50:
        print("\n⚠️ VEREDICTO: PLAN DE NEGOCIO ACEPTABLE, NECESITA MEJORAS SIGNIFICATIVAS")
    else:
        print("\n❌ VEREDICTO: PLAN DE NEGOCIO INCOMPLETO, REQUIERE TRABAJO SUSTANCIAL")
    
    return {
        'parrafos_totales': len(doc.paragraphs),
        'secciones_principales': len(secciones_principales),
        'calidad_general': calidad_general,
        'completitud': completitud,
        'fortalezas': fortalezas,
        'debilidades': debilidades,
        'recomendaciones': recomendaciones
    }

if __name__ == '__main__':
    analizar_documento_completo()