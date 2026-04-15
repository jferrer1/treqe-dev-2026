import docx
from collections import defaultdict

def analyze_document():
    doc = docx.Document('Plan_Negocio_Treqe_CON_SOLUCIONES_2026.docx')
    
    print('=== ANALISIS DEL DOCUMENTO ===')
    print(f'Total parrafos: {len(doc.paragraphs)}')
    
    # Contar parrafos tachados
    strikethrough_count = 0
    sections_with_strikethrough = []
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
            
        # Check for strikethrough
        has_strikethrough = False
        for run in para.runs:
            if run.font.strike:
                has_strikethrough = True
                break
        
        if has_strikethrough:
            strikethrough_count += 1
            # Get context (previous heading)
            heading = 'Sin titulo'
            for j in range(i-1, max(-1, i-10), -1):
                if doc.paragraphs[j].style.name.startswith('Heading'):
                    heading = doc.paragraphs[j].text.strip()
                    break
            
            sections_with_strikethrough.append({
                'index': i,
                'heading': heading,
                'text': text[:100] + '...' if len(text) > 100 else text
            })
    
    print(f'Parrafos tachados: {strikethrough_count}')
    print()
    
    # Group by heading
    grouped = defaultdict(list)
    for section in sections_with_strikethrough:
        grouped[section['heading']].append(section)
    
    print('=== SECCIONES CON TEXTO TACHADO ===')
    for heading, sections in grouped.items():
        print(f'\n**{heading}** ({len(sections)} secciones):')
        for section in sections[:3]:
            print(f'  - {section["text"]}')
        if len(sections) > 3:
            print(f'  ... y {len(sections)-3} mas')
    
    # Identify topics
    print('\n=== TEMAS IDENTIFICADOS ===')
    topics = {
        'chat_grupal': ['chat', 'grupo', 'conversacion', 'negociacion', 'oficial', 'estructurada'],
        'devoluciones': ['devolucion', 'desistimiento', '24 horas', 'garantia', 'fondo'],
        'envios': ['envio', 'transporte', 'logistica', 'correos', 'paquete', 'entrega'],
        'garantias': ['garantia', 'seguro', 'escrow', 'deposito', 'compensacion'],
        'algoritmo': ['algoritmo', 'matching', 'k=', 'ciclo', 'rueda', 'intercambio'],
        'reputacion': ['reputacion', 'scoring', 'puntuacion', 'confianza', 'historial']
    }
    
    topic_sections = defaultdict(list)
    for section in sections_with_strikethrough:
        text_lower = section['text'].lower()
        topic_found = 'otros'
        
        for topic, keywords in topics.items():
            for keyword in keywords:
                if keyword in text_lower:
                    topic_found = topic
                    break
            if topic_found != 'otros':
                break
        
        topic_sections[topic_found].append(section)
    
    for topic, sections in topic_sections.items():
        if sections:
            print(f'{topic}: {len(sections)} secciones')
    
    return doc, sections_with_strikethrough, topic_sections

if __name__ == '__main__':
    doc, strikethrough_sections, topic_sections = analyze_document()