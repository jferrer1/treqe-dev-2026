#!/usr/bin/env python3
"""
Actualizar la sección de marketing en el documento con la versión mejorada
"""

from docx import Document
import re

def actualizar_marketing():
    print("ACTUALIZANDO SECCIÓN DE MARKETING CON VERSIÓN MEJORADA...")
    
    # Cargar documento
    doc = Document('Plan_Negocio_Treqe_HUMANIZADO_ACTUALIZADO.docx')
    print(f"Documento cargado: {len(doc.paragraphs)} párrafos")
    
    # Buscar sección 10: ESTRATEGIA DE MARKETING
    print("\nBuscando sección de marketing...")
    
    inicio_marketing = -1
    fin_marketing = -1
    
    for i, para in enumerate(doc.paragraphs):
        texto = para.text.strip()
        
        # Buscar inicio de sección 10
        if "10. ESTRATEGIA DE MARKETING" in texto or "10. ESTRATEGIA DE MARKETING DETALLADA" in texto:
            inicio_marketing = i
            print(f"  Encontrado inicio en párrafo {i}: {texto}")
            
            # Buscar fin de sección (inicio de sección 11)
            for j in range(i + 1, len(doc.paragraphs)):
                texto_j = doc.paragraphs[j].text.strip()
                if texto_j.startswith("11.") or "11. ASPECTOS LEGALES" in texto_j:
                    fin_marketing = j
                    print(f"  Encontrado fin en párrafo {j}: {texto_j}")
                    break
            
            if fin_marketing == -1:
                fin_marketing = len(doc.paragraphs) - 1
            break
    
    if inicio_marketing == -1:
        print("ERROR: No se encontró la sección de marketing")
        return False
    
    print(f"  Sección marketing: párrafos {inicio_marketing} a {fin_marketing}")
    
    # Leer contenido mejorado de marketing
    with open('seccion_marketing_mejorada.md', 'r', encoding='utf-8') as f:
        contenido_mejorado = f.read()
    
    # Dividir en párrafos
    lineas = contenido_mejorado.split('\n')
    
    # Eliminar contenido antiguo de marketing
    print(f"  Eliminando {fin_marketing - inicio_marketing - 1} párrafos antiguos...")
    
    # Mantener el título de la sección (inicio_marketing)
    # Eliminar desde inicio_marketing + 1 hasta fin_marketing - 1
    for j in range(fin_marketing - 1, inicio_marketing, -1):
        if j < len(doc.paragraphs):
            # Eliminar párrafo
            p = doc._body._element
            p.remove(doc.paragraphs[j]._element)
    
    # Insertar nuevo contenido
    print("  Insertando contenido mejorado...")
    
    # Insertar después del título
    posicion_insercion = inicio_marketing + 1
    
    for linea in lineas:
        if linea.strip():  # Solo insertar líneas no vacías
            # Insertar párrafo
            nuevo_parrafo = doc.paragraphs[posicion_insercion].insert_paragraph_before(linea)
            # Mantener estilo similar
            if linea.startswith('##') or linea.startswith('###'):
                for run in nuevo_parrafo.runs:
                    run.bold = True
    
    # Guardar documento actualizado
    output_path = 'Plan_Negocio_Treqe_MARKETING_MEJORADO.docx'
    doc.save(output_path)
    
    print(f"\n✅ Documento guardado: {output_path}")
    print("✅ Sección de marketing actualizada con:")
    print("   - Estrategia de 5 fases mejorada")
    print("   - Psicología de marketing aplicada")
    print("   - Tácticas específicas para Treqe")
    print("   - Optimización de conversión detallada")
    print("   - Estrategia de retención robusta")
    print("   - Presupuesto y cronograma detallados")
    
    return output_path

if __name__ == '__main__':
    actualizar_marketing()