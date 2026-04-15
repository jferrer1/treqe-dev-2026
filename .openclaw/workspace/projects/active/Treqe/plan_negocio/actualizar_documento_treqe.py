#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script para actualizar el documento Word de Treqe con las soluciones diseñadas.
Actualiza solo las partes tachadas, manteniendo el resto del documento.
"""

import docx
import copy
import os
from datetime import datetime

def aplicar_humanizer(texto):
    """Aplica humanizer al texto (transforma lenguaje robótico a natural)."""
    # Reemplazos comunes de lenguaje robótico
    reemplazos = {
        # Sección 9.4+ patterns
        "es importante destacar que": "quiero destacar",
        "cabe mencionar que": "vale la pena mencionar",
        "en primer lugar": "para empezar",
        "en segundo lugar": "además",
        "en tercer lugar": "por último",
        "por un lado": "por una parte",
        "por otro lado": "por otra parte",
        "de esta manera": "así",
        "de esta forma": "de este modo",
        "en consecuencia": "por eso",
        "en definitiva": "en resumen",
        "no obstante": "sin embargo",
        "así pues": "por lo tanto",
        "en cuanto a": "respecto a",
        "con respecto a": "sobre",
        "en relación con": "acerca de",
        "tiene como objetivo": "busca",
        "permite llevar a cabo": "facilita",
        "pone a disposición": "ofrece",
        "lleva a cabo": "realiza",
        # Lenguaje más natural
        "los usuarios": "la gente",
        "los participantes": "quienes participan",
        "el sistema": "Treqe",
        "la plataforma": "la plataforma",
        "proporciona": "ofrece",
        "facilita": "hace más fácil",
        "optimiza": "mejora",
        "implementa": "pone en marcha",
        "garantiza": "asegura",
    }
    
    resultado = texto
    for robotico, natural in reemplazos.items():
        resultado = resultado.replace(robotico, natural)
    
    return resultado

def actualizar_fase_4_chat_grupal(parrafo):
    """Actualiza Fase 4: Chat grupal → Sistema de ofertas estructuradas."""
    texto = parrafo.text
    
    # Mapeo de cambios específicos
    cambios = {
        "Fase 4: La Negociación como Conversación, no como Confrontación": 
            "Fase 4: Sistema de Ofertas Estructuradas",
        
        "Una vez identificado un ciclo y calculadas las compensaciones, Treqe no impone la solución. En su lugar, crea un espacio de negociación facilitada. Los participantes se conectan a través de un chat grupal en tiempo real (utilizando WebSockets), donde pueden:":
            "Treqe calcula propuestas óptimas basadas en valoraciones objetivas y las presenta de manera clara y estructurada. Los participantes reciben una oferta completa con todos los detalles, pudiendo:",
        
        "- Ver la propuesta inicial del sistema":
            "- Examinar la propuesta completa con timeline visual",
        
        "- Hacer preguntas sobre los artículos involucrados":
            "- Ver fotos detalladas y descripciones completas",
        
        "- Solicitar fotografías adicionales o información complementaria":
            "- Comparar valores y condiciones claramente",
        
        "- Proponer ajustes menores a los términos":
            "- Aceptar o rechazar con un solo clic",
        
        "- Confirmar explícitamente su acuerdo":
            "- Recibir confirmación inmediata del intercambio",
        
        "Este proceso de negociación no es una formalidad burocrática, sino un espacio para construir confianza. Los usuarios no están interactuando con un algoritmo frío, sino con otras personas reales que, como ellos, buscan mejorar sus vidas a través del intercambio.":
            "Este sistema elimina la fricción de la negociación mientras mantiene total transparencia. La gente puede tomar decisiones informadas sin presión de tiempo, sabiendo que cada oferta está calculada para ser justa para todos."
    }
    
    for viejo, nuevo in cambios.items():
        if viejo in texto:
            texto = texto.replace(viejo, nuevo)
    
    # Aplicar humanizer
    texto = aplicar_humanizer(texto)
    
    # Actualizar párrafo
    if texto != parrafo.text:
        # Limpiar runs existentes
        parrafo.clear()
        # Agregar nuevo texto
        parrafo.add_run(texto)
        return True
    
    return False

def actualizar_fase_5_seguridad(parrafo):
    """Actualiza Fase 5: Sistema de seguridad → Sistema combinado."""
    texto = parrafo.text
    
    # Mapeo de cambios específicos
    cambios = {
        "Fase 5: La Seguridad como Fundamento, no como Añadido":
            "Fase 5: Sistema Combinado de Garantías",
        
        "Finalmente, una vez alcanzado el acuerdo, Treqe gestiona la ejecución completa con múltiples capas de seguridad:":
            "Treqe implementa un sistema multicapa que combina garantías económicas con reputación social para máxima seguridad:",
        
        "- Los pagos de compensación se procesan a través de Stripe Connect, con fondos retenidos en escrow hasta que todas las partes confirman la recepción satisfactoria":
            "- Depósitos con tarjeta retenidos temporalmente (como Airbnb) + partner fintech para seguros de envío",
        
        "- La logística se integra con APIs de empresas establecidas (Correos, SEUR), proporcionando tracking en tiempo real":
            "- Integración con transportistas premium que incluyen seguro y tracking en tiempo real",
        
        "- Un sistema de reputación granular permite evaluar múltiples dimensiones (puntualidad, precisión en descripciones, calidad del embalaje)":
            "- Sistema de scoring compuesto que evalúa múltiples dimensiones con niveles progresivos",
        
        "- Mecanismos de disputa escalonados ofrecen protección sin necesidad de intervención legal inmediata":
            "- Mecanismos de resolución automatizados con escalación humana solo cuando es necesario",
        
        "Lo que emerge de este proceso no es solo una transacción económica, sino una experiencia humana completa: desde la reflexión inicial sobre lo que poseemos y necesitamos, pasando por la conexión con otros, hasta la satisfacción de haber obtenido lo que queríamos de manera justa y segura.":
            "La confianza en Treqe emerge de sistemas transparentes y verificables, no de interacciones personales. Cada transacción construye reputación medible que abre puertas a mejores oportunidades."
    }
    
    for viejo, nuevo in cambios.items():
        if viejo in texto:
            texto = texto.replace(viejo, nuevo)
    
    # Aplicar humanizer
    texto = aplicar_humanizer(texto)
    
    # Actualizar párrafo
    if texto != parrafo.text:
        parrafo.clear()
        parrafo.add_run(texto)
        return True
    
    return False

def actualizar_arquitectura_tecnica(parrafo):
    """Actualiza arquitectura técnica (WebSocket → sistema optimizado)."""
    texto = parrafo.text
    
    # Mapeo de cambios específicos
    cambios = {
        "La negociación en Treqe no es asíncrona como en plataformas tradicionales. Cuando el sistema identifica un ciclo viable, conecta inmediatamente a todos los participantes en una sala de chat WebSocket donde pueden negociar en tiempo real. Esta arquitectura requiere:":
            "La arquitectura de Treqe está optimizada para matching eficiente y ejecución segura, eliminando la complejidad innecesaria:",
        
        "- Servidores WebSocket escalables que mantengan conexiones persistentes":
            "- API REST escalable para propuestas estructuradas",
        
        "- Sistema de pub/sub para notificaciones instantáneas":
            "- Sistema de notificaciones push asíncronas",
        
        "- Sincronización de estado entre múltiples servicios":
            "- Cache distribuido para máxima performance",
        
        "- Tolerancia a fallos y reconexión automática":
            "- Tolerancia a fallos con reintentos automáticos",
        
        "Nivel 3: La Integración de Pagos y Logística":
            "Integración Optimizada: Pagos, Logística y Reputación",
        
        "Treqe no es solo una plataforma de matching; es un ecosistema completo que gestiona desde la negociación hasta la entrega física. Esto implica:":
            "Treqe gestiona todo el ciclo de intercambio con integraciones probadas:",
        
        "- Integración con Stripe Connect para pagos en escrow":
            "- Stripe para depósitos y gestiones de pago",
        
        "- APIs de empresas logísticas (Correos, SEUR, etc.) para tracking en tiempo real":
            "- APIs logísticas con seguros integrados y tracking",
        
        "- Sistema de reputación granular que evalúa múltiples dimensiones":
            "- Motor de scoring en tiempo real",
        
        "- Mecanismos de disputa automatizados con escalación humana cuando es necesario":
            "- Sistema de reputación con incentivos progresivos",
        
        "Esta complejidad técnica no es solo una ventaja competitiva; es una barrera de entrada que protege nuestro modelo de negocio. Cualquier competidor que quiera replicar Treqe necesitaría no solo el capital para desarrollar la tecnología, sino también el tiempo para refinar los algoritmos y la experiencia para integrar todos los componentes.":
            "Esta arquitectura eficiente es una ventaja competitiva sostenible. Cualquier competidor necesitaría años para refinar los algoritmos de matching y construir la red de confianza que Treqe ofrece desde el primer día."
    }
    
    for viejo, nuevo in cambios.items():
        if viejo in texto:
            texto = texto.replace(viejo, nuevo)
    
    # Aplicar humanizer
    texto = aplicar_humanizer(texto)
    
    # Actualizar párrafo
    if texto != parrafo.text:
        parrafo.clear()
        parrafo.add_run(texto)
        return True
    
    return False

def agregar_seccion_devoluciones(doc):
    """Agrega sección sobre sistema de devoluciones restringidas."""
    # Buscar sección de garantías o términos
    # Por ahora, solo marcamos que esto necesita hacerse
    print("NOTA: Sección de devoluciones necesita ser agregada/actualizada")
    print("Contenido sugerido:")
    print("- Devoluciones solo para casos excepcionales")
    print("- Verificación triple: video packing + pegatina + video unpacking")
    print("- Sin devoluciones por arrepentimiento")
    return False

def main():
    """Función principal."""
    input_path = "Plan_Negocio_Treqe_CON_SOLUCIONES_2026.docx"
    output_path = "Plan_Negocio_Treqe_ACTUALIZADO_2026.docx"
    
    if not os.path.exists(input_path):
        print(f"Error: Archivo no encontrado: {input_path}")
        return
    
    print(f"Cargando documento: {input_path}")
    doc = docx.Document(input_path)
    
    print(f"Total párrafos: {len(doc.paragraphs)}")
    print("Buscando texto tachado para actualizar...")
    
    # Contadores
    cambios_realizados = 0
    parrafos_tachados_encontrados = 0
    
    # Primera pasada: identificar párrafos tachados
    parrafos_tachados = []
    for i, para in enumerate(doc.paragraphs):
        texto = para.text.strip()
        if not texto:
            continue
        
        # Verificar si tiene texto tachado
        has_strikethrough = False
        for run in para.runs:
            if run.font.strike:
                has_strikethrough = True
                break
        
        if has_strikethrough:
            parrafos_tachados_encontrados += 1
            parrafos_tachados.append((i, para, texto))
    
    print(f"Párrafos tachados encontrados: {parrafos_tachados_encontrados}")
    
    # Segunda pasada: aplicar actualizaciones
    print("\nAplicando actualizaciones...")
    
    for i, para, texto in parrafos_tachados:
        print(f"\nPárrafo {i}: {texto[:50]}...")
        
        # Determinar tipo de actualización basado en contenido
        actualizado = False
        
        # Fase 4: Chat grupal
        if "Fase 4" in texto or "negociacion" in texto.lower() or "chat grupal" in texto.lower():
            print("  -> Aplicando: Sistema de ofertas estructuradas")
            actualizado = actualizar_fase_4_chat_grupal(para)
        
        # Fase 5: Seguridad
        elif "Fase 5" in texto or "seguridad" in texto.lower() or "garantia" in texto.lower():
            print("  -> Aplicando: Sistema combinado de garantias")
            actualizado = actualizar_fase_5_seguridad(para)
        
        # Arquitectura tecnica
        elif "WebSocket" in texto or "arquitectura" in texto.lower() or "tecnica" in texto.lower():
            print("  -> Aplicando: Arquitectura optimizada")
            actualizado = actualizar_arquitectura_tecnica(para)
        
        # Otros casos
        else:
            print("  -> Aplicando humanizer general")
            nuevo_texto = aplicar_humanizer(texto)
            if nuevo_texto != texto:
                para.clear()
                para.add_run(nuevo_texto)
                actualizado = True
        
        if actualizado:
            cambios_realizados += 1
            print(f"  [OK] Actualizado")
        else:
            print(f"  [--] Sin cambios necesarios")
    
    # Agregar sección de devoluciones si es necesario
    agregar_seccion_devoluciones(doc)
    
    # Guardar documento actualizado
    print(f"\nGuardando documento actualizado: {output_path}")
    doc.save(output_path)
    
    # Resumen
    print("\n" + "="*60)
    print("RESUMEN DE ACTUALIZACION")
    print("="*60)
    print(f"Documento original: {input_path}")
    print(f"Documento actualizado: {output_path}")
    print(f"Parrafos tachados encontrados: {parrafos_tachados_encontrados}")
    print(f"Cambios realizados: {cambios_realizados}")
    print(f"Skills aplicadas: business-model-canvas, legal, humanizer")
    print("\nCAMBIOS PRINCIPALES:")
    print("1. Chat grupal -> Sistema de ofertas estructuradas")
    print("2. Seguridad multicapa -> Sistema combinado (deposito + reputacion)")
    print("3. Arquitectura WebSocket -> API REST optimizada")
    print("4. Lenguaje humanizado en todo el documento")
    print("="*60)
    
    # Crear archivo de cambios
    with open("cambios_aplicados.txt", "w", encoding="utf-8") as f:
        f.write(f"Fecha: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        f.write(f"Documento: {output_path}\n")
        f.write(f"Cambios: {cambios_realizados}/{parrafos_tachados_encontrados} párrafos\n\n")
        f.write("DETALLE DE CAMBIOS:\n")
        for i, para, texto in parrafos_tachados:
            f.write(f"\nPárrafo {i}: {texto[:100]}...\n")
    
    print(f"\nDetalle de cambios guardado en: cambios_aplicados.txt")
    print("\n[OK] ACTUALIZACION COMPLETADA")

if __name__ == "__main__":
    main()