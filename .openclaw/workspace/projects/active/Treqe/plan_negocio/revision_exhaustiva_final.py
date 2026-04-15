#!/usr/bin/env python3
"""
Revisión exhaustiva del documento final para verificar:
1. Humanizer aplicado correctamente
2. TODO el contenido incluido
3. Es el más completo hasta ahora
"""

from docx import Document
import re

def verificar_humanizer_aplicado(doc):
    """Verificar que humanizer está aplicado correctamente"""
    print("\n1. VERIFICANDO HUMANIZER APLICADO:")
    
    # Patrones de humanizer que deberían estar aplicados
    patrones_humanizer = [
        (r'\bsinergia\b', 'debería ser "colaboración"'),
        (r'\bparadigma\b', 'debería ser "forma de hacer las cosas"'),
        (r'\becosistema\b', 'debería ser "conjunto"'),
        (r'\boptimización\b', 'debería ser "mejora"'),
        (r'\brobusta\b', 'debería ser "sólida"'),
    ]
    
    problemas_humanizer = []
    
    for i, para in enumerate(doc.paragraphs):
        texto = para.text
        for patron, mensaje in patrones_humanizer:
            if re.search(patron, texto, re.IGNORECASE):
                problemas_humanizer.append((i, texto[:100], mensaje))
    
    if problemas_humanizer:
        print(f"  ⚠️ PROBLEMAS ENCONTRADOS: {len(problemas_humanizer)}")
        for j, (num, texto, mensaje) in enumerate(problemas_humanizer[:3]):
            print(f"    Párrafo {num}: {texto}...")
            print(f"    {mensaje}")
    else:
        print("  ✅ HUMANIZER APLICADO CORRECTAMENTE")
    
    return len(problemas_humanizer)

def verificar_contenido_completo(doc_final, doc_original):
    """Verificar que TODO el contenido del original está incluido"""
    print("\n2. VERIFICANDO CONTENIDO COMPLETO:")
    
    # Contar párrafos
    print(f"  Original: {len(doc_original.paragraphs)} párrafos")
    print(f"  Final: {len(doc_final.paragraphs)} párrafos")
    
    # Verificar contenido clave del marketing mejorado
    contenido_clave = [
        "SEO Programático",
        "5-Phase Launch Strategy",
        "Psychology & Mental Models",
        "Conversion Optimization",
        "Retention & LTV Strategy",
        "CAC por canal",
        "LTV:CAC",
        "presupuesto trimestral",
        "A/B testing",
        "funnel optimization"
    ]
    
    contenido_faltante = []
    
    for clave in contenido_clave:
        encontrado_original = any(clave.lower() in p.text.lower() for p in doc_original.paragraphs)
        encontrado_final = any(clave.lower() in p.text.lower() for p in doc_final.paragraphs)
        
        if encontrado_original and not encontrado_final:
            contenido_faltante.append(clave)
    
    if contenido_faltante:
        print(f"  ⚠️ CONTENIDO FALTANTE: {len(contenido_faltante)} elementos")
        for clave in contenido_faltante[:5]:
            print(f"    - {clave}")
    else:
        print("  ✅ TODO EL CONTENIDO CLAVE ESTÁ INCLUIDO")
    
    return len(contenido_faltante)

def comparar_con_otros_documentos(doc_final):
    """Comparar con otros documentos para verificar que es el más completo"""
    print("\n3. COMPARANDO CON OTROS DOCUMENTOS:")
    
    documentos = [
        ('Plan_Negocio_Treqe_HUMANIZADO_ACTUALIZADO.docx', 'Humanizado Actualizado (10:28)'),
        ('Plan_Negocio_Treqe_MARKETING_MEJORADO.docx', 'Marketing Mejorado (base)'),
        ('Plan_Negocio_Treqe_MARKETING_MEJORADO_HUMANIZADO.docx', 'Marketing Mejorado Humanizado'),
        ('Plan_Negocio_Treqe_10h28_HUMANIZADO_SELECTIVO.docx', '10:28 Humanizado Selectivo'),
    ]
    
    resultados = []
    
    for archivo, nombre in documentos:
        try:
            doc = Document(archivo)
            resultados.append((nombre, len(doc.paragraphs), archivo))
        except:
            pass
    
    # Ordenar por número de párrafos (más completo)
    resultados.sort(key=lambda x: x[1], reverse=True)
    
    print("  Ranking de completitud:")
    for i, (nombre, parrafos, archivo) in enumerate(resultados):
        indicador = "✅" if archivo == 'Plan_Negocio_Treqe_FINAL_CORREGIDO.docx' else "  "
        print(f"    {indicador} {i+1}. {nombre}: {parrafos} párrafos")
    
    # Verificar si el final es el más completo
    if resultados[0][2] == 'Plan_Negocio_Treqe_FINAL_CORREGIDO.docx':
        print("  ✅ ES EL DOCUMENTO MÁS COMPLETO")
        return True
    else:
        print(f"  ⚠️ NO ES EL MÁS COMPLETO. El más completo es: {resultados[0][0]}")
        return False

def verificar_estructura_profesional(doc):
    """Verificar que mantiene estructura profesional"""
    print("\n4. VERIFICANDO ESTRUCTURA PROFESIONAL:")
    
    secciones_esperadas = [
        "INTRODUCCIÓN",
        "PROBLEMA NO RESUELTO",
        "SOLUCIÓN TREQE",
        "VENTAJA COMPETITIVA",
        "MODELO DE NEGOCIO",
        "PROYECCIONES FINANCIERAS",
        "EQUIPO Y PLAN DE EJECUCIÓN",
        "ANÁLISIS DE RIESGOS",
        "CONCLUSIONES",
        "ESTRATEGIA DE MARKETING",
        "ASPECTOS LEGALES",
        "ANEXOS"
    ]
    
    secciones_encontradas = []
    
    for para in doc.paragraphs:
        texto = para.text.upper()
        for seccion in secciones_esperadas:
            if seccion in texto and len(texto) < 100:  # Probablemente es título
                secciones_encontradas.append(seccion)
                break
    
    print(f"  Secciones esperadas: {len(secciones_esperadas)}")
    print(f"  Secciones encontradas: {len(secciones_encontradas)}")
    
    faltantes = [s for s in secciones_esperadas if s not in secciones_encontradas]
    
    if faltantes:
        print(f"  ⚠️ SECCIONES FALTANTES: {len(faltantes)}")
        for seccion in faltantes[:3]:
            print(f"    - {seccion}")
    else:
        print("  ✅ TODAS LAS SECCIONES PROFESIONALES PRESENTES")
    
    return len(faltantes)

def crear_informe_final(doc_final, problemas_humanizer, contenido_faltante, es_mas_completo, secciones_faltantes):
    """Crear informe final detallado"""
    
    informe = f"""# INFORME DE REVISIÓN EXHAUSTIVA
## Documento: Plan_Negocio_Treqe_FINAL_CORREGIDO.docx

### 📊 ESTADÍSTICAS GENERALES:
- **Párrafos totales:** {len(doc_final.paragraphs)}
- **Problemas de humanizer:** {problemas_humanizer}
- **Contenido faltante:** {contenido_faltante} elementos
- **Es el más completo:** {'SÍ' if es_mas_completo else 'NO'}
- **Secciones faltantes:** {secciones_faltantes}

### 🎯 VERIFICACIÓN HUMANIZER:
{"✅ APLICADO CORRECTAMENTE" if problemas_humanizer == 0 else f"⚠️ {problemas_humanizer} PROBLEMAS DETECTADOS"}

### 📚 VERIFICACIÓN CONTENIDO COMPLETO:
{"✅ TODO EL CONTENIDO INCLUIDO" if contenido_faltante == 0 else f"⚠️ {contenido_faltante} ELEMENTOS FALTANTES"}

### 🏆 VERIFICACIÓN COMPLETITUD:
{"✅ ES EL DOCUMENTO MÁS COMPLETO" if es_mas_completo else "⚠️ NO ES EL MÁS COMPLETO"}

### 🏢 VERIFICACIÓN ESTRUCTURA PROFESIONAL:
{"✅ ESTRUCTURA COMPLETA" if secciones_faltantes == 0 else f"⚠️ {secciones_faltantes} SECCIONES FALTANTES"}

### 🔍 CONTENIDO CLAVE VERIFICADO:

#### MARKETING MEJORADO (skill `marketing-mode`):
- [{"✓" if any("SEO Programático" in p.text for p in doc_final.paragraphs) else "✗"}] SEO Programático
- [{"✓" if any("5-Phase" in p.text for p in doc_final.paragraphs) else "✗"}] 5-Phase Launch Strategy
- [{"✓" if any("Psychology" in p.text for p in doc_final.paragraphs) else "✗"}] Psychology & Mental Models
- [{"✓" if any("Conversion Optimization" in p.text for p in doc_final.paragraphs) else "✗"}] Conversion Optimization
- [{"✓" if any("Retention" in p.text for p in doc_final.paragraphs) else "✗"}] Retention & LTV Strategy

#### CONTENIDO TÉCNICO:
- [{"✓" if any("algoritmo" in p.text.lower() for p in doc_final.paragraphs) else "✗"}] Algoritmos y tecnología
- [{"✓" if any("métricas" in p.text.lower() for p in doc_final.paragraphs) else "✗"}] Métricas y KPIs
- [{"✓" if any("presupuesto" in p.text.lower() for p in doc_final.paragraphs) else "✗"}] Presupuesto detallado
- [{"✓" if any("proyecciones" in p.text.lower() for p in doc_final.paragraphs) else "✗"}] Proyecciones financieras

### 📋 RECOMENDACIONES:

{"🎉 DOCUMENTO LISTO PARA USO - Cumple todos los criterios" 
 if problemas_humanizer == 0 and contenido_faltante == 0 and es_mas_completo and secciones_faltantes == 0
 else "🔧 SE REQUIEREN AJUSTES - Ver detalles arriba"}

### 🚀 PRÓXIMOS PASOS:

1. {"Usar directamente para presentaciones/inversores" 
    if problemas_humanizer == 0 and contenido_faltante == 0 and es_mas_completo and secciones_faltantes == 0
    else "Corregir los problemas identificados arriba"}

2. Realizar prueba final de lectura
3. Preparar versión ejecutiva (1-2 páginas)
4. Preparar presentación de inversores

---
**Fecha de revisión:** 26 de febrero 2026, 20:20
**Revisor:** Asistente OpenClaw
**Documento revisado:** Plan_Negocio_Treqe_FINAL_CORREGIDO.docx
"""
    
    with open('INFORME_REVISION_EXHAUSTIVA.md', 'w', encoding='utf-8') as f:
        f.write(informe)
    
    print(f"\n✅ INFORME CREADO: INFORME_REVISION_EXHAUSTIVA.md")
    
    return informe

def main():
    print("REVISIÓN EXHAUSTIVA DEL DOCUMENTO FINAL")
    print("=" * 50)
    
    # Cargar documentos
    try:
        doc_final = Document('Plan_Negocio_Treqe_FINAL_CORREGIDO.docx')
        doc_original = Document('Plan_Negocio_Treqe_MARKETING_MEJORADO.docx')
    except Exception as e:
        print(f"Error cargando documentos: {e}")
        return
    
    # Ejecutar verificaciones
    problemas_humanizer = verificar_humanizer_aplicado(doc_final)
    contenido_faltante = verificar_contenido_completo(doc_final, doc_original)
    es_mas_completo = comparar_con_otros_documentos(doc_final)
    secciones_faltantes = verificar_estructura_profesional(doc_final)
    
    # Crear informe final
    crear_informe_final(doc_final, problemas_humanizer, contenido_faltante, es_mas_completo, secciones_faltantes)
    
    print("\n" + "=" * 50)
    print("REVISIÓN COMPLETADA")
    
    # Resumen ejecutivo
    if problemas_humanizer == 0 and contenido_faltante == 0 and es_mas_completo and secciones_faltantes == 0:
        print("🎉 ¡DOCUMENTO APROBADO! Cumple todos los criterios.")
    else:
        print("⚠️ DOCUMENTO REQUIERE AJUSTES. Ver informe para detalles.")

if __name__ == '__main__':
    main()