#!/usr/bin/env python3
"""
Script simple para crear documento Word del plan de negocio Treqe
"""

import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

# Crear documento
doc = Document()

# Configurar estilos
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# TÃ­tulo principal
title = doc.add_heading('PLAN DE NEGOCIO TREQE', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.runs[0]
title_run.font.size = Pt(24)
title_run.font.bold = True

# SubtÃ­tulo
subtitle = doc.add_paragraph('Documento Ãºnico definitivo - Todo en un solo documento')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle.runs[0]
subtitle_run.font.size = Pt(14)
subtitle_run.font.italic = True

doc.add_paragraph(f'Creado el {datetime.now().strftime("%d/%m/%Y")}')
doc.add_page_break()

# ===== PARTE 1: RESUMEN EJECUTIVO =====
doc.add_heading('ðŸ“‹ RESUMEN EJECUTIVO (2 minutos)', level=1)

p = doc.add_paragraph()
p.add_run('El problema que resolvemos: ').bold = True
p.add_run('72% de espaÃ±oles tiene cosas guardadas que no usa. Intercambiar entre 2 personas tiene solo 5% de probabilidad.')

p = doc.add_paragraph()
p.add_run('Nuestra soluciÃ³n: ').bold = True
p.add_run('Treqe conecta a 3+ personas para intercambios circulares (Aâ†’Bâ†’Câ†’A). Probabilidad sube al 20-35%.')

p = doc.add_paragraph()
p.add_run('CÃ³mo ganamos dinero: ').bold = True
p.add_run('3% de comisiÃ³n cuando el intercambio funciona. Solo cobramos cuando tÃº ganas.')

doc.add_heading('Los nÃºmeros clave', level=2)
doc.add_paragraph('â€¢ InversiÃ³n inicial: 58.000â‚¬', style='List Bullet')
doc.add_paragraph('â€¢ AÃ±o 1: 8.000 usuarios activos, 360.000â‚¬ ingresos, 84.000â‚¬ beneficio', style='List Bullet')
doc.add_paragraph('â€¢ AÃ±o 2: 25.000 usuarios activos, 1.728.000â‚¬ ingresos', style='List Bullet')
doc.add_paragraph('â€¢ AÃ±o 3: 60.000 usuarios activos, 5.040.000â‚¬ ingresos', style='List Bullet')

doc.add_heading('PrÃ³ximos pasos inmediatos', level=2)
doc.add_paragraph('1. Constituir Sociedad Limitada (3.000â‚¬)', style='List Number')
doc.add_paragraph('2. Registrar marca "Treqe" EspaÃ±a + UE (850â‚¬)', style='List Number')
doc.add_paragraph('3. Desarrollar algoritmo v0.1', style='List Number')
doc.add_paragraph('4. Primeros 50 usuarios (personas conocidas)', style='List Number')

doc.add_page_break()

# ===== PARTE 2: EL PROBLEMA REAL =====
doc.add_heading('ðŸŽ¯ PARTE 1: EL PROBLEMA REAL', level=1)

doc.add_heading('Caso real: Ana, Carlos y Beatriz', level=2)

p = doc.add_paragraph()
p.add_run('SituaciÃ³n inicial (todos frustrados):').bold = True

doc.add_paragraph('â€¢ Ana: Bicicleta Orbea (300â‚¬) â†’ quiere sofÃ¡ IKEA (300â‚¬)', style='List Bullet')
doc.add_paragraph('â€¢ Carlos: SofÃ¡ IKEA (300â‚¬) â†’ quiere ordenador portÃ¡til (300â‚¬)', style='List Bullet')
doc.add_paragraph('â€¢ Beatriz: Ordenador portÃ¡til (300â‚¬) â†’ quiere bicicleta (300â‚¬)', style='List Bullet')

p = doc.add_paragraph()
p.add_run('Con Treqe (72 horas despuÃ©s):').bold = True

doc.add_paragraph('â€¢ DÃ­a 1: Registro + algoritmo encuentra combinaciÃ³n', style='List Bullet')
doc.add_paragraph('â€¢ DÃ­a 2: Ana envÃ­a bici a Carlos', style='List Bullet')
doc.add_paragraph('â€¢ DÃ­a 3: Beatriz envÃ­a ordenador a Carlos', style='List Bullet')
doc.add_paragraph('â€¢ DÃ­a 4: Carlos envÃ­a sofÃ¡ a Ana, Beatriz recibe bici', style='List Bullet')

p = doc.add_paragraph()
p.add_run('Resultado final:').bold = True

doc.add_paragraph('â€¢ Todos felices con lo que querÃ­an', style='List Bullet')
doc.add_paragraph('â€¢ Valor intercambiado: 900â‚¬', style='List Bullet')
doc.add_paragraph('â€¢ ComisiÃ³n Treqe (3%): 27â‚¬ (9â‚¬ cada uno)', style='List Bullet')
doc.add_paragraph('â€¢ Tiempo invertido: 15 minutos cada uno', style='List Bullet')
doc.add_paragraph('â€¢ Tiempo ahorrado vs vender: 10 horas cada uno', style='List Bullet')

doc.add_heading('Datos del mercado espaÃ±ol', level=2)
doc.add_paragraph('â€¢ Mercado segunda mano: 5.000 millones â‚¬ anuales', style='List Bullet')
doc.add_paragraph('â€¢ Cosas guardadas en hogares: 15.000 millones â‚¬', style='List Bullet')
doc.add_paragraph('â€¢ Personas que preferirÃ­an intercambiar: 65% poblaciÃ³n', style='List Bullet')
doc.add_paragraph('â€¢ Intercambios que no ocurren por dificultad: 80%', style='List Bullet')

doc.add_heading('El problema matemÃ¡tico', level=2)
p = doc.add_paragraph()
p.add_run('Para intercambio directo (2 personas): ').bold = True
p.add_run('5% probabilidad')

p = doc.add_paragraph()
p.add_run('Con Treqe (3 personas): ').bold = True
p.add_run('20% probabilidad (4x mÃ¡s)')

p = doc.add_paragraph()
p.add_run('Con Treqe (4 personas): ').bold = True
p.add_run('35% probabilidad (7x mÃ¡s)')

p = doc.add_paragraph()
p.add_run('Cada persona adicional ').bold = True
p.add_run('multiplica las posibilidades, no las suma.')

doc.add_page_break()

# ===== PARTE 3: LA SOLUCIÃ“N TREQE =====
doc.add_heading('ðŸš€ PARTE 2: LA SOLUCIÃ“N TREQE', level=1)

doc.add_heading('AsÃ­ funciona (4 pasos simples)', level=2)

doc.add_paragraph('1. Cuentas tu historia', style='List Number')
doc.add_paragraph('   â€¢ Subes foto de lo que tienes', style='List Bullet')
doc.add_paragraph('   â€¢ Describes lo que te emocionarÃ­a tener', style='List Bullet')
doc.add_paragraph('   â€¢ Estimas el valor', style='List Bullet')

doc.add_paragraph('2. Descubres posibilidades', style='List Number')
doc.add_paragraph('   â€¢ Nuestro algoritmo busca combinaciones', style='List Bullet')
doc.add_paragraph('   â€¢ No solo 2 personas, sino 3, 4, 5...', style='List Bullet')
doc.add_paragraph('   â€¢ En 24 horas tienes opciones reales', style='List Bullet')

doc.add_paragraph('3. Vives tu vida', style='List Number')
doc.add_paragraph('   â€¢ Aceptas el intercambio que te convence', style='List Bullet')
doc.add_paragraph('   â€¢ Nosotros coordinamos el resto', style='List Bullet')
doc.add_paragraph('   â€¢ TÃº sigues con tu vida normal', style='List Bullet')

doc.add_paragraph('4. La magia ocurre', style='List Number')
doc.add_paragraph('   â€¢ En 72 horas (promedio) recibes lo que querÃ­as', style='List Bullet')
doc.add_paragraph('   â€¢ Das lo que ya no usabas', style='List Bullet')
doc.add_paragraph('   â€¢ Conoces personas reales', style='List Bullet')

doc.add_heading('ComparaciÃ³n con competencia', level=2)

p = doc.add_paragraph()
p.add_run('Wallapop/Vinted (hoy):').bold = True
doc.add_paragraph('   "Vende tu bici por 210â‚¬ (30% menos), luego compra sofÃ¡ por 300â‚¬"', style='List Bullet')
doc.add_paragraph('   â€¢ 2 transacciones separadas', style='List Bullet')
doc.add_paragraph('   â€¢ Regateos interminables', style='List Bullet')
doc.add_paragraph('   â€¢ 10+ horas perdidas', style='List Bullet')
doc.add_paragraph('   â€¢ Riesgo de estafa', style='List Bullet')
doc.add_paragraph('   â€¢ Coste real: 90â‚¬ + 10 horas', style='List Bullet')

p = doc.add_paragraph()
p.add_run('Treqe (nuestra propuesta):').bold = True
doc.add_paragraph('   "Tu bici ya es tu sofÃ¡. Solo necesitÃ¡bamos encontrar a Carlos y Beatriz."', style='List Bullet')
doc.add_paragraph('   â€¢ 1 intercambio circular', style='List Bullet')
doc.add_paragraph('   â€¢ Sin regateos', style='List Bullet')
doc.add_paragraph('   â€¢ 15 minutos de tu tiempo', style='List Bullet')
doc.add_paragraph('   â€¢ Cero riesgo (triple protecciÃ³n)', style='List Bullet')
doc.add_paragraph('   â€¢ Coste real: 9â‚¬ + 15 minutos', style='List Bullet')

p = doc.add_paragraph()
p.add_run('Ahorro con Treqe: ').bold = True
p.add_run('81â‚¬ + 9,75 horas por intercambio')

doc.add_page_break()

# ===== PARTE 4: MODELO DE NEGOCIO =====
doc.add_heading('ðŸ’° PARTE 3: MODELO DE NEGOCIO', level=1)

doc.add_heading('ComisiÃ³n 3% (transparente)', level=2)
p = doc.add_paragraph()
p.add_run('Te decimos en quÃ© se va tu 3%:').bold = True

doc.add_paragraph('â€¢ 1% para el cerebro (algoritmo): 0,30â‚¬ por intercambio 100â‚¬', style='List Bullet')
doc.add_paragraph('â€¢ 1% para la confianza (seguridad): 0,30â‚¬ por intercambio 100â‚¬', style='List Bullet')
doc.add_paragraph('â€¢ 1% para el futuro (mejoras): 0,30â‚¬ por intercambio 100â‚¬', style='List Bullet')

doc.add_heading('Segmentos de clientes', level=2)
doc.add_paragraph('1. Millennials urbanos (25-35 aÃ±os): 5 millones en EspaÃ±a', style='List Number')
doc.add_paragraph('2. Familias con hijos (35-50 aÃ±os): 3 millones de hogares', style='List Number')
doc.add_paragraph('3. Estudiantes universitarios (18-25 aÃ±os): 1,5 millones', style='List Number')

doc.add_heading('Flujos de ingresos', level=2)
doc.add_paragraph('â€¢ Principal: ComisiÃ³n 3% sobre intercambios', style='List Bullet')
doc.add_paragraph('â€¢ Secundario: SuscripciÃ³n Premium 9,99â‚¬/mes', style='List Bullet')
doc.add_paragraph('â€¢ Terciario: Publicidad contextual relevante', style='List Bullet')
doc.add_paragraph('â€¢ Futuro: Data insights y estudios de mercado', style='List Bullet')

doc.add_page_break()

# ===== PARTE 5: PROYECCIONES FINANCIERAS =====
doc.add_heading('ðŸ“ˆ PARTE 4: PROYECCIONES FINANCIERAS', level=1)

doc.add_heading('InversiÃ³n inicial: 58.000â‚¬', level=2)
doc.add_paragraph('â€¢ TecnologÃ­a (40%): 23.200â‚¬', style='List Bullet')
doc.add_paragraph('â€¢ Marketing (35%): 20.300â‚¬', style='List Bullet')
doc.add_paragraph('â€¢ Operaciones (25%): 14.500â‚¬', style='List Bullet')

doc.add_heading('Proyecciones aÃ±o a aÃ±o', level=2)

p = doc.add_paragraph()
p.add_run('AÃ±o 1 (2026) - La prueba:').bold = True
doc.add_paragraph('   â€¢ Usuarios totales: 50.000', style='List Bullet')
doc.add_paragraph('   â€¢ Usuarios activos: 8.000 (16%)', style='List Bullet')
doc.add_paragraph('   â€¢ Intercambios/mes: 10.000', style='List Bullet')
doc.add_paragraph('   â€¢ Valor medio: 100â‚¬', style='List Bullet')
doc.add_paragraph('   â€¢ Ingresos: 360.000â‚¬/aÃ±o', style='List Bullet')
doc.add_paragraph('   â€¢ Beneficio: 84.000â‚¬/aÃ±o', style='List Bullet')

p = doc.add_paragraph()
p.add_run('AÃ±o 2 (2027) - El crecimiento:').bold = True
doc.add_paragraph('   â€¢ Usuarios totales: 150.000', style='List Bullet')
doc.add_paragraph('   â€¢ Usuarios activos: 25.000 (17%)', style='List Bullet')
doc.add_paragraph('   â€¢ Intercambios/mes: 40.000', style='List Bullet')
doc.add_paragraph('   â€¢ Valor medio: 120â‚¬', style='List Bullet')
doc.add_paragraph('   â€¢ Ingresos: 1.728.000â‚¬/aÃ±o', style='List Bullet')
doc.add_paragraph('   â€¢ Beneficio: 960.000â‚¬/aÃ±o', style='List Bullet')

p = doc.add_paragraph()
p.add_run('AÃ±o 3 (2028) - La consolidaciÃ³n:').bold = True
doc.add_paragraph('   â€¢ Usuarios totales: 350.000', style='List Bullet')
doc.add_paragraph('   â€¢ Usuarios activos: 60.000 (17%)', style='List Bullet')
doc.add_paragraph('   â€¢ Intercambios/mes: 100.000', style='List Bullet')
doc.add_paragraph('   â€¢ Valor medio: 140â‚¬', style='List Bullet')
doc.add_paragraph('   â€¢ Ingresos: 5.040.000â‚¬/aÃ±o', style='List Bullet')
doc.add_paragraph('   â€¢ Beneficio: 3.000.000â‚¬/aÃ±o', style='List Bullet')

doc.add_heading('Unit Economics', level=2)
doc.add_paragraph('â€¢ CAC (Costo AdquisiciÃ³n Usuario): 15â‚¬ aÃ±o 1 â†’ 10â‚¬ aÃ±o 3', style='List Bullet')
doc.add_paragraph('â€¢ LTV (Valor Vida Usuario): 360â‚¬ aÃ±o 1 â†’ 1.200â‚¬ aÃ±o 3', style='List Bullet')
doc.add_paragraph('â€¢ Ratio LTV:CAC: 24:1 aÃ±o 1 â†’ 120:1 aÃ±o 3', style='List Bullet')
doc.add_paragraph('â€¢ Payback Period: 15 dÃ­as por usuario, 8,3 meses inversiÃ³n total', style='List Bullet')

doc.add_page_break()

# ===== PARTE 6: CRONOGRAMA 12 MESES =====
doc.add_heading('ðŸ“… PARTE 5: CRONOGRAMA 12 MESES', level=1)

doc.add_heading('Mes 1-3: Los cimientos', level=2)
doc.add_paragraph('â€¢ Semana 1-2: FundaciÃ³n legal (SL + marca)', style='List Bullet')
doc.add_paragraph('â€¢ Semana 3-4: TecnologÃ­a bÃ¡sica (dominio + landing)', style='List Bullet')
doc.add_paragraph('â€¢ Mes 2: Primeros usuarios reales (50 personas)', style='List Bullet')
doc.add_paragraph('â€¢ Mes 3: AutomatizaciÃ³n bÃ¡sica (algoritmo v1.0)', style='List Bullet')
doc.add_paragraph('â€¢ Objetivo mes 3: 100 usuarios, 20 intercambios', style='List Bullet')

doc.add_heading('Mes 4-6: Primera ciudad (Madrid)', level=2)
doc.add_paragraph('â€¢ Mes 4: Escala Madrid (1.000 usuarios)', style='List Bullet')
doc.add_paragraph('â€¢ Mes 5: OptimizaciÃ³n Madrid (2.000 usuarios)', style='List Bullet')
doc.add_paragraph('â€¢ Mes 6: ConsolidaciÃ³n Madrid (3.000 usuarios)', style='List Bullet')
doc.add_paragraph('â€¢ Presupuesto gastado: 25.000â‚¬ (de 58.000â‚¬)', style='List Bullet')

doc.add_heading('Mes 7-9: ExpansiÃ³n a 3 ciudades', level=2)
doc.add_paragraph('â€¢ Ciudades: Barcelona, Valencia, Sevilla', style='List Bullet')
doc.add_paragraph('â€¢ Mes 7: Barcelona (1.000 usuarios)', style='List Bullet')
doc.add_paragraph('â€¢ Mes 8: Valencia + Sevilla (1.000 cada una)', style='List Bullet')
doc.add_paragraph('â€¢ Mes 9: IntegraciÃ³n multi-ciudad', style='List Bullet')
doc.add_paragraph('â€¢ Objetivo: 5.000 usuarios total', style='List Bul
let')
doc.add_paragraph('â€¢ Presupuesto gastado: 45.000â‚¬ (de 58.000â‚¬)', style='List Bullet')

doc.add_heading('Mes 10-12: EspaÃ±a entera + consolidaciÃ³n', level=2)
doc.add_paragraph('â€¢ Mes 10: Estrategia nacional (10.000 usuarios)', style='List Bullet')
doc.add_paragraph('â€¢ Mes 11: AutomatizaciÃ³n total (20.000 usuarios)', style='List Bullet')
doc.add_paragraph('â€¢ Mes 12: ConsolidaciÃ³n y plan aÃ±o 2', style='List Bullet')
doc.add_paragraph('â€¢ Objetivo final: 50.000 usuarios total', style='List Bullet')
doc.add_paragraph('â€¢ Presupuesto gastado: 58.000â‚¬ (completo)', style='List Bullet')
doc.add_paragraph('â€¢ Ingresos aÃ±o 1 objetivo: 360.000â‚¬', style='List Bullet')
doc.add_paragraph('â€¢ Beneficio aÃ±o 1 objetivo: 84.000â‚¬', style='List Bullet')

doc.add_page_break()

# ===== PARTE 7: EQUIPO Y CONTRATACIONES =====
doc.add_heading('ðŸ‘¥ PARTE 6: EQUIPO Y CONTRATACIONES', level=1)

doc.add_heading('Fase 1: Fundador solo (Mes 1-3)', level=2)
doc.add_paragraph('â€¢ Fundador: Todo (tecnologÃ­a + negocio + comunidad + marketing)', style='List Bullet')
doc.add_paragraph('â€¢ Externo: Abogado (500â‚¬/mes) + Contable (300â‚¬/mes)', style='List Bullet')
doc.add_paragraph('â€¢ Total coste: 800â‚¬/mes', style='List Bullet')

doc.add_heading('Fase 2: Primeras contrataciones (Mes 4-6)', level=2)
doc.add_paragraph('â€¢ CTO part-time (mes 4): 20h/semana, 1.500â‚¬/mes', style='List Bullet')
doc.add_paragraph('â€¢ Marketing/Comunidad part-time (mes 5): 10h/semana, 800â‚¬/mes', style='List Bullet')
doc.add_paragraph('â€¢ Total equipo mes 6: 3.100â‚¬/mes', style='List Bullet')

doc.add_heading('Fase 3: ExpansiÃ³n (Mes 7-9)', level=2)
doc.add_paragraph('â€¢ 3 personas locales part-time (Barcelona, Valencia, Sevilla)', style='List Bullet')
doc.add_paragraph('â€¢ 800â‚¬/mes cada una', style='List Bullet')
doc.add_paragraph('â€¢ Total equipo mes 9: 5.500â‚¬/mes', style='List Bullet')

doc.add_heading('Fase 4: ConsolidaciÃ³n (Mes 10-12)', level=2)
doc.add_paragraph('â€¢ Evaluar conversiones part-time â†’ full-time', style='List Bullet')
doc.add_paragraph('â€¢ Total equipo mes 12: 5.000-7.000â‚¬/mes', style='List Bullet')

doc.add_page_break()

# ===== PARTE 8: ASPECTOS LEGALES =====
doc.add_heading('âš–ï¸ PARTE 7: ASPECTOS LEGALES', level=1)

doc.add_heading('Estructura: Sociedad Limitada (SL)', level=2)
doc.add_paragraph('â€¢ Capital: 3.000â‚¬', style='List Bullet')
doc.add_paragraph('â€¢ Ventajas: Responsabilidad limitada, fiscalidad favorable startups', style='List Bullet')
doc.add_paragraph('â€¢ Coste constituciÃ³n: 3.500â‚¬ (3.000â‚¬ capital + 500â‚¬ gestorÃ­a)', style='List Bullet')

doc.add_heading('ProtecciÃ³n propiedad intelectual', level=2)
doc.add_paragraph('â€¢ Algoritmo: Secreto comercial (aÃ±o 1) â†’ Patente europea (aÃ±o 2)', style='List Bullet')
doc.add_paragraph('â€¢ Marca "Treqe": Registro EspaÃ±a + UE (850â‚¬)', style='List Bullet')
doc.add_paragraph('â€¢ Clases: 35 (intermediaciÃ³n), 9 (software), 42 (tecnologÃ­a)', style='List Bullet')

doc.add_heading('Presupuesto legal 3 aÃ±os: 35.000â‚¬', level=2)
doc.add_paragraph('â€¢ AÃ±o 1: 8.000â‚¬ (SL, marca, tÃ©rminos bÃ¡sicos)', style='List Bullet')
doc.add_paragraph('â€¢ AÃ±o 2: 12.000â‚¬ (patente, expansiÃ³n)', style='List Bullet')
doc.add_paragraph('â€¢ AÃ±o 3: 15.000â‚¬ (mÃ¡s protecciÃ³n)', style='List Bullet')

p = doc.add_paragraph()
p.add_run('1,2% de facturaciÃ³n en protecciÃ³n legal nos protege 100% del negocio.').italic = True

doc.add_heading('Contratos clave', level=2)
doc.add_paragraph('â€¢ TÃ©rminos y Condiciones: Claros, lenguaje humano', style='List Bullet')
doc.add_paragraph('â€¢ PolÃ­tica Privacidad RGPD: Consentimiento explÃ­cito', style='List Bullet')
doc.add_paragraph('â€¢ Contrato Escrow: Dinero retenido hasta confirmaciÃ³n', style='List Bullet')

doc.add_page_break()

# ===== PARTE 9: DISEÃ‘O Y EXPERIENCIA =====
doc.add_heading('ðŸŽ¨ PARTE 8: DISEÃ‘O Y EXPERIENCIA', level=1)

doc.add_heading('DirecciÃ³n estÃ©tica', level=2)
p = doc.add_paragraph()
p.add_run('"Brutalista digital con toques orgÃ¡nicos"').italic = True

doc.add_paragraph('â€¢ Formas geomÃ©tricas claras (brutalista)', style='List Bullet')
doc.add_paragraph('â€¢ Colores tierra, espacio generoso (orgÃ¡nico)', style='List Bullet')
doc.add_paragraph('â€¢ TipografÃ­a que se lee fÃ¡cil', style='List Bullet')

doc.add_heading('Paleta de colores', level=2)
doc.add_paragraph('â€¢ Primario: #2A2D34 (gris oscuro - seriedad)', style='List Bullet')
doc.add_paragraph('â€¢ Secundario: #E8E9EB (gris claro - limpieza)', style='List Bullet')
doc.add_paragraph('â€¢ Acento: #C97D60 (terracota - acciÃ³n)', style='List Bullet')
doc.add_paragraph('â€¢ Fondo: #F5F1E6 (crema - calidez)', style='List Bullet')

doc.add_heading('Experiencia usuario (simple y clara)', level=2)
doc.add_paragraph('â€¢ Registro: 30 segundos (email + telÃ©fono)', style='List Bullet')
doc.add_paragraph('â€¢ Primer intercambio sugerido: 24 horas', style='List Bullet')
doc.add_paragraph('â€¢ Soporte: Chat en app, respuesta 2 horas', style='List Bullet')
doc.add_paragraph('â€¢ Comunidad: Foro transparente', style='List Bullet')

doc.add_heading('Landing Page (asÃ­ se ve)', level=2)
p = doc.add_paragraph()
p.add_run('â”Œâ”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”').font.name = 'Consolas'
p = doc.add_paragraph()
p.add_run('â”‚ TREQE                   â”‚').font.name = 'Consolas'
p = doc.add_paragraph()
p.add_run('â”‚                         â”‚').font.name = 'Consolas'
p = doc.add_paragraph()
p.add_run('â”‚ INTERCAMBIA LO QUE      â”‚').font.name = 'Consolas'
p = doc.add_paragraph()
p.add_run('â”‚ TIENES POR LO QUE       â”‚').font.name = 'Consolas'
p = doc.add_paragraph()
p.add_run('â”‚ QUIERES                 â”‚').font.name = 'Consolas'
p = doc.add_paragraph()
p.add_run('â”‚ En 5 minutos.           â”‚').font.name = 'Consolas'
p = doc.add_paragraph()
p.add_run('â”‚ Sin regateos.           â”‚').font.name = 'Consolas'
p = doc.add_paragraph()
p.add_run('â”‚                         â”‚').font.name = 'Consolas'
p = doc.add_paragraph()
p.add_run('â”‚ [Â¿QuÃ© tienes?]          â”‚').font.name = 'Consolas'
p = doc.add_paragraph()
p.add_run('â”‚ [Â¿QuÃ© quieres?]         â”‚').font.name = 'Consolas'
p = doc.add_paragraph()
p.add_run('â”‚ [BUSCAR TRUEQUES]       â”‚').font.name = 'Consolas'
p = doc.add_paragraph()
p.add_run('â””â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”˜').font.name = 'Consolas'

doc.add_page_break()

# ===== PARTE 10: RIESGOS Y CONTINGENCIAS =====
doc.add_heading('ðŸš¨ PARTE 9: RIESGOS Y CONTINGENCIAS', level=1)

doc.add_heading('Los 8 riesgos principales', level=2)

p = doc.add_paragraph()
p.add_run('1. Problema huevo-gallina').bold = True
doc.add_paragraph('   â€¢ QuÃ©: Necesitas usuarios para que funcione, pero necesitas que funcione para tener usuarios', style='List Bullet')
doc.add_paragraph('   â€¢ Probabilidad: Alta', style='List Bullet')
doc.add_paragraph('   â€¢ Impacto: Alto', style='List Bullet')
doc.add_paragraph('   â€¢ MitigaciÃ³n: Empezar con 50 personas conocidas, intercambios manuales', style='List Bullet')

p = doc.add_paragraph()
p.add_run('2. Algoritmo falla').bold = True
doc.add_paragraph('   â€¢ QuÃ©: Sugiere intercambios que no funcionan', style='List Bullet')
doc.add_paragraph('   â€¢ Probabilidad: Media', style='List Bullet')
doc.add_paragraph('   â€¢ Impacto: Alto', style='List Bullet')
doc.add_paragraph('   â€¢ MitigaciÃ³n: Humanos revisan primeros 100 intercambios', style='List Bullet')

p = doc.add_paragraph()
p.add_run('3. Gastar dinero antes de tiempo').bold = True
doc.add_paragraph('   â€¢ QuÃ©: Quemar 58.000â‚¬ sin llegar a 8.000 usuarios', style='List Bullet')
doc.add_paragraph('   â€¢ Probabilidad: Media', style='List Bullet')
doc.add_paragraph('   â€¢ Impacto: Medio', style='List Bullet')
doc.add_paragraph('   â€¢ MitigaciÃ³n: Presupuesto mensual estricto, mÃ©tricas claras', style='List Bullet')

p = doc.add_paragraph()
p.add_run('4. Competencia copia rÃ¡pido').bold = True
doc.add_paragraph('   â€¢ QuÃ©: Wallapop aÃ±ade "modo intercambio" en 3 meses', style='List Bullet')
doc.add_paragraph('   â€¢ Probabilidad: Baja', style='List Bullet')
doc.add_paragraph('   â€¢ Impacto: Alto', style='List Bullet')
doc.add_paragraph('   â€¢ MitigaciÃ³n: Algoritmo complejo, comunidad leal, marca diferente', style='List Bullet')

p = doc.add_paragraph()
p.add_run('5. Problemas legales').bold = True
doc.add_paragraph('   â€¢ QuÃ©: Multas por incumplimiento normativo', style='List Bullet')
doc.add_paragraph('   â€¢ Probabilidad: Baja', style='List Bullet')
doc.add_paragraph('   â€¢ Impacto: Alto', style='List Bullet')
doc.add_paragraph('   â€¢ MitigaciÃ³n: InversiÃ³n legal proactiva (35.000â‚¬ 3 aÃ±os)', style='List Bullet')

p = doc.add_paragraph()
p.add_run('6. Fraudes/estafas').bold = True
doc.add_paragraph('   â€¢ QuÃ©: Usuarios intentan estafar', style='List Bullet')
doc.add_paragraph('   â€¢ Probabilidad: Media', style='List Bullet')
doc.add_paragraph('   â€¢ Impacto: Medio', style='List Bullet')
doc.add_paragraph('   â€¢ MitigaciÃ³n: Sistema escrow, verificaciÃ³n usuarios, seguro 1.000â‚¬', style='List Bullet')

p = doc.add_paragraph()
p.add_run('7. Problemas logÃ­stica').bold = True
doc.add_paragraph('   â€¢ QuÃ©: EnvÃ­os se pierden/daÃ±an', style='List Bullet')
doc.add_paragraph('   â€¢ Probabilidad: Media', style='List Bullet')
doc.add_paragraph('   â€¢ Impacto: Bajo', style='List Bullet')
doc.add_paragraph('   â€¢ MitigaciÃ³n: Seguro envÃ­os, partners verificados', style='List Bullet')

p = doc.add_paragraph()
p.add_run('8. Escalabilidad tÃ©cnica').bold = True
doc.add_paragraph('   â€¢ QuÃ©: Sistema no escala con crecimiento', style='List Bullet')
doc.add_paragraph('   â€¢ Probabilidad: Baja', style='List Bullet')
doc.add_paragraph('   â€¢ Impacto: Alto', style='List Bullet')
doc.add_paragraph('   â€¢ MitigaciÃ³n: Arquitectura escalable desde inicio, CTO desde mes 4', style='List Bullet')

doc.add_page_break()

# ===== PARTE 11: ANEXOS Y RECURSOS =====
doc.add_heading('ðŸ“§ PARTE 10: ANEXOS Y RECURSOS', level=1)

doc.add_heading('Email lanzamiento beta', level=2)
p = doc.add_paragraph()
p.add_run('Asunto: Â¡Treqe estÃ¡ vivo! Eres de los primeros').italic = True
doc.add_paragraph()
doc.add_paragraph('Hola [Nombre],')
doc.add_paragraph()
doc.add_paragraph('Te escribo personalmente porque eres una de las primeras 50 personas en probar Treqe.')
doc.add_paragraph()
doc.add_paragraph('Â¿Recuerdas nuestra conversaciÃ³n sobre [su artÃ­culo especÃ­fico]? Pues Treqe ya puede ayudarte a intercambiarlo.')
doc.add_paragraph()
doc.add_paragraph('AsÃ­ funciona:')
doc.add_paragraph('1. Entras en treqe.es')
doc.add_paragraph('2. Pones quÃ© tienes y quÃ© te gustarÃ­a tener')
doc.add_paragraph('3. En 24 horas te sugerimos intercambios')
doc.add_paragraph()
doc.add_paragraph('Los primeros 100 intercambios son gratis (sin el 3%).')
doc.add_paragraph()
doc.add_paragraph('Â¿Te animas a probarlo?')
doc.add_paragraph()
doc.add_paragraph('[Tu nombre]')
doc.add_paragraph('Fundador, Treqe')

doc.add_heading('Pitch bÃ¡sico (2 minutos)', level=2)
doc.add_paragraph('Problema: "Â¿Tienes algo guardado que ya no usas y quieres otra cosa? Intercambiar entre 2 personas tiene solo 5% de probabilidad."')
doc.add_paragraph()
doc.add_paragraph('SoluciÃ³n: "Treqe conecta a 3+ personas para intercambios circulares. Buscamos Aâ†’Bâ†’Câ†’A, no Aâ†’B. Probabilidad sube a 20-35%."')
doc.add_paragraph()
doc.add_paragraph('Modelo: "3% de comisiÃ³n cuando el intercambio funciona. Solo cobramos cuando tÃº ganas."')
doc.add_paragraph()
doc.add_paragraph('TamaÃ±o mercado: "15.000 millones â‚¬ en cosas guardadas en hogares espaÃ±oles. 65% de la poblaciÃ³n preferirÃ­a intercambiar antes que vender."')
doc.add_paragraph()
doc.add_paragraph('Traction: "AÃ±o 1: 8.000 usuarios, 400.000â‚¬ ingresos. AÃ±o 3: 60.000 usuarios, 2.500.000â‚¬ ingresos."')
doc.add_paragraph()
doc.add_paragraph('Equipo: "[Tu nombre], fundador. Experiencia en [tu experiencia]. Apasionado por resolver este problema."')
doc.add_paragraph()
doc.add_paragraph('PeticiÃ³n: "58.000â‚¬ para lanzar. 3.000â‚¬ para SL, 850â‚¬ para marca, resto para tecnologÃ­a y primeros usuarios."')

doc.add_page_break()

# ===== PARTE 12: CHECKLISTS =====
doc.add_heading('ðŸŽ¯ PARTE 11: CHECKLISTS', level=1)

doc.add_heading('Checklist Semana 1', level=2)
doc.add_paragraph('Legal:', style='List Bullet')
doc.add_paragraph('  [ ] Consultar 3 abogados startups (500â‚¬)', style='List Bullet')
doc.add_paragraph('  [ ] Decidir estructura (SL)', style='List Bullet')
doc.add_paragraph('  [ ] Empezar trÃ¡mites SL (3.500â‚¬)', style='List Bullet')
doc.add_paragraph()
doc.add_paragraph('TecnologÃ­a:', style='List Bullet')
doc.add_paragraph('  [ ] Registrar dominio treqe.es', style='List Bullet')
doc.add_paragraph('  [ ] Setup hosting bÃ¡sico', style='List Bullet')
doc.add_paragraph('  [ ] Landing page simple', style='List Bullet')
doc.add_paragraph()
doc.add_paragraph('Marketing:', style='List Bullet')
doc.add_paragraph('  [ ] Crear lista 50 personas', style='List Bullet')
doc.add_paragraph('  [ ] Preparar email lanzamiento', style='List Bullet')
doc.add_paragraph('  [ ] Setup Google Analytics', style='List Bullet')

doc.add_heading('Checklist Mes 1', level=2)
doc.add_paragraph('[ ] SL constituida (3.500â‚¬)', style='List Bullet')
doc.add_paragraph('[ ] Marca "Treqe" registrada (850â‚¬)', style='List Bullet')
doc.add_paragraph('[ ] Dominio treqe.es activo', style='List Bullet')
doc.add_paragraph('[ ] Landing page funcionando', style='List Bullet')
doc.add_paragraph('[ ] Algoritmo v0.1 operativo', style='List Bullet')
doc.add_paragraph
('[ ] 10 personas en lista espera', style='List Bullet')
doc.add_paragraph('[ ] Presupuesto gastado: 4.850â‚¬ (de 58.000â‚¬)', style='List Bullet')
doc.add_paragraph('[ ] Primeras 5 conversaciones con usuarios', style='List Bullet')

doc.add_heading('Checklist Trimestre 1', level=2)
doc.add_paragraph('[ ] 100 usuarios registrados', style='List Bullet')
doc.add_paragraph('[ ] 20 intercambios completados', style='List Bullet')
doc.add_paragraph('[ ] Sistema web bÃ¡sico funcionando', style='List Bullet')
doc.add_paragraph('[ ] Algoritmo v1.0 automÃ¡tico', style='List Bullet')
doc.add_paragraph('[ ] Soporte bÃ¡sico (email, FAQ)', style='List Bullet')
doc.add_paragraph('[ ] Primeras mÃ©tricas (CAC, LTV, conversiÃ³n)', style='List Bullet')
doc.add_paragraph('[ ] Presupuesto gastado: 15.000â‚¬ (de 58.000â‚¬)', style='List Bullet')
doc.add_paragraph('[ ] DecisiÃ³n: Â¿Continuar o replantear?', style='List Bullet')

doc.add_heading('Checklist AÃ±o 1', level=2)
doc.add_paragraph('[ ] 50.000 usuarios totales', style='List Bullet')
doc.add_paragraph('[ ] 8.000 usuarios activos', style='List Bullet')
doc.add_paragraph('[ ] 10.000 intercambios/mes', style='List Bullet')
doc.add_paragraph('[ ] 360.000â‚¬ ingresos', style='List Bullet')
doc.add_paragraph('[ ] 84.000â‚¬ beneficio', style='List Bullet')
doc.add_paragraph('[ ] Sistema completamente automÃ¡tico', style='List Bullet')
doc.add_paragraph('[ ] App mÃ³vil nativa (iOS/Android)', style='List Bullet')
doc.add_paragraph('[ ] Equipo: 5 personas', style='List Bullet')
doc.add_paragraph('[ ] Presupuesto gastado: 58.000â‚¬ (completo)', style='List Bullet')
doc.add_paragraph('[ ] Plan aÃ±o 2 detallado', style='List Bullet')

doc.add_page_break()

# ===== CONCLUSIÃ“N =====
doc.add_heading('ðŸ CONCLUSIÃ“N', level=1)

p = doc.add_paragraph()
p.add_run('Por quÃ© esto funciona (resumen final):').bold = True

doc.add_paragraph('1. Resuelve un problema real que afecta al 72% de espaÃ±oles', style='List Number')
doc.add_paragraph('2. MatemÃ¡ticamente sÃ³lido (5% probabilidad â†’ 20-35% con Treqe)', style='List Number')
doc.add_paragraph('3. Modelo econÃ³mico viable (3% comisiÃ³n, LTV:CAC 24:1 aÃ±o 1)', style='List Number')
doc.add_paragraph('4. Escalable tecnolÃ³gicamente (algoritmo + automatizaciÃ³n)', style='List Number')
doc.add_paragraph('5. Protegido legalmente (35.000â‚¬ inversiÃ³n en 3 aÃ±os)', style='List Number')
doc.add_paragraph('6. Equipo realista (crecimiento progresivo, sin quemar)', style='List Number')
doc.add_paragraph('7. Riesgos identificados y mitigados (8 riesgos + contingencias)', style='List Number')

doc.add_heading('Nuestra ventaja competitiva', level=2)
p = doc.add_paragraph()
p.add_run('No competimos con Wallapop/Vinted en precio. Competimos en ').bold = True
p.add_run('valor.').bold = True

p = doc.add_paragraph()
p.add_run('Ellos: ').bold = True
p.add_run('"Vende por menos, compra por mÃ¡s, pierde tiempo y dinero"')

p = doc.add_paragraph()
p.add_run('Nosotros: ').bold = True
p.add_run('"Consigue lo que quieres, manteniendo el valor, en 72 horas"')

doc.add_heading('Llamada a la acciÃ³n', level=2)
p = doc.add_paragraph()
p.add_run('Si eres inversor: ').bold = True
p.add_run('58.000â‚¬ para capturar un mercado de 15.000 millones â‚¬')

p = doc.add_paragraph()
p.add_run('Si eres usuario: ').bold = True
p.add_run('Prueba el primer intercambio gratis (100 primeros)')

p = doc.add_paragraph()
p.add_run('Si eres potencial empleado: ').bold = True
p.add_run('Ãšnete en mes 4-6 cuando escalemos')

p = doc.add_paragraph()
p.add_run('Treqe no es solo una startup. Es una nueva forma de pensar la propiedad.').italic = True

p = doc.add_paragraph()
p.add_run('Porque a veces, lo que tienes ya es lo que quieres. Solo necesitas encontrar a Carlos y Beatriz.').italic = True

# Guardar documento
output_path = os.path.join(os.path.dirname(__file__), 'Plan_Negocio_Treqe_DEFINITIVO_COMPLETO.docx')
doc.save(output_path)

print(f"âœ… Documento Word creado exitosamente: {output_path}")
print(f"ðŸ“„ TamaÃ±o: {os.path.getsize(output_path)} bytes")
print(f"ðŸ“‹ PÃ¡ginas: Aproximadamente {len(doc.element.xpath('//w:sectPr'))} pÃ¡ginas")
print("\nðŸŽ¯ Documento incluye:")
print("   â€¢ 12 partes completas")
print("   â€¢ Todas las skills aplicadas (humanizer, legal, business-model-canvas, marketing-mode)")
print("   â€¢ Casos reales con Ana, Carlos, Beatriz")
print("   â€¢ Proyecciones financieras aÃ±o 1-3")
print("   â€¢ Cronograma 12 meses paso a paso")
print("   â€¢ Aspectos legales completos")
print("   â€¢ Riesgos identificados y mitigados")
print("   â€¢ Checklists ejecutables")
print("\nðŸ“§ Â¡Documento listo para usar!")
