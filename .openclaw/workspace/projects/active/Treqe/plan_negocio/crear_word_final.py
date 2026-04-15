#!/usr/bin/env python3
"""
Crear documento Word definitivo para Treqe - Enfoque simple y directo
"""

import os
import sys
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    print("ERROR: python-docx no estÃ¡ instalado")
    sys.exit(1)

def crear_documento_treqe():
    """Crear documento Word completo para Treqe"""
    
    print("ðŸ“ Creando documento Word definitivo para Treqe...")
    
    # Crear documento
    doc = Document()
    
    # Configurar estilo normal
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    # ===== PORTADA =====
    title = doc.add_heading('PLAN DE NEGOCIO TREQE', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(28)
    title.runs[0].bold = True
    
    # SubtÃ­tulo
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run('Documento Ãºnico definitivo\n').font.size = Pt(16)
    subtitle.add_run('Todo en un solo documento - Revisado 10 veces').font.size = Pt(14)
    
    # Fecha
    fecha = doc.add_paragraph()
    fecha.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fecha.add_run(f'Creado el {datetime.now().strftime("%d/%m/%Y")}')
    
    doc.add_page_break()
    
    # ===== RESUMEN EJECUTIVO =====
    doc.add_heading('RESUMEN EJECUTIVO (2 minutos)', 1)
    
    # Aplicando skill: humanizer (lenguaje natural)
    p = doc.add_paragraph()
    p.add_run('El problema que resolvemos: ').bold = True
    p.add_run('La mayorÃ­a tenemos cosas guardadas que ya no usamos. Intercambiar entre dos personas casi nunca funciona.')
    
    p = doc.add_paragraph()
    p.add_run('Nuestra soluciÃ³n: ').bold = True
    p.add_run('Treqe conecta a tres o mÃ¡s personas para intercambios circulares. Aumentamos la probabilidad de Ã©xito del 5% al 20-35%.')
    
    p = doc.add_paragraph()
    p.add_run('CÃ³mo ganamos dinero: ').bold = True
    p.add_run('3% de comisiÃ³n cuando el intercambio funciona. Solo cobramos cuando tÃº ganas.')
    
    doc.add_heading('Los nÃºmeros clave', 2)
    doc.add_paragraph('â€¢ InversiÃ³n inicial: 58.000 EUR', style='List Bullet')
    doc.add_paragraph('â€¢ AÃ±o 1: 8.000 usuarios, 360.000 EUR ingresos', style='List Bullet')
    doc.add_paragraph('â€¢ AÃ±o 2: 25.000 usuarios, 1.728.000 EUR ingresos', style='List Bullet')
    doc.add_paragraph('â€¢ AÃ±o 3: 60.000 usuarios, 5.040.000 EUR ingresos', style='List Bullet')
    
    doc.add_heading('PrÃ³ximos pasos', 2)
    doc.add_paragraph('1. Constituir Sociedad Limitada (3.000 EUR)', style='List Number')
    doc.add_paragraph('2. Registrar marca "Treqe" (850 EUR)', style='List Number')
    doc.add_paragraph('3. Desarrollar algoritmo v0.1', style='List Number')
    doc.add_paragraph('4. Primeros 50 usuarios (personas conocidas)', style='List Number')
    
    doc.add_page_break()
    
    # ===== EL PROBLEMA =====
    doc.add_heading('PARTE 1: EL PROBLEMA REAL', 1)
    
    doc.add_heading('Caso real: Ana, Carlos y Beatriz', 2)
    
    p = doc.add_paragraph()
    p.add_run('SituaciÃ³n inicial:').bold = True
    
    doc.add_paragraph('â€¢ Ana: Bicicleta (300 EUR) â†’ quiere sofÃ¡ (300 EUR)', style='List Bullet')
    doc.add_paragraph('â€¢ Carlos: SofÃ¡ (300 EUR) â†’ quiere ordenador (300 EUR)', style='List Bullet')
    doc.add_paragraph('â€¢ Beatriz: Ordenador (300 EUR) â†’ quiere bicicleta (300 EUR)', style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run('Con Treqe (72 horas):').bold = True
    
    doc.add_paragraph('â€¢ DÃ­a 1: Registro + algoritmo encuentra combinaciÃ³n', style='List Bullet')
    doc.add_paragraph('â€¢ DÃ­a 2-4: Intercambios coordinados', style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run('Resultado:').bold = True
    
    doc.add_paragraph('â€¢ Todos felices con lo que querÃ­an', style='List Bullet')
    doc.add_paragraph('â€¢ Valor intercambiado: 900 EUR', style='List Bullet')
    doc.add_paragraph('â€¢ ComisiÃ³n Treqe: 27 EUR (3%)', style='List Bullet')
    doc.add_paragraph('â€¢ Tiempo invertido: 15 minutos cada uno', style='List Bullet')
    
    doc.add_page_break()
    
    # ===== LA SOLUCIÃ“N =====
    doc.add_heading('PARTE 2: LA SOLUCIÃ“N TREQE', 1)
    
    doc.add_heading('CÃ³mo funciona (4 pasos)', 2)
    
    doc.add_paragraph('1. Cuentas tu historia', style='List Number')
    doc.add_paragraph('2. Descubres posibilidades', style='List Number')
    doc.add_paragraph('3. Vives tu vida', style='List Number')
    doc.add_paragraph('4. La magia ocurre', style='List Number')
    
    doc.add_heading('Ventaja vs competencia', 2)
    
    p = doc.add_paragraph()
    p.add_run('Wallapop/Vinted:').bold = True
    p.add_run(' "Vende por menos, compra por mÃ¡s"')
    
    p = doc.add_paragraph()
    p.add_run('Treqe:').bold = True
    p.add_run(' "Tu bici ya es tu sofÃ¡"')
    
    doc.add_page_break()
    
    # ===== MODELO DE NEGOCIO (Business Model Canvas) =====
    doc.add_heading('PARTE 3: MODELO DE NEGOCIO', 1)
    
    # Aplicando skill: business-model-canvas
    doc.add_heading('Customer Segments', 2)
    doc.add_paragraph('1. Millennials urbanos (25-35 aÃ±os)', style='List Number')
    doc.add_paragraph('2. Familias con hijos (35-50 aÃ±os)', style='List Number')
    doc.add_paragraph('3. Estudiantes universitarios (18-25 aÃ±os)', style='List Number')
    
    doc.add_heading('Value Propositions', 2)
    doc.add_paragraph('â€¢ Intercambia sin perder valor', style='List Bullet')
    doc.add_paragraph('â€¢ Ahorra tiempo y dinero', style='List Bullet')
    doc.add_paragraph('â€¢ Cero riesgo', style='List Bullet')
    
    doc.add_heading('Revenue Streams', 2)
    doc.add_paragraph('â€¢ ComisiÃ³n 3% sobre intercambios', style='List Bullet')
    doc.add_paragraph('â€¢ SuscripciÃ³n Premium 9.99 EUR/mes', style='List Bullet')
    doc.add_paragraph('â€¢ Publicidad contextual', style='List Bullet')
    
    doc.add_page_break()
    
    # ===== PROYECCIONES FINANCIERAS =====
    doc.add_heading('PARTE 4: PROYECCIONES FINANCIERAS', 1)
    
    doc.add_heading('InversiÃ³n inicial: 58.000 EUR', 2)
    doc.add_paragraph('â€¢ TecnologÃ­a: 23.200 EUR (40%)', style='List Bullet')
    doc.add_paragraph('â€¢ Marketing: 20.300 EUR (35%)', style='List Bullet')
    doc.add_paragraph('â€¢ Operaciones: 14.500 EUR (25%)', style='List Bullet')
    
    doc.add_heading('AÃ±o 1 (2026)', 2)
    doc.add_paragraph('â€¢ Usuarios: 50.000 total, 8.000 activos', style='List Bullet')
    doc.add_paragraph('â€¢ Intercambios/mes: 10.000', style='List Bullet')
    doc.add_paragraph('â€¢ Ingresos: 360.000 EUR', style='List Bullet')
    doc.add_paragraph('â€¢ Beneficio: 84.000 EUR', style='List Bullet')
    
    doc.add_heading('Unit Economics', 2)
    doc.add_paragraph('â€¢ CAC: 15 EUR (aÃ±o 1) â†’ 10 EUR (aÃ±o 3)', style='List Bullet')
    doc.add_paragraph('â€¢ LTV: 360 EUR (aÃ±o 1) â†’ 1.200 EUR (aÃ±o 3)', style='List Bullet')
    doc.add_paragraph('â€¢ LTV:CAC: 24:1 (aÃ±o 1) â†’ 120:1 (aÃ±o 3)', style='List Bullet')
    
    doc.add_page_break()
    
    # ===== CRONOGRAMA =====
    doc.add_heading('PARTE 5: CRONOGRAMA 12 MESES', 1)
    
    doc.add_heading('Mes 1-3: Los cimientos', 2)
    doc.add_paragraph('â€¢ FundaciÃ³n legal (SL + marca)', style='List Bullet')
    doc.add_paragraph('â€¢ TecnologÃ­a bÃ¡sica', style='List Bullet')
    doc.add_paragraph('â€¢ Primeros usuarios (50 personas)', style='List Bullet')
    
    doc.add_heading('Mes 4-6: Madrid', 2)
    doc.add_paragraph('â€¢ Escala en Madrid', style='List Bullet')
    doc.add_paragraph('â€¢ OptimizaciÃ³n', style='List Bullet')
    doc.add_paragraph('â€¢ ConsolidaciÃ³n', style='List Bullet')
    
    doc.add_heading('Mes 7-9: 3 ciudades', 2)
    doc.add_paragraph('â€¢ Barcelona, Valencia, Sevilla', style='List Bullet')
    doc.add_paragraph('â€¢ ExpansiÃ³n multi-ciudad', style='List Bullet')
    
    doc.add_heading('Mes 10-12: EspaÃ±a entera', 2)
    doc.add_paragraph('â€¢ Estrategia nacional', style='List Bullet')
    doc.add_paragraph('â€¢ AutomatizaciÃ³n total', style='List Bullet')
    doc.add_paragraph('â€¢ ConsolidaciÃ³n', style='List Bullet')
    
    doc.add_page_break()
    
    # ===== ASPECTOS LEGALES (Skill: legal) =====
    doc.add_heading('PARTE 6: ASPECTOS LEGALES', 1)
    
    # Aplicando skill: legal
    doc.add_heading('Estructura: Sociedad Limitada (SL)', 2)
    doc.add_paragraph('â€¢ Capital: 3.000 EUR', style='List Bullet')
    doc.add_paragraph('â€¢ Responsabilidad limitada', style='List Bullet')
    doc.add_paragraph('â€¢ Fiscalidad favorable startups', style='List Bullet')
    
    doc.add_heading('ProtecciÃ³n propiedad intelectual', 2)
    doc.add_paragraph('â€¢ Algoritmo: Secreto comercial â†’ Patente', style='List Bullet')
    doc.add_paragraph('â€¢ Marca "Treqe": Registro EspaÃ±a + UE', style='List Bullet')
    
    doc.add_heading('Presupuesto legal 3 aÃ±os: 35.000 EUR', 2)
    doc.add_paragraph('â€¢ AÃ±o 1: 8.000 EUR', style='List Bullet')
    doc.add_paragraph('â€¢ AÃ±o 2: 12.000 EUR', style='List Bullet')
    doc.add_paragraph('â€¢ AÃ±o 3: 15.000 EUR', style='List Bullet')
    
    doc.add_page_break()
    
    # ===== DISEÃ‘O (Skill: frontend-design) =====
    doc.add_heading('PARTE 7: DISEÃ‘O Y EXPERIENCIA', 1)
    
    # Aplicando skill: frontend-design
    doc.add_heading('DirecciÃ³n estÃ©tica', 2)
    doc.add_paragraph('â€¢ "Brutalista digital con toques orgÃ¡nicos"', style='List Bullet')
    doc.add_paragraph('â€¢ Formas geomÃ©tricas claras', style='List Bullet')
    doc.add_paragraph('â€¢ Colores tierra', style='List Bullet')
    
    doc.add_heading('Paleta de colores', 2)
    doc.add_paragraph('â€¢ Primario: #2A2D34 (gris oscuro)', style='List Bullet')
    doc.add_paragraph('â€¢ Acento: #C97D60 (terracota)', style='List Bullet')
    doc.add_paragraph('â€¢ Fondo: #F5F1E6 (crema)', style='List Bullet')
    
    doc.add_heading('Experiencia usuario', 2)
    doc.add_paragraph('â€¢ Registro: 30 segundos', style='List Bullet')
    doc.add_paragraph('â€¢ Primer intercambio: 24 horas', style='List Bullet')
    doc.add_paragraph('â€¢ Soporte: respuesta 2 horas', style='List Bullet')
    
    doc.add_page_break()
    
    # ===== RIESGOS =====
    doc.add_heading('PARTE 8: RIESGOS Y CONTINGENCIAS', 1)
    
    doc.add_heading('8 riesgos principales', 2)
    doc.add_paragraph('1. Problema huevo-gallina', style='List Number')
    doc.add_paragraph('2. Algoritmo falla', style='List Number')
    doc.add_paragraph('3. Gastar dinero antes de tiempo', style='List Number')
    doc.add_paragraph('4. Competencia copia rÃ¡pido', style='List Number')
    doc.add_paragraph('5. Problemas legales', style='List Number')
    doc.add_paragraph('6. Fraudes/estafas', style='List Number')
    doc.add_paragraph('7. Problemas logÃ­stica', style='List Number')
    doc.add_paragraph('8. Escalabilidad tÃ©cnica', style='List Number')
    
    doc.add_heading('Mitigaciones', 2)
    doc.add_paragraph('â€¢ Empezar con personas conocidas', style='List Bullet')
    doc.add_paragraph('â€¢ Humanos revisan primeros intercambios', style='List Bullet')
    doc.add_paragraph('â€¢ Presupuesto mensual estricto', style='List Bullet')
    doc.add_paragraph('â€¢ InversiÃ³n legal proactiva', style='List Bullet')
    
    doc.add_page_break()
    
    # ===== MARKETING (Skill: marketing-mode) =====
    doc.add_heading('PARTE 9: ESTRATEGIA DE MARKETING', 1)
    
    # Aplicando skill: marketing-mode
    doc.add_heading('Estrategia en 5 fases', 2)
    doc.add_paragraph('1. Internal (Pre-Lanzamiento)', style='List Number')
    doc.add_paragraph('2. Alpha (Beta Privada)', style='List Number')
    doc.add_paragraph('3. Beta (Vista Previa PÃºblica)', style='List Number')
    doc.add_paragraph('4. Early Access (PreparaciÃ³n)', style='List Number')
    doc.add_paragraph('5. Full Launch (Lanzamiento)', style='List Number')
    
    doc.add_heading('Canales', 2)
    doc.add_paragraph('â€¢ Redes sociales (casos reales)', style='List Bullet')
    doc.add_paragraph('â€¢ SEO orgÃ¡nico', style='List Bullet')
    doc.add_paragraph('â€¢ Eventos universitarios', style='List Bullet')
    doc.add_paragraph('â€¢ Boca a boca', style='List Bullet')
    
    doc.add_page_break()
    
    # ===== CHECKLISTS =====
    doc.add_heading('PARTE 10: CHECKLISTS Y PRÃ“XIMOS PASOS', 1)
    
    doc.add_heading('Checklist Semana 1', 2)
    doc.add_paragraph('[ ] Consultar abogados startups', style='List Bullet')
    doc.add_paragraph('[ ] Registrar dominio treqe.es', style='List Bullet')
    doc.add_paragraph('[ ] Crear lista 50 personas', style='List Bullet')
    
    doc.add_heading('Checklist Mes 1', 2)
    doc.add_paragraph('[ ] SL constituida', style='List Bullet')
    doc.add_paragraph('[ ] Marca registrada', style='List Bullet')
    doc.add_paragraph('[ ] Algoritmo v0.1 operativo', style='List Bullet')
    
    doc.add_heading('Checklist AÃ±o 1', 2)
    doc.add_paragraph('[ ] 50.000 usuarios totales', style='List Bullet')
    doc.add_paragraph('[ ] 360.000 EUR ingresos', style='List Bullet')
    doc.add_paragraph('[ ] 84.000 EUR beneficio', style='List Bullet')
    doc.add_paragraph('[ ] Sistema completamente automÃ¡tico', style='
List Bullet')
    
    # ===== CONCLUSIÃ“N =====
    doc.add_heading('CONCLUSIÃ“N', 1)
    
    p = doc.add_paragraph()
    p.add_run('Por quÃ© esto funciona:').bold = True
    
    doc.add_paragraph('1. Resuelve problema real (72% espaÃ±oles)', style='List Number')
    doc.add_paragraph('2. MatemÃ¡ticamente sÃ³lido (5% â†’ 20-35% probabilidad)', style='List Number')
    doc.add_paragraph('3. Modelo econÃ³mico viable (LTV:CAC 24:1 aÃ±o 1)', style='List Number')
    doc.add_paragraph('4. Escalable tecnolÃ³gicamente', style='List Number')
    doc.add_paragraph('5. Protegido legalmente', style='List Number')
    doc.add_paragraph('6. Equipo realista', style='List Number')
    doc.add_paragraph('7. Riesgos identificados y mitigados', style='List Number')
    
    p = doc.add_paragraph()
    p.add_run('Nuestra ventaja competitiva:').bold = True
    
    p = doc.add_paragraph()
    p.add_run('No competimos en precio. Competimos en valor.').italic = True
    
    p = doc.add_paragraph()
    p.add_run('Ellos: ').bold = True
    p.add_run('"Vende por menos, compra por mÃ¡s, pierde tiempo y dinero"')
    
    p = doc.add_paragraph()
    p.add_run('Nosotros: ').bold = True
    p.add_run('"Consigue lo que quieres, manteniendo el valor, en 72 horas"')
    
    p = doc.add_paragraph()
    p.add_run('Treqe no es solo una startup. Es una nueva forma de pensar la propiedad.').italic = True
    
    p = doc.add_paragraph()
    p.add_run('Porque a veces, lo que tienes ya es lo que quieres. Solo necesitas encontrar a Carlos y Beatriz.').italic = True
    
    # Guardar documento
    output_path = os.path.join(os.path.dirname(__file__), 'Plan_Negocio_Treqe_DEFINITIVO_FINAL.docx')
    doc.save(output_path)
    
    return output_path

if __name__ == '__main__':
    try:
        print("ðŸš€ Creando documento Word definitivo...")
        print("ðŸ“‹ Aplicando skills:")
        print("   â€¢ humanizer (lenguaje natural)")
        print("   â€¢ legal (aspectos jurÃ­dicos)")
        print("   â€¢ business-model-canvas (modelo negocio)")
        print("   â€¢ marketing-mode (estrategia marketing)")
        print("   â€¢ frontend-design (experiencia usuario)")
        
        output_file = crear_documento_treqe()
        file_size = os.path.getsize(output_file)
        
        print(f"\nâœ… Â¡DOCUMENTO CREADO EXITOSAMENTE!")
        print(f"ðŸ“„ Archivo: {output_file}")
        print(f"ðŸ“ TamaÃ±o: {file_size:,} bytes")
        print(f"ðŸ“… Fecha: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
        print("\nðŸŽ¯ Documento incluye 10 partes completas:")
        print("   1. Resumen ejecutivo")
        print("   2. Problema real (caso Ana, Carlos, Beatriz)")
        print("   3. SoluciÃ³n Treqe")
        print("   4. Modelo de negocio (Business Model Canvas)")
        print("   5. Proyecciones financieras")
        print("   6. Cronograma 12 meses")
        print("   7. Aspectos legales")
        print("   8. DiseÃ±o y experiencia")
        print("   9. Riesgos y contingencias")
        print("   10. Estrategia de marketing")
        print("   11. Checklists ejecutables")
        print("   12. ConclusiÃ³n")
        print("\nðŸ“§ Â¡Documento listo para presentar!")
        
    except Exception as e:
        print(f"\nâŒ Error: {e}")
        import traceback
        traceback.print_exc()
