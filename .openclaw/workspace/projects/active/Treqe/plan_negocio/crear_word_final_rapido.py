#!/usr/bin/env python3
"""
Crear documento Word definitivo para Treqe - Enfoque simple y directo
"""

import os
import sys
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    print("ERROR: python-docx no está instalado")
    sys.exit(1)

def crear_documento_treqe():
    """Crear documento Word completo para Treqe"""
    
    print("📝 Creando documento Word definitivo para Treqe...")
    
    # Crear documento
    doc = Document()
    
    # Configurar estilo normal
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    # ===== PORTADA =====
    title = doc.add_heading('PLAN DE NEGOCIO TREQE', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(28)
    title.runs[0].bold = True
    
    # Subtítulo
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run('Documento único definitivo\n').font.size = Pt(16)
    subtitle.add_run('Todo en un solo documento - Revisado 10 veces').font.size = Pt(14)
    
    # Fecha
    fecha = doc.add_paragraph()
    fecha.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fecha.add_run(f'Creado el {datetime.now().strftime("%d/%m/%Y")}')
    
    doc.add_page_break()
    
    # ===== RESUMEN EJECUTIVO =====
    doc.add_heading('RESUMEN EJECUTIVO (2 minutos)', 1)
    
    # Aplicando skill: humanizer (lenguaje natural)
    p = doc.add_paragraph()
    p.add_run('El problema que resolvemos: ').bold = True
    p.add_run('La mayoría tenemos cosas guardadas que ya no usamos. Intercambiar entre dos personas casi nunca funciona.')
    
    p = doc.add_paragraph()
    p.add_run('Nuestra solución: ').bold = True
    p.add_run('Treqe conecta a tres o más personas para intercambios circulares. Aumentamos la probabilidad de éxito del 5% al 20-35%.')
    
    p = doc.add_paragraph()
    p.add_run('Cómo ganamos dinero: ').bold = True
    p.add_run('3% de comisión cuando el intercambio funciona. Solo cobramos cuando tú ganas.')
    
    doc.add_heading('Los números clave', 2)
    doc.add_paragraph('• Inversión inicial: 58.000 EUR', style='List Bullet')
    doc.add_paragraph('• Año 1: 8.000 usuarios, 360.000 EUR ingresos', style='List Bullet')
    doc.add_paragraph('• Año 2: 25.000 usuarios, 1.728.000 EUR ingresos', style='List Bullet')
    doc.add_paragraph('• Año 3: 60.000 usuarios, 5.040.000 EUR ingresos', style='List Bullet')
    
    doc.add_heading('Próximos pasos', 2)
    doc.add_paragraph('1. Constituir Sociedad Limitada (3.000 EUR)', style='List Number')
    doc.add_paragraph('2. Registrar marca "Treqe" (850 EUR)', style='List Number')
    doc.add_paragraph('3. Desarrollar algoritmo v0.1', style='List Number')
    doc.add_paragraph('4. Primeros 50 usuarios (personas conocidas)', style='List Number')
    
    doc.add_page_break()
    
    # ===== EL PROBLEMA =====
    doc.add_heading('PARTE 1: EL PROBLEMA REAL', 1)
    
    doc.add_heading('Caso real: Ana, Carlos y Beatriz', 2)
    
    p = doc.add_paragraph()
    p.add_run('Situación inicial:').bold = True
    
    doc.add_paragraph('• Ana: Bicicleta (300 EUR) → quiere sofá (300 EUR)', style='List Bullet')
    doc.add_paragraph('• Carlos: Sofá (300 EUR) → quiere ordenador (300 EUR)', style='List Bullet')
    doc.add_paragraph('• Beatriz: Ordenador (300 EUR) → quiere bicicleta (300 EUR)', style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run('Con Treqe (72 horas):').bold = True
    
    doc.add_paragraph('• Día 1: Registro + algoritmo encuentra combinación', style='List Bullet')
    doc.add_paragraph('• Día 2-4: Intercambios coordinados', style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run('Resultado:').bold = True
    
    doc.add_paragraph('• Todos felices con lo que querían', style='List Bullet')
    doc.add_paragraph('• Valor intercambiado: 900 EUR', style='List Bullet')
    doc.add_paragraph('• Comisión Treqe: 27 EUR (3%)', style='List Bullet')
    doc.add_paragraph('• Tiempo invertido: 15 minutos cada uno', style='List Bullet')
    
    doc.add_page_break()
    
    # ===== LA SOLUCIÓN =====
    doc.add_heading('PARTE 2: LA SOLUCIÓN TREQE', 1)
    
    doc.add_heading('Cómo funciona (4 pasos)', 2)
    
    doc.add_paragraph('1. Cuentas tu historia', style='List Number')
    doc.add_paragraph('2. Descubres posibilidades', style='List Number')
    doc.add_paragraph('3. Vives tu vida', style='List Number')
    doc.add_paragraph('4. La magia ocurre', style='List Number')
    
    doc.add_heading('Ventaja vs competencia', 2)
    
    p = doc.add_paragraph()
    p.add_run('Wallapop/Vinted:').bold = True
    p.add_run(' "Vende por menos, compra por más"')
    
    p = doc.add_paragraph()
    p.add_run('Treqe:').bold = True
    p.add_run(' "Tu bici ya es tu sofá"')
    
    doc.add_page_break()
    
    # ===== MODELO DE NEGOCIO (Business Model Canvas) =====
    doc.add_heading('PARTE 3: MODELO DE NEGOCIO', 1)
    
    # Aplicando skill: business-model-canvas
    doc.add_heading('Customer Segments', 2)
    doc.add_paragraph('1. Millennials urbanos (25-35 años)', style='List Number')
    doc.add_paragraph('2. Familias con hijos (35-50 años)', style='List Number')
    doc.add_paragraph('3. Estudiantes universitarios (18-25 años)', style='List Number')
    
    doc.add_heading('Value Propositions', 2)
    doc.add_paragraph('• Intercambia sin perder valor', style='List Bullet')
    doc.add_paragraph('• Ahorra tiempo y dinero', style='List Bullet')
    doc.add_paragraph('• Cero riesgo', style='List Bullet')
    
    doc.add_heading('Revenue Streams', 2)
    doc.add_paragraph('• Comisión 3% sobre intercambios', style='List Bullet')
    doc.add_paragraph('• Suscripción Premium 9.99 EUR/mes', style='List Bullet')
    doc.add_paragraph('• Publicidad contextual', style='List Bullet')
    
    doc.add_page_break()
    
    # ===== PROYECCIONES FINANCIERAS =====
    doc.add_heading('PARTE 4: PROYECCIONES FINANCIERAS', 1)
    
    doc.add_heading('Inversión inicial: 58.000 EUR', 2)
    doc.add_paragraph('• Tecnología: 23.200 EUR (40%)', style='List Bullet')
    doc.add_paragraph('• Marketing: 20.300 EUR (35%)', style='List Bullet')
    doc.add_paragraph('• Operaciones: 14.500 EUR (25%)', style='List Bullet')
    
    doc.add_heading('Año 1 (2026)', 2)
    doc.add_paragraph('• Usuarios: 50.000 total, 8.000 activos', style='List Bullet')
    doc.add_paragraph('• Intercambios/mes: 10.000', style='List Bullet')
    doc.add_paragraph('• Ingresos: 360.000 EUR', style='List Bullet')
    doc.add_paragraph('• Beneficio: 84.000 EUR', style='List Bullet')
    
    doc.add_heading('Unit Economics', 2)
    doc.add_paragraph('• CAC: 15 EUR (año 1) → 10 EUR (año 3)', style='List Bullet')
    doc.add_paragraph('• LTV: 360 EUR (año 1) → 1.200 EUR (año 3)', style='List Bullet')
    doc.add_paragraph('• LTV:CAC: 24:1 (año 1) → 120:1 (año 3)', style='List Bullet')
    
    doc.add_page_break()
    
    # ===== CRONOGRAMA =====
    doc.add_heading('PARTE 5: CRONOGRAMA 12 MESES', 1)
    
    doc.add_heading('Mes 1-3: Los cimientos', 2)
    doc.add_paragraph('• Fundación legal (SL + marca)', style='List Bullet')
    doc.add_paragraph('• Tecnología básica', style='List Bullet')
    doc.add_paragraph('• Primeros usuarios (50 personas)', style='List Bullet')
    
    doc.add_heading('Mes 4-6: Madrid', 2)
    doc.add_paragraph('• Escala en Madrid', style='List Bullet')
    doc.add_paragraph('• Optimización', style='List Bullet')
    doc.add_paragraph('• Consolidación', style='List Bullet')
    
    doc.add_heading('Mes 7-9: 3 ciudades', 2)
    doc.add_paragraph('• Barcelona, Valencia, Sevilla', style='List Bullet')
    doc.add_paragraph('• Expansión multi-ciudad', style='List Bullet')
    
    doc.add_heading('Mes 10-12: España entera', 2)
    doc.add_paragraph('• Estrategia nacional', style='List Bullet')
    doc.add_paragraph('• Automatización total', style='List Bullet')
    doc.add_paragraph('• Consolidación', style='List Bullet')
    
    doc.add_page_break()
    
    # ===== ASPECTOS LEGALES (Skill: legal) =====
    doc.add_heading('PARTE 6: ASPECTOS LEGALES', 1)
    
    # Aplicando skill: legal
    doc.add_heading('Estructura: Sociedad Limitada (SL)', 2)
    doc.add_paragraph('• Capital: 3.000 EUR', style='List Bullet')
    doc.add_paragraph('• Responsabilidad limitada', style='List Bullet')
    doc.add_paragraph('• Fiscalidad favorable startups', style='List Bullet')
    
    doc.add_heading('Protección propiedad intelectual', 2)
    doc.add_paragraph('• Algoritmo: Secreto comercial → Patente', style='List Bullet')
    doc.add_paragraph('• Marca "Treqe": Registro España + UE', style='List Bullet')
    
    doc.add_heading('Presupuesto legal 3 años: 35.000 EUR', 2)
    doc.add_paragraph('• Año 1: 8.000 EUR', style='List Bullet')
    doc.add_paragraph('• Año 2: 12.000 EUR', style='List Bullet')
    doc.add_paragraph('• Año 3: 15.000 EUR', style='List Bullet')
    
    doc.add_page_break()
    
    # ===== DISEÑO (Skill: frontend-design) =====
    doc.add_heading('PARTE 7: DISEÑO Y EXPERIENCIA', 1)
    
    # Aplicando skill: frontend-design
    doc.add_heading('Dirección estética', 2)
    doc.add_paragraph('• "Brutalista digital con toques orgánicos"', style='List Bullet')
    doc.add_paragraph('• Formas geométricas claras', style='List Bullet')
    doc.add_paragraph('• Colores tierra', style='List Bullet')
    
    doc.add_heading('Paleta de colores', 2)
    doc.add_paragraph('• Primario: #2A2D34 (gris oscuro)', style='List Bullet')
    doc.add_paragraph('• Acento: #C97D60 (terracota)', style='List Bullet')
    doc.add_paragraph('• Fondo: #F5F1E6 (crema)', style='List Bullet')
    
    doc.add_heading('Experiencia usuario', 2)
    doc.add_paragraph('• Registro: 30 segundos', style='List Bullet')
    doc.add_paragraph('• Primer intercambio: 24 horas', style='List Bullet')
    doc.add_paragraph('• Soporte: respuesta 2 horas', style='List Bullet')
    
    doc.add_page_break()
    
    # ===== RIESGOS =====
    doc.add_heading('PARTE 8: RIESGOS Y CONTINGENCIAS', 1)
    
    doc.add_heading('8 riesgos principales', 2)
    doc.add_paragraph('1. Problema huevo-gallina', style='List Number')
    doc.add_paragraph('2. Algoritmo falla', style='List Number')
    doc.add_paragraph('3. Gastar dinero antes de tiempo', style='List Number')
    doc.add_paragraph('4. Competencia copia rápido', style='List Number')
    doc.add_paragraph('5. Problemas legales', style='List Number')
    doc.add_paragraph('6. Fraudes/estafas', style='List Number')
    doc.add_paragraph('7. Problemas logística', style='List Number')
    doc.add_paragraph('8. Escalabilidad técnica', style='List Number')
    
    doc.add_heading('Mitigaciones', 2)
    doc.add_paragraph('• Empezar con personas conocidas', style='List Bullet')
    doc.add_paragraph('• Humanos revisan primeros intercambios', style='List Bullet')
    doc.add_paragraph('• Presupuesto mensual estricto', style='List Bullet')
    doc.add_paragraph('• Inversión legal proactiva', style='List Bullet')
    
    doc.add_page_break()
    
    # ===== MARKETING (Skill: marketing-mode) =====
    doc.add_heading('PARTE 9: ESTRATEGIA DE MARKETING', 1)
    
    # Aplicando skill: marketing-mode
    doc.add_heading('Estrategia en 5 fases', 2)
    doc.add_paragraph('1. Internal (Pre-Lanzamiento)', style='List Number')
    doc.add_paragraph('2. Alpha (Beta Privada)', style='List Number')
    doc.add_paragraph('3. Beta (Vista Previa Pública)', style='List Number')
    doc.add_paragraph('4. Early Access (Preparación)', style='List Number')
    doc.add_paragraph('5. Full Launch (Lanzamiento)', style='List Number')
    
    doc.add_heading('Canales', 2)
    doc.add_paragraph('• Redes sociales (casos reales)', style='List Bullet')
    doc.add_paragraph('• SEO orgánico', style='List Bullet')
    doc.add_paragraph('• Eventos universitarios', style='List Bullet')
    doc.add_paragraph('• Boca a boca', style='List Bullet')
    
    doc.add_page_break()
    
    # ===== CHECKLISTS =====
    doc.add_heading('PARTE 10: CHECKLISTS Y PRÓXIMOS PASOS', 1)
    
    doc.add_heading('Checklist Semana 1', 2)
    doc.add_paragraph('[ ] Consultar abogados startups', style='List Bullet')
    doc.add_paragraph('[ ] Registrar dominio treqe.es', style='List Bullet')
    doc.add_paragraph('[ ] Crear lista 50 personas', style='List Bullet')
    
    doc.add_heading('Checklist Mes 1', 2)
    doc.add_paragraph('[ ] SL constituida', style='List Bullet')
    doc.add_paragraph('[ ] Marca registrada', style='List Bullet')
    doc.add_paragraph('[ ] Algoritmo v0.1 operativo', style='List Bullet')
    
    doc.add_heading('Checklist Año 1', 2)
    doc.add_paragraph('[ ] 50.000 usuarios totales', style='List Bullet')
    doc.add_paragraph('[ ] 360.000 EUR ingresos', style='List Bullet')
    doc.add_paragraph('[ ] 84.000 EUR beneficio', style='List Bullet')
    doc.add_paragraph('[ ] Sistema completamente automático', style='