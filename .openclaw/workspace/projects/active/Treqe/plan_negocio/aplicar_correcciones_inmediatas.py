#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script para aplicar correcciones inmediatas identificadas en el documento.
SOLO corrige errores obvios y actualiza listas para que coincidan con los nuevos títulos.
"""

import docx
import os
from datetime import datetime

def aplicar_correcciones_inmediatas():
    """Aplica correcciones críticas identificadas en el análisis."""
    input_path = "Plan_Negocio_Treqe_ACTUALIZADO_2026.docx"
    output_path = "Plan_Negocio_Treqe_CORREGIDO_2026.docx"
    
    if not os.path.exists(input_path):
        print(f"Error: Archivo no encontrado: {input_path}")
        return None
    
    print("Cargando documento para correcciones inmediatas...")
    doc = docx.Document(input_path)
    
    cambios_aplicados = []
    
    # === CORRECCIÓN 1: Párrafo 102 (error tipográfico) ===
    if 102 < len(doc.paragraphs):
        para_102 = doc.paragraphs[102]
        texto_102 = para_102.text.strip()
        
        if "dTreqe" in texto_102:
            nuevo_texto = texto_102.replace("dTreqe", "del sistema")
            para_102.clear()
            para_102.add_run(nuevo_texto)
            cambios_aplicados.append(("102", "Corregido 'dTreqe' → 'del sistema'"))
    
    # === CORRECCIÓN 2: Actualizar lista Fase 4 (102-106) ===
    # La lista actual sigue siendo de chat grupal, no coincide con "Sistema de Ofertas Estructuradas"
    
    nueva_lista_fase4 = [
        "- Examinar la propuesta completa con timeline visual",
        "- Ver fotografías en alta resolución y descripciones validadas", 
        "- Comparar valores de mercado para cada artículo",
        "- Aceptar o rechazar con un solo clic",
        "- Recibir confirmación inmediata y etiquetas de envío"
    ]
    
    # Actualizar párrafos 102-106 (ya 102 corregido arriba, actualizar 103-106)
    for i, nuevo_item in enumerate(nueva_lista_fase4, start=103):
        if i < len(doc.paragraphs):
            para = doc.paragraphs[i]
            texto_actual = para.text.strip()
            
            # Solo actualizar si es un item de lista (comienza con "-")
            if texto_actual.startswith("-"):
                para.clear()
                para.add_run(nuevo_item)
                cambios_aplicados.append((str(i), f"Actualizado item lista Fase 4: {texto_actual[:30]}..."))
    
    # === CORRECCIÓN 3: Actualizar lista Fase 5 (110-113) ===
    # Los puntos actuales no reflejan "Sistema Combinado de Garantías"
    
    nueva_lista_fase5 = [
        "- Depósitos con tarjeta gestionados vía Stripe (retenidos, no cobrados)",
        "- Seguro de envío incluido para artículos >€300 mediante partner fintech",
        "- Sistema de scoring en tiempo real con niveles progresivos",
        "- Resolución de disputas automatizada con escalación opcional"
    ]
    
    # Actualizar párrafos 110-113
    for i, nuevo_item in enumerate(nueva_lista_fase5, start=110):
        if i < len(doc.paragraphs):
            para = doc.paragraphs[i]
            texto_actual = para.text.strip()
            
            if texto_actual.startswith("-"):
                para.clear()
                para.add_run(nuevo_item)
                cambios_aplicados.append((str(i), f"Actualizado item lista Fase 5: {texto_actual[:30]}..."))
    
    # === CORRECCIÓN 4: Asegurar coherencia párrafo 107 ===
    if 107 < len(doc.paragraphs):
        para_107 = doc.paragraphs[107]
        texto_107 = para_107.text.strip()
        
        # El párrafo 107 todavía habla de "negociación" y "chat"
        # Debería hablar de "decisión informada" y "transparencia"
        if "negociación" in texto_107.lower() or "chat" in texto_107.lower():
            nuevo_texto_107 = "Este sistema elimina la fricción de la negociación mientras mantiene total transparencia. Los usuarios pueden tomar decisiones informadas sin presión de tiempo, sabiendo que cada oferta está calculada para ser justa para todos."
            para_107.clear()
            para_107.add_run(nuevo_texto_107)
            cambios_aplicados.append(("107", "Actualizado para coherencia con Sistema de Ofertas"))
    
    # === CORRECCIÓN 5: Asegurar coherencia párrafo 114 ===
    if 114 < len(doc.paragraphs):
        para_114 = doc.paragraphs[114]
        texto_114 = para_114.text.strip()
        
        # El párrafo 114 es muy filosófico, podría ser más concreto
        if "experiencia humana completa" in texto_114:
            nuevo_texto_114 = "La confianza en Treqe emerge de sistemas transparentes y verificables, no de interacciones personales. Cada transacción construye reputación medible que abre puertas a mejores oportunidades de intercambio."
            para_114.clear()
            para_114.add_run(nuevo_texto_114)
            cambios_aplicados.append(("114", "Actualizado para coherencia con Sistema de Garantías"))
    
    # === GUARDAR DOCUMENTO CORREGIDO ===
    print(f"\nGuardando documento corregido: {output_path}")
    doc.save(output_path)
    
    # === GENERAR REPORTE ===
    reporte_path = "reporte_correcciones_inmediatas.txt"
    with open(reporte_path, "w", encoding="utf-8") as f:
        f.write(f"REPORTE DE CORRECCIONES INMEDIATAS\n")
        f.write(f"Fecha: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        f.write(f"Documento original: {input_path}\n")
        f.write(f"Documento corregido: {output_path}\n")
        f.write(f"Cambios aplicados: {len(cambios_aplicados)}\n\n")
        
        f.write("DETALLE DE CAMBIOS:\n")
        f.write("="*60 + "\n")
        
        if cambios_aplicados:
            for num, desc in cambios_aplicados:
                f.write(f"Párrafo {num}: {desc}\n")
        else:
            f.write("No se aplicaron correcciones (¿ya estaban correctos?)\n")
        
        f.write("\n" + "="*60 + "\n")
        f.write("RESUMEN:\n")
        f.write(f"1. Corrección error tipográfico párrafo 102: {'SÍ' if any('102' in c[0] for c in cambios_aplicados) else 'NO'}\n")
        f.write(f"2. Lista Fase 4 actualizada (103-106): {'SÍ' if any(c[0] in ['103','104','105','106'] for c in cambios_aplicados) else 'NO'}\n")
        f.write(f"3. Lista Fase 5 actualizada (110-113): {'SÍ' if any(c[0] in ['110','111','112','113'] for c in cambios_aplicados) else 'NO'}\n")
        f.write(f"4. Coherencia párrafo 107: {'SÍ' if any('107' in c[0] for c in cambios_aplicados) else 'NO'}\n")
        f.write(f"5. Coherencia párrafo 114: {'SÍ' if any('114' in c[0] for c in cambios_aplicados) else 'NO'}\n")
    
    # === MOSTRAR RESUMEN ===
    print("\n" + "="*60)
    print("RESUMEN DE CORRECCIONES APLICADAS")
    print("="*60)
    
    if cambios_aplicados:
        print(f"Total correcciones: {len(cambios_aplicados)}")
        for num, desc in cambios_aplicados:
            print(f"  • Párrafo {num}: {desc}")
    else:
        print("No se aplicaron correcciones (posiblemente ya estaban correctos)")
    
    print("\n" + "="*60)
    print(f"Documento guardado: {output_path}")
    print(f"Reporte detallado: {reporte_path}")
    print("="*60)
    
    return doc, cambios_aplicados

if __name__ == "__main__":
    doc, cambios = aplicar_correcciones_inmediatas()
    
    if cambios:
        print("\n✅ CORRECCIONES INMEDIATAS COMPLETADAS")
        print("\nPRÓXIMOS PASOS RECOMENDADOS:")
        print("1. Revisar documento corregido")
        print("2. Evaluar si se necesitan las expansiones detalladas")
        print("3. Validar coherencia global del documento")
    else:
        print("\n⚠ NO SE APLICARON CORRECCIONES")
        print("(Posiblemente los errores ya estaban corregidos o el documento tiene estructura diferente)")