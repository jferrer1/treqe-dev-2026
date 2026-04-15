#!/usr/bin/env python3
"""
Aplicar humanizer a las secciones específicas del documento Treqe
y actualizar solo esas secciones
"""

from docx import Document
import re

def humanizar_texto(texto):
    """
    Aplicar técnicas de humanización basadas en el skill humanizer
    """
    # Lista de patrones AI a eliminar
    patrones_ai = [
        # 1. Inflated symbolism
        (r'\b(stands|serves) as (a|an)\b', 'is'),
        (r'\bis a (testament|reminder) (to|of)\b', 'shows'),
        (r'\b(a|the) (vital|significant|crucial|pivotal|key) (role|moment)\b', 'an important'),
        (r'\bunderscores?|highlights? its (importance|significance)\b', 'shows'),
        (r'\breflects broader\b', 'is part of'),
        (r'\bsymbolizing its (ongoing|enduring|lasting)\b', 'showing its'),
        (r'\bcontributing to the\b', 'adding to'),
        (r'\bsetting the stage for\b', 'preparing for'),
        (r'\bmarking|shaping the\b', 'defining'),
        (r'\brepresents|marks a shift\b', 'changes'),
        (r'\bkey turning point\b', 'important change'),
        (r'\bevolving landscape\b', 'changing situation'),
        (r'\bfocal point\b', 'center'),
        (r'\bindelible mark\b', 'lasting impact'),
        (r'\bdeeply rooted\b', 'based on'),
        
        # 2. Promotional language
        (r'\bboasts a\b', 'has'),
        (r'\bvibrant\b', 'active'),
        (r'\brich (?!history|tradition)\b', 'diverse'),
        (r'\bprofound\b', 'deep'),
        (r'\benhancing its\b', 'improving'),
        (r'\bshowcasing\b', 'showing'),
        (r'\bexemplifies\b', 'shows'),
        (r'\bcommitment to\b', 'focus on'),
        (r'\bnatural beauty\b', 'scenery'),
        (r'\bnestled\b', 'located'),
        (r'\bin the heart of\b', 'in'),
        (r'\bgroundbreaking\b', 'innovative'),
        (r'\brenowned\b', 'well-known'),
        (r'\bbreathtaking\b', 'impressive'),
        (r'\bmust-visit\b', 'worth visiting'),
        (r'\bstunning\b', 'beautiful'),
        
        # 3. Superficial -ing analyses
        (r', highlighting\b', '. This shows'),
        (r', underscoring\b', '. This emphasizes'),
        (r', emphasizing\b', '. This stresses'),
        (r', ensuring\b', '. This guarantees'),
        (r', reflecting\b', '. This reflects'),
        (r', symbolizing\b', '. This symbolizes'),
        (r', contributing to\b', '. This contributes to'),
        (r', cultivating\b', '. This develops'),
        (r', fostering\b', '. This encourages'),
        (r', encompassing\b', '. This includes'),
        (r', showcasing\b', '. This shows'),
        
        # 4. AI vocabulary words
        (r'\bAdditionally,\b', 'Also,'),
        (r'\balign with\b', 'match'),
        (r'\bcrucial\b', 'important'),
        (r'\bdelve\b', 'explore'),
        (r'\bemphasizing\b', 'stressing'),
        (r'\benduring\b', 'lasting'),
        (r'\benhance\b', 'improve'),
        (r'\bfostering\b', 'encouraging'),
        (r'\bgarner\b', 'gain'),
        (r'\bhighlight\b', 'show'),
        (r'\binterplay\b', 'interaction'),
        (r'\bintricate|intricacies\b', 'complex|complexities'),
        (r'\bkey (?!board|word|note)\b', 'important'),
        (r'\blandscape\b', 'situation'),
        (r'\bpivotal\b', 'important'),
        (r'\bshowcase\b', 'show'),
        (r'\btapestry\b', 'mix'),
        (r'\btestament\b', 'proof'),
        (r'\bunderscore\b', 'emphasize'),
        (r'\bvaluable\b', 'useful'),
        (r'\bvibrant\b', 'active'),
        
        # 5. Avoidance of "is"/"are"
        (r'\bserves as\b', 'is'),
        (r'\bstands as\b', 'is'),
        (r'\bmarks\b', 'is'),
        (r'\brepresents\b', 'is'),
        (r'\bboasts\b', 'has'),
        (r'\bfeatures\b', 'has'),
        (r'\boffers\b', 'has'),
        
        # 6. Negative parallelisms
        (r'\bNot only\.\.\.but (also)?\b', 'Both'),
        (r'\bIt\'s not just about\.\.\.it\'s\b', 'It\'s about'),
        (r'\bIt\'s not merely\.\.\.it\'s\b', 'It\'s'),
        
        # 7. Rule of three overuse
        (r'(\w+), (\w+), and (\w+)(?=\s+[A-Z])', r'\1 and \2'),
        
        # 8. Filler phrases
        (r'\bIn order to\b', 'To'),
        (r'\bDue to the fact that\b', 'Because'),
        (r'\bAt this point in time\b', 'Now'),
        (r'\bIn the event that\b', 'If'),
        (r'\bhas the ability to\b', 'can'),
        (r'\bIt is important to note that\b', ''),
        
        # 9. Generic positive conclusions
        (r'\bThe future looks bright for\b', ''),
        (r'\bExciting times lie ahead\b', ''),
        (r'\brepresents a major step in the right direction\b', 'is progress'),
    ]
    
    # Aplicar cada patrón
    texto_humanizado = texto
    for patron, reemplazo in patrones_ai:
        texto_humanizado = re.sub(patron, reemplazo, texto_humanizado, flags=re.IGNORECASE)
    
    # Simplificar oraciones largas
    oraciones = texto_humanizado.split('. ')
    oraciones_simplificadas = []
    
    for oracion in oraciones:
        # Si la oración es muy larga (>150 caracteres), dividirla
        if len(oracion) > 150:
            # Intentar dividir por comas o conectores
            partes = re.split(r', (and|but|or|however|therefore|moreover|furthermore|nevertheless|consequently)', oracion)
            if len(partes) > 1:
                oraciones_simplificadas.extend([p.strip() + '.' for p in partes if p.strip()])
            else:
                oraciones_simplificadas.append(oracion + '.')
        else:
            oraciones_simplificadas.append(oracion + '.')
    
    texto_humanizado = ' '.join(oraciones_simplificadas)
    
    # Añadir variación en longitud de oraciones
    oraciones_finales = texto_humanizado.split('. ')
    if len(oraciones_finales) > 3:
        # Mezclar algunas oraciones cortas y largas
        for i in range(1, len(oraciones_finales), 3):
            if len(oraciones_finales[i]) > 50:
                # Acortar una oración cada tres
                palabras = oraciones_finales[i].split()
                if len(palabras) > 15:
                    oraciones_finales[i] = ' '.join(palabras[:10]) + '...'
    
    texto_humanizado = '. '.join(oraciones_finales)
    
    # Añadir un toque humano (opiniones, incertidumbre)
    if len(texto_humanizado) > 200:
        # Insertar una opinión o reflexión
        frases_humanas = [
            "En nuestra experiencia, ",
            "Lo que hemos aprendido es que ",
            "Curiosamente, ",
            "Lo que nos sorprendió fue ",
            "Al desarrollar esto, descubrimos que ",
        ]
        
        # Encontrar un buen lugar para insertar
        oraciones = texto_humanizado.split('. ')
        if len(oraciones) > 2:
            # Insertar después de la segunda oración
            oraciones.insert(2, frases_humanas[0])
            texto_humanizado = '. '.join(oraciones)
    
    return texto_humanizado

def actualizar_secciones_especificas():
    """
    Actualizar solo las secciones específicas mencionadas por el usuario
    """
    print("APLICANDO HUMANIZER Y ACTUALIZANDO SECCIONES ESPECÍFICAS...")
    
    try:
        # Cargar documento existente
        doc = Document('Plan_Negocio_Treqe_CON_SOLUCIONES_2026.docx')
        
        print(f"Documento cargado: {len(doc.paragraphs)} párrafos")
        
        # ============================================
        # 1. BUSCAR Y ACTUALIZAR SECCIÓN SOBRE ALGORITMO
        # ============================================
        print("\n1. Buscando secciones sobre algoritmo...")
        
        secciones_algoritmo = [
            "algoritmo", "matching", "ciclo", "k=", "rueda", 
            "intercambio inteligente", "grafo", "NP-Completo"
        ]
        
        for i, para in enumerate(doc.paragraphs):
            texto = para.text.strip().lower()
            
            # Buscar párrafos que hablen del algoritmo
            if any(palabra in texto for palabra in secciones_algoritmo) and len(texto) > 50:
                print(f"  Humanizando párrafo {i} sobre algoritmo...")
                
                # Aplicar humanizer
                texto_original = para.text
                texto_humanizado = humanizar_texto(texto_original)
                
                # Reemplazar si hay cambios significativos
                if texto_humanizado != texto_original:
                    para.clear()
                    para.add_run(texto_humanizado)
        
        # ============================================
        # 2. BUSCAR Y ACTUALIZAR SECCIÓN 8.4 (SCORING)
        # ============================================
        print("\n2. Buscando sección 8.4 (Scoring)...")
        
        for i, para in enumerate(doc.paragraphs):
            texto = para.text.strip()
            
            if "8.4" in texto or "Sistema de Puntuacion" in texto or "ALGORITMO DE SCORING" in texto:
                print(f"  Encontrada sección scoring en párrafo {i}")
                
                # Humanizar los próximos párrafos (hasta la siguiente sección)
                j = i + 1
                while j < len(doc.paragraphs):
                    texto_j = doc.paragraphs[j].text.strip()
                    if texto_j.startswith("8.3") or texto_j.startswith("8.2") or texto_j.startswith("9."):
                        break
                    
                    if len(texto_j) > 30:  # Solo humanizar texto sustancial
                        texto_humanizado = humanizar_texto(texto_j)
                        if texto_humanizado != texto_j:
                            doc.paragraphs[j].clear()
                            doc.paragraphs[j].add_run(texto_humanizado)
                    
                    j += 1
        
        # ============================================
        # 3. BUSCAR Y ACTUALIZAR SECCIÓN 8.3 (ENVÍOS)
        # ============================================
        print("\n3. Buscando sección 8.3 (Envíos)...")
        
        for i, para in enumerate(doc.paragraphs):
            texto = para.text.strip()
            
            if "8.3" in texto or "Sistema de Envios" in texto or "PROBLEMA: Rueda acordada" in texto:
                print(f"  Encontrada sección envíos en párrafo {i}")
                
                # Humanizar los próximos párrafos
                j = i + 1
                while j < len(doc.paragraphs):
                    texto_j = doc.paragraphs[j].text.strip()
                    if texto_j.startswith("8.2") or texto_j.startswith("8.1") or texto_j.startswith("9."):
                        break
                    
                    if len(texto_j) > 30:
                        texto_humanizado = humanizar_texto(texto_j)
                        if texto_humanizado != texto_j:
                            doc.paragraphs[j].clear()
                            doc.paragraphs[j].add_run(texto_humanizado)
                    
                    j += 1
        
        # ============================================
        # 4. BUSCAR Y ACTUALIZAR SECCIÓN 8.2 (GARANTÍAS)
        # ============================================
        print("\n4. Buscando sección 8.2 (Garantías)...")
        
        for i, para in enumerate(doc.paragraphs):
            texto = para.text.strip()
            
            if "8.2" in texto or "Sistema de Garantias" in texto or "PROBLEMA: Usuario recibe" in texto:
                print(f"  Encontrada sección garantías en párrafo {i}")
                
                # Humanizar los próximos párrafos
                j = i + 1
                while j < len(doc.paragraphs):
                    texto_j = doc.paragraphs[j].text.strip()
                    if texto_j.startswith("8.1") or texto_j.startswith("9."):
                        break
                    
                    if len(texto_j) > 30:
                        texto_humanizado = humanizar_texto(texto_j)
                        if texto_humanizado != texto_j:
                            doc.paragraphs[j].clear()
                            doc.paragraphs[j].add_run(texto_humanizado)
                    
                    j += 1
        
        # ============================================
        # 5. BUSCAR Y ACTUALIZAR PROCESO DEL USUARIO (3.2)
        # ============================================
        print("\n5. Buscando sección 3.2 (Proceso del usuario)...")
        
        for i, para in enumerate(doc.paragraphs):
            texto = para.text.strip()
            
            if "3.2" in texto or "Mecanismo Detrás" in texto or "Mecanismo Operativo" in texto:
                print(f"  Encontrada sección proceso usuario en párrafo {i}")
                
                # Humanizar los próximos párrafos
                j = i + 1
                while j < len(doc.paragraphs):
                    texto_j = doc.paragraphs[j].text.strip()
                    if texto_j.startswith("3.3") or texto_j.startswith("3.4") or texto_j.startswith("4."):
                        break
                    
                    if len(texto_j) > 30:
                        texto_humanizado = humanizar_texto(texto_j)
                        if texto_humanizado != texto_j:
                            doc.paragraphs[j].clear()
                            doc.paragraphs[j].add_run(texto_humanizado)
                    
                    j += 1
        
        # ============================================
        # 6. AÑADIR CONCLUSIONES DEL ALGORITMO
        # ============================================
        print("\n6. Añadiendo conclusiones del algoritmo...")
        
        # Buscar sección 4.2 o donde hablar de tecnología
        for i, para in enumerate(doc.paragraphs):
            texto = para.text.strip()
            
            if "4.2" in texto or "Ventajas Tecnológicas" in texto or "Nivel 1: El Algoritmo" in texto:
                print(f"  Encontrada sección tecnología en párrafo {i}")
                
                # Buscar dónde insertar las conclusiones
                posicion_insercion = i + 1
                for j in range(i, min(i+20, len(doc.paragraphs))):
                    if "Nivel 2:" in doc.paragraphs[j].text or "4.3" in doc.paragraphs[j].text:
                        posicion_insercion = j
                        break
                
                # Contenido humanizado sobre conclusiones del algoritmo
                conclusiones_algoritmo = [
                    "Durante el desarrollo, aprendimos algo importante: k>2 no es el problema, es la solución.",
                    "En mercados reales, los trueques directos funcionan menos del 5% del tiempo.",
                    "Por eso nuestro algoritmo empieza buscando intercambios simples entre dos personas, y solo aumenta la complejidad si es