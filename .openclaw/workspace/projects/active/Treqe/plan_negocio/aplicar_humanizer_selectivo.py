#!/usr/bin/env python3
"""
Aplicar humanizer selectivamente al documento existente
"""

from docx import Document
import re

def es_texto_sintetico(texto):
    """Detectar si el texto es demasiado sintético/AI"""
    indicadores = [
        # Vocabulario AI
        r'\b(adicionalmente|profundizar|intrincado|pivotal|destacar|subrayar|vibrante|rico|innovador)\b',
        # Construcciones infladas
        r'(sirve como|se erige como|marca un hito|representa un|cuenta con|ofrece)',
        # Lenguaje corporativo excesivo
        r'\b(sinergia|paradigma|ecosistema|soluciones|optimización|eficientizar)\b',
        # Listas excesivamente estructuradas
        r'^\d+\.\s+[A-Z]',
        # Checkmarks y emojis de marketing
        r'[✅🔍🎯📊🚀💡📈🔧🎨⚖️📋]',
    ]
    
    for patron in indicadores:
        if re.search(patron, texto, re.IGNORECASE):
            return True
    
    # Texto muy corto o muy largo sin personalidad
    palabras = texto.split()
    if len(palabras) < 5 or len(palabras) > 50:
        return False  # No juzgar por longitud sola
    
    # Falta de lenguaje personal
    if not any(palabra in texto.lower() for palabra in ['creo', 'pensamos', 'creemos', 'quizás', 'probablemente', 'la verdad']):
        if len(palabras) > 20:  # Textos largos sin personalidad son sospechosos
            return True
    
    return False

def humanizar_texto(texto):
    """Aplicar humanizer de verdad"""
    
    # Si es muy corto, dejarlo como está
    if len(texto.split()) < 10:
        return texto
    
    # Transformaciones específicas
    transformaciones = [
        # Vocabulario AI -> natural
        (r'\bAdicionalmente\b', 'Además'),
        (r'\bprofundizar en\b', 'analizar'),
        (r'\bintrincado\b', 'complejo'),
        (r'\bpivotal\b', 'importante'),
        (r'\bdestacar que\b', 'señalar que'),
        (r'\bsubrayar la importancia\b', 'indicar que es importante'),
        
        # Construcciones infladas -> simples
        (r'sirve como', 'es'),
        (r'se erige como', 'es'),
        (r'marca un hito', 'es importante'),
        (r'representa un', 'es un'),
        (r'cuenta con', 'tiene'),
        (r'ofrece', 'tiene'),
        
        # Lenguaje corporativo -> natural
        (r'\bsinergia\b', 'colaboración'),
        (r'\bparadigma\b', 'forma de hacer las cosas'),
        (r'\becosistema\b', 'conjunto'),
        (r'\bsoluciones\b', 'formas de resolverlo'),
        (r'\boptimización\b', 'mejora'),
        (r'\beficientizar\b', 'hacer más eficiente'),
        
        # Checkmarks y emojis -> texto
        (r'✅', '✓ '),
        (r'🔍', '(importante) '),
        (r'🎯', '(objetivo) '),
        (r'📊', '(datos) '),
        (r'🚀', '(rápido) '),
        (r'💡', '(idea) '),
        (r'📈', '(crecimiento) '),
        (r'🔧', '(herramienta) '),
        (r'🎨', '(diseño) '),
        (r'⚖️', '(legal) '),
        (r'📋', '(lista) '),
    ]
    
    for patron, reemplazo in transformaciones:
        texto = re.sub(patron, reemplazo, texto, flags=re.IGNORECASE)
    
    # Añadir algo de personalidad ocasionalmente
    import random
    if random.random() > 0.7 and len(texto.split()) > 15:
        frases_personales = [
            ' La verdad es que ',
            ' Creemos que ',
            ' En nuestra experiencia, ',
            ' Lo que hemos visto es que ',
            ' Probablemente ',
        ]
        
        # Encontrar un buen lugar para insertar
        palabras = texto.split()
        if len(palabras) > 10:
            posicion = len(palabras) // 2
            # Buscar punto o coma cerca
            texto_mod = ' '.join(palabras[:posicion]) + random.choice(frases_personales) + ' '.join(palabras[posicion:])
            texto = texto_mod
    
    # Simplificar oraciones muy largas
    if len(texto.split()) > 40:
        # Buscar conjunciones para dividir
        for conjuncion in [' y ', ' pero ', ' aunque ', ' sin embargo ', ' además ']:
            if conjuncion in texto:
                partes = texto.split(conjuncion, 1)
                if len(partes[0].split()) > 15 and len(partes[1].split()) > 15:
                    texto = partes[0] + '.' + conjuncion + partes[1]
                    break
    
    return texto

def procesar_documento():
    print("Aplicando humanizer selectivo al documento existente...")
    
    # Cargar documento original
    doc = Document('Plan_Negocio_Treqe_MARKETING_MEJORADO.docx')
    print(f"Documento cargado: {len(doc.paragraphs)} párrafos")
    
    # Identificar secciones que necesitan humanización
    secciones_a_humanizar = []
    
    for i, para in enumerate(doc.paragraphs):
        texto = para.text.strip()
        
        # Sección 5: Modelo de negocio
        if '5. MODELO DE NEGOCIO' in texto:
            secciones_a_humanizar.append(('modelo_negocio', i))
        
        # Sección 10: Marketing
        if '10. ESTRATEGIA DE MARKETING' in texto:
            secciones_a_humanizar.append(('marketing', i))
        
        # Sección 11: Legal
        if '11. ASPECTOS LEGALES' in texto:
            secciones_a_humanizar.append(('legal', i))
        
        # Anexos
        if 'ANEXOS' in texto and len(texto) < 20:  # Solo el título, no contenido
            secciones_a_humanizar.append(('anexos', i))
    
    print(f"Secciones identificadas para humanizar: {len(secciones_a_humanizar)}")
    
    # Procesar cada sección
    cambios_realizados = 0
    
    for seccion, inicio in secciones_a_humanizar:
        print(f"\nProcesando sección: {seccion} (inicio párrafo {inicio})")
        
        # Determinar fin de la sección
        fin = len(doc.paragraphs) - 1
        for j in range(inicio + 1, len(doc.paragraphs)):
            texto_j = doc.paragraphs[j].text.strip()
            # Buscar siguiente sección numerada
            if re.match(r'^\d+\.', texto_j) and j > inicio + 5:
                fin = j - 1
                break
            # O ANEXOS si es sección 11
            if 'ANEXOS' in texto_j and j > inicio + 5 and seccion == 'legal':
                fin = j - 1
                break
        
        print(f"  Sección va del párrafo {inicio} al {fin}")
        
        # Procesar párrafos de esta sección
        for j in range(inicio, fin + 1):
            if j < len(doc.paragraphs):
                texto_original = doc.paragraphs[j].text
                
                # Solo procesar si no es título principal y es sintético
                if j > inicio and es_texto_sintetico(texto_original):
                    texto_humanizado = humanizar_texto(texto_original)
                    
                    if texto_humanizado != texto_original:
                        # Reemplazar texto
                        doc.paragraphs[j].text = texto_humanizado
                        cambios_realizados += 1
                        
                        if cambios_realizados <= 3:  # Mostrar primeros ejemplos
                            print(f"  Ejemplo cambio {cambios_realizados}:")
                            print(f"    Original: {texto_original[:80]}...")
                            print(f"    Humanizado: {texto_humanizado[:80]}...")
    
    # Guardar documento humanizado
    output_path = 'Plan_Negocio_Treqe_HUMANIZADO_SELECTIVO.docx'
    doc.save(output_path)
    
    print(f"\n✅ Documento guardado: {output_path}")
    print(f"✅ Cambios realizados: {cambios_realizados}")
    print("\nMétodo: Humanizer aplicado selectivamente a:")
    print("1. Sección 5: Modelo de negocio")
    print("2. Sección 10: Marketing")
    print("3. Sección 11: Legal")
    print("4. Anexos")
    print("\nSe mantuvo todo el contenido original, solo se humanizó lo sintético.")
    
    return output_path

if __name__ == '__main__':
    procesar_documento()