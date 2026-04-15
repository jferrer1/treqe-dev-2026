#!/usr/bin/env python3
"""
Crear documento final verdaderamente humanizado
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def agregar_titulo(doc, texto, nivel=1):
    """Agregar titulo"""
    if nivel == 1:
        p = doc.add_heading(texto, level=0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif nivel == 2:
        p = doc.add_heading(texto, level=1)
    else:
        p = doc.add_heading(texto, level=2)
    return p

def agregar_texto(doc, texto):
    """Agregar texto normal"""
    if texto.strip():
        # Mantener saltos de linea naturales
        lineas = texto.split('\n')
        for linea in lineas:
            if linea.strip():
                p = doc.add_paragraph(linea.strip())
                p.style = 'Normal'
    return True

def leer_contenido(archivo):
    """Leer contenido humanizado"""
    with open(archivo, 'r', encoding='utf-8') as f:
        return f.read()

def crear_documento_humanizado():
    print("Creando documento verdaderamente humanizado...")
    
    # Crear documento
    doc = Document()
    
    # Estilo basico
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # PORTADA
    print("1. Portada...")
    agregar_titulo(doc, "PLAN DE NEGOCIO", 1)
    agregar_titulo(doc, "TREQE", 1)
    agregar_titulo(doc, "Cuando el trueque por fin funciona", 2)
    
    doc.add_paragraph()
    agregar_texto(doc, "Febrero 2026")
    agregar_texto(doc, "Para: Gente que cree que las cosas pueden ser más simples")
    agregar_texto(doc, "De: Carlos, Ana y Javier (los que lo estamos haciendo)")
    
    doc.add_page_break()
    
    # INDICE
    print("2. Índice...")
    agregar_titulo(doc, "Lo que vas a encontrar aquí", 1)
    
    indice = [
        "1. La idea simple (que nadie había hecho)",
        "2. El problema que todos tenemos",
        "3. Cómo lo resolvemos (sin ser magos)",
        "4. Por qué podemos hacerlo (y otros no)",
        "5. Cómo ganamos dinero (sin que duela)",
        "6. Los números (los aburridos pero necesarios)",
        "7. Quiénes somos (y por qué nos importa)",
        "8. Lo que podría salir mal (somos realistas)",
        "9. Lo que haremos si sale bien",
        "10. Cómo daremos a conocer esto",
        "11. Lo aburrido pero importante (legal)",
        "Lo que no cabía en el resto (anexos)"
    ]
    
    for item in indice:
        agregar_texto(doc, f"• {item}")
    
    doc.add_page_break()
    
    # SECCION 1-4 (basicas)
    print("3. Secciones 1-4...")
    
    secciones_iniciales = [
        ("1. La idea simple (que nadie había hecho)",
         "Todo el mundo tiene cosas que ya no usa. Y cosas que quiere. El problema es que casi nunca coinciden.\n\n" +
         "Ana quiere la bici de Beatriz. Beatriz quiere el ordenador de Carlos. Carlos quiere el sofá de Ana. Y así se queda todo parado.\n\n" +
         "Nosotros encontramos esos círculos. Eso es todo."),
        
        ("2. El problema que todos tenemos",
         "Has intentado intercambiar algo alguna vez, ¿verdad? En Wallapop, en grupos de Facebook, entre amigos.\n\n" +
         "Funciona así: publicas lo que tienes. Esperas. A lo mejor alguien está interesado. Chateáis. Negociáis. Y al final, casi nunca cuadra.\n\n" +
         "No es que la gente no quiera intercambiar. Es que las herramientas actuales no sirven para intercambiar."),
        
        ("3. Cómo lo resolvemos (sin ser magos)",
         "No usamos magia. Usamos matemáticas. Matemáticas simples, la verdad.\n\n" +
         "Tú nos dices: \"Tengo esto, quiero esto otro.\" Nosotros buscamos a más personas como tú. Personas que tienen lo que tú quieres, y quieren lo que otros tienen.\n\n" +
         "Cuando encontramos un círculo que funciona (3, 4, 5 personas), os lo proponemos. Decidís si os gusta. Si a todos os gusta, intercambiáis."),
        
        ("4. Por qué podemos hacerlo (y otros no)",
         "Wallapop no hace esto. Vinted no hace esto. Facebook Marketplace no hace esto.\n\n" +
         "No es que sean tontos. Es que están ocupados siendo Wallapop, Vinted y Facebook Marketplace.\n\n" +
         "Nosotros solo hacemos esto. Buscar círculos de intercambio. Nada más. Y creemos que hacer solo una cosa, pero hacerla bien, es suficiente.")
    ]
    
    for titulo, contenido in secciones_iniciales:
        agregar_titulo(doc, titulo, 2)
        agregar_texto(doc, contenido)
        doc.add_paragraph()
    
    # SECCION 5: MODELO DE NEGOCIO HUMANIZADO
    print("4. Sección 5: Modelo de negocio humanizado...")
    contenido_modelo = leer_contenido('modelo_negocio_verdaderamente_humanizado.md')
    
    # Extraer titulo principal
    lineas = contenido_modelo.split('\n')
    titulo_principal = lineas[0] if lineas else "5. CÓMO GANAMOS DINERO (Y POR QUÉ FUNCIONA)"
    
    agregar_titulo(doc, titulo_principal, 1)
    agregar_texto(doc, '\n'.join(lineas[1:]))
    
    doc.add_page_break()
    
    # SECCION 6-9 (continuacion basica)
    print("5. Secciones 6-9...")
    
    secciones_continuacion = [
        ("6. Los números (los aburridos pero necesarios)",
         "Año 1: Perdemos dinero. A propósito. Gastamos 114.000€ más de lo que ingresamos.\n\n" +
         "Año 2: Empezamos a ganar. 60.000 usuarios, 162.000€ de beneficio.\n\n" +
         "Año 3: Si todo va bien. 150.000 usuarios, 540.000€ de beneficio.\n\n" +
         "No son promesas. Son objetivos. Y sabemos que los objetivos a veces no se cumplen."),
        
        ("7. Quiénes somos (y por qué nos importa)",
         "Carlos: 8 años haciendo marketplaces que casi funcionan. Aprendió más de los fracasos que de los éxitos.\n\n" +
         "Ana: PhD en algoritmos. Podría estar en Google. Prefiere resolver problemas reales.\n\n" +
         "Javier: 12 años en marketing. Cree que el mejor marketing es hacer algo que la gente quiera contar.\n\n" +
         "Nos importa porque creemos que el trueque puede funcionar mejor. Y porque nos gusta resolver problemas difíciles."),
        
        ("8. Lo que podría salir mal (somos realistas)",
         "1. Nadie viene. Mitigación: Empezar con gente que conocemos.\n" +
         "2. El algoritmo falla. Mitigación: Ana está obsesionada con probarlo.\n" +
         "3. Nos quedamos sin dinero. Mitigación: Parar a los 12 meses si no vemos luz.\n" +
         "4. Los grandes nos copian. Mitigación: Significaría que teníamos razón."),
        
        ("9. Lo que haremos si sale bien",
         "Si funciona en Barcelona y Madrid, iremos a otras ciudades.\n\n" +
         "Si funciona en España, miraremos otros países.\n\n" +
         "Si realmente funciona, quizás algún grande quiera comprarnos. O quizás sigamos solos, haciendo crecer esto.\n\n" +
         "No lo sabemos. Y está bien no saberlo.")
    ]
    
    for titulo, contenido in secciones_continuacion:
        agregar_titulo(doc, titulo, 2)
        agregar_texto(doc, contenido)
        doc.add_paragraph()
    
    # SECCION 10: MARKETING HUMANIZADO
    print("6. Sección 10: Marketing humanizado...")
    contenido_marketing = leer_contenido('marketing_verdaderamente_humanizado.md')
    
    lineas_marketing = contenido_marketing.split('\n')
    titulo_marketing = lineas_marketing[0] if lineas_marketing else "10. CÓMO DAREMOS A CONOCER ESTO"
    
    agregar_titulo(doc, titulo_marketing, 1)
    agregar_texto(doc, '\n'.join(lineas_marketing[1:]))
    
    doc.add_page_break()
    
    # SECCION 11: LEGAL HUMANIZADO
    print("7. Sección 11: Legal humanizado...")
    contenido_legal = leer_contenido('aspectos_legales_verdaderamente_humanizados.md')
    
    lineas_legal = contenido_legal.split('\n')
    titulo_legal = lineas_legal[0] if lineas_legal else "11. LO ABURRIDO PERO IMPORTANTE (LEGAL)"
    
    agregar_titulo(doc, titulo_legal, 1)
    agregar_texto(doc, '\n'.join(lineas_legal[1:]))
    
    doc.add_page_break()
    
    # ANEXOS HUMANIZADOS
    print("8. Anexos humanizados...")
    contenido_anexos = leer_contenido('anexos_verdaderamente_humanizados.md')
    
    lineas_anexos = contenido_anexos.split('\n')
    titulo_anexos = lineas_anexos[0] if lineas_anexos else "ANEXOS (LO QUE NO CABÍA EN EL RESTO)"
    
    agregar_titulo(doc, titulo_anexos, 1)
    agregar_texto(doc, '\n'.join(lineas_anexos[1:]))
    
    # Guardar
    output_path = 'Plan_Negocio_Treqe_VERDADERAMENTE_HUMANIZADO.docx'
    doc.save(output_path)
    
    print(f"Documento creado: {output_path}")
    print("Este documento aplica humanizer de verdad:")
    print("- Tiene opiniones, no solo datos")
    print("- Reconoce incertidumbre")
    print("- Usa lenguaje natural, no corporativo")
    print("- Cuenta una historia, no solo lista hechos")
    print("- Es honesto sobre lo que no sabemos")
    
    return output_path

if __name__ == '__main__':
    crear_documento_humanizado()