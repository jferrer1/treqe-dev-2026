#!/usr/bin/env python3
"""
Crear wireframes con skill 'frontend-design' para anexos del plan
"""

from docx import Document

def crear_wireframes_treqe():
    """Crear wireframes para plataforma Treqe"""
    print("Creando wireframes con skill 'frontend-design'...")
    
    wireframes = """
## ANEXO A: WIREFRAMES Y DISEÑO UI/UX PLATAFORMA TREQE

### 🎨 DIRECCIÓN ESTÉTICA ELEGIDA: "BRUTALISTA DIGITAL CON TOQUES ORGÁNICOS"

**Concepto:** Interfaz que combina la crudeza del brutalismo (formas geométricas, tipografía fuerte) con elementos orgánicos (colores tierra, texturas naturales) para transmitir confianza + innovación.

**Paleta de colores:**
- Primario: #2A2D34 (gris oscuro casi negro - seriedad)
- Secundario: #E8E9EB (gris claro casi blanco - limpieza)
- Acento: #C97D60 (terracota - calidez humana)
- Énfasis: #8CB369 (verde oliva - crecimiento, sostenibilidad)
- Fondo: #F5F1E6 (crema - organicidad)

**Tipografía:**
- Títulos: "Space Grotesk" (geométrica, tech, fuerte)
- Cuerpo: "Inter" (legibilidad excelente, neutral)
- Énfasis: "IBM Plex Mono" (código, algoritmo, precisión)

### 📱 WIREFRAME 1: LANDING PAGE (VERSIÓN DESKTOP)

```
┌─────────────────────────────────────────────────────────────┐
│ TREQE                                                        │
│ [Logo]                    [Iniciar Sesión] [Registrarse]    │
├─────────────────────────────────────────────────────────────┤
│                                                            │
│   INTERCAMBIA LO QUE TIENES                                │
│   POR LO QUE QUIERES                                       │
│   En 5 minutos. Sin regateos.                              │
│                                                            │
│   [¿Qué tienes?] [¿Qué quieres?] [BUSCAR TRUEQUES]        │
│                                                            │
│   "Ana intercambió su bicicleta por un sofá por un         │
│    ordenador en 72 horas" - Ver caso completo →            │
│                                                            │
├─────────────────────────────────────────────────────────────┤
│   CÓMO FUNCIONA                                            │
│   ┌─────┐ ┌─────┐ ┌─────┐ ┌─────┐                         │
│   │ 1.  │ │ 2.  │ │ 3.  │ │ 4.  │                         │
│   │Cuenta││Descubre││Vive tu││La magia│                    │
│   │tu    ││lo que  ││vida   ││ocurre  │                    │
│   │historia│te emociona│      │        │                    │
│   └─────┘ └─────┘ └─────┘ └─────┘                         │
│                                                            │
│   TRUEQUES IMPOSIBLES RESUELTOS                            │
│   ┌─────────────────────────────────────────────────────┐  │
│   │ 🚲 → 🛋️ → 💻                                     │  │
│   │ Bicicleta → Sofá → Ordenador                         │  │
│   │ 3 personas, 72 horas, 0€ intercambiados              │  │
│   └─────────────────────────────────────────────────────┘  │
│   ┌─────────────────────────────────────────────────────┐  │
│   │ 📚 → 🎮 → 📱                                      │  │
│   │ Libros raros → Consola → Smartphone                  │  │
│   │ 4 personas, 5 días, todos contentos                  │  │
│   └─────────────────────────────────────────────────────┘  │
│                                                            │
└─────────────────────────────────────────────────────────────┘
```

**Características clave landing:**
1. **Hero section minimalista** con CTA claro
2. **Proceso de 4 pasos** visual simple
3. **Casos reales** con iconografía y timeline
4. **Brutalista:** Bordes rectos, contraste alto, jerarquía clara
5. **Orgánico:** Colores tierra, espacio generoso, tipografía humana

### 📱 WIREFRAME 2: DASHBOARD DE USUARIO

```
┌─────────────────────────────────────────────────────────────┐
│ [Logo] Hola, Ana 👋                     [Notificaciones: 3] │
├─────────────────────────────────────────────────────────────┤
│                                                            │
│   MIS TRUEQUES ACTIVOS (2)                                │
│   ┌─────────────────────────────────────────────────────┐  │
│   │ 🔄 EN PROCESO                                      │  │
│   │ Bicicleta Orbea → Sofá IKEA → Ordenador MacBook    │  │
│   │ ✅ Ana → Beatriz (completado)                      │  │
│   │ 🔄 Carlos → Ana (en camino)                        │  │
│   │ ⏳ Beatriz → Carlos (pendiente)                    │  │
│   │ Tiempo estimado: 48 horas restantes                │  │
│   └─────────────────────────────────────────────────────┘  │
│   ┌─────────────────────────────────────────────────────┐  │
│   │ ⏳ PENDIENTE DE MATCH                               │  │
│   │ Guitarra española → ??                              │  │
│   │ "Buscando matches... 15 personas interesadas"       │  │
│   │ [Cancelar búsqueda] [Aumentar tiempo: +24h]         │  │
│   └─────────────────────────────────────────────────────┘  │
│                                                            │
│   ¿QUÉ QUIERES INTERCAMBIAR HOY?                          │
│   ┌─────┬─────┬─────┬─────┐                               │
│   │ 📱 │ 💻 │ 🚲 │ 🎮 │                               │
│   │Tech │Hogar│Deporte│Hobbies│                            │
│   └─────┴─────┴─────┴─────┘                               │
│   [Añadir nuevo artículo...]                              │
│                                                            │
│   TU REPUTACIÓN: ⭐⭐⭐⭐⭐ (4.8/5)                        │
│   "Miembro Confiable - 12 trueques completados"           │
│                                                            │
│   TRUEQUES SUGERIDOS PARA TI                              │
│   ┌─────────────────────────────────────────────────────┐  │
│   │ 🎧 → 📚                                            │  │
│   │ Auriculares Sony → Colección Stephen King           │  │
│   │ Match probability: 92%                              │  │
│   │ [Ver detalles] [Iniciar trueque]                    │  │
│   └─────────────────────────────────────────────────────┘  │
│                                                            │
└─────────────────────────────────────────────────────────────┘
```

**Características dashboard:**
1. **Estado visual claro** de trueques (✅🔄⏳)
2. **Progreso circular** mostrado como timeline
3. **Categorías con iconos** para rápido acceso
4. **Reputación prominente** como motivador
5. **Sugerencias inteligentes** basadas en historial

### 📱 WIREFRAME 3: PROCESO DE MATCHING (ALGORITMO VISUALIZADO)

```
┌─────────────────────────────────────────────────────────────┐
│ ENCONTRANDO TU TRUEQUE PERFECTO...                         │
├─────────────────────────────────────────────────────────────┤
│                                                            │
│   BUSCANDO K=2 (MÁS SIMPLE)                               │
│   ┌───┐      ┌───┐                                        │
│   │ A │ ───→ │ B │   ❌ No match directo                 │
│   └───┘      └───┘                                        │
│                                                            │
│   BUSCANDO K=3 (CÍRCULO PEQUEÑO)                          │
│   ┌───┐      ┌───┐      ┌───┐                            │
│   │ A │ ───→ │ B │ ───→ │ C │                            │
│   └───┘      └───┘      └───┘                            │
│      ↑                    ↓                                │
│      └────────────────────┘   ✅ ¡MATCH ENCONTRADO!       │
│                                                            │
│   TRUEQUE CIRCULAR ENCONTRADO:                            │
│   Ana (bicicleta) → Beatriz (sofá) → Carlos (ordenador)   │
│                                                            │
│   DETALLES DEL TRUEQUE:                                   │
│   ┌─────────────────────────────────────────────────────┐  │
│   │ Participantes: 3                                    │  │
│   │ Tiempo estimado: 72 horas                           │  │
│   │ Valor total: 1.850€                                 │  │
│   │ Comisión Treqe: 55.50€ (3%)                         │  │
│   │ Garantía: Activada (hasta 1.000€)                   │  │
│   │ [Aceptar trueque] [Rechazar] [Negociar]             │  │
│   └─────────────────────────────────────────────────────┘  │
│                                                            │
│   ¿POR QUÉ ESTE TRUEQUE FUNCIONA?                         │
│   • Densidad de mercado: 5.2% (óptimo)                    │
│   • Reputación media: 4.6/5 (alta confianza)              │
│   • Distancia máxima: 42km (logística viable)             │
│   • Tiempo algoritmo: 2.3 segundos                        │  │
│                                                            │
└─────────────────────────────────────────────────────────────┘
```

**Características matching:**
1. **Visualización algoritmo** en tiempo real
2. **Explicación progresiva** de k=2 → k=3 → match
3. **Transparencia total** en cálculos y razones
4. **Datos técnicos** presentados de forma comprensible
5. **Opciones claras** con consecuencias visibles

### 📱 WIREFRAME 4: MÓVIL (VERSIÓN ADAPTADA)

```
┌─────────────────────────────┐
│ TREQE                      ≡│
├─────────────────────────────┤
│                             │
│   Hola, Ana 👋             │
│   Trueques activos: 2      │
│                             │
│   ┌─────────────────────┐  │
│   │ 🔄 Bicicleta → Sofá │  │
│   │    → Ordenador      │  │
│   │ 48h restantes       │  │
│   └─────────────────────┘  │
│                             │
│   ＋ Añadir artículo        │
│                             │
│   🔍 Buscar trueques        │
│                             │
│   ⭐ Mi reputación: 4.8/5  │
│                             │
│   💬 Comunidad             │
│                             │
│   ⚙️ Ajustes               │
│                             │
└─────────────────────────────┘
```

**Características móvil:**
1. **Navegación hamburguesa** para espacio
2. **Información esencial** arriba
3. **Acciones principales** con iconos grandes
4. **Diseño táctil** con áreas de toque generosas
5. **Mismo lenguaje visual** que desktop

### 🎨 SISTEMA DE DISEÑO (DESIGN TOKENS)

**Espaciado (8px grid):**
- xs: 4px (elementos muy juntos)
- sm: 8px (espaciado interno)
- md: 16px (espaciado estándar)
- lg: 24px (separación secciones)
- xl: 32px (márgenes grandes)

**Sombras:**
- sm: 0 1px 3px rgba(0,0,0,0.12)
- md: 0 4px 6px rgba(0,0,0,0.1)
- lg: 0 10px 15px rgba(0,0,0,0.1)

**Bordes:**
- sm: 2px (botones, inputs)
- md: 4px (tarjetas, contenedores)
- lg: 8px (elementos destacados)

**Animaciones:**
- Fast: 150ms (hover, focus)
- Medium: 300ms (transiciones UI)
- Slow: 500ms (cambios de página)

### 🚀 PROTOTIPO FUNCIONAL (HTML/CSS/JS)

**Tecnología:** Vanilla HTML/CSS/JS + minimal framework
**Razón:** Rapidez, control total, sin dependencias
**Características:**
- Progressive enhancement (funciona sin JS)
- Responsive por defecto (mobile-first)
- Accessibility (ARIA labels, keyboard nav)
- Performance (lazy loading, optimized assets)

**Estructura código:**
```
treqe-prototype/
├── index.html          # Landing page
├── dashboard.html      # Dashboard usuario
├── matching.html       # Proceso matching
├── styles/
│   ├── base.css       # Reset, variables
│   ├── components.css # Botones, cards, forms
│   └── layout.css     # Grid, flexbox
├── scripts/
│   ├── main.js        # Lógica principal
│   └── matching.js    # Simulación algoritmo
└── assets/
    ├── icons/         # SVG icons
    └── images/        # Optimized images
```

### 💡 PRINCIPIOS DE UX APLICADOS:

1. **Clarity over cleverness:** Interfaz obvia, no adivinar
2. **Progressive disclosure:** Información cuando se necesita
3. **Forgiving design:** Errores fáciles de corregir
4. **Consistency:** Mismos patrones en toda la plataforma
5. **Feedback inmediato:** Acciones con respuesta visual clara
6. **Human-centered:** Lenguaje natural, no técnico

### 🎯 CONCLUSIÓN DISEÑO:

**"Treqe necesita una interfaz que transmita confianza (brutalismo) pero también calidez humana (orgánico). El diseño debe hacer visible lo invisible: el algoritmo trabajando, la reputación creciendo, la comunidad activa."**

**Este sistema de diseño:**
1. **Escala** de 1 usuario a 1 millón
2. **Mantiene coherencia** entre web y móvil
3. **Educa al usuario** sobre el proceso
4. **Construye confianza** mediante transparencia
5. **Diferencia Treqe** de competidores genéricos

**Los wireframes están listos para implementación o presentación a inversores como parte del anexo visual del plan de negocio.**
"""
    
    return wireframes

def insertar_wireframes_en_documento():
    """Insertar wireframes en el documento del plan"""
    print("Insertando wireframes en documento...")
    
    # Cargar documento actual
    try:
        doc = Document('Plan_Negocio_Treqe_BMC_MEJORADO.docx')
        print(f"Documento cargado: {len(doc.paragraphs)} parrafos")
    except:
        print("Error cargando documento")
        return
    
    # Crear wireframes
    wireframes_content = crear_wireframes_treqe()
    
    # Buscar donde insertar (al final, como anexo)
    print("Insertando como anexo al final del documento...")
    
    # Añadir título de anexo
    p = doc.add_paragraph()
    p.add_run("ANEXO A: WIREFRAMES Y DISEÑO UI/UX - PLATAFORMA TREQE").bold = True
    p.style = 'Heading 1'
    
    # Añadir descripción
    p = doc.add_paragraph()
    p.add_run("Este anexo contiene los wireframes y sistema de diseño creados con la skill 'frontend-design'. Muestra la dirección estética, componentes UI, y experiencia de usuario planeada para la plataforma Treqe.")
    
    # Insertar contenido
    lineas = wireframes_content.strip().split('\n')
    for linea in lineas:
        if linea.strip():
            p = doc.add_paragraph()
            p.text = linea.strip()
    
    # Guardar documento
    output_path = 'Plan_Negocio_Treqe_COMPLETO_FINAL.docx'
    doc.save(output_path)
    
    print(f"\n✅ DOCUMENTO COMPLETO CREADO: {output_path}")
