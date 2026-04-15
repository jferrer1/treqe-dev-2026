#!/usr/bin/env python3
"""
Analizar estructura completa del documento para identificar secciones específicas
"""

import re
from docx import Document

def analizar_estructura():
    print("ANALIZANDO ESTRUCTURA DEL DOCUMENTO TREQE...")
    
    try:
        # Cargar documento
        doc = Document('Plan_Negocio_Treqe_CON_SOLUCIONES_2026.docx')
        
        print(f"Documento cargado: {len(doc.paragraphs)} párrafos")
        
        # Buscar secciones específicas
        secciones_buscar = [
            'algoritmo', 'matching', 'intercambio', 'ciclo',
            'envío', 'envios', 'logística', 'recepcion',
            'garantía', 'garantias', 'seguro', 'seguridad',
            'scoring', 'reputación', 'puntuación',
            'proceso', 'flujo', 'usuario', 'experiencia',
            'tecnología', 'tecnico', 'sistema'
        ]
        
        print("\n=== SECCIONES IDENTIFICADAS ===")
        
        for i, para in enumerate(doc.paragraphs):
            texto = para.text.strip()
            if texto and len(texto) > 10:
                texto_lower = texto.lower()
                
                # Buscar títulos de sección (texto en mayúsculas o con números)
                es_titulo = (
                    len(texto) < 100 and 
                    (texto.isupper() or re.search(r'^\d+\.', texto) or 
                     any(palabra in texto_lower for palabra in secciones_buscar))
                )
                
                if es_titulo:
                    print(f"\n[{i:4d}] TÍTULO: {texto}")
                    
                    # Mostrar las primeras líneas del contenido
                    contenido_lines = []
                    for j in range(1, 6):
                        if i + j < len(doc.paragraphs):
                            contenido = doc.paragraphs[i + j].text.strip()
                            if contenido:
                                contenido_lines.append(contenido[:100])
                    
                    if contenido_lines:
                        print(f"     Contenido inicial: {' | '.join(contenido_lines[:2])}")
        
        # Buscar específicamente secciones sobre algoritmo
        print("\n=== BUSCANDO SECCIONES SOBRE ALGORITMO ===")
        for i, para in enumerate(doc.paragraphs):
            texto = para.text.strip().lower()
            if any(palabra in texto for palabra in ['algoritmo', 'matching', 'ciclo', 'k=', 'rueda']):
                print(f"\n[{i:4d}] Posible sección algoritmo: {para.text[:80]}...")
                
                # Mostrar contexto
                for j in range(-2, 3):
                    idx = i + j
                    if 0 <= idx < len(doc.paragraphs):
                        contexto = doc.paragraphs[idx].text.strip()
                        if contexto:
                            print(f"     [{idx:4d}] {contexto[:80]}{'...' if len(contexto) > 80 else ''}")
        
        # Buscar sección de conclusiones o anexos
        print("\n=== BUSCANDO SECCIÓN DE CONCLUSIONES ===")
        for i, para in enumerate(doc.paragraphs):
            texto = para.text.strip().lower()
            if any(palabra in texto for palabra in ['conclusión', 'conclusiones', 'anexo', 'apéndice']):
                print(f"\n[{i:4d}] {para.text}")
                
        return True
        
    except Exception as e:
        print(f"ERROR: {e}")
        import traceback
        traceback.print_exc()
        return False

if __name__ == '__main__':
    analizar_estructura()