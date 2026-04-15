#!/usr/bin/env python3
"""
Script simple para crear documento Word profesional para Treqe
"""

import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Crear documento
doc = Document()

# Estilo normal
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# ===== PORTADA =====
title = doc.add_heading('PLAN DE NEGOCIO PROFESIONAL', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.size = Pt(36)
title.runs[0].bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.add_run('TREQE\n').font.size = Pt(28)
subtitle.add_run('Plataforma de Trueque Inteligente\n').font.size = Pt(18)

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run(f'Fecha: {datetime.now().strftime("%d de %B de %Y")}\n').font.size = Pt(12)
info.add_run('Versión: 3.0 - Documento Profesional Mejorado\n').font.size = Pt(12)
info.add_run('Estado: CONFIDENCIAL\n').font.size = Pt(12)
info.add_run('Páginas: 30-35\n').font.size = Pt(12)

doc.add_page_break()

# ===== ÍNDICE =====
doc.add_heading('ÍNDICE', 1)

# Tabla índice
table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'SECCIÓN'
hdr[1].text = 'DESCRIPCIÓN'
hdr[2].text = 'PÁGINA'

secciones = [
    ('1', 'INTRODUCCIÓN', '3'),
    ('1.1', 'Transformación Sector', '3'),
    ('1.2', 'Datos 2025-2026', '4'),
    ('1.3', 'Competencia', '6'),
    ('1.4', 'Tendencias', '8'),
    ('2', 'PROBLEMA', '10'),
    ('2.1', 'Paradoja Liquidez', '10'),
    ('2.2', 'Opciones No Óptimas', '11'),
    ('2.3', 'Limitación Matemática', '12'),
    ('2.4', 'Oportunidad', '13'),
    ('3', 'SOLUCIÓN TREQE', '15'),
    ('3.1', 'Concepto Revolucionario', '15'),
    ('3.2', 'Mecanismo Operativo', '16'),
    ('3.3', 'Ejemplo Práctico', '19'),
    ('3.4', 'Innovaciones', '22'),
    ('4', 'VENTAJA COMPETITIVA', '24'),
    ('4.1', 'Posicionamiento', '24'),
    ('4.2', 'Ventajas Tecnológicas', '26'),
    ('4.3', 'Ventajas Económicas', '28'),
    ('4.4', 'Barreras Entrada', '30'),
    ('5', 'MODELO NEGOCIO', '32'),
    ('5.1', 'Filosofía', '32'),
    ('5.2', 'Flujos Ingresos', '33'),
    ('5.3', 'Inversión', '35'),
    ('5.4', 'Financiación', '37'),
    ('6', 'FINANZAS 2026-2029', '39'),
    ('6.1', 'Supuestos', '39'),
    ('6.2', 'Proyecciones', '41'),
    ('6.3', 'Pérdidas/Ganancias', '43'),
    ('6.4', 'Cash Flow', '45'),
    ('6.5', 'Ratios', '47'),
    ('7', 'EQUIPO Y EJECUCIÓN', '49'),
    ('7.1', 'Equipo Fundador', '49'),
    ('7.2', 'Plan Fases', '51'),
    ('7.3', 'Próximos Pasos', '53'),
    ('8', 'RIESGOS', '55'),
    ('8.1', 'Matriz Riesgos', '55'),
    ('8.2', 'Mitigación', '57'),
    ('9', 'CONCLUSIONES', '59'),
    ('APÉNDICES', '', '61')
]

for s in secciones:
    row = table.add_row().cells
    row[0].text = s[0]
    row[1].text = s[1]
    row[2].text = s[2]

doc.add_page_break()

# ===== SECCIÓN 1 =====
doc.add_heading('1. INTRODUCCIÓN: CONTEXTO MERCADO', 1)

doc.add_heading('1.1 Transformación Sector', 2)
doc.add_paragraph('El mercado de segunda mano en España ha evolucionado de respuesta a crisis a movimiento cultural. Valores emergentes: sostenibilidad, consumo consciente, relación inteligente con objetos.')

doc.add_heading('1.2 Datos Cuantitativos 2025-2026', 2)
doc.add_paragraph('• Volumen 2026: 8.200M€ (+42% desde 2020)', style='List Bullet')
doc.add_paragraph('• Usuarios: 28M españoles (47% población)', style='List Bullet')
doc.add_paragraph('• Gasto medio: 1.850€/usuario/año', style='List Bullet')
doc.add_paragraph('• Mobile-first: 94% transacciones desde móvil', style='List Bullet')

doc.add_page_break()

# ===== SECCIÓN 2 =====
doc.add_heading('2. PROBLEMA: PARADOJA DE LA LIQUIDEZ', 1)

doc.add_heading('2.1 Situación Usuario', 2)
doc.add_paragraph('Millones tienen valor atrapado en posesiones no deseadas, pero carecen de liquidez para lo que necesitan.')

doc.add_heading('Ejemplo: Ana, arquitecta 32 años', 3)
doc.add_paragraph('Tiene: Bicicleta (450€), sofá (600€), libros (450€) = 1.500€', style='List Bullet')
doc.add_paragraph('Necesita: Escritorio, estanterías, sofá moderno = 2.000€', style='List Bullet')
doc.add_paragraph('Problema: Valor existe, necesidad existe, falta mecanismo', style='List Bullet')

doc.add_heading('2.2 Opciones No Óptimas', 2)
doc.add_paragraph('1. Mantener objetos: Ocupa espacio, depreciación, coste psicológico', style='List Number')
doc.add_paragraph('2. Vender por menos: Pérdida 30-50% valor real', style='List Number')
doc.add_paragraph('3. Trueque directo: Coincidencia perfecta casi imposible (5% probabilidad)', style='List Number')

doc.add_page_break()

# ===== SECCIÓN 3 =====
doc.add_heading('3. SOLUCIÓN TREQE: RUEDAS INTERCAMBIO', 1)

doc.add_heading('3.1 Concepto Revolucionario', 2)
doc.add_paragraph('Treqe resuelve paradoja con ruedas de intercambio inteligente. Algoritmo identifica ciclos entre 3+ usuarios.')

doc.add_heading('3.2 Mecanismo 4 Pasos', 2)
doc.add_paragraph('1. Registro: Usuario cataloga artículos, indica preferencias', style='List Number')
doc.add_paragraph('2. Algoritmo: Analiza preferencias cruzadas, encuentra ciclos', style='List Number')
doc.add_paragraph('3. Validación: Propuesta a participantes, cada uno acepta', style='List Number')
doc.add_paragraph('4. Ejecución: Intercambios coordinados, garantía activa', style='List Number')

doc.add_heading('3.3 Ejemplo Práctico', 2)
doc.add_paragraph('Carlos: Bicicleta → quiere Consola', style='List Bullet')
doc.add_paragraph('Beatriz: Consola → quiere Sofá', style='List Bullet')
doc.add_paragraph('David: Sofá → quiere Ordenador', style='List Bullet')
doc.add_paragraph('Elena: Ordenador → quiere Bicicleta', style='List Bullet')
doc.add_paragraph('Solución: Carlos→Beatriz→David→Elena→Carlos', style='List Bullet')
doc.add_paragraph('Resultado: Todos felices. Valor: 2.000€. Comisión: 60€ (3%).', style='List Bullet')

doc.add_page_break()

# ===== SECCIÓN 4 =====
doc.add_heading('4. VENTAJA COMPETITIVA', 1)

doc.add_heading('4.1 Posicionamiento Único', 2)
doc.add_paragraph('Espacio vacante: trueque estructurado y escalable. Competidores: compraventa monetaria.')

doc.add_heading('Wallapop/Vinted vs Treqe', 3)
doc.add_paragraph('Ellos: "Vende por menos, compra por más" - Usuario pierde valor', style='List Bullet')
doc.add_paragraph('Nosotros: "Mantén valor, obtén lo que quieres" - Usuario preserva valor', style='List Bullet')

doc.add_heading('4.2 Ventajas Tecnológicas', 2)
doc.add_paragraph('• Algoritmo patentable: Resuelve NP-Completo matching circular', style='List Bullet')
doc.add_paragraph('• Eficiente: <5 minutos para 1.000 usuarios', style='List Bullet')
doc.add_paragraph('• Sistema reputación: Blockchain, transparencia, reduce fraude 90%', style='List Bullet')

doc.add_page_break()

# ===== SECCIÓN 5 =====
doc.add_heading('5. MODELO DE NEGOCIO', 1)

doc.add_heading('5.1 Filosofía', 2)
doc.add_paragraph('Solo ganamos cuando usuarios ganan. Alineación perfecta incentivos.')

doc.add_heading('5.2 Flujos Ingresos Multicapa', 2)
doc.add_paragraph('1. Comisión 3% por éxito (solo si intercambio funciona)', style='List Number')
doc.add_paragraph('2. Suscripción Premium 9,99€/mes (prioridad, valoraciones, seguro)', style='List Number')
doc.add_paragraph('3. Servicios empresas (white-label, análisis datos, consultoría)', style='List Number')

doc.add_heading('5.3 Inversión Inicial: 58.000€', 2)
doc.add_paragraph('• Desarrollo: 23.200€ (40%)', style='List Bullet')
doc.add_paragraph('• Marketing: 20.300€ (35%)', style='List Bullet')
doc.add_paragraph('• Operaciones: 14.500€ (25%)', style='List Bullet')

doc.add_page_break()

# ===== SECCIÓN 6 =====
doc.add_heading('6. PROYECCIONES FINANCIERAS 2026-2029', 1)

# Tabla proyecciones
table = doc.add_table(rows=5, cols=5)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'AÑO'
hdr[1].text = 'USUARIOS'
hdr[2].text = 'ACTIVOS'
hdr[3].text = 'INTERCAMBIOS/MES'
hdr[4].text = 'INGRESOS'

data = [
    ('2026', '50.000', '8.000', '10.000', '360.000€'),
    ('2027', '150.000', '25.000', '30.000', '1.728.000€'),
    ('2028', '350.000', '60.000', '70.000', '5.040.000€'),
    ('2029', '750.000', '120.000', '150.000', '12.960.000€')
]

for i, d in enumerate(data, 1):
    row = table.rows[i].cells
    row[0].text = d[0]
    row[1].text = d[1]
    row[2].text = d[2]
    row[3].text = d[3]
    row[4].text = d[4]

doc.add_paragraph()

doc.add_heading('6.3 Estado Pérdidas/Ganancias Año 1', 2)
doc.add_paragraph('Ingresos: 360.000€', style='List Bullet')
doc.add_paragraph('Costes: 276.000€', style='List Bullet')
doc.add_paragraph('Beneficio: 84.000€ (23% margen)', style='List Bullet')

doc.add_page_break()

# ===== SECCIÓN 7 =====
doc.add_heading('7. EQUIPO Y EJECUCIÓN', 1)

doc.add_heading('7.1 Equipo Fundador', 2)
doc.add_paragraph('• CEO: 10+ años scale-ups, ex-director producto', style='List Bullet')
doc.add_paragraph('• CTO: PhD Ciencias Computación, especialista algoritmos', style='List Bullet')
doc.add_paragraph('• CMO: Experto crecimiento orgánico, métricas-driven', style='List Bullet')

doc.add_heading('7.2 Plan 3 Fases', 2)
doc.add_paragraph('1. MVP y Validación (Meses 1-3): Algoritmo básico, 100 usuarios beta', style='List Number')
doc.add_paragraph('2. Escala Madrid (Meses 4-9): Lanzamiento público, optimización', style='List Number')
doc.add_paragraph('3. Expansión Nacional (Meses 10-18): 5 ciudades, equipo comercial', style='List Number')

doc.add_page_break()

# ===== SECCIÓN 8 =====
doc.add_heading('8. RIESGOS Y MITIGACIÓN', 1)

# Tabla riesgos
table = doc.add_table(rows=6, cols=4)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'RIESGO'
hdr[1].text = 'PROB'
hdr[2].text = 'IMPACTO'
hdr[3].text = 'MITIGACIÓN'

riesgos = [
    ('Huevo-gallina', 'Alta', 'Alto', 'Comunidad cerrada inicial'),
    ('Fallo algoritmo', 'Media', 'Alto', 'Testing + backup manual'),
    ('Fraudes', 'Media', 'Alto', 'Reputación + garantía'),
    ('Competencia', 'Baja', 'Medio', 'Patentes + first-mover'),
    ('Legales', 'Baja', 'Alto', 'Asesoría proactiva'),
    ('Escalabilidad', 'Media', 'Medio', 'Arquitectura serverless')
]

for i, r in enumerate(riesgos, 1):
    row = table.rows[i].cells
    row[0].text = r[0]
    row[1].text = r[1]
    row[2].text = r[2]
    row[3].text = r[3]

doc.add_page_break()

# ===== SECCIÓN 9 =====
doc.add_heading('9. CONCLUSIONES', 1)

doc.add_paragraph('Treqe representa oportunidad única porque:', style='List Bullet')
doc.add_paragraph('1. Resuelve problema real no atendido', style='List Number')
doc.add_paragraph('2. Ocupa espacio mercado vacante', style='List Number')
doc.add_paragraph('3. Modelo viable desde día 1', style='List Number')
doc.add_paragraph('4. Ventajas competitivas sostenibles', style='List Number')
doc.add_paragraph('5. Equipo experimentado', style='List Number')
doc.add_paragraph('6. Plan ejecución realista', style='List Number')
doc.add_paragraph('7. Riesgos mitigados', style='List Number')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Recomendación: ').bold = True
p.add_run('Inversión 58.000€ para MVP. Retorno esperado: 5-7x en 3 años.')

p = doc.add_paragraph()
p.add_run('Próximos pasos: ').bold = True
p.add_run('Constituir SL, registrar marca, algoritmo v0.1, primeros 50 usuarios.')

# ===== FIN =====
doc.add_page_break()
final = doc.add_paragraph()
final.alignment = WD_ALIGN_PARAGRAPH.CENTER
final.add_run('--- FIN DEL DOCUMENTO ---\n').font.size = Pt(14)
final.add_run('Treqe SL - Plataforma Trueque Inteligente\n').font.size = Pt(12)
final.add_run(f'Generado: {datetime.now().strftime("%d/%m/%Y %H:%M")}\n').font.size = Pt(10)
final.add_run('Confidencial\n').font.size = Pt(10)

# Guardar
output_path = os.path.join(os.path.dirname(__file__), 'Plan_Negocio_Treqe_PROF