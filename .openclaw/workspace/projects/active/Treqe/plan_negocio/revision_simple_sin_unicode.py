#!/usr/bin/env python3
"""
Revisión simple sin Unicode
"""

from docx import Document
import re

def main():
    print("REVISION EXHAUSTIVA DEL DOCUMENTO FINAL")
    print("========================================")
    
    # Cargar documentos
    try:
        doc_final = Document('Plan_Negocio_Treqe_FINAL_CORREGIDO.docx')
        doc_original = Document('Plan_Negocio_Treqe_MARKETING_MEJORADO.docx')
        print(f"Documento final cargado: {len(doc_final.paragraphs)} parrafos")
        print(f"Documento original cargado: {len(doc_original.paragraphs)} parrafos")
    except Exception as e:
        print(f"Error: {e}")
        return
    
    print("\n1. VERIFICACION HUMANIZER:")
    
    # Buscar lenguaje corporativo que deberia estar humanizado
    lenguaje_corporativo = ['sinergia', 'paradigma', 'ecosistema', 'optimizacion']
    problemas = []
    
    for i, para in enumerate(doc_final.paragraphs):
        texto = para.text.lower()
        for palabra in lenguaje_corporativo:
            if palabra in texto:
                problemas.append((i, palabra, para.text[:80]))
                break
    
    if problemas:
        print(f"  PROBLEMAS: {len(problemas)} palabras corporativas encontradas")
        for j, (num, palabra, texto) in enumerate(problemas[:3]):
            print(f"    Parrafo {num}: '{palabra}' en: {texto}...")
    else:
        print("  OK: No se encontro lenguaje corporativo (humanizer aplicado)")
    
    print("\n2. VERIFICACION CONTENIDO COMPLETO:")
    
    # Contenido clave que debe estar
    contenido_clave = [
        "SEO Programatico",
        "5-Phase",
        "Psychology",
        "Conversion",
        "Retention",
        "CAC",
        "LTV:CAC",
        "presupuesto trimestral"
    ]
    
    faltante = []
    for clave in contenido_clave:
        encontrado = any(clave.lower() in p.text.lower() for p in doc_final.paragraphs)
        if not encontrado:
            faltante.append(clave)
    
    if faltante:
        print(f"  FALTANTE: {len(faltante)} elementos")
        for clave in faltante:
            print(f"    - {clave}")
    else:
        print("  OK: Todo el contenido clave esta presente")
    
    print("\n3. COMPARACION CON OTROS DOCUMENTOS:")
    
    # Lista de documentos para comparar
    docs_comparar = [
        ('Plan_Negocio_Treqe_HUMANIZADO_ACTUALIZADO.docx', 'Humanizado Actualizado'),
        ('Plan_Negocio_Treqe_MARKETING_MEJORADO.docx', 'Marketing Mejorado'),
        ('Plan_Negocio_Treqe_MARKETING_MEJORADO_HUMANIZADO.docx', 'Marketing Humanizado'),
        ('Plan_Negocio_Treqe_10h28_HUMANIZADO_SELECTIVO.docx', '10:28 Selectivo'),
    ]
    
    resultados = []
    for archivo, nombre in docs_comparar:
        try:
            doc = Document(archivo)
            resultados.append((nombre, len(doc.paragraphs), archivo))
        except:
            pass
    
    # Ordenar
    resultados.sort(key=lambda x: x[1], reverse=True)
    
    print("  Ranking por parrafos:")
    for i, (nombre, parrafos, archivo) in enumerate(resultados):
        es_final = (archivo == 'Plan_Negocio_Treqe_FINAL_CORREGIDO.docx')
        marca = ">>>" if es_final else "   "
        print(f"    {marca} {i+1}. {nombre}: {parrafos} parrafos")
    
    # Verificar si es el mas completo
    mas_completo = resultados[0][2] == 'Plan_Negocio_Treqe_FINAL_CORREGIDO.docx'
    if mas_completo:
        print("  OK: Es el documento mas completo")
    else:
        print(f"  NO ES EL MAS COMPLETO. El mas completo es: {resultados[0][0]}")
    
    print("\n4. ESTRUCTURA PROFESIONAL:")
    
    secciones = [
        "INTRODUCCION",
        "PROBLEMA",
        "SOLUCION",
        "VENTAJA COMPETITIVA",
        "MODELO DE NEGOCIO",
        "PROYECCIONES FINANCIERAS",
        "EQUIPO",
        "RIESGOS",
        "CONCLUSIONES",
        "MARKETING",
        "LEGAL",
        "ANEXOS"
    ]
    
    encontradas = []
    for para in doc_final.paragraphs:
        texto = para.text.upper()
        for seccion in secciones:
            if seccion in texto and len(texto) < 100:
                if seccion not in encontradas:
                    encontradas.append(seccion)
    
    print(f"  Secciones encontradas: {len(encontradas)}/{len(secciones)}")
    
    faltantes_secciones = [s for s in secciones if s not in encontradas]
    if faltantes_secciones:
        print(f"  FALTAN: {len(faltantes_secciones)} secciones")
        for s in faltantes_secciones[:3]:
            print(f"    - {s}")
    else:
        print("  OK: Todas las secciones profesionales presentes")
    
    print("\n5. RESUMEN FINAL:")
    print("=" * 40)
    
    # Calcular puntuacion
    puntuacion = 0
    max_puntuacion = 4
    
    if len(problemas) == 0:
        print("✅ Humanizer: APLICADO CORRECTAMENTE")
        puntuacion += 1
    else:
        print(f"❌ Humanizer: {len(problemas)} problemas")
    
    if len(faltante) == 0:
        print("✅ Contenido: COMPLETO")
        puntuacion += 1
    else:
        print(f"❌ Contenido: {len(faltante)} elementos faltantes")
    
    if mas_completo:
        print("✅ Completitud: ES EL MAS COMPLETO")
        puntuacion += 1
    else:
        print("❌ Completitud: NO ES EL MAS COMPLETO")
    
    if len(faltantes_secciones) == 0:
        print("✅ Estructura: PROFESIONAL COMPLETA")
        puntuacion += 1
    else:
        print(f"❌ Estructura: {len(faltantes_secciones)} secciones faltantes")
    
    print(f"\nPUNTUACION: {puntuacion}/{max_puntuacion}")
    
    if puntuacion == max_puntuacion:
        print("\n🎉 ¡DOCUMENTO APROBADO! Listo para uso profesional.")
    else:
        print(f"\n⚠️ SE REQUIEREN AJUSTES ({max_puntuacion - puntuacion} areas por corregir)")
    
    # Crear informe simple
    with open('INFORME_REVISION_SIMPLE.md', 'w', encoding='utf-8') as f:
        f.write(f"""# INFORME DE REVISION
Documento: Plan_Negocio_Treqe_FINAL_CORREGIDO.docx

## RESULTADOS:

1. HUMANIZER: {"APLICADO" if len(problemas)==0 else f"{len(problemas)} PROBLEMAS"}
2. CONTENIDO: {"COMPLETO" if len(faltante)==0 else f"{len(faltante)} FALTANTES"}
3. COMPLETITUD: {"MAS COMPLETO" if mas_completo else "NO MAS COMPLETO"}
4. ESTRUCTURA: {"COMPLETA" if len(faltantes_secciones)==0 else f"{len(faltantes_secciones)} FALTANTES"}

## PUNTUACION: {puntuacion}/4

{"✅ DOCUMENTO LISTO PARA USO" if puntuacion==4 else "⚠️ REQUIERE AJUSTES"}

## RECOMENDACION:
{"Usar directamente para presentaciones e inversores." if puntuacion==4 else "Corregir los problemas identificados antes de usar."}
""")
    
    print("\nInforme guardado: INFORME_REVISION_SIMPLE.md")

if __name__ == '__main__':
    main()