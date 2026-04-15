import docx

def get_detailed_strikethrough_info():
    doc = docx.Document('Plan_Negocio_Treqe_CON_SOLUCIONES_2026.docx')
    
    print("DETALLES DE TEXTO TACHADO:\n")
    
    strikethrough_paragraphs = []
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        
        # Check for strikethrough
        has_strikethrough = False
        strikethrough_text = ""
        
        for run in para.runs:
            if run.font.strike:
                has_strikethrough = True
                strikethrough_text += run.text
        
        if has_strikethrough:
            # Find heading
            heading = "Sin titulo"
            for j in range(i-1, max(-1, i-5), -1):
                if doc.paragraphs[j].style.name.startswith('Heading'):
                    heading = doc.paragraphs[j].text.strip()
                    break
            
            # Get context (next 2 paragraphs)
            context_after = []
            for j in range(i+1, min(len(doc.paragraphs), i+3)):
                if doc.paragraphs[j].text.strip():
                    context_after.append(doc.paragraphs[j].text.strip())
            
            strikethrough_paragraphs.append({
                'index': i,
                'heading': heading,
                'text': text,
                'strikethrough_text': strikethrough_text,
                'context_after': context_after[:2]
            })
    
    # Print first 10 for review
    for i, para in enumerate(strikethrough_paragraphs[:15]):
        print(f"\n--- PARRAFO TACHADO {i+1} ---")
        print(f"Seccion: {para['heading']}")
        print(f"Texto completo: {para['text'][:200]}...")
        print(f"Texto tachado: {para['strikethrough_text'][:200]}...")
        if para['context_after']:
            print("Contexto siguiente:")
            for ctx in para['context_after']:
                print(f"  - {ctx[:100]}...")
        print("-" * 50)
    
    # Analyze by content
    print("\n\nANALISIS TEMATICO:")
    
    topics = {
        'chat_grupal': 0,
        'devoluciones': 0,
        'envios': 0,
        'garantias': 0,
        'algoritmo': 0,
        'reputacion': 0,
        'otros': 0
    }
    
    keyword_map = {
        'chat_grupal': ['chat', 'grupo', 'conversacion', 'negociacion', 'oficial', 'estructurada', 'propuesta'],
        'devoluciones': ['devolucion', 'desistimiento', '24 horas', 'garantia', 'fondo', 'reembolso'],
        'envios': ['envio', 'transporte', 'logistica', 'correos', 'paquete', 'entrega', 'shipping'],
        'garantias': ['garantia', 'seguro', 'escrow', 'deposito', 'compensacion', 'proteccion'],
        'algoritmo': ['algoritmo', 'matching', 'k=', 'ciclo', 'rueda', 'intercambio', 'optimizacion'],
        'reputacion': ['reputacion', 'scoring', 'puntuacion', 'confianza', 'historial', 'rating']
    }
    
    for para in strikethrough_paragraphs:
        text_lower = para['text'].lower()
        topic_found = 'otros'
        
        for topic, keywords in keyword_map.items():
            for keyword in keywords:
                if keyword in text_lower:
                    topic_found = topic
                    break
            if topic_found != 'otros':
                break
        
        topics[topic_found] += 1
    
    print("\nDistribucion de temas:")
    for topic, count in topics.items():
        if count > 0:
            print(f"  {topic}: {count} parrafos")
    
    return strikethrough_paragraphs

if __name__ == '__main__':
    get_detailed_strikethrough_info()