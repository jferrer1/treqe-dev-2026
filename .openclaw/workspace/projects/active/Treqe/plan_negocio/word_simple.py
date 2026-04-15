#!/usr/bin/env python3
"""
Script simple para crear documento Word de Treqe
"""

import os
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Crear documento
doc = Document()

# Estilo normal
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# Título
title = doc.add_heading('PLAN DE NEGOCIO TREQE', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.size = Pt(28)
title.runs[0].bold = True

# Subtítulo
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.add_run('Documento único definitivo\n').font.size = Pt(16)
sub.add_run('Todo en un solo documento').font.size = Pt(14)

# Fecha
fecha = doc.add_paragraph()
fecha.alignment = WD_ALIGN_PARAGRAPH.CENTER
fecha.add_run(f'Creado el {datetime.now().strftime("%d/%m/%Y")}')

doc.add_page_break()

# ===== RESUMEN EJECUTIVO =====
doc.add_heading('RESUMEN EJECUTIVO', 1)

p = doc.add_paragraph()
p.add_run('Problema: ').bold = True
p.add_run('72% de españoles tiene cosas guardadas que no usa. Intercambiar entre 2 personas tiene solo 5% de probabilidad.')

p = doc.add_paragraph()
p.add_run('Solución: ').bold = True
p.add_run('Treqe conecta a 3+ personas para intercambios circulares. Probabilidad sube al 20-35%.')

p = doc.add_paragraph()
p.add_run('Modelo: ').bold = True
p.add_run('3% de comisión cuando el intercambio funciona. Solo cobramos cuando tú ganas.')

doc.add_heading('Números clave', 2)
doc.add_paragraph('• Inversión: 58.000 EUR', style='List Bullet')
doc.add_paragraph('• Año 1: 8.000 usuarios, 360.000 EUR ingresos', style='List Bullet')
doc.add_paragraph('• Año 2: 25.000 usuarios, 1.728.000 EUR ingresos', style='List Bullet')
doc.add_paragraph('• Año 3: 60.000 usuarios, 5.040.000 EUR ingresos', style='List Bullet')

doc.add_heading('Próximos pasos', 2)
doc.add_paragraph('1. Constituir SL (3.000 EUR)', style='List Number')
doc.add_paragraph('2. Registrar marca "Treqe" (850 EUR)', style='List Number')
doc.add_paragraph('3. Desarrollar algoritmo v0.1', style='List Number')
doc.add_paragraph('4. Primeros 50 usuarios', style='List Number')

doc.add_page_break()

# ===== PROBLEMA =====
doc.add_heading('PARTE 1: EL PROBLEMA', 1)

doc.add_heading('Caso real: Ana, Carlos, Beatriz', 2)
doc.add_paragraph('• Ana: Bicicleta (300 EUR) → quiere sofá (300 EUR)', style='List Bullet')
doc.add_paragraph('• Carlos: Sofá (300 EUR) → quiere ordenador (300 EUR)', style='List Bullet')
doc.add_paragraph('• Beatriz: Ordenador (300 EUR) → quiere bicicleta (300 EUR)', style='List Bullet')

p = doc.add_paragraph()
p.add_run('Con Treqe (72 horas):').bold = True
p.add_run(' Todos felices. Valor intercambiado: 900 EUR. Comisión: 27 EUR (3%). Tiempo: 15 minutos cada uno.')

doc.add_page_break()

# ===== SOLUCIÓN =====
doc.add_heading('PARTE 2: LA SOLUCIÓN', 1)

doc.add_heading('Cómo funciona', 2)
doc.add_paragraph('1. Cuentas tu historia', style='List Number')
doc.add_paragraph('2. Descubres posibilidades', style='List Number')
doc.add_paragraph('3. Vives tu vida', style='List Number')
doc.add_paragraph('4. La magia ocurre', style='List Number')

doc.add_page_break()

# ===== MODELO NEGOCIO =====
doc.add_heading('PARTE 3: MODELO DE NEGOCIO', 1)

doc.add_heading('Segmentos clientes', 2)
doc.add_paragraph('1. Millennials urbanos', style='List Number')
doc.add_paragraph('2. Familias con hijos', style='List Number')
doc.add_paragraph('3. Estudiantes universitarios', style='List Number')

doc.add_heading('Flujos ingresos', 2)
doc.add_paragraph('• Comisión 3% sobre intercambios', style='List Bullet')
doc.add_paragraph('• Suscripción Premium 9.99 EUR/mes', style='List Bullet')
doc.add_paragraph('• Publicidad contextual', style='List Bullet')

doc.add_page_break()

# ===== FINANZAS =====
doc.add_heading('PARTE 4: PROYECCIONES FINANCIERAS', 1)

doc.add_heading('Inversión inicial: 58.000 EUR', 2)
doc.add_paragraph('• Tecnología: 23.200 EUR', style='List Bullet')
doc.add_paragraph('• Marketing: 20.300 EUR', style='List Bullet')
doc.add_paragraph('• Operaciones: 14.500 EUR', style='List Bullet')

doc.add_heading('Año 1 (2026)', 2)
doc.add_paragraph('• Usuarios: 50.000 total, 8.000 activos', style='List Bullet')
doc.add_paragraph('• Ingresos: 360.000 EUR', style='List Bullet')
doc.add_paragraph('• Beneficio: 84.000 EUR', style='List Bullet')

doc.add_page_break()

# ===== CRONOGRAMA =====
doc.add_heading('PARTE 5: CRONOGRAMA', 1)

doc.add_heading('Mes 1-3: Cimientos', 2)
doc.add_paragraph('• Fundación legal', style='List Bullet')
doc.add_paragraph('• Tecnología básica', style='List Bullet')
doc.add_paragraph('• Primeros usuarios', style='List Bullet')

doc.add_heading('Mes 4-6: Madrid', 2)
doc.add_paragraph('• Escala Madrid', style='List Bullet')
doc.add_paragraph('• Optimización', style='List Bullet')
doc.add_paragraph('• Consolidación', style='List Bullet')

doc.add_page_break()

# ===== LEGAL =====
doc.add_heading('PARTE 6: ASPECTOS LEGALES', 1)

doc.add_heading('Estructura: Sociedad Limitada', 2)
doc.add_paragraph('• Capital: 3.000 EUR', style='List Bullet')
doc.add_paragraph('• Responsabilidad limitada', style='List Bullet')

doc.add_heading('Protección', 2)
doc.add_paragraph('• Algoritmo: Secreto comercial', style='List Bullet')
doc.add_paragraph('• Marca: Registro España + UE', style='List Bullet')

doc.add_page_break()

# ===== DISEÑO =====
doc.add_heading('PARTE 7: DISEÑO', 1)

doc.add_heading('Estética', 2)
doc.add_paragraph('• "Brutalista digital con toques orgánicos"', style='List Bullet')
doc.add_paragraph('• Colores tierra', style='List Bullet')

doc.add_heading('Experiencia', 2)
doc.add_paragraph('• Registro: 30 segundos', style='List Bullet')
doc.add_paragraph('• Soporte: respuesta 2 horas', style='List Bullet')

doc.add_page_break()

# ===== RIESGOS =====
doc.add_heading('PARTE 8: RIESGOS', 1)

doc.add_heading('8 riesgos principales', 2)
doc.add_paragraph('1. Problema huevo-gallina', style='List Number')
doc.add_paragraph('2. Algoritmo falla', style='List Number')
doc.add_paragraph('3. Gastar dinero antes de tiempo', style='List Number')
doc.add_paragraph('4. Competencia copia', style='List Number')
doc.add_paragraph('5. Problemas legales', style='List Number')
doc.add_paragraph('6. Fraudes', style='List Number')
doc.add_paragraph('7. Logística', style='List Number')
doc.add_paragraph('8. Escalabilidad', style='List Number')

doc.add_page_break()

# ===== MARKETING =====
doc.add_heading('PARTE 9: MARKETING', 1)

doc.add_heading('Estrategia 5 fases', 2)
doc.add_paragraph('1. Pre-lanzamiento', style='List Number')
doc.add_paragraph('2. Beta privada', style='List Number')
doc.add_paragraph('3. Vista previa', style='List Number')
doc.add_paragraph('4. Early access', style='List Number')
doc.add_paragraph('5. Lanzamiento completo', style='List Number')

doc.add_page_break()

# ===== CHECKLISTS =====
doc.add_heading('PARTE 10: CHECKLISTS', 1)

doc.add_heading('Semana 1', 2)
doc.add_paragraph('[ ] Consultar abogados', style='List Bullet')
doc.add_paragraph('[ ] Registrar dominio', style='List Bullet')
doc.add_paragraph('[ ] Crear lista personas', style='List Bullet')

doc.add_heading('Mes 1', 2)
doc.add_paragraph('[ ] SL constituida', style='List Bullet')
doc.add_paragraph('[ ] Marca registrada', style='List Bullet')
doc.add_paragraph('[ ] Algoritmo v0.1', style='List Bullet')

doc.add_page_break()

# ===== CONCLUSIÓN =====
doc.add_heading('CONCLUSIÓN', 1)

p = doc.add_paragraph()
p.add_run('Por qué funciona:').bold = True

doc.add_paragraph('1. Resuelve problema real', style='List Number')
doc.add_paragraph('2. Matemáticamente sólido', style='List Number')
doc.add_paragraph('3. Modelo económico viable', style='List Number')
doc.add_paragraph('4. Escalable', style='List Number')
doc.add_paragraph('5. Protegido', style='List Number')

p = doc.add_paragraph()
p.add_run('Ventaja competitiva:').bold = True

p = doc.add_paragraph()
p.add_run('No competimos en precio. Competimos en valor.').italic = True

p = doc.add_paragraph()
p.add_run('Treqe no es solo una startup. Es una nueva forma de pensar la propiedad.').italic = True

# Guardar
output_path = os.path.join(os.path.dirname(__file__), 'Plan_Negocio_Treqe_DEFINITIVO.docx')
doc.save(output_path)

print(f"DOCUMENTO CREADO: {output_path}")
print(f"TAMANO: {os.path.getsize(output_path):,} bytes")
print("INCLUYE 10 partes completas")
print("SKILLS aplicadas: humanizer, legal, business-model-canvas, marketing-mode, frontend-design")