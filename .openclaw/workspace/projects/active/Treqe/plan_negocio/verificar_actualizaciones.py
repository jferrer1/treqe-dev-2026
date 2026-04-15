#!/usr/bin/env python3
"""
Verificar qué secciones se actualizaron en el documento
"""

from docx import Document

def verificar_actualizaciones():
    print("VERIFICANDO ACTUALIZACIONES REALIZADAS...")
    
    # Cargar documento actualizado
    try:
        doc = Document('Plan_Negocio_Treqe_HUMANIZADO_ACTUALIZADO.docx')
    except:
        print("ERROR: No se pudo cargar el documento actualizado")
        return
    
    print(f"Documento cargado: {len(doc.paragraphs)} parrafos")
    
    # Verificar secciones clave
    secciones_verificar = [
        ("ALGORITMO", ["algoritmo", "matching", "k=", "NP", "conclusiones"]),
        ("SCORING 8.4", ["8.4", "puntuacion", "scoring", "reputacion"]),
        ("ENVIOS 8.3", ["8.3", "envios", "triple proteccion"]),
        ("GARANTIAS 8.2", ["8.2", "garantias", "desistimiento"]),
        ("PROCESO USUARIO 3.2", ["3.2", "mecanismo", "proceso", "usuario"])
    ]
    
    for nombre_seccion, palabras_clave in secciones_verificar:
        print(f"\n--- {nombre_seccion} ---")
        
        encontrados = []
        for i, para in enumerate(doc.paragraphs):
            texto = para.text.strip().lower()
            if any(palabra in texto for palabra in palabras_clave):
                # Mostrar un extracto
                extracto = texto[:100] + "..." if len(texto) > 100 else texto
                encontrados.append((i, extracto))
        
        if encontrados:
            print(f"Encontrados {len(encontrados)} parrafos:")
            for i, extracto in encontrados[:3]:  # Mostrar solo primeros 3
                print(f"  [{i}] {extracto}")
            if len(encontrados) > 3:
                print(f"  ... y {len(encontrados)-3} mas")
        else:
            print("No encontrado")
    
    # Verificar si se añadieron conclusiones del algoritmo
    print("\n--- CONCLUSIONES DEL ALGORITMO ---")
    
    for i, para in enumerate(doc.paragraphs):
        texto = para.text.strip()
        if "lo que aprendimos" in texto.lower() or "durante el desarrollo" in texto.lower():
            print(f"Encontrado en parrafo {i}:")
            print(f"  {texto[:150]}...")
            break
    else:
        print("No se encontraron conclusiones explicitas")
    
    # Contar párrafos humanizados (más cortos, menos lenguaje AI)
    print("\n--- RESUMEN DE CAMBIOS ---")
    
    parrafos_largos = 0
    parrafos_medianos = 0
    parrafos_cortos = 0
    
    for para in doc.paragraphs:
        texto = para.text.strip()
        if not texto:
            continue
        
        longitud = len(texto)
        if longitud > 200:
            parrafos_largos += 1
        elif longitud > 50:
            parrafos_medianos += 1
        else:
            parrafos_cortos += 1
    
    print(f"Parrafos largos (>200 chars): {parrafos_largos}")
    print(f"Parrafos medianos (50-200): {parrafos_medianos}")
    print(f"Parrafos cortos (<50): {parrafos_cortos}")
    print(f"Total parrafos con texto: {parrafos_largos + parrafos_medianos + parrafos_cortos}")
    
    # Verificar lenguaje simplificado
    palabras_ai = ["ademas", "crucial", "destacar", "subrayar", "profundo", "vibrante"]
    palabras_simple = ["tambien", "importante", "mostrar", "enfatizar", "profundo", "activo"]
    
    print("\n--- LENGUAJE SIMPLIFICADO ---")
    
    for i in range(min(5, len(doc.paragraphs))):
        texto = doc.paragraphs[i].text.strip().lower()
        if texto:
            tiene_ai = any(palabra in texto for palabra in palabras_ai)
            tiene_simple = any(palabra in texto for palabra in palabras_simple)
            
            if tiene_ai or tiene_simple:
                print(f"Parrafo {i}:")
                if tiene_ai:
                    print("  Tiene lenguaje AI")
                if tiene_simple:
                    print("  Tiene lenguaje simplificado")
                print(f"  Extracto: {texto[:80]}...")

if __name__ == '__main__':
    verificar_actualizaciones()