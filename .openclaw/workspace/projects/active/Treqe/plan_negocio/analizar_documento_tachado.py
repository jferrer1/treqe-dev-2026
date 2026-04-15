#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script para analizar el documento Word con partes tachadas
y actualizar con las soluciones diseñadas para Treqe.
"""

import docx
import re
import os
from pathlib import Path

def analizar_documento_tachado(input_path):
    """Analiza el documento Word y extrae partes tachadas."""
    print("ANALIZANDO documento Word con partes tachadas...")
    
    # Cargar documento
    doc = docx.Document(input_path)
    print(f"Documento cargado: {len(doc.paragraphs)} párrafos")
    
    # Analizar estructura
    secciones_tachadas = []
    secciones_completas = []
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        
        # Detectar texto tachado
        has_strikethrough = False
        tachado_text = ""
        for run in para.runs:
            if run.font.strike:
                has_strikethrough = True
                tachado_text += run.text
        
        # Detectar estilo
        style = para.style.name
        
        if has_strikethrough:
            secciones_tachadas.append({
                'indice': i,
                'texto': text,
                'tachado': tachado_text,
                'estilo': style,
                'contexto': []
            })
            
            # Obtener contexto (párrafos anteriores y siguientes)
            contexto = []
            for j in range(max(0, i-2), min(len(doc.paragraphs), i+3)):
                if j != i:
                    contexto.append({
                        'indice': j,
                        'texto': doc.paragraphs[j].text.strip(),
                        'estilo': doc.paragraphs[j].style.name
                    })
            secciones_tachadas[-1]['contexto'] = contexto
    
    print(f"🎯 Secciones tachadas encontradas: {len(secciones_tachadas)}")
    
    # Agrupar por temas
    temas = {
        'chat_grupal': [],
        'devoluciones': [],
        'envios': [],
        'garantias': [],
        'algoritmo': [],
        'reputacion': [],
        'otros': []
    }
    
    palabras_clave = {
        'chat_grupal': ['chat', 'grupo', 'conversación', 'negociación', 'oficial', 'estructurada'],
        'devoluciones': ['devolución', 'desistimiento', '24 horas', 'garantía', 'fondo'],
        'envios': ['envío', 'transporte', 'logística', 'correos', 'paquete', 'entrega'],
        'garantias': ['garantía', 'seguro', 'escrow', 'depósito', 'compensación'],
        'algoritmo': ['algoritmo', 'matching', 'k=', 'ciclo', 'rueda', 'intercambio'],
        'reputacion': ['reputación', 'scoring', 'puntuación', 'confianza', 'historial']
    }
    
    for seccion in secciones_tachadas:
        texto = seccion['texto'].lower()
        tema_encontrado = 'otros'
        
        for tema, palabras in palabras_clave.items():
            for palabra in palabras:
                if palabra in texto:
                    tema_encontrado = tema
                    break
            if tema_encontrado != 'otros':
                break
        
        temas[tema_encontrado].append(seccion)
    
    # Mostrar resumen
    print("\n📊 DISTRIBUCIÓN DE SECCIONES TACHADAS:")
    for tema, secciones in temas.items():
        if secciones:
            print(f"  • {tema}: {len(secciones)} secciones")
    
    return temas, doc

def identificar_soluciones_necesarias(temas):
    """Identifica qué soluciones necesitan actualización."""
    print("\n🔧 IDENTIFICANDO SOLUCIONES NECESARIAS:")
    
    soluciones_necesarias = []
    
    # Mapeo de temas a soluciones
    mapeo_soluciones = {
        'chat_grupal': {
            'nombre': 'Sistema de ofertas estructuradas',
            'skill': 'business-model-canvas + frontend-design',
            'descripcion': 'Reemplazar chat grupal por sistema automático de propuestas'
        },
        'devoluciones': {
            'nombre': 'Desistimiento en casos excepcionales',
            'skill': 'legal + humanizer',
            'descripcion': 'Restringir devoluciones a casos claros con verificación'
        },
        'envios': {
            'nombre': 'Sistema de envíos con garantías',
            'skill': 'legal + business-model-canvas',
            'descripcion': 'Depósito con tarjeta + sistema de reputación'
        },
        'garantias': {
            'nombre': 'Sistema combinado de garantías',
            'skill': 'legal + business-model-canvas',
            'descripcion': 'Combinación depósito + reputación + partner fintech'
        },
        'algoritmo': {
            'nombre': 'Algoritmo ascendente k=2→k_max',
            'skill': 'business-model-canvas',
            'descripcion': 'Estrategia ascendente con timeout inteligente'
        },
        'reputacion': {
            'nombre': 'Sistema de scoring/reputación',
            'skill': 'business-model-canvas + marketing-mode',
            'descripcion': 'Métrica compuesta con niveles y beneficios'
        }
    }
    
    for tema, secciones in temas.items():
        if secciones and tema in mapeo_soluciones:
            solucion = mapeo_soluciones[tema]
            soluciones_necesarias.append({
                'tema': tema,
                'secciones': secciones,
                'solucion': solucion
            })
            
            print(f"\n🎯 {solucion['nombre']}:")
            print(f"   • Skill: {solucion['skill']}")
            print(f"   • Descripción: {solucion['descripcion']}")
            print(f"   • Secciones afectadas: {len(secciones)}")
    
    return soluciones_necesarias

def generar_plan_actualizacion(soluciones_necesarias):
    """Genera un plan de actualización detallado."""
    print("\n📋 PLAN DE ACTUALIZACIÓN:")
    
    plan = []
    
    for item in soluciones_necesarias:
        tema = item['tema']
        solucion = item['solucion']
        secciones = item['secciones']
        
        # Agrupar secciones por contexto
        secciones_agrupadas = []
        for seccion in secciones:
            # Buscar títulos cercanos
            titulo = "Sin título"
            for ctx in seccion['contexto']:
                if ctx['estilo'].startswith('Heading'):
                    titulo = ctx['texto']
                    break
            
            secciones_agrupadas.append({
                'indice': seccion['indice'],
                'titulo': titulo,
                'texto_tachado': seccion['tachado'][:100] + "..." if len(seccion['tachado']) > 100 else seccion['tachado']
            })
        
        plan.append({
            'tema': tema,
            'solucion': solucion,
            'secciones': secciones_agrupadas,
            'acciones': [
                f"Leer sección: {solucion['nombre']}",
                f"Aplicar skill: {solucion['skill']}",
                f"Actualizar con: {solucion['descripcion']}",
                "Mantener contexto y estructura original",
                "Aplicar humanizer al lenguaje"
            ]
        })
        
        print(f"\n📝 {solucion['nombre'].upper()}:")
        print(f"   Skills: {solucion['skill']}")
        for accion in plan[-1]['acciones']:
            print(f"   • {accion}")
        print(f"   Secciones a actualizar: {len(secciones_agrupadas)}")
        for sec in secciones_agrupadas[:3]:  # Mostrar solo 3
            print(f"     - [{sec['indice']}] {sec['titulo']}: {sec['texto_tachado']}")
        if len(secciones_agrupadas) > 3:
            print(f"     ... y {len(secciones_agrupadas) - 3} más")
    
    return plan

def main():
    """Función principal."""
    input_path = "Plan_Negocio_Treqe_CON_SOLUCIONES_2026.docx"
    
    if not os.path.exists(input_path):
        print(f"❌ Archivo no encontrado: {input_path}")
        return
    
    # 1. Analizar documento
    temas, doc = analizar_documento_tachado(input_path)
    
    # 2. Identificar soluciones necesarias
    soluciones_necesarias = identificar_soluciones_necesarias(temas)
    
    # 3. Generar plan de actualización
    plan = generar_plan_actualizacion(soluciones_necesarias)
    
    # 4. Guardar plan
    output_path = "plan_actualizacion_soluciones_treqe.md"
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write("# PLAN DE ACTUALIZACIÓN - SOLUCIONES TREQE\n\n")
        f.write("## 📋 Resumen Ejecutivo\n\n")
        f.write(f"Documento analizado: {input_path}\n")
        f.write(f"Secciones tachadas encontradas: {sum(len(temas[t]) for t in temas)}\n")
        f.write(f"Soluciones identificadas: {len(soluciones_necesarias)}\n\n")
        
        f.write("## 🎯 Soluciones a Implementar\n\n")
        for item in plan:
            solucion = item['solucion']
            f.write(f"### {solucion['nombre']}\n\n")
            f.write(f"**Skills requeridas:** {solucion['skill']}\n\n")
            f.write(f"**Descripción:** {solucion['descripcion']}\n\n")
            f.write("**Acciones:**\n")
            for accion in item['acciones']:
                f.write(f"- {accion}\n")
            f.write("\n**Secciones afectadas:**\n")
            for sec in item['secciones']:
                f.write(f"- [{sec['indice']}] {sec['titulo']}\n")
            f.write("\n---\n\n")
        
        f.write("## 📝 Notas Importantes\n\n")
        f.write("1. **NO BORRAR** contenido existente, solo actualizar\n")
        f.write("2. **MANTENER** estructura y contexto original\n")
        f.write("3. **APLICAR HUMANIZER** a todo el lenguaje\n")
        f.write("4. **USAR SKILLS POR SECCIÓN** según corresponda\n")
        f.write("5. **VALIDAR** con usuario antes de implementación final\n")
    
    print(f"\n✅ Plan de actualización guardado en: {output_path}")
    print("\n🚀 **PRÓXIMOS PASOS:**")
    print("1. Revisar el plan de actualización")
    print("2. Aplicar skills según cada sección")
    print("3. Mantener todo el contenido no tachado")
    print("4. Validar actualizaciones con el usuario")

if __name__ == "__main__":
    main()