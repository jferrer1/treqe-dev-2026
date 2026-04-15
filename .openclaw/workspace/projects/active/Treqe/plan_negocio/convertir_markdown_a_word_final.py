#!/usr/bin/env python3
"""
Convertir el contenido Markdown final a documento Word
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import re

def convertir_a_word():
    print("CONVIRTIENDO CONTENIDO FINAL A DOCUMENTO WORD...")
    
    # Leer contenido Markdown
    with open('contenido_final_actualizado.md', 'r', encoding='utf-8') as f:
        contenido = f.read()
    
    # Crear documento
    doc = Document()
    
    # Configurar estilos
    styles = doc.styles
    
    # Estilo para título principal
    if 'TituloPrincipal' not in styles:
        title_style = styles.add_style('TituloPrincipal', WD_STYLE_TYPE.PARAGRAPH)
        title_style.font.name = 'Calibri'
        title_style.font.size = Pt(20)
        title_style.font.bold = True
    
    # Estilo para títulos de sección
    if 'TituloSeccion' not in styles:
        section_style = styles.add_style('TituloSeccion', WD_STYLE_TYPE.PARAGRAPH)
        section_style.font.name = 'Calibri'
        section_style.font.size = Pt(16)
        section_style.font.bold = True
    
    # Estilo para subtítulos
    if 'Subtitulo' not in styles:
        subtitle_style = styles.add_style('Subtitulo', WD_STYLE_TYPE.PARAGRAPH)
        subtitle_style.font.name = 'Calibri'
        subtitle_style.font.size = Pt(14)
        subtitle_style.font.bold = True
    
    # Estilo para texto normal
    if 'TextoNormal' not in styles:
        normal_style = styles.add_style('TextoNormal', WD_STYLE_TYPE.PARAGRAPH)
        normal_style.font.name = 'Calibri'
        normal_style.font.size = Pt(11)
    
    # Estilo para listas
    if 'Lista' not in styles:
        list_style = styles.add_style('Lista', WD_STYLE_TYPE.PARAGRAPH)
        list_style.font.name = 'Calibri'
        list_style.font.size = Pt(11)
        list_style.paragraph_format.left_indent = Inches(0.25)
    
    # Procesar contenido
    lines = contenido.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        if not line:
            i += 1
            continue
        
        # Título principal
        if line.startswith('# ') and 'PLAN DE NEGOCIO' in line:
            # Título principal centrado
            title = doc.add_paragraph()
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title.add_run(line.replace('# ', ''))
            title_run.font.size = Pt(20)
            title_run.font.bold = True
            title_run.font.name = 'Calibri'
            
            # Información debajo
            i += 1
            while i < len(lines) and lines[i].strip() and not lines[i].startswith('---'):
                info_line = lines[i].strip()
                if info_line:
                    info = doc.add_paragraph()
                    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    info_run = info.add_run(info_line)
                    info_run.font.size = Pt(10)
                    info_run.font.name = 'Calibri'
                    if 'CONFIDENCIAL' in info_line:
                        info_run.bold = True
                i += 1
            
            doc.add_page_break()
        
        # Títulos de sección (##)
        elif line.startswith('## '):
            p = doc.add_paragraph()
            p_run = p.add_run(line.replace('## ', ''))
            p_run.font.size = Pt(16)
            p_run.font.bold = True
            p_run.font.name = 'Calibri'
        
        # Subtítulos (###)
        elif line.startswith('### '):
            p = doc.add_paragraph()
            p_run = p.add_run(line.replace('### ', ''))
            p_run.font.size = Pt(14)
            p_run.font.bold = True
            p_run.font.name = 'Calibri'
        
        # Listas con guiones
        elif line.startswith('- **') or line.startswith('- '):
            # Procesar lista
            p = doc.add_paragraph(style='Lista')
            
            # Manejar texto en negrita
            if '**' in line:
                parts = line.split('**')
                for j, part in enumerate(parts):
                    if j % 2 == 0:  # Texto normal
                        if part.startswith('- '):
                            part = part[2:]  # Quitar el guión
                        p_run = p.add_run(part)
                        p_run.font.name = 'Calibri'
                        p_run.font.size = Pt(11)
                    else:  # Texto en negrita
                        p_run = p.add_run(part)
                        p_run.font.name = 'Calibri'
                        p_run.font.size = Pt(11)
                        p_run.bold = True
            else:
                p_run = p.add_run(line[2:])  # Quitar el guión
                p_run.font.name = 'Calibri'
                p_run.font.size = Pt(11)
        
        # Texto en negrita (pero no títulos)
        elif line.startswith('**') and line.endswith('**'):
            p = doc.add_paragraph()
            p_run = p.add_run(line[2:-2])  # Quitar **
            p_run.font.name = 'Calibri'
            p_run.font.size = Pt(11)
            p_run.bold = True
        
        # Texto normal
        elif line and not line.startswith('---'):
            p = doc.add_paragraph(style='TextoNormal')
            p_run = p.add_run(line)
            p_run.font.name = 'Calibri'
            p_run.font.size = Pt(11)
        
        # Separador de página (---)
        elif line.startswith('---'):
            doc.add_page_break()
        
        i += 1
    
    # Guardar documento
    output_path = 'Plan_Negocio_Treqe_FINAL_ACTUALIZADO.docx'
    doc.save(output_path)
    
    print(f"Documento creado: {output_path}")
    print(f"Tamaño del contenido: {len(contenido):,} caracteres")
    
    return output_path

if __name__ == '__main__':
    convertir_a_word()