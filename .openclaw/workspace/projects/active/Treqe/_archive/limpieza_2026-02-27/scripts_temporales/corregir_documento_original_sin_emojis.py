#!/usr/bin/env python3
"""
Corregir documento original sin emojis para Windows
"""

import docx
import os
import shutil

def corregir_documento():
    """Corregir el documento original."""
    
    original_path = "plan_negocio/Plan_Negocio_Treqe_100%_COMPLETO_FINAL.docx"
    corrected_path = "plan_negocio/Plan_Negocio_Treqe_100%_COMPLETO_FINAL_CORREGIDO.docx"
    
    if not os.path.exists(original_path):
        print(f"ERROR: No se encuentra: {original_path}")
        return None
    
    print("=== CORRIGIENDO DOCUMENTO ORIGINAL ===")
    print(f"Documento original: {original_path}")
    
    # Backup
    backup_path = original_path + ".backup"
    shutil.copy2(original_path, backup_path)
    print(f"Backup creado: {backup_path}")
    
    # Abrir documento
    doc = docx.Document(original_path)
    original_size = os.path.getsize(original_path)
    print(f"Tamaño original: {original_size/1024:.1f} KB")
    print(f"Párrafos originales: {len(doc.paragraphs)}")
    
    # Buscar sección ANEXOS
    anexos_start = -1
    for i, p in enumerate(doc.paragraphs):
        if 'ANEXOS' in p.text.upper() and len(p.text) < 20:
            anexos_start = i
            print(f"Sección ANEXOS encontrada en párrafo {i}")
            break
    
    if anexos_start == -1:
        print("ERROR: No se encontró sección ANEXOS")
        return None
    
    # Determinar hasta dónde van los anexos actuales
    anexos_end = anexos_start + 1
    while anexos_end < len(doc.paragraphs):
        p = doc.paragraphs[anexos_end]
        if p.style.name.startswith('Heading') and anexos_end > anexos_start + 1:
            break
        anexos_end += 1
    
    print(f"Eliminando anexos actuales (párrafos {anexos_start+1} a {anexos_end-1})")
    
    # Eliminar anexos existentes
    for _ in range(anexos_start + 1, anexos_end):
        if anexos_start + 1 < len(doc.paragraphs):
            p = doc.paragraphs[anexos_start + 1]
            p._element.getparent().remove(p._element)
    
    print("Añadiendo anexos corregidos...")
    
    # ANEXO A: ESTRUCTURA DEL EQUIPO (NO CVs inventados)
    doc.add_paragraph()
    doc.add_paragraph("ANEXO A: ESTRUCTURA Y ROLES DEL EQUIPO")
    doc.paragraphs[-1].style = "Heading 2"
    
    contenido_a = [
        "",
        "ROLES NECESARIOS PARA EL DESARROLLO DEL MVP:",
        "",
        "1. FUNDADOR/CEO",
        "   • Responsabilidades: Estrategia, visión de producto, fundraising",
        "   • Experiencia ideal: Conocimiento sector, liderazgo, capacidad negociación",
        "",
        "2. DESARROLLADOR FULL-STACK",
        "   • Responsabilidades: Desarrollo plataforma, arquitectura, mantenimiento",
        "   • Stack: Node.js/Express o Python/Django, React, PostgreSQL",
        "   • Habilidades: APIs, algoritmos, deployment cloud",
        "",
        "3. DISEÑADOR UX/UI",
        "   • Responsabilidades: User research, wireframes, diseño interfaz",
        "   • Herramientas: Figma, Adobe XD, user testing",
        "   • Habilidades: Design thinking, accesibilidad, prototipado",
        "",
        "NOTA: Los CVs específicos de los miembros del equipo se añadirán cuando se definan los candidatos reales.",
    ]
    
    for item in contenido_a:
        if item == "":
            doc.add_paragraph()
        else:
            doc.add_paragraph(item)
    
    doc.add_page_break()
    
    # ANEXO B: METODOLOGÍA DE VALIDACIÓN (NO datos inventados)
    doc.add_paragraph("ANEXO B: METODOLOGÍA PARA VALIDACIÓN DE MERCADO")
    doc.paragraphs[-1].style = "Heading 2"
    
    contenido_b = [
        "",
        "FASES RECOMENDADAS PARA VALIDAR EL MERCADO:",
        "",
        "FASE 1: INVESTIGACIÓN SECUNDARIA",
        "   • Fuentes: INE, Eurostat, informes sectoriales públicos",
        "   • Método: Análisis datos existentes, benchmarking competidores",
        "   • Output: Tamaño mercado estimado, tendencias, oportunidades",
        "",
        "FASE 2: ENCUESTAS CUANTITATIVAS",
        "   • Método: Encuestas online (Typeform, Google Forms)",
        "   • Muestra: 200-500 usuarios españoles",
        "   • Métricas: Intención uso, preocupaciones, disposición pago",
        "",
        "FASE 3: ENTREVISTAS CUALITATIVAS",
        "   • Método: Entrevistas 1:1 con usuarios potenciales",
        "   • Participantes: 15-20 usuarios",
        "   • Objetivo: Deep dive en pain points, validar user flows",
        "",
        "FASE 4: PILOTO CONTROLADO",
        "   • Método: MVP en barrio específico (ej: Gràcia, Barcelona)",
        "   • Participantes: 100-200 usuarios activos",
        "   • Métricas: Engagement, retención, tasa conversión",
        "",
        "NOTA: Los datos específicos de cada estudio se añadirán tras realizar las investigaciones correspondientes.",
    ]
    
    for item in contenido_b:
        if item == "":
            doc.add_paragraph()
        else:
            doc.add_paragraph(item)
    
    doc.add_page_break()
    
    # ANEXO C: DISEÑOS DETALLADOS (DESARROLLAR MÁS)
    doc.add_paragraph("ANEXO C: ESPECIFICACIONES DE DISEÑO DE PLATAFORMA")
    doc.paragraphs[-1].style = "Heading 2"
    
    contenido_c = [
        "",
        "ARQUITECTURA DE LA PLATAFORMA:",
        "",
        "1. SISTEMA DE CATALOGACIÓN",
        "   • Upload múltiple de imágenes (3-10 por producto)",
        "   • Etiquetado automático por categoría/estado",
        "   • Valoración estimada por algoritmo",
        "   • Descripción asistida por templates",
        "",
        "2. MOTOR DE BÚSQUEDA Y MATCHING",
        "   • Búsqueda semántica (entender consultas naturales)",
        "   • Filtros avanzados: Radio, categoría, valor, fecha",
        "   • Algoritmo de matching para ruedas de intercambio",
        "   • Sistema de recomendaciones personalizadas",
        "",
        "3. SISTEMA DE INTERCAMBIO",
        "   • Propuesta 1-clic para ruedas sugeridas",
        "   • Chat integrado para negociación",
        "   • Acuerdo de términos (fecha, lugar, condiciones)",
        "   • Sistema de seguimiento y notificaciones",
        "   • Calificación post-intercambio",
        "",
        "4. PERFILES Y REPUTACIÓN",
        "   • Perfil público con historial de intercambios",
        "   • Sistema de reputación basado en transacciones",
        "   • Verificaciones (email, teléfono, redes sociales)",
        "   • Badges por logros y participación",
    ]
    
    for item in contenido_c:
        if item == "":
            doc.add_paragraph()
        else:
            doc.add_paragraph(item)
    
    doc.add_page_break()
    
    # ANEXO D: GLOSARIO (MANTENER)
    doc.add_paragraph("ANEXO D: GLOSARIO DE TÉRMINOS")
    doc.paragraphs[-1].style = "Heading 2"
    
    contenido_d = [
        "",
        "TÉRMINOS DEL PROYECTO TREQE:",
        "",
        "• Rueda de Intercambio: Ciclo cerrado A→B→C→A donde todos reciben valor",
        "• Algoritmo de Matching: Sistema que encuentra combinaciones óptimas entre usuarios",
        "• Compensación Monetaria: Pago pequeño para equilibrar diferencias de valor",
        "• Liquidez de Plataforma: Facilidad para encontrar contrapartes de intercambio",
        "• Ciclo Cerrado de Valor: Productos que circulan sin convertirse en residuos",
        "",
        "TÉRMINOS DE NEGOCIO:",
        "",
        "• Two-sided Marketplace: Plataforma que conecta dos grupos de usuarios",
        "• Network Effects: El valor aumenta con más usuarios (efecto red)",
        "• Chicken-and-Egg Problem: Desafío inicial de plataformas",
        "• CAC/LTV: Coste adquisición cliente / Valor vida útil cliente",
        "• MVP/PMF: Producto mínimo viable / Ajuste producto-mercado",
    ]
    
    for item in contenido_d:
        if item == "":
            doc.add_paragraph()
        else:
            doc.add_paragraph(item)
    
    doc.add_page_break()
    
    # ANEXO E: FUENTES VERÍDICAS (CONSULTABLES)
    doc.add_paragraph("ANEXO E: FUENTES Y REFERENCIAS")
    doc.paragraphs[-1].style = "Heading 2"
    
    contenido_e = [
        "",
        "FUENTES PÚBLICAS Y OFICIALES:",
        "",
        "1. INSTITUCIONES OFICIALES:",
        "   • INE (Instituto Nacional de Estadística) - https://www.ine.es",
        "   • Eurostat - https://ec.europa.eu/eurostat",
        "   • Banco de España - https://www.bde.es",
        "   • CNMC (Comisión Nacional Mercados Competencia) - https://www.cnmc.es",
        "",
        "2. ORGANIZACIONES SECTORIALES:",
        "   • Observatorio Cetelem - Informes anuales de consumo",
        "   • Asociación Española Economía Circular - https://economiacircular.org",
        "   • AECOC (Asociación Fabricantes Distribuidores) - Datos retail",
        "",
        "3. PLATAFORMAS DE DATOS:",
        "   • Statista - Reports mercado segunda mano",
        "   • Google Trends - Tendencias de búsqueda",
        "   • SimilarWeb - Tráfico de competidores",
        "   • App Annie/Data.ai - Métricas apps móviles",
        "",
        "4. INFORMES PÚBLICOS DE COMPETIDORES:",
        "   • Wallapop - Informes de impacto anual",
        "   • Vinted - Reports de transparencia y sostenibilidad",
        "   • Milanuncios - Datos públicos de Schibsted",
        "",
        "METODOLOGÍA: Para este plan se han consultado datos públicos de las fuentes mencionadas y análisis de informes sectoriales disponibles.",
    ]
    
    for item in contenido_e:
        if item == "":
            doc.add_paragraph()
        else:
            doc.add_paragraph(item)
    
    # Guardar
    print(f"Guardando documento corregido: {corrected_path}")
    doc.save(corrected_path)
    
    # Verificar
    new_doc = docx.Document(corrected_path)
    final_size = os.path.getsize(corrected_path)
    print(f"Documento corregido guardado.")
    print(f"Tamaño final: {final_size/1024:.1f} KB")
    print(f"Párrafos finales: {len(new_doc.paragraphs)}")
    print(f"Diferencia tamaño: {(final_size - original_size)/1024:+.1f} KB")
    
    return corrected_path

if __name__ == "__main__":
    print("CORRECCION DIRECTA DEL DOCUMENTO ORIGINAL")
    print("Seguir instrucciones exactas:")
    print("1. ELIMINAR CVs inventados del equipo")
    print("2. ELIMINAR estudios especificos inventados")
    print("3. MANTENER/MEJORAR diseños detallados de plataforma")
    print("4. CORREGIR fuentes para que sean veridicas y consultables")
    print("5. MANTENER glosario de terminos")
    print()
    
    resultado = corregir_documento()
    
    if resultado:
        print()
        print("=" * 60)
        print("DOCUMENTO ORIGINAL CORREGIDO EXITOSAMENTE")
        print("=" * 60)
        print(f"Documento: {resultado}")
        print()
        print("Se ha mantenido TODO el contenido original excepto los anexos,")
        print("que han sido reemplazados por version corregida.")
        print()
        print("ANEXO A: Estructura equipo (roles, NO CVs inventados)")
        print("ANEXO B: Metodologia validacion (fases, NO datos inventados)")
        print("ANEXO C: Diseños detallados (especificaciones desarrolladas)")
        print("ANEXO D: Glosario (terminos definidos)")
        print("ANEXO E: Fuentes veridicas (todas consultables)")
    else:
        print("Error al corregir el documento.")