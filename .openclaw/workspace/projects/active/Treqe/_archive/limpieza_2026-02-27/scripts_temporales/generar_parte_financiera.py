#!/usr/bin/env python3
"""
Generar la PARTE FINANCIERA COMPLETA del plan de negocio Treqe
Incluye: inversión inicial, proyecciones, cash flow, punto equilibrio, ratios
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
from pathlib import Path

def create_financial_section():
    """Crear documento con la parte financiera completa."""
    
    doc = Document()
    
    # Título del documento
    doc.add_heading('PLAN FINANCIERO COMPLETO - TREQE', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Parte 9 del Plan de Negocio Profesional', style='Subtitle').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f'Fecha: {datetime.now().strftime("%d de %B de %Y")}').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    doc.add_page_break()
    
    # ========== 9.1 INVERSIÓN INICIAL ==========
    doc.add_heading('9.1 Inversión Inicial y Estructura de Costes', 1)
    doc.add_paragraph()
    
    doc.add_heading('Desglose Detallado de la Inversión Inicial (58.000€)', 2)
    
    # Tabla de inversión inicial
    inv_table = doc.add_table(rows=12, cols=3)
    inv_table.style = 'Light Grid Accent 1'
    
    # Encabezados
    headers = ['Concepto', 'Detalle', 'Importe (€)']
    for i, header in enumerate(headers):
        inv_table.cell(0, i).text = header
        inv_table.cell(0, i).paragraphs[0].runs[0].bold = True
    
    # Datos
    investment_data = [
        ['DESARROLLO TECNOLÓGICO (40%)', 'Plataforma completa + algoritmos', '23.200'],
        ['  • Frontend (Next.js + React)', 'Desarrollo 3 meses', '8.000'],
        ['  • Backend (Node.js + Python)', 'Microservicios + APIs', '7.500'],
        ['  • Algoritmos de matching', 'Desarrollo propietario', '4.200'],
        ['  • Infraestructura cloud', 'Vercel + Railway + servicios', '3.500'],
        ['MARKETING INICIAL (35%)', 'Lanzamiento + adquisición usuarios', '20.300'],
        ['  • Campañas digitales', 'Google Ads + Social Media', '9.000'],
        ['  • Contenido + SEO', 'Blog, guías, contenido viral', '5.300'],
        ['  • Eventos + partnerships', 'Lanzamiento Barcelona', '6.000'],
        ['OPERACIONES Y EQUIPO (25%)', 'Gastos generales + equipo fundador', '14.500'],
        ['  • Equipo fundador (3 personas)', '6 meses a media jornada', '9.000'],
        ['  • Legal + administrativo', 'SL, contratos, protección IP', '3.500'],
        ['  • Oficina + herramientas', 'Coworking + software', '2.000'],
        ['TOTAL INVERSIÓN INICIAL', '', '58.000']
    ]
    
    for i, (concepto, detalle, importe) in enumerate(investment_data, 1):
        inv_table.cell(i, 0).text = concepto
        inv_table.cell(i, 1).text = detalle
        inv_table.cell(i, 2).text = importe
    
    doc.add_paragraph()
    doc.add_paragraph('**Financiación propuesta:**')
    doc.add_paragraph('• 40.000€ - Inversores ángeles / business angels', style='List Bullet')
    doc.add_paragraph('• 10.000€ - Préstamo participativo ENISA', style='List Bullet')
    doc.add_paragraph('• 8.000€ - Aportación equipo fundador', style='List Bullet')
    
    doc.add_page_break()
    
    # ========== 9.2 PROYECCIONES DE INGRESOS ==========
    doc.add_heading('9.2 Proyecciones de Ingresos Detalladas', 1)
    doc.add_paragraph()
    
    doc.add_paragraph('**Supuestos clave para las proyecciones:**')
    assumptions = [
        'Comisión del 1% sobre valor del artículo adquirido',
        'Valor medio por transacción: 150€ (año 1), 160€ (año 2), 170€ (año 3)',
        'Tasa de conversión usuarios activos → transacciones: 10% mensual',
        'Crecimiento orgánico: 15% mensual (año 1), 10% mensual (año 2), 5% mensual (año 3)',
        'Retención de usuarios: 70% mensual'
    ]
    
    for assumption in assumptions:
        doc.add_paragraph(f'• {assumption}', style='List Bullet')
    
    doc.add_paragraph()
    doc.add_heading('Proyecciones Mensuales Detalladas - Año 1', 2)
    
    # Tabla proyecciones mensuales año 1
    monthly_table = doc.add_table(rows=13, cols=7)
    monthly_table.style = 'Light Grid Accent 2'
    
    m_headers = ['Mes', 'Usuarios Activos', 'Nuevos Usuarios', 'Transacciones', 'Volumen (€)', 'Comisiones (€)', 'Acumulado (€)']
    for i, header in enumerate(m_headers):
        monthly_table.cell(0, i).text = header
        monthly_table.cell(0, i).paragraphs[0].runs[0].bold = True
    
    monthly_data = [
        ['1', '500', '500', '50', '7.500', '75', '75'],
        ['2', '850', '350', '85', '12.750', '128', '203'],
        ['3', '1.275', '425', '128', '19.200', '192', '395'],
        ['4', '1.850', '575', '185', '27.750', '278', '673'],
        ['5', '2.600', '750', '260', '39.000', '390', '1.063'],
        ['6', '3.600', '1.000', '360', '54.000', '540', '1.603'],
        ['7', '4.900', '1.300', '490', '73.500', '735', '2.338'],
        ['8', '6.500', '1.600', '650', '97.500', '975', '3.313'],
        ['9', '8.400', '1.900', '840', '126.000', '1.260', '4.573'],
        ['10', '10.700', '2.300', '1.070', '160.500', '1.605', '6.178'],
        ['11', '13.400', '2.700', '1.340', '201.000', '2.010', '8.188'],
        ['12', '16.600', '3.200', '1.660', '249.000', '2.490', '10.678']
    ]
    
    for i, row in enumerate(monthly_data, 1):
        for j, value in enumerate(row):
            monthly_table.cell(i, j).text = value
    
    doc.add_paragraph()
    doc.add_heading('Proyecciones Anuales Consolidadas', 2)
    
    annual_table = doc.add_table(rows=4, cols=5)
    annual_table.style = 'Light Shading Accent 1'
    
    a_headers = ['Año', 'Usuarios Finales', 'Transacciones', 'Volumen Total (€)', 'Ingresos Comisiones (€)']
    for i, header in enumerate(a_headers):
        annual_table.cell(0, i).text = header
        annual_table.cell(0, i).paragraphs[0].runs[0].bold = True
    
    annual_data = [
        ['1', '16.600', '7.150', '1.072.500', '10.725'],
        ['2', '75.000', '30.000', '4.800.000', '48.000'],
        ['3', '150.000', '60.000', '10.200.000', '102.000']
    ]
    
    for i, row in enumerate(annual_data, 1):
        for j, value in enumerate(row):
            annual_table.cell(i, j).text = value
    
    doc.add_page_break()
    
    # ========== 9.3 ESTADO DE PÉRDIDAS Y GANANCIAS ==========
    doc.add_heading('9.3 Estado de Pérdidas y Ganancias Proyectado', 1)
    doc.add_paragraph()
    
    pyg_table = doc.add_table(rows=15, cols=4)
    pyg_table.style = 'Light Grid Accent 1'
    
    pyg_headers = ['Concepto', 'Año 1 (€)', 'Año 2 (€)', 'Año 3 (€)']
    for i, header in enumerate(pyg_headers):
        pyg_table.cell(0, i).text = header
        pyg_table.cell(0, i).paragraphs[0].runs[0].bold = True
    
    pyg_data = [
        ['**INGRESOS**', '', '', ''],
        ['Comisiones plataforma', '10.725', '48.000', '102.000'],
        ['Servicios premium', '-', '12.000', '30.000'],
        ['Publicidad segmentada', '-', '6.000', '18.000'],
        ['**TOTAL INGRESOS**', '**10.725**', '**66.000**', '**150.000**'],
        ['', '', '', ''],
        ['**GASTOS OPERATIVOS**', '', '', ''],
        ['Costes desarrollo', '23.200', '15.000', '20.000'],
        ['Marketing y ventas', '20.300', '25.000', '30.000'],
        ['Personal', '9.000', '18.000', '30.000'],
        ['Alquiler y servicios', '2.000', '3.000', '4.000'],
        ['Administración y legal', '3.500', '4.000', '5.000'],
        ['Amortizaciones', '1.000', '1.500', '2.000'],
        ['**TOTAL GASTOS**', '**58.000**', '**66.500**', '**91.000**'],
        ['', '', '', ''],
        ['**RESULTADO OPERATIVO (EBITDA)**', '**-47.275**', '**-500**', '**59.000**'],
        ['Amortizaciones', '(1.000)', '(1.500)', '(2.000)'],
        ['**RESULTADO ANTES IMPUESTOS**', '**-48.275**', '**-2.000**', '**57.000**'],
        ['Impuestos (25%)', '-', '-', '14.250'],
        ['**RESULTADO NETO**', '**-48.275**', '**-2.000**', '**42.750**']
    ]
    
    for i, row in enumerate(pyg_data, 1):
        for j, value in enumerate(row):
            pyg_table.cell(i, j).text = value
    
    doc.add_paragraph()
    doc.add_paragraph('**Análisis de márgenes:**')
    doc.add_paragraph('• Margen bruto: 100% (comisiones son ingreso puro)', style='List Bullet')
    doc.add_paragraph('• Margen EBITDA año 3: 39,3%', style='List Bullet')
    doc.add_paragraph('• Margen neto año 3: 28,5%', style='List Bullet')
    
    doc.add_page_break()
    
    # ========== 9.4 CASH FLOW PROYECTADO ==========
    doc.add_heading('9.4 Cash Flow Proyectado (Mensual y Anual)', 1)
    doc.add_paragraph()
    
    doc.add_heading('Cash Flow Mensual - Primeros 12 Meses', 2)
    
    cf_monthly_table = doc.add_table(rows=13, cols=6)
    cf_monthly_table.style = 'Light Grid Accent 2'
    
    cf_headers = ['Mes', 'Entradas (€)', 'Salidas (€)', 'Cash Flow Mensual', 'Acumulado', 'Comentario']
    for i, header in enumerate(cf_headers):
        cf_monthly_table.cell(0, i).text = header
        cf_monthly_table.cell(0, i).paragraphs[0].runs[0].bold = True
    
    cf_monthly_data = [
        ['0', '58.000', '0', '58.000', '58.000', 'Inversión inicial'],
        ['1', '75', '12.000', '-11.925', '46.075', 'Lanzamiento + desarrollo'],
        ['2', '128', '8.000', '-7.872', '38.203', 'Marketing inicial'],
        ['3', '192', '6.000', '-5.808', '32.395', 'Optimización plataforma'],
        ['4', '278', '5.000', '-4.722', '27.673', 'Primera comunidad activa'],
        ['5', '390', '4.000', '-3.610', '24.063', 'Crecimiento orgánico'],
        ['6', '540', '3.500', '-2.960', '21.103', 'Mejora retención'],
        ['7', '735', '3.000', '-2.265', '18.838', 'Expansión categorías'],
        ['8', '975', '2.500', '-1.525', '17.313', 'Primeros partnerships'],
        ['9', '1.260', '2.500', '-1.240', '16.073', 'Optimización costes'],
        ['10', '1.605', '2.500', '-895', '15.178', 'Punto equilibrio cercano'],
        ['11', '2.010', '2.500', '-490', '14.688', 'Cash flow casi neutral'],
        ['12', '2.490', '2.500', '-10', '14.678', 'Equilibrio operativo']
    ]
    
    for i, row in enumerate(cf_monthly_data, 1):
        for j, value in enumerate(row):
            cf_monthly_table.cell(i, j).text = value
    
    doc.add_paragraph()
    doc.add_heading('Cash Flow Anual Consolidado', 2)
    
    cf_annual_table = doc.add_table(rows=5, cols=4)
    cf_annual_table.style = 'Light Shading Accent 1'
    
    cf_a_headers = ['Concepto', 'Año 1 (€)', 'Año 2 (€)', 'Año 3 (€)']
    for i, header in enumerate(cf_a_headers):
        cf_annual_table.cell(0, i).text = header
        cf_annual_table.cell(0, i).paragraphs[0].runs[0].bold = True
    
    cf_annual_data = [
        ['Cash inicial', '58.000', '14.678', '26.678'],
        ['Entradas de efectivo', '10.725', '66.000', '150.000'],
        ['Salidas de efectivo', '(54.047)', '(54.000)', '(94.000)'],
        ['Cash flow operativo', '-43.322', '12.000', '56.000'],
        ['Cash final período', '14.678', '26.678', '82.678']
    ]
    
    for i, row in enumerate(cf_annual_data, 1):
        for j, value in enumerate(row):
            cf_annual_table.cell(i, j).text = value
    
    doc.add_page_break()
    
    # ========== 9.5 PUNTO DE EQUILIBRIO ==========
    doc.add_heading('9.5 Punto de Equilibrio y Análisis de Sensibilidad', 1)
    doc.add_paragraph()
    
    doc.add_paragraph('**Cálculo del punto de equilibrio operativo:**')
    
    equilibrio_text = """
**Fórmula:** Punto de Equilibrio = Costes Fijos / (Precio - Coste Variable Unitario)

**Datos:**
- Costes fijos mensuales: 4.500€ (desarrollo, marketing, personal, administrativo)
- Comisión media por transacción: 1,50€ (1% sobre transacción media de 150€)
- Coste variable unitario: 0,15€ (coste de procesamiento por transacción)

**Cálculo:**
Punto de Equilibrio = 4.500€ / (1,50€ - 0,15€) = 4.500€ / 1,35€ = 3