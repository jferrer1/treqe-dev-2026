#!/usr/bin/env python3
"""
Convertir Markdown a Word Document
Versión simple pero funcional
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
from datetime import datetime

def markdown_to_word(markdown_file, output_file):
    """Convertir archivo Markdown a documento Word."""
    
    print(f"Convirtiendo {markdown_file} a {output_file}...")
    
    # Leer archivo Markdown
    with open(markdown_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Crear documento Word
    doc = Document()
    
    # Configurar márgenes
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Procesar contenido línea por línea
    lines = content.split('\n')
    
    for line in lines:
        line = line.rstrip()
        
        if not line:
            doc.add_paragraph()
            continue
        
        # Detectar encabezados
        if line.startswith('# '):
            # Título principal
            p = doc.add_heading(line[2:], 0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('## '):
            # Título sección
            doc.add_heading(line[3:], 1)
        elif line.startswith('### '):
            # Subtítulo
            doc.add_heading(line[4:], 2)
        elif line.startswith('#### '):
            # Sub-subtítulo
            doc.add_heading(line[5:], 3)
        elif line.startswith('**') and line.endswith('**'):
            # Texto en negrita
            p = doc.add_paragraph()
            run = p.add_run(line[2:-2])
            run.bold = True
        elif line.startswith('- ') or line.startswith('* '):
            # Lista con viñetas
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(line[2:])
        elif line.startswith('|') and '|' in line[1:]:
            # Tabla Markdown - manejar simple
            if '---' in line or '===' in line:
                continue  # Saltar separadores de tabla
            # Para simplificar, convertimos a texto normal
            cells = [cell.strip() for cell in line.split('|') if cell.strip()]
            if cells:
                p = doc.add_paragraph()
                p.add_run(' | '.join(cells))
        elif line.startswith('```'):
            # Código - saltar por ahora
            continue
        else:
            # Texto normal
            doc.add_paragraph(line)
    
    # Guardar documento
    doc.save(output_file)
    print(f"Documento guardado: {output_file}")
    print(f"Tamaño: {len(content)} caracteres convertidos")

if __name__ == '__main__':
    # Convertir resumen ejecutivo
    markdown_to_word(
        'Resumen_Ejecutivo_Completo.md',
        'plan_negocio/Resumen_Ejecutivo_Treqe_Detallado.docx'
    )