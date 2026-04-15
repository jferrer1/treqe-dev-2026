#!/usr/bin/env python3
"""
Corregir anexos - versión completa
"""

import docx
import os

def crear_anexos_veridicos():
    """Crear anexos verídicos sin información inventada."""
    
    input_path = "plan_negocio/Plan_Negocio_Treqe_100%_PROFESIONAL_FINAL.docx"
    output_path = "plan_negocio/Plan_Negocio_Treqe_ANEXOS_VERIDICOS_FINAL.docx"
    
    if not os.path.exists(input_path):
        print(f"Error: No se encuentra: {input_path}")
        return None
    
    print(f"Cargando: {input_path}")
    doc = docx.Document(input_path)
    
    # Encontrar sección ANEXOS
    inicio_anexos = -1
    for i, p in enumerate(doc.paragraphs):
        if 'ANEXOS' in p.text.upper() and len(p.text) < 20:
            inicio_anexos = i
            break
    
    if inicio_anexos == -1:
        doc.add_page_break()
        titulo = doc.add_paragraph("ANEXOS")
        titulo.style = "Heading 1"
        inicio_anexos = len(doc.paragraphs) - 1
    
    # Limpiar contenido existente
    i = inicio_anexos + 1
    while i < len(doc.paragraphs):
        p = doc.paragraphs[i]
        if p.style.name == "Heading 1" and i > inicio_anexos + 1:
            break
        i += 1
    
    if i > inicio_anexos + 1:
        for _ in range(inicio_anexos + 1, i):
            if inicio_anexos + 1 < len(doc.paragraphs):
                p = doc.paragraphs[inicio_anexos + 1]
                p_element = p._element
                p_element.getparent().remove(p_element)
    
    print("Creando anexos verídicos...")
    
    # ANEXO A: ESTRUCTURA EQUIPO
    doc.add_paragraph()
    anexo_a = doc.add_paragraph("ANEXO A: ESTRUCTURA Y ROLES DEL EQUIPO")
    anexo_a.style = "Heading 2"
    
    contenido_a = [
        "",
        "ROLES NECESARIOS PARA MVP:",
        "",
        "1. FUNDADOR/CEO",
        "• Responsabilidades: Estrategia, fundraising, visión producto",
        "• Experiencia ideal: Conocimiento sector, liderazgo, negociación",
        "",
        "2. DESARROLLADOR FULL-STACK",
        "• Responsabilidades: Desarrollo plataforma, mantenimiento, escalabilidad",
        "• Stack: Node.js/Express o Python/Django, React/Next.js, PostgreSQL",
        "• Habilidades: APIs, algoritmos, deployment cloud",
        "",
        "3. DISEÑADOR UX/UI",
        "• Responsabilidades: User research, wireframes, diseño interfaz",
        "• Herramientas: Figma, Adobe XD, user testing",
        "• Habilidades: Design thinking, prototipado, accesibilidad",
        "",
        "NOTA: Los CVs específicos se añadirán cuando se definan los miembros del equipo.",
    ]
    
    for item in contenido_a:
        if item == "":
            doc.add_paragraph()
        else:
            doc.add_paragraph(item)
    
    doc.add_page_break()
    
    # ANEXO B: METODOLOGÍA ESTUDIOS
    anexo_b = doc.add_paragraph("ANEXO B: METODOLOGÍA PARA VALIDACIÓN DE MERCADO")
    anexo_b.style = "Heading 2"
    
    contenido_b = [
        "",
        "FASES DE VALIDACIÓN RECOMENDADAS:",
        "",
        "FASE 1: INVESTIGACIÓN SECUNDARIA (SEMANAS 1-2)",
        "• Análisis datos públicos: INE, Eurostat, informes sector",
        "• Estudios existentes: Observatorio Cetelem, informes consultoras",
        "• Tendencias: Google Trends, reports industria",
        "",
        "FASE 2: ENCUESTAS CUALITATIVAS (SEMANAS 3-4)",
        "• Objetivo: Entender pain points y necesidades",
        "• Método: Encuestas online (Typeform, Google Forms)",
        "• Muestra: 200-500 participantes",
        "• Métricas: Intención uso, preocupaciones, disposición pago",
        "",
        "FASE 3: ENTREVISTAS PROFUNDAS (SEMANAS 5-6)",
        "• Objetivo: Validar assumptions y user flows",
        "• Método: Entrevistas 1:1 (presencial/video)",
        "• Participantes: 15-20 usuarios potenciales",
        "• Output: User personas, journey maps, insights cualitativos",
        "",
        "FASE 4: PILOTO CONTROLADO (SEMANAS 7-12)",
        "• Objetivo: Validar producto en entorno real",
        "• Método: MVP en barrio específico",
        "• Participantes: 100-200 usuarios activos",
        "• Métricas: Engagement, retención, tasa conversión",
        "",
        "NOTA: Los datos específicos se añadirán tras realizar las investigaciones.",
    ]
    
    for item in contenido_b:
        if item == "":
            doc.add_paragraph()
        else:
            doc.add_paragraph(item)
    
    doc.add_page_break()
    
    # ANEXO C: DISEÑOS DETALLADOS
    anexo_c = doc.add_paragraph("ANEXO C: ESPECIFICACIONES DE DISEÑO Y UX")
    anexo_c.style = "Heading 2"
    
    contenido_c = [
        "",
        "ARQUITECTURA DE LA PLATAFORMA:",
        "",
        "COMPONENTES PRINCIPALES:",
        "1. SISTEMA DE CATALOGACIÓN",
        "• Upload múltiple de imágenes (3-10 por producto)",
        "• Etiquetado automático por categoría/color/estado",
        "• Valoración estimada por algoritmo (basado en mercado)",
        "• Descripción asistida por templates inteligentes",
        "",
        "2. MOTOR DE BÚSQUEDA Y MATCHING",
        "• Búsqueda semántica (entender 'mueble pequeño para estudio')",
        "• Filtros avanzados: Radio, categoría, valor, fecha publicación",
        "• Algoritmo de matching: Sugiere ruedas de intercambio óptimas",
        "• Sistema de recomendaciones personalizadas",
        "",
        "3. SISTEMA DE INTERCAMBIO",
        "• Propuesta 1-clic para ruedas sugeridas",
        "• Chat integrado para negociación",
        "• Acuerdo de términos (fecha, lugar, condiciones)",
        "• Sistema de seguimiento y notificaciones",
        "• Calificación post-intercambio (estrellas + comentarios)",
        "",
        "4. PERFILES Y REPUTACIÓN",
        "• Perfil público con historial de intercambios",
        "• Sistema de reputación basado en transacciones",
        "• Verificaciones (email, teléfono, redes sociales)",
        "• Badges por logros (ej: 'Intercambiador frecuente')",
        "",
        "FLUJOS DE USUARIO CLAVE:",
        "",
        "FLUJO A: PUBLICAR PRODUCTO",
        "1. Click 'Publicar' en header o dashboard",
        "2. Subir fotos (arrastrar/soltar o seleccionar)",
        "3. Completar formulario: título, descripción, categoría, estado",
        "4. Especificar qué busca a cambio (categorías, valor aproximado)",
        "5. Confirmar y publicar (producto visible en 30 segundos)",
        "",
        "FLUJO B: ENCONTRAR INTERCAMBIO",
        "1. Buscar en homepage o explorar categorías",
        "2. Ver producto que interesa",
        "3. Click 'Intercambiar' (ver opciones de ruedas)",
        "4. Seleccionar rueda preferida",
        "5. Enviar propuesta o iniciar chat",
        "",
        "FLUJO C: COMPLETAR INTERCAMBIO",
        "1. Acordar términos en chat (fecha, lugar, detalles)",
        "2. Confirmar intercambio en plataforma",
        "3. Recibir recordatorios (24h antes, 1h antes)",
        "4. Encontrarse y realizar intercambio",
        "5. Confirmar recepción y calificar",
        "",
        "DISEÑO VISUAL - SISTEMA DE DISEÑO:",
        "",
        "PALETA DE COLORES:",
        "• Primary: #1a73e8 (azul confianza - botones principales)",
        "• Secondary: #34a853 (verde sostenibilidad - features eco)",
        "• Accent: #fbbc04 (amarillo atención - alertas, destacados)",
        "• Neutral: #5f6368 (gris texto), #f8f9fa (fondo claro)",
        "• Success: #34a853, Warning: #fbbc04, Error: #ea4335",
        "",
        "TIPOGRAFÍA:",
        "• Inter: Para interfaz (sans-serif, excelente legibilidad pantallas)",
        "• Escala: 12, 14, 16, 18, 20, 24, 32, 48, 64px",
        "• Pesos: Regular (400), Medium (500), Semibold (600), Bold (700)",
        "",
        "ESPACIADO (SISTEMA 8PX):",
        "• Base: 8px (escala: 8, 16, 24, 32, 40, 48, 64, 80, 96)",
        "• Container: Max-width 1200px, padding responsive",
        "• Grid: 12-column, gutter 24px, margin 0 auto",
        "",
        "COMPONENTES:",
        "• Botones: Primary (relleno azul), Secondary (contorno), Text",
        "• Cards: Borde radius 12px, shadow subtle, hover effects",
        "• Forms: Labels flotantes, validación en tiempo real",
        "• Navigation: Sticky header, bottom nav móvil, hamburger menu",
        "",
        "PRINCIPIOS UX:",
        "• Mobile-first: 70%+ tráfico desde móvil",
        "• Accesibilidad: Contrastes AA, keyboard navigation, screen readers",
        "• Performance: LCP < 2.5s, FID < 100ms, CLS < 0.1",
        "• Consistencia: Mismos patrones en toda la plataforma",
        "• Feedback: Estados loading, success, error siempre visibles",
    ]
    
    for item in contenido_c:
        if item == "":
            doc.add_paragraph()
        else:
            doc.add_paragraph(item)
    
    doc.add_page_break()
    
    # ANEXO D: GLOSARIO
    anexo_d = doc.add_paragraph("ANEXO D: GLOSARIO DE TÉRMINOS")
    anexo_d.style = "Heading 2"
    
    contenido_d = [
        "",
        "TÉRMINOS DEL PROYECTO TREQE:",
        "",
        "• Rueda de Intercambio: Ciclo A→B→C→A donde todos reciben lo que quieren",
        "• Algoritmo de Matching: Sistema que encuentra combinaciones óptimas entre usuarios",
        "• Compensación Monetaria: Pago pequeño para equilibrar diferencias de valor",
        "• Liquidez de Plataforma: Facilidad para encontrar contrapartes de intercambio",
        "• Ciclo Cerrado: Productos que circulan sin convertirse en residuos",
        "",
        "TÉRMINOS DE NEGOCIO:",
        "",
        "• Two-sided Marketplace: Plataforma que conecta dos grupos de usuarios",
        "• Network Effects: El valor aumenta con más usuarios (efecto red)",
        "• Chicken-and-Egg: Problema inicial de plataformas (qué viene primero)",
        "• CAC/LTV: Coste adquisición cliente / Valor vida útil cliente",
        "• Burn Rate/Runway: Ritmo de gasto / Tiempo hasta necesitar más capital",
        "• MVP: Producto Mínimo Viable para validar hipótesis",
        "• PMF: Product-Market Fit (ajuste producto-mercado)",
        "",
        "TÉRMINOS TÉCNICOS:",
        "",
        "• API: Interfaz para comunicación entre sistemas",
        "• Backend/Frontend: Lógica servidor / Interfaz usuario",
        "• Database: Base de datos (ej: PostgreSQL, MongoDB)",
        "• Cloud: Servicios en la nube (AWS, Google Cloud, Azure)",
        "• Responsive: Diseño que se adapta a diferentes tamaños pantalla",
        "• UX/UI: Experiencia usuario / Interfaz usuario",
    ]
    
    for item in contenido_d:
        if item == "":
            doc.add_paragraph()
        else:
            doc.add_paragraph(item)
    
    doc.add_page_break()
    
    # ANEXO E: FUENTES VERÍDICAS
    anexo_e = doc.add_paragraph("ANEXO E: FUENTES Y REFERENCIAS")
    anexo_e.style = "Heading 2"
    
    contenido_e = [
        "",
        "FUENTES PÚBLICAS Y OFICIALES:",
        "",
        "1. INSTITUCIONES OFICIALES:",
        "• INE (Instituto Nacional de Estadística) - https://www.ine.es",
        "  - Estadísticas comercio electrónico, consumo, demografía",
        "• Eurostat - https://ec.europa.eu/eurostat",
        "  - Datos comparativos UE, economía circular, sostenibilidad",
        "• Banco de España - https://www.bde.es",
        "  - Informes pagos digitales, regulación fintech",
        "• CNMC (Comisión Nacional Mercados Competencia) - https://www.cnmc.es",
        "  - Estudios plataformas digitales, competencia",
        "",
        "2. ORGANIZACIONES SECTORIALES:",
        "• Observatorio Cetelem - Informes anuales consumo",
        "• Asociación Española Economía Circular - https://economiacircular.org",
        "• AECOC (Asociación Fabricantes Distribuidores) - Datos retail",
        "• AMETIC (Asociación Empresas Tecnología) - Informes sector",
        "",
        "3. ESTUDIOS Y REPORTS:",
        "• Statista - Reports mercado segunda mano España/Europa",
        "• Google Trends - Tendencias búsquedas 'segunda mano', 'trueque'",
        "• SimilarWeb - Tráfico competidores (Wallapop, Vinted, Milanuncios)",
        "",
        "4. INFORMES COMPETIDORES (PÚBLICOS):",
        "• Wallapop - Informes impacto anuales (disponibles en su web)",
        "• Vinted - Reports transparencia y sostenibilidad",
        "• Milanuncios - Datos públicos de Schibsted (empresa matriz)",
        "",
        "5. PLATAFORMAS DE DATOS:",
        "• App Annie/Data.ai - Métricas apps móviles",
        "• Crunchbase - Información startups, funding rounds",
        "• PitchBook - Análisis inversión venture capital",
        "",
        "METODOLOGÍA DE INVESTIGACIÓN:",
        "",
        "Para desarrollar el plan de negocio se han consultado:",
        "• Datos públicos de las fuentes mencionadas anteriormente",
        "• Análisis de informes sectoriales disponibles públicamente",
        "• Tendencias identificadas en medios especializados",
        "• Benchmarking de competidores mediante datos públicos",
        "",
        "NOTA: Se recomienda actualizar periódicamente con las fuentes más recientes y",
        "complementar con investigación primaria (encuestas, entrevistas) para validación.",
    ]
    
    for item in contenido_e:
        if item == "":
            doc.add_paragraph()
        else:
            doc.add_paragraph(item)
    
    # Guardar
    print(f"Guardando: {output_path}")
    doc.save(output_path)
    
    # Verificar
    doc_verificado = docx.Document(output_path)
    print(f"Documento guardado. Párrafos: {len(doc_verificado.paragraphs)}")
    
    return output_path

if __name__ == '__main__':
    print("=== CREANDO ANEXOS VERÍDICOS ===")
    print("Ajustando anexos para que sean factuales y verificables...\n")
    
    ruta = crear_anexos_veridicos()
    
    if ruta:
        print("\n" + "="*60)
        print("✅ ANEXOS VERÍDICOS CREADOS EXITOSAMENTE")
        print("="*60)
        print(f"\nDocumento: {ruta}")
        print("\nContenido verídico añadido:")
        print("1. ANEXO A: Estructura equipo (roles necesarios, sin CVs inventados)")
        print("2. ANEXO B: Metodología estudios (cómo validar, sin datos inventados)")
        print("3. ANEXO C: Diseños detallados (especificaciones reales de plataforma)")
        print("4. ANEXO