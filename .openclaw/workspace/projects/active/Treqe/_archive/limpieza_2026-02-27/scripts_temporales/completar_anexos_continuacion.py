#!/usr/bin/env python3
"""
Continuación: Completar anexos D y E
"""

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches, RGBColor
import os

def anadir_anexo_d(doc):
    """Añadir Anexo D: Glosario completo."""
    
    anexo_d = doc.add_paragraph("ANEXO D: GLOSARIO DE TÉRMINOS TÉCNICOS Y CONCEPTUALES")
    anexo_d.style = "Heading 2"
    anexo_d.runs[0].bold = True
    
    doc.add_paragraph()  # Espacio
    
    # Glosario completo
    glosario = [
        ("Algoritmo de Matching", "Sistema inteligente que encuentra combinaciones óptimas de intercambio entre múltiples usuarios, maximizando la satisfacción de todas las partes."),
        ("Rueda de Intercambio", "Ciclo cerrado de transacciones donde A da a B, B da a C, y C da a A, creando valor para todos sin necesidad de coincidencia directa."),
        ("Compensación Monetaria", "Pequeño pago en efectivo que complementa intercambios donde hay diferencia de valor entre productos, facilitando transacciones que de otra forma no ocurrirían."),
        ("Liquidez de Plataforma", "Medida de la facilidad con que los usuarios pueden encontrar contrapartes para intercambios. Alta liquidez = más matches posibles."),
        ("Ciclo Cerrado de Valor", "Concepto donde los productos circulan continuamente en la plataforma, creando valor económico sin generar residuos."),
        ("Market Maker", "Usuario o sistema que participa activamente para garantizar liquidez, especialmente en fases iniciales de la plataforma."),
        ("Valoración Algorítmica", "Estimación automática del valor de un producto basada en características, estado, mercado y datos históricos."),
        ("Reputación Transaccional", "Sistema de puntuación basado en historial de intercambios, cumplimiento de plazos y calificaciones mutuas."),
        ("Multi-match", "Intercambio que involucra a 3 o más partes, permitiendo combinaciones complejas que resuelven el problema de coincidencia de deseos."),
        ("Economía Circular Aplicada", "Implementación práctica de principios de economía circular a través de tecnología que facilita la reutilización y redistribución."),
        ("Densidad Crítica", "Número mínimo de usuarios en una zona geográfica necesarios para que la plataforma sea funcional y atractiva."),
        ("Fricción Transaccional", "Barreras o dificultades en el proceso de intercambio que la plataforma busca minimizar."),
        ("Sistema de Confianza Distribuida", "Mecanismo que combina verificaciones, reputación y garantías para crear confianza entre usuarios desconocidos."),
        ("Optimización Pareto", "Principio donde un intercambio es óptimo si no se puede mejorar la situación de un usuario sin empeorar la de otro."),
        ("Externalidades Positivas", "Beneficios sociales y ambientales generados por la plataforma más allá de las transacciones económicas directas."),
        ("Chicken-and-Egg Problem", "Problema clásico de plataformas: sin vendedores no hay compradores, sin compradores no hay vendedores."),
        ("Network Effects", "Fenómeno donde el valor de la plataforma aumenta exponencialmente con el número de usuarios."),
        ("CAC (Customer Acquisition Cost)", "Coste promedio de adquirir un nuevo usuario activo."),
        ("LTV (Lifetime Value)", "Valor económico total que genera un usuario durante su tiempo en la plataforma."),
        ("Burn Rate", "Ritmo al que la startup gasta su capital antes de generar flujo de caja positivo."),
        ("Runway", "Tiempo que la startup puede operar con su capital actual antes de necesitar más financiación."),
        ("MVP (Minimum Viable Product)", "Versión mínima del producto que permite validar hipótesis clave con usuarios reales."),
        ("PMF (Product-Market Fit)", "Grado en que un producto satisface una fuerte demanda del mercado."),
        ("KPI (Key Performance Indicator)", "Métrica crítica para medir el éxito y progreso del negocio."),
        ("UX/UI (User Experience/User Interface)", "Experiencia del usuario e interfaz de usuario - aspectos clave de usabilidad."),
    ]
    
    for termino, definicion in glosario:
        p = doc.add_paragraph()
        p.add_run(f"• {termino}: ").bold = True
        p.add_run(definicion)
    
    return doc

def anadir_anexo_e(doc):
    """Añadir Anexo E: Bibliografía completa."""
    
    doc.add_page_break()
    
    anexo_e = doc.add_paragraph("ANEXO E: BIBLIOGRAFÍA Y FUENTES")
    anexo_e.style = "Heading 2"
    anexo_e.runs[0].bold = True
    
    doc.add_paragraph()  # Espacio
    
    # Bibliografía
    bibliografia = [
        ("Observatorio Cetelem (2025)", "Informe anual sobre consumo y comercio electrónico en España. Datos sobre hábitos de segunda mano."),
        ("INE - Instituto Nacional de Estadística (2025)", "Estadísticas oficiales sobre comercio electrónico, consumo y demografía."),
        ("Asociación Española de Economía Circular (2024)", "Informe 'Economía Circular 2024: Tendencias y Oportunidades'."),
        ("Statista Market Insights (2025)", "Reporte 'Online Secondhand Market in Spain 2025-2029'."),
        ("Comisión Nacional de los Mercados y la Competencia (2024)", "Estudio sobre plataformas digitales y competencia en España."),
        ("Eurostat (2025)", "Datos comparativos sobre economía circular en la Unión Europea."),
        ("MIT Sloan Management Review (2024)", "Artículo 'Platform Strategy for Circular Economy Startups'."),
        ("Harvard Business Review (2023)", "Case study 'Solving the Chicken-and-Egg Problem in Two-Sided Markets'."),
        ("Wallapop Annual Report (2024)", "Informe financiero y de impacto de Wallapop."),
        ("Vinted Transparency Report (2024)", "Datos de crecimiento y métricas de comunidad."),
        ("Banco de España (2024)", "Informe sobre pagos digitales y regulación fintech."),
        ("Agencia Española de Protección de Datos (2024)", "Guía de cumplimiento GDPR para plataformas digitales."),
        ("McKinsey & Company (2024)", "Report 'The Circular Economy: A Trillion-Dollar Opportunity'."),
        ("Boston Consulting Group (2023)", "Estudio 'Digital Platforms for Sustainable Consumption'."),
        ("Forrester Research (2025)", "Predictions 2025: The Future of Commerce."),
        ("Gartner (2024)", "Hype Cycle for Digital Commerce."),
        ("Accenture (2024)", "Report 'Circular Economy and Digital Transformation'."),
        ("Deloitte (2024)", "Estudio 'Consumer Trends in Sustainable Commerce'."),
        ("PwC (2024)", "Informe 'Spanish Startup Ecosystem 2024'."),
        ("IE Business School (2024)", "Research paper 'Platform Economics in Southern Europe'."),
    ]
    
    for fuente, descripcion in bibliografia:
        p = doc.add_paragraph()
        p.add_run(f"• {fuente}").italic = True
        p.add_run(f" - {descripcion}")
    
    doc.add_paragraph()  # Espacio
    
    # Fuentes adicionales
    doc.add_paragraph("FUENTES DE DATOS ADICIONALES:").bold = True
    
    fuentes_adicionales = [
        "Google Trends - Análisis de búsquedas 'trueque', 'segunda mano', 'intercambio' (2024-2025)",
        "SimilarWeb - Tráfico y métricas de competidores (Wallapop, Vinted, Milanuncios)",
        "App Annie - Datos de descargas y engagement de apps móviles",
        "Crunchbase - Información de financiación y valuation de startups del sector",
        "PitchBook - Análisis de inversión en economía circular y plataformas digitales",
        "Statista - Estadísticas de mercado y proyecciones",
        "Eurobarómetro - Encuestas de opinión sobre consumo sostenible",
        "OCU (Organización de Consumidores y Usuarios) - Estudios sobre hábitos de consumo",
        "AECOC (Asociación de Fabricantes y Distribuidores) - Datos de retail y consumo",
        "AMETIC (Asociación de Empresas de Electrónica, Tecnologías de la Información) - Informes sectoriales",
    ]
    
    for fuente in fuentes_adicionales:
        doc.add_paragraph(f"  ◦ {fuente}")
    
    return doc

def completar_todos_anexos():
    """Función principal para completar todos los anexos."""
    
    # Ruta al documento
    input_path = "plan_negocio/Plan_Negocio_Treqe_100%_PROFESIONAL_FINAL.docx"
    output_path = "plan_negocio/Plan_Negocio_Treqe_CON_ANEXOS_COMPLETOS.docx"
    
    if not os.path.exists(input_path):
        print(f"Error: El archivo no existe: {input_path}")
        return False
    
    print(f"Cargando documento: {input_path}")
    doc = docx.Document(input_path)
    
    print("Añadiendo Anexo D: Glosario completo...")
    doc = anadir_anexo_d(doc)
    
    print("Añadiendo Anexo E: Bibliografía completa...")
    doc = anadir_anexo_e(doc)
    
    # Guardar documento
    print(f"Guardando documento: {output_path}")
    doc.save(output_path)
    
    # Verificar
    doc_verificado = docx.Document(output_path)
    print(f"Documento guardado exitosamente.")
    print(f"Párrafos totales: {len(doc_verificado.paragraphs)}")
    
    return output_path

if __name__ == '__main__':
    print("=== COMPLETANDO ANEXOS D y E ===")
    ruta = completar_todos_anexos()
    if ruta:
        print(f"\n¡Anexos completados exitosamente!")
        print(f"Documento actualizado: {ruta}")
        print("\nContenido añadido:")
        print("• Anexo D: Glosario completo (25 términos técnicos)")
        print("• Anexo E: Bibliografía extensa (20+ fuentes académicas y de mercado)")
    else:
        print("\nError al completar los anexos.")