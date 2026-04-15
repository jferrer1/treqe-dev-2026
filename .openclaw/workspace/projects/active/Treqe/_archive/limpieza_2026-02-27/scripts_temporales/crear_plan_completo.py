#!/usr/bin/env python3
"""
Crear Plan de Negocio COMPLETO y ACTUALIZADO para Treqe.
Incluye: datos 2026, arquitectura técnica, estrategia DIY.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import sys

def add_section_title(doc, title):
    """Añade título de sección."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0, 32, 96)
    doc.add_paragraph()

def create_complete_plan():
    doc = Document()
    
    # Portada
    doc.add_heading('PLAN DE NEGOCIO TREQE 2026', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    doc.add_paragraph('Versión Mejorada - ' + datetime.now().strftime('%d/%m/%Y')).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    doc.add_paragraph('• Datos mercado actualizados 2026').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('• Arquitectura técnica completa').alignment = WD_ALIGN_PARAGRAPH.CENTER  
    doc.add_paragraph('• Estrategia "Hazlo Tú Mismo"').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()
    
    # 1. RESUMEN EJECUTIVO ACTUALIZADO
    doc.add_heading('1. RESUMEN EJECUTIVO 2026', 1)
    doc.add_paragraph('**Contexto actualizado Febrero 2026:**')
    
    puntos = [
        "✅ Mercado segunda mano España: 8.200M€ (+42% vs 2020)",
        "✅ 28 millones usuarios activos (47% población)",
        "✅ Gasto medio anual: 1.850€ (+142% vs 2016)",
        "✅ Mobile-first: 94% usa apps vs 62% en 2021",
        "✅ Crecimiento segmento lujo: +125% (2023-2025)",
        "✅ Competencia: Wallapop (15M), Vinted (4.5M), Facebook Marketplace (20M potencial)"
    ]
    
    for punto in puntos:
        doc.add_paragraph(punto, style='List Bullet')
    
    doc.add_paragraph()
    doc.add_paragraph('**Propuesta Treqe 2026:** Plataforma de trueque inteligente con:')
    doc.add_paragraph('• Ruedas de intercambio (3+ usuarios) con matching algorítmico', style='List Bullet')
    doc.add_paragraph('• Compensación económica automática (diferencias de valor)', style='List Bullet')
    doc.add_paragraph('• Comisión 1% (vs 5-10% competencia)', style='List Bullet')
    doc.add_paragraph('• Enfoque mobile-first + PWA instalable', style='List Bullet')
    
    doc.add_page_break()
    
    # 2. ARQUITECTURA TÉCNICA - RUEDAS DE INTERCAMBIO
    doc.add_heading('2. ARQUITECTURA DEL SISTEMA COMPLEJO', 1)
    
    doc.add_paragraph('**2.1 Stack Tecnológico Propuesto:**')
    stack = doc.add_table(rows=5, cols=2)
    stack.cell(0, 0).text = 'Frontend'
    stack.cell(0, 1).text = 'Next.js 14 + React 19 + TypeScript + Tailwind + PWA'
    stack.cell(1, 0).text = 'Backend API'
    stack.cell(1, 1).text = 'Node.js 20 + Express + WebSockets (negociación real-time)'
    stack.cell(2, 0).text = 'Microservicio Matching'
    stack.cell(2, 1).text = 'Python + NetworkX (grafos) + PuLP (optimización)'
    stack.cell(3, 0).text = 'Base de datos'
    stack.cell(3, 1).text = 'PostgreSQL (transacciones) + Redis (caché)'
    stack.cell(4, 0).text = 'Infraestructura'
    stack.cell(4, 1).text = 'Vercel (frontend) + Railway/Render (backend) + Cloudinary (imágenes)'
    
    doc.add_paragraph()
    doc.add_paragraph('**2.2 Algoritmo de Ruedas de Intercambio:**')
    
    algoritmo = """
**Problema:** Dado N usuarios con {artículo_ofrecido, artículos_deseados, valoración}, encontrar ciclos de intercambio óptimos.

**Solución en 4 pasos:**

1. **Construcción grafo dirigido:**
   - Nodos = usuarios
   - Arco U→V si V tiene algo que U quiere
   - Peso = |valor(U.ofrece) - valor(V.ofrece)| × (1 - preferencia)

2. **Búsqueda de ciclos (3-5 nodos):**
   - DFS limitado en profundidad
   - Timeout: 500ms por usuario
   - Redis cache de grafos parciales

3. **Optimización económica:**
   ```
   Minimizar: Σ|compensación_i|
   Sujeto a: valor(recibido_i) ≥ valor(entregado_i) + compensación_i
            compensación_i ∈ [-max_diff, max_diff]
   ```
   - Resolver con PuLP (programación lineal)

4. **Presentación y negociación:**
   - Mostrar top 3 ciclos
   - Permitir ajustes manuales
   - Confirmación multi-firma
   """
    
    doc.add_paragraph(algoritmo)
    
    doc.add_paragraph('**2.3 Flujo de estados de una rueda:**')
    estados = doc.add_table(rows=7, cols=2)
    estados.cell(0, 0).text = 'Estado'
    estados.cell(0, 1).text = 'Descripción'
    estados.cell(1, 0).text = '1. OFERTA'
    estados.cell(1, 1).text = 'Usuario publica artículo + lo que quiere a cambio'
    estados.cell(2, 0).text = '2. MATCHING'
    estados.cell(2, 1).text = 'Sistema busca ciclos cada 5 min (caché Redis)'
    estados.cell(3, 0).text = '3. NEGOCIACIÓN'
    estados.cell(3, 1).text = 'WebSocket chat grupal + ajuste compensaciones'
    estados.cell(4, 0).text = '4. ACUERDO'
    estados.cell(4, 1).text = 'Todos confirman → Stripe retiene compensaciones'
    estados.cell(5, 0).text = '5. ENVÍO'
    estados.cell(5, 1).text = 'Cada usuario envía su artículo (tracking Correos/SEUR)'
    estados.cell(6, 0).text = '6. COMPLETADO'
    estados.cell(6, 1).text = 'Todos reciben → liberación pagos + reputación'
    
    doc.add_page_break()
    
    # 3. ESTRATEGIA "HAZLO TÚ MISMO"
    doc.add_heading('3. ESTRATEGIA "HAZLO TÚ MISMO" (SIN CONTRATAR)', 1)
    
    doc.add_paragraph('**3.1 Lo que podemos hacer con OpenClaw/AI:**')
    
    podemos = [
        ("Desarrollo completo", "Next.js + Node.js + Python scripts", "100% automatizable"),
        ("Diseño UI", "Tailwind templates + AI (Claude/ChatGPT)", "Zero diseñador"),
        ("Base de datos", "PostgreSQL schema automático", "Scripts de migración"),
        ("DevOps", "GitHub Actions + Vercel CLI", "CI/CD automático"),
        ("Testing", "Playwright E2E + Jest unit tests", "Tests generados por AI"),
        ("Documentación", "Swagger auto-generado + TypeDoc", "Docs siempre actualizados"),
        ("Marketing inicial", "Landing page + SEO básico", "Contenido AI + plantillas")
    ]
    
    podemos_table = doc.add_table(rows=len(podemos)+1, cols=3)
    podemos_table.cell(0, 0).text = 'Área'
    podemos_table.cell(0, 1).text = 'Cómo'
    podemos_table.cell(0, 2).text = 'Automatización'
    
    for i, (area, como, auto) in enumerate(podemos, 1):
        podemos_table.cell(i, 0).text = area
        podemos_table.cell(i, 1).text = como
        podemos_table.cell(i, 2).text = auto
    
    doc.add_paragraph()
    doc.add_paragraph('**3.2 Timeline desarrollo (1 persona, 3 meses):**')
    
    timeline = [
        ("Semana 1-2", "Landing page + waitlist", "Next.js + Vercel + Stripe"),
        ("Semana 3-4", "Auth + perfil usuario", "NextAuth + PostgreSQL"),
        ("Semana 5-6", "Subida artículos + imágenes", "Cloudinary + forms"),
        ("Semana 7-8", "Algoritmo matching básico", "Python microservice"),
        ("Semana 9-10", "Negociación WebSocket", "Socket.io + Redis"),
        ("Semana 11-12", "Pagos Stripe Connect", "Escrow + webhooks"),
        ("Semana 13-14", "Logística APIs", "Correos/SEUR integration"),
        ("Semana 15-16", "Testing + deploy", "Playwright + Vercel prod")
    ]
    
    for tiempo, tarea, tech in timeline:
        doc.add_paragraph(f'**{tiempo}:** {tarea} ({tech})', style='List Bullet')
    
    doc.add_page_break()
    
    # 4. ELEMENTOS CLAVE DEL BUSINESS PLAN ANTIGUO
    doc.add_heading('4. ELEMENTOS IMPORTANTES DEL PLAN ANTIGUO (INCORPORADOS)', 1)
    
    elementos = [
        ("**Modelo de Peter Drucker**", "Innovación como 'nueva dimensión de desempeño'", "Mantener filosofía de cambio de reglas"),
        ("**Cadena de valor centrada en usuario**", "Seguridad, información, logística, garantía", "Incorporar como principios de diseño"),
        ("**7Ps del marketing de servicios**", "Personas, procesos, evidencia física", "Aplicar a experiencia de trueque"),
        ("**Segmentación por motivaciones**", "Económica, espacio, actualización, ecológica", "Mantener en targeting"),
        ("**Ejemplo Juan y Helena**", "Casos de uso concretos y comprensibles", "Usar en onboarding y marketing"),
        ("**Estrategia de precios psicológicos**", "0,99% vs 1% (simbólico)", "Considerar para comisiones"),
        ("**Enfoque omnichannel**", "Web + móvil + posible físico", "PWA como solución móvil")
    ]
    
    for elemento, desc, aplicacion in elementos:
        doc.add_paragraph(f'• **{elemento}:** {desc}', style='List Bullet')
        doc.add_paragraph(f'  → Aplicación: {aplicacion}', style='List Bullet 2')
    
    doc.add_paragraph()
    doc.add_paragraph('**4.1 Mejoras respecto al plan antiguo:**')
    mejoras = [
        "✅ Datos actualizados 2026 vs 2016/2021",
        "✅ Arquitectura serverless vs hosting tradicional",
        "✅ Mobile-first (PWA) vs web responsive",
        "✅ AI matching vs búsqueda manual",
        "✅ Stripe Connect vs transferencias bancarias manuales",
        "✅ APIs logística vs acuerdos bilaterales"
    ]
    
    for mejora in mejoras:
        doc.add_paragraph(mejora, style='List Bullet')
    
    # 5. PRESUPUESTO DIY
    doc.add_page_break()
    doc.add_heading('5. PRESUPUESTO "HAZLO TÚ MISMO" (3 MESES)', 1)
    
    presupuesto = [
        ("Herramientas desarrollo", "Vercel Pro, Railway, Cloudinary", "50€/mes", "150€"),
        ("APIs externas", "Stripe, SendGrid, Twilio, Correos API", "100€/mes", "300€"),
        ("Dominio + SSL", "treqe.es + certificado", "20€/año", "20€"),
        ("Diseño/assets", "Canva Pro, stock images", "30€/mes", "90€"),
        ("Legal básico", "Plantillas + consulta puntual", "Una vez", "500€"),
        ("Marketing inicial", "Google Ads crédito trial", "100€", "100€"),
        ("**TOTAL ESTIMADO**", "", "", "**1.160€**")
    ]
    
    pres_table = doc.add_table(rows=len(presupuesto)+1, cols=4)
    pres_table.cell(0, 0).text = 'Concepto'
    pres_table.cell(0, 1).text = 'Detalle'
    pres_table.cell(0, 2).text = 'Periodicidad'
    pres_table.cell(0, 3).text = 'Coste'
    
    for i, (concepto, detalle, periodo, coste) in enumerate(presupuesto, 1):
        pres_table.cell(i, 0).text = concepto
        pres_table.cell(i, 1).text = detalle
        pres_table.cell(i, 2).text = periodo
        pres_table.cell(i, 3).text = coste
    
    doc.add_paragraph()
    doc.add_paragraph('**Ahorro vs contratar equipo:**')
    doc.add_paragraph('• Desarrollador senior: 4.000€/mes × 3 = 12.000€', style='List Bullet')
    doc.add_paragraph('• Diseñador UX/UI: 3.000€/mes × 3 = 9.000€', style='List Bullet')
    doc.add_paragraph('• **Total tradicional: ~21.000€**', style='List Bullet')
    doc.add_paragraph('• **Nuestro enfoque: 1.160€ (95% ahorro)**', style='List Bullet')
    
    # 6. PRÓXIMOS PASOS INMEDIATOS
    doc.add_page_break()
    doc.add_heading('6. PRÓXIMOS PASOS INMEDIATOS (SEMANA 1)', 1)
    
    pasos = [
        ("Día 1", "Registrar dominio treqe.es + configurar email"),
        ("Día 2", "Landing page con Waitlist (Next.js + Vercel)"),
        ("Día 3", "Stripe Connect setup + términos legales básicos"),
        ("Día 4", "Script Python algoritmo matching básico (proof-of-concept)"),
        ("Día 5", "Base de datos PostgreSQL schema (usuarios, artículos)"),
        ("Día 6", "Auth con NextAuth (Google + email + redes)"),
        ("Día 7", "Deploy beta cerrado + primeros 50 testers")
    ]
    
    for dia, accion in pasos:
        doc.add_paragraph(f'**{dia}:** {accion}', style='List Bullet')
    
    doc.add_paragraph()
    doc.add_paragraph('**Métrica éxito semana 1:** 100+ suscriptores waitlist')
    
    # Guardar
    output_path = Path(__file__).parent / 'Plan_Negocio_Treqe_Completo_2026.docx'
    doc.save(str(output_path))
    
    print(f"Documento creado: {output_path}")
    print("Contenido:")
    print("- Datos mercado 2026 actualizados")
    print("- Arquitectura técnica ruedas intercambio")
    print("- Algoritmo matching detallado")
    print("- Estrategia DIY (sin contratar)")
    print("- Presupuesto 1.160€ (3 meses)")
    print("- Timeline desarrollo semana a semana")
    
    return str(output_path)

if __name__ == '__main__':
    from pathlib import Path
    create_complete_plan()