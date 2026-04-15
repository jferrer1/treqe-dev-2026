#!/usr/bin/env python3
"""
Continuación de la integración de secciones
"""

from docx import Document
import os

def continuar_integracion():
    """Continuar integración desde donde quedó."""
    
    print("Continuando integración de secciones...")
    
    # Cargar documento existente
    doc_path = os.path.join(os.path.dirname(__file__), 'plan_negocio', 'Plan_Negocio_Treqe_Completo.docx')
    doc = Document(doc_path)
    
    # Continuar desde sección 5.1
    text_5_1_cont = """**Justificación:** Competitividad extrema vs competencia, transparencia total, alineación perfecta de incentivos (Treqe gana cuando usuarios completan intercambios satisfactorios)

**Fase 2: Crecimiento y Profundización (Año 2) - Servicios Premium**
Una vez establecida la base de usuarios, se introducen servicios de valor añadido:

1. **Cuenta Verificada (4,99€/mes):**
   - Verificación de identidad: Documento oficial + selfie
   - Badge de confianza: Indicador visual en perfil y listings
   - Prioridad en matching: Los algoritmos priorizan usuarios verificados
   - Soporte prioritario: Atención al cliente dedicada
   - Objetivo penetración: 20% de usuarios activos

2. **Destacados en Búsquedas (2,99€/artículo):**
   - Mayor visibilidad: Posicionamiento superior en resultados
   - Duración: 30 días de destacado
   - Estadísticas: Acceso a datos de visualizaciones y contactos
   - Segmentación: Por categoría, ubicación, valor

3. **Logística Premium (+3€/envío):**
   - Recogida a domicilio: Sin necesidad de ir a punto de entrega
   - Seguro ampliado: Cobertura hasta 1.000€
   - Tracking avanzado: Actualizaciones en tiempo real
   - Embalaje profesional: Incluye materiales de protección

4. **Servicios de Valoración (1,99€/valoración):**
   - Valoración profesional: Estimación por expertos por categoría
   - Certificado de autenticidad: Para artículos de marca
   - Informe de estado: Evaluación detallada de conservación
   - Recomendaciones: Precio óptimo para intercambio

**Fase 3: Madurez y Diversificación (Año 3) - Publicidad Segmentada y Partnerships**

1. **Publicidad para Marcas Sostenibles:**
   - Criterios de aceptación: Certificación B Corp, prácticas circulares verificadas
   - Formatos: Banners en categorías relevantes, contenido patrocinado
   - Segmentación: Por intereses de usuario, ubicación, historial
   - Modelo: CPM (coste por mil impresiones) + CPC (coste por clic)

2. **Promociones de Categorías Específicas:**
   - Ejemplos: "Moda sostenible de primavera", "Tecnología reacondicionada"
   - Colaboraciones: Con influencers de sostenibilidad y consumo consciente
   - Eventos virtuales: Ferias de intercambio temáticas
   - Guías de contenido: "Cómo intercambiar muebles de diseño"

3. **Partnerships Estratégicas:**
   - Empresas de reparación: Descuentos para usuarios Treqe
   - Plataformas de upcycling: Intercambio de usuarios y contenido
   - Instituciones educativas: Programas de economía circular
   - Administraciones públicas: Proyectos de sostenibilidad local"""
    
    # Buscar el último párrafo y continuar
    # Por simplicidad, agregaré un nuevo párrafo
    doc.add_paragraph(text_5_1_cont)
    
    # 5.2 Inversión Inicial
    doc.add_heading('5.2 Inversión Inicial', 1)
    
    text_5_2 = """**Inversión inicial total: 58.000€**

| Concepto | Importe | % Total | Detalle |
|----------|---------|---------|---------|
| **Desarrollo tecnológico** | 23.200€ | 40% | Frontend, backend, algoritmos, infraestructura |
| **Marketing inicial** | 20.300€ | 35% | Campañas digitales, contenido, eventos, PR |
| **Operaciones y equipo** | 14.500€ | 25% | Equipo fundador, legal, oficina, herramientas |
| **TOTAL** | **58.000€** | **100%** | |

**Desglose detallado de la inversión:**

**Desarrollo tecnológico (23.200€):**
- Frontend (Next.js + React + TypeScript): 8.000€
- Backend (Node.js + APIs + WebSockets): 6.500€
- Algoritmo matching (Python + optimización): 4.200€
- Infraestructura (Vercel + Railway + Cloudinary): 2.500€
- Testing y calidad: 2.000€

**Marketing inicial (20.300€):**
- Campañas digitales (Google Ads, Social Media): 12.000€
- Contenido y PR: 4.500€
- Eventos y comunidades: 2.800€
- Influencers y partnerships: 1.000€

**Operaciones y equipo (14.500€):**
- Equipo fundador (3 personas x 3 meses): 9.000€
- Legal y constitución: 3.000€
- Oficina y herramientas: 1.500€
- Contingencia: 1.000€"""
    
    doc.add_paragraph(text_5_2)
    
    # 5.3 Financiación Propuesta
    doc.add_heading('5.3 Financiación Propuesta', 1)
    
    text_5_3 = """**Estructura de financiación propuesta:**

| Fuente | Importe | % Total | Condiciones |
|--------|---------|---------|-------------|
| **Inversores ángeles / business angels** | 40.000€ | 69% | Equity 15-20% |
| **Préstamo participativo ENISA** | 10.000€ | 17% | 5 años, interés reducido |
| **Aportación equipo fundador** | 8.000€ | 14% | Capital propio |
| **TOTAL** | **58.000€** | **100%** | |

**Valoración pre-money:** 200.000€
**Equity ofrecido a inversores:** 15-20%
**ROI esperado para inversores:** 3-5x en 3-5 años

**Justificación de la valoración:**
- **Mercado potencial:** 8 millones de usuarios en España
- **Modelo escalable:** Arquitectura serverless, costes marginales bajos
- **Primer mover ventaja:** Segmento no atendido por competencia
- **Equipo fundador:** Experiencia combinada en tecnología, marketing y emprendimiento
- **Trayectoria crecimiento:** Proyecciones realistas con alto potencial

**Uso de fondos detallado:**
1. **Meses 1-3 (15.000€):** Desarrollo MVP + marketing inicial
2. **Meses 4-6 (18.000€):** Optimización + expansión geográfica
3. **Meses 7-9 (15.000€):** Servicios premium + equipo
4. **Meses 10-12 (10.000€):** Crecimiento acelerado + reserva

**Métricas clave para desembolsos:**
- **Desembolso 2:** Alcanzar 5.000 usuarios activos
- **Desembolso 3:** Alcanzar 15.000 usuarios activos
- **Desembolso 4:** Alcanzar 25.000 usuarios activos"""
    
    doc.add_paragraph(text_5_3)
    
    doc.add_page_break()
    
    # ========== SECCIÓN 6: PROYECCIONES FINANCIERAS ==========
    
    doc.add_heading('6. PROYECCIONES FINANCIERAS 2026-2029', 0)
    
    # 6.1 Supuestos Clave
    doc.add_heading('6.1 Supuestos Clave', 1)
    
    text_6_1 = """Las proyecciones financieras se basan en supuestos conservadores y realistas, validados mediante análisis de mercado y benchmarking de plataformas similares:

**Supuestos de crecimiento:**
- **Comisión efectiva:** 1% sobre el valor del artículo adquirido
- **Valor medio por transacción:** 150€ (año 1), 160€ (año 2), 170€ (año 3)
- **Tasa de conversión usuarios activos → transacciones:** 10% mensual
- **Crecimiento mensual de usuarios:** 15% (año 1), 10% (año 2), 5% (año 3)
- **Retención mensual de usuarios:** 70% (líder en sector)

**Supuestos de costes:**
- **Coste de adquisición de cliente (CAC):** 2,50€ (año 1), 3,00€ (año 2), 3,50€ (año 3)
- **Valor de vida del cliente (LTV):** 45€ (año 1), 78€ (año 2), 91€ (año 3)
- **Ratio LTV:CAC objetivo:** >3:1 (excelente: >3:1, bueno: >1:1)

**Supuestos operativos:**
- **Coste variable por transacción:** 0,15€ (procesamiento pagos, soporte)
- **Costes fijos mensuales:** 4.500€ (equipo, infraestructura, oficina)
- **Incremento costes:** 15% anual (inflación + crecimiento)"""
    
    doc.add_paragraph(text_6_1)
    
    # 6.2 Proyecciones Anuales
    doc.add_heading('6.2 Proyecciones Anuales', 1)
    
    text_6_2 = """**Proyecciones de crecimiento 2026-2029:**

| Métrica | Año 1 (2026) | Año 2 (2027) | Año 3 (2028) |
|---------|--------------|--------------|--------------|
| **Usuarios activos finales** | 25.000 | 75.000 | 150.000 |
| **Transacciones anuales** | 15.000 | 60.000 | 120.000 |
| **Volumen transaccional** | 2.250.000€ | 9.000.000€ | 18.000.000€ |
| **Ingresos comisiones (1%)** | **22.500€** | **90.000€** | **180.000€** |
| **Ingresos servicios premium** | 0€ | 18.000€ | 48.000€ |
| **Ingresos publicidad** | 0€ | 6.000€ | 18.000€ |
| **TOTAL INGRESOS** | **22.500€** | **114.000€** | **246.000€** |

**Análisis de las proyecciones:**

**Año 1 (2026) - Validación y crecimiento inicial:**
- Enfoque en Barcelona y Madrid como mercados piloto
- Objetivo: Demostrar viabilidad del modelo y validar algoritmo
- Ingresos principalmente de comisiones básicas
- Inversión significativa en adquisición de usuarios

**Año 2 (2027) - Escalabilidad y diversificación:**
- Expansión a 10 ciudades principales españolas
- Introducción de servicios premium (20% penetración objetivo)
- Primeros ingresos por publicidad segmentada
- Optimización de costes y mejora de márgenes

**Año 3 (2028) - Consolidación y rentabilidad:**
- Cobertura nacional completa
- Penetración premium aumentada a 30%
- Publicidad como fuente de ingresos significativa
- Rentabilidad EBITDA positiva sostenida"""
    
    doc.add_paragraph(text_6_2)
    
    # 6.3 Estado de Pérdidas y Ganancias
    doc.add_heading('6.3 Estado de Pérdidas y Ganancias', 1)
    
    text_6_3 = """**Estado de Pérdidas y Ganancias Proyectado:**

| Concepto | Año 1 (2026) | Año 2 (2027) | Año 3 (2028) |
|----------|--------------|--------------|--------------|
| **Ingresos totales** | 22.500€ | 114.000€ | 246.000€ |
| **Costes desarrollo** | (23.200€) | (25.000€) | (30.000€) |
| **Costes marketing** | (20.300€) | (35.000€) | (45.000€) |
| **Costes personal** | (9.000€) | (18.000€) | (30.000€) |
| **Costes operativos** | (5.500€) | (8.000€) | (12.000€) |
| **EBITDA** | **(35.500€)** | **28.000€** | **129.000€** |
| **Amortizaciones** | (1.000€) | (1.500€) | (2.000€) |
| **Intereses financieros** | (500€) | (1.000€) | (1.500€) |
| **Resultado antes impuestos** | **(37.000€)** | **25.500€** | **125.500€** |
| **Impuestos (25%)** | 0€ | 0€ | (31.375€) |
| **RESULTADO NETO** | **(37.000€)** | **25.500€** | **94.125€** |

**Análisis de márgenes:**
- **Margen EBITDA año 3:** 52,4% (129.000€ / 246.000€)
- **Margen neto año 3:** 38,3% (94.125€ / 246.000€)
- **Evolución márgenes:** Mejora consistente año tras año
- **Comparativa sector:** Márgenes superiores a plataformas establecidas

**Punto de equilibrio operativo:**
- **Cálculo:** 4.500€ costes fijos / (1,50€ comisión - 0,15€ coste) = **3.333 transacciones/mes**
- **Equivalente en usuarios:** 33.330 usuarios activos (10% tasa conversión)
- **Fecha objetivo:** Mes 10 del año 1"""
    
    doc.add_paragraph(text_6_3)
    
    # 6.4 Cash Flow Proyectado
    doc.add_heading('6.4 Cash Flow Proyectado', 1)
    
    text_6_4 = """**Cash Flow Operativo Proyectado:**

| Período | Cash Flow Operativo | Acumulado |
|---------|---------------------|-----------|
| **Año 1** | -28.000€ | -28.000€ |
| **Año 2** | +12.000€ | -16.000€ |
| **Año 3** | +58.000€ | +42.000€ |

**Análisis de tesorería:**
- **Tesorería inicial:** 58.000€ (inversión)
- **Tesorería fin año 1:** 14.678€
- **Tesorería fin año 2:** 26.678€
- **Tesorería fin año 3:** 84.678€

**Runway (autonomía financiera):**
- **Después año 1:** 14 meses de operación con tesorería disponible
- **Después año 2:** 26 meses de operación
- **Después año 3:** 84 meses de operación

**Necesidades de financiación adicional:**
- **Año 1:** Ninguna (cubierto por inversión inicial)
- **Año 2:** Posible ronda Serie A de 100.000€ para expansión internacional
- **Año 3:** Autofinanciación con cash flow positivo

**Ratios de liquidez:**
- **Current ratio año 3:** 5,8 (excelente, >2 es bueno)
- **Quick ratio año 3:** 4,2 (liquidez inmediata alta)
- **Cash conversion cycle:** Negativo (modelo favorable)"""
    
    doc.add_paragraph(text_6_4)
    
    # 6.5 Ratios Financieros Clave
    doc.add_heading('6.5 Ratios Financieros Clave', 1)
    
    text_6_5 = """**Ratios financieros clave (Año 3 - 2028):**

| Ratio | Valor | Interpretación | Benchmark sector |
|-------|-------|----------------|------------------|
| **Margen EBITDA** | 52,4% | Excelente rentabilidad operativa | 30-40% (bueno) |
| **Margen neto** | 38,3% | Rentabilidad final muy saludable | 20-30% (bueno) |
| **ROI inversión** | 162% | Retorno excelente sobre inversión | >100% (excelente) |
| **LTV:CAC** | 24: