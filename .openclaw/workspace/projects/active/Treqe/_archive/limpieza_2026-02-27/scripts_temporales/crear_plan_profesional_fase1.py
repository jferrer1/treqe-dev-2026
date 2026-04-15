#!/usr/bin/env python3
"""
FASE 1: Crear Plan de Negocio Profesional para Treqe
Secciones 1-5: Resumen ejecutivo, concepto, mercado, competencia, modelo
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime
from pathlib import Path

def create_professional_plan_phase1():
    """Crear primera fase del plan profesional (secciones 1-5)."""
    
    doc = Document()
    
    # ========== CONFIGURACIÓN DE ESTILOS ==========
    styles = doc.styles
    
    # Estilo para títulos principales
    title_style = styles.add_style('TituloPrincipal', WD_STYLE_TYPE.PARAGRAPH)
    title_font = title_style.font
    title_font.name = 'Calibri'
    title_font.size = Pt(20)
    title_font.bold = True
    title_font.color.rgb = RGBColor(0, 32, 96)  # Azul corporativo
    
    # Estilo para subtítulos
    subtitle_style = styles.add_style('Subtitulo', WD_STYLE_TYPE.PARAGRAPH)
    subtitle_font = subtitle_style.font
    subtitle_font.name = 'Calibri'
    subtitle_font.size = Pt(14)
    subtitle_font.bold = True
    subtitle_font.color.rgb = RGBColor(0, 32, 96)
    
    # Estilo para párrafos normales
    normal_style = styles.add_style('NormalProfesional', WD_STYLE_TYPE.PARAGRAPH)
    normal_font = normal_style.font
    normal_font.name = 'Calibri'
    normal_font.size = Pt(11)
    normal_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    normal_style.paragraph_format.space_after = Pt(6)
    
    # ========== PORTADA PROFESIONAL ==========
    doc.add_paragraph('PLAN DE NEGOCIO', style='Title').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    doc.add_paragraph('TREQE', style='Title').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Plataforma de Trueque Inteligente', style='Subtitle').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Información del documento
    fecha = datetime.now().strftime('%d de %B de %Y')
    info_table = doc.add_table(rows=4, cols=2)
    info_table.style = 'Light Shading Accent 1'
    
    info_table.cell(0, 0).text = 'Documento:'
    info_table.cell(0, 1).text = 'Plan de Negocio Profesional v3.0'
    
    info_table.cell(1, 0).text = 'Fecha:'
    info_table.cell(1, 1).text = fecha
    
    info_table.cell(2, 0).text = 'Versión:'
    info_table.cell(2, 1).text = 'Fase 1 - Secciones 1-5'
    
    info_table.cell(3, 0).text = 'Estado:'
    info_table.cell(3, 1).text = 'CONFIDENCIAL - Propiedad de Treqe'
    
    doc.add_page_break()
    
    # ========== ÍNDICE ==========
    doc.add_paragraph('ÍNDICE', style='TituloPrincipal').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    indice = [
        ("1. RESUMEN EJECUTIVO", 3),
        ("2. CONCEPTO DE NEGOCIO", 5),
        ("3. ANÁLISIS DE MERCADO", 8),
        ("4. ANÁLISIS DE LA COMPETENCIA", 12),
        ("5. MODELO DE NEGOCIO", 15),
        ("6. ESTRATEGIA DE MARKETING", 18),
        ("7. OPERACIONES Y TECNOLOGÍA", 21),
        ("8. EQUIPO Y ORGANIZACIÓN", 24),
        ("9. PLAN FINANCIERO", 27),
        ("10. ANÁLISIS DE RIESGOS", 31),
        ("11. HOJA DE RUTA", 34),
        ("12. CONCLUSIONES Y PRÓXIMOS PASOS", 37),
        ("ANEXOS", 39)
    ]
    
    for seccion, pagina in indice:
        p = doc.add_paragraph()
        p.add_run(seccion).bold = True
        p.add_run(f'\t{pagina}')
    
    doc.add_page_break()
    
    # ========== 1. RESUMEN EJECUTIVO ==========
    doc.add_paragraph('1. RESUMEN EJECUTIVO', style='TituloPrincipal')
    doc.add_paragraph()
    
    # 1.1 La Oportunidad
    doc.add_paragraph('1.1 La Oportunidad de Mercado', style='Subtitulo')
    doc.add_paragraph(
        'El mercado de segunda mano en España ha experimentado una transformación radical en los últimos cinco años. '
        'De ser un sector marginal asociado a crisis económicas, ha pasado a convertirse en una opción de consumo '
        'consciente, sostenible y económicamente inteligente para millones de españoles. '
        'En 2026, el mercado supera los 8.200 millones de euros, con 28 millones de usuarios activos que realizan '
        'una media de 6,6 transacciones anuales por persona.',
        style='NormalProfesional'
    )
    
    doc.add_paragraph(
        'Sin embargo, este crecimiento exponencial ha dejado al descubierto una limitación fundamental en las '
        'plataformas existentes: la falta de flexibilidad en los modelos de intercambio. Los usuarios que desean '
        'renovar sus posesiones pero carecen de liquidez monetaria se encuentran atrapados en un dilema: '
        'mantener artículos que ya no necesitan o venderlos por debajo de su valor real.',
        style='NormalProfesional'
    )
    
    # 1.2 La Solución Treqe
    doc.add_paragraph('1.2 La Solución Treqe', style='Subtitulo')
    doc.add_paragraph(
        'Treqe es una plataforma digital innovadora que introduce el concepto de "trueque inteligente" en el '
        'mercado de segunda mano español. A diferencia del trueque tradicional 1:1, Treqe implementa un sistema '
        'de "ruedas de intercambio" que permite a tres o más usuarios participar en cadenas de valor circulares, '
        'donde cada persona recibe exactamente lo que desea a cambio de lo que ofrece.',
        style='NormalProfesional'
    )
    
    doc.add_paragraph(
        'La innovación clave reside en la integración de compensaciones económicas dentro de estas ruedas. '
        'Cuando existe una diferencia de valor entre los artículos intercambiados, el sistema calcula '
        'automáticamente la compensación monetaria necesaria para equilibrar la transacción, creando así un '
        'modelo híbrido que combina las ventajas del trueque (conservación de valor) con la flexibilidad del '
        'mercado monetario (ajuste de diferencias).',
        style='NormalProfesional'
    )
    
    # 1.3 Ventaja Competitiva
    doc.add_paragraph('1.3 Ventaja Competitiva Sostenible', style='Subtitulo')
    doc.add_paragraph(
        'Treqe se posiciona como primer mover en el segmento de trueque estructurado en España, un nicho que '
        'permanece inexplorado por los principales actores del mercado. Mientras plataformas como Wallapop, '
        'Vinted y Milanuncios se centran exclusivamente en la compraventa monetaria, Treqe introduce una '
        'nueva dimensión de valor que resuelve problemas específicos de los usuarios:',
        style='NormalProfesional'
    )
    
    ventajas = [
        "**Liquidez limitada:** Permite renovar posesiones sin necesidad de desembolso monetario completo",
        "**Valoración subjetiva:** Facilita intercambios donde el valor emocional o de uso supera el valor de mercado",
        "**Sostenibilidad:** Promueve la economía circular maximizando la vida útil de los productos",
        "**Comunidad:** Crea redes de intercambio que generan engagement y fidelización orgánica"
    ]
    
    for ventaja in ventajas:
        doc.add_paragraph(f'• {ventaja}', style='List Bullet')
    
    # 1.4 Modelo de Ingresos
    doc.add_paragraph('1.4 Modelo de Ingresos', style='Subtitulo')
    doc.add_paragraph(
        'Treqe opera con un modelo de comisiones transparente y competitivo. Por cada transacción completada '
        'en la plataforma, se aplica una comisión del 1% sobre el valor del artículo adquirido, pagada por '
        'el usuario que recibe el bien. Esta estructura presenta varias ventajas estratégicas:',
        style='NormalProfesional'
    )
    
    puntos_comision = [
        "**Competitividad:** Significativamente inferior a las comisiones del 5-10% de la competencia",
        "**Transparencia:** El usuario paga solo cuando recibe valor tangible",
        "**Escalabilidad:** El modelo crece proporcionalmente al volumen de transacciones",
        "**Sostenibilidad:** Genera ingresos recurrentes sin afectar la liquidez de los usuarios"
    ]
    
    for punto in puntos_comision:
        doc.add_paragraph(f'• {punto}', style='List Bullet')
    
    # 1.5 Proyecciones Financieras
    doc.add_paragraph('1.5 Proyecciones Financieras Clave', style='Subtitulo')
    
    proyecciones_table = doc.add_table(rows=4, cols=3)
    proyecciones_table.style = 'Light Grid Accent 1'
    
    proyecciones_table.cell(0, 0).text = 'Métrica'
    proyecciones_table.cell(0, 1).text = 'Año 1'
    proyecciones_table.cell(0, 2).text = 'Año 3'
    
    proyecciones_table.cell(1, 0).text = 'Usuarios activos'
    proyecciones_table.cell(1, 1).text = '25.000'
    proyecciones_table.cell(1, 2).text = '150.000'
    
    proyecciones_table.cell(2, 0).text = 'Transacciones anuales'
    proyecciones_table.cell(2, 1).text = '15.000'
    proyecciones_table.cell(2, 2).text = '120.000'
    
    proyecciones_table.cell(3, 0).text = 'Volumen transaccional'
    proyecciones_table.cell(3, 1).text = '2,25M€'
    proyecciones_table.cell(3, 2).text = '18M€'
    
    doc.add_paragraph(
        'El punto de equilibrio se proyecta para el mes 10 de operaciones, con una inversión inicial total '
        'de 58.000€ distribuida en desarrollo tecnológico, marketing inicial y operaciones durante los '
        'primeros 12 meses.',
        style='NormalProfesional'
    )
    
    doc.add_page_break()
    
    # ========== 2. CONCEPTO DE NEGOCIO ==========
    doc.add_paragraph('2. CONCEPTO DE NEGOCIO', style='TituloPrincipal')
    doc.add_paragraph()
    
    # 2.1 Definición del Concepto
    doc.add_paragraph('2.1 Definición del Concepto', style='Subtitulo')
    doc.add_paragraph(
        'Treqe es más que una plataforma de intercambio; es un ecosistema digital diseñado para resolver '
        'el problema fundamental de la falta de liquidez en el mercado de segunda mano. El concepto se '
        'basa en una premisa simple pero poderosa: el valor de un artículo no reside únicamente en su '
        'precio de mercado, sino en su utilidad para una persona específica en un momento determinado.',
        style='NormalProfesional'
    )
    
    doc.add_paragraph(
        'Esta perspectiva permite crear un mercado donde las transacciones no están limitadas por la '
        'disponibilidad monetaria, sino por la complementariedad de necesidades y posesiones entre '
        'usuarios. Treqe actúa como facilitador inteligente de estas conexiones, utilizando algoritmos '
        'avanzados para identificar oportunidades de intercambio mutuamente beneficiosas que de otra '
        'forma permanecerían ocultas.',
        style='NormalProfesional'
    )
    
    # 2.2 Las Ruedas de Intercambio
    doc.add_paragraph('2.2 Las Ruedas de Intercambio: Innovación Central', style='Subtitulo')
    doc.add_paragraph(
        'El mecanismo central de Treqe son las "ruedas de intercambio", un sistema que supera las '
        'limitaciones del trueque tradicional mediante la creación de cadenas circulares de tres o '
        'más participantes. Este enfoque resuelve el problema clásico de la "doble coincidencia de '
        'deseos" que ha impedido la escalabilidad del trueque durante siglos.',
        style='NormalProfesional'
    )
    
    doc.add_paragraph(
        '**Ejemplo práctico de una rueda de intercambio:**',
        style='NormalProfesional'
    )
    
    ejemplo_table = doc.add_table(rows=4, cols=4)
    ejemplo_table.style = 'Light Shading Accent 2'
    
    ejemplo_table.cell(0, 0).text = 'Participante'
    ejemplo_table.cell(0, 1).text = 'Ofrece'
    ejemplo_table.cell(0, 2).text = 'Desea'
    ejemplo_table.cell(0, 3).text = 'Compensación'
    
    ejemplo_table.cell(1, 0).text = 'Ana'
    ejemplo_table.cell(1, 1).text = 'Bicicleta de montaña (valor: 450€)'
    ejemplo_table.cell(1, 2).text = 'Sofá de diseño (valor: 600€)'
    ejemplo_table.cell(1, 3).text = '+150€'
    
    ejemplo_table.cell(2, 0).text = 'Carlos'
    ejemplo_table.cell(2, 1).text = 'Sofá de diseño (valor: 600€)'
    ejemplo_table.cell(2, 2).text = 'Ordenador portátil (valor: 800€)'
    ejemplo_table.cell(2, 3).text = '+200€'
    
    ejemplo_table.cell(3, 0).text = 'Beatriz'
    ejemplo_table.cell(3, 1).text = 'Ordenador portátil (valor: 800€)'
    ejemplo_table.cell(3, 2).text = 'Bicicleta de montaña (valor: 450€)'
    ejemplo_table.cell(3, 3).text = '-350€'
    
    doc.add_paragraph(
        'En este ejemplo, cada participante obtiene exactamente lo que desea mediante una combinación '
        'de intercambio físico y ajuste monetario. El sistema calcula automáticamente las compensaciones '
        'óptimas y gestiona los pagos a través de Stripe Connect, garantizando seguridad y transparencia.',
        style='NormalProfesional'
    )
    
    # 2.3 Propuesta de Valor
    doc.add_paragraph('2.3 Propuesta de Valor para el Usuario', style='Subtitulo')
    
    propuestas = [
        ("**Para usuarios con liquidez limitada:**", 
         "Acceso a bienes de mayor valor mediante la combinación de trueque y pequeñas compensaciones monetarias"),
        
        ("**Para usuarios conscientes medioambientalmente:**",
         "Participación activa en la economía circular, maximizando la vida útil de los productos"),
        
        ("**Para coleccionistas y buscadores de rarezas:**",
         "Oportunidad de acceder a artículos únicos que raramente aparecen en el mercado monetario"),
        
        ("**Para personas en transición vital:**",
         "Flexibilidad para adaptar posesiones a nuevas circunstancias (mudanzas, cambios familiares, etc.)"),
        
        ("**Para la comunidad en general:**",
         "Creación de redes de confianza y colaboración más allá de la transacción comercial")
    ]
    
    for titulo, descrip