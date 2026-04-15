#!/usr/bin/env python3
"""
Script simple para completar los anexos vacíos del documento principal
"""

import docx
import os

def completar_anexos():
    # Documento principal
    doc_path = "plan_negocio/Plan_Negocio_Treqe_100%_COMPLETO_FINAL.docx"
    output_path = "plan_negocio/Plan_Negocio_Treqe_CON_ANEXOS_LLENOS.docx"
    
    if not os.path.exists(doc_path):
        print(f"Error: No se encuentra {doc_path}")
        return
    
    print(f"Abriendo documento: {doc_path}")
    doc = docx.Document(doc_path)
    
    # Buscar donde dice "ANEXOS"
    anexos_index = -1
    for i, p in enumerate(doc.paragraphs):
        if "ANEXOS" in p.text and len(p.text) < 20:
            anexos_index = i
            print(f"Encontrado 'ANEXOS' en párrafo {i}")
            break
    
    if anexos_index == -1:
        print("No se encontró la sección ANEXOS")
        return
    
    # Los siguientes párrafos deberían ser la lista de anexos
    # Vamos a reemplazar desde anexos_index+1 hasta encontrar algo que no sea un anexo
    start_replace = anexos_index + 1
    end_replace = start_replace
    
    while end_replace < len(doc.paragraphs):
        p = doc.paragraphs[end_replace]
        text = p.text.strip()
        # Si es un anexo (empieza con "Anexo" o tiene formato de lista)
        if text and ("Anexo" in text or text.startswith("Anexo")):
            end_replace += 1
        else:
            # Si encontramos algo que no es anexo, paramos
            break
    
    print(f"Reemplazando párrafos {start_replace} a {end_replace-1}")
    
    # Eliminar los párrafos vacíos de anexos
    for _ in range(start_replace, end_replace):
        if start_replace < len(doc.paragraphs):
            p = doc.paragraphs[start_replace]
            p._element.getparent().remove(p._element)
    
    # Ahora añadimos contenido real de anexos
    print("Añadiendo contenido real de anexos...")
    
    # Añadir espacio
    doc.add_paragraph()
    
    # ANEXO A
    doc.add_paragraph("ANEXO A: ESTRUCTURA DEL EQUIPO")
    doc.paragraphs[-1].style = "Heading 2"
    doc.add_paragraph("Roles mínimos requeridos para desarrollar el MVP:")
    doc.add_paragraph("1. Fundador/CEO - Estrategia, visión, fundraising")
    doc.add_paragraph("2. Desarrollador Full-stack - Plataforma técnica completa")
    doc.add_paragraph("3. Diseñador UX/UI - Experiencia de usuario e interfaz")
    doc.add_paragraph("Nota: Los CVs específicos se añadirán cuando se definan los miembros del equipo.")
    
    doc.add_page_break()
    
    # ANEXO B
    doc.add_paragraph("ANEXO B: METODOLOGÍA DE VALIDACIÓN")
    doc.paragraphs[-1].style = "Heading 2"
    doc.add_paragraph("Fases recomendadas para validar el mercado:")
    doc.add_paragraph("1. Investigación secundaria (datos públicos, informes sectoriales)")
    doc.add_paragraph("2. Encuestas online (200-500 participantes)")
    doc.add_paragraph("3. Entrevistas cualitativas (15-20 usuarios potenciales)")
    doc.add_paragraph("4. Piloto controlado en barrio específico")
    doc.add_paragraph("Nota: Los datos específicos se añadirán tras realizar las investigaciones.")
    
    doc.add_page_break()
    
    # ANEXO C
    doc.add_paragraph("ANEXO C: ESPECIFICACIONES TÉCNICAS")
    doc.paragraphs[-1].style = "Heading 2"
    doc.add_paragraph("Componentes principales de la plataforma:")
    doc.add_paragraph("• Sistema de catalogación (upload imágenes, etiquetado automático)")
    doc.add_paragraph("• Motor de búsqueda y matching (algoritmo de ruedas de intercambio)")
    doc.add_paragraph("• Sistema de intercambio (propuestas, chat, acuerdos, seguimiento)")
    doc.add_paragraph("• Perfiles y reputación (historial, calificaciones, verificaciones)")
    doc.add_paragraph("• Sistema de pagos (compensaciones monetarias, comisiones)")
    
    doc.add_page_break()
    
    # ANEXO D
    doc.add_paragraph("ANEXO D: GLOSARIO DE TÉRMINOS")
    doc.paragraphs[-1].style = "Heading 2"
    doc.add_paragraph("• Rueda de Intercambio: Ciclo cerrado A→B→C→A donde todos reciben valor")
    doc.add_paragraph("• Algoritmo de Matching: Sistema que encuentra combinaciones óptimas")
    doc.add_paragraph("• Compensación Monetaria: Pago pequeño para equilibrar diferencias de valor")
    doc.add_paragraph("• Liquidez: Facilidad para encontrar contrapartes de intercambio")
    doc.add_paragraph("• Two-sided Marketplace: Plataforma que conecta dos grupos de usuarios")
    doc.add_paragraph("• Network Effects: El valor aumenta exponencialmente con más usuarios")
    
    doc.add_page_break()
    
    # ANEXO E
    doc.add_paragraph("ANEXO E: FUENTES Y REFERENCIAS")
    doc.paragraphs[-1].style = "Heading 2"
    doc.add_paragraph("Fuentes públicas consultables:")
    doc.add_paragraph("• INE (Instituto Nacional de Estadística) - ine.es")
    doc.add_paragraph("• Eurostat - ec.europa.eu/eurostat")
    doc.add_paragraph("• Banco de España - bde.es")
    doc.add_paragraph("• CNMC (Comisión Nacional de Mercados y Competencia) - cnmc.es")
    doc.add_paragraph("• Observatorio Cetelem - Informes anuales de consumo")
    doc.add_paragraph("• Statista - Reports mercado segunda mano España/Europa")
    doc.add_paragraph("• Google Trends - Tendencias de búsqueda")
    doc.add_paragraph("• SimilarWeb - Tráfico de competidores")
    
    # Guardar
    print(f"Guardando documento: {output_path}")
    doc.save(output_path)
    
    # Verificar
    new_doc = docx.Document(output_path)
    print(f"Documento guardado. Párrafos totales: {len(new_doc.paragraphs)}")
    
    # Contar anexos
    anexos_count = 0
    for p in new_doc.paragraphs:
        if "ANEXO" in p.text and ":" in p.text:
            anexos_count += 1
    
    print(f"Anexos añadidos: {anexos_count}")
    
    return output_path

if __name__ == "__main__":
    print("=== COMPLETANDO ANEXOS VACÍOS ===")
    print("Este script añadirá contenido real a los anexos que están vacíos.")
    print()
    
    result = completar_anexos()
    
    if result:
        print()
        print("=" * 60)
        print("✅ ANEXOS COMPLETADOS EXITOSAMENTE")
        print("=" * 60)
        print(f"Documento actualizado: {result}")
        print()
        print("Contenido añadido:")
        print("• ANEXO A: Estructura del equipo (roles necesarios)")
        print("• ANEXO B: Metodología de validación (fases de investigación)")
        print("• ANEXO C: Especificaciones técnicas (componentes plataforma)")
        print("• ANEXO D: Glosario de términos (definiciones clave)")
        print("• ANEXO E: Fuentes y referencias (fuentes verificables)")
        print()
        print("Los anexos ahora tienen contenido real y útil.")
    else:
        print("❌ Error al completar los anexos.")