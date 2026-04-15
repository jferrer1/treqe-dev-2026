#!/usr/bin/env python3
"""
Reescribir el contenido adjunto con redacción extensa y humanizada
Manteniendo el estilo profesional y elaborado de las secciones 1-2
"""

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from datetime import datetime

def reescribir_con_humanizacion():
    """Reescribir contenido adjunto con estilo extenso y humanizado."""
    
    print("Reescribiendo contenido adjunto con redacción extensa y humanizada...")
    
    # Cargar documento base (con secciones 1-2 ya humanizadas)
    base_path = os.path.join(os.path.dirname(__file__), 'plan_negocio', 'Plan_Negocio_Treqe_Completo.docx')
    doc = Document(base_path)
    
    # Continuar desde donde termina el documento actual
    # (después de la sección 2.4)
    
    # ========== SECCIÓN 3: SOLUCIÓN INNOVADORA (HUMANIZADA) ==========
    
    doc.add_page_break()
    doc.add_heading('3. LA SOLUCIÓN TREQE: CUANDO LAS MATEMÁTICAS ENCUENTRAN SOLUCIONES DONDE LOS HUMANOS VEN PROBLEMAS', 0)
    
    # 3.1
    doc.add_heading('3.1 Más Allá del Trueque Tradicional: Un Salto Conceptual', 1)
    
    text_3_1 = """Si hay algo que la historia nos ha enseñado sobre el trueque es que, a pesar de su aparente simplicidad, presenta una limitación fundamental que ha impedido su escalabilidad como modelo comercial viable. Me refiero, por supuesto, a lo que los economistas denominan "el problema de la doble coincidencia de deseos". 

Durante siglos, esta limitación ha confinado el trueque a escalas reducidas, a intercambios ocasionales entre personas que casualmente querían exactamente lo que la otra ofrecía. Pero, ¿y si pudiéramos superar esta limitación? ¿Y si pudiéramos transformar lo que históricamente ha sido un proceso azaroso en uno sistemático y escalable?

Es precisamente aquí donde Treqe introduce lo que consideramos una innovación conceptual significativa: las "ruedas de intercambio inteligente". La idea es tan elegante en su concepción como poderosa en su aplicación. En lugar de limitarnos a buscar coincidencias directas entre dos personas (una tarea estadísticamente improbable), permitimos que tres, cuatro o incluso cinco usuarios participen en cadenas circulares donde cada uno recibe exactamente lo que desea, aunque no directamente de la persona a la que entrega su artículo.

Lo que hace especial a este sistema no es solo su capacidad para resolver un problema matemático, sino su comprensión profunda de la naturaleza humana. Reconoce que nuestras necesidades son complejas, que los objetos que poseemos tienen tanto valor sentimental como económico, y que la falta de liquidez no debería ser una barrera insalvable para satisfacer nuestras necesidades de renovación y cambio."""
    
    doc.add_paragraph(text_3_1)
    
    # 3.2
    doc.add_heading('3.2 El Mecanismo Detrás de la Magia: Un Proceso Pensado para Personas Reales', 1)
    
    text_3_2 = """Para comprender cómo Treqe transforma una idea conceptual en una experiencia tangible, es necesario analizar el mecanismo operativo paso a paso. Cada fase ha sido diseñada no solo pensando en la eficiencia técnica, sino especialmente en la experiencia humana.

**Fase 1: El Arte de Declarar lo que Tenemos y lo que Deseamos**
El proceso comienza con algo aparentemente simple pero profundamente significativo: invitar a los usuarios a reflexionar sobre lo que poseen y lo que realmente necesitan. No se trata de un inventario frío, sino de un ejercicio de autoconocimiento práctico.

Cuando un usuario ingresa a Treqe, se le pide que complete dos listas:
- **"Esto tengo":** Los artículos que actualmente posee pero que ya no se alinean con su vida actual
- **"Esto quiero":** Los objetos que necesita para mejorar su espacio, su comodidad, su funcionalidad

Para cada artículo, el sistema solicita información que va más allá de lo meramente descriptivo. Se piden fotografías desde múltiples ángulos, no por mero formalismo, sino porque entendemos que ver un objeto es comenzar a establecer una relación con él. Se solicita una valoración económica, pero el sistema ofrece sugerencias basadas en datos de mercado reales, evitando que los usuarios subestimen o sobreestimen el valor de sus posesiones.

**Fase 2: Cuando las Matemáticas se Convierten en Conectores Humanos**
Una vez que los usuarios han declarado sus preferencias, el sistema transforma esta información en una estructura matemática: un grafo dirigido. Puede sonar técnico, pero en esencia es bastante simple: cada usuario se convierte en un nodo, y las flechas (aristas) van desde lo que una persona tiene hacia lo que otra persona quiere.

Utilizando algoritmos de teoría de grafos (específicamente búsqueda en profundidad optimizada), el sistema busca ciclos cerrados de 3 a 5 nodos. Cada ciclo representa una posible cadena de intercambio donde todos obtienen lo que desean. Lo notable es la velocidad: cada búsqueda tiene un timeout de 500 milisegundos, garantizando que el sistema sea rápido y responsivo incluso a medida que crece.

**Fase 3: La Equidad como Principio de Diseño**
En el mundo real, es raro que dos artículos tengan exactamente el mismo valor económico. Por eso, cuando Treqe identifica un ciclo viable, no se limita a proponer intercambios directos. En su lugar, utiliza programación lineal (con el algoritmo PuLP) para calcular las compensaciones monetarias óptimas.

La optimización considera múltiples factores:
- Minimizar el flujo monetario total (reducir al máximo las transferencias)
- Distribuir las compensaciones de manera equitativa
- Respetar las preferencias individuales (algunos usuarios prefieren pagar menos aunque reciban menos valor, otros prefieren lo contrario)
- Garantizar que nadie se sienta en desventaja en la transacción

**Fase 4: La Negociación como Conversación, no como Confrontación**
Una vez identificado un ciclo y calculadas las compensaciones, Treqe no impone la solución. En su lugar, crea un espacio de negociación facilitada. Los participantes se conectan a través de un chat grupal en tiempo real (utilizando WebSockets), donde pueden:
- Ver la propuesta inicial del sistema
- Hacer preguntas sobre los artículos involucrados
- Solicitar fotografías adicionales o información complementaria
- Proponer ajustes menores a los términos
- Confirmar explícitamente su acuerdo

Este proceso de negociación no es una formalidad burocrática, sino un espacio para construir confianza. Los usuarios no están interactuando con un algoritmo frío, sino con otras personas reales que, como ellos, buscan mejorar sus vidas a través del intercambio.

**Fase 5: La Seguridad como Fundamento, no como Añadido**
Finalmente, una vez alcanzado el acuerdo, Treqe gestiona la ejecución completa con múltiples capas de seguridad:
- Los pagos de compensación se procesan a través de Stripe Connect, con fondos retenidos en escrow hasta que todas las partes confirman la recepción satisfactoria
- La logística se integra con APIs de empresas establecidas (Correos, SEUR), proporcionando tracking en tiempo real
- Un sistema de reputación granular permite evaluar múltiples dimensiones (puntualidad, precisión en descripciones, calidad del embalaje)
- Mecanismos de disputa escalonados ofrecen protección sin necesidad de intervención legal inmediata

Lo que emerge de este proceso no es solo una transacción económica, sino una experiencia humana completa: desde la reflexión inicial sobre lo que poseemos y necesitamos, pasando por la conexión con otros, hasta la satisfacción de haber obtenido lo que queríamos de manera justa y segura."""
    
    doc.add_paragraph(text_3_2)
    
    # 3.3
    doc.add_heading('3.3 Un Caso que Ilumina el Concepto: Ana, Carlos y Beatriz', 1)
    
    text_3_3 = """Para comprender realmente cómo funciona Treqe, nada mejor que analizar un caso concreto. No se trata de un ejemplo teórico, sino de una situación que refleja la realidad de millones de personas.

**Los Protagonistas y sus Historias:**

**Ana, 32 años, arquitecta en Barcelona:**
Ana compró su bicicleta de montaña Specialized Rockhopper durante sus años universitarios, cuando el ciclismo era más que un hobby: era una pasión. Ahora, con un trabajo exigente y menos tiempo libre, la bicicleta pasa la mayor parte del tiempo aparcada en el pequeño balcón de su apartamento. Lo que Ana realmente necesita es un sofá de diseño contemporáneo que se adapte mejor a su espacio y estilo actual. El problema no es de deseo, sino de recursos: aunque la bicicleta tiene un valor considerable (450€), Ana no tiene los 600€ que costaría el sofá que quiere.

**Carlos, 41 años, profesor universitario en Hospitalet de Llobregat:**
Carlos invirtió en un sofá de diseño de calidad hace dos años, pero su vida ha cambiado. Ha comenzado un proyecto personal: un canal educativo en YouTube sobre historia del arte. Para producir contenido de calidad, necesita un ordenador portátil potente para edición de video (valorado en 800€). El sofá, aunque perfectamente funcional y valioso (600€), ya no es una prioridad. Como muchos proyectos personales, el suyo tiene un presupuesto limitado.

**Beatriz, 28 años, diseñadora gráfica freelance en Badalona:**
Beatriz invirtió en un MacBook Pro M2 para su trabajo como diseñadora freelance. Recientemente, ha comenzado a trabajar de forma híbrida, yendo a una oficina compartida varios días a la semana. Necesita una bicicleta para sus desplazamientos urbanos (presupuesto: 450€), pero como freelance con ingresos variables, un desembolso de esa magnitud afectaría significativamente su liquidez mensual.

**El Problema Tradicional:**
Por separado, cada uno enfrenta una situación frustrante:
- Ana no puede conseguir su sofá porque Carlos no quiere su bicicleta
- Carlos no puede conseguir su ordenador porque Beatriz no quiere su sofá  
- Beatriz no puede conseguir su bicicleta porque Ana no quiere su ordenador

Es el clásico problema de la doble coincidencia de deseos: cada par podría intercambiar si sus deseos coincidieran, pero no lo hacen.

**La Solución Treqe:**
El algoritmo identifica que, aunque ningún intercambio directo es posible, existe un ciclo cerrado perfecto entre los tres. Propone:

**Intercambios físicos:**
1. Ana entrega su bicicleta a Beatriz
2. Carlos entrega su sofá a Ana
3. Beatriz entrega su ordenador a Carlos

**Compensaciones monetarias:**
- Ana paga 150€ a Carlos (la diferencia entre el sofá que recibe y la bicicleta que entrega)
- Carlos paga 200€ a Beatriz (la diferencia entre el ordenador que recibe y el sofá que entrega)
- Beatriz recibe 350€ neto (150€ de Ana + 200€ de Carlos)

**Los Resultados Transformadores:**
- **Ana:** Obtiene el sofá de diseño que necesita por solo 150€, en lugar de los 600€ que costaría nuevo. Su ahorro del 75% no es solo económico: es la posibilidad de tener un espacio que realmente refleja quién es ahora.
- **Carlos:** Consigue el ordenador que necesita para su proyecto YouTube por 200€, en lugar de 800€. El ahorro del 75% representa meses de trabajo que puede dedicar a crear contenido en lugar de ahorrar para equipamiento.
- **Beatriz:** Recibe la bicicleta que necesita para sus desplazamientos más 350€ de liquidez adicional. El valor total recibido (800€) le permite mejorar tanto su movilidad como su situación financiera.

**Lo que este Caso nos Enseña:**
Este ejemplo no es una anomalía estadística, sino una ilustración de cómo Treqe transforma problemas aparentemente insolubles en oportunidades de ganancia mutua. Lo notable no es solo el ahorro económico (aunque es significativo), sino cómo el sistema permite a cada persona avanzar en sus proyectos de vida sin las restricciones tradicionales de liquidez.

La belleza del sistema reside en su simplicidad final: tres personas, tres necesidades diferentes, una solución elegante que beneficia a todos. Es la materialización de la idea de que, a veces, lo que necesitamos no es más dinero, sino una forma más inteligente de utilizar el valor que ya poseemos."""
    
    doc.add_paragraph(text_3_3)
    
    doc.add_page_break()
    
    # ========== SECCIÓN 4: VENTAJA COMPETITIVA (HUMANIZADA) ==========
    
    doc.add_heading('4. POR QUÉ TREQE ES DIFERENTE: MÁS ALLÁ DE LA COMPETENCIA, HACIA UN NUEVO PARADIGMA', 0)
    
    # 4.1
    doc.add_heading('4.1 No Competir, sino Crear: El Arte de Encontrar Espacios Vacíos', 1)
    
    text_4_1 = """En el mundo empresarial contemporáneo, existe una tendencia casi obsesiva por analizar a la competencia, por posicionarse frente a actores establecidos, por competir en mercados saturados. Treqe adopta un enfoque radicalmente diferente: en lugar de competir en espacios ya ocupados, crea un nuevo espacio.

Cuando analizamos el panorama del mercado de segunda mano en España, vemos actores consolidados que han definido claramente sus territorios:
- **Wallapop** es el generalista masivo, el lugar donde encuentras de todo
- **Vinted** es la especialista en moda, la comunidad en torno a la ropa de segunda mano
- **Facebook Marketplace** es la opción casual, integrada en la red social
- **Milanuncios** es el tradicional, el que perdura por inercia y reconocimiento

Lo que resulta fascinante (y lo que representa nuestra oportunidad) es que ninguno de estos actores ha desarrollado soluciones robustas para el trueque. No es que no puedan hacerlo técnicamente; es que no les interesa hacerlo. Su modelo de negocio se fundamenta en transacciones monetarias: cobran comisiones sobre ventas. Desarrollar sistemas de trueque complejos podría, paradójicamente, reducir sus ingresos al permitir a los usuarios satisfacer necesidades sin transacciones monetarias.

Esta omisión estratégica por parte de los gigantes establecidos no es un descuido, sino una consecuencia lógica de sus modelos de negocio. Y es precisamente en este espacio vacío donde Treqe encuentra su razón de ser.

**Nuestro posicionamiento no es "somos mejores que Wallapop", sino "ofrecemos algo que Wallapop no ofrece".** No competimos en precio (aunque somos significativamente más baratos), ni en catálogo (aunque el nuestro es diverso), ni en experiencia de usuario (aunque la nuestra está optimizada para móvil). Competimos en concepto: ofrecemos trueque estructurado y escalable en un mercado donde el trueque ha sido históricamente marginal e informal.

Esta posición de "primer mover" en trueque estructurado nos otorga varias ventajas:
- **Ausencia de competencia directa inicial:** No necesitamos convencer a usuarios de que nos elijan en lugar de Wallapop; les ofrecemos algo que Wallapop no les ofrece
- **Captura de un segmento específico:** Atendemos a usuarios que prefieren intercambiar antes que vender, que tienen limitaciones de liquidez, que valoran la sostenibilidad sobre la transacción monetaria
- **Construcción de comunidad desde valores compartidos:** Nuestros primeros usuarios no vienen solo por utilidad, sino por identificación con un concepto

Lo que estamos construyendo no es solo otra plataforma de segunda mano, sino un nuevo paradigma de consumo: uno donde el valor se mueve no solo a través del dinero, sino a través de cadenas inteligentes de intercambio."""
    
    doc.add_paragraph(text_4_1)
    
    # Continuaré con más secciones...
    # Por ahora guardo el documento
    
    # Guardar documento
    output_dir = os.path.join(os.path.dirname(__file__), 'plan_negocio')
    os.makedirs(output_dir, exist_ok=True)
    
    output_path = os.path.join(output_dir, 'Plan_Negocio_Treqe_HUMANIZADO.docx')
    doc.save(output_path)
    
    print(f"Documento humanizado creado: