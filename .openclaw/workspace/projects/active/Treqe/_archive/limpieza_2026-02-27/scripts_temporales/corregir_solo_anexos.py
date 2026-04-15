#!/usr/bin/env python3
"""
Corregir SOLO los anexos del documento bueno, manteniendo todo lo demás intacto.
"""

import docx
import os
import shutil

def corregir_solo_anexos():
    """Corregir solo los anexos del documento original, manteniendo todo lo demás."""
    
    # Documento BUENO (el que ya estaba bien)
    documento_bueno = "plan_negocio/Plan_Negocio_Treqe_100%_COMPLETO_FINAL.docx"
    documento_corregido = "plan_negocio/Plan_Negocio_Treqe_100%_COMPLETO_FINAL_ANEXOS_CORREGIDOS.docx"
    
    if not os.path.exists(documento_bueno):
        print(f"ERROR: No se encuentra el documento bueno: {documento_bueno}")
        return None
    
    print("=== CORRIGIENDO SOLO LOS ANEXOS ===")
    print(f"Documento original (bueno): {documento_bueno}")
    print(f"Tamaño original: {os.path.getsize(documento_bueno)/1024:.1f} KB")
    
    # Crear backup del documento original
    backup_path = documento_bueno + ".backup_pre_correccion"
    shutil.copy2(documento_bueno, backup_path)
    print(f"Backup creado: {backup_path}")
    
    # Abrir documento
    doc = docx.Document(documento_bueno)
    print(f"Párrafos originales: {len(doc.paragraphs)}")
    
    # Buscar sección ANEXOS
    anexos_start = -1
    for i, p in enumerate(doc.paragraphs):
        if 'ANEXOS' in p.text.upper() and len(p.text) < 20:
            anexos_start = i
            print(f"Sección ANEXOS encontrada en párrafo {i}")
            break
    
    if anexos_start == -1:
        print("ERROR: No se encontró sección ANEXOS")
        return None
    
    # Determinar hasta dónde van los anexos actuales
    anexos_end = anexos_start + 1
    while anexos_end < len(doc.paragraphs):
        p = doc.paragraphs[anexos_end]
        # Si encontramos otra sección principal, paramos
        if p.style.name.startswith('Heading') and anexos_end > anexos_start + 1:
            break
        anexos_end += 1
    
    print(f"Eliminando anexos actuales (párrafos {anexos_start+1} a {anexos_end-1})")
    
    # Eliminar anexos existentes
    for _ in range(anexos_start + 1, anexos_end):
        if anexos_start + 1 < len(doc.paragraphs):
            p = doc.paragraphs[anexos_start + 1]
            p._element.getparent().remove(p._element)
    
    print("Añadiendo anexos corregidos...")
    
    # ANEXO A: ESTRUCTURA DEL EQUIPO (NO CVs inventados)
    doc.add_paragraph()
    doc.add_paragraph("ANEXO A: ESTRUCTURA Y ROLES DEL EQUIPO")
    doc.paragraphs[-1].style = "Heading 2"
    
    contenido_a = [
        "",
        "ROLES NECESARIOS PARA EL DESARROLLO DEL MVP:",
        "",
        "1. FUNDADOR/CEO",
        "   • Responsabilidades: Estrategia, visión de producto, fundraising, relaciones con inversores",
        "   • Experiencia ideal: Conocimiento del sector segunda mano, liderazgo, capacidad negociación",
        "   • Tiempo requerido: Dedicación completa desde inicio",
        "",
        "2. DESARROLLADOR FULL-STACK",
        "   • Responsabilidades: Desarrollo plataforma completa, arquitectura, deployment, mantenimiento",
        "   • Stack técnico: Node.js/Express o Python/Django, React/Next.js, PostgreSQL",
        "   • Habilidades clave: APIs REST, algoritmos de matching, deployment cloud, CI/CD",
        "   • Tiempo requerido: 80% dedicación inicial, reduciendo a 50% tras MVP",
        "",
        "3. DISEÑADOR UX/UI",
        "   • Responsabilidades: User research, wireframes, diseño interfaz, prototipado, user testing",
        "   • Herramientas: Figma, Adobe XD, user testing platforms",
        "   • Habilidades clave: Design thinking, accesibilidad, mobile-first design, prototipado rápido",
        "   • Tiempo requerido: 60% dedicación inicial, reduciendo a 30% tras diseño base",
        "",
        "4. COMMUNITY/GROWTH MANAGER (Fase 2)",
        "   • Responsabilidades: Adquisición usuarios, gestión comunidad, contenidos, marketing digital",
        "   • Habilidades: Growth hacking, community management, analítica digital, copywriting",
        "   • Tiempo requerido: 50% dedicación desde mes 4",
        "",
        "NOTA IMPORTANTE:",
        "Los CVs específicos de los miembros del equipo se añadirán cuando se definan los candidatos reales.",
        "Esta estructura representa los roles necesarios, no personas específicas.",
    ]
    
    for item in contenido_a:
        if item == "":
            doc.add_paragraph()
        else:
            doc.add_paragraph(item)
    
    doc.add_page_break()
    
    # ANEXO B: METODOLOGÍA DE VALIDACIÓN (NO datos inventados)
    doc.add_paragraph("ANEXO B: METODOLOGÍA PARA VALIDACIÓN DE MERCADO")
    doc.paragraphs[-1].style = "Heading 2"
    
    contenido_b = [
        "",
        "FASES RECOMENDADAS PARA VALIDAR EL MERCADO DE INTERCAMBIO INTELIGENTE:",
        "",
        "FASE 1: INVESTIGACIÓN SECUNDARIA (Semanas 1-2)",
        "   • Fuentes principales: INE, Eurostat, informes sectoriales públicos, datos competidores",
        "   • Método: Análisis datos existentes, benchmarking, identificación gaps de mercado",
        "   • Output esperado: Tamaño mercado estimado, tendencias, oportunidades específicas",
        "   • Coste estimado: €0 (fuentes públicas) - €500 (informes premium)",
        "",
        "FASE 2: ENCUESTAS CUANTITATIVAS (Semanas 3-4)",
        "   • Método: Encuestas online (Typeform, Google Forms) distribuida via redes sociales",
        "   • Muestra objetivo: 200-500 usuarios españoles activos en segunda mano",
        "   • Métricas clave: Intención uso, preocupaciones principales, disposición a pagar",
        "   • Output esperado: Datos cuantitativos sobre aceptación modelo, pricing optimal",
        "   • Coste estimado: €200-€800 (dependiendo de alcance y plataforma)",
        "",
        "FASE 3: ENTREVISTAS CUALITATIVAS (Semanas 5-6)",
        "   • Método: Entrevistas 1:1 semi-estructuradas con usuarios potenciales",
        "   • Participantes objetivo: 15-20 usuarios (mix vendedores/compradores/trueque)",
        "   • Objetivo principal: Deep dive en pain points, validar user flows, identificar fricciones",
        "   • Output esperado: Insights cualitativos, user personas detalladas, priorización features",
        "   • Coste estimado: €300-€600 (incentivos participantes + tiempo análisis)",
        "",
        "FASE 4: PILOTO CONTROLADO (Meses 2-3)",
        "   • Método: MVP funcional en barrio específico (ej: Gràcia, Barcelona)",
        "   • Participantes objetivo: 100-200 usuarios activos en el área definida",
        "   • Métricas clave: Engagement, retención D7/D30, tasa conversión, NPS",
        "   • Output esperado: Validación real del modelo, datos de uso, ajustes producto",
        "   • Coste estimado: €1,000-€2,000 (desarrollo MVP + marketing local)",
        "",
        "NOTA METODOLÓGICA:",
        "Los datos específicos de cada estudio se añadirán tras realizar las investigaciones correspondientes.",
        "Esta metodología proporciona un framework para validación progresiva del mercado.",
    ]
    
    for item in contenido_b:
        if item == "":
            doc.add_paragraph()
        else:
            doc.add_paragraph(item)
    
    doc.add_page_break()
    
    # ANEXO C: DISEÑOS DETALLADOS (DESARROLLAR MÁS)
    doc.add_paragraph("ANEXO C: ESPECIFICACIONES DE DISEÑO DE PLATAFORMA")
    doc.paragraphs[-1].style = "Heading 2"
    
    contenido_c = [
        "",
        "ARQUITECTURA COMPLETA DE LA PLATAFORMA TREQE:",
        "",
        "1. SISTEMA DE CATALOGACIÓN INTELIGENTE",
        "   • Upload múltiple de imágenes (3-10 por producto) con compresión automática",
        "   • Etiquetado automático por categoría/estado usando visión por computadora",
        "   • Valoración estimada por algoritmo basado en mercado, condición, demanda",
        "   • Descripción asistida por templates y sugerencias de IA",
        "   • Categorización jerárquica: Categoría principal → Subcategoría → Atributos específicos",
        "",
        "2. MOTOR DE BÚSQUEDA Y MATCHING AVANZADO",
        "   • Búsqueda semántica que entiende consultas naturales (ej: 'bicicleta de montaña usada')",
        "   • Filtros avanzados: Radio geográfico, categoría, rango de valor, fecha publicación, estado",
        "   • Algoritmo de matching para ruedas de intercambio:",
        "     - Identifica ciclos cerrados A→B→C→A",
        "     - Optimiza para maximizar satisfacción global",
        "     - Considera preferencias de usuarios (ej: 'prefiero electrónica por ropa')",
        "   • Sistema de recomendaciones personalizadas basado en historial y preferencias",
        "",
        "3. SISTEMA DE INTERCAMBIO Y NEGOCIACIÓN",
        "   • Propuesta 1-clic para ruedas sugeridas por el algoritmo",
        "   • Chat integrado para negociación de detalles específicos",
        "   • Acuerdo de términos digital: Fecha intercambio, lugar, condiciones específicas",
        "   • Sistema de seguimiento: Recordatorios automáticos, confirmaciones, check-ins",
        "   • Calificación post-intercambio: Sistema de 5 estrellas + comentarios específicos",
        "   • Sistema de disputas: Mediación escalonada (automática → soporte → arbitraje)",
        "",
        "4. PERFILES, REPUTACIÓN Y COMUNIDAD",
        "   • Perfil público completo: Historial intercambios, productos publicados, calificaciones",
        "   • Sistema de reputación multi-dimensional:",
        "     - Puntuación general (1-5 estrellas)",
        "     - Badges por logros (ej: 'Intercambiador ávido', 'Productos de calidad')",
        "     - Niveles de confianza basados en actividad y feedback",
        "   • Verificaciones progresivas: Email → Teléfono → Redes sociales → Documento identidad",
        "   • Elementos sociales: Seguimientos entre usuarios, likes productos, comentarios públicos",
        "",
        "5. SISTEMA DE PAGOS Y COMPENSACIONES",
        "   • Integración Stripe/PayPal para compensaciones monetarias",
        "   • Cálculo automático de diferencias de valor en ruedas de intercambio",
        "   • Historial completo de transacciones financieras",
        "   • Sistema de seguridad antifraude y verificación transacciones",
        "",
        "6. PANEL DE ADMINISTRACIÓN Y ANALYTICS",
        "   • Dashboard administrador: Métricas clave, moderación, gestión usuarios",
        "   • Analytics avanzado: Funnel conversión, engagement, retención, revenue",
        "   • Herramientas moderación: Reportes automáticos, revisión manual, acciones correctivas",
        "   • Sistema de notificaciones push/email personalizables",
    ]
    
    for item in contenido_c:
        if item == "":
            doc.add_paragraph()
        else:
            doc.add_paragraph(item)
    
    doc.add_page_break()
    
    # ANEXO D: GLOSARIO (MANTENER Y MEJORAR)
    doc.add_paragraph("ANEXO D: GLOSARIO DE TÉRMINOS")
    doc.paragraphs[-1].style = "Heading 2"
    
    contenido_d = [
        "",
        "TÉRMINOS ESPECÍFICOS DEL PROYECTO TREQE:",
        "",
        "• **Rueda de Intercambio:** Ciclo cerrado A→B→C→A donde cada usuario recibe un producto deseado y entrega uno disponible, creando liquidez donde antes había estancamiento.",
        "• **Algoritmo de Matching:** Sistema computacional que analiza preferencias, disponibilidades y ubicaciones para encontrar combinaciones óptimas de intercambio entre múltiples usuarios.",
        "• **Compensación Monetaria:** Pago pequeño (€1-€20) que equilibra diferencias de valor entre productos en una rueda de intercambio, facilitando transacciones que de otra forma no serían posibles.",
        "• **Liquidez de Plataforma:** Facilidad con la que los usuarios encuentran contrapartes de intercambio adecuadas; medida clave del éxito de una plataforma de trueque.",
        "• **Ciclo Cerrado de Valor:** Concepto económico donde productos circulan entre usuarios manteniendo su utilidad, evitando convertirse en residuos y maximizando vida útil.",
        "• **Market Maker (Creador de Mercado):** Rol de Treqe como facilitador que crea liquidez donde no existía, conectando oferta y demanda de forma inteligente.",
        "",
        "TÉRMINOS DE NEGOCIO Y TECNOLOGÍA:",
        "",
        "• **Two-sided Marketplace:** Plataforma que conecta dos grupos distintos de usuarios (en Treqe: oferentes y demandantes de productos de segunda mano).",
        "• **Network Effects (Efectos de Red):** Fenómeno donde el valor de la plataforma aumenta exponencialmente con cada usuario adicional, creando barreras de entrada naturales.",
        "• **Chicken-and-Egg Problem:** Desafío clásico de plataformas: ¿qué viene primero, los oferentes o los demandantes? Treqe lo resuelve con intercambios multi-partes.",
        "• **CAC/LTV:** Coste de Adquisición de Cliente / Valor de Vida del Cliente; métrica fundamental para sostenibilidad del negocio.",
        "• **MVP/PMF:** Producto Mínimo Viable / Ajuste Producto-Mercado; fases críticas de validación y crecimiento startup.",
        "• **Serverless Architecture:** Infraestructura cloud donde el proveedor gestiona servidores automáticamente, permitiendo escalabilidad y reduciendo costes fijos.",
        "• **CI/CD:** Integración Continua / Entrega Continua; metodología para desarrollo ágil con despliegues frecuentes y automatizados.",
        "",
        "TÉRMINOS DE ECONOMÍA CIRCULAR:",
        "",
        "• **Economía Circular:** Sistema económico que busca eliminar residuos mediante reutilización, reparación, refabricación y reciclaje continuo de materiales y productos.",
        "• **Obsolescencia Programada:** Práctica de diseñar productos con vida útil artificialmente limitada; Treqe combate mediante prolongación vida útil.",
        "• **Huella de Carbono:** Total emisiones gases efecto invernadero asociadas a un producto; reducida significativamente mediante reutilización.",
        "• **Upcycling:** Proceso de transformar productos desechados en nuevos de mayor valor; facilitado por intercambios que encuentran nuevos usos para objetos existentes.",
    ]
    
    for item in contenido_d:
        if item == "":
            doc.add_paragraph()
        else:
            doc.add_paragraph(item)
    
    doc.add_page_break()
    
    # ANEXO E: FUENTES VERÍDICAS (CONSULTABLES)
    doc.add_paragraph("ANEXO E: FUENTES Y REFERENCIAS")
    doc.paragraphs[-1].style = "Heading 2"
    
    contenido_e = [
        "",
        "FUENTES PÚBLICAS, OFICIALES Y VERIFICABLES:",
        "",
        "1. INSTITUCIONES OFICIALES ESPAÑOLAS Y EUROPEAS:",
        "   • **INE (Instituto Nacional de Estadística)** - https://www.ine.es",
        "     - Datos población, consumo, hábitos compra, indicadores económicos",
        "   • **Eurostat (Oficina Estadística Europea)** - https://ec.europa.eu/e