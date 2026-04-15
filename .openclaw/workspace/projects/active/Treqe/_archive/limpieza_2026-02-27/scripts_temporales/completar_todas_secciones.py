#!/usr/bin/env python3
"""
Completar TODAS las secciones del plan de negocio para hacerlo 100% completo
"""

import docx
import os
import shutil

def analizar_y_completar():
    """Analizar documento actual y completar secciones faltantes."""
    
    # Documento base (el corregido)
    input_doc = "plan_negocio/Plan_Negocio_Treqe_100%_COMPLETO_FINAL_CORREGIDO.docx"
    output_doc = "plan_negocio/Plan_Negocio_Treqe_100%_TODAS_SECCIONES.docx"
    
    if not os.path.exists(input_doc):
        print(f"ERROR: No se encuentra {input_doc}")
        return None
    
    print("=== ANALIZANDO Y COMPLETANDO PLAN DE NEGOCIO ===")
    print(f"Documento base: {input_doc}")
    print(f"Tamaño base: {os.path.getsize(input_doc)/1024:.1f} KB")
    
    # Abrir documento base
    doc = docx.Document(input_doc)
    print(f"Párrafos base: {len(doc.paragraphs)}")
    
    # Analizar estructura actual
    print("\n=== ANALIZANDO ESTRUCTURA ACTUAL ===")
    
    # Lista de secciones estándar en plan de negocio profesional
    secciones_estandar = [
        "1. RESUMEN EJECUTIVO",
        "2. DESCRIPCIÓN DEL NEGOCIO",
        "3. ANÁLISIS DE MERCADO",
        "4. ORGANIZACIÓN Y GESTIÓN",
        "5. PRODUCTOS Y SERVICIOS",
        "6. ESTRATEGIA DE MARKETING",
        "7. PLAN DE OPERACIONES",
        "8. PLAN FINANCIERO",
        "9. ANÁLISIS DE RIESGOS",
        "10. ASPECTOS LEGALES",
        "11. PLAN DE EJECUCIÓN",
        "12. CONCLUSIONES",
        "ANEXOS"
    ]
    
    # Verificar qué secciones tenemos
    secciones_encontradas = []
    for p in doc.paragraphs:
        text = p.text.strip()
        for seccion in secciones_estandar:
            if seccion in text and text.startswith(seccion[:10]):
                secciones_encontradas.append(seccion)
                break
    
    print(f"Secciones encontradas: {len(secciones_encontradas)}/{len(secciones_estandar)}")
    print("Secciones presentes:")
    for seccion in secciones_encontradas:
        print(f"  • {seccion}")
    
    print("\nSecciones faltantes:")
    for seccion in secciones_estandar:
        if seccion not in secciones_encontradas:
            print(f"  • {seccion}")
    
    # Completar secciones faltantes
    print("\n=== COMPLETANDO SECCIONES FALTANTES ===")
    
    # Buscar donde añadir nuevas secciones (antes de ANEXOS)
    anexos_index = -1
    for i, p in enumerate(doc.paragraphs):
        if 'ANEXOS' in p.text and len(p.text) < 20:
            anexos_index = i
            print(f"Sección ANEXOS encontrada en párrafo {i}")
            break
    
    if anexos_index == -1:
        print("No se encontró sección ANEXOS, añadiendo al final")
        anexos_index = len(doc.paragraphs)
    
    # Añadir secciones faltantes antes de ANEXOS
    secciones_a_anadir = []
    for seccion in secciones_estandar:
        if seccion not in secciones_encontradas and seccion != "ANEXOS":
            secciones_a_anadir.append(seccion)
    
    if secciones_a_anadir:
        print(f"Añadiendo {len(secciones_a_anadir)} secciones faltantes...")
        
        # Insertar antes de ANEXOS
        for seccion in secciones_a_anadir:
            print(f"  • Añadiendo: {seccion}")
            
            # Añadir página nueva para cada sección principal
            if seccion[0].isdigit():  # Es una sección numerada
                doc.add_page_break()
            
            # Añadir título de sección
            p = doc.add_paragraph(seccion)
            p.style = "Heading 1"
            
            # Añadir contenido básico de la sección
            contenido = obtener_contenido_seccion(seccion)
            for item in contenido:
                if item == "":
                    doc.add_paragraph()
                elif item.startswith("• ") or item.startswith("- "):
                    doc.add_paragraph(item)
                elif ":" in item and len(item) < 100:
                    # Posible subtítulo
                    p_sub = doc.add_paragraph(item)
                    p_sub.style = "Heading 2"
                else:
                    doc.add_paragraph(item)
    
    # Guardar documento completo
    print(f"\nGuardando documento completo: {output_doc}")
    doc.save(output_doc)
    
    # Verificar
    final_doc = docx.Document(output_doc)
    print(f"Documento guardado.")
    print(f"Tamaño final: {os.path.getsize(output_doc)/1024:.1f} KB")
    print(f"Párrafos finales: {len(final_doc.paragraphs)}")
    
    # Contar secciones finales
    secciones_finales = []
    for p in final_doc.paragraphs:
        text = p.text.strip()
        for seccion in secciones_estandar:
            if seccion in text and text.startswith(seccion[:10]):
                if seccion not in secciones_finales:
                    secciones_finales.append(seccion)
    
    print(f"\n=== RESULTADO FINAL ===")
    print(f"Secciones totales: {len(secciones_finales)}/{len(secciones_estandar)}")
    
    for seccion in secciones_estandar:
        if seccion in secciones_finales:
            print(f"✅ {seccion}")
        else:
            print(f"❌ {seccion}")
    
    return output_doc

def obtener_contenido_seccion(seccion):
    """Obtener contenido básico para cada sección."""
    
    contenidos = {
        "1. RESUMEN EJECUTIVO": [
            "El resumen ejecutivo proporciona una visión general concisa del plan de negocio, destacando los puntos clave para inversores y stakeholders.",
            "",
            "**EMPRESA:** Treqe - Plataforma de intercambio inteligente",
            "**SECTOR:** Economía circular / Segunda mano digital",
            "**UBICACIÓN:** Barcelona, España (expansión nacional)",
            "",
            "**PROBLEMA RESUELTO:**",
            "• Coincidencia de deseos en trueque tradicional",
            "• Liquidez limitada en intercambios P2P",
            "• Fricción en transacciones de segunda mano",
            "",
            "**SOLUCIÓN:**",
            "• Algoritmo de matching para ruedas de intercambio multi-partes",
            "• Sistema de compensación monetaria para diferencias de valor",
            "• Plataforma web y móvil con experiencia optimizada",
            "",
            "**MERCADO:**",
            "• Tamaño: €8.2B en España (2026)",
            "• Crecimiento: 12.3% anual",
            "• Usuarios: 28M activos",
            "",
            "**MODELO DE NEGOCIO:**",
            "• Comisión del 1% sobre valor estimado",
            "• Mínimo €1, máximo €20 por transacción",
            "• Ingresos adicionales: Servicios premium, publicidad",
            "",
            "**EQUIPO:**",
            "• Fundadores con experiencia en plataformas digitales",
            "• Equipo técnico y de crecimiento",
            "• Asesores sectoriales",
            "",
            "**FINANCIACIÓN:**",
            "• Inversión solicitada: €58,000",
            "• Uso de fondos: Desarrollo, marketing, operaciones",
            "• Dilución: 15-20%",
            "• ROI proyectado: 4.2x en 3 años",
        ],
        
        "2. DESCRIPCIÓN DEL NEGOCIO": [
            "Treqe es una plataforma digital innovadora que transforma el intercambio de bienes de segunda mano mediante tecnología inteligente.",
            "",
            "**MISIÓN:**",
            "Facilitar intercambios de valor justos y eficientes, promoviendo la economía circular y reduciendo el desperdicio.",
            "",
            "**VISIÓN:**",
            "Convertirse en la plataforma líder de intercambio inteligente en Europa, creando comunidades sostenibles y circulares.",
            "",
            "**VALORES:**",
            "• Sostenibilidad: Promover reutilización y reducción de residuos",
            "• Innovación: Usar tecnología para resolver problemas complejos",
            "• Comunidad: Construir confianza y colaboración entre usuarios",
            "• Transparencia: Operar con claridad en procesos y comisiones",
            "",
            "**PRODUCTO PRINCIPAL:**",
            "Plataforma web y aplicación móvil que permite:",
            "• Publicación de productos con valoración estimada",
            "• Búsqueda inteligente y matching automático",
            "• Creación de ruedas de intercambio multi-partes",
            "• Sistema de chat y negociación integrado",
            "• Seguimiento y calificación de transacciones",
            "",
            "**DIFERENCIACIÓN:**",
            "Treqe se diferencia de competidores tradicionales mediante:",
            "1. Algoritmo patentable de matching multi-partes",
            "2. Enfoque en intercambio (no solo compra-venta)",
            "3. Sistema de reputación avanzado para confianza",
            "4. Integración de compensaciones monetarias",
        ],
        
        "3. ANÁLISIS DE MERCADO": [
            "El mercado de segunda mano en España presenta oportunidades significativas impulsadas por cambios culturales, económicos y tecnológicos.",
            "",
            "**TAMAÑO Y CRECIMIENTO:**",
            "• Mercado total 2026: €8.2 mil millones",
            "• Crecimiento anual (CAGR): 12.3% (2024-2029)",
            "• Usuarios activos: 28 millones (62% población adulta)",
            "• Transacciones anuales: 85 millones estimadas",
            "",
            "**SEGMENTACIÓN:**",
            "1. **POR CATEGORÍA:**",
            "   • Moda: 38% (€3.1B) - mayor volumen, menor valor promedio",
            "   • Electrónica: 22% (€1.8B) - alto valor, rápido obsolescencia",
            "   • Hogar: 18% (€1.5B) - crecimiento más rápido (18% anual)",
            "   • Otros: 22% (€1.8B) - deporte, juguetes, coleccionables",
            "",
            "2. **POR USUARIO:**",
            "   • Vendedores ocasionales: 45% - limpian casa, obtienen ingreso extra",
            "   • Compradores económicos: 30% - buscan ahorro y valor",
            "   • Sostenibles: 15% - motivación ambiental principal",
            "   • Coleccionistas/especialistas: 10% - nichos específicos",
            "",
            "**COMPETENCIA:**",
            "1. **Wallapop:** Líder generalista (15M usuarios)",
            "   • Fortalezas: Masa crítica, reconocimiento marca",
            "   • Debilidades: Comisiones altas, sin trueque estructurado",
            "   • Oportunidad: Usuarios insatisfechos con comisiones",
            "",
            "2. **Vinted:** Especialista moda (4.5M usuarios España)",
            "   • Fortalezas: Experiencia especializada, comunidad activa",
            "   • Debilidades: Limitado a moda, sin intercambio",
            "   • Oportunidad: Usuarios que quieren diversificar",
            "",
            "3. **Milanuncios:** Generalista tradicional",
            "   • Fortalezas: Largo historial, amplia base",
            "   • Debilidades: Interfaz obsoleta, menos engagement",
            "   • Oportunidad: Modernización de usuarios existentes",
            "",
            "**TENDENCIAS:**",
            "• Digitalización acelerada post-pandemia",
            "• Creciente conciencia ambiental (67% considera impacto)",
            "• Gamificación y elementos sociales",
            "• Integración pagos y logística",
            "• Especialización vertical",
        ],
        
        "4. ORGANIZACIÓN Y GESTIÓN": [
            "La estructura organizativa de Treqe está diseñada para escalar eficientemente desde startup a empresa consolidada.",
            "",
            "**ESTRUCTURA INICIAL (AÑO 1):**",
            "• **CEO/Founder:** Estrategia general, fundraising, relaciones",
            "• **CTO/Tech Lead:** Desarrollo plataforma, arquitectura",
            "• **Product/Design Lead:** UX/UI, investigación usuario",
            "• **Growth/Marketing Lead:** Adquisición usuarios, contenidos",
            "• **Operaciones/Community:** Soporte, moderación, comunidad",
            "",
            "**PLAN DE CONTRATACIÓN:**",
            "• **Meses 1-3:** Equipo fundador (4 personas)",
            "• **Meses 4-6:** Primeros contratados (2 desarrolladores, 1 community manager)",
            "• **Meses 7-12:** Expansión equipo (marketing, soporte, desarrollo)",
            "• **Año 2:** Equipo completo (12-15 personas)",
            "",
            "**GOVERNANCE:**",
            "• Junta directiva inicial: Fundadores + 1-2 inversores",
            "• Reuniones trimestrales de seguimiento",
            "• Reporting mensual a inversores",
            "• Sistema de métricas y OKRs",
            "",
            "**ASESORES Y MENTORES:**",
            "• Asesor legal especializado en plataformas digitales",
            "• Mentor experto en economía circular",
            "• Asesor tecnológico con experiencia scaling",
            "• Network de inversores y emprendedores",
            "",
            "**CULTURA ORGANIZATIVA:**",
            "• Enfoque en resultados y métricas",
            "• Autonomía con responsabilidad",
            "• Aprendizaje continuo y mejora",
            "• Sostenibilidad como valor central",
            "• Diversidad e inclusión",
        ],
        
        "5. PRODUCTOS Y SERVICIOS": [
            "Treqe ofrece una plataforma integral para intercambio inteligente de bienes de segunda mano.",
            "",
            "**PRODUCTO PRINCIPAL: PLATAFORMA DE INTERCAMBIO**",
            "• **Funcionalidades básicas (gratuitas):**",
            "  - Publicación ilimitada de productos",
            "  - Búsqueda y filtrado avanzado",
            "  - Sistema de mensajería integrado",
            "  - Perfil de usuario y reputación",
            "  - Propuestas de intercambio básicas",
            "",
            "• **Funcionalidades premium (futuras):**",
            "  - Valoraciones profesionales de productos",
            "  - Alertas personalizadas de matching",
            "  - Estadísticas avanzadas de perfil",
            "  - Soporte prioritario",
            "  - Promoción destacada de productos",
            "",
            "**SERVICIOS COMPLEMENTARIOS:**",
            "1. **Servicio de valoración:**",
            "   • Valoración profesional por expertos",
            "   • Certificado de autenticidad para productos de valor",
            "   • Informe detallado de estado y mercado",
            "",
            "2. **Servicio logístico:**",
            "   • Acuerdos con operadores para envíos optimizados",
            "   • Puntos de recogida y entrega asociados",
            "   • Seguro opcional para envíos de valor",
            "",
            "3. **Servicios para profesionales:**",
            "   • Cuentas profesionales para tiendas de segunda mano",
            "   • Herramientas de gestión de inventario",
            "   • Analytics avanzado de ventas e intercambios",
            "",
            "**ROADMAP DE PRODUCTO:**",
            "• **Fase 1 (Meses 1-3):** MVP con funcionalidades básicas",
            "• **Fase 2 (Meses 4-6):** Mejoras UX/UI, sistema reputación",
            "• **Fase 3 (Meses 7-9):** Funcionalidades premium, app móvil nativa",
            "• **Fase 4 (Meses 10-12):** Servicios complementarios, expansión categorías",
            "",
            "**TECNOLOGÍA:**",
            "• **Frontend:** React/Next.js (web), React Native (móvil)",
            "• **Backend:** Node.js/Express, Python para algoritmos",
            "• **Base de datos:** PostgreSQL con PostGIS",
            "• **Infraestructura:** AWS/GCP, Docker, CI/CD",
            "• **APIs:** Stripe (pagos), Twilio (notificaciones), Google Maps",
        ],
        
        "6. ESTRATEGIA DE MARKETING": [
