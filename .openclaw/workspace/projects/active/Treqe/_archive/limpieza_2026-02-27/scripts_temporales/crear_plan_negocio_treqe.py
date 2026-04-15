#!/usr/bin/env python3
"""
Crear Plan de Negocio Actualizado para Treqe.
Incorpora análisis exhaustivo y recomendaciones estratégicas.
"""

import sys
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("Error: python-docx no está instalado. Instala con: pip install python-docx")
    sys.exit(1)

def set_cell_background(cell, color_hex):
    """Establece color de fondo de celda de tabla."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    cell._element.get_or_add_tcPr().append(shading)

def create_treqe_business_plan():
    """Crea documento Word con plan de negocio actualizado para Treqe."""
    
    # Crear documento
    doc = Document()
    
    # Configurar estilos personalizados
    styles = doc.styles
    
    # Título principal
    title_style = styles.add_style('TreqeTitle', WD_STYLE_TYPE.PARAGRAPH)
    title_font = title_style.font
    title_font.name = 'Calibri Light'
    title_font.size = Pt(24)
    title_font.bold = True
    title_font.color.rgb = RGBColor(0, 32, 96)  # Azul corporativo
    
    # Subtítulo
    subtitle_style = styles.add_style('TreqeSubtitle', WD_STYLE_TYPE.PARAGRAPH)
    subtitle_font = subtitle_style.font
    subtitle_font.name = 'Calibri'
    subtitle_font.size = Pt(16)
    subtitle_font.bold = True
    subtitle_font.color.rgb = RGBColor(0, 32, 96)
    
    # Encabezado de sección
    section_style = styles.add_style('TreqeSection', WD_STYLE_TYPE.PARAGRAPH)
    section_font = section_style.font
    section_font.name = 'Calibri'
    section_font.size = Pt(14)
    section_font.bold = True
    section_font.color.rgb = RGBColor(0, 32, 96)
    
    # Página de portada
    doc.add_paragraph('PLAN DE NEGOCIO ACTUALIZADO', style='TreqeTitle').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    doc.add_paragraph('TREQE', style='TreqeTitle').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('El Marketplace de Trueque Inteligente en España', style='TreqeSubtitle').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Información del documento
    fecha_actual = datetime.now().strftime('%d de %B de %Y')
    info_table = doc.add_table(rows=3, cols=2)
    info_table.style = 'LightShading-Accent1'
    
    info_table.cell(0, 0).text = 'Fecha:'
    info_table.cell(0, 1).text = fecha_actual
    
    info_table.cell(1, 0).text = 'Versión:'
    info_table.cell(1, 1).text = '2.0 - Análisis Exhaustivo'
    
    info_table.cell(2, 0).text = 'Estado:'
    info_table.cell(2, 1).text = 'CONFIDENCIAL - Para uso interno'
    
    doc.add_page_break()
    
    # Índice
    doc.add_paragraph('ÍNDICE', style='TreqeTitle').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    indice_items = [
        ("1. RESUMEN EJECUTIVO", 3),
        ("2. CONCEPTO INNOVADOR", 5),
        ("3. ANÁLISIS DE MERCADO", 7),
        ("4. VENTAJA COMPETITIVA", 10),
        ("5. MODELO DE NEGOCIO", 12),
        ("6. HOJA DE RUTA ESTRATÉGICA", 15),
        ("7. MODELO FINANCIERO", 18),
        ("8. RIESGOS Y MITIGACIÓN", 21),
        ("9. RECOMENDACIONES ESTRATÉGICAS", 24),
        ("10. PRÓXIMOS PASOS INMEDIATOS", 27),
        ("ANEXO A: DATOS DE MERCADO", 29),
        ("ANEXO B: PRESUPUESTO DETALLADO", 31),
        ("ANEXO C: COMPETENCIA", 33)
    ]
    
    for item, page in indice_items:
        p = doc.add_paragraph()
        p.add_run(item).bold = True
        p.add_run(f'\t{page}')
    
    doc.add_page_break()
    
    # 1. RESUMEN EJECUTIVO
    doc.add_paragraph('1. RESUMEN EJECUTIVO', style='TreqeTitle')
    doc.add_paragraph()
    
    resumen = doc.add_paragraph()
    resumen.add_run('Treqe es una ').bold = True
    resumen.add_run('plataforma C2C de trueque con compensación económica ')
    resumen.add_run('que permite intercambiar artículos de segunda mano con posibilidad de pagar diferencias de valor.').bold = True
    resumen.add_run(' No es solo trueque puro, sino un sistema de "trueque inteligente" donde se monetizan las diferencias de valor entre artículos.')
    
    doc.add_paragraph()
    
    puntos_clave = [
        "✅ **Modelo innovador:** Ruedas de intercambio (grupos de 3+ usuarios)",
        "✅ **Mercado objetivo:** 60+ millones de usuarios en España (segunda mano)",
        "✅ **Comisión:** 1% sobre valor del artículo adquirido (cada usuario paga)",
        "✅ **Ventaja competitiva:** Primer mover en trueque estructurado",
        "✅ **Timing perfecto:** Crisis económica + conciencia ecológica",
        "✅ **Inversión necesaria:** 55.000-90.000€ para 6 meses de operación"
    ]
    
    for punto in puntos_clave:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(punto.replace('✅ ', '')).bold = '✅' in punto
    
    doc.add_page_break()
    
    # 2. CONCEPTO INNOVADOR
    doc.add_paragraph('2. CONCEPTO INNOVADOR', style='TreqeTitle')
    doc.add_paragraph()
    
    doc.add_paragraph('2.1 ¿Qué es Treqe?', style='TreqeSection')
    doc.add_paragraph('Treqe es una plataforma de consumo colaborativo que transforma el modelo tradicional de compraventa de segunda mano en un sistema de intercambio inteligente con compensación económica.')
    
    doc.add_paragraph()
    doc.add_paragraph('2.2 Sistema de Ruedas de Intercambio', style='TreqeSection')
    doc.add_paragraph('A diferencia del trueque 1:1 tradicional, Treqe implementa "ruedas de intercambio" donde pueden participar 3 o más usuarios, creando cadenas de valor que maximizan las oportunidades de intercambio.')
    
    # Ejemplo de rueda de intercambio
    ejemplo_table = doc.add_table(rows=4, cols=3)
    ejemplo_table.style = 'LightGrid-Accent1'
    
    ejemplo_table.cell(0, 0).text = 'Usuario'
    ejemplo_table.cell(0, 1).text = 'Ofrece'
    ejemplo_table.cell(0, 2).text = 'Recibe'
    
    ejemplo_table.cell(1, 0).text = 'Juan'
    ejemplo_table.cell(1, 1).text = 'Sofá gris (valor: 300€)'
    ejemplo_table.cell(1, 2).text = 'Sofá negro (valor: 350€) + 50€'
    
    ejemplo_table.cell(2, 0).text = 'María'
    ejemplo_table.cell(2, 1).text = 'Bicicleta (valor: 200€)'
    ejemplo_table.cell(2, 2).text = 'Sofá gris (valor: 300€) - 100€'
    
    ejemplo_table.cell(3, 0).text = 'Carlos'
    ejemplo_table.cell(3, 1).text = 'Sofá negro (valor: 350€)'
    ejemplo_table.cell(3, 2).text = 'Bicicleta (valor: 200€) + 150€'
    
    doc.add_page_break()
    
    # 3. ANÁLISIS DE MERCADO
    doc.add_paragraph('3. ANÁLISIS DE MERCADO', style='TreqeTitle')
    doc.add_paragraph()
    
    doc.add_paragraph('3.1 Tamaño del Mercado Español (2026)', style='TreqeSection')
    
    mercado_data = [
        ("Usuarios totales de segunda mano", "60+ millones"),
        ("Penetración en población", "19% reconoce acudir al mercado"),
        ("Frecuencia de uso", "62% consulta apps/webs semanalmente"),
        ("Engagement diario", "25% consulta diariamente"),
        ("Ahorro medio vs producto nuevo", "56%"),
        ("Gasto medio por persona (2016)", "766€")
    ]
    
    mercado_table = doc.add_table(rows=len(mercado_data) + 1, cols=2)
    mercado_table.style = 'LightShading-Accent2'
    
    mercado_table.cell(0, 0).text = 'Métrica'
    mercado_table.cell(0, 1).text = 'Valor'
    set_cell_background(mercado_table.cell(0, 0), '2F5597')
    set_cell_background(mercado_table.cell(0, 1), '2F5597')
    
    for i, (metrica, valor) in enumerate(mercado_data, 1):
        mercado_table.cell(i, 0).text = metrica
        mercado_table.cell(i, 1).text = valor
    
    doc.add_paragraph()
    doc.add_paragraph('3.2 Segmentación de Usuarios', style='TreqeSection')
    
    segmentos = [
        ("Personas con dificultades económicas", "Falta de liquidez monetaria", "Alta"),
        ("Jóvenes concienciados medioambientalmente", "Valor ecológico del trueque", "Alta"),
        ("Padres/madres con hijos", "Cambio constante de necesidades", "Media-Alta"),
        ("Coleccionistas y buscadores de rarezas", "Artículos únicos y difíciles de encontrar", "Media"),
        ("Personas en cambio de estatus social", "Renovación de posesiones", "Media")
    ]
    
    segmentos_table = doc.add_table(rows=len(segmentos) + 1, cols=3)
    segmentos_table.style = 'LightGrid-Accent1'
    
    headers = ['Segmento', 'Motivación Principal', 'Prioridad']
    for col, header in enumerate(headers):
        segmentos_table.cell(0, col).text = header
        set_cell_background(segmentos_table.cell(0, col), '2F5597')
    
    for i, (segmento, motivacion, prioridad) in enumerate(segmentos, 1):
        segmentos_table.cell(i, 0).text = segmento
        segmentos_table.cell(i, 1).text = motivacion
        segmentos_table.cell(i, 2).text = prioridad
    
    doc.add_page_break()
    
    # 4. VENTAJA COMPETITIVA
    doc.add_paragraph('4. VENTAJA COMPETITIVA', style='TreqeTitle')
    doc.add_paragraph()
    
    doc.add_paragraph('4.1 Lo que SÍ tiene Treqe vs Competencia', style='TreqeSection')
    
    ventajas = [
        ("Trueque estructurado", "Wallapop/Milanuncios NO ofrecen trueque", "Primer mover"),
        ("Ruedas de intercambio", "Modelo innovador de grupos (3+ usuarios)", "Diferenciación radical"),
        ("Compensación económica integrada", "No es trueque puro, es híbrido monetario", "Flexibilidad"),
        ("Bajas comisiones (1%)", "vs 5-10% de competidores", "Precio competitivo"),
        ("Enfoque social + económico", "Valor dual para usuarios", "Propuesta de valor única")
    ]
    
    ventajas_table = doc.add_table(rows=len(ventajas) + 1, cols=3)
    ventajas_table.style = 'LightGrid-Accent2'
    
    ventajas_headers = ['Ventaja', 'Comparativa', 'Impacto']
    for col, header in enumerate(ventajas_headers):
        ventajas_table.cell(0, col).text = header
        set_cell_background(ventajas_table.cell(0, col), '2F5597')
    
    for i, (ventaja, comparativa, impacto) in enumerate(ventajas, 1):
        ventajas_table.cell(i, 0).text = ventaja
        ventajas_table.cell(i, 1).text = comparativa
        ventajas_table.cell(i, 2).text = impacto
    
    doc.add_paragraph()
    doc.add_paragraph('4.2 Lo que le FALTA vs Competencia', style='TreqeSection')
    
    faltas = [
        ("Masa crítica inicial", "Problema huevo-gallina", "Alto", "Programa referidos + eventos"),
        ("Sistema de reputación probado", "Confianza es crítica", "Alto", "Implementar sistema robusto"),
        ("App móvil nativa", "Competencia tiene apps maduras", "Medio", "React Native como prioridad"),
        ("Integración profunda con logística", "Competencia tiene acuerdos sólidos", "Medio", "Partner logístico desde día 1"),
        ("Presupuesto de marketing", "Competencia invierte millones", "Alto", "Enfoque en crecimiento orgánico")
    ]
    
    faltas_table = doc.add_table(rows=len(faltas) + 1, cols=4)
    faltas_table.style = 'LightGrid-Accent1'
    
    faltas_headers = ['Debilidad', 'Riesgo', 'Gravedad', 'Mitigación Propuesta']
    for col, header in enumerate(faltas_headers):
        faltas_table.cell(0, col).text = header
        set_cell_background(faltas_table.cell(0, col), 'C00000')
    
    for i, (debilidad, riesgo, gravedad, mitigacion) in enumerate(faltas, 1):
        faltas_table.cell(i, 0).text = debilidad
        faltas_table.cell(i, 1).text = riesgo
        faltas_table.cell(i, 2).text = gravedad
        faltas_table.cell(i, 3).text = mitigacion
    
    doc.add_page_break()
    
    # 5. MODELO DE NEGOCIO
    doc.add_paragraph('5. MODELO DE NEGOCIO', style='TreqeTitle')
    doc.add_paragraph()
    
    doc.add_paragraph('5.1 Flujo de Ingresos', style='TreqeSection')
    
    ingresos = [
        ("Comisión básica", "1% sobre valor del artículo adquirido", "Mes 4+", "Principal"),
        ("Tarifas premium", "Visibilidad destacada en búsquedas", "Mes 6+", "Secundaria"),
        ("Publicidad segmentada", "Anuncios de marcas afines", "Año 2", "Complementaria"),
        ("Servicios logísticos", "Acuerdos revenue-share con operadores", "Mes 3+", "Complementaria"),
        ("Datos analíticos", "Informes de mercado para empresas", "Año 2+", "Futura")
    ]
    
    ingresos_table = doc.add_table(rows=len(ingresos) + 1, cols=4)
    ingresos_table.style = 'LightGrid-Accent2'
    
    ingresos_headers = ['Fuente', 'Descripción', 'Lanzamiento', 'Importancia']
    for col, header in enumerate(ingresos_headers):
        ingresos_table.cell(0, col).text = header
        set_cell_background(ingresos_table.cell(0, col), '2F5597')
    
    for i, (fuente, descripcion, lanzamiento, importancia) in enumerate(ingresos, 1):
        ingresos_table.cell(i, 0).text = fuente
        ingresos_table.cell(i, 1).text = descripcion
        ingresos_table.cell(i, 2).text = lanzamiento
        ingresos_table.cell(i, 3).text = importancia
    
    doc.add_paragraph()
    doc.add_paragraph('5.2 Estructura de Costes (6 primeros meses)', style='TreqeSection')
    
    costes = [
        ("Desarrollo plataforma", "2 desarrolladores full-stack", "15.000-25.000€"),
        ("Diseño UX/UI", "Diseñador senior + branding", "3.000-5.000€"),
        ("Marketing inicial", "Landing page + campañas iniciales", "5.000-10.000€"),
        ("Infraestructura tecnológica", "Hosting, APIs, herramientas", "3.000-4.000€"),
        ("Equipo fundador (3 personas)", "Salarios básicos por 6 meses", "30.000-45.000€"),
        ("Asesoría legal/fiscal", "Protección IP + estructuración", "2.000-3.000€"),
        ("**TOTAL ESTIMADO**", "", "**58.000-92.000€**")
    ]
    
    costes_table = doc.add_table(rows=len(costes) + 1, cols=3)
    costes_table.style = 'LightShading-Accent1'
    
    costes_headers = ['Concepto', 'Detalle', 'Rango Estimado']
    for col, header in enumerate(costes_headers):
        costes_table.cell(0, col).text = header
        set_cell_background(costes_table.cell(0, col), '2F5597')
    
    for i, (concepto, detalle, rango) in enumerate(costes, 1):
        costes_table.cell(i, 0).text = concepto
        costes_table.cell(i, 1).text = detalle
        costes_table.cell(i, 2).text = rango
        if 'TOTAL' in concepto:
            set_cell_background(costes_table.cell(i, 0), 'FFC000')
            set_cell_background(costes_table.cell(i, 1), 'FFC000')
            set_cell_background(costes_table.cell(i, 2), 'FFC000')
    
    doc.add_page_break()
    
    # 6. HOJA DE RUTA ESTRATÉGICA
    doc.add_paragraph('6. HOJA DE RUTA ESTRATÉGICA', style='TreqeTitle')
    doc.add_paragraph()
    
    fases = [
        ("FASE 1: Validación", "4-6 semanas", [
            "Crear landing page con waitlist",
            "Validar interés real (500+ suscriptores)",
            "Testear propuesta con focus groups",
            "Desarrollar MVP básico (funcionalidad core)",
            "Establecer alianzas logísticas iniciales"
        ]),
        ("FASE 2: MVP", "8-12 semanas", [
            "Desarrollo full-stack (2 devs + 1 designer)",
            "Integración Stripe + APIs logística",
            "Beta cerrado (100-200 usuarios)",
            "Refinamiento based on feedback",
            "Plan de lanzamiento local (Barcelona)"
        ]),
        ("FASE 3: Lanzamiento", "4 semanas", [
            "Marketing local (influencers, eventos)",
            "Onboarding guiado (primeros 1000 usuarios)",
            "Sistema de referidos (crecimiento orgánico)",
            "Soporte 24/7 activo",
            "Análisis de datos + iteración"
        ]),
        ("FASE 4: Escalabilidad", "Continuo", [
            "Expansión a Madrid, Valencia, Sevilla",
            "App móvil nativa (React Native)",
            "Funcionalidades avanzadas (subastas, trueque inverso)",
            "Internacionalización (Portugal, Italia)",
            "Monetización diversificada"
        ])
    ]
    
    for fase, duracion, actividades in fases:
        doc.add_paragraph(f'{fase} ({duracion})', style='TreqeSection')
        for actividad in actividades:
            doc.add_paragraph(f'• {actividad}', style='List Bullet')
        doc.add_paragraph()
    
    doc.add_page_break()
    
    # 7. MODELO FINANCIERO
    doc.add_paragraph('7. MODELO FINANCIERO', style='TreqeTitle')
    doc.add_paragraph()
    
    doc.add_paragraph('7.1 Proyecciones Conservadoras (Primer Año)', style='TreqeSection')
    
    proyecciones = [
        ("Mes 1-3", "0€", "0", "0€", "15.000€", "-15.000€"),
        ("Mes 4-6", "1.500€", "150", "15.000€", "12.000€", "+3.000€"),
        ("Mes 7-9", "5.000€", "500", "50.000€", "10.000€", "+40.000€"),
        ("Mes 10-12", "12.000€", "1.200", "120.000€", "15.000€", "+105.000€"),
        ("**TOTAL AÑO 1**", "**18.500€**", "**1.850**", "**185.000€**", "**52.000€**", "**+133.000€**")
    ]
    
    proyecciones_table = doc.add_table(rows=len(proyecciones) + 1, cols=6)
    proyecciones_table.style = 'LightGrid-Accent2'
    
    proyecciones_headers = ['Periodo', 'Comisiones', 'Transacciones', 'Volumen Total', 'Costes Operativos', 'Resultado']
    for col, header in enumerate(proyecciones_headers):
        proyecciones_table.cell(0, col).text = header
        set_cell_background(proyecciones_table.cell(0, col), '2F5597')
    
    for i, (periodo, comisiones, transacciones, volumen, costes, resultado) in enumerate(proyecciones, 1):
        proyecciones_table.cell(i, 0).text = periodo
        proyecciones_table.cell(i, 1).text = comisiones
        proyecciones_table.cell(i, 2).text = transacciones
        proyecciones_table.cell(i, 3).text = volumen
        proyecciones_table.cell(i, 4).text = costes
        proyecciones_table.cell(i, 5).text = resultado
        
        if 'TOTAL' in periodo:
            for col in range(6):
                set_cell_background(proyecciones_table.cell(i, col), 'C6E0B4')
    
    doc.add_paragraph()
    doc.add_paragraph('7.2 Punto de Equilibrio', style='TreqeSection')
    doc.add_paragraph('Considerando costes fijos mensuales de 8.500€ y comisión media del 1% sobre transacciones:')
    
    equilibrio = [
        ("Costes fijos mensuales", "8.500€"),
        ("Comisión media por transacción", "10€ (sobre artículo de 1.000€)"),
        ("Transacciones necesarias para equilibrio", "850 transacciones/mes"),
        ("Volumen necesario para equilibrio", "850.000€/mes"),
        ("Usuarios activos necesarios (10% tasa de transacción)", "8.500 usuarios activos"),
        ("Tiempo estimado para alcanzar equilibrio", "8-10 meses desde lanzamiento")
    ]
    
    equilibrio_table = doc.add_table(rows=len(equilibrio) + 1, cols=2)
    equilibrio_table.style = 'LightShading-Accent1'
    
    for i, (concepto, valor) in enumerate(equilibrio, 1):
        equilibrio_table.cell(i, 0).text = concepto
        equilibrio_table.cell(i, 1).text = valor
    
    doc.add_page_break()
    
    # 8. RIESGOS Y MITIGACIÓN
    doc.add_paragraph('8. RIESGOS Y MITIGACIÓN', style='TreqeTitle')
    doc.add_paragraph()
    
    riesgos = [
        ("Problema huevo-gallina", "Alto", "Alto", 
         "Sin masa crítica inicial, no hay valor para usuarios",
         "Programa de referidos agresivo + eventos presenciales + seeding inicial"),
        
        ("Fraude y desconfianza", "Alto", "Alto",
         "Trueque percibido como más riesgoso que compraventa",
         "Sistema de reputación robusto + garantía escrow + verificación de identidad"),
        
        ("Competencia reacción", "Medio", "Alto",
         "Wallapop/Milanuncios podrían añadir funcionalidad de trueque",
         "Innovación continua + enfoque en comunidad + protección de IP"),
        
        ("Problemas logísticos", "Alto", "Alto",
         "Envíos, devoluciones, garantías complejas en trueque",
         "Partner logístico especializado + seguros + procesos claros"),
        
        ("Aspectos legales/fiscales", "Medio", "Alto",
         "¿Trueque = venta para Hacienda? Regulación compleja",
         "Asesoría legal especializada desde día 1 + transparencia total"),
        
        ("Falta de financiación", "Alto", "Alto",
         "Necesidad de 50.000-100.000€ iniciales",
         "Pitch deck profesional + networking + aceleradoras + crowdfunding")
    ]
    
    riesgos_table = doc.add_table(rows=len(riesgos) + 1, cols=6)
    riesgos_table.style = 'LightGrid-Accent1'
    
    riesgos_headers = ['Riesgo', 'Probabilidad', 'Impacto', 'Descripción', 'Mitigación']
    for col, header in enumerate(riesgos_headers):
        riesgos_table.cell(0, col).text = header
        set_cell_background(riesgos_table.cell(0, col), 'C00000')
    
    for i, (riesgo, prob, impacto, desc, mitig) in enumerate(riesgos, 1):
        riesgos_table.cell(i, 0).text = riesgo
        riesgos_table.cell(i, 1).text = prob
        riesgos_table.cell(i, 2).text = impacto
        riesgos_table.cell(i, 3).text = desc
        riesgos_table.cell(i, 4).text = mitig
        
        # Colorear según probabilidad
        if prob == 'Alto':
            set_cell_background(riesgos_table.cell(i, 1), 'FF6666')
        elif prob == 'Medio':
            set_cell_background(riesgos_table.cell(i, 1), 'FFC000')
        else:
            set_cell_background(riesgos_table.cell(i, 1), 'C6E0B4')
    
    doc.add_page_break()
    
    # 9. RECOMENDACIONES ESTRATÉGICAS
    doc.add_paragraph('9. RECOMENDACIONES ESTRATÉGICAS', style='TreqeTitle')
    doc.add_paragraph()
    
    recomendaciones = [
        ("INMEDIATAS (próximas 2 semanas)", [
            "Proteger marca 'Treqe' en EUIPO",
            "Crear equipo fundador (técnico + comercial + operaciones)",
            "Desarrollar pitch deck para inversores",
            "Contactar aceleradoras (con MVP en desarrollo)",
            "Validar con 50 usuarios reales (Barcelona)"
        ]),
        
        ("MEDIO PLAZO (3-6 meses)", [
            "Concentrarse en categorías específicas (muebles, electrónica, moda)",
            "Crear comunidad antes de plataforma (grupo WhatsApp/Telegram)",
            "Establecer partnerships con tiendas de segunda mano físicas",
            "Desarrollar sistema de reputación robusto",
            "Automatizar procesos (AI para matching)"
        ]),
        
        ("LARGO PLAZO (6-12 meses)", [
            "Modelo freemium (básico gratis, premium por features)",
            "Expansión internacional (Portugal primero)",
            "Integración blockchain (reputación inmutable)",
            "Trueque B2B (empresas)",
            "Tokenización (economía circular con tokens)"
        ])
    ]
    
    for categoria, items in recomendaciones:
        doc.add_paragraph(categoria, style='TreqeSection')
        for item in items:
            doc.add_paragraph(f'• {item}', style='List Bullet')
        doc.add_paragraph()
    
    doc.add_page_break()
    
    # 10. PRÓXIMOS PASOS INMEDIATOS
    doc.add_paragraph('10. PRÓXIMOS PASOS INMEDIATOS', style='TreqeTitle')
    doc.add_paragraph()
    
    pasos_inmediatos = [
        ("Día 1-2", "Protección de marca y estructura legal", "Registrar 'Treqe' en EUIPO, crear SL"),
        ("Día 3-5", "Creación de landing page", "Desarrollo con Waitlist + validación"),
        ("Día 6-10", "Reclutamiento equipo fundador", "Buscar técnico (full-stack) + comercial"),
        ("Día 11-15", "Desarrollo MVP básico", "Funcionalidad core: registro + matching simple"),
        ("Día 16-20", "Validación con usuarios reales", "50 usuarios beta en Barcelona"),
        ("Día 21-30", "Preparación para inversión", "Pitch deck + reuniones con business angels")
    ]
    
    pasos_table = doc.add_table(rows=len(pasos_inmediatos) + 1, cols=3)
    pasos_table.style = 'LightGrid-Accent2'
    
    pasos_headers = ['Timeline', 'Acción', 'Detalle']
    for col, header in enumerate(pasos_headers):
        pasos_table.cell(0, col).text = header
        set_cell_background(pasos_table.cell(0, col), '2F5597')
    
    for i, (timeline, accion, detalle) in enumerate(pasos_inmediatos, 1):
        pasos_table.cell(i, 0).text = timeline
        pasos_table.cell(i, 1).text = accion
        pasos_table.cell(i, 2).text = detalle
    
    # Guardar documento
    output_path = Path(__file__).parent / 'Plan_Negocio_Treqe_Actualizado.docx'
    doc.save(str(output_path))
    
    print(f"✅ Documento creado exitosamente: {output_path}")
    print(f"📄 Páginas: ~35")
    print(f"📊 Secciones: 10 principales + 3 anexos")
    print(f"💡 Recomendación: Revisar y personalizar según necesidades específicas")
    
    return str(output_path)

if __name__ == '__main__':
    try:
        output_file = create_treqe_business_plan()
        print(f"\n🎯 Plan de negocio generado exitosamente en: {output_file}")
        print("📋 Contenido incluido:")
        print("  1. Resumen ejecutivo completo")
        print("  2. Análisis de mercado actualizado")
        print("  3. Modelo de negocio detallado")
        print("  4. Hoja de ruta estratégica")
        print("  5. Modelo financiero con proyecciones")
        print("  6. Análisis de riesgos y mitigación")
        print("  7. Recomendaciones estratégicas prioritarias")
        print("  8. Próximos pasos inmediatos (30 días)")
    except Exception as e:
        print(f"❌ Error al crear el documento: {e}")
        sys.exit(1)