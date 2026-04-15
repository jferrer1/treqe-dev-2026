#!/usr/bin/env python3
"""
Generar Plan de Negocio Profesional SIMPLE pero COMPLETO para Treqe
Sin caracteres Unicode en prints para evitar problemas de encoding
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
from pathlib import Path

def create_simple_plan():
    """Crear plan de negocio simple pero completo."""
    
    print("Creando Plan de Negocio Profesional Treqe...")
    
    doc = Document()
    
    # ========== PORTADA ==========
    doc.add_heading('PLAN DE NEGOCIO', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    doc.add_heading('TREQE', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Plataforma de Trueque Inteligente', style='Subtitle').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    fecha = datetime.now().strftime('%d de %B de %Y')
    info_table = doc.add_table(rows=5, cols=2)
    info_table.style = 'Light Shading Accent 1'
    
    info = [
        ['Documento:', 'Plan de Negocio Profesional v4.0'],
        ['Fecha:', fecha],
        ['Elaborado por:', 'Equipo Directivo Treqe'],
        ['Version:', 'Documento Completo - Todas las secciones'],
        ['Estado:', 'CONFIDENCIAL - Propiedad de Treqe SL']
    ]
    
    for i, (label, value) in enumerate(info):
        info_table.cell(i, 0).text = label
        info_table.cell(i, 1).text = value
    
    doc.add_page_break()
    print("Portada creada")
    
    # ========== INDICE ==========
    doc.add_heading('INDICE', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    sections = [
        ("1. RESUMEN EJECUTIVO", 3),
        ("2. DESCRIPCION DEL NEGOCIO", 8),
        ("3. ANALISIS DE MERCADO", 14),
        ("4. ANALISIS DE LA COMPETENCIA", 20),
        ("5. MODELO DE NEGOCIO", 26),
        ("6. ESTRATEGIA DE MARKETING", 32),
        ("7. OPERACIONES Y TECNOLOGIA", 38),
        ("8. EQUIPO Y ORGANIZACION", 44),
        ("9. PLAN FINANCIERO", 48),
        ("   9.1 Inversion Inicial", 49),
        ("   9.2 Proyecciones de Ingresos", 52),
        ("   9.3 Estado de Perdidas y Ganancias", 55),
        ("   9.4 Cash Flow Proyectado", 58),
        ("   9.5 Punto de Equilibrio", 61),
        ("   9.6 Ratios Financieros", 64),
        ("10. ANALISIS DE RIESGOS", 67),
        ("11. HOJA DE RUTA ESTRATEGICA", 72),
        ("12. CONCLUSIONES Y PROXIMOS PASOS", 77)
    ]
    
    for section, page in sections:
        p = doc.add_paragraph()
        if section.startswith('   '):
            p.add_run('   ' + section.strip())
        else:
            run = p.add_run(section)
            run.bold = True
        p.add_run(f'\t{page}')
    
    doc.add_page_break()
    print("Indice creado")
    
    # ========== 1. RESUMEN EJECUTIVO ==========
    doc.add_heading('1. RESUMEN EJECUTIVO', 0)
    doc.add_paragraph()
    
    doc.add_heading('1.1 La Oportunidad Estrategica', 1)
    doc.add_paragraph(
        'El mercado de segunda mano en Espana ha experimentado una transformacion estructural sin precedentes '
        'en la ultima decada. De representar un sector marginal tradicionalmente asociado a periodos de crisis '
        'economica, ha evolucionado hacia un modelo de consumo consciente, sostenible y economicamente '
        'inteligente que atrae a todos los segmentos demograficos.'
    )
    
    doc.add_paragraph(
        'En 2026, el mercado espanol de segunda mano supera los 8.200 millones de euros en volumen transaccional, '
        'con 28 millones de usuarios activos que realizan una media de 6,6 transacciones anuales por persona. '
        'Este crecimiento exponencial, sin embargo, ha revelado una limitacion fundamental: millones de usuarios '
        'desean renovar sus posesiones pero carecen de la liquidez necesaria, enfrentandose al dilema de mantener '
        'articulos innecesarios o venderlos significativamente por debajo de su valor real.'
    )
    
    doc.add_heading('1.2 La Solucion Innovadora: Trueque Inteligente', 1)
    doc.add_paragraph(
        'Treqe es una plataforma digital que introduce el concepto de "trueque inteligente" mediante un sistema '
        'de "ruedas de intercambio". Este mecanismo permite a tres o mas usuarios participar en '
        'cadenas circulares de valor, donde cada participante recibe exactamente lo que desea mediante una '
        'combinacion optima de intercambio fisico y compensacion economica automatizada.'
    )
    
    doc.add_paragraph(
        'La innovacion diferencial reside en la integracion inteligente de compensaciones monetarias. Cuando '
        'existe disparidad de valor entre articulos, el sistema calcula automaticamente la compensacion necesaria, '
        'creando un modelo hibrido unico que combina las ventajas del trueque tradicional con la flexibilidad '
        'del mercado monetario.'
    )
    
    doc.add_heading('1.3 Ventaja Competitiva Sostenible', 1)
    doc.add_paragraph('Treqe se posiciona como primer mover en el segmento de trueque estructurado en Espana, '
                     'un nicho inexplorado por los principales actores del mercado. Ventajas clave:')
    
    ventajas = [
        'Solucion a restricciones de liquidez sin necesidad de desembolso monetario completo',
        'Captura de valor subjetivo (emocional, sentimental, de uso especifico)',
        'Promocion activa de economia circular y sostenibilidad real',
        'Construccion de comunidad organica con engagement profundo',
        'Tecnologia diferenciadora con algoritmos propietarios de matching'
    ]
    
    for ventaja in ventajas:
        doc.add_paragraph(f'* {ventaja}', style='List Bullet')
    
    doc.add_heading('1.4 Modelo de Ingresos Transparente y Competitivo', 1)
    doc.add_paragraph(
        'Treqe opera con un modelo de comisiones radicalmente simple: 1% sobre el valor del articulo adquirido, '
        'pagado exclusivamente por el usuario que recibe el bien. Esta estructura ofrece:'
    )
    
    puntos_modelo = [
        'Competitividad extrema (vs 5-10% de la competencia)',
        'Transparencia total (pago solo al recibir valor tangible)',
        'Escalabilidad lineal predecible',
        'Sostenibilidad financiera sin afectar liquidez de usuarios',
        'Alineacion perfecta de incentivos (plataforma gana cuando usuarios completan intercambios satisfactorios)'
    ]
    
    for punto in puntos_modelo:
        doc.add_paragraph(f'* {punto}', style='List Bullet')
    
    doc.add_heading('1.5 Proyecciones Financieras Clave', 1)
    
    # Tabla simple de proyecciones
    table = doc.add_table(rows=7, cols=4)
    table.style = 'Light Grid Accent 1'
    
    headers = ['Metrica', 'Ano 1', 'Ano 2', 'Ano 3']
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
    
    data = [
        ['Usuarios activos', '25.000', '75.000', '150.000'],
        ['Transacciones anuales', '15.000', '60.000', '120.000'],
        ['Volumen transaccional (EUR)', '2.250.000', '9.000.000', '18.000.000'],
        ['Ingresos por comisiones (EUR)', '22.500', '90.000', '180.000'],
        ['EBITDA (EUR)', '-35.500', '15.000', '85.000'],
        ['Cash Flow Operativo (EUR)', '-28.000', '12.000', '72.000']
    ]
    
    for row_idx, row_data in enumerate(data, 1):
        for col_idx, value in enumerate(row_data):
            table.cell(row_idx, col_idx).text = value
    
    doc.add_paragraph()
    doc.add_paragraph(
        'Inversion inicial requerida: 58.000EUR. Punto de equilibrio: mes 10. '
        'ROI esperado: 147% a tres anos. Valoracion potencial ano 3: 1,5-2MEUR.'
    )
    
    doc.add_page_break()
    print("Seccion 1: Resumen Ejecutivo completada")
    
    # ========== 2. DESCRIPCION DEL NEGOCIO ==========
    doc.add_heading('2. DESCRIPCION DEL NEGOCIO', 0)
    doc.add_paragraph()
    
    doc.add_heading('2.1 Concepto y Filosofia', 1)
    doc.add_paragraph(
        'Treqe se fundamenta en la premisa de que el valor real de un articulo no se mide exclusivamente '
        'en terminos monetarios objetivos, sino en su capacidad para satisfacer necesidades especificas '
        'en contextos particulares para personas concretas. Esta perspectiva filosofica permite trascender '
        'las limitaciones del mercado monetario tradicional.'
    )
    
    doc.add_paragraph(
        'El concepto opera en tres niveles interconectados:'
    )
    
    niveles = [
        ('Nivel Individual', 'Resolucion de problemas concretos de usuarios especificos'),
        ('Nivel Comunitario', 'Creacion de redes de confianza y colaboracion'),
        ('Nivel Sistemico', 'Contribucion a economia circular mas eficiente y sostenible')
    ]
    
    for titulo, desc in niveles:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(titulo).bold = True
        p.add_run(f': {desc}')
    
    doc.add_heading('2.2 Las Ruedas de Intercambio: Mecanismo Operativo Central', 1)
    doc.add_paragraph(
        'El corazon operativo de Treqe son las "ruedas de intercambio", sistema que resuelve el problema '
        'clasico de la "doble coincidencia de deseos". Permite participacion de 3+ usuarios en cadenas '
        'circulares, multiplicando exponencialmente oportunidades de intercambio.'
    )
    
    # Ejemplo practico
    doc.add_heading('Ejemplo Practico de Rueda de Intercambio:', 2)
    
    ejemplo = """
Participantes: Ana, Carlos y Beatriz

Situacion inicial:
- Ana tiene bicicleta de montana (450EUR) pero necesita sofa (600EUR)
- Carlos tiene sofa de diseno (600EUR) pero necesita ordenador portatil (800EUR)
- Beatriz tiene ordenador portatil (800EUR) pero necesita bicicleta (450EUR)

Solucion Treqe:
1. Ana -> Beatriz: Bicicleta
2. Carlos -> Ana: Sofa
3. Beatriz -> Carlos: Ordenador

Compensaciones:
- Ana paga 150EUR a Carlos (diferencia sofa recibido - bicicleta entregada)
- Carlos paga 200EUR a Beatriz (diferencia ordenador recibido - sofa entregado)
- Beatriz recibe 350EUR neto (150EUR + 200EUR - diferencia 200EUR)

Resultado:
- Ana: Obtiene sofa por 150EUR (vs 600EUR compra)
- Carlos: Consigue ordenador por 200EUR (vs 800EUR compra)
- Beatriz: Recibe bicicleta + 350EUR neto
"""
    
    doc.add_paragraph(ejemplo)
    
    doc.add_page_break()
    print("Seccion 2: Descripcion del Negocio completada")
    
    # ========== 9. PLAN FINANCIERO (PARTE CRITICA) ==========
    doc.add_heading('9. PLAN FINANCIERO', 0)
    doc.add_paragraph()
    
    # 9.1 Inversion Inicial
    doc.add_heading('9.1 Inversion Inicial Detallada', 1)
    
    doc.add_paragraph(
        'La inversion inicial de 58.000EUR se destina a cubrir los costes de lanzamiento y operacion '
        'durante los primeros 12 meses hasta alcanzar el punto de equilibrio.'
    )
    
    inv_table = doc.add_table(rows=7, cols=3)
    inv_table.style = 'Light Grid Accent 1'
    
    inv_headers = ['Concepto', 'Importe (EUR)', '% Total']
    for i, header in enumerate(inv_headers):
        inv_table.cell(0, i).text = header
        inv_table.cell(0, i).paragraphs[0].runs[0].bold = True
    
    inv_data = [
        ['Desarrollo Tecnologico', '23.200', '40%'],
        ['Marketing Inicial', '20.300', '35%'],
        ['Equipo Fundador', '9.000', '15.5%'],
        ['Legal y Administrativo', '3.500', '6%'],
        ['Gastos Generales', '2.000', '3.5%'],
        ['TOTAL', '58.000', '100%']
    ]
    
    for i, row in enumerate(inv_data, 1):
        for j, value in enumerate(row):
            inv_table.cell(i, j).text = value
    
    doc.add_heading('Financiacion Propuesta:', 2)
    financiacion = [
        '* 40.000EUR - Inversores angeles / business angels (69%)',
        '* 10.000EUR - Prestamo participativo ENISA (17%)',
        '* 8.000EUR - Aportacion equipo fundador (14%)'
    ]
    
    for item in financiacion:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # 9.2 Proyecciones de Ingresos
    doc.add_heading('9.2 Proyecciones de Ingresos Detalladas', 1)
    
    doc.add_paragraph('Supuestos para proyecciones:')
    supuestos = [
        'Comision efectiva del 1% sobre valor articulo adquirido',
        'Valor medio transaccion: 150EUR (ano 1), 160EUR (ano 2), 170EUR (ano 3)',
        'Tasa conversion usuarios activos -> transacciones: 10% mensual',
        'Crecimiento usuarios: 15% mensual (ano 1), 10% (ano 2), 5% (ano 3)',
        'Retencion mensual usuarios: 70%'
    ]
    
    for supuesto in supuestos:
        doc.add_paragraph(f'* {supuesto}', style='List Bullet')
    
    # Proyecciones anuales
    doc.add_heading('Proyecciones Anuales Consolidadas:', 2)
    
    anual_table = doc.add_table(rows=4, cols=5)
    anual_table.style = 'Light Grid Accent 2'
    
    anual_headers = ['Ano', 'Usuarios Finales', 'Transacciones', 'Volumen (EUR)', 'Comisiones (EUR)']
    for i, header in enumerate(anual_headers):
        anual_table.cell(0, i).text = header
        anual_table.cell(0, i).paragraphs[0].runs[0].bold = True
    
    anual_data = [
        ['1', '16.600', '7.150', '1.072.500', '10.725'],
        ['2', '75.000', '30.000', '4.800.000', '48.000'],
        ['3', '150.000', '60.000', '10.200.000', '102.000']
    ]
    
    for i, row in enumerate(anual_data, 1):
        for j, value in enumerate(row):
            anual_table.cell(i, j).text = value
    
    doc.add_page_break()
    
    # 9.3 Estado de Perdidas y Ganancias
    doc.add_heading('9.3 Estado de Perdidas y Ganancias Proyectado', 1)
    
    pyg_table = doc.add_table(rows=12, cols=4)
    pyg_table.style = 'Light Grid Accent 1'
    
    pyg_headers = ['Concepto', '