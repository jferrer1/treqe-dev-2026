#!/usr/bin/env python3
"""
Corregir SOLO los anexos del documento bueno - versión simple
"""

import docx
import os
import shutil

def corregir_anexos():
    # Documento BUENO
    doc_bueno = "plan_negocio/Plan_Negocio_Treqe_100%_COMPLETO_FINAL.docx"
    doc_corregido = "plan_negocio/Plan_Negocio_Treqe_100%_COMPLETO_FINAL_ANEXOS_OK.docx"
    
    if not os.path.exists(doc_bueno):
        print(f"ERROR: No se encuentra {doc_bueno}")
        return None
    
    print("CORRIGIENDO SOLO LOS ANEXOS DEL DOCUMENTO BUENO")
    print(f"Documento: {doc_bueno}")
    print(f"Tamaño: {os.path.getsize(doc_bueno)/1024:.1f} KB")
    
    # Backup
    shutil.copy2(doc_bueno, doc_bueno + ".backup_original")
    
    # Abrir documento
    doc = docx.Document(doc_bueno)
    print(f"Párrafos: {len(doc.paragraphs)}")
    
    # Buscar ANEXOS
    anexos_idx = -1
    for i, p in enumerate(doc.paragraphs):
        if 'ANEXOS' in p.text and len(p.text) < 15:
            anexos_idx = i
            print(f"ANEXOS encontrado en párrafo {i}")
            break
    
    if anexos_idx == -1:
        print("No se encontró ANEXOS")
        return None
    
    # Eliminar anexos existentes (desde anexos_idx+1 hasta siguiente sección)
    to_remove = []
    for i in range(anexos_idx + 1, len(doc.paragraphs)):
        p = doc.paragraphs[i]
        if p.style.name.startswith('Heading') and i > anexos_idx + 1:
            break
        to_remove.append(i)
    
    print(f"Eliminando {len(to_remove)} párrafos de anexos antiguos")
    
    # Eliminar en orden inverso
    for i in reversed(to_remove):
        if i < len(doc.paragraphs):
            p = doc.paragraphs[i]
            p._element.getparent().remove(p._element)
    
    # Añadir anexos corregidos
    print("Añadiendo anexos corregidos...")
    
    # ANEXO A
    doc.add_paragraph()
    p = doc.add_paragraph("ANEXO A: ESTRUCTURA DEL EQUIPO")
    p.style = "Heading 2"
    
    doc.add_paragraph("Roles necesarios para desarrollo MVP:")
    doc.add_paragraph("1. Fundador/CEO - Estrategia, fundraising, visión")
    doc.add_paragraph("2. Desarrollador Full-Stack - Plataforma, arquitectura")
    doc.add_paragraph("3. Diseñador UX/UI - Experiencia usuario, interfaces")
    doc.add_paragraph("NOTA: CVs específicos se añadirán con miembros reales.")
    
    doc.add_page_break()
    
    # ANEXO B
    p = doc.add_paragraph("ANEXO B: METODOLOGÍA DE VALIDACIÓN")
    p.style = "Heading 2"
    
    doc.add_paragraph("Fases recomendadas:")
    doc.add_paragraph("1. Investigación secundaria (datos públicos)")
    doc.add_paragraph("2. Encuestas cuantitativas (200-500 usuarios)")
    doc.add_paragraph("3. Entrevistas cualitativas (15-20 usuarios)")
    doc.add_paragraph("4. Piloto controlado (barrio específico)")
    doc.add_paragraph("NOTA: Datos específicos tras realizar investigaciones.")
    
    doc.add_page_break()
    
    # ANEXO C
    p = doc.add_paragraph("ANEXO C: ESPECIFICACIONES TÉCNICAS")
    p.style = "Heading 2"
    
    doc.add_paragraph("Sistemas de plataforma:")
    doc.add_paragraph("• Catalogación inteligente (imágenes, etiquetado, valoración)")
    doc.add_paragraph("• Motor búsqueda y matching (algoritmo ruedas intercambio)")
    doc.add_paragraph("• Sistema intercambio (chat, acuerdos, seguimiento)")
    doc.add_paragraph("• Perfiles y reputación (verificaciones, calificaciones)")
    doc.add_paragraph("• Pagos y compensaciones (Stripe/PayPal integrados)")
    
    doc.add_page_break()
    
    # ANEXO D
    p = doc.add_paragraph("ANEXO D: GLOSARIO")
    p.style = "Heading 2"
    
    doc.add_paragraph("Términos clave:")
    doc.add_paragraph("• Rueda de Intercambio: Ciclo cerrado A→B→C→A")
    doc.add_paragraph("• Algoritmo de Matching: Sistema que encuentra combinaciones óptimas")
    doc.add_paragraph("• Compensación Monetaria: Pago para equilibrar diferencias de valor")
    doc.add_paragraph("• Liquidez de Plataforma: Facilidad para encontrar contrapartes")
    
    doc.add_page_break()
    
    # ANEXO E
    p = doc.add_paragraph("ANEXO E: FUENTES Y REFERENCIAS")
    p.style = "Heading 2"
    
    doc.add_paragraph("Fuentes verificables:")
    doc.add_paragraph("• INE (Instituto Nacional de Estadística) - ine.es")
    doc.add_paragraph("• Eurostat - ec.europa.eu/eurostat")
    doc.add_paragraph("• Banco de España - bde.es")
    doc.add_paragraph("• CNMC - cnmc.es")
    doc.add_paragraph("• Observatorio Cetelem - informes consumo")
    doc.add_paragraph("• Statista - datos mercado")
    
    # Guardar
    print(f"Guardando: {doc_corregido}")
    doc.save(doc_corregido)
    
    # Verificar
    final_size = os.path.getsize(doc_corregido)
    print(f"Tamaño final: {final_size/1024:.1f} KB")
    
    return doc_corregido

if __name__ == "__main__":
    print("=== CORRECCIÓN DE ANEXOS ===")
    print("Manteniendo documento original intacto")
    print("Solo modificando anexos según instrucciones:")
    print("1. ELIMINAR CVs inventados")
    print("2. ELIMINAR estudios inventados")
    print("3. MANTENER/MEJORAR diseños")
    print("4. CORREGIR fuentes (verificables)")
    print("5. MANTENER glosario")
    print()
    
    resultado = corregir_anexos()
    
    if resultado:
        print()
        print("=" * 60)
        print("ANEXOS CORREGIDOS EXITOSAMENTE")
        print("=" * 60)
        print(f"Documento original preservado: Plan_Negocio_Treqe_100%_COMPLETO_FINAL.docx")
        print(f"Documento con anexos corregidos: {resultado}")
        print()
        print("Se ha mantenido TODO el contenido original del documento bueno.")
        print("Solo se han reemplazado los anexos según especificaciones.")
    else:
        print("Error al corregir anexos.")