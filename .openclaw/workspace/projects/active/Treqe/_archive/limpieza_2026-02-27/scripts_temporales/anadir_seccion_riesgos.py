#!/usr/bin/env python3
"""
Añadir Sección 9: Análisis de Riesgos al documento existente
"""

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches, RGBColor

def crear_contenido_riesgos():
    """Crear contenido estructurado para la sección de riesgos."""
    
    contenido = [
        ("9. ANÁLISIS DE RIESGOS Y MITIGACIÓN", "Título Principal", True),
        ("", "Espacio", False),
        
        ("9.1 Riesgos de Mercado y Competencia", "Título Sección", True),
        ("", "Espacio", False),
        
        ("El dilema del huevo y la gallina en plataformas de intercambio", "Subtítulo", True),
        ("El riesgo más fundamental que enfrenta Treqe es el clásico problema de las plataformas de doble cara: sin vendedores no hay compradores, y sin compradores no hay vendedores. En el contexto de una plataforma de intercambio inteligente, este desafío se multiplica, ya que no solo necesitamos usuarios, sino usuarios dispuestos a participar en ruedas de intercambio complejas donde la liquidez depende de múltiples participantes simultáneamente.", "Texto Normal", False),
        ("", "Espacio", False),
        
        ("La experiencia de plataformas como Wallapop nos enseña que los primeros 10,000 usuarios son los más difíciles de conseguir, pero una vez alcanzada esa masa crítica, el crecimiento se acelera exponencialmente. El riesgo aquí no es solo la falta de usuarios iniciales, sino la posibilidad de que usuarios tempranos se desencanten por la falta de contrapartes disponibles, creando un ciclo negativo de abandono que se retroalimenta.", "Texto Normal", False),
        ("", "Espacio", False),
        
        ("La sombra de los gigantes establecidos", "Subtítulo", True),
        ("Wallapop, con sus 15 millones de usuarios activos en España, y Vinted, con su enfoque especializado en moda de segunda mano, representan competidores formidables no solo por su tamaño, sino por su capacidad de respuesta. Ambas plataformas tienen equipos de producto que monitorizan constantemente el mercado y pueden implementar características similares a las de Treqe en cuestión de meses si perciben una amenaza real.", "Texto Normal", False),
        ("", "Espacio", False),
        
        ("El riesgo no es que copien nuestra tecnología (que es relativamente sencilla de replicar), sino que utilicen su escala masiva para ofrecer funcionalidades similares de manera gratuita o con comisiones más bajas, aprovechando sus economías de escala y su base de usuarios ya establecida. Para un usuario promedio, la decisión entre usar una plataforma nueva con características innovadoras pero poca actividad, versus una plataforma establecida con millones de usuarios pero funcionalidades más básicas, suele inclinarse hacia lo segundo.", "Texto Normal", False),
        ("", "Espacio", False),
        
        ("La volatilidad del mercado de segunda mano", "Subtítulo", True),
        ("Los datos del Observatorio Cetelem muestran que el 62% de los españoles consulta aplicaciones de segunda mano al menos una vez por semana, pero solo el 19% realiza compras regulares. Esta brecha entre intención y acción representa un riesgo fundamental: ¿qué sucede si la tendencia hacia la economía circular se estanca o retrocede?", "Texto Normal", False),
        ("", "Espacio", False),
        
        ("Factores macroeconómicos como la inflación, el desempleo o cambios en la política fiscal pueden alterar drásticamente los patrones de consumo. En períodos de recesión económica profunda, paradójicamente, el mercado de segunda mano puede florecer (como sucedió durante la crisis de 2008), pero en recesiones moderadas o períodos de recuperación, los consumidores pueden volver a preferir productos nuevos como señal de estatus o confianza en la economía.", "Texto Normal", False),
        ("", "Espacio", False),
        
        ("9.2 Riesgos Operacionales y Tecnológicos", "Título Sección", True),
        ("", "Espacio", False),
        
        ("La complejidad oculta del algoritmo de emparejamiento", "Subtítulo", True),
        ("El corazón de Treqe es su algoritmo de emparejamiento inteligente que crea 'ruedas de intercambio' óptimas. Este algoritmo, aunque conceptualmente elegante, presenta riesgos operacionales significativos:", "Texto Normal", False),
        ("", "Espacio", False),
        
        ("1. Escalabilidad computacional: A medida que crece el número de usuarios y productos, el problema de encontrar emparejamientos óptimos se vuelve exponencialmente más complejo. Lo que funciona para 100 usuarios puede colapsar con 10,000.", "Lista", False),
        ("2. Latencia percibida: Los usuarios de aplicaciones móviles esperan respuestas en menos de 2 segundos. Si nuestro algoritmo tarda 10 segundos en encontrar emparejamientos óptimos, la experiencia de usuario se degrada significativamente.", "Lista", False),
        ("3. Errores de emparejamiento: Un algoritmo que sugiere intercambios desequilibrados (por ejemplo, proponiendo intercambiar un iPhone por un libro usado) erosionará rápidamente la confianza de los usuarios.", "Lista", False),
        ("", "Espacio", False),
        
        ("La infraestructura como punto único de fallo", "Subtítulo", True),
        ("Nuestra arquitectura serverless, aunque moderna y escalable, introduce dependencias críticas:", "Texto Normal", False),
        ("", "Espacio", False),
        
        ("• Proveedores de cloud: Una interrupción en AWS, Google Cloud o Azure podría dejar la plataforma completamente inoperativa.", "Lista", False),
        ("• APIs de terceros: Dependemos de servicios de pago (Stripe), mensajería (Twilio), geolocalización (Google Maps) y autenticación (OAuth). La falla de cualquiera de estos servicios afecta funcionalidades core.", "Lista", False),
        ("• Costes impredecibles: El modelo de pago por uso de serverless puede generar facturas sorpresa si hay picos de tráfico inesperados o bugs que generen llamadas repetitivas.", "Lista", False),
        ("", "Espacio", False),
        
        ("La seguridad como riesgo existencial", "Subtítulo", True),
        ("Para una plataforma que maneja transacciones económicas y datos personales, una brecha de seguridad no es solo un inconveniente técnico, es una amenaza existencial:", "Texto Normal", False),
        ("", "Espacio", False),
        
        ("1. Robo de datos personales: Nombres, direcciones, números de teléfono, preferencias de consumo.", "Lista", False),
        ("2. Fraude en transacciones: Suplantación de identidad, pagos fraudulentos, chargebacks.", "Lista", False),
        ("3. Ataques de denegación de servicio: Que podrían dejar la plataforma inaccesible durante horas o días críticos.", "Lista", False),
        ("4. Vulnerabilidades en el código: Desde inyecciones SQL hasta cross-site scripting que podrían comprometer todo el sistema.", "Lista", False),
        ("", "Espacio", False),
        
        ("9.3 Riesgos Financieros y de Liquidez", "Título Sección", True),
        ("", "Espacio", False),
        
        ("El valle de la muerte del cash flow", "Subtítulo", True),
        ("Las proyecciones financieras muestran que Treqe alcanzará el punto de equilibrio alrededor de los 3,333 transacciones mensuales, lo que según nuestro modelo ocurrirá aproximadamente en el mes 9-10 de operaciones. El riesgo financiero crítico es el período anterior a ese punto: el 'valle de la muerte' donde los gastos operativos superan consistentemente a los ingresos.", "Texto Normal", False),
        ("", "Espacio", False),
        
        ("Con una inversión inicial de €58,000, tenemos un colchón de aproximadamente 6 meses de operaciones a todo gas. Pero este colchón es más delgado de lo que parece:", "Texto Normal", False),
        ("", "Espacio", False),
        
        ("• Gastos fijos mensuales: €4,200 en servidores, APIs, marketing básico y soporte.", "Lista", False),
        ("• Gastos variables: Comisiones de procesamiento de pagos (2.9% + €0.30 por transacción), costes de adquisición de usuarios (estimados en €3-€5 por usuario activo).", "Lista", False),
        ("• Imprevistos: Desde multas regulatorias hasta costes legales por disputas entre usuarios.", "Lista", False),
        ("", "Espacio", False),
        
        ("Si el crecimiento de usuarios es más lento de lo proyectado (un escenario muy probable en plataformas sociales), podríamos agotar nuestro capital antes de alcanzar el punto de equilibrio, forzándonos a una ronda de financiación de emergencia en condiciones desfavorables o, en el peor caso, al cierre.", "Texto Normal", False),
        ("", "Espacio", False),
        
        ("9.4 Riesgos Legales y Regulatorios", "Título Sección", True),
        ("", "Espacio", False),
        
        ("El laberinto de la regulación de pagos", "Subtítulo", True),
        ("Treqe procesará pagos entre usuarios, lo que nos coloca bajo el escrutinio de múltiples reguladores:", "Texto Normal", False),
        ("", "Espacio", False),
        
        ("• Banco de España: Como entidad que facilita pagos, estamos sujetos a la Ley 10/2014 de pagos.", "Lista", False),
        ("• Ley de Servicios de Pago (PSD2): Requiere autenticación reforzada y cumplimiento de estándares de seguridad.", "Lista", False),
        ("• Prevención de blanqueo de capitales: Debemos implementar sistemas KYC (Know Your Customer) para transacciones por encima de ciertos umbrales.", "Lista", False),
        ("", "Espacio", False),
        
        ("El incumplimiento de estas regulaciones no solo conlleva multas (que pueden alcanzar el 10% de la facturación anual), sino también la posible revocación de licencias operativas, lo que sería fatal para el negocio.", "Texto Normal", False),
        ("", "Espacio", False),
        
        ("9.5 Plan de Mitigación y Contingencia", "Título Sección", True),
        ("", "Espacio", False),
        
        ("Estrategia de mitigación por categoría de riesgo", "Subtítulo", True),
        ("Para riesgos de mercado:", "Texto Normal", False),
        ("", "Espacio", False),
        
        ("1. Enfoque geográfico hiperlocal: Comenzar en un solo barrio de Barcelona (ej: Gràcia) donde podamos alcanzar densidad crítica rápidamente.", "Lista", False),
        ("2. Programa de embajadores: Reclutar usuarios influyentes en comunidades específicas que actúen como evangelizadores orgánicos.", "Lista", False),
        ("3. Garantía de liquidez inicial: En los primeros meses, el equipo de Treqe actuará como 'market maker', participando activamente en intercambios.", "Lista", False),
        ("", "Espacio", False),
        
        ("Para riesgos operacionales:", "Texto Normal", False),
        ("", "Espacio", False),
        
        ("1. Arquitectura de fallo gradual: Diseñar el sistema para que, si el algoritmo complejo falla, pueda recurrir a un algoritmo simplificado.", "Lista", False),
        ("2. Monitorización en tiempo real: Implementar dashboards que alerten sobre degradación de rendimiento antes de que los usuarios lo noten.", "Lista", False),
        ("3. Backups multi-cloud: Mantener copias de seguridad en al menos dos proveedores cloud diferentes.", "Lista", False),
        ("", "Espacio", False),
        
        ("9.6 Matriz de Riesgos Priorizados", "Título Sección", True),
        ("", "Espacio", False),
        
        ("| Riesgo | Probabilidad | Impacto | Puntuación | Estrategia Principal |", "Tabla Encabezado", False),
        ("|--------|--------------|---------|------------|---------------------|", "Tabla Línea", False),
        ("| Falta de masa crítica inicial | Alta (70%) | Crítico (9) | 6.3 | Enfoque hiperlocal + embajadores |", "Tabla Fila", False),
        ("| Competencia establecida copia features | Media (50%) | Alto (7) | 3.5 | Desarrollo rápido + patentes defensivas |", "Tabla Fila", False),
        ("| Problemas de escalabilidad algoritmo | Media (40%) | Alto (8) | 3.2 | Arquitectura de fallo gradual |", "Tabla Fila", False),
        ("| Incumplimiento regulatorio | Baja (20%) | Crítico (10) | 2.0 | Asesoría legal proactiva |", "Tabla Fila", False),
        ("| Agotamiento de capital pre-breakeven | Media (45%) | Crítico (9) | 4.1 | Presupuesto escalonado + múltiples escenarios |", "Tabla Fila", False),
        ("", "Espacio", False),
        
        ("Conclusión del análisis de riesgos", "Subtítulo", True),
        ("Los riesgos identificados para Treqe son significativos pero no insuperables. La naturaleza misma de una plataforma de intercambio innovadora implica enfrentar desafíos complejos, desde el clásico problema del huevo y la gallina hasta las intricadas regulaciones financieras.", "Texto Normal", False),
        ("", "Espacio", False),
        
        ("Lo que distingue a Treqe no es la ausencia de riesgos, sino la claridad con que los identificamos y la robustez de nuestros planes de mitigación. Hemos diseñado la empresa no como un castillo de naipes que colapsa ante la primera brisa, sino como una estructura flexible que puede adaptarse, pivotar y perseverar frente a adversidades.", "Texto Normal", False),
        ("", "Espacio", False),
        
        ("El riesgo más peligroso, curiosamente, no está en esta lista: es el riesgo de no intentarlo. En un mundo donde la economía circular se vuelve cada vez más urgente, donde las comunidades buscan nuevas formas de conectar y crear valor, y donde la tecnología permite soluciones antes imposibles, el mayor riesgo sería dejar que el miedo a lo que podría salir mal nos impida construir lo que podría salir extraordinariamente bien.", "Texto Normal", False),
    ]
    
    return contenido

def anadir_seccion_riesgos():
    """Añadir sección de riesgos al documento existente."""
    
    try:
        # Cargar documento existente
        input_path = "plan_negocio/Plan_Negocio_Treqe_COMPLETAMENTE_HUMANIZADO_FINAL.docx"
        output_path = "plan_negocio/Plan_Negocio_Treqe_CON_RIESGOS.docx"
        
        print(f"Cargando documento: {input_path}")
        doc = docx.Document(input_path)
        
        print(f"Documento cargado. Párrafos actuales: {len(doc.paragraphs)}")
        
        # Encontrar donde termina la sección 8 (Conclusiones)
        # Buscamos "8. CONCLUSIONES" o similar
        indice_insercion = len(doc.paragraphs)  # Por defecto al final
        
        for i, para in enumerate(doc.paragraphs):
            if "8." in para.text and ("CONCLUSIONES" in para.text.upper() or "Conclusiones" in para.text):
                print(f"Encontrada sección 8 en párrafo {i}: {para.text[:50]}...")
                # Buscar el final de la sección (próximo título o fin de documento)
                for j in range(i + 1, len(doc.paragraphs)):
                    if any(marker in doc.paragraphs[j].text for marker in ["9.", "ANEXO", "Anexo", "BIBLIOGRAFÍA"]):
                        indice_insercion = j
                        break
                else:
                    indice_insercion = len(doc.paragraphs)
                break
        
        print(f"Insertando nueva sección después del párrafo {indice_insercion}")
        
        # Obtener contenido de riesgos
        contenido_riesgos = crear_contenido_riesgos()
        
        # Insertar contenido
        for texto, tipo, negrita in contenido_riesgos:
            if tipo == "Espacio":
                # Añadir párrafo vacío
                doc.add_paragraph()
            elif tipo == "Título Principal":
                p = doc.add_paragraph(texto)
                p.style = "Heading 1"
                p.runs[0].bold = True
                p.runs[0].font.size = Pt(16)
            elif tipo == "Título Sección":
                p = doc.add_paragraph(texto)
                p.style = "Heading 2"
                p.runs[0].bold = True
                p.runs[0].font.size = Pt(14)
            elif tipo ==