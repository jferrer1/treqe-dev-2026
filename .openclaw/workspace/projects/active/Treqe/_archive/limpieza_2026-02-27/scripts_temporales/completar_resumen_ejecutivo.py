#!/usr/bin/env python3
"""
Completar el Resumen Ejecutivo en el documento Word existente
Agregar las secciones faltantes con redacción profesional
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def complete_executive_summary():
    """Completar el resumen ejecutivo en el documento existente."""
    
    print("Completando resumen ejecutivo en documento existente...")
    
    # Cargar documento existente
    doc_path = os.path.join(os.path.dirname(__file__), 'plan_negocio', 'Plan_Negocio_Treqe_Elaborado.docx')
    doc = Document(doc_path)
    
    # Encontrar dónde termina el documento actual
    # Vamos a agregar después de la sección 1.2
    
    # ========== CONTINUAR DESPUÉS DE LA SECCIÓN 1.2 ==========
    
    # 1.3 El Panorama Competitivo Actual
    doc.add_heading('1.3 El Panorama Competitivo Actual', 1)
    
    competitivo_text = """El mercado español de segunda mano presenta una estructura competitiva consolidada, con actores claramente definidos que han establecido posiciones diferenciadas a lo largo de los últimos años. Esta consolidación no implica estancamiento, sino más bien una especialización creciente y una profundización en segmentos específicos.

**Wallapop: El Líder Indiscutible del Mercado Generalista**
Con aproximadamente 15 millones de usuarios en España, Wallapop ha logrado establecer lo que en términos de negocio digital se denomina un "efecto de red" masivo. Su modelo se fundamenta en la amplitud: cualquier categoría de producto encuentra cabida en su plataforma, desde muebles hasta electrónica, pasando por vehículos y servicios. 

La fortaleza principal de Wallapop reside precisamente en esta amplitud. Un usuario que busca cualquier tipo de artículo de segunda mano tiene altas probabilidades de encontrarlo en su plataforma. Sin embargo, esta ventaja conlleva también limitaciones. El enfoque en volumen de transacciones monetarias ha llevado a una estructura de comisiones que muchos usuarios consideran elevada (5% más 0,90€ fijo por venta), y la experiencia de usuario puede resultar impersonal en transacciones de cierto valor.

**Vinted: La Especialista en Moda que Creó su Propio Ecosistema**
Si Wallapop es el generalista, Vinted es la especialista. Con 4,5 millones de usuarios activos en España, ha conseguido crear no solo una plataforma de compraventa, sino toda una comunidad en torno a la moda de segunda mano. Su catálogo se concentra en un 95% en ropa, calzado y complementos, lo que le permite ofrecer funcionalidades específicas para este segmento.

El sistema de reputación de Vinted es particularmente sofisticado, con evaluaciones detalladas que van más allá del simple "positivo/negativo". Los usuarios pueden valorar aspectos como la precisión en las descripciones, la calidad del embalaje, o la rapidez en la comunicación. Esta especialización tiene un coste: las comisiones totales pueden alcanzar el 8-9% del valor de la transacción, una cifra significativa para artículos de moda que suelen tener valores moderados.

**Facebook Marketplace: El Gigante con Potencial No Realizado**
La integración de Marketplace dentro de la red social Facebook le otorga un potencial teórico enorme: aproximadamente 20 millones de usuarios españoles podrían acceder a la plataforma de forma inmediata. Su principal ventaja es la gratuidad para transacciones entre particulares, lo que la hace atractiva para ventas esporádicas.

Sin embargo, esta integración con la red social también presenta limitaciones. La experiencia de usuario es notablemente más básica que la de plataformas especializadas, los sistemas de seguridad y garantía son menos robustos, y la mezcla entre contenido social y transaccional puede resultar incómoda para algunos usuarios. Marketplace funciona mejor como complemento para ventas ocasionales que como plataforma principal para usuarios frecuentes.

**Milanuncios: El Tradicional que Encuentra su Nicho**
Con aproximadamente el 10% de cuota de mercado, Milanuncios representa la continuidad de un modelo anterior a la explosión de las apps móviles. Su base de usuarios tiende a ser de mayor edad y menos digitalizada que la de otras plataformas, lo que se traduce en transacciones de mayor valor medio pero menor frecuencia.

La ventaja de Milanuncios reside en el reconocimiento de marca histórico y en la confianza que genera en segmentos poblacionales que valoran la estabilidad por encima de la innovación. Su interfaz, menos modernizada que la de sus competidores, resulta familiar para usuarios que llevan años utilizando el servicio.

**El Espacio Vacío: Trueque Estructurado y Escalable**
Lo que resulta evidente al analizar este panorama competitivo es que todas las plataformas existentes se centran exclusivamente en el modelo de compraventa monetaria. Ninguna ha desarrollado soluciones robustas para el trueque, a pesar de que numerosos estudios de mercado indican una demanda significativa por este tipo de intercambios.

Esta omisión no es casual. Los modelos de negocio de las plataformas establecidas se fundamentan en el volumen de transacciones monetarias. Desarrollar sistemas de trueque complejos reduciría potencialmente su facturación, al permitir a los usuarios satisfacer necesidades sin transacciones monetarias. Además, la complejidad técnica requerida para implementar sistemas de matching eficientes es considerable y no se alinea necesariamente con sus roadmaps de desarrollo actuales.

Es precisamente en este espacio vacío donde Treqe encuentra su oportunidad estratégica. No compitiendo frontalmente con los gigantes establecidos, sino ocupando un nicho que ellos han elegido no atender: el trueque estructurado, escalable y tecnológicamente sofisticado."""
    
    doc.add_paragraph(competitivo_text)
    
    # 1.4 Tendencias Emergentes que Definen el Futuro
    doc.add_heading('1.4 Tendencias Emergentes que Definen el Futuro', 1)
    
    tendencias_text = """El análisis de la evolución reciente del mercado permite identificar cinco tendencias clave que están reconfigurando el sector y que, muy probablemente, definirán su desarrollo en los próximos años. Comprender estas tendencias no es solo un ejercicio académico, sino una necesidad estratégica para cualquier actor que pretenda competir en este espacio.

**1. La Premiumización Acelerada: Cuando la Segunda Mano deja de ser "de Segunda"**
Un cambio fundamental que estamos observando es la creciente demanda de artículos de segunda mano de alta calidad. Ya no se trata exclusivamente de productos básicos o de bajo coste. El segmento de artículos premium crece a un ritmo del 125% interanual (datos 2023-2025), superando con creces el crecimiento del mercado general.

Marcas de lujo, productos tecnológicos de gama alta, muebles de diseño, instrumentos musicales profesionales... todos encuentran un mercado ávido en plataformas de segunda mano. Los compradores en este segmento no buscan principalmente ahorro económico (aunque lo valoran), sino acceso a productos de calidad que de otra forma estarían fuera de su alcance, ya sea por precio o por disponibilidad.

**2. La Sostenibilidad como Driver Principal, no como Accesorio**
El 68% de los usuarios actuales menciona explícitamente la motivación ecológica como factor determinante en su participación en el mercado de segunda mano. Esta cifra, que ha crecido consistentemente en los últimos cinco años, indica un cambio profundo en la mentalidad del consumidor.

La economía circular ha dejado de ser un concepto teórico discutido en foros especializados para convertirse en una práctica cotidiana para millones de españoles. Los usuarios no solo buscan ahorrar dinero; buscan conscientemente reducir su huella ambiental, extender la vida útil de los productos, y contribuir a un modelo de consumo más responsable.

**3. La Paradoja de la Hiperlocalización en un Mundo Digital Global**
Contrariamente a lo que podría pensarse en un mundo digital aparentemente globalizado, las transacciones más exitosas y con mayor índice de satisfacción ocurren en radios geográficos reducidos. Los datos indican que los intercambios dentro de un radio de 50km tienen una tasa de finalización un 40% superior a los de mayor distancia.

Esta tendencia responde a múltiples factores: menores costes logísticos, mayor facilidad para la inspección previa de los artículos, mayor rapidez en la entrega, y, quizás lo más importante, mayor confianza entre usuarios que comparten un contexto geográfico común. Las plataformas que logren potenciar estas comunidades locales sin perder la escala global tendrán una ventaja competitiva significativa.

**4. La Regulación Emergente: De la Informalidad a la Profesionalización**
Las autoridades fiscales y regulatorias están desarrollando marcos específicos para las ventas entre particulares. A partir de 2025, nuevas normativas establecen umbrales de declaración y obligaciones formales para vendedores frecuentes, lo que está profesionalizando progresivamente el sector.

Este proceso de regulación, aunque inicialmente pueda percibirse como una carga burocrática, representa en realidad una oportunidad para la consolidación del sector. Establece reglas claras, reduce la incertidumbre jurídica, y diferencia a los actores serios de los ocasionales. Las plataformas que integren estas consideraciones regulatorias desde el diseño tendrán una ventaja operativa importante.

**5. La Experiencia Mobile-First Absoluta: No es una Opción, es el Standard**
La generación que ha crecido con smartphones en la mano (millennials y generación Z) representa actualmente el 75% del volumen transaccional en plataformas de segunda mano. Su expectativa es una experiencia perfectamente optimizada para móvil, con procesos que no requieran transición a ordenador en ningún punto.

Esto va más allá del simple diseño responsive. Implica repensar completamente los flujos de usuario, los tiempos de carga, los métodos de pago, y la gestión de notificaciones. Una plataforma que ofrezca una experiencia móvil inferior, por pequeña que sea la diferencia, perderá progresivamente usuarios frente a competidores que prioricen este aspecto.

**Implicaciones para Treqe**
Estas tendencias no son observaciones abstractas, sino guías concretas para el desarrollo de Treqe. La plataforma se diseña específicamente para:
- Facilitar el intercambio de artículos premium mediante sistemas de valoración robustos
- Incorporar la sostenibilidad como elemento central de la propuesta de valor
- Potenciar las comunidades locales mediante algoritmos de matching geográficamente inteligentes
- Integrar desde el inicio las consideraciones regulatorias emergentes
- Ofrecer una experiencia mobile-first que supere las expectativas del usuario contemporáneo

La convergencia de estas tendencias crea un contexto particularmente favorable para la introducción de Treqe. No se trata de competir en un mercado saturado, sino de ofrecer una solución específicamente adaptada a las demandas emergentes del consumidor actual."""
    
    doc.add_paragraph(tendencias_text)
    
    doc.add_page_break()
    
    # ========== SECCIÓN 2: EL PROBLEMA NO RESUELTO ==========
    
    doc.add_heading('2. EL PROBLEMA NO RESUELTO: LA PARADOJA DE LA LIQUIDEZ', 0)
    
    # 2.1 La Situación del Usuario Contemporáneo
    doc.add_heading('2.1 La Situación del Usuario Contemporáneo', 1)
    
    situacion_text = """Para comprender la oportunidad que representa Treqe, es necesario analizar en profundidad la situación real del usuario contemporáneo del mercado de segunda mano. Esta situación, que afecta a millones de españoles, constituye lo que denominamos "la paradoja de la liquidez": tener valor atrapado en posesiones que ya no se desean, mientras se carece del capital necesario para adquirir lo que realmente se necesita.

Imaginemos por un momento a un usuario típico. Podría ser Ana, una arquitecta de 32 años que vive en Barcelona. En su apartamento de 70m², Ana acumula varios objetos que ya no utiliza pero que conservan tanto valor económico como sentimental:

- Una bicicleta de montaña Specialized Rockhopper que compró durante su etapa universitaria, cuando practicaba ciclismo con regularidad. La bicicleta está en excelente estado, con mantenimiento regular, pero ahora Ana trabaja largas horas y rara vez tiene tiempo para salir al monte.
- Un sofá de tres plazas que heredó de sus abuelos. Es un mueble sólido, bien construido, pero su estilo tradicional no encaja con la decoración contemporánea que Ana ha elegido para su hogar.
- Una colección de libros especializados en arquitectura moderna que consultó durante sus estudios de posgrado. Los libros son valiosos como referencia, pero ocupan un estantería completa que Ana preferiría dedicar a otros usos.

Ana desea renovar su espacio. Sus necesidades actuales son diferentes:
- Necesita un escritorio ergonómico para teletrabajar cómodamente varias horas al día
- Requiere unas estanterías modulares que le permitan organizar mejor sus pertenencias en un espacio limitado
- Quisiera cambiar ese sofá heredado por uno más moderno y acorde con su estilo actual
- Le gustaría disponer de una mesa plegable para poder recibir a amigos en su pequeño salón

El problema de Ana no es de deseo, sino de recursos. Aunque los objetos que posee tienen un valor económico significativo (estimado en aproximadamente 1.500€ en total), carece de la liquidez necesaria para realizar las compras que desea (estimadas en 2.000€). Su salario como arquitecta junior, combinado con los costes de vida en Barcelona, no le permiten realizar este desembolso sin afectar significativamente sus finanzas personales.

Esta situación no es excepcional. Según estudios recientes, el 63% de los españoles entre 25 y 45 años declara tener en su hogar al menos tres artículos que ya no utilizan pero que conservan valor económico. La media de valor "atrapado" en estos artículos se estima en 1.200€ por hogar. Multiplicado por los aproximadamente 8 millones de hogares en esta franja demográfica, obtenemos un volumen económico de casi 10.000 millones de euros en valor no realizado.

La paradoja es evidente: hay valor, hay necesidad, pero falta el mecanismo eficiente para convertir uno en otro. Los sistemas actuales ofrecen soluciones imperfectas que, como veremos a continuación, dejan insatisfecha a una proporción significativa de usuarios."""
    
    doc.add_paragraph(situacion_text)
    
    # 2.2 Las Opciones No Óptimas Disponibles
    doc.add_heading('2.2 Las Opciones No Óptimas Disponibles', 1)
    
    opciones_text = """Frente a la situación descrita, Ana (y millones de usuarios en circunstancias similares) se enfrenta a lo que en economía se denomina un "trilema": tres opciones disponibles, cada una con sus propios costes y limitaciones, sin que ninguna ofrezca una solución satisfactoria.

**Opción A: Mantener los Objetos Innecesarios**
Esta es, estadísticamente, la opción más común. El 58% de los usuarios en la situación de Ana elige simplemente conservar los objetos que ya no utiliza. Las consecuencias de esta elección son múltiples:

- **Ocupación de espacio valioso:** En ciudades como Barcelona, donde el metro cuadrado tiene un coste elevadísimo (entre 15€ y 20€ mensuales por m² de alquiler), dedicar espacio a objetos no utilizados representa un coste de oportunidad significativo. Los 2m² que ocupa el sofá heredado podrían tener un valor de uso muy superior.
  
- **Depreciación continua:** Los objetos, especialmente los tecnológicos y los de moda, pierden valor con el tiempo. Una bicicleta que hoy vale 450€ podría valer 300€ dentro de dos años. Cada mes que pasa sin utilizarla o sin venderla representa una pérdida económica real, aunque no inmediatamente visible.
  
- **Coste psicológico:** Convivir diariamente con objetos que no se desean pero que no se pueden cambiar genera una insatisfacción sorda pero constante. Es lo que algunos psicólogos denominan "carga cognitiva ambiental": el entorno físico envía mensajes contradictorios sobre lo que somos y lo que queremos ser.
  
- **Inercia acumulativa:** Cuanto más tiempo pasa, más difícil resulta tomar la decisión de cambiar. Los objetos se integran en el paisaje doméstico, se normaliza su presencia, y la urgencia de la renovación disminuye, aunque la insatisfacción subyacente persista.

**Opción B: Vender por Debajo del Valor Real**
La segunda opción disponible es vender los objetos, pero esta solución presenta sus propias limitaciones:

- **La realidad del mercado de segunda mano:** Para vender con rapidez en plataformas como Wallapop o Vinted, es necesario fijar un precio