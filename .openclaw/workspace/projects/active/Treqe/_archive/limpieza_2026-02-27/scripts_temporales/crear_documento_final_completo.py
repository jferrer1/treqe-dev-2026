#!/usr/bin/env python3
"""
Crear documento final completo con todas las secciones
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from datetime import datetime

def crear_documento_completo():
    """Crear documento completo con todas las secciones."""
    
    print("Creando documento completo con todas las secciones...")
    
    # Crear nuevo documento
    doc = Document()
    
    # Configurar márgenes
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # ========== PORTADA ==========
    title = doc.add_heading('PLAN DE NEGOCIO PROFESIONAL', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    company = doc.add_heading('TREQE', 0)
    company.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Plataforma de Trueque Inteligente')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    fecha = datetime.now().strftime('%d de %B de %Y')
    date_para = doc.add_paragraph(f'Fecha: {fecha}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    version_para = doc.add_paragraph('Versión: 3.0 - Documento Completo')
    version_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    status_para = doc.add_paragraph('Estado: CONFIDENCIAL')
    status_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # ========== ÍNDICE ==========
    doc.add_heading('ÍNDICE', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    sections = [
        ("1. INTRODUCCIÓN", 3),
        ("   1.1 Transformación del Sector", 3),
        ("   1.2 Datos Cuantitativos 2025-2026", 4),
        ("   1.3 Panorama Competitivo", 6),
        ("   1.4 Tendencias Emergentes", 8),
        ("2. PROBLEMA NO RESUELTO", 10),
        ("   2.1 Situación Usuario Contemporáneo", 10),
        ("   2.2 Opciones No Óptimas", 11),
        ("   2.3 Limitación Matemática", 12),
        ("   2.4 Oportunidad Cuantificada", 13),
        ("3. SOLUCIÓN INNOVADORA", 15),
        ("   3.1 Concepto Fundamental", 15),
        ("   3.2 Mecanismo Operativo", 16),
        ("   3.3 Ejemplo Práctico", 19),
        ("4. VENTAJA COMPETITIVA", 22),
        ("   4.1 Posicionamiento Único", 22),
        ("   4.2 Ventajas Tecnológicas", 24),
        ("   4.3 Ventajas Económicas", 26),
        ("   4.4 Ventajas Sostenibilidad", 28),
        ("   4.5 Barreras Entrada", 30),
        ("5. MODELO DE NEGOCIO", 32),
        ("   5.1 Flujos Ingresos", 32),
        ("   5.2 Inversión Inicial", 34),
        ("   5.3 Financiación", 36),
        ("6. PROYECCIONES FINANCIERAS", 38),
        ("   6.1 Supuestos Clave", 38),
        ("   6.2 Proyecciones Anuales", 40),
        ("   6.3 Estado Pérdidas y Ganancias", 42),
        ("   6.4 Cash Flow", 44),
        ("   6.5 Ratios Financieros", 46),
        ("7. EQUIPO Y EJECUCIÓN", 48),
        ("   7.1 Equipo Fundador", 48),
        ("   7.2 Plan por Fases", 50),
        ("8. CONCLUSIONES", 52)
    ]
    
    for section, page in sections:
        p = doc.add_paragraph()
        if section.startswith('   '):
            p.add_run(section)
        else:
            p.add_run(section).bold = True
        p.add_run(f'\t{page}')
    
    doc.add_page_break()
    
    # ========== SECCIÓN 1: INTRODUCCIÓN ==========
    doc.add_heading('1. INTRODUCCIÓN', 0)
    
    # 1.1
    doc.add_heading('1.1 Transformación del Sector', 1)
    doc.add_paragraph('El mercado de segunda mano en España ha experimentado una transformación estructural sin precedentes en la última década. De representar un sector marginal tradicionalmente asociado a periodos de crisis económica, ha evolucionado hacia un modelo de consumo consciente, sostenible y económicamente inteligente que atrae a todos los segmentos demográficos.')
    
    # 1.2
    doc.add_heading('1.2 Datos Cuantitativos 2025-2026', 1)
    doc.add_paragraph('• Mercado total España: 8.200M€ (+42% vs 2020)')
    doc.add_paragraph('• Usuarios activos: 28 millones (47% población)')
    doc.add_paragraph('• Gasto medio anual: 1.850€ (+142% vs 2016)')
    doc.add_paragraph('• Penetración móvil: 94% usa apps (mobile-first)')
    doc.add_paragraph('• Segmento lujo: +125% crecimiento (2023-2025)')
    
    # 1.3
    doc.add_heading('1.3 Panorama Competitivo', 1)
    doc.add_paragraph('• Wallapop: 15M usuarios, 5% comisión + 0,90€ fijo')
    doc.add_paragraph('• Vinted: 4,5M usuarios, especializada moda, comisiones 8-9%')
    doc.add_paragraph('• Facebook Marketplace: Potencial 20M usuarios, gratuita pero experiencia básica')
    doc.add_paragraph('• Milanuncios: 10% cuota, usuarios menos digitalizados')
    doc.add_paragraph('• Espacio vacío: Trueque estructurado no atendido')
    
    # 1.4
    doc.add_heading('1.4 Tendencias Emergentes', 1)
    doc.add_paragraph('1. Premiumización acelerada (+125% crecimiento artículos premium)')
    doc.add_paragraph('2. Sostenibilidad como driver principal (68% usuarios menciona motivación ecológica)')
    doc.add_paragraph('3. Importancia comunidades locales (transacciones <50km más exitosas)')
    doc.add_paragraph('4. Regulación emergente (normativas fiscales 2025+)')
    doc.add_paragraph('5. Experiencia mobile-first absoluta (75% volumen millennials/gen Z)')
    
    doc.add_page_break()
    
    # ========== SECCIÓN 2: PROBLEMA ==========
    doc.add_heading('2. PROBLEMA NO RESUELTO', 0)
    
    # 2.1
    doc.add_heading('2.1 Situación Usuario Contemporáneo', 1)
    doc.add_paragraph('Millones de usuarios españoles enfrentan la "paradoja de la liquidez": tener valor atrapado en posesiones que ya no desean, mientras carecen del capital necesario para adquirir lo que realmente necesitan.')
    doc.add_paragraph('Ejemplo: Ana, arquitecta 32 años Barcelona')
    doc.add_paragraph('• Tiene: Bicicleta (450€), sofá (600€), libros (450€) - Total: 1.500€')
    doc.add_paragraph('• Necesita: Escritorio, estanterías, sofá nuevo - Total: 2.000€')
    doc.add_paragraph('• Problema: Aunque tiene valor, carece de liquidez')
    
    # 2.2
    doc.add_heading('2.2 Opciones No Óptimas', 1)
    doc.add_paragraph('Opción A: Mantener objetos innecesarios (58% elige)')
    doc.add_paragraph('• Ocupación espacio valioso')
    doc.add_paragraph('• Depreciación continua')
    doc.add_paragraph('• Coste psicológico')
    doc.add_paragraph('• Inercia acumulativa')
    
    doc.add_paragraph('Opción B: Vender por debajo valor real')
    doc.add_paragraph('• Pérdida 30-50% del valor')
    doc.add_paragraph('• Frustración por "regalar"')
    doc.add_paragraph('• Efecto acumulativo negativo')
    
    doc.add_paragraph('Opción C: Postergar renovación')
    doc.add_paragraph('• Impacto calidad de vida')
    doc.add_paragraph('• Obsolescencia funcional')
    doc.add_paragraph('• Oportunidades perdidas')
    
    # 2.3
    doc.add_heading('2.3 Limitación Matemática', 1)
    doc.add_paragraph('Doble coincidencia de deseos: Para que dos personas intercambien directamente, ambas deben querer exactamente lo que la otra ofrece.')
    doc.add_paragraph('Estadísticas:')
    doc.add_paragraph('• Tasa éxito intercambios directos: <5%')
    doc.add_paragraph('• Tiempo medio búsqueda: 2-3 meses')
    doc.add_paragraph('• Abandono por frustración: 45% después 1 mes')
    
    # 2.4
    doc.add_heading('2.4 Oportunidad Cuantificada', 1)
    doc.add_paragraph('• Usuarios que preferirían intercambiar: 8 millones de españoles')
    doc.add_paragraph('• Frecuencia deseada renovación: Cada 2-3 años (vs 5-7 actual)')
    doc.add_paragraph('• Volumen económico artículos "atrapados": 15.000M€')
    doc.add_paragraph('• Brecha satisfacción: 73% usuarios frustrados por no poder intercambiar')
    
    doc.add_page_break()
    
    # ========== SECCIÓN 3: SOLUCIÓN ==========
    doc.add_heading('3. SOLUCIÓN INNOVADORA', 0)
    
    # 3.1
    doc.add_heading('3.1 Concepto Fundamental', 1)
    doc.add_paragraph('Treqe introduce "ruedas de intercambio" que permiten participación de 3-5 usuarios en cadenas circulares de valor, resolviendo matemáticamente el problema de doble coincidencia.')
    
    # 3.2
    doc.add_heading('3.2 Mecanismo Operativo', 1)
    doc.add_paragraph('Paso 1: Registro Preferencias')
    doc.add_paragraph('• Usuarios indican "Tengo" + "Quiero"')
    doc.add_paragraph('• Especifican valor, condiciones, fotos')
    
    doc.add_paragraph('Paso 2: Algoritmo Matching')
    doc.add_paragraph('• Construye grafo dirigido')
    doc.add_paragraph('• Busca ciclos 3-5 nodos (DFS optimizado)')
    doc.add_paragraph('• Timeout: 500ms por búsqueda')
    
    doc.add_paragraph('Paso 3: Optimización Económica')
    doc.add_paragraph('• Calcula compensaciones óptimas (algoritmo PuLP)')
    doc.add_paragraph('• Minimiza transferencias totales')
    doc.add_paragraph('• Maximiza satisfacción global')
    
    doc.add_paragraph('Paso 4: Negociación Facilitada')
    doc.add_paragraph('• Chat grupal tiempo real (WebSockets)')
    doc.add_paragraph('• Sistema sugiere términos')
    doc.add_paragraph('• Acuerdo requiere confirmación todos')
    
    doc.add_paragraph('Paso 5: Ejecución Segura')
    doc.add_paragraph('• Pagos vía Stripe Connect (escrow)')
    doc.add_paragraph('• Logística integrada APIs Correos/SEUR')
    doc.add_paragraph('• Sistema reputación + garantías')
    
    # 3.3
    doc.add_heading('3.3 Ejemplo Práctico', 1)
    doc.add_paragraph('Situación:')
    doc.add_paragraph('• Ana: Bicicleta (450€) → necesita sofá (600€)')
    doc.add_paragraph('• Carlos: Sofá (600€) → necesita ordenador (800€)')
    doc.add_paragraph('• Beatriz: Ordenador (800€) → necesita bicicleta (450€)')
    
    doc.add_paragraph('Solución:')
    doc.add_paragraph('1. Ana → Beatriz: Bicicleta')
    doc.add_paragraph('2. Carlos → Ana: Sofá')
    doc.add_paragraph('3. Beatriz → Carlos: Ordenador')
    
    doc.add_paragraph('Compensaciones:')
    doc.add_paragraph('• Ana paga 150€ a Carlos')
    doc.add_paragraph('• Carlos paga 200€ a Beatriz')
    doc.add_paragraph('• Beatriz recibe 350€ neto')
    
    doc.add_paragraph('Resultados:')
    doc.add_paragraph('• Ana: Sofá 600€ por 150€ → Ahorro 450€ (75%)')
    doc.add_paragraph('• Carlos: Ordenador 800€ por 200€ → Ahorro 600€ (75%)')
    doc.add_paragraph('• Beatriz: Bicicleta 450€ + 350€ → Valor total 800€')
    
    doc.add_page_break()
    
    # ========== SECCIÓN 4: VENTAJA COMPETITIVA ==========
    doc.add_heading('4. VENTAJA COMPETITIVA', 0)
    
    # 4.1
    doc.add_heading('4.1 Posicionamiento Único', 1)
    doc.add_paragraph('Primer mover en trueque estructurado en España - nicho inexplorado por competencia.')
    
    # 4.2
    doc.add_heading('4.2 Ventajas Tecnológicas', 1)
    doc.add_paragraph('Algoritmos:')
    doc.add_paragraph('• Matching teoría grafos (NetworkX)')
    doc.add_paragraph('• Optimización lineal (PuLP)')
    doc.add_paragraph('• Reputación machine learning')
    doc.add_paragraph('• Detección fraudes')
    
    doc.add_paragraph('Arquitectura:')
    doc.add_paragraph('• Frontend: Next.js 14 + React 19 + TypeScript + PWA')
    doc.add_paragraph('• Backend: Node.js + Express + WebSockets')
    doc.add_paragraph('• Matching: Python microservicio + Redis')
    doc.add_paragraph('• Base datos: PostgreSQL + TimescaleDB')
    doc.add_paragraph('• Infra: Serverless (Vercel + Railway)')
    
    # 4.3
    doc.add_heading('4.3 Ventajas Económicas', 1)
    doc.add_paragraph('Comisiones:')
    doc.add_paragraph('• Treqe: 1% (ningún coste adicional)')
    doc.add_paragraph('• Wallapop: 5% + 0,90€ fijo')
    doc.add_paragraph('• Vinted: 5% + 0,70€ + 3% protección')
    doc.add_paragraph('• Ventaja Treqe: 80-90% más barato')
    
    # 4.4
    doc.add_heading('4.4 Ventajas Sostenibilidad', 1)
    doc.add_paragraph('Impacto medioambiental:')
    doc.add_paragraph('• Extensión vida útil: +3-5 años')
    doc.add_paragraph('• Reducción residuos: ~150kg CO2eq/transacción')
    doc.add_paragraph('• Economía circular real')
    
    doc.add_paragraph('Impacto social:')
    doc.add_paragraph('• Acceso bienes sin liquidez')
    doc.add_paragraph('• Reducción desigualdades')
    doc.add_paragraph('• Construcción comunidades')
    
    # 4.5
    doc.add_heading('4.5 Barreras Entrada', 1)
    doc.add_paragraph('1. Complejidad algorítmica: 6-9 meses desarrollo')
    doc.add_paragraph('2. Efecto red local: Masa crítica comunidades')
    doc.add_paragraph('3. Base datos preferencias: Activo intangible')
    
    doc.add_page_break()
    
    # ========== SECCIÓN 5: MODELO DE NEGOCIO ==========
    doc.add_heading('5. MODELO DE NEGOCIO', 0)
    
    # 5.1
    doc.add_heading('5.1 Flujos Ingresos', 1)
    doc.add_paragraph('Fase 1 (Año 1): Comisión 1%')
    doc.add_paragraph('• Sobre valor artículo adquirido')
    doc.add_paragraph('• Pagada por receptor')
    doc.add_paragraph('• Ejemplo: 500€ → 5€')
    
    doc.add_paragraph('Fase