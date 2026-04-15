#!/usr/bin/env python3
"""
Añadir Sección 9: Análisis de Riesgos al documento existente - Versión completa
"""

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches, RGBColor

def crear_contenido_riesgos():
    """Crear contenido estructurado para la sección de riesgos."""
    
    contenido = [
        ("9. ANÁLISIS DE RIESGOS Y MITIGACIÓN", "Título Principal", True),
        ("", "Espacio", False),
        
        ("9.1 Riesgos de Mercado y Competencia", "Título Sección", True),
        ("", "Espacio", False),
        
        ("El dilema del huevo y la gallina en plataformas de intercambio", "Subtítulo", True),
        ("El riesgo más fundamental que enfrenta Treqe es el clásico problema de las plataformas de doble cara: sin vendedores no hay compradores, y sin compradores no hay vendedores. En el contexto de una plataforma de intercambio inteligente, este desafío se multiplica, ya que no solo necesitamos usuarios, sino usuarios dispuestos a participar en ruedas de intercambio complejas donde la liquidez depende de múltiples participantes simultáneamente.", "Texto Normal", False),
        ("", "Espacio", False),
        
        ("La experiencia de plataformas como Wallapop nos enseña que los primeros 10,000 usuarios son los más difíciles de conseguir, pero una vez alcanzada esa masa crítica, el crecimiento se acelera exponencialmente. El riesgo aquí no es solo la falta de usuarios iniciales, sino la posibilidad de que usuarios tempranos se desencanten por la falta de contrapartes disponibles, creando un ciclo negativo de abandono que se retroalimenta.", "Texto Normal", False),
        ("", "Espacio", False),
        
        ("La sombra de los gigantes establecidos", "Subtítulo", True),
        ("Wallapop, con sus 15 millones de usuarios activos en España, y Vinted, con su enfoque especializado en moda de segunda mano, representan competidores formidables no solo por su tamaño, sino por su capacidad de respuesta. Ambas plataformas tienen equipos de producto que monitorizan constantemente el mercado y pueden implementar características similares a las de Treqe en cuestión de meses si perciben una amenaza real.", "Texto Normal", False),
        ("", "Espacio", False),
        
        ("El riesgo no es que copien nuestra tecnología (que es relativamente sencilla de replicar), sino que utilicen su escala masiva para ofrecer funcionalidades similares de manera gratuita o con comisiones más bajas, aprovechando sus economías de escala y su base de usuarios ya establecida. Para un usuario promedio, la decisión entre usar una plataforma nueva con características innovadoras pero poca actividad, versus una plataforma establecida con millones de usuarios pero funcionalidades más básicas, suele inclinarse hacia lo segundo.", "Texto Normal", False),
        ("", "Espacio", False),
        
        ("La volatilidad del mercado de segunda mano", "Subtítulo", True),
        ("Los datos del Observatorio Cetelem muestran que el 62% de los españoles consulta aplicaciones de segunda mano al menos una vez por semana, pero solo el 19% realiza compras regulares. Esta brecha entre intención y acción representa un riesgo fundamental: ¿qué sucede si la tendencia hacia la economía circular se estanca o retrocede?", "Texto Normal", False),
        ("", "Espacio", False),
        
        ("Factores macroeconómicos como la inflación, el desempleo o cambios en la política fiscal pueden alterar drásticamente los patrones de consumo. En períodos de recesión económica profunda, paradójicamente, el mercado de segunda mano puede florecer (como sucedió durante la crisis de 2008), pero en recesiones moderadas o períodos de recuperación, los consumidores pueden volver a preferir productos nuevos como señal de estatus o confianza en la economía.", "Texto Normal", False),
        ("", "Espacio", False),
        
        ("9.2 Riesgos Operacionales y Tecnológicos", "Título Sección", True),
        ("", "Espacio", False),
        
        ("La complejidad oculta del algoritmo de emparejamiento", "Subtítulo", True),
        ("El corazón de Treqe es su algoritmo de emparejamiento inteligente que crea 'ruedas de intercambio' óptimas. Este algoritmo, aunque conceptualmente elegante, presenta riesgos operacionales significativos:", "Texto Normal", False),
        ("", "Espacio", False),
        
        ("1. Escalabilidad computacional: A medida que crece el número de usuarios y productos, el problema de encontrar emparejamientos óptimos se vuelve exponencialmente más complejo. Lo que funciona para 100 usuarios puede colapsar con 10,000.", "Lista", False),
        ("2. Latencia percibida: Los usuarios de aplicaciones móviles esperan respuestas en menos de 2 segundos. Si nuestro algoritmo tarda 10 segundos en encontrar emparejamientos óptimos, la experiencia de usuario se degrada significativamente.", "Lista", False),
        ("3. Errores de emparejamiento: Un algoritmo que sugiere intercambios desequilibrados (por ejemplo, proponiendo intercambiar un iPhone por un libro usado) erosionará rápidamente la confianza de los usuarios.", "Lista", False),
        ("", "Espacio", False),
        
        ("La infraestructura como punto único de fallo", "Subtítulo", True),
        ("Nuestra arquitectura serverless, aunque moderna y escalable, introduce dependencias críticas:", "Texto Normal", False),
        ("", "Espacio", False),
        
        ("• Proveedores de cloud: Una interrupción en AWS, Google Cloud o Azure podría dejar la plataforma completamente inoperativa.", "Lista", False),
        ("• APIs de terceros: Dependemos de servicios de pago (Stripe), mensajería (Twilio), geolocalización (Google Maps) y autenticación (OAuth). La falla de cualquiera de estos servicios afecta funcionalidades core.", "Lista", False),
        ("• Costes impredecibles: El modelo de pago por uso de serverless puede generar facturas sorpresa si hay picos de tráfico inesperados o bugs que generen llamadas repetitivas.", "Lista", False),
        ("", "Espacio", False),
        
        ("La seguridad como riesgo existencial", "Subtítulo", True),
        ("Para una plataforma que maneja transacciones económicas y datos personales, una brecha de seguridad no es solo un inconveniente técnico, es una amenaza existencial:", "Texto Normal", False),
        ("", "Espacio", False),
        
        ("1. Robo de datos personales: Nombres, direcciones, números de teléfono, preferencias de consumo.", "Lista", False),
        ("2. Fraude en transacciones: Suplantación de identidad, pagos fraudulentos, chargebacks.", "Lista", False),
        ("3. Ataques de denegación de servicio: Que podrían dejar la plataforma inaccesible durante horas o días críticos.", "Lista", False),
        ("4. Vulnerabilidades en el código: Desde inyecciones SQL hasta cross-site scripting que podrían comprometer todo el sistema.", "Lista", False),
        ("", "Espacio", False),
        
        ("9.3 Riesgos Financieros y de Liquidez", "Título Sección", True),
        ("", "Espacio", False),
        
        ("El valle de la muerte del cash flow", "Subtítulo", True),
        ("Las proyecciones financieras muestran que Treqe alcanzará el punto de equilibrio alrededor de los 3,333 transacciones mensuales, lo que según nuestro modelo ocurrirá aproximadamente en el mes 9-10 de operaciones. El riesgo financiero crític", "Texto Normal", False),
    ]
    
    return contenido

def anadir_seccion_riesgos():
    """Añadir sección de riesgos al documento existente."""
    
    try:
        # Cargar documento existente
        input_path = "plan_negocio/Plan_Negocio_Treqe_COMPLETAMENTE_HUMANIZADO_FINAL.docx"
        output_path = "plan_negocio/Plan_Negocio_Treqe_CON_RIESGOS.docx"
        
        print(f"Cargando documento: {input_path}")
        doc = docx.Document(input_path)
        
        print(f"Documento cargado. Párrafos actuales: {len(doc.paragraphs)}")
        
        # Añadir nueva página para la sección 9
        doc.add_page_break()
        
        # Añadir título de la sección
        titulo = doc.add_paragraph("9. ANÁLISIS DE RIESGOS Y MITIGACIÓN")
        titulo.style = "Heading 1"
        titulo.runs[0].bold = True
        
        # Añadir contenido
        doc.add_paragraph()  # Espacio
        
        # 9.1 Riesgos de Mercado
        subtitulo1 = doc.add_paragraph("9.1 Riesgos de Mercado y Competencia")
        subtitulo1.style = "Heading 2"
        subtitulo1.runs[0].bold = True
        
        doc.add_paragraph()  # Espacio
        
        # Contenido 9.1
        contenido_9_1 = [
            "El dilema del huevo y la gallina en plataformas de intercambio",
            "El riesgo más fundamental que enfrenta Treqe es el clásico problema de las plataformas de doble cara: sin vendedores no hay compradores, y sin compradores no hay vendedores. En el contexto de una plataforma de intercambio inteligente, este desafío se multiplica, ya que no solo necesitamos usuarios, sino usuarios dispuestos a participar en ruedas de intercambio complejas donde la liquidez depende de múltiples participantes simultáneamente.",
            "",
            "La experiencia de plataformas como Wallapop nos enseña que los primeros 10,000 usuarios son los más difíciles de conseguir, pero una vez alcanzada esa masa crítica, el crecimiento se acelera exponencialmente. El riesgo aquí no es solo la falta de usuarios iniciales, sino la posibilidad de que usuarios tempranos se desencanten por la falta de contrapartes disponibles, creando un ciclo negativo de abandono que se retroalimenta.",
            "",
            "La sombra de los gigantes establecidos",
            "Wallapop, con sus 15 millones de usuarios activos en España, y Vinted, con su enfoque especializado en moda de segunda mano, representan competidores formidables no solo por su tamaño, sino por su capacidad de respuesta. Ambas plataformas tienen equipos de producto que monitorizan constantemente el mercado y pueden implementar características similares a las de Treqe en cuestión de meses si perciben una amenaza real.",
            "",
            "El riesgo no es que copien nuestra tecnología (que es relativamente sencilla de replicar), sino que utilicen su escala masiva para ofrecer funcionalidades similares de manera gratuita o con comisiones más bajas, aprovechando sus economías de escala y su base de usuarios ya establecida.",
            "",
            "La volatilidad del mercado de segunda mano",
            "Los datos del Observatorio Cetelem muestran que el 62% de los españoles consulta aplicaciones de segunda mano al menos una vez por semana, pero solo el 19% realiza compras regulares. Esta brecha entre intención y acción representa un riesgo fundamental.",
        ]
        
        for texto in contenido_9_1:
            if texto == "":
                doc.add_paragraph()
            elif texto.startswith("El dilema") or texto.startswith("La sombra") or texto.startswith("La volatilidad"):
                p = doc.add_paragraph(texto)
                p.runs[0].bold = True
            else:
                doc.add_paragraph(texto)
        
        # 9.2 Riesgos Operacionales
        doc.add_paragraph()  # Espacio
        subtitulo2 = doc.add_paragraph("9.2 Riesgos Operacionales y Tecnológicos")
        subtitulo2.style = "Heading 2"
        subtitulo2.runs[0].bold = True
        
        doc.add_paragraph()  # Espacio
        
        # Contenido 9.2
        contenido_9_2 = [
            "La complejidad oculta del algoritmo de emparejamiento",
            "El corazón de Treqe es su algoritmo de emparejamiento inteligente que crea 'ruedas de intercambio' óptimas. Este algoritmo presenta riesgos operacionales significativos:",
            "",
            "• Escalabilidad computacional: A medida que crece el número de usuarios, el problema se vuelve exponencialmente más complejo.",
            "• Latencia percibida: Los usuarios esperan respuestas en menos de 2 segundos.",
            "• Errores de emparejamiento: Sugerencias desequilibradas erosionan la confianza.",
            "",
            "La infraestructura como punto único de fallo",
            "Nuestra arquitectura serverless introduce dependencias críticas:",
            "",
            "• Proveedores de cloud: Interrupciones podrían dejar la plataforma inoperativa.",
            "• APIs de terceros: Stripe, Twilio, Google Maps, OAuth.",
            "• Costes impredecibles: Modelo de pago por uso puede generar facturas sorpresa.",
            "",
            "La seguridad como riesgo existencial",
            "Para una plataforma que maneja transacciones y datos personales:",
            "",
            "• Robo de datos personales",
            "• Fraude en transacciones",
            "• Ataques de denegación de servicio",
            "• Vulnerabilidades en el código",
        ]
        
        for texto in contenido_9_2:
            if texto == "":
                doc.add_paragraph()
            elif texto.startswith("La complejidad") or texto.startswith("La infraestructura") or texto.startswith("La seguridad"):
                p = doc.add_paragraph(texto)
                p.runs[0].bold = True
            else:
                doc.add_paragraph(texto)
        
        # 9.3 Riesgos Financieros
        doc.add_paragraph()  # Espacio
        subtitulo3 = doc.add_paragraph("9.3 Riesgos Financieros y de Liquidez")
        subtitulo3.style = "Heading 2"
        subtitulo3.runs[0].bold = True
        
        doc.add_paragraph()  # Espacio
        
        # Contenido 9.3
        contenido_9_3 = [
            "El valle de la muerte del cash flow",
            "Treqe alcanzará el punto de equilibrio alrededor de los 3,333 transacciones mensuales (mes 9-10). El riesgo es el período anterior:",
            "",
            "• Inversión inicial: €58,000 (colchón de ~6 meses)",
            "• Gastos fijos mensuales: €4,200",
            "• Gastos variables: Comisiones de pagos + adquisición de usuarios",
            "• Imprevistos: Multas, costes legales",
            "",
            "Si el crecimiento es más lento de lo proyectado, podríamos agotar capital antes del breakeven.",
        ]
        
        for texto in contenido_9_3:
            if texto == "":
                doc.add_paragraph()
            elif texto.startswith("El valle"):
                p = doc.add_paragraph(texto)
                p.runs[0].bold = True
            else:
                doc.add_paragraph(texto)
        
        # 9.4 Riesgos Legales
        doc.add_paragraph()  # Espacio
        subtitulo4 = doc.add_paragraph("9.4 Riesgos Legales y Regulatorios")
        subtitulo4.style = "Heading 2"
        subtitulo4.runs[0].bold = True
        
        doc.add_paragraph()  # Espacio
        
        # Contenido 9.4
        contenido_9_4 = [
            "El laberinto de la regulación de pagos",
            "Treqe procesará pagos entre usuarios, sujeto a:",
            "",
            "• Banco de España: Ley 10/2014 de pagos",
            "• PSD2: Autenticación reforzada",
            "• Prevención de blanqueo: Sistemas KYC",
            "",
            "Incumplimiento conlleva multas (hasta 10% facturación) y posible revocación de licencias.",
        ]
        
        for texto in contenido_9_4:
            if texto == "":
                doc.add_paragraph()
            elif texto.startswith("El laberinto"):
                p = doc.add_paragraph(texto)
                p.runs[0].bold = True
            else:
                doc.add_paragraph(texto)
        
        # 9.5 Plan de Mitigación
        doc.add_paragraph()  # Espacio
        subtitulo5 = doc.add_paragraph("9.5 Plan de Mitigación y Contingencia")
        subtitulo5.style = "Heading 2"
        subtitulo5.runs[0].bold = True
        
        doc.add_paragraph()  # Espacio
        
        # Contenido 9.5
        contenido_9_5 = [
            "Estrategias clave de mitigación:",
            "",
            "Para riesgos de mercado:",
            "• Enfoque hiperlocal (ej: barrio Gràcia en Barcelona)",
            "• Programa de embajadores en comunidades específicas",
            "• Garantía de liquidez inicial (equipo como market maker)",
            "",
            "Para riesgos operacionales:",
            "• Arquitectura de fallo