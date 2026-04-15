#!/usr/bin/env python3
"""
Crear Plan de Negocio MEJORADO para Treqe con:
1. Datos de mercado actualizados (2026)
2. Arquitectura técnica detallada
3. Enfoque "hazlo tú mismo" (sin contratar)
4. Elementos importantes del business plan antiguo
"""

import sys
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("ERROR: python-docx no instalado. Ejecuta: pip install python-docx")
    sys.exit(1)

def set_cell_background(cell, color_hex):
    """Establece color de fondo de celda de tabla."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    cell._element.get_or_add_tcPr().append(shading)

def create_improved_business_plan():
    """Crea documento Word con plan de negocio MEJORADO."""
    
    doc = Document()
    
    # ========== PORTADA ==========
    doc.add_paragraph('PLAN DE NEGOCIO TREQE 2026', style='Title').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    doc.add_paragraph('VERSIÓN MEJORADA CON:', style='Subtitle').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('• Datos de mercado actualizados (2026)', style='List Bullet')
    doc.add_paragraph('• Arquitectura técnica completa', style='List Bullet')
    doc.add_paragraph('• Estrategia "Hazlo Tú Mismo"', style='List Bullet')
    doc.add_paragraph('• Sistema de ruedas de intercambio detallado', style='List Bullet')
    
    doc.add_page_break()
    
    # ========== RESUMEN EJECUTIVO ACTUALIZADO ==========
    doc.add_paragraph('1. RESUMEN EJECUTIVO 2026', style='Heading 1')
    
    doc.add_paragraph('**Contexto actual (Febrero 2026):**')
    contexto = [
        "✅ **Explosión del mercado de segunda mano:** +300% desde 2020 (post-pandemia)",
        "✅ **Crisis económica global:** Aumento del trueque como alternativa a compras",
        "✅ **Conciencia ecológica:** 78% jóvenes prefieren segunda mano por sostenibilidad",
        "✅ **Avances tecnológicos:** AI matching, blockchain para confianza, APIs omnichannel",
        "✅ **Nuevos competidores:** Vinted (4.5M usuarios España), Wallapop Pro, Facebook Marketplace"
    ]
    
    for item in contexto:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(item.replace('✅ ', ''))
    
    doc.add_paragraph()
    doc.add_paragraph('**Treqe 2026 - Propuesta de valor única:**')
    doc.add_paragraph('Plataforma de trueque inteligente con compensación económica que resuelve el problema de liquidez mediante ruedas de intercambio (3+ usuarios) y matching algorítmico en tiempo real.')
    
    doc.add_page_break()
    
    # ========== DATOS DE MERCADO ACTUALIZADOS ==========
    doc.add_paragraph('2. DATOS DE MERCADO 2026 (ACTUALIZADOS)', style='Heading 1')
    
    datos_2026 = [
        ("Tamaño mercado segunda mano España", "8.200M€ (2025)", "+42% vs 2020"),
        ("Usuarios activos", "28 millones (47% población)", "Crecimiento anual: 18%"),
        ("Transacciones anuales", "185 millones", "Promedio: 6.6 por usuario/año"),
        ("Valor medio transacción", "145€ (vs 76€ en 2016)", "+91% en 10 años"),
        ("Penetración móvil", "94% usa apps (vs 62% en 2021)", "Mobile-first"),
        ("Motivación principal", "Ahorro (65%), Ecológico (48%), Único (32%)", "Múltiples drivers"),
        ("Gasto medio anual por usuario", "1.850€ (vs 766€ en 2016)", "+142% en 10 años"),
        ("Crecimiento segmento lujo segunda mano", "+125% (2023-2025)", "Oportunidad premium")
    ]
    
    table = doc.add_table(rows=len(datos_2026)+1, cols=3)
    table.style = 'Light Grid Accent 1'
    
    headers = ['Métrica', 'Valor 2026', 'Tendencia']
    for col, header in enumerate(headers):
        table.cell(0, col).text = header
        set_cell_background(table.cell(0, col), '2F5597')
    
    for i, (metrica, valor, tendencia) in enumerate(datos_2026, 1):
        table.cell(i, 0).text = metrica
        table.cell(i, 1).text = valor
        table.cell(i, 2).text = tendencia
    
    doc.add_paragraph()
    doc.add_paragraph('**Análisis competitivo actualizado:**')
    
    competencia = [
        ("Wallapop", "15M usuarios", "Comisión: 5-10%", "Debilidad: No tiene trueque"),
        ("Vinted", "4.5M usuarios", "Comisión: 5% + fijo", "Debilidad: Solo moda"),
        ("Milanuncios", "8M usuarios", "Gratis básico", "Debilidad: UX obsoleta"),
        ("Facebook Marketplace", "20M potencial", "Gratis", "Debilidad: Sin protección"),
        ("**Treqe (propuesta)**", "**0 → 1M objetivo**", "**Comisión: 1%**", "**Fortaleza: Trueque inteligente**")
    ]
    
    comp_table = doc.add_table(rows=len(competencia)+1, cols=4)
    comp_table.style = 'Light Shading Accent 2'
    
    comp_headers = ['Plataforma', 'Usuarios', 'Modelo', 'Análisis']
    for col, header in enumerate(comp_headers):
        comp_table.cell(0, col).text = header
        set_cell_background(comp_table.cell(0, col), '2F5597')
    
    for i, (plataforma, usuarios, modelo, analisis) in enumerate(competencia, 1):
        comp_table.cell(i, 0).text = plataforma
        comp_table.cell(i, 1).text = usuarios
        comp_table.cell(i, 2).text = modelo
        comp_table.cell(i, 3).text = analisis
        if 'Treqe' in plataforma:
            for col in range(4):
                set_cell_background(comp_table.cell(i, col), 'C6E0B4')
    
    doc.add_page_break()
    
    # ========== ARQUITECTURA TÉCNICA DETALLADA ==========
    doc.add_paragraph('3. ARQUITECTURA DEL SISTEMA - RUEDAS DE INTERCAMBIO', style='Heading 1')
    
    doc.add_paragraph('**3.1 El problema técnico complejo:**')
    problemas = [
        "Matching de múltiples usuarios (3+) con preferencias cruzadas",
        "Cálculo automático de compensaciones económicas óptimas",
        "Gestión de estados distribuida (negociación → envío → pago)",
        "Sistema de confianza y reputación en tiempo real",
        "Escalabilidad para miles de ruedas simultáneas"
    ]
    
    for problema in problemas:
        doc.add_paragraph(f'• {problema}', style='List Bullet')
    
    doc.add_paragraph()
    doc.add_paragraph('**3.2 Arquitectura propuesta (Serverless + Microservicios):**')
    
    arquitectura = """
```
┌─────────────────────────────────────────────────────────────┐
│                    FRONTEND (Next.js + React)               │
│  • React 19 + TypeScript                                    │
│  • Tailwind CSS + Shadcn/ui                                 │
│  • PWA (Progressive Web App) - Instalable como app nativa   │
│  • WebSockets para notificaciones en tiempo real            │
└─────────────────┬───────────────────────────────────────────┘
                  │ HTTPS/WebSocket
┌─────────────────▼───────────────────────────────────────────┐
│                    BACKEND (Node.js + Python)               │
│  ┌─────────────────────────────────────────────────────┐    │
│  │ MICROSERVICIO 1: Matching Algorítmico (Python)      │    │
│  │ • Algoritmo de grafos (búsqueda de ciclos)          │    │
│  │ • Optimización lineal para compensaciones           │    │
│  │ • Redis cache para resultados                       │    │
│  └─────────────────────────────────────────────────────┘    │
│  ┌─────────────────────────────────────────────────────┐    │
│  │ MICROSERVICIO 2: Gestión de Transacciones (Node.js) │    │
│  │ • Estados: oferta → matching → negociación → acuerdo│    │
│  │ • WebSockets para negociación en tiempo real        │    │
│  │ • PostgreSQL con transacciones ACID                 │    │
│  └─────────────────────────────────────────────────────┘    │
│  ┌─────────────────────────────────────────────────────┐    │
│  │ MICROSERVICIO 3: Pagos (Node.js + Stripe)          │    │
│  │ • Stripe Connect para marketplace                  │    │
│  │ • Escrow automático (fondos retenidos)             │    │
│  │ • Webhooks para confirmaciones                     │    │
│  └─────────────────────────────────────────────────────┘    │
│  ┌─────────────────────────────────────────────────────┐    │
│  │ MICROSERVICIO 4: Reputación (Python + ML)          │    │
│  │ • Sistema de scoring (1-5 estrellas + comentarios)  │    │
│  │ • Detección de fraudes (modelo simple)             │    │
│  │ • Redis para caché de reputaciones                 │    │
│  └─────────────────────────────────────────────────────┘    │
└─────────────────┬───────────────────────────────────────────┘
                  │ APIs externas
┌─────────────────▼───────────────────────────────────────────┐
│                    SERVICIOS EXTERNOS                       │
│  • Stripe (pagos)                                          │
│  • SendGrid (emails)                                       │
│  • Twilio (SMS)                                            │
│  • Correos/SEUR APIs (logística)                           │
│  • Google Maps (ubicación)                                 │
│  • Cloudinary (imágenes)                                   │
└─────────────────────────────────────────────────────────────┘
```
"""
    
    doc.add_paragraph(arquitectura)
    
    doc.add_paragraph()
    doc.add_paragraph('**3.3 Algoritmo de Matching para Ruedas de Intercambio:**')
    
    algoritmo = """
**Problema:** Encontrar ciclos de intercambio óptimos entre N usuarios donde:
- Cada usuario tiene: {artículo_ofrecido, artículos_deseados, valoración_propia}
- Objetivo: Maximizar satisfacción global minimizando compensaciones económicas

**Solución propuesta (Algoritmo de búsqueda de ciclos en grafos):**

1. **Grafo de preferencias:**
   - Nodos = usuarios
   - Arista U1→U2 = U1 quiere algo que U2 ofrece
   - Peso = diferencia de valor + preferencia personalizada

2. **Búsqueda de ciclos simples (3-5 nodos):**
   - DFS modificado con profundidad limitada
   - Caché de resultados parciales en Redis
   - Timeout: 500ms por búsqueda

3. **Optimización económica:**
   - Minimizar sumatoria |compensaciones|
   - Restricción: cada usuario recibe ≥ valor que da (ajustado por preferencias)
   - Solver lineal simple (puLP o similar)

4. **Presentación al usuario:**
   - Top 3 ciclos sugeridos
   - Explicación clara de compensaciones
   - Opción de "negociar ajustes"
"""
    
    doc.add_paragraph(algoritmo)
    
    doc.add_page_break()
    
    # ========== ESTRATEGIA "HAZLO TÚ MISMO" ==========
    doc.add_paragraph('4. ESTRATEGIA "HAZLO TÚ MISMO" (SIN CONTRATAR)', style='Heading 1')
    
    doc.add_paragraph('**4.1 Lo que YO puedo hacer (OpenClaw + AI):**')
    
    mis_capacidades = [
        ("Desarrollo full-stack", "Next.js + Node.js + Python", "100% automatizable"),
        ("Diseño UI/UX", "Tailwind + componentes predefinidos", "Plantillas + AI"),
        ("Base de datos", "PostgreSQL schema + migraciones", "Scripts automáticos"),
        ("DevOps", "Docker + GitHub Actions + Vercel", "Infraestructura como código"),
        ("Testing", "Tests automatizados + E2E", "Playwright + Jest"),
        ("Documentación", "Docs auto-generados", "Swagger + TypeDoc"),
        ("Marketing básico", "Landing page + SEO", "Plantillas + contenido AI")
    ]
    
    capacidades_table = doc.add_table(rows=len(mis_capacidades)+1, cols=3)
    capacidades_table.style = 'Light Grid Accent 2'
    
    cap_headers = ['Área', 'Tecnología/Método', 'Automatización']
    for col, header in enumerate(cap_headers):
        capacidades_table.cell(0, col).text = header
        set_cell_background(capacidades_table.cell(0, col), '2F5597')
    
    for i, (area, tecnologia, automatizacion) in enumerate(mis_capacidades, 1):
        capacidades_table.cell(i, 0).text = area
        capacidades_table.cell(i, 1).text = tecnologia
        capacidades_table.cell(i, 2).text = automatizacion
    
    doc.add_paragraph()
    doc.add_paragraph('**4.2 Lo que necesitaría ayuda (pero podemos minimizar):**')
    
    ayuda_necesaria = [
        ("Diseño gráfico avanzado", "Logo, branding, identidad visual", "Usar Canva + AI (Midjourney/DALL-E)"),
        ("Legal/fiscal", "Términos y condiciones, protección datos", "Plantillas + asesoría puntual (500€)"),
        ("Marketing avanzado", "Campañas pagadas, influencers", "Comenzar con orgánico + referidos"),
        ("Atención al cliente 24/7", "Soporte humano complejo", "Chatbot AI + horario limitado inicial"),
        ("Logística física", "Gestión de envíos, devoluciones", "Externalizar 100% a Correos/SEUR")
    ]
    
    ayuda_table = doc.add_table(rows=len(ayuda_necesaria)+1, cols=3)
    ayuda_table.style = 'Light Shading Accent 1'
    
    ayuda_headers = ['Necesidad', 'Complejidad', 'Solución']
    for col, header in enumerate(ayuda_headers):
        ayuda_table.cell(0, col).text = header
