#!/usr/bin/env python3
"""
Corregir anexos para que sean verídicos - sin información inventada
"""

import docx
import os

def corregir_anexos_veridicos():
    """Crear anexos verídicos sin información inventada."""
    
    # Ruta al documento original
    input_path = "plan_negocio/Plan_Negocio_Treqe_100%_PROFESIONAL_FINAL.docx"
    output_path = "plan_negocio/Plan_Negocio_Treqe_ANEXOS_VERIDICOS.docx"
    
    if not os.path.exists(input_path):
        print(f"Error: No se encuentra: {input_path}")
        return None
    
    print(f"Cargando documento: {input_path}")
    doc = docx.Document(input_path)
    
    print("Buscando sección de anexos...")
    
    # Encontrar donde empiezan los anexos
    inicio_anexos = -1
    for i, p in enumerate(doc.paragraphs):
        if 'ANEXOS' in p.text.upper() and len(p.text) < 20:
            inicio_anexos = i
            print(f"Sección ANEXOS encontrada en párrafo {i}")
            break
    
    if inicio_anexos == -1:
        print("Añadiendo sección ANEXOS al final...")
        doc.add_page_break()
        titulo = doc.add_paragraph("ANEXOS")
        titulo.style = "Heading 1"
        inicio_anexos = len(doc.paragraphs) - 1
    
    # Eliminar contenido existente de anexos
    print("Limpiando contenido existente...")
    i = inicio_anexos + 1
    while i < len(doc.paragraphs):
        p = doc.paragraphs[i]
        if p.style.name == "Heading 1" and i > inicio_anexos + 1:
            break
        i += 1
    
    if i > inicio_anexos + 1:
        for _ in range(inicio_anexos + 1, i):
            if inicio_anexos + 1 < len(doc.paragraphs):
                p = doc.paragraphs[inicio_anexos + 1]
                p_element = p._element
                p_element.getparent().remove(p_element)
    
    print("Creando anexos verídicos...")
    
    # ========== ANEXO A: ESTRUCTURA DEL EQUIPO (NO CVs INVENTADOS) ==========
    print("  - Anexo A: Estructura del equipo (roles necesarios)")
    doc.add_paragraph()
    
    anexo_a = doc.add_paragraph("ANEXO A: ESTRUCTURA Y ROLES DEL EQUIPO")
    anexo_a.style = "Heading 2"
    
    contenido_a = [
        "",
        "ROLES NECESARIOS PARA EL LANZAMIENTO:",
        "",
        "1. CEO/FUNDADOR",
        "• Responsabilidades: Visión estratégica, fundraising, relaciones inversores",
        "• Experiencia requerida: Conocimiento sector segunda mano, liderazgo, negociación",
        "• Tiempo estimado: Full-time",
        "",
        "2. CTO/TECH LEAD", 
        "• Responsabilidades: Desarrollo plataforma, arquitectura, equipo técnico",
        "• Habilidades requeridas: Full-stack development, algoritmos, escalabilidad",
        "• Stack recomendado: Node.js/Python, React, PostgreSQL, AWS/Cloud",
        "• Tiempo estimado: Full-time",
        "",
        "3. PRODUCT/DESIGN LEAD",
        "• Responsabilidades: UX/UI, investigación usuario, roadmap producto",
        "• Habilidades requeridas: Design thinking, Figma/Adobe XD, user testing",
        "• Tiempo estimado: Full-time o contractor inicial",
        "",
        "4. GROWTH/MARKETING",
        "• Responsabilidades: Adquisición usuarios, community building, contenidos",
        "• Habilidades requeridas: Marketing digital, analytics, redes sociales",
        "• Tiempo estimado: Part-time inicial, scaling con crecimiento",
        "",
        "NOTA: Los CVs específicos del equipo fundador se añadirán cuando se definan los miembros reales.",
    ]
    
    for item in contenido_a:
        if item == "":
            doc.add_paragraph()
        else:
            doc.add_paragraph(item)
    
    doc.add_page_break()
    
    # ========== ANEXO B: METODOLOGÍA DE ESTUDIO (NO DATOS INVENTADOS) ==========
    print("  - Anexo B: Metodología para estudios de mercado")
    anexo_b = doc.add_paragraph("ANEXO B: METODOLOGÍA PARA ESTUDIOS DE MERCADO")
    anexo_b.style = "Heading 2"
    
    contenido_b = [
        "",
        "METODOLOGÍAS RECOMENDADAS PARA VALIDAR EL MERCADO:",
        "",
        "1. INVESTIGACIÓN SECUNDARIA (INMEDIATA)",
        "• Fuentes públicas: INE, Eurostat, informes sectoriales",
        "• Análisis competencia: SimilarWeb, App Annie, reports públicos",
        "• Tendencias: Google Trends, informes consultoras (McKinsey, BCG, etc.)",
        "",
        "2. ENCUESTAS ONLINE (FASE 1 - 2-4 SEMANAS)",
        "• Plataformas: Typeform, Google Forms, SurveyMonkey",
        "• Muestra objetivo: 500-1000 usuarios españoles",
        "• Coste estimado: 500-2000€ (depende de plataforma y targeting)",
        "• Métricas clave: Intención de uso, preocupaciones, disposición a pagar",
        "",
        "3. ENTREVISTAS/USER TESTING (FASE 2 - 4-6 SEMANAS)",
        "• Participantes: 20-30 usuarios potenciales",
        "• Método: Entrevistas semi-estructuradas, testing prototipos",
        "• Objetivo: Deep dive en pain points, validar user flows",
        "• Coste estimado: 1000-3000€ (incentivos participantes)",
        "",
        "4. PILOTO CONTROLADO (FASE 3 - 8-12 SEMANAS)",
        "• Ubicación: Barrio específico (ej: Gràcia, Barcelona)",
        "• Participantes: 200-500 usuarios activos",
        "• Métricas: Engagement, retención, tasa conversión, NPS",
        "• Coste estimado: 5000-15000€ (desarrollo + marketing local)",
        "",
        "NOTA: Los datos específicos de estudios se añadirán cuando se realicen las investigaciones.",
    ]
    
    for item in contenido_b:
        if item == "":
            doc.add_paragraph()
        else:
            doc.add_paragraph(item)
    
    doc.add_page_break()
    
    # ========== ANEXO C: DISEÑOS DETALLADOS (DESARROLLADOS) ==========
    print("  - Anexo C: Diseños detallados de plataforma")
    anexo_c = doc.add_paragraph("ANEXO C: DISEÑOS DETALLADOS DE LA PLATAFORMA")
    anexo_c.style = "Heading 2"
    
    contenido_c = [
        "",
        "ARQUITECTURA DE INFORMACIÓN - SITEMAP:",
        "",
        "NIVEL 1: PÁGINAS PÚBLICAS",
        "• Homepage (/) - Value prop, búsqueda, categorías, testimonios",
        "• Cómo funciona (/como-funciona) - Explicación visual del proceso",
        "• Categorías (/categorias) - Exploración por tipo de producto",
        "• Blog (/blog) - Contenido educativo sobre economía circular",
        "• FAQ (/faq) - Preguntas frecuentes",
        "• Contacto (/contacto) - Formulario y información",
        "",
        "NIVEL 2: ÁREA DE USUARIO (REQUIERE LOGIN)",
        "• Dashboard (/dashboard) - Resumen actividad, notificaciones",
        "• Mi perfil (/perfil) - Información personal, configuración",
        "• Mis productos (/productos) - Inventario, publicar nuevo",
        "• Mis intercambios (/intercambios) - Activos, historial",
        "• Mensajes (/mensajes) - Chat integrado",
        "• Configuración (/configuracion) - Preferencias, privacidad",
        "",
        "NIVEL 3: PÁGINAS DE PRODUCTO",
        "• Detalle producto (/producto/{id}) - Galería, descripción, valoración",
        "• Propuesta intercambio (/intercambiar/{id}) - Formulario propuesta",
        "• Ruedas sugeridas (/ruedas/{id}) - Algoritmo de matching",
        "",
        "USER FLOW PRINCIPAL - INTERCAMBIO:",
        "",
        "1. USUARIO VISITANTE → REGISTRO",
        "• Entrada: Homepage, redes sociales, referidos",
        "• Acción: Click 'Empieza gratis' o 'Regístrate'",
        "• Opciones: Email, Google, Facebook, Apple",
        "• Verificación: Email/móvil (2 pasos)",
        "• Onboarding: Tutorial interactivo (3 pasos)",
        "",
        "2. USUARIO REGISTRADO → PUBLICAR PRODUCTO",
        "• Acceso: Botón '+' en header o dashboard",
        "• Paso 1: Fotos (mínimo 3, máximo 10)",
        "  - Cámara directa o subir desde galería",
        "  - Sugerencias IA: encuadre, iluminación",
        "  - Etiquetado automático de categorías",
        "• Paso 2: Descripción",
        "  - Campos: Título, descripción, categoría, estado",
        "  - Asistente IA: Sugiere descripciones basadas en fotos",
        "  - Valoración estimada: Algoritmo sugiere rango",
        "• Paso 3: Preferencias intercambio",
        "  - Qué busca a cambio: Categorías, valor aproximado",
        "  - Radio de intercambio: Local (5km), ciudad, provincia",
        "  - Disponibilidad: Fechas/horarios",
        "",
        "3. BUSCAR/PROPONER INTERCAMBIO",
        "• Búsqueda:",
        "  - Semántica: 'mueble pequeño para estudio'",
        "  - Filtros: Categoría, radio, valor, estado",
        "  - Guardar búsquedas: Alertas para nuevos matches",
        "• Descubrimiento:",
        "  - Recomendaciones: Basadas en historial y preferencias",
        "  - Ruedas sugeridas: 'Estas personas buscan lo que tú tienes'",
        "  - Categorías trending: Lo más buscado en tu zona",
        "• Propuesta:",
        "  - 1-clic: 'Intercambiar' muestra ruedas posibles",
        "  - Multi-opción: Ver todas las combinaciones posibles",
        "  - Negociación: Chat integrado para ajustar términos",
        "",
        "4. COMPLETAR INTERCAMBIO",
        "• Acuerdo:",
        "  - Términos: Fecha, lugar, condiciones",
        "  - Compensación: Si aplica, monto acordado",
        "  - Garantía: Sistema de depósito en garantía opcional",
        "• Ejecución:",
        "  - Recordatorios: Notificaciones push/email",
        "  - Check-in: Confirmación llegada al punto encuentro",
        "  - Verificación: Código QR para confirmar entrega",
        "• Post-intercambio:",
        "  - Calificación: Sistema de estrellas + comentarios",
        "  - Reputación: Puntos acumulativos visibles en perfil",
        "  - Recompensas: Badges por intercambios exitosos",
        "",
        "ESPECIFICACIONES DE DISEÑO - UI/UX:",
        "",
        "SISTEMA DE DISEÑO:",
        "• Colores primarios:",
        "  - Azul confianza: #1a73e8 (primary actions)",
        "  - Verde sostenibilidad: #34a853 (eco features)",
        "  - Gris neutro: #5f6368 (text, backgrounds)",
        "  - Blanco: #ffffff (cards, surfaces)",
        "• Tipografía:",
        "  - Inter: Sans-serif para UI (legibilidad en pantallas)",
        "  - Merriweather: Serif para contenido largo (blog, descripciones)",
        "  - Escalas: 12, 14, 16, 18, 20, 24, 32, 48, 64px",
        "• Espaciado:",
        "  - Base: 4px (escala: 4, 8, 12, 16, 24, 32, 48, 64, 96)",
        "  - Container: Max-width 1280px, padding 24px móvil / 48px desktop",
        "• Componentes:",
        "  - Botones: Primary (relleno), Secondary (contorno), Text (solo texto)",
        "  - Cards: Sombras sutiles, bordes redondeados 8px",
        "  - Forms: Estados claro (default, hover, focus, error, disabled)",
        "  - Navigation: Fijo en móvil, sticky en desktop",
        "",
        "PRINCIPIOS UX:",
        "• Mobile-first: 70% tráfico esperado desde móvil",
        "• Accesibilidad: Nivel AA WCAG 2.1 (contraste 4.5:1 mínimo)",
        "• Performance: First Contentful Paint < 1.5s, Time to Interactive < 3s",
        "• Navegación: Máximo 3 clics para acciones principales",
        "• Feedback: Estados loading, success, error claros e inmediatos",
        "• Error prevention: Validación en tiempo real, confirmaciones destructivas",
        "",
        "PATRONES DE INTERACCIÓN:",
        "• Gestos móvil: Swipe para acciones rápidas (archivar, favorito)",
        "• Pull-to-refresh: Actualización contenido en listados",
        "• Infinite scroll: Paginación automática en resultados búsqueda",
        "• Progressive disclosure: Información compleja revelada gradualmente",
        "• Undo actions: Posibilidad deshacer acciones importantes (5 segundos)",
    ]
    
    for item in contenido_c:
        if item == "":
            doc.add_paragraph()
        else:
            doc.add_paragraph(item)
    
    doc.add_page_break()
    
    # ========== ANEXO D: GLOSARIO ==========
    print("  - Anexo D: Glosario de términos del proyecto")
    anexo_d = doc.add_paragraph("ANEXO D: GLOSARIO DE TÉRMINOS DEL PROYECTO")
    anexo_d.style = "Heading 2"
    
    contenido_d = [
        "",
        "TÉRMINOS ESPECÍFICOS DE TREQE:",
        "",
        "• Rueda de Intercambio: Ciclo cerrado donde A da a B, B da a C, y C da a A",
        "• Algoritmo de Matching: Sistema que encuentra combinaciones óptimas entre múltiples usuarios",
        "• Compensación Monetaria: Pago pequeño que complementa diferencias de valor en intercambios",
        "• Liquidez de Plataforma: Facilidad para encontrar contrapartes (más usuarios = más liquidez)",
        "• Ciclo Cerrado de Valor: Productos que circulan continuamente sin convertirse en residuos",
        "",
        "TÉRMINOS GENERALES DE PLATAFORMAS:",
        "",
        "• Two-sided Marketplace: Plataforma que conecta dos grupos (vendedores/compradores)",
        "• Network Effects: Valor aumenta con número de usuarios (más usuarios = mejor experiencia)",
        "• Chicken-and-Egg Problem: Desafío inicial: sin vendedores no hay compradores, sin compradores no hay vendedores",
        "• CAC (Customer Acquisition Cost): Coste promedio de adquirir un usuario",
        "• LTV (Lifetime Value): Valor económico total de un usuario durante su tiempo en plataforma",
        "• Burn Rate: Ritmo de gasto de capital antes de generar cash flow positivo",
        "• Runway: Tiempo que puede operar con capital actual antes de necesitar más financiación",
        "",
        "TÉRMINOS DE ECONOMÍA CIRCULAR:",
        "",
        "• Economía Circular: Sistema donde recursos se mantienen en uso el mayor tiempo posible",
        "• Reutilización: Usar un producto nuevamente para su propósito original",
        "• Redistrib