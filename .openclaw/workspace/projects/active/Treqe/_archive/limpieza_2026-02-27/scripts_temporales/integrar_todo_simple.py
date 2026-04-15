#!/usr/bin/env python3
"""
Integrar todo el contenido adjunto de manera simple
"""

from docx import Document
import os

def integrar_todo():
    """Integrar todas las secciones adjuntas."""
    
    print("Integrando contenido adjunto completo...")
    
    # Cargar documento existente
    doc_path = os.path.join(os.path.dirname(__file__), 'plan_negocio', 'Plan_Negocio_Treqe_Completo.docx')
    doc = Document(doc_path)
    
    # Agregar sección 3
    doc.add_heading('3. SOLUCIÓN INNOVADORA: RUEDAS DE INTERCAMBIO INTELIGENTE', 0)
    
    # 3.1
    doc.add_heading('3.1 Concepto Fundamental', 1)
    doc.add_paragraph('Treqe introduce "ruedas de intercambio" que permiten participación de 3-5 usuarios en cadenas circulares de valor, resolviendo matemáticamente el problema de doble coincidencia.')
    
    # 3.2
    doc.add_heading('3.2 Mecanismo Operativo Detallado', 1)
    doc.add_paragraph('Paso 1: Registro Preferencias')
    doc.add_paragraph('• Usuarios indican "Tengo" (artículos disponibles) + "Quiero" (artículos deseados)')
    doc.add_paragraph('• Especifican valor estimado, condiciones, fotos, descripción')
    
    doc.add_paragraph('Paso 2: Algoritmo Matching (Teoría Grafos)')
    doc.add_paragraph('• Construye grafo dirigido: usuarios → artículos deseados')
    doc.add_paragraph('• Busca ciclos cerrados 3-5 nodos usando DFS optimizado')
    doc.add_paragraph('• Timeout: 500ms por búsqueda (escalabilidad)')
    
    doc.add_paragraph('Paso 3: Optimización Económica (Programación Lineal)')
    doc.add_paragraph('• Calcula compensaciones monetarias óptimas (algoritmo PuLP)')
    doc.add_paragraph('• Minimiza transferencias totales')
    doc.add_paragraph('• Maximiza satisfacción global')
    doc.add_paragraph('• Considera preferencias individuales')
    
    doc.add_paragraph('Paso 4: Negociación Facilitada')
    doc.add_paragraph('• Chat grupal tiempo real (WebSockets)')
    doc.add_paragraph('• Sistema sugiere términos basados optimización')
    doc.add_paragraph('• Acuerdo requiere confirmación todos participantes')
    
    doc.add_paragraph('Paso 5: Ejecución Segura')
    doc.add_paragraph('• Pagos vía Stripe Connect (fondos escrow)')
    doc.add_paragraph('• Logística integrada APIs Correos/SEUR')
    doc.add_paragraph('• Sistema reputación + garantías')
    doc.add_paragraph('• Soporte disputas automatizado')
    
    # 3.3
    doc.add_heading('3.3 Ejemplo Práctico Detallado', 1)
    doc.add_paragraph('Situación Inicial:')
    doc.add_paragraph('• Ana: Bicicleta montaña (450€) → necesita sofá diseño (600€)')
    doc.add_paragraph('• Carlos: Sofá diseño (600€) → necesita ordenador gaming (800€)')
    doc.add_paragraph('• Beatriz: Ordenador gaming (800€) → necesita bicicleta montaña (450€)')
    
    doc.add_paragraph('Solución Treqe:')
    doc.add_paragraph('Intercambios físicos:')
    doc.add_paragraph('1. Ana → Beatriz: Bicicleta')
    doc.add_paragraph('2. Carlos → Ana: Sofá')
    doc.add_paragraph('3. Beatriz → Carlos: Ordenador')
    
    doc.add_paragraph('Compensaciones monetarias:')
    doc.add_paragraph('• Ana paga 150€ a Carlos (600€ - 450€)')
    doc.add_paragraph('• Carlos paga 200€ a Beatriz (800€ - 600€)')
    doc.add_paragraph('• Beatriz recibe 350€ neto (150€ + 200€)')
    
    doc.add_paragraph('Resultados:')
    doc.add_paragraph('• Ana: Sofá 600€ por 150€ → Ahorro 450€ (75%)')
    doc.add_paragraph('• Carlos: Ordenador 800€ por 200€ → Ahorro 600€ (75%)')
    doc.add_paragraph('• Beatriz: Bicicleta 450€ + 350€ → Valor total 800€')
    
    doc.add_paragraph('Innovación diferencial: Sistema híbrido único combina trueque físico con compensación económica optimizada.')
    
    doc.add_page_break()
    
    # Sección 4
    doc.add_heading('4. VENTAJA COMPETITIVA SOSTENIBLE', 0)
    
    # 4.1
    doc.add_heading('4.1 Posicionamiento Único', 1)
    doc.add_paragraph('Primer mover en trueque estructurado en España - nicho inexplorado por competencia.')
    
    # 4.2
    doc.add_heading('4.2 Ventajas Tecnológicas', 1)
    doc.add_paragraph('Algoritmos propietarios:')
    doc.add_paragraph('• Matching basado teoría grafos (NetworkX)')
    doc.add_paragraph('• Optimización lineal compensaciones (PuLP)')
    doc.add_paragraph('• Sistema reputación machine learning')
    doc.add_paragraph('• Detección fraudes análisis patrones')
    
    doc.add_paragraph('Arquitectura moderna:')
    doc.add_paragraph('• Frontend: Next.js 14 + React 19 + TypeScript + PWA')
    doc.add_paragraph('• Backend: Node.js + Express + WebSockets')
    doc.add_paragraph('• Matching Service: Python microservicio + Redis')
    doc.add_paragraph('• Base datos: PostgreSQL + TimescaleDB')
    doc.add_paragraph('• Infraestructura: Serverless (Vercel + Railway)')
    doc.add_paragraph('• APIs integradas: Stripe, Correos/SEUR, Google Maps, SendGrid')
    
    # 4.3
    doc.add_heading('4.3 Ventajas Económicas', 1)
    doc.add_paragraph('Modelo comisiones disruptivo:')
    doc.add_paragraph('Plataforma | Comisión | Coste adicional')
    doc.add_paragraph('Treqe | 1% | Ninguno')
    doc.add_paragraph('Wallapop | 5% | +0,90€ fijo')
    doc.add_paragraph('Vinted | 5% | +0,70€ + 3% protección')
    doc.add_paragraph('Ventaja Treqe: 80-90% más barato')
    
    doc.add_paragraph('Transparencia radical:')
    doc.add_paragraph('• Comisión única predecible')
    doc.add_paragraph('• Sin costes ocultos')
    doc.add_paragraph('• Pago solo al recibir valor')
    doc.add_paragraph('• Desglose detallado compensaciones')
    
    # 4.4
    doc.add_heading('4.4 Ventajas Sostenibilidad', 1)
    doc.add_paragraph('Impacto medioambiental:')
    doc.add_paragraph('• Extensión vida útil productos: +3-5 años')
    doc.add_paragraph('• Reducción residuos: ~150kg CO2eq/transacción')
    doc.add_paragraph('• Economía circular real (no greenwashing)')
    
    doc.add_paragraph('Impacto social:')
    doc.add_paragraph('• Acceso bienes sin liquidez')
    doc.add_paragraph('• Reducción desigualdades económicas')
    doc.add_paragraph('• Construcción comunidades locales')
    doc.add_paragraph('• Fomento confianza entre ciudadanos')
    
    doc.add_paragraph('Contribución ODS:')
    doc.add_paragraph('• ODS 12: Producción consumo responsables')
    doc.add_paragraph('• ODS 13: Acción por el clima')
    doc.add_paragraph('• ODS 11: Ciudades comunidades sostenibles')
    
    # 4.5
    doc.add_heading('4.5 Barreras Entrada', 1)
    doc.add_paragraph('1. Complejidad algorítmica: 6-9 meses desarrollo equipo especializado')
    doc.add_paragraph('2. Efecto red local: Masa crítica comunidades geográficas')
    doc.add_paragraph('3. Base datos preferencias: Activo intangible que crece con tiempo')
    
    doc.add_page_break()
    
    # Sección 5
    doc.add_heading('5. MODELO DE NEGOCIO', 0)
    
    # 5.1
    doc.add_heading('5.1 Flujos Ingresos Multicapa', 1)
    doc.add_paragraph('Fase 1 (Año 1): Comisión Básica 1%')
    doc.add_paragraph('• Sobre valor artículo adquirido')
    doc.add_paragraph('• Pagada exclusivamente por receptor')
    doc.add_paragraph('• Ejemplo: Recibe artículo 500€ → paga 5€')
    
    doc.add_paragraph('Fase 2 (Año 2): Servicios Premium')
    doc.add_paragraph('• Cuenta Verificada: 4,99€/mes')
    doc.add_paragraph('• Destacados búsquedas: 2,99€/artículo')
    doc.add_paragraph('• Logística Premium: +3€/envío')
    doc.add_paragraph('• Objetivo: 20% usuarios premium')
    
    doc.add_paragraph('Fase 3 (Año 3): Publicidad Segmentada')
    doc.add_paragraph('• Anuncios marcas sostenibles (certificación B Corp)')
    doc.add_paragraph('• Promociones categorías específicas')
    doc.add_paragraph('• Partnerships empresas circulares')
    doc.add_paragraph('• Modelo: CPM + CPC')
    
    # 5.2
    doc.add_heading('5.2 Inversión Inicial: 58.000€', 1)
    doc.add_paragraph('Concepto | Importe | % Total | Detalle')
    doc.add_paragraph('Desarrollo tecnológico | 23.200€ | 40% | Frontend, backend, algoritmos, infra')
    doc.add_paragraph('Marketing inicial | 20.300€ | 35% | Campañas digitales, contenido, eventos')
    doc.add_paragraph('Operaciones y equipo | 14.500€ | 25% | Equipo fundador, legal, oficina')
    doc.add_paragraph('TOTAL | 58.000€ | 100%')
    
    # 5.3
    doc.add_heading('5.3 Financiación Propuesta', 1)
    doc.add_paragraph('• 40.000€ (69%): Inversores ángeles / business angels')
    doc.add_paragraph('• 10.000€ (17%): Préstamo participativo ENISA')
    doc.add_paragraph('• 8.000€ (14%): Aportación equipo fundador')
    doc.add_paragraph('Valoración pre-money: 200.000€')
    doc.add_paragraph('Equity ofrecido: 15-20%')
    doc.add_paragraph('ROI esperado inversores: 3-5x en 3-5 años')
    
    doc.add_page_break()
    
    # Sección 6
    doc.add_heading('6. PROYECCIONES FINANCIERAS 2026-2029', 0)
    
    # 6.1
    doc.add_heading('6.1 Supuestos Clave', 1)
    doc.add_paragraph('• Comisión efectiva: 1% sobre valor artículo adquirido')
    doc.add_paragraph('• Valor medio transacción: 150€ (año 1), 160€ (año 2), 170€ (año 3)')
    doc.add_paragraph('• Tasa conversión usuarios activos → transacciones: 10% mensual')
    doc.add_paragraph('• Crecimiento usuarios: 15% mensual (año 1), 10% (año 2), 5% (año 3)')
    doc.add_paragraph('• Retención mensual usuarios: 70%')
    
    # 6.2
    doc.add_heading('6.2 Proyecciones Anuales', 1)
    doc.add_paragraph('Métrica | Año 1 | Año 2 | Año 3')
    doc.add_paragraph('Usuarios activos finales | 25.000 | 75.000 | 150.000')
    doc.add_paragraph('Transacciones anuales | 15.000 | 60.000 | 120.000')
    doc.add_paragraph('Volumen transaccional | 2.250.000€ | 9.000.000€ | 18.000.000€')
    doc.add_paragraph('Ingresos comisiones (1%) | 22.500€ | 90.000€ | 180.000€')
    doc.add_paragraph('Ingresos servicios premium | 0€ | 18.000€ | 48.000€')
    doc.add_paragraph('Ingresos publicidad | 0€ | 6.000€ | 18.000€')
    doc.add_paragraph('TOTAL INGRESOS | 22.500€ | 114.000€ | 246.000€')
    
    # 6.3
    doc.add_heading('6.3 Estado Pérdidas y Ganancias', 1)
    doc.add_paragraph('Concepto | Año 1 | Año 2 | Año 3')
    doc.add_paragraph('Ingresos totales | 22.500€ | 114.000€ | 246.000€')
    doc.add_paragraph('Costes desarrollo | (23.200€) | (25.000€) | (30.000€)')
    doc.add_paragraph('Costes marketing | (20.300€) | (35.000€) | (45.000€)')
    doc.add_paragraph('Costes personal | (9.000€) | (18.000€) | (30.000€)')
    doc.add_paragraph('Costes operativos | (5.500€) | (8.000€) | (12.000€)')
    doc.add_paragraph('EBITDA | (35.500€) | 28.000€ | 129.000€')
    doc.add_paragraph('Amortizaciones | (1.000€) | (1.500€) | (2.000€)')
    doc.add_paragraph('Intereses financieros | (500€) | (1.000€) | (1.500€)')
    doc.add_paragraph('Resultado antes impuestos | (37.000€) | 25.500€ | 125.500€')
    doc.add_paragraph('Impuestos (25%) | 0€ | 0€ | (31.375€)')
    doc.add_paragraph('RESULTADO NETO | (37.000€) | 25.500€ | 94.125€')
    
    # 6.4
    doc.add_heading('6.4 Cash Flow Proyectado', 1)
    doc.add_paragraph('Cash Flow Operativo:')
    doc.add_paragraph('• Año 1: -28.000€ (inversión)')
    doc.add_paragraph('• Año 2: +12.000€ (positivo)')
    doc.add_paragraph('• Año 3: +58.000€ (sólido)')
    
    doc.add_paragraph('Punto de equilibrio:')
    doc.add_paragraph('• Cálculo: 4.500€ costes fijos / (1,50€ comisión - 0,15€ coste) = 3.333 transacciones/mes')
    doc.add_paragraph('• Equivalente: 33.330 usuarios activos (10% tasa conversión)')
    doc.add_paragraph('• Fecha objetivo: Mes 10')
    
    doc.add_paragraph('Tesorería proyectada:')
    doc.add_paragraph('• Fin año 1: 14.678€')
    doc.add_paragraph('• Fin año 2: 26.678€')
    doc.add_paragraph('• Fin año 3: 84.678€')
    doc.add_paragraph('Runway después año 1: 14 meses operación')
    
    # 6.5
    doc.add_heading('6.5 Ratios Financieros Clave (Año 3)', 1)
    doc.add_paragraph('• Margen EBITDA: 52,4% (129.000€ / 246.000€)')
    doc.add_paragraph('• Margen neto: 38,3% (94.125€ / 246.000€)')
    doc.add_paragraph('• ROI inversión: 162% (94.125€ / 58.000€)')
    doc.add_paragraph('• LTV:CAC: 24:1 (excelente, >3:1 es bueno)')
    doc.add_paragraph('• Current ratio: 5,8 (liquidez alta, >2 es bueno)')
    doc.add_paragraph('• Burn rate año 1: 4.143€/mes')
    doc.add_paragraph('• Payback period: 2,1 años')
    
    doc.add_page_break