#!/usr/bin/env python3
"""
Integrar las secciones adjuntas (3-7) en el documento Word completo
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def integrar_secciones():
    """Integrar secciones 3-7 del contenido adjunto."""
    
    print("Integrando secciones 3-7 en documento completo...")
    
    # Cargar documento existente
    doc_path = os.path.join(os.path.dirname(__file__), 'plan_negocio', 'Plan_Negocio_Treqe_Completo.docx')
    doc = Document(doc_path)
    
    # ========== SECCIÓN 3: SOLUCIÓN INNOVADORA ==========
    
    doc.add_heading('3. SOLUCIÓN INNOVADORA: RUEDAS DE INTERCAMBIO INTELIGENTE', 0)
    
    # 3.1 Concepto Fundamental
    doc.add_heading('3.1 Concepto Fundamental', 1)
    
    text_3_1 = """Treqe introduce el concepto de **"ruedas de intercambio"** que permiten la participación de 3 a 5 usuarios en cadenas circulares de valor, resolviendo matemáticamente el problema histórico de la doble coincidencia de deseos que ha limitado el trueque tradicional.

Esta innovación representa un salto cualitativo en la forma en que concebimos el intercambio de bienes. Mientras que el trueque tradicional requiere que dos personas quieran exactamente lo que la otra ofrece (una coincidencia estadísticamente improbable), las ruedas de intercambio de Treqe permiten que múltiples participantes se conecten en ciclos cerrados donde cada uno recibe exactamente lo que desea, aunque no directamente de la persona a la que le entrega su artículo.

La clave de este sistema reside en su capacidad para identificar y optimizar estas cadenas circulares mediante algoritmos avanzados, transformando lo que históricamente ha sido un proceso manual, lento y limitado en uno automatizado, rápido y escalable."""
    
    doc.add_paragraph(text_3_1)
    
    # 3.2 Mecanismo Operativo Detallado
    doc.add_heading('3.2 Mecanismo Operativo Detallado', 1)
    
    text_3_2 = """El sistema de Treqe opera a través de un flujo estructurado en cinco fases fundamentales, cada una diseñada para garantizar eficiencia, seguridad y satisfacción del usuario:

**Paso 1: Registro de Preferencias**
Los usuarios ingresan al sistema y declaran dos categorías fundamentales:
- **"Tengo":** Los artículos que poseen y están dispuestos a intercambiar
- **"Quiero":** Los artículos que desean obtener

Para cada ítem, se solicita información detallada que incluye:
- Descripción completa y estado de conservación
- Fotografías desde múltiples ángulos
- Valoración económica estimada (con sugerencias del sistema basadas en datos de mercado)
- Preferencias específicas (radio geográfico máximo, estado mínimo aceptable, etc.)

**Paso 2: Algoritmo de Matching (Teoría de Grafos)**
El sistema transforma estas preferencias en una estructura matemática: un grafo dirigido donde:
- **Nodos:** Representan a los usuarios
- **Aristas dirigidas:** Van desde lo que un usuario tiene hacia lo que otro usuario quiere

Utilizando la librería NetworkX, el sistema aplica algoritmos de búsqueda en profundidad (DFS) optimizados para identificar ciclos cerrados de 3 a 5 nodos que representan posibles cadenas de intercambio. Cada búsqueda tiene un timeout de 500ms para garantizar escalabilidad.

**Paso 3: Optimización Económica (Programación Lineal)**
Cuando los valores de los artículos en un ciclo no coinciden exactamente (lo que ocurre en el 95% de los casos), el sistema calcula compensaciones monetarias óptimas. Utilizando el algoritmo PuLP de programación lineal, determina las transferencias mínimas necesarias para equilibrar el intercambio, considerando:
- Minimización del flujo monetario total
- Equidad en la distribución de compensaciones
- Preferencias individuales sobre pagos/recepciones

**Paso 4: Negociación Facilitada**
Una vez identificado un ciclo viable, se crea un espacio de negociación grupal con las siguientes características:
- **Chat en tiempo real:** Utilizando WebSockets para comunicación instantánea
- **Panel de condiciones:** Visualización clara de lo que cada uno da y recibe
- **Sugerencias del sistema:** Propuestas basadas en la optimización económica
- **Modificación flexible:** Los usuarios pueden ajustar términos dentro de ciertos límites
- **Acuerdo requerido:** Todas las partes deben confirmar explícitamente

**Paso 5: Ejecución Segura**
Una vez alcanzado el acuerdo, el sistema gestiona la ejecución completa:
- **Pagos en escrow:** Utilizando Stripe Connect, los fondos de compensación se retienen hasta confirmación de recepción
- **Logística integrada:** APIs con Correos, SEUR y otras transportistas para envíos rastreables
- **Seguimiento en tiempo real:** Todos los participantes pueden monitorizar el estado de cada envío
- **Sistema de disputas:** Mecanismo escalonado para resolver incidencias
- **Liberación de fondos:** Los pagos se completan automáticamente tras confirmaciones positivas"""
    
    doc.add_paragraph(text_3_2)
    
    # 3.3 Ejemplo Práctico Detallado
    doc.add_heading('3.3 Ejemplo Práctico Detallado', 1)
    
    text_3_3 = """Para comprender la potencia del sistema, analicemos un caso concreto con mayor profundidad:

**Situación Inicial:**
- **Ana:** Posee una bicicleta de montaña valorada en 450€, pero necesita un sofá de diseño valorado en 600€
- **Carlos:** Posee un sofá de diseño valorado en 600€, pero necesita un ordenador gaming valorado en 800€
- **Beatriz:** Posee un ordenador gaming valorado en 800€, pero necesita una bicicleta de montaña valorada en 450€

**Problema Tradicional:**
Ningún intercambio directo 1:1 es viable porque:
- Ana quiere el sofá de Carlos, pero Carlos no quiere la bicicleta de Ana
- Carlos quiere el ordenador de Beatriz, pero Beatriz no quiere el sofá de Carlos
- Beatriz quiere la bicicleta de Ana, pero Ana no quiere el ordenador de Beatriz

**Solución Treqe:**
El algoritmo identifica un ciclo cerrado perfecto y propone:

**Intercambios físicos:**
1. Ana → Beatriz: Bicicleta de montaña
2. Carlos → Ana: Sofá de diseño
3. Beatriz → Carlos: Ordenador gaming

**Compensaciones monetarias:**
- Ana paga 150€ a Carlos (diferencia: 600€ - 450€)
- Carlos paga 200€ a Beatriz (diferencia: 800€ - 600€)
- Beatriz recibe 350€ neto (150€ + 200€)

**Resultados finales:**
- **Ana:** Obtiene sofá valor 600€ por 150€ → Ahorro: 450€ (75%)
- **Carlos:** Obtiene ordenador valor 800€ por 200€ → Ahorro: 600€ (75%)
- **Beatriz:** Obtiene bicicleta valor 450€ + 350€ → Valor total recibido: 800€
- **Treqe:** Recibe 15€ en comisiones (1% de 1.500€ de valor intercambiado)

**Innovación diferencial:**
Este sistema híbrido único combina trueque físico con compensación económica optimizada, permitiendo a los usuarios obtener lo que necesitan sin requerir liquidez completa, mientras maximizan el valor obtenido de sus posesiones actuales."""
    
    doc.add_paragraph(text_3_3)
    
    doc.add_page_break()
    
    # ========== SECCIÓN 4: VENTAJA COMPETITIVA ==========
    
    doc.add_heading('4. VENTAJA COMPETITIVA SOSTENIBLE', 0)
    
    # 4.1 Posicionamiento Único
    doc.add_heading('4.1 Posicionamiento Único', 1)
    
    text_4_1 = """Treqe se posiciona como **primer mover en el segmento de trueque estructurado en España**, ocupando un nicho que las plataformas establecidas han elegido no atender. Esta posición estratégica ofrece varias ventajas:

**Ausencia de competencia directa inicial:**
Los gigantes actuales (Wallapop, Vinted) no tienen incentivo para desarrollar sistemas de trueque complejos porque:
- Su modelo de negocio se basa en volumen de transacciones monetarias
- Los sistemas de trueque reducirían potencialmente su facturación (menos ventas monetarias)
- La complejidad técnica requerida es significativa y no alineada con sus roadmaps de desarrollo

**Captura de un segmento no atendido:**
Treqe atiende específicamente a usuarios que:
- Prefieren intercambiar antes que vender (8 millones de españoles según estudios)
- Carecen de liquidez para compras directas pero tienen valor atrapado en posesiones
- Valoran la sostenibilidad sobre la transacción monetaria pura
- Buscan soluciones más creativas que la simple compraventa

**Efectos de red locales más fuertes:**
A diferencia de plataformas globales, Treqe potencia comunidades locales:
- Intercambios preferentemente en radios reducidos (menos de 50km)
- Construcción de confianza entre vecinos
- Reducción de costes logísticos
- Mayor probabilidad de intercambios repetidos"""
    
    doc.add_paragraph(text_4_1)
    
    # 4.2 Ventajas Tecnológicas
    doc.add_heading('4.2 Ventajas Tecnológicas', 1)
    
    text_4_2 = """**Algoritmos propietarios avanzados:**
- **Matching basado en teoría de grafos:** Utilizando NetworkX para identificar ciclos óptimos
- **Optimización lineal de compensaciones:** Algoritmo PuLP para minimizar transferencias monetarias
- **Sistema de reputación con machine learning:** Evaluaciones predictivas basadas en historial
- **Detección de fraudes mediante análisis de patrones:** Identificación proactiva de comportamientos sospechosos

**Arquitectura tecnológica moderna y escalable:**
- **Frontend:** Next.js 14 + React 19 + TypeScript + PWA (Progressive Web App)
- **Backend:** Node.js + Express + WebSockets para comunicación en tiempo real
- **Servicio de Matching:** Microservicio Python especializado + Redis para caché
- **Base de datos:** PostgreSQL con extensión TimescaleDB para datos temporales
- **Infraestructura:** Serverless en Vercel (frontend) y Railway (backend/microservicios)
- **APIs integradas:** Stripe Connect (pagos), Correos/SEUR (logística), Google Maps (geolocalización), SendGrid (email)

**Ventajas de la arquitectura serverless:**
- Costes operativos reducidos en 90% vs infraestructura tradicional
- Escalabilidad automática según demanda
- Mantenimiento simplificado (sin gestión de servidores)
- Deployment continuo y actualizaciones instantáneas"""
    
    doc.add_paragraph(text_4_2)
    
    # 4.3 Ventajas Económicas
    doc.add_heading('4.3 Ventajas Económicas', 1)
    
    text_4_3 = """**Modelo de comisiones disruptivo:**

| Plataforma | Comisión Base | Costes Adicionales | Total Efectivo |
|------------|---------------|-------------------|----------------|
| **Treqe** | **1%** | Ninguno | **1%** |
| Wallapop | 5% | +0,90€ fijo | 5-15% dependiendo valor |
| Vinted | 5% | +0,70€ + 3% protección | 8-9% promedio |
| eBay | 10-12% | +coste listing | 12-15% |

**Análisis de la ventaja competitiva:**
- **Para transacción de 100€:** Treqe cobra 1€ vs 5-15€ de competencia
- **Para usuario que recibe artículo:** Paga solo al obtener valor tangible
- **Transparencia radical:** Sin costes ocultos, sin sorpresas
- **Alineación de incentivos:** Treqe gana cuando los usuarios completan intercambios satisfactorios

**Propuesta de valor económica por segmento:**
- **Usuarios con limitaciones de liquidez:** Acceso a bienes necesarios sin desembolso completo
- **Vendedores ocasionales:** Maximización del valor obtenido de posesiones
- **Defensores de sostenibilidad:** Contribución tangible a economía circular con beneficio económico
- **Comunidades locales:** Retención de valor económico dentro de la zona geográfica"""
    
    doc.add_paragraph(text_4_3)
    
    # 4.4 Ventajas de Sostenibilidad
    doc.add_heading('4.4 Ventajas de Sostenibilidad', 1)
    
    text_4_4 = """**Impacto medioambiental cuantificable:**
- **Extensión de vida útil de productos:** Estimado en +3-5 años por artículo intercambiado
- **Reducción de residuos:** Aproximadamente 150kg de CO2 equivalente evitados por transacción
- **Economía circular real:** No greenwashing, impacto medible y verificable
- **Conservación de recursos:** Menor extracción de materias primas y energía de fabricación

**Impacto social positivo:**
- **Acceso a bienes sin liquidez:** Reducción de barreras económicas para renovación de posesiones
- **Reducción de desigualdades económicas:** Sistema que beneficia especialmente a segmentos con limitaciones presupuestarias
- **Construcción de comunidades locales:** Fortalecimiento de lazos entre vecinos
- **Fomento de confianza entre ciudadanos:** Sistema reputacional que premia comportamientos honestos

**Contribución a Objetivos de Desarrollo Sostenible (ODS):**
- **ODS 12:** Producción y consumo responsables (meta 12.5: reducción residuos)
- **ODS 13:** Acción por el clima (meta 13.3: educación y concienciación)
- **ODS 11:** Ciudades y comunidades sostenibles (meta 11.6: impacto ambiental urbano)
- **ODS 8:** Trabajo decente y crecimiento económico (meta 8.4: eficiencia recursos)"""
    
    doc.add_paragraph(text_4_4)
    
    # 4.5 Barreras de Entrada
    doc.add_heading('4.5 Barreras de Entrada', 1)
    
    text_4_5 = """**1. Complejidad algorítmica (Alta):**
- **Tiempo de desarrollo:** 6-9 meses para equipo especializado
- **Conocimiento requerido:** Teoría de grafos, optimización lineal, sistemas distribuidos
- **Testing extensivo:** Necesario para garantizar estabilidad y equidad
- **Optimización continua:** El algoritmo mejora con cada transacción completada

**2. Efecto de red local (Media-Alta):**
- **Masa crítica necesaria:** Comunidades geográficas requieren usuarios suficientes para matching efectivo
- **Confianza acumulativa:** Cada transacción exitosa fortalece la red local
- **Ventaja del primer mover:** Difícil desplazar una comunidad establecida
- **Data network effects:** Los datos de preferencias mejoran el matching con el tiempo

**3. Base de datos de preferencias (Activo intangible):**
- **Valor creciente:** Cada usuario añadido mejora las posibilidades de matching
- **Dificultad de replicación:** Competidores necesitarían tiempo para construir base comparable
- **Aprendizaje continuo:** El sistema mejora con cada interacción
- **Personalización:** Mejor experiencia usuario con más datos disponibles"""
    
    doc.add_paragraph(text_4_5)
    
    doc.add_page_break()
    
    # ========== SECCIÓN 5: MODELO DE NEGOCIO ==========
    
    doc.add_heading('5. MODELO DE NEGOCIO', 0)
    
    # 5.1 Flujos de Ingresos Multicapa
    doc.add_heading('5.1 Flujos de Ingresos Multicapa', 1)
    
    text_5_1 = """El modelo de negocio de Treqe evoluciona a través de tres fases estratégicas, cada una diseñada para maximizar ingresos mientras se mantiene la propuesta de valor central:

**Fase 1: Lanzamiento y Validación (Año 1) - Comisión Básica 1%**
- **Mecanismo:** 1% sobre el valor del artículo adquirido
- **Pagador:** Exclusivamente el usuario que recibe el bien
- **Ejemplo:** Usuario recibe artículo valorado en 500€ → paga 5€ de comisión
-