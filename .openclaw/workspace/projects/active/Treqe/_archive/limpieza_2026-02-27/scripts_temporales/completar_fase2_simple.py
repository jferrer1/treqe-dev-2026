#!/usr/bin/env python3
"""
Completar Fase 2 de manera simple
"""

from docx import Document
import os

def completar_fase2():
    """Completar Fase 2 simple."""
    
    print("COMPLETANDO FASE 2 SIMPLE...")
    
    # Cargar documento
    doc_path = os.path.join(os.path.dirname(__file__), 'plan_negocio', 'Plan_Negocio_Treqe_FASE1_COMPLETADA.docx')
    doc = Document(doc_path)
    
    # Agregar sección 10
    doc.add_page_break()
    doc.add_heading('10. ESTRATEGIA DE MARKETING DETALLADA', 0)
    
    doc.add_heading('10.1 Plan de Lanzamiento Fase por Fase', 1)
    doc.add_paragraph('Fase Alpha (Meses 1-2): Validacion con 50 usuarios invitados, matching asistido, feedback intensivo.')
    doc.add_paragraph('Fase Beta (Meses 3-4): Waitlist publica, onboarding por invitacion, programas de referidos.')
    doc.add_paragraph('Fase Lanzamiento (Meses 5-6): Lanzamiento Barcelona/Madrid, campanas digitales, partnerships.')
    doc.add_paragraph('Fase Expansion (Meses 7-12): Expansion a principales ciudades espanolas, optimizacion de canales.')
    
    doc.add_heading('10.2 Canales de Adquisicion', 1)
    doc.add_paragraph('Marketing de Contenidos/SEO: Blog sobre economia circular, guias de trueque.')
    doc.add_paragraph('Redes Sociales: Grupos Facebook de segunda mano, subreddits especializados.')
    doc.add_paragraph('Influencers: Micro-influencers de sostenibilidad, programas de embajadores.')
    doc.add_paragraph('Publicidad Digital: Google Ads, Facebook Ads con targeting avanzado.')
    doc.add_paragraph('Partnerships: Tiendas de segunda mano fisica, ONGs ambientales.')
    
    doc.add_heading('10.3 Metricas Clave', 1)
    doc.add_paragraph('CAC (Coste Adquisicion Cliente): Objetivo < 15€')
    doc.add_paragraph('LTV (Valor Vida Cliente): Objetivo > 45€')
    doc.add_paragraph('Retencion a 30 dias: Objetivo > 40%')
    doc.add_paragraph('LTV:CAC Ratio: Objetivo > 3:1')
    
    doc.add_heading('10.4 Presupuesto Marketing Año 1', 1)
    doc.add_paragraph('Total: 120.000€')
    doc.add_paragraph('Meses 1-2: 2.000€ (Alpha)')
    doc.add_paragraph('Meses 3-4: 6.000€ (Beta)')
    doc.add_paragraph('Meses 5-6: 16.000€ (Lanzamiento)')
    doc.add_paragraph('Meses 7-12: 96.000€ (Expansion)')
    
    doc.add_heading('10.5 Cronograma', 1)
    doc.add_paragraph('Mes 1-2: Desarrollo contenido, configuracion canales, identificacion early adopters.')
    doc.add_paragraph('Mes 3-4: Onboarding primeros 100 usuarios, programas referidos, optimizacion.')
    doc.add_paragraph('Mes 5-6: Eventos lanzamiento Barcelona/Madrid, campanas digitales intensivas.')
    doc.add_paragraph('Mes 7-12: Expansion geografica progresiva, consolidacion nacional.')
    
    # Agregar sección 11
    doc.add_page_break()
    doc.add_heading('11. ASPECTOS LEGALES Y REGULATORIOS', 0)
    
    doc.add_heading('11.1 Estructura Legal', 1)
    doc.add_paragraph('Sociedad Limitada (SL) con capital social 3.000€.')
    doc.add_paragraph('Proteccion patrimonial fundadores, credibilidad frente inversores.')
    doc.add_paragraph('Flexibilidad para futuras rondas inversion, compatibilidad con programas ENISA.')
    
    doc.add_heading('11.2 Propiedad Intelectual', 1)
    doc.add_paragraph('Algoritmo de Matching: Proteccion como patente de metodo o secreto comercial.')
    doc.add_paragraph('Sistema Optimizacion Economica: Derechos autor sobre implementacion especifica.')
    doc.add_paragraph('Marca "Treqe": Registro marca comunitaria, proteccion elementos identidad.')
    
    doc.add_heading('11.3 Cumplimiento Normativo', 1)
    doc.add_paragraph('GDPR/LOPD-GDD: Registro actividades tratamiento, delegado proteccion datos.')
    doc.add_paragraph('Ley de Servicios de la Sociedad de la Informacion (LSSI): Aviso legal, condiciones uso.')
    doc.add_paragraph('Normativa consumo: Derechos desistimiento 14 dias, garantias legales.')
    doc.add_paragraph('Aspectos fiscales: IVA no aplicable a comisiones entre particulares (art. 20 LIVA).')
    
    doc.add_heading('11.4 Contratos y Terminos', 1)
    doc.add_paragraph('Terminos y Condiciones: Delimitacion responsabilidades, limites plataforma.')
    doc.add_paragraph('Politica Privacidad: Tratamiento datos transparente, derechos usuarios.')
    doc.add_paragraph('Contratos Proveedores: SLAs con servicios criticos (Stripe, APIs logistica).')
    doc.add_paragraph('Acuerdos Confidencialidad: Con empleados, colaboradores, partners.')
    
    doc.add_heading('11.5 Seguros y Garantias', 1)
    doc.add_paragraph('Seguro Responsabilidad Civil: Cobertura posibles reclamaciones usuarios.')
    doc.add_paragraph('Seguro Ciberriesgo: Proteccion frente a brechas seguridad, ataques informaticos.')
    doc.add_paragraph('Fondo Garantia: Reserva para compensaciones casos excepcionales.')
    doc.add_paragraph('Seguro D&O (Directors and Officers): Proteccion personal directivos.')
    
    # Guardar documento
    output_path = os.path.join(os.path.dirname(__file__), 'plan_negocio', 'Plan_Negocio_Treqe_FASE2_COMPLETADA.docx')
    doc.save(output_path)
    
    size_kb = os.path.getsize(output_path) / 1024
    print("FASE 2 COMPLETADA")
    print("Documento:", output_path)
    print("Tamaño:", size_kb, "KB")
    print("")
    print("SECCIONES AGREGADAS:")
    print("10. ESTRATEGIA DE MARKETING DETALLADA")
    print("   - 10.1 Plan de Lanzamiento Fase por Fase")
    print("   - 10.2 Canales de Adquisicion")
    print("   - 10.3 Metricas Clave")
    print("   - 10.4 Presupuesto Marketing")
    print("   - 10.5 Cronograma")
    print("")
    print("11. ASPECTOS LEGALES Y REGULATORIOS")
    print("   - 11.1 Estructura Legal")
    print("   - 11.2 Propiedad Intelectual")
    print("   - 11.3 Cumplimiento Normativo")
    print("   - 11.4 Contratos y Terminos")
    print("   - 11.5 Seguros y Garantias")
    print("")
    print("ESTADO ACTUAL: 95% COMPLETO")
    print("¿CONTINUAR CON FASE 3?")
    
    return output_path

if __name__ == '__main__':
    try:
        nuevo_doc = completar_fase2()
        print("EXITO: Fase 2 completada")
    except Exception as e:
        print(f"ERROR: {e}")
        import traceback
        traceback.print_exc()