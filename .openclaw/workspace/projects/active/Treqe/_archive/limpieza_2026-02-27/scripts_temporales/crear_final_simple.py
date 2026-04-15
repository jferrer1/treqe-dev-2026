#!/usr/bin/env python3
"""
Crear documento final simple con todo el contenido
"""

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from datetime import datetime

def main():
    print("Creando documento final completo...")
    
    # Crear documento
    doc = Document()
    
    # Márgenes
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Portada
    doc.add_heading('PLAN DE NEGOCIO TREQE', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    doc.add_paragraph('Plataforma de Trueque Inteligente').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    doc.add_paragraph(f'Fecha: {datetime.now().strftime("%d/%m/%Y")}').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Versión: Final Completo').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Estado: Confidencial').alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # Índice
    doc.add_heading('ÍNDICE', 0)
    doc.add_paragraph('1. Introducción')
    doc.add_paragraph('2. Problema no Resuelto')
    doc.add_paragraph('3. Solución Innovadora')
    doc.add_paragraph('4. Ventaja Competitiva')
    doc.add_paragraph('5. Modelo de Negocio')
    doc.add_paragraph('6. Proyecciones Financieras')
    doc.add_paragraph('7. Equipo y Ejecución')
    doc.add_paragraph('8. Conclusiones')
    
    doc.add_page_break()
    
    # 1. Introducción
    doc.add_heading('1. INTRODUCCIÓN', 0)
    doc.add_paragraph('Mercado segunda mano España 2026:')
    doc.add_paragraph('• 8.200M€ volumen (+42% vs 2020)')
    doc.add_paragraph('• 28M usuarios activos (47% población)')
    doc.add_paragraph('• Gasto medio: 1.850€/año')
    doc.add_paragraph('• Mobile-first: 94% apps')
    
    doc.add_paragraph('Competencia:')
    doc.add_paragraph('• Wallapop: 15M usuarios, 5% + 0,90€')
    doc.add_paragraph('• Vinted: 4,5M usuarios, 8-9% comisión')
    doc.add_paragraph('• Facebook Marketplace: Gratuita, básica')
    doc.add_paragraph('• Milanuncios: 10% cuota, tradicional')
    
    doc.add_paragraph('Tendencias:')
    doc.add_paragraph('1. Premiumización (+125% crecimiento)')
    doc.add_paragraph('2. Sostenibilidad (68% motivación ecológica)')
    doc.add_paragraph('3. Comunidades locales (<50km más exitosas)')
    doc.add_paragraph('4. Regulación emergente (2025+)')
    doc.add_paragraph('5. Mobile-first absoluto (75% millennials/gen Z)')
    
    doc.add_page_break()
    
    # 2. Problema
    doc.add_heading('2. PROBLEMA NO RESUELTO', 0)
    doc.add_paragraph('Paradoja de la liquidez:')
    doc.add_paragraph('• Usuarios tienen valor atrapado en posesiones no deseadas')
    doc.add_paragraph('• Carecen de liquidez para adquirir lo que necesitan')
    doc.add_paragraph('• Ejemplo: Ana (Barcelona) - 1.500€ valor atrapado vs 2.000€ necesidades')
    
    doc.add_paragraph('Opciones no óptimas:')
    doc.add_paragraph('A. Mantener (58%): Coste espacio, depreciación, insatisfacción')
    doc.add_paragraph('B. Vender barato: Pérdida 30-50% valor, frustración')
    doc.add_paragraph('C. Postergar: Calidad vida reducida, oportunidades perdidas')
    
    doc.add_paragraph('Limitación matemática:')
    doc.add_paragraph('• Doble coincidencia deseos: <5% éxito intercambios directos')
    doc.add_paragraph('• Tiempo medio: 2-3 meses búsqueda')
    doc.add_paragraph('• Abandono: 45% después 1 mes frustración')
    
    doc.add_paragraph('Oportunidad cuantificada:')
    doc.add_paragraph('• 8M españoles prefieren intercambiar antes que vender')
    doc.add_paragraph('• 15.000M€ valor atrapado en artículos no utilizados')
    doc.add_paragraph('• 73% usuarios frustrados por no poder intercambiar')
    
    doc.add_page_break()
    
    # 3. Solución
    doc.add_heading('3. SOLUCIÓN INNOVADORA', 0)
    doc.add_paragraph('Ruedas de intercambio inteligente:')
    doc.add_paragraph('• 3-5 usuarios en cadenas circulares')
    doc.add_paragraph('• Resuelve matemáticamente doble coincidencia')
    
    doc.add_paragraph('Mecanismo operativo:')
    doc.add_paragraph('1. Registro preferencias ("Tengo" + "Quiero")')
    doc.add_paragraph('2. Algoritmo matching (grafos, DFS optimizado, 500ms timeout)')
    doc.add_paragraph('3. Optimización económica (PuLP, minimizar transferencias)')
    doc.add_paragraph('4. Negociación facilitada (WebSockets, chat grupal)')
    doc.add_paragraph('5. Ejecución segura (Stripe Connect escrow, APIs logística)')
    
    doc.add_paragraph('Ejemplo práctico:')
    doc.add_paragraph('• Ana: Bicicleta 450€ → Sofá 600€')
    doc.add_paragraph('• Carlos: Sofá 600€ → Ordenador 800€')
    doc.add_paragraph('• Beatriz: Ordenador 800€ → Bicicleta 450€')
    
    doc.add_paragraph('Solución:')
    doc.add_paragraph('• Intercambios: Ana→Beatriz (bici), Carlos→Ana (sofá), Beatriz→Carlos (ordenador)')
    doc.add_paragraph('• Compensaciones: Ana paga 150€ a Carlos, Carlos paga 200€ a Beatriz')
    doc.add_paragraph('• Resultados: Ana ahorra 450€ (75%), Carlos ahorra 600€ (75%), Beatriz recibe 800€ valor total')
    
    doc.add_page_break()
    
    # 4. Ventaja Competitiva
    doc.add_heading('4. VENTAJA COMPETITIVA', 0)
    doc.add_paragraph('Posicionamiento único:')
    doc.add_paragraph('• Primer mover trueque estructurado España')
    doc.add_paragraph('• Nicho inexplorado por competencia')
    
    doc.add_paragraph('Ventajas tecnológicas:')
    doc.add_paragraph('• Algoritmos: Grafos (NetworkX), optimización (PuLP), ML reputación')
    doc.add_paragraph('• Arquitectura: Next.js 14 + React 19 + TypeScript + PWA')
    doc.add_paragraph('• Backend: Node.js + WebSockets + Python microservicios')
    doc.add_paragraph('• Infra: Serverless (Vercel + Railway), PostgreSQL + Redis')
    
    doc.add_paragraph('Ventajas económicas:')
    doc.add_paragraph('• Comisión 1% (vs 5-15% competencia)')
    doc.add_paragraph('• 80-90% más barato que Wallapop/Vinted')
    doc.add_paragraph('• Transparencia radical: sin costes ocultos')
    
    doc.add_paragraph('Ventajas sostenibilidad:')
    doc.add_paragraph('• Extensión vida útil: +3-5 años')
    doc.add_paragraph('• Reducción CO2: ~150kg/transacción')
    doc.add_paragraph('• Economía circular real')
    doc.add_paragraph('• Contribución ODS 12, 13, 11')
    
    doc.add_paragraph('Barreras entrada:')
    doc.add_paragraph('1. Complejidad algorítmica (6-9 meses desarrollo)')
    doc.add_paragraph('2. Efecto red local (masa crítica comunidades)')
    doc.add_paragraph('3. Base datos preferencias (activo intangible)')
    
    doc.add_page_break()
    
    # 5. Modelo de Negocio
    doc.add_heading('5. MODELO DE NEGOCIO', 0)
    doc.add_paragraph('Flujos ingresos:')
    doc.add_paragraph('• Fase 1 (Año 1): Comisión 1% sobre valor recibido')
    doc.add_paragraph('• Fase 2 (Año 2): Servicios premium (verificación 4,99€/mes, destacados 2,99€, logística +3€)')
    doc.add_paragraph('• Fase 3 (Año 3): Publicidad segmentada (marcas sostenibles, CPM+CPC)')
    
    doc.add_paragraph('Inversión inicial: 58.000€')
    doc.add_paragraph('• Desarrollo: 23.200€ (40%)')
    doc.add_paragraph('• Marketing: 20.300€ (35%)')
    doc.add_paragraph('• Operaciones: 14.500€ (25%)')
    
    doc.add_paragraph('Financiación:')
    doc.add_paragraph('• Inversores ángeles: 40.000€ (69%)')
    doc.add_paragraph('• Préstamo ENISA: 10.000€ (17%)')
    doc.add_paragraph('• Equipo fundador: 8.000€ (14%)')
    doc.add_paragraph('• Valoración: 200.000€ pre-money')
    doc.add_paragraph('• Equity: 15-20%')
    doc.add_paragraph('• ROI esperado: 3-5x en 3-5 años')
    
    doc.add_page_break()
    
    # 6. Proyecciones Financieras
    doc.add_heading('6. PROYECCIONES FINANCIERAS', 0)
    doc.add_paragraph('Supuestos clave:')
    doc.add_paragraph('• Comisión: 1%')
    doc.add_paragraph('• Valor medio: 150€ (a1), 160€ (a2), 170€ (a3)')
    doc.add_paragraph('• Crecimiento usuarios: 15% (a1), 10% (a2), 5% (a3)')
    doc.add_paragraph('• Retención: 70% mensual')
    
    doc.add_paragraph('Proyecciones:')
    doc.add_paragraph('Año 1: 25.000 usuarios, 15.000 transacciones, 2.250.000€ volumen, 22.500€ ingresos')
    doc.add_paragraph('Año 2: 75.000 usuarios, 60.000 transacciones, 9.000.000€ volumen, 114.000€ ingresos')
    doc.add_paragraph('Año 3: 150.000 usuarios, 120.000 transacciones, 18.000.000€ volumen, 246.000€ ingresos')
    
    doc.add_paragraph('Estado pérdidas y ganancias:')
    doc.add_paragraph('Año 1: EBITDA -35.500€, Resultado neto -37.000€')
    doc.add_paragraph('Año 2: EBITDA +28.000€, Resultado neto +25.500€')
    doc.add_paragraph('Año 3: EBITDA +129.000€, Resultado neto +94.125€')
    
    doc.add_paragraph('Cash flow:')
    doc.add_paragraph('• Año 1: -28.000€')
    doc.add_paragraph('• Año 2: +12.000€')
    doc.add_paragraph('• Año 3: +58.000€')
    doc.add_paragraph('• Punto equilibrio: 3.333 transacciones/mes (mes 10)')
    doc.add_paragraph('• Tesorería año 3: 84.678€')
    doc.add_paragraph('• Runway año 1: 14 meses')
    
    doc.add_paragraph('Ratios año 3:')
    doc.add_paragraph('• Margen EBITDA: 52,4%')
    doc.add_paragraph('• Margen neto: 38,3%')
    doc.add_paragraph('• ROI: 162%')
    doc.add_paragraph('• LTV:CAC: 24:1 (excelente)')
    doc.add_paragraph('• Current ratio: 5,8 (liquidez alta)')
    
    doc.add_page_break()
    
    # 7. Equipo y Ejecución
    doc.add_heading('7. EQUIPO Y EJECUCIÓN', 0)
    doc.add_paragraph('Equipo fundador:')
    doc.add_paragraph('• CEO: 10+ años emprendimiento digital, scale-ups, economía colaborativa')
    doc.add_paragraph('• CTO: PhD Ciencias Computación, ML, sistemas distribuidos')
    doc.add_paragraph('• CMO: Growth marketing, comunidades, sostenibilidad')
    
    doc.add_paragraph('Plan ejecución:')
    doc.add_paragraph('• Fase 1 (meses 1-3): Validación - landing page, algoritmo básico, 500 early adopters Barcelona')
    doc.add_paragraph('• Fase 2 (meses 4-6): MVP - plataforma completa, Stripe Connect, expansión Madrid+Valencia, 5.000 usuarios')
    doc.add_paragraph('• Fase 3 (meses 7-12): Crecimiento - optimización, servicios premium, expansión nacional, 25.000 usuarios')
    
    doc.add_page_break()
    
    # 8. Conclusiones
    doc.add_heading('8. CONCLUSIONES', 0)
    doc.add_paragraph('Treqe ofrece:')
    doc.add_paragraph('1. Innovación real que resuelve problema histórico')
    doc.add_paragraph('2. Mercado validado de 8M usuarios potenciales')
    doc.add_paragraph('3. Ventaja competitiva sostenible (primer mover, algoritmos)')
    doc.add_paragraph('4. Modelo escalable con costes marginales bajos')
    doc.add_paragraph('5. Impacto positivo verificable (sostenibilidad)')
    doc.add_paragraph('6. Viabilidad financiera demostrada (rentabilidad año 3)')
    
    doc.add_paragraph('Próximos pasos:')
    doc.add_paragraph('1. Registrar dominio treqe.es')
    doc.add_paragraph('2. Landing page con waitlist')
    doc.add_paragraph('3. Algoritmo matching POC')
    doc.add_paragraph('4. 500 early adopters Barcelona')
    doc.add_paragraph('5. Validación transacciones reales')
    
    doc.add_paragraph('Métrica éxito 3 meses: 1.000+ waitlist, 100+ transacciones validadas')
    
    # Guardar
    output_dir = os.path.join(os.path.dirname(__file__), 'plan_negocio')
    os.makedirs(output_dir, exist_ok=True)
    
    output_path = os.path.join(output_dir, 'Plan_Negocio_Treqe_FINAL_COMPLETO.docx')
    doc.save(output_path)
    
    print(f"✅ Documento creado: {output_path}")
    print("Contenido incluido:")
    print("- 8 secciones completas")
    print("- Todas las secciones del contenido adjunto")
    print("- Redacción profesional y estructurada")
    print("- Datos financieros completos")
    print("- Plan de ejecución detallado")
    
    return output_path

if __name__ == '__main__':
    try:
        output_file = main()
        print("\n🎉 Documento final completo creado exitosamente.")
    except Exception as e:
        print(f"❌ Error: {e}")
        import traceback
        traceback.print_exc()