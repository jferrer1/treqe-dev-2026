#!/usr/bin/env python3
"""
Continuación del documento completo
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def continuar_documento():
    """Continuar documento desde donde quedó."""
    
    print("Continuando documento completo...")
    
    # Cargar documento existente
    doc_path = os.path.join(os.path.dirname(__file__), 'plan_negocio', 'Plan_Negocio_Treqe_FINAL_COMPLETO.docx')
    doc = Document(doc_path)
    
    # Continuar desde 5.1
    doc.add_paragraph('Fase 2 (Año 2): Servicios Premium')
    doc.add_paragraph('• Cuenta Verificada: 4,99€/mes')
    doc.add_paragraph('• Destacados búsquedas: 2,99€/artículo')
    doc.add_paragraph('• Logística Premium: +3€/envío')
    doc.add_paragraph('• Objetivo: 20% usuarios premium')
    
    doc.add_paragraph('Fase 3 (Año 3): Publicidad Segmentada')
    doc.add_paragraph('• Anuncios marcas sostenibles')
    doc.add_paragraph('• Promociones categorías')
    doc.add_paragraph('• Partnerships circulares')
    doc.add_paragraph('• Modelo: CPM + CPC')
    
    # 5.2
    doc.add_heading('5.2 Inversión Inicial', 1)
    doc.add_paragraph('Total: 58.000€')
    doc.add_paragraph('• Desarrollo tecnológico: 23.200€ (40%)')
    doc.add_paragraph('• Marketing inicial: 20.300€ (35%)')
    doc.add_paragraph('• Operaciones y equipo: 14.500€ (25%)')
    
    # 5.3
    doc.add_heading('5.3 Financiación', 1)
    doc.add_paragraph('• Inversores ángeles: 40.000€ (69%)')
    doc.add_paragraph('• Préstamo ENISA: 10.000€ (17%)')
    doc.add_paragraph('• Equipo fundador: 8.000€ (14%)')
    doc.add_paragraph('Valoración pre-money: 200.000€')
    doc.add_paragraph('Equity ofrecido: 15-20%')
    doc.add_paragraph('ROI esperado: 3-5x en 3-5 años')
    
    doc.add_page_break()
    
    # ========== SECCIÓN 6: PROYECCIONES FINANCIERAS ==========
    doc.add_heading('6. PROYECCIONES FINANCIERAS', 0)
    
    # 6.1
    doc.add_heading('6.1 Supuestos Clave', 1)
    doc.add_paragraph('• Comisión efectiva: 1%')
    doc.add_paragraph('• Valor medio transacción: 150€ (a1), 160€ (a2), 170€ (a3)')
    doc.add_paragraph('• Tasa conversión: 10% mensual')
    doc.add_paragraph('• Crecimiento usuarios: 15% (a1), 10% (a2), 5% (a3)')
    doc.add_paragraph('• Retención: 70% mensual')
    
    # 6.2
    doc.add_heading('6.2 Proyecciones Anuales', 1)
    doc.add_paragraph('Año 1 | Año 2 | Año 3')
    doc.add_paragraph('Usuarios: 25.000 | 75.000 | 150.000')
    doc.add_paragraph('Transacciones: 15.000 | 60.000 | 120.000')
    doc.add_paragraph('Volumen: 2.250.000€ | 9.000.000€ | 18.000.000€')
    doc.add_paragraph('Ingresos comisiones: 22.500€ | 90.000€ | 180.000€')
    doc.add_paragraph('Ingresos premium: 0€ | 18.000€ | 48.000€')
    doc.add_paragraph('Ingresos publicidad: 0€ | 6.000€ | 18.000€')
    doc.add_paragraph('TOTAL INGRESOS: 22.500€ | 114.000€ | 246.000€')
    
    # 6.3
    doc.add_heading('6.3 Estado Pérdidas y Ganancias', 1)
    doc.add_paragraph('Año 1 | Año 2 | Año 3')
    doc.add_paragraph('Ingresos: 22.500€ | 114.000€ | 246.000€')
    doc.add_paragraph('Costes desarrollo: (23.200€) | (25.000€) | (30.000€)')
    doc.add_paragraph('Costes marketing: (20.300€) | (35.000€) | (45.000€)')
    doc.add_paragraph('Costes personal: (9.000€) | (18.000€) | (30.000€)')
    doc.add_paragraph('Costes operativos: (5.500€) | (8.000€) | (12.000€)')
    doc.add_paragraph('EBITDA: (35.500€) | 28.000€ | 129.000€')
    doc.add_paragraph('Resultado neto: (37.000€) | 25.500€ | 94.125€')
    
    # 6.4
    doc.add_heading('6.4 Cash Flow', 1)
    doc.add_paragraph('Cash Flow Operativo:')
    doc.add_paragraph('• Año 1: -28.000€')
    doc.add_paragraph('• Año 2: +12.000€')
    doc.add_paragraph('• Año 3: +58.000€')
    
    doc.add_paragraph('Punto de equilibrio:')
    doc.add_paragraph('• 3.333 transacciones/mes')
    doc.add_paragraph('• 33.330 usuarios activos')
    doc.add_paragraph('• Mes 10')
    
    doc.add_paragraph('Tesorería:')
    doc.add_paragraph('• Fin año 1: 14.678€')
    doc.add_paragraph('• Fin año 2: 26.678€')
    doc.add_paragraph('• Fin año 3: 84.678€')
    doc.add_paragraph('• Runway año 1: 14 meses')
    
    # 6.5
    doc.add_heading('6.5 Ratios Financieros (Año 3)', 1)
    doc.add_paragraph('• Margen EBITDA: 52,4%')
    doc.add_paragraph('• Margen neto: 38,3%')
    doc.add_paragraph('• ROI inversión: 162%')
    doc.add_paragraph('• LTV:CAC: 24:1 (excelente)')
    doc.add_paragraph('• Current ratio: 5,8 (liquidez alta)')
    doc.add_paragraph('• Burn rate año 1: 4.143€/mes')
    doc.add_paragraph('• Payback period: 2,1 años')
    
    doc.add_page_break()
    
    # ========== SECCIÓN 7: EQUIPO Y EJECUCIÓN ==========
    doc.add_heading('7. EQUIPO Y EJECUCIÓN', 0)
    
    # 7.1
    doc.add_heading('7.1 Equipo Fundador', 1)
    doc.add_paragraph('CEO - Estrategia y Producto:')
    doc.add_paragraph('• 10+ años emprendimiento digital')
    doc.add_paragraph('• Experiencia scale-ups tecnológicas')
    doc.add_paragraph('• Especialización economía colaborativa')
    
    doc.add_paragraph('CTO - Tecnología y Algoritmos:')
    doc.add_paragraph('• PhD Ciencias Computación')
    doc.add_paragraph('• Experiencia machine learning')
    doc.add_paragraph('• Arquitectura sistemas distribuidos')
    
    doc.add_paragraph('CMO - Marketing y Comunidad:')
    doc.add_paragraph('• Especialista growth marketing')
    doc.add_paragraph('• Experiencia construcción comunidades')
    doc.add_paragraph('• Conocimiento sector sostenibilidad')
    
    # 7.2
    doc.add_heading('7.2 Plan por Fases', 1)
    doc.add_paragraph('Fase 1 - Validación (meses 1-3):')
    doc.add_paragraph('• Landing page + waitlist')
    doc.add_paragraph('• Algoritmo matching básico')
    doc.add_paragraph('• 500 early adopters Barcelona')
    
    doc.add_paragraph('Fase 2 - MVP (meses 4-6):')
    doc.add_paragraph('• Plataforma funcional completa')
    doc.add_paragraph('• Integración Stripe Connect')
    doc.add_paragraph('• Expansión Madrid + Valencia')
    doc.add_paragraph('• 5.000 usuarios')
    
    doc.add_paragraph('Fase 3 - Crecimiento (meses 7-12):')
    doc.add_paragraph('• Optimización algoritmo')
    doc.add_paragraph('• Servicios premium')
    doc.add_paragraph('• Expansión nacional')
    doc.add_paragraph('• 25.000 usuarios activos')
    
    doc.add_page_break()
    
    # ========== SECCIÓN 8: CONCLUSIONES ==========
    doc.add_heading('8. CONCLUSIONES', 0)
    
    doc.add_paragraph('Treqe representa una oportunidad única en el mercado español:')
    doc.add_paragraph('1. **Innovación real:** Resuelve problema histórico (doble coincidencia deseos)')
    doc.add_paragraph('2. **Mercado validado:** 8 millones de usuarios potenciales, 15.000M€ valor atrapado')
    doc.add_paragraph('3. **Ventaja competitiva sostenible:** Primer mover, algoritmos propietarios, comisiones disruptivas')
    doc.add_paragraph('4. **Modelo escalable:** Arquitectura serverless, costes marginales bajos')
    doc.add_paragraph('5. **Impacto positivo:** Sostenibilidad ambiental + social verificable')
    doc.add_paragraph('6. **Viabilidad financiera:** Rentabilidad año 3, ROI 162%, cash flow positivo año 2')
    
    doc.add_paragraph('**Próximos pasos inmediatos:**')
    doc.add_paragraph('1. Registrar dominio treqe.es')
    doc.add_paragraph('2. Desarrollar landing page con waitlist')
    doc.add_paragraph('3. Implementar algoritmo matching básico (proof-of-concept)')
    doc.add_paragraph('4. Captar primeros 500 early adopters Barcelona')
    doc.add_paragraph('5. Validar modelo con transacciones reales')
    
    doc.add_paragraph('**Métrica éxito 3 meses:** 1.000+ suscriptores waitlist, 100+ transacciones validadas')
    
    # Guardar documento
    output_path = os.path.join(os.path.dirname(__file__), 'plan_negocio', 'Plan_Negocio_Treqe_FINAL_COMPLETO.docx')
    doc.save(output_path)
    
    print(f"Documento completado: {output_path}")
    return output_path

if __name__ == '__main__':
    try:
        # Primero crear documento base
        from crear_documento_final_completo import crear_documento_completo
        crear_documento_completo()
        
        # Luego continuar
        output_file = continuar_documento()
        print("✅ Documento final completo creado exitosamente.")
    except Exception as e:
        print(f"❌ Error: {e}")
        import traceback
        traceback.print_exc()