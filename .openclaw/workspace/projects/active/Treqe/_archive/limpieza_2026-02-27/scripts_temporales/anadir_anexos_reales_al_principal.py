#!/usr/bin/env python3
"""
Añadir contenido real de anexos al documento principal Treqe
"""

import docx
import os

def anadir_anexos_completos():
    """Añadir contenido completo de anexos al documento principal."""
    
    # Documento principal
    input_path = "plan_negocio/Plan_Negocio_Treqe_100%_COMPLETO_FINAL.docx"
    output_path = "plan_negocio/Plan_Negocio_Treqe_CON_ANEXOS_REALES.docx"
    
    if not os.path.exists(input_path):
        print(f"Error: No se encuentra {input_path}")
        return None
    
    print(f"Cargando documento principal: {input_path}")
    doc = docx.Document(input_path)
    
    # Encontrar donde empiezan los anexos
    anexos_start = -1
    for i, p in enumerate(doc.paragraphs):
        if 'ANEXOS' in p.text.upper() and len(p.text) < 20:
            anexos_start = i
            print(f"Sección ANEXOS encontrada en párrafo {i}")
            break
    
    if anexos_start == -1:
        print("No se encontró sección ANEXOS")
        return None
    
    # Buscar hasta donde termina la lista de anexos
    anexos_end = anexos_start + 1
    while anexos_end < len(doc.paragraphs):
        p = doc.paragraphs[anexos_end]
        # Si encontramos contenido que no es parte de la lista de anexos
        if p.text.strip() and not ('Anexo' in p.text or p.text.strip().startswith('Anexo')):
            # Verificar si es el inicio de otra sección
            if len(p.text) < 100 and ('SECCIÓN' in p.text or p.style.name.startswith('Heading')):
                break
        anexos_end += 1
    
    print(f"Lista de anexos ocupa párrafos {anexos_start} a {anexos_end-1}")
    
    # Eliminar la lista de anexos existente (solo títulos con números de página)
    print("Eliminando lista de anexos vacía...")
    for _ in range(anexos_start + 1, anexos_end):
        if anexos_start + 1 < len(doc.paragraphs):
            p = doc.paragraphs[anexos_start + 1]
            p_element = p._element
            p_element.getparent().remove(p_element)
    
    print("Añadiendo anexos con contenido real...")
    
    # Añadir página nueva para anexos
    doc.add_page_break()
    
    # ========== ANEXO A: ESTRUCTURA DEL EQUIPO ==========
    doc.add_paragraph("ANEXO A: ESTRUCTURA Y ROLES DEL EQUIPO")
    doc.paragraphs[-1].style = "Heading 2"
    
    contenido_a = [
        "",
        "ROLES NECESARIOS PARA EL DESARROLLO DEL MVP:",
        "",
        "1. FUNDADOR/CEO",
        "• Responsabilidades: Visión estratégica, desarrollo del modelo de negocio, fundraising, relaciones con inversores",
        "• Experiencia ideal: Conocimiento del sector de segunda mano y economía colaborativa, liderazgo, capacidad de negociación",
        "• Tiempo requerido: Full-time desde el inicio",
        "",
        "2. DESARROLLADOR FULL-STACK",
        "• Responsabilidades: Desarrollo de la plataforma web y móvil, arquitectura técnica, mantenimiento, escalabilidad",
        "• Stack tecnológico recomendado:",
        "  - Frontend: React/Next.js con TypeScript",
        "  - Backend: Node.js/Express o Python/Django",
        "  - Base de datos: PostgreSQL con PostGIS para geolocalización",
        "  - Infraestructura: AWS/GCP, Docker, CI/CD",
        "• Habilidades clave: Desarrollo de APIs, algoritmos de matching, deployment en cloud",
        "• Tiempo requerido: Full-time desde fase de desarrollo",
        "",
        "3. DISEÑADOR UX/UI",
        "• Responsabilidades: Investigación de usuario, wireframes, diseño de interfaz, user testing",
        "• Herramientas: Figma, Adobe XD, herramientas de prototipado",
        "• Habilidades: Design thinking, investigación cualitativa, accesibilidad, diseño responsive",
        "• Tiempo requerido: Full-time o contractor durante fases de diseño",
        "",
        "4. ESPECIALISTA EN CRECIMIENTO (GROWTH)",
        "• Responsabilidades: Adquisición de usuarios, estrategia de contenidos, community management",
        "• Habilidades: Marketing digital, analytics, redes sociales, SEO",
        "• Tiempo requerido: Part-time inicial, escalando con el crecimiento",
        "",
        "NOTA: Los CVs específicos de los miembros del equipo se añadirán una vez se definan los candidatos reales.",
        "La estructura presentada representa los roles mínimos necesarios para desarrollar y lanzar el MVP.",
    ]
    
    for item in contenido_a:
        if item == "":
            doc.add_paragraph()
        else:
            doc.add_paragraph(item)
    
    doc.add_page_break()
    
    # ========== ANEXO B: METODOLOGÍA DE VALIDACIÓN ==========
    doc.add_paragraph("ANEXO B: METODOLOGÍA PARA VALIDACIÓN DE MERCADO")
    doc.paragraphs[-1].style = "Heading 2"
    
    contenido_b = [
        "",
        "PLAN DE INVESTIGACIÓN Y VALIDACIÓN EN 4 FASES:",
        "",
        "FASE 1: INVESTIGACIÓN SECUNDARIA (SEMANAS 1-2)",
        "• Objetivo: Comprender el mercado, competencia y tendencias",
        "• Métodos:",
        "  - Análisis de datos públicos: INE, Eurostat, informes sectoriales",
        "  - Estudios existentes: Observatorio Cetelem, informes de consultoras",
        "  - Tendencias: Google Trends, reports de industria",
        "  - Benchmarking: Análisis de competidores (Wallapop, Vinted, Milanuncios)",
        "• Output: Mapa competitivo, tamaño de mercado estimado, tendencias clave",
        "",
        "FASE 2: ENCUESTAS CUANTITATIVAS (SEMANAS 3-4)",
        "• Objetivo: Validar hipótesis con muestra representativa",
        "• Métodos:",
        "  - Encuestas online: Plataformas como Typeform, Google Forms",
        "  - Muestra objetivo: 500-1000 usuarios españoles",
        "  - Segmentación: Por edad, ubicación, frecuencia de uso de segunda mano",
        "• Métricas clave:",
        "  - Intención de uso de la plataforma",
        "  - Preocupaciones principales sobre intercambios",
        "  - Disposición a pagar comisiones",
        "  - Preferencias sobre funcionalidades",
        "• Coste estimado: 500-2000€ (dependiendo de plataforma y targeting)",
        "",
        "FASE 3: ENTREVISTAS CUALITATIVAS (SEMANAS 5-6)",
        "• Objetivo: Deep dive en pain points y validar user flows",
        "• Métodos:",
        "  - Entrevistas 1:1 (presenciales o por video)",
        "  - Participantes: 15-20 usuarios potenciales",
        "  - Duración: 45-60 minutos por entrevista",
        "• Temas:",
        "  - Experiencias con plataformas de segunda mano existentes",
        "  - Dificultades con trueque/intercambio tradicional",
        "  - Reacciones a prototipos/wireframes",
        "  - Sugerencias para mejora",
        "• Output: User personas, customer journey maps, insights cualitativos",
        "• Coste estimado: 1000-3000€ (incentivos para participantes)",
        "",
        "FASE 4: PILOTO CONTROLADO (SEMANAS 7-12)",
        "• Objetivo: Validar producto en entorno real con métricas de negocio",
        "• Métodos:",
        "  - MVP funcional en barrio específico (ej: Gràcia, Barcelona)",
        "  - Participantes: 200-500 usuarios activos",
        "  - Duración: 4-6 semanas",
        "• Métricas de validación:",
        "  - Tasa de registro → publicación de producto",
        "  - Tasa de match (propuestas de intercambio)",
        "  - Tasa de completación (intercambios realizados)",
        "  - Net Promoter Score (NPS)",
        "  - Retención D1, D7, D30",
        "• Coste estimado: 5000-15000€ (desarrollo + marketing local)",
        "",
        "NOTA: Los datos específicos de cada estudio se añadirán tras realizar las investigaciones correspondientes.",
        "Esta metodología asegura validación progresiva con incremento gradual de inversión.",
    ]
    
    for item in contenido_b:
        if item == "":
            doc.add_paragraph()
        else:
            doc.add_paragraph(item)
    
    doc.add_page_break()
    
    # ========== ANEXO C: ESPECIFICACIONES TÉCNICAS ==========
    doc.add_paragraph("ANEXO C: ESPECIFICACIONES TÉCNICAS DE LA PLATAFORMA")
    doc.paragraphs[-1].style = "Heading 2"
    
    contenido_c = [
        "",
        "ARQUITECTURA TÉCNICA DEL SISTEMA:",
        "",
        "COMPONENTES PRINCIPALES:",
        "",
        "1. SISTEMA DE CATALOGACIÓN Y GESTIÓN DE PRODUCTOS",
        "• Upload múltiple de imágenes (3-10 por producto, máximo 10MB total)",
        "• Procesamiento de imágenes: Redimensionamiento automático, optimización, generación de thumbnails",
        "• Etiquetado automático: Detección de categoría, color, estado mediante visión por computadora",
        "• Valoración estimada: Algoritmo que sugiere rango de valor basado en características y mercado",
        "• Descripción asistida: Templates inteligentes que guían al usuario",
        "• Sistema de categorías: Jerárquico con subcategorías específicas",
        "",
        "2. MOTOR DE BÚSQUEDA Y MATCHING",
        "• Búsqueda semántica: Entiende consultas naturales ('mueble pequeño para estudio')",
        "• Filtros avanzados:",
        "  - Geolocalización: Radio personalizable (1km, 5km, 10km, ciudad, provincia)",
        "  - Categoría y subcategoría",
        "  - Rango de valor estimado",
        "  - Estado del producto (nuevo, como nuevo, bueno, aceptable)",
        "  - Fecha de publicación",
        "• Algoritmo de matching para ruedas de intercambio:",
        "  - Input: Producto publicado + preferencias de intercambio",
        "  - Proceso: Búsqueda de ciclos cerrados A→B→C→A que maximicen satisfacción",
        "  - Optimización: Considera valor, categorías deseadas, ubicación, disponibilidad",
        "  - Output: Lista de ruedas sugeridas ordenadas por probabilidad de éxito",
        "• Sistema de recomendaciones: Basado en historial de usuario y comportamiento similar",
        "",
        "3. SISTEMA DE INTERCAMBIO Y NEGOCIACIÓN",
        "• Propuesta 1-clic: Para ruedas sugeridas por el algoritmo",
        "• Chat integrado:",
        "  - Mensajería en tiempo real (WebSockets)",
        "  - Soporte para imágenes en chat",
        "  - Plantillas de mensaje para negociación",
        "  - Historial completo de conversaciones",
        "• Acuerdo de términos:",
        "  - Fecha y hora del intercambio",
        "  - Lugar de encuentro (sugerencias basadas en ubicación)",
        "  - Condiciones específicas",
        "  - Compensación monetaria si aplica",
        "• Sistema de seguimiento:",
        "  - Estados: Propuesto → Aceptado → Programado → Completado/Cancelado",
        "  - Notificaciones: Email y push para cada cambio de estado",
        "  - Recordatorios: 24h y 1h antes del intercambio",
        "• Calificación post-intercambio:",
        "  - Sistema de 5 estrellas + comentarios",
        "  - Calificación mutua obligatoria para completar intercambio",
        "  - Comentarios públicos en perfil",
        "",
        "4. SISTEMA DE PERFILES Y REPUTACIÓN",
        "• Perfil público:",
        "  - Foto, nombre, ubicación",
        "  - Historial de intercambios (completados, en progreso)",
        "  - Promedio de calificaciones",
        "  - Badges por logros",
        "• Sistema de reputación:",
        "  - Puntuación basada en transacciones exitosas",
        "  - Penalizaciones por cancelaciones tardías o incumplimientos",
        "  - Niveles de confianza (Novato → Experto → Maestro)",
        "• Verificaciones:",
        "  - Email (obligatorio)",
        "  - Teléfono móvil (opcional, aumenta confianza)",
        "  - Redes sociales (opcional, verificación cruzada)",
        "  - Documento de identidad (para transacciones de alto valor)",
        "",
        "5. SISTEMA DE PAGOS Y COMPENSACIONES",
        "• Integración con pasarelas de pago: Stripe, PayPal",
        "• Sistema de depósitos en garantía: Para transacciones de alto valor",
        "• Compensaciones monetarias:",
        "  - Calculadas automáticamente basadas en diferencia de valor",
        "  - Límites máximos por categoría",
        "  - Procesamiento seguro con protección al comprador",
        "• Comisiones:",
        "  - 1% sobre valor estimado (mínimo 1€, máximo 20€)",
        "  - Cobrada solo al completar intercambio exitoso",
        "  - Exenciones para primeros intercambios (promoción)",
        "",
        "FLUJOS DE USUARIO PRINCIPALES:",
        "",
        "FLUJO A: PUBLICACIÓN DE PRODUCTO (OBJETIVO: < 2 MINUTOS)",
        "1. Click 'Publicar' en header o dashboard",
        "2. Subida de fotos (arrastrar/soltar o seleccionar de galería)",
        "3. Asistente paso a paso:",
        "   a. Categoría (sugerida automáticamente de fotos)",
        "   b. Título (sugerencias basadas en categoría)",
        "   c. Descripción (plantillas inteligentes)",
        "   d. Estado (nuevo, como nuevo, bueno, aceptable)",
        "   e. Valor estimado (sugerido por algoritmo, editable)",
        "4. Preferencias de intercambio:",
        "   a. Qué busca a cambio (categorías, valor aproximado)",
        "   b. Radio de intercambio (local, ciudad, provincia)",
        "   c. Disponibilidad (fechas/horarios preferidos)",
        "5. Vista previa y confirmación",
        "6. Publicación (visible en 30 segundos)",
        "",
        "FLUJO B: ENCONTRAR Y PROPOR INTERCAMBIO (OBJETIVO: < 1 MINUTO)",
        "1. Búsqueda o exploración de categorías",
        "2. Ver producto de interés",
        "3. Click 'Intercambiar' → muestra ruedas sugeridas",
        "4. Selección de rueda preferida",
        "5. Vista de detalles de la rueda (quién da qué a quién)",
        "6. Envío de propuesta o inicio de chat",
        "",
        "FLUJO C: COMPLETAR INTERCAMBIO (OBJETIVO: SEGUIMIENTO SIN FRICCIÓN)",
        "1. Notificación de propuesta recibida/aceptada",
        "2.