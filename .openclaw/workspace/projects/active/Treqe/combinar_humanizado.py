#!/usr/bin/env python3
"""
Combinar documento base (secciones 1-2 humanizadas) con sección 3 humanizada
"""

from docx import Document
import os

def combinar_documentos():
    """Combinar secciones humanizadas."""
    
    print("Combinando documento humanizado...")
    
    # Cargar documento base (con secciones 1-2 ya humanizadas)
    base_path = os.path.join(os.path.dirname(__file__), 'plan_negocio', 'Plan_Negocio_Treqe_Completo.docx')
    doc = Document(base_path)
    
    # Agregar sección 3 humanizada
    doc.add_page_break()
    
    # Título sección 3
    doc.add_heading('3. LA SOLUCIÓN TREQE: CUANDO LAS MATEMÁTICAS ENCUENTRAN PUENTES HUMANOS', 0)
    
    # 3.1
    doc.add_heading('3.1 Más Allá del Trueque Tradicional: Un Salto Conceptual', 1)
    doc.add_paragraph('Si hay algo que la historia nos ha enseñado sobre el trueque es que, a pesar de su aparente simplicidad, presenta una limitación fundamental que ha impedido su escalabilidad como modelo comercial viable. Me refiero, por supuesto, a lo que los economistas denominan "el problema de la doble coincidencia de deseos".')
    doc.add_paragraph('Durante siglos, esta limitación ha confinado el trueque a escalas reducidas, a intercambios ocasionales entre personas que casualmente querían exactamente lo que la otra ofrecía. Pero, ¿y si pudiéramos superar esta limitación? ¿Y si pudiéramos transformar lo que históricamente ha sido un proceso azaroso en uno sistemático y escalable?')
    doc.add_paragraph('Es precisamente aquí donde Treqe introduce lo que consideramos una innovación conceptual significativa: las "ruedas de intercambio inteligente". La idea es tan elegante en su concepción como poderosa en su aplicación. En lugar de limitarnos a buscar coincidencias directas entre dos personas (una tarea estadísticamente improbable), permitimos que tres, cuatro o incluso cinco usuarios participen en cadenas circulares donde cada uno recibe exactamente lo que desea, aunque no directamente de la persona a la que entrega su artículo.')
    doc.add_paragraph('Lo que hace especial a este sistema no es solo su capacidad para resolver un problema matemático, sino su comprensión profunda de la naturaleza humana. Reconoce que nuestras necesidades son complejas, que los objetos que poseemos tienen tanto valor sentimental como económico, y que la falta de liquidez no debería ser una barrera insalvable para satisfacer nuestras necesidades de renovación y cambio.')
    
    # 3.2
    doc.add_heading('3.2 El Mecanismo Detrás de la Magia: Un Proceso Pensado para Personas Reales', 1)
    doc.add_paragraph('Para comprender cómo Treqe transforma una idea conceptual en una experiencia tangible, es necesario analizar el mecanismo operativo paso a paso. Cada fase ha sido diseñada no solo pensando en la eficiencia técnica, sino especialmente en la experiencia humana.')
    
    doc.add_paragraph('Fase 1: El Arte de Declarar lo que Tenemos y lo que Deseamos')
    doc.add_paragraph('El proceso comienza con algo aparentemente simple pero profundamente significativo: invitar a los usuarios a reflexionar sobre lo que poseen y lo que realmente necesitan. No se trata de un inventario frío, sino de un ejercicio de autoconocimiento práctico.')
    doc.add_paragraph('Cuando un usuario ingresa a Treqe, se le pide que complete dos listas:')
    doc.add_paragraph('- "Esto tengo": Los artículos que actualmente posee pero que ya no se alinean con su vida actual')
    doc.add_paragraph('- "Esto quiero": Los objetos que necesita para mejorar su espacio, su comodidad, su funcionalidad')
    doc.add_paragraph('Para cada artículo, el sistema solicita información que va más allá de lo meramente descriptivo. Se piden fotografías desde múltiples ángulos, no por mero formalismo, sino porque entendemos que ver un objeto es comenzar a establecer una relación con él. Se solicita una valoración económica, pero el sistema ofrece sugerencias basadas en datos de mercado reales, evitando que los usuarios subestimen o sobreestimen el valor de sus posesiones.')
    
    doc.add_paragraph('Fase 2: Cuando las Matemáticas se Convierten en Conectores Humanos')
    doc.add_paragraph('Una vez que los usuarios han declarado sus preferencias, el sistema transforma esta información en una estructura matemática: un grafo dirigido. Puede sonar técnico, pero en esencia es bastante simple: cada usuario se convierte en un nodo, y las flechas (aristas) van desde lo que una persona tiene hacia lo que otra persona quiere.')
    doc.add_paragraph('Utilizando algoritmos de teoría de grafos (específicamente búsqueda en profundidad optimizada), el sistema busca ciclos cerrados de 3 a 5 nodos. Cada ciclo representa una posible cadena de intercambio donde todos obtienen lo que desean. Lo notable es la velocidad: cada búsqueda tiene un timeout de 500 milisegundos, garantizando que el sistema sea rápido y responsivo incluso a medida que crece.')
    
    doc.add_paragraph('Fase 3: La Equidad como Principio de Diseño')
    doc.add_paragraph('En el mundo real, es raro que dos artículos tengan exactamente el mismo valor económico. Por eso, cuando Treqe identifica un ciclo viable, no se limita a proponer intercambios directos. En su lugar, utiliza programación lineal (con el algoritmo PuLP) para calcular las compensaciones monetarias óptimas.')
    doc.add_paragraph('La optimización considera múltiples factores:')
    doc.add_paragraph('- Minimizar el flujo monetario total (reducir al máximo las transferencias)')
    doc.add_paragraph('- Distribuir las compensaciones de manera equitativa')
    doc.add_paragraph('- Respetar las preferencias individuales (algunos usuarios prefieren pagar menos aunque reciban menos valor, otros prefieren lo contrario)')
    doc.add_paragraph('- Garantizar que nadie se sienta en desventaja en la transacción')
    
    doc.add_paragraph('Fase 4: La Negociación como Conversación, no como Confrontación')
    doc.add_paragraph('Una vez identificado un ciclo y calculadas las compensaciones, Treqe no impone la solución. En su lugar, crea un espacio de negociación facilitada. Los participantes se conectan a través de un chat grupal en tiempo real (utilizando WebSockets), donde pueden:')
    doc.add_paragraph('- Ver la propuesta inicial del sistema')
    doc.add_paragraph('- Hacer preguntas sobre los artículos involucrados')
    doc.add_paragraph('- Solicitar fotografías adicionales o información complementaria')
    doc.add_paragraph('- Proponer ajustes menores a los términos')
    doc.add_paragraph('- Confirmar explícitamente su acuerdo')
    doc.add_paragraph('Este proceso de negociación no es una formalidad burocrática, sino un espacio para construir confianza. Los usuarios no están interactuando con un algoritmo frío, sino con otras personas reales que, como ellos, buscan mejorar sus vidas a través del intercambio.')
    
    doc.add_paragraph('Fase 5: La Seguridad como Fundamento, no como Añadido')
    doc.add_paragraph('Finalmente, una vez alcanzado el acuerdo, Treqe gestiona la ejecución completa con múltiples capas de seguridad:')
    doc.add_paragraph('- Los pagos de compensación se procesan a través de Stripe Connect, con fondos retenidos en escrow hasta que todas las partes confirman la recepción satisfactoria')
    doc.add_paragraph('- La logística se integra con APIs de empresas establecidas (Correos, SEUR), proporcionando tracking en tiempo real')
    doc.add_paragraph('- Un sistema de reputación granular permite evaluar múltiples dimensiones (puntualidad, precisión en descripciones, calidad del embalaje)')
    doc.add_paragraph('- Mecanismos de disputa escalonados ofrecen protección sin necesidad de intervención legal inmediata')
    doc.add_paragraph('Lo que emerge de este proceso no es solo una transacción económica, sino una experiencia humana completa: desde la reflexión inicial sobre lo que poseemos y necesitamos, pasando por la conexión con otros, hasta la satisfacción de haber obtenido lo que queríamos de manera justa y segura.')
    
    # 3.3
    doc.add_heading('3.3 Un Caso que Ilumina el Concepto: Ana, Carlos y Beatriz', 1)
    doc.add_paragraph('Para comprender realmente cómo funciona Treqe, nada mejor que analizar un caso concreto. No se trata de un ejemplo teórico, sino de una situación que refleja la realidad de millones de personas.')
    
    doc.add_paragraph('Los Protagonistas y sus Historias:')
    doc.add_paragraph('Ana, 32 años, arquitecta en Barcelona:')
    doc.add_paragraph('Ana compró su bicicleta de montaña Specialized Rockhopper durante sus años universitarios, cuando el ciclismo era más que un hobby: era una pasión. Ahora, con un trabajo exigente y menos tiempo libre, la bicicleta pasa la mayor parte del tiempo aparcada en el pequeño balcón de su apartamento. Lo que Ana realmente necesita es un sofá de diseño contemporáneo que se adapte mejor a su espacio y estilo actual. El problema no es de deseo, sino de recursos: aunque la bicicleta tiene un valor considerable (450€), Ana no tiene los 600€ que costaría el sofá que quiere.')
    
    doc.add_paragraph('Carlos, 41 años, profesor universitario en Hospitalet de Llobregat:')
    doc.add_paragraph('Carlos invirtió en un sofá de diseño de calidad hace dos años, pero su vida ha cambiado. Ha comenzado un proyecto personal: un canal educativo en YouTube sobre historia del arte. Para producir contenido de calidad, necesita un ordenador portátil potente para edición de video (valorado en 800€). El sofá, aunque perfectamente funcional y valioso (600€), ya no es una prioridad. Como muchos proyectos personales, el suyo tiene un presupuesto limitado.')
    
    doc.add_paragraph('Beatriz, 28 años, diseñadora gráfica freelance en Badalona:')
    doc.add_paragraph('Beatriz invirtió en un MacBook Pro M2 para su trabajo como diseñadora freelance. Recientemente, ha comenzado a trabajar de forma híbrida, yendo a una oficina compartida varios días a la semana. Necesita una bicicleta para sus desplazamientos urbanos (presupuesto: 450€), pero como freelance con ingresos variables, un desembolso de esa magnitud afectaría significativamente su liquidez mensual.')
    
    doc.add_paragraph('El Problema Tradicional:')
    doc.add_paragraph('Por separado, cada uno enfrenta una situación frustrante:')
    doc.add_paragraph('- Ana no puede conseguir su sofá porque Carlos no quiere su bicicleta')
    doc.add_paragraph('- Carlos no puede conseguir su ordenador porque Beatriz no quiere su sofá')
    doc.add_paragraph('- Beatriz no puede conseguir su bicicleta porque Ana no quiere su ordenador')
    doc.add_paragraph('Es el clásico problema de la doble coincidencia de deseos: cada par podría intercambiar si sus deseos coincidieran, pero no lo hacen.')
    
    doc.add_paragraph('La Solución Treqe:')
    doc.add_paragraph('El algoritmo identifica que, aunque ningún intercambio directo es posible, existe un ciclo cerrado perfecto entre los tres. Propone:')
    doc.add_paragraph('Intercambios físicos:')
    doc.add_paragraph('1. Ana entrega su bicicleta a Beatriz')
    doc.add_paragraph('2. Carlos entrega su sofá a Ana')
    doc.add_paragraph('3. Beatriz entrega su ordenador a Carlos')
    
    doc.add_paragraph('Compensaciones monetarias:')
    doc.add_paragraph('- Ana paga 150€ a Carlos (la diferencia entre el sofá que recibe y la bicicleta que entrega)')
    doc.add_paragraph('- Carlos paga 200€ a Beatriz (la diferencia entre el ordenador que recibe y el sofá que entrega)')
    doc.add_paragraph('- Beatriz recibe 350€ neto (150€ de Ana + 200€ de Carlos)')
    
    doc.add_paragraph('Los Resultados Transformadores:')
    doc.add_paragraph('- Ana: Obtiene el sofá de diseño que necesita por solo 150€, en lugar de los 600€ que costaría nuevo. Su ahorro del 75% no es solo económico: es la posibilidad de tener un espacio que realmente refleja quién es ahora.')
    doc.add_paragraph('- Carlos: Consigue el ordenador que necesita para su proyecto YouTube por 200€, en lugar de 800€. El ahorro del 75% representa meses de trabajo que puede dedicar a crear contenido en lugar de ahorrar para equipamiento.')
    doc.add_paragraph('- Beatriz: Recibe la bicicleta que necesita para sus desplazamientos más 350€ de liquidez adicional. El valor total recibido (800€) le permite mejorar tanto su movilidad como su situación financiera.')
    
    doc.add_paragraph('Lo que este Caso nos Enseña:')
    doc.add_paragraph('Este ejemplo no es una anomalía estadística, sino una ilustración de cómo Treqe transforma problemas aparentemente insolubles en oportunidades de ganancia mutua. Lo notable no es solo el ahorro económico (aunque es significativo), sino cómo el sistema permite a cada persona avanzar en sus proyectos de vida sin las restricciones tradicionales de liquidez.')
    doc.add_paragraph('La belleza del sistema reside en su simplicidad final: tres personas, tres necesidades diferentes, una solución elegante que beneficia a todos. Es la materialización de la idea de que, a veces, lo que necesitamos no es más dinero, sino una forma más inteligente de utilizar el valor que ya poseemos.')
    
    # Guardar documento
    output_path = os.path.join(os.path.dirname(__file__), 'plan_negocio', 'Plan_Negocio_Treqe_HUMANIZADO_FINAL.docx')
    doc.save(output_path)
    
    print(f"Documento humanizado final creado: {output_path}")
    print("Contenido incluido:")
    print("- Secciones 1-2: Redacción humanizada original")
    print("- Sección 3: Redacción completamente humanizada y extensa")
    print("- Estilo narrativo y profundo")
    print("- Ejemplos detallados y casos prácticos")
    
    return output_path

if __name__ == '__main__':
    try:
        output_file = combinar_documentos()
        print("EXITO: Documento humanizado creado exitosamente.")
    except Exception as e:
        print(f"ERROR: {e}")
        import traceback
        traceback.print_exc()