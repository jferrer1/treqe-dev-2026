#!/usr/bin/env python3
"""
Iniciar Fase 2: Estrategia de Marketing Detallada y Aspectos Legales
"""

from docx import Document
import os

def iniciar_fase2():
    """Iniciar Fase 2 de completación."""
    
    print("INICIANDO FASE 2: ESTRATEGIA DE MARKETING Y ASPECTOS LEGALES")
    print("=" * 60)
    
    # Cargar documento de Fase 1
    doc_path = os.path.join(os.path.dirname(__file__), 'plan_negocio', 'Plan_Negocio_Treqe_FASE1_COMPLETADA.docx')
    doc = Document(doc_path)
    
    print("Documento cargado:", doc_path)
    print("Tamaño actual:", os.path.getsize(doc_path) / 1024, "KB")
    
    # ========== SECCIÓN 10: ESTRATEGIA DE MARKETING DETALLADA ==========
    print("\nAgregando SECCION 10: ESTRATEGIA DE MARKETING DETALLADA...")
    
    doc.add_page_break()
    doc.add_heading('10. ESTRATEGIA DE MARKETING DETALLADA: CONSTRUYENDO COMUNIDAD, IMPULSANDO CRECIMIENTO', 0)
    
    # 10.1
    doc.add_heading('10.1 Plan de Lanzamiento Fase por Fase: Del Primer Usuario a la Masa Crítica', 1)
    doc.add_paragraph('El lanzamiento de Treqe no es un evento único, sino un proceso cuidadosamente orquestado en fases, cada una con objetivos específicos, tácticas adaptadas y métricas claras.')
    
    doc.add_paragraph('Fase Alpha (Meses 1-2): Validación del Concepto')
    doc.add_paragraph('Objetivo: Demostrar que el algoritmo funciona y que los usuarios encuentran valor en el trueque estructurado.')
    doc.add_paragraph('- Tácticas: 50 usuarios invitados manualmente, matching asistido por el equipo, feedback intensivo')
    doc.add_paragraph('- Métricas: Tasa de éxito en intercambios, satisfacción usuario, tiempo para matching')
    doc.add_paragraph('- Presupuesto: 1.000€ (principalmente para incentivos a early adopters)')
    
    doc.add_paragraph('Fase Beta (Meses 3-4): Refinamiento y Primer Crecimiento')
    doc.add_paragraph('Objetivo: Refinar la experiencia basada en feedback y comenzar crecimiento orgánico controlado.')
    doc.add_paragraph('- Tácticas: Waitlist pública, onboarding por invitación, programas de referidos')
    doc.add_paragraph('- Métricas: Tasa de conversión waitlist, crecimiento orgánico, retención a 30 días')
    doc.add_paragraph('- Presupuesto: 3.000€ (marketing de contenidos, programas de referidos)')
    
    doc.add_paragraph('Fase de Lanzamiento Público (Meses 5-6): Aceleración Controlada')
    doc.add_paragraph('Objetivo: Alcanzar masa crítica inicial en mercados objetivo prioritarios.')
    doc.add_paragraph('- Tácticas: Lanzamiento en Barcelona y Madrid, campañas digitales dirigidas, partnerships locales')
    doc.add_paragraph('- Métricas: 5.000 usuarios activos, 1.000 transacciones mensuales, CAC (coste de adquisición)')
    doc.add_paragraph('- Presupuesto: 8.000€ (campañas digitales, eventos locales, PR)')
    
    doc.add_paragraph('Fase de Expansión (Meses 7-12): Crecimiento Sostenido')
    doc.add_paragraph('Objetivo: Expandir a principales ciudades españolas y optimizar adquisición.')
    doc.add_paragraph('- Tácticas: Expansión geográfica progresiva, optimización de canales, programas de fidelización')
    doc.add_paragraph('- Métricas: 25.000 usuarios, LTV (valor de vida del cliente), ROI marketing')
    doc.add_paragraph('- Presupuesto: 15.000€ (marketing performance, expansión geográfica)')
    
    # 10.2
    doc.add_heading('10.2 Canales de Adquisición y Costes: Donde Encontramos a Nuestros Usuarios', 1)
    doc.add_paragraph('Nuestra estrategia de adquisición combina canales orgánicos y pagados, priorizando aquellos con mejor fit para nuestro perfil de usuario y mayor potencial de escalabilidad.')
    
    doc.add_paragraph('Canal 1: Marketing de Contenidos y SEO (Costo Medio, Alto Engagement)')
    doc.add_paragraph('- Tácticas: Blog sobre economía circular, guías de trueque, casos de éxito')
    doc.add_paragraph('- Público objetivo: Early adopters conscientes, interesados en sostenibilidad')
    doc.add_paragraph('- Coste estimado: 2.000€/mes (creación de contenido, SEO técnico)')
    doc.add_paragraph('- Métricas objetivo: 5.000 visitas/mes, 10% tasa de conversión a waitlist')
    
    doc.add_paragraph('Canal 2: Redes Sociales y Comunidades (Bajo Costo, Alto Potencial Viral)')
    doc.add_paragraph('- Tácticas: Grupos Facebook de segunda mano, subreddits especializados, Instagram de sostenibilidad')
    doc.add_paragraph('- Público objetivo: Usuarios activos en plataformas de segunda mano')
    doc.add_paragraph('- Coste estimado: 1.500€/mes (gestión comunidades, contenido visual)')
    doc.add_paragraph('- Métricas objetivo: 2.000 seguidores/mes, 15% tasa de engagement')
    
    doc.add_paragraph('Canal 3: Marketing de Influencers y Early Adopters (Costo Variable, Alta Credibilidad)')
    doc.add_paragraph('- Tácticas: Colaboraciones con micro-influencers de sostenibilidad, programas de embajadores')
    doc.add_paragraph('- Público objetivo: Seguidores de influencers de estilo de vida consciente')
    doc.add_paragraph('- Coste estimado: 3.000€/mes (honorarios influencers, incentivos embajadores)')
    doc.add_paragraph('- Métricas objetivo: 500 referidos/mes por influencer, 25% tasa de conversión')
    
    doc.add_paragraph('Canal 4: Publicidad Digital Dirigida (Alto Costo, Alta Precisión)')
    doc.add_paragraph('- Tácticas: Google Ads para búsquedas relacionadas, Facebook Ads con targeting avanzado')
    doc.add_paragraph('- Público objetivo: Usuarios con intención de compra/venta de segunda mano')
    doc.add_paragraph('- Coste estimado: 5.000€/mes (presupuesto publicitario, gestión campañas)')
    doc.add_paragraph('- Métricas objetivo: CAC < 15€, ROI > 300% a 12 meses')
    
    doc.add_paragraph('Canal 5: Partnerships y Colaboraciones (Bajo Costo, Alto Impacto)')
    doc.add_paragraph('- Tácticas: Acuerdos con tiendas de segunda mano física, colaboraciones con ONGs ambientales')
    doc.add_paragraph('- Público objetivo: Clientes existentes de comercios de segunda mano')
    doc.add_paragraph('- Coste estimado: 1.000€/mes (eventos conjuntos, materiales promocionales)')
    doc.add_paragraph('- Métricas objetivo: 200 usuarios/mes por partnership, 30% tasa de retención')
    
    # 10.3
    doc.add_heading('10.3 Métricas de Crecimiento y Retención: Lo que Realmente Importa', 1)
    doc.add_paragraph('En Treqe, no solo medimos crecimiento bruto, sino crecimiento saludable y sostenible. Nuestras métricas clave reflejan esta filosofía.')
    
    doc.add_paragraph('Métricas de Adquisición (Top of Funnel):')
    doc.add_paragraph('- CAC (Coste de Adquisición de Cliente): Objetivo < 15€')
    doc.add_paragraph('- Tasa de Conversión Waitlist: Objetivo > 40%')
    doc.add_paragraph('- Tiempo desde registro hasta primer intercambio: Objetivo < 7 días')
    doc.add_paragraph('- Fuentes de tráfico por calidad: Segmentación por canal y calidad')
    
    doc.add_paragraph('Métricas de Activación y Engagement (Middle of Funnel):')
    doc.add_paragraph('- Tasa de Activación (usuarios que completan primer perfil): Objetivo > 70%')
    doc.add_paragraph('- Número promedio de artículos listados por usuario: Objetivo > 3')
    doc.add_paragraph('- Frecuencia de uso (sesiones/semana): Objetivo > 2')
    doc.add_paragraph('- Tiempo en plataforma por sesión: Objetivo > 8 minutos')
    
    doc.add_paragraph('Métricas de Retención y Valor (Bottom of Funnel):')
    doc.add_paragraph('- Retención a 30 días: Objetivo > 40%')
    doc.add_paragraph('- Retención a 90 días: Objetivo > 25%')
    doc.add_paragraph('- LTV (Valor de Vida del Cliente): Objetivo > 45€')
    doc.add_paragraph('- LTV:CAC Ratio: Objetivo > 3:1')
    
    doc.add_paragraph('Métricas de Viralidad y Referidos:')
    doc.add_paragraph('- Tasa de Viralidad (K-factor): Objetivo > 0.5')
    doc.add_paragraph('- Número de referidos por usuario activo: Objetivo > 1.2')
    doc.add_paragraph('- Tasa de conversión de referidos: Objetivo > 30%')
    doc.add_paragraph('- Coste de referido vs CAC orgánico: Objetivo 50% más barato')
    
    # 10.4
    doc.add_heading('10.4 Presupuesto de Marketing Detallado: Cada Euro con Propósito', 1)
    doc.add_paragraph('Nuestro presupuesto de marketing no es un gasto, sino una inversión cuidadosamente asignada para maximizar crecimiento sostenible.')
    
    doc.add_paragraph('Presupuesto Año 1 (Total: 30.000€):')
    doc.add_paragraph('- Meses 1-2 (Alpha): 2.000€ (1.000€/mes)')
    doc.add_paragraph('- Meses 3-4 (Beta): 6.000€ (3.000€/mes)')
    doc.add_paragraph('- Meses 5-6 (Lanzamiento): 16.000€ (8.000€/mes)')
    doc.add_paragraph('- Meses 7-12 (Expansión): 90.000€ (15.000€/mes)')
    
    doc.add_paragraph('Distribución por Canal (Año 1):')
    doc.add_paragraph('- Marketing de Contenidos/SEO: 24.000€ (20%)')
    doc.add_paragraph('- Redes Sociales/Comunidades: 18.000€ (15%)')
    doc.add_paragraph('- Influencers/Early Adopters: 36.000€ (30%)')
    doc.add_paragraph('- Publicidad Digital: 30.000€ (25%)')
    doc.add_paragraph('- Partnerships/Colaboraciones: 12.000€ (10%)')
    
    doc.add_paragraph('Retorno Esperado (Año 1):')
    doc.add_paragraph('- Usuarios adquiridos: 25.000')
    doc.add_paragraph('- CAC promedio: 12€')
    doc.add_paragraph('- LTV promedio (a 12 meses): 45€')
    doc.add_paragraph('- ROI marketing: 275%')
    doc.add_paragraph('- Ingresos generados por usuarios adquiridos: 82.500€')
    
    # 10.5
    doc.add_heading('10.5 Cronograma de Actividades: El Camino Concreto', 1)
    doc.add_paragraph('Cada actividad de marketing está programada en un cronograma detallado que asegura ejecución coordinada y maximización de impacto.')
    
    doc.add_paragraph('Mes 1-2 (Pre-lanzamiento):')
    doc.add_paragraph('- Semana 1-2: Desarrollo contenido base (blog, guías, FAQ)')
    doc.add_paragraph('- Semana 3-4: Configuración canales sociales y SEO técnico')
    doc.add_paragraph('- Semana 5-6: Identificación y contacto early adopters')
    doc.add_paragraph('- Semana 7-8: Lanzamiento waitlist y primeros tests')
    
    doc.add_paragraph('Mes 3-4 (Lanzamiento Beta):')
    doc.add_paragraph('- Semana 9-10: Onboarding primeros 100 usuarios invitados')
    doc.add_paragraph('- Semana 11-12: Programas de referidos y feedback intensivo')
    doc.add_paragraph('- Semana 13-14: Optimización basada en datos iniciales')
    doc.add_paragraph('- Semana 15-16: Preparación lanzamiento público')
    
    doc.add_paragraph('Mes 5-6 (Lanzamiento Público Barcelona/Madrid):')
    doc.add_paragraph('- Semana 17: Evento lanzamiento Barcelona')
    doc.add_paragraph('- Semana 18: Campaña digital intensiva Barcelona')
    doc.add_paragraph('- Semana 19: Evento lanzamiento Madrid')
    doc.add_paragraph('- Semana 20-24: Campañas sostenidas ambas ciudades')
    
    doc.add_paragraph('Mes 7-12 (Expansión Nacional):')
    doc.add_paragraph('- Mes 7: Valencia y Sevilla')
    doc.add_paragraph('- Mes 8: Bilbao y Málaga')
    doc.add_paragraph('- Mes 9-10: Otras capitales provinciales')
    doc.add_paragraph('- Mes 11-12: Consolidación nacional y optimización')
    
    # ========== SECCIÓN 11: ASPECTOS LEGALES Y REGULATORIOS ==========
    print("\nAgregando SECCION 11: ASPECTOS LEGALES Y REGULATORIOS...")
    
    doc.add_page_break()
    doc.add_heading('11. ASPECTOS LEGALES Y REGULATORIOS: CONSTRUYENDO SOBRE CIMIENTOS SÓLIDOS', 0)
    
    # 11.1
    doc.add_heading('11.1 Estructura Legal Óptima: La Fundación Jurídica de Treqe', 1)
    doc.add_paragraph('La elección de la estructura legal no es una formalidad, sino una decisión estratégica que afecta fiscalidad, responsabilidad, capacidad de financiación y escalabilidad futura.')
    
    doc.add_paragraph('Opción Seleccionada: Sociedad Limitada (SL)')
    doc.add_paragraph('- Capital social mínimo: 3.000€ (suficiente para operaciones iniciales)')
    doc.add_paragraph('- Responsabilidad: Limitada al capital social (protección patrimonial fundadores)')
    doc.add_paragraph('- Fiscalidad: Impuesto de Sociedades (25%), posibilidad de tipos reducidos para startups')
    doc.add_paragraph('- Flexibilidad: Fácil entrada/salida de socios, adaptabilidad a crecimiento')
    
    doc.add_paragraph('Razones para esta elección:')
    doc.add_paragraph('1. Protección patrimonial de los fundadores frente a deudas sociales')
    doc.add_paragraph('2. Credibilidad frente a inversores y partners')
    doc.add_paragraph('3. Flexibilidad para futuras rondas de inversión')
    doc.add_paragraph('4. Compatibilidad con programas de apoyo a startups (ENISA, etc.)')
    doc.add_paragraph('5. Simplicidad administrativa relativa vs otras opciones')
    
    doc.add_paragraph('Plan de Evolución Legal:')
    doc.add_paragraph('- Año 1: Sociedad Limitada estándar')
    doc.add_paragraph('- Año 2-3: Transformación a Sociedad Limitada de Formación Sucesiva si cumple requisitos')
    doc.add_paragraph('- Año 3+: Evaluación transformación a Sociedad Anónima si se planea salida a bolsa o entrada de inversores institucionales')
    
    # 11.2
    doc.add_heading('11.2 Propiedad Intelectual y Protección: Defendiendo Nuestro Valor Fundamental', 1)
    doc.add_paragraph('En Treqe, nuestro valor no está solo en el código, sino en los algoritmos, procesos y metodologías únicas que hemos desarrollado. Proteger esta propiedad intelectual es proteger el corazón de nuestro negocio.')
    
    doc.add_paragraph('Elementos a Proteger:')
    doc.add_paragraph('1. Algoritmo de Matching de Ciclos Múltiples')
    doc.add_paragraph('   - Protección: Patente de método (si cumple requisitos de novedad)')
    doc.add_paragraph('   - Alternativa: Secreto comercial con acuerdos de confidencialidad')
    doc.add_paragraph('   - Documentación: Código fuente con timestamp notarial')
    
    doc.add_paragraph('2. Sistema de Optimización Económica')
    doc.add_paragraph('   - Protección: Derechos de autor sobre implementación específica')
    doc.add_paragraph('   - Documentación: Especificaciones técnicas detalladas')
    doc.add_paragraph('   - Contratos: Acuerdos con desarrolladores que asignen IP a la empresa')
    
    doc.add_paragraph('3. Marca "Treqe" y Elementos de Identidad')
    doc